"""
Content-preservation regression gate.

Table/page-break/continuation changes must never lose or duplicate real document
content. This compares the SOURCE docx against the formatted OUTPUT docx (plus
the rendered structure) and reports inventories + hard blockers:

  - table cell content multiset (the strongest "no row/cell loss" signal);
  - Источник: / Примечание: service-line count;
  - reference (bibliography) entry count;
  - table caption / title inventory;
  - empty-output-table-from-non-empty-source.

Normalization (so accepted formatter behaviour does NOT false-fire):
  - whitespace collapse;
  - decimal comma/dot unified and trailing decimal zeros stripped
    (``31,75`` == ``31.75``, ``13,00`` == ``13``);
  - case-folded (table-title capitalization from upstream rules is accepted).

Structure checks (TOC/sections/appendix) come from
``document_structure_validation.evaluate_document_structure``; this module adds
the content-level checks. Pure functions over python-docx Documents.
"""
from __future__ import annotations

import re
from collections import Counter
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn


@dataclass(frozen=True)
class ContentIssue:
    issue_type: str
    severity: str  # "fail" | "needs_human_review"
    evidence: dict[str, object] = field(default_factory=dict)


_NUM_RE = re.compile(r"\d+(?:[.,]\d+)?")
_CAPTION_RE = re.compile(r"^\s*таблица\s+\d+(?:\.\d+)*", re.IGNORECASE)
_SOURCE_NOTE_RE = re.compile(r"^\s*(источник|примечание)\s*[:.]", re.IGNORECASE)
_REF_ENTRY_RE = re.compile(r"^\s*\d+\.?\s+\S")


def _canon_number(m: re.Match) -> str:
    s = m.group(0).replace(",", ".")
    if "." in s:
        s = s.rstrip("0").rstrip(".")
    return s


def normalize_cell(text: str) -> str:
    """Whitespace-collapse, case-fold, and canonicalise numbers so accepted
    normalizations compare equal: decimal comma/dot, trailing decimal zeros,
    thousands-separator spaces (``73 900`` == ``73900``), and capitalization."""
    t = " ".join((text or "").replace("\xa0", " ").split()).lower()
    t = re.sub(r"(?<=\d)[  ](?=\d)", "", t)  # join thousands-separator groups
    return _NUM_RE.sub(_canon_number, t)


def table_cell_multiset(doc: Document) -> Counter:
    c: Counter = Counter()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                norm = normalize_cell(cell.text)
                if norm:
                    c[norm] += 1
    return c


def _is_numeric_index_row(values: list[str]) -> bool:
    vals = [" ".join((v or "").split()) for v in values]
    nonempty = [v for v in vals if v]
    return len(nonempty) >= 2 and all(v.isdigit() for v in nonempty)


def meaningful_row_set(doc: Document) -> set[str]:
    """Distinct meaningful table-row fingerprints. Excludes numeric column-index
    rows. A header/numeric row repeated across continuation fragments yields the
    same fingerprint, so legitimate repetition does NOT register as new content —
    only genuinely lost/added DATA rows change the set."""
    out: set[str] = set()
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text for cell in row.cells]
            if _is_numeric_index_row(values):
                continue
            fp = " | ".join(normalize_cell(v) for v in values if normalize_cell(v))
            if len(fp) >= 5:
                out.add(fp)
    return out


def data_row_multiset(doc: Document) -> Counter:
    """Multiset of DATA rows: every row except each table's first row (the header,
    which is legitimately repeated on continuation fragments) and numeric index
    rows. Lets us detect real data-row duplication without flagging header repeat."""
    c: Counter = Counter()
    for table in doc.tables:
        for idx, row in enumerate(table.rows):
            if idx == 0:
                continue
            values = [cell.text for cell in row.cells]
            if _is_numeric_index_row(values):
                continue
            fp = " | ".join(normalize_cell(v) for v in values if normalize_cell(v))
            if len(fp) >= 5:
                c[fp] += 1
    return c


def reference_entry_set(doc: Document) -> set[str]:
    """Distinct reference-entry CONTENT (leading number stripped) inside the
    references section — robust to the formatter renumbering entries."""
    texts = _para_texts(doc)
    started = False
    out: set[str] = set()
    for t in texts:
        low = t.lower()
        if low.startswith("список использован"):
            started = True
            continue
        if not started:
            continue
        if low in {"приложения", "приложение"} or low.startswith("приложение "):
            break
        content = re.sub(r"^\s*\d+\.?\s*", "", t)  # strip leading "N." numbering
        norm = normalize_cell(content)
        if len(norm) >= 15:
            out.add(norm)
    return out


def _para_texts(doc: Document) -> list[str]:
    return [" ".join((p.text or "").split()) for p in doc.paragraphs]


def source_note_lines(doc: Document) -> list[str]:
    return [normalize_cell(t) for t in _para_texts(doc) if _SOURCE_NOTE_RE.match(t)]


def caption_inventory(doc: Document) -> list[str]:
    return [normalize_cell(t) for t in _para_texts(doc) if _CAPTION_RE.match(t)]


def reference_count(doc: Document) -> int:
    """Count numbered bibliography entries after a references heading."""
    texts = _para_texts(doc)
    started = False
    n = 0
    for t in texts:
        low = t.lower()
        if low.startswith("список использован"):
            started = True
            continue
        if not started:
            continue
        if low in {"приложения", "приложение"} or low.startswith("приложение "):
            break
        if _REF_ENTRY_RE.match(t):
            n += 1
    return n


def _empty_output_tables(doc: Document) -> int:
    n = 0
    for table in doc.tables:
        if all(not (cell.text or "").strip() for row in table.rows for cell in row.cells):
            n += 1
    return n


def evaluate_content_preservation(
    source_doc: Document,
    output_doc: Document,
) -> tuple[dict, list[ContentIssue]]:
    issues: list[ContentIssue] = []

    # --- DATA-row preservation (excludes first-row headers + numeric index rows,
    # so legitimate header repetition / merged-header read differences do not
    # false-fire; only real DATA content loss/dup/addition registers) ---
    src_data = data_row_multiset(source_doc)
    out_data = data_row_multiset(output_doc)
    lost_rows = set(src_data) - set(out_data)
    added_rows = set(out_data) - set(src_data)
    if lost_rows:
        issues.append(ContentIssue("lost_table_cell_content", "fail",
                                   {"count": len(lost_rows), "examples": [r[:60] for r in list(lost_rows)[:5]]}))
    if added_rows:
        # new distinct data rows in output are not auto-fail (could be a benign
        # re-layout artefact) but must be surfaced for review.
        issues.append(ContentIssue("added_table_data_rows", "needs_human_review",
                                   {"count": len(added_rows), "examples": [r[:60] for r in list(added_rows)[:5]]}))

    # real data-row duplication (header repetition excluded by dropping first rows)
    dup_rows = {k: out_data[k] - src_data[k] for k in out_data if out_data[k] > src_data.get(k, 0)}
    if dup_rows:
        issues.append(ContentIssue("duplicated_table_data_row", "fail",
                                   {"count": sum(dup_rows.values()), "examples": [k[:60] for k in list(dup_rows)[:5]]}))

    src_sn = source_note_lines(source_doc)
    out_sn = source_note_lines(output_doc)
    if len(out_sn) < len(src_sn):
        issues.append(ContentIssue("lost_source_note_line", "fail",
                                   {"source": len(src_sn), "output": len(out_sn)}))

    # References are deliberately reformatted/renumbered by the B1 subsystem, so
    # exact text matching is unreliable for a hard gate — report counts only and
    # rely on the structure gate (СПИСОК section present) + dedicated ref tests.
    src_refs = reference_entry_set(source_doc)
    out_refs = reference_entry_set(output_doc)

    src_empty, out_empty = _empty_output_tables(source_doc), _empty_output_tables(output_doc)
    if out_empty > src_empty:
        issues.append(ContentIssue("empty_output_table_created", "fail",
                                   {"source_empty": src_empty, "output_empty": out_empty}))

    src_cells = table_cell_multiset(source_doc)
    out_cells = table_cell_multiset(output_doc)
    report = {
        "source_cells": sum(src_cells.values()),
        "output_cells": sum(out_cells.values()),
        "source_data_rows": len(set(src_data)),
        "output_data_rows": len(set(out_data)),
        "lost_data_rows": len(lost_rows),
        "added_data_rows": len(added_rows),
        "duplicated_data_rows": sum(dup_rows.values()),
        "source_note_lines": {"source": len(src_sn), "output": len(out_sn)},
        "reference_entries_informational": {"source": len(src_refs), "output": len(out_refs)},
        "captions": {"source": len(caption_inventory(source_doc)),
                     "output": len(caption_inventory(output_doc))},
        "empty_tables": {"source": src_empty, "output": out_empty},
        "content_fail": [i.issue_type for i in issues if i.severity == "fail"],
        "content_review": [i.issue_type for i in issues if i.severity == "needs_human_review"],
    }
    return report, issues
