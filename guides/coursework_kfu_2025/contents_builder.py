from __future__ import annotations

import logging
import re
import shutil
import tempfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .classifier import clean_spaces, find_body_start_index, parse_heading1, parse_heading2
from .layout_render import render_docx_to_pdf
from .page_breaks import apply_page_breaks
from .page_numbering import apply_page_numbering_policy
from .pdf_layout_analyzer import analyze_pdf_lines

logger = logging.getLogger(__name__)

_CONTENTS_HEADING_STRICT_RE = re.compile(
    r"^\s*(содержание|оглавление)\s*$", re.IGNORECASE
)
# Accepts trailing punctuation (`.`, `:`, `;`), dot leaders, ellipses, mid-dots
# and an optional page number — typical malformed TOC heading variants such as
# "Содержание.", "СОДЕРЖАНИЕ ……… 1", "Оглавление........... 2".
_CONTENTS_HEADING_LOOSE_RE = re.compile(
    r"^\s*(содержание|оглавление)"
    r"\s*[.:;]?"
    r"[.․‥…·•\s]*"
    r"\d{0,4}\s*$",
    re.IGNORECASE,
)
_SOURCE_NOTE_RE = re.compile(r"^\s*(источник|составлено по|рассчитано по|примечание)\s*:", re.IGNORECASE)
_TABLE_RE = re.compile(r"^\s*таблица\s+\d+(?:\.\d+){0,2}\b", re.IGNORECASE)
_FIGURE_RE = re.compile(r"^\s*(рис\.|рисунок)\s*\d+(?:\.\d+){0,2}\b", re.IGNORECASE)

_TOC_PAGE_TAIL_RE = re.compile(r"[\s.․‥…·•]+\d{1,4}\s*$")
_APPENDIX_LOCAL_RE = re.compile(r"^приложение\s+(?:\d{1,3}|[a-zа-яё])\b")
_H1_TOC_ENTRY_RE = re.compile(r"^\d+\.\s+\S")
_H2_TOC_ENTRY_RE = re.compile(r"^\d+\.\d+\.?\s+\S")

_STRUCTURAL_HEADINGS = {
    "введение": "ВВЕДЕНИЕ",
    "заключение": "ЗАКЛЮЧЕНИЕ",
    "список использованных источников": "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    "список использованной литературы": "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    "приложения": "ПРИЛОЖЕНИЯ",
}

_PAGE_PLACEHOLDER = "000"


@dataclass(frozen=True)
class TocEntry:
    title: str
    bookmark_name: str
    paragraph_index: int


def _norm(text: str) -> str:
    return clean_spaces(text).lower().rstrip(".")


def _first_paragraph_segment(text: str) -> str:
    """
    Return the first non-empty soft-break segment of a paragraph. Old TOC
    blocks sometimes survive as one paragraph with line-break-separated
    entries; the first segment is the heading we want to recognize.
    """
    cleaned = clean_spaces(text or "")
    if not cleaned:
        return ""
    return re.split(r"[\n\v]+", cleaned, maxsplit=1)[0].strip()


def _paragraph_segments(text: str) -> list[str]:
    cleaned = clean_spaces(text or "")
    return [s.strip() for s in re.split(r"[\n\v]+", cleaned) if s.strip()]


def _is_contents_heading(text: str) -> bool:
    """
    Pure-text Содержание/Оглавление detector. Tolerant of trailing punctuation,
    dot leaders, and a trailing page number, and of soft-break-joined input
    where the heading is the first line. Pure text check only — callers must
    apply `_paragraph_has_following_toc_evidence` to reject false positives
    (standalone Содержание-like text that does not lead a real TOC block).
    """
    if not text:
        return False
    first = _first_paragraph_segment(text)
    if not first:
        return False
    return bool(_CONTENTS_HEADING_LOOSE_RE.match(first))


def _is_intro_heading(text: str) -> bool:
    return _norm(text) == "введение"


def _is_toc_like_entry(text: str) -> bool:
    """
    True when *text* (a single non-empty line) reads like a TOC entry:
    a structural heading (ВВЕДЕНИЕ/ЗАКЛЮЧЕНИЕ/СПИСОК.../ПРИЛОЖЕНИЯ),
    a numbered chapter or sub-chapter (`1. ...`, `1.1. ...`), or a specific
    appendix label (`ПРИЛОЖЕНИЕ 1/2/3`). Trailing dot leaders and a trailing
    page number are tolerated.
    """
    t = clean_spaces(text or "").strip()
    if not t:
        return False
    t = _TOC_PAGE_TAIL_RE.sub("", t).strip()
    if not t:
        return False
    low = t.lower().rstrip(".").strip()
    if low in {"введение", "заключение", "приложения"}:
        return True
    if low.startswith("список использованных") or low.startswith("список использованной"):
        return True
    if _APPENDIX_LOCAL_RE.match(low):
        return True
    if _H2_TOC_ENTRY_RE.match(t):
        return True
    if _H1_TOC_ENTRY_RE.match(t):
        return True
    return False


def _paragraph_has_following_toc_evidence(
    document: Document, paragraph_text: str, idx: int, upper_bound: int
) -> bool:
    """
    Confirm that a Содержание-like heading paragraph leads a real TOC block by
    requiring at least one TOC-like entry either inside the same paragraph
    (soft-break segments) or in one of the subsequent paragraphs up to
    *upper_bound* exclusive. The scan stops at the next Содержание/Оглавление
    paragraph so a stray loose-only heading cannot borrow evidence from a
    later canonical TOC block. Used by body_start computation where the upper
    bound spans the whole document.
    """
    segments = _paragraph_segments(paragraph_text)
    for seg in segments[1:]:
        if _is_toc_like_entry(seg):
            return True
    paragraphs = document.paragraphs
    end = min(upper_bound, len(paragraphs))
    for j in range(idx + 1, end):
        next_text = paragraphs[j].text or ""
        if _is_contents_heading(next_text):
            break
        for seg in _paragraph_segments(next_text):
            if _is_toc_like_entry(seg):
                return True
    return False


def _is_safe_to_remove_pre_body_block(
    document: Document, paragraph_text: str, idx: int, body_start: int
) -> bool:
    """
    Removal-safety check for a Содержание-like paragraph at *idx*: every
    paragraph between *idx* (exclusive) and *body_start* (exclusive) must be
    either blank or TOC-like. This blocks aggressive deletion of unrelated
    front-matter body text while still allowing removal of an old TOC block
    that has been emptied of its entries (e.g. побитая Роман — heading +
    trailing blanks).
    """
    segments = _paragraph_segments(paragraph_text)
    for seg in segments[1:]:
        if not _is_toc_like_entry(seg):
            return False
    paragraphs = document.paragraphs
    end = min(body_start, len(paragraphs))
    for j in range(idx + 1, end):
        next_text = paragraphs[j].text or ""
        if _is_contents_heading(next_text):
            break
        cleaned = clean_spaces(next_text).strip()
        if not cleaned:
            continue
        for seg in _paragraph_segments(next_text):
            if not _is_toc_like_entry(seg):
                return False
    return True


def _find_body_start_index_for_contents(document: Document) -> int | None:
    """
    Find the real body introduction for TOC rebuild.

    The generic classifier intentionally treats a standalone "ВВЕДЕНИЕ" as the
    body start. Old hand-made TOCs sometimes contain exactly that standalone
    entry before the real body, so for contents replacement we skip everything
    after the last pre-body contents/oglavlenie heading and use the last
    standalone intro after it.
    """
    paragraphs = document.paragraphs
    upper = len(paragraphs)
    contents_indices: list[int] = []
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text or ""
        if not _is_contents_heading(text):
            continue
        if not _paragraph_has_following_toc_evidence(document, text, idx, upper):
            continue
        contents_indices.append(idx)

    if contents_indices:
        last_contents_idx = max(contents_indices)
        intro_candidates = [
            idx
            for idx in range(last_contents_idx + 1, len(paragraphs))
            if _is_intro_heading(paragraphs[idx].text)
        ]
        if intro_candidates:
            return intro_candidates[-1]

    return find_body_start_index(document)


def _is_appendix_local_heading(text: str) -> bool:
    return bool(re.match(r"^\s*приложение\s+(?:\d{1,3}|[a-zа-яё])\b", text, re.IGNORECASE))


def _is_excluded_service_line(text: str) -> bool:
    return bool(_SOURCE_NOTE_RE.match(text) or _TABLE_RE.match(text) or _FIGURE_RE.match(text))


def _style_name(paragraph) -> str:
    try:
        return (paragraph.style.name or "").lower()
    except Exception:
        return ""


def _is_heading1_style(paragraph) -> bool:
    return _style_name(paragraph) in {"heading 1", "заголовок 1"}


def _is_heading2_style(paragraph) -> bool:
    return _style_name(paragraph) in {"heading 2", "заголовок 2"}


def _paragraph_has_word_numbering(paragraph) -> bool:
    """True when the paragraph carries Word-list numbering (w:numPr)."""
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return False
    return p_pr.find(qn("w:numPr")) is not None


_SENTENCE_BOUNDARY_RE = re.compile(r"[.!?]\s+[А-ЯЁA-Z]")
_FALLBACK_TITLE_MAX_LEN = 200


def _heading_title_looks_like_body_prose(title: str) -> bool:
    """
    True when the title text after a "N." / "N.M." prefix reads like body
    prose rather than a heading. Used only by the fallback path.
    """
    t = clean_spaces(title)
    if not t:
        return True
    if len(t) > _FALLBACK_TITLE_MAX_LEN:
        return True
    if t.endswith((".", ":", ";", "?", "!")):
        return True
    if _SENTENCE_BOUNDARY_RE.search(t):
        return True
    return False


def _prev_nonempty_body_text(paragraphs, idx: int, body_start: int) -> str:
    j = idx - 1
    while j >= body_start:
        text = clean_spaces(paragraphs[j].text)
        if text:
            return text
        j -= 1
    return ""


def _is_fallback_heading1_candidate(
    paragraph,
    parsed_h1: dict,
    paragraphs,
    idx: int,
    body_start: int,
    last_h1_chapter: int | None,
) -> bool:
    """
    Accept a real `N. Title` paragraph as a TOC heading even when safe_formatter
    could not promote it to Heading 1 style. Guards mirror the documented plan
    so false positives stay blocked.
    """
    if _paragraph_has_word_numbering(paragraph):
        return False
    if _heading_title_looks_like_body_prose(parsed_h1.get("title") or ""):
        return False

    prev_text = _prev_nonempty_body_text(paragraphs, idx, body_start)
    if prev_text.endswith(":"):
        return False

    chapter_num = parsed_h1["chapter_num"]
    expected = 1 if last_h1_chapter is None else last_h1_chapter + 1
    return chapter_num == expected


def _is_fallback_heading2_candidate(
    paragraph,
    parsed_h2: dict,
    paragraphs,
    idx: int,
    body_start: int,
    last_h1_chapter: int | None,
    last_h2_in_chapter: int | None,
) -> bool:
    """
    Accept a real `N.M. Title` paragraph as a TOC heading even when
    safe_formatter could not promote it to Heading 2 style.
    """
    if _paragraph_has_word_numbering(paragraph):
        return False
    if _heading_title_looks_like_body_prose(parsed_h2.get("title") or ""):
        return False

    prev_text = _prev_nonempty_body_text(paragraphs, idx, body_start)
    if prev_text.endswith(":"):
        return False

    if last_h1_chapter is None:
        return False
    if parsed_h2["chapter_num"] != last_h1_chapter:
        return False

    expected = 1 if last_h2_in_chapter is None else last_h2_in_chapter + 1
    return parsed_h2["paragraph_num"] == expected


def _collect_body_entries(document: Document, body_start: int) -> list[TocEntry]:
    entries: list[TocEntry] = []
    paragraphs = document.paragraphs

    last_h1_chapter: int | None = None
    last_h2_in_chapter: int | None = None

    for idx, paragraph in enumerate(paragraphs[body_start:], start=body_start):
        text = clean_spaces(paragraph.text)
        if not text or _is_excluded_service_line(text):
            continue

        low = _norm(text)
        if low in _STRUCTURAL_HEADINGS:
            canonical = _STRUCTURAL_HEADINGS[low]
            entries.append(TocEntry(canonical, f"kpfu_toc_{len(entries) + 1}", idx))
            if canonical == "ПРИЛОЖЕНИЯ":
                break
            continue

        if _is_appendix_local_heading(text):
            continue

        parsed_h1 = parse_heading1(text)
        if parsed_h1 and parsed_h1["kind"] == "heading1_chapter":
            if _is_heading1_style(paragraph):
                entries.append(TocEntry(text, f"kpfu_toc_{len(entries) + 1}", idx))
                last_h1_chapter = parsed_h1["chapter_num"]
                last_h2_in_chapter = None
                continue
            if _is_fallback_heading1_candidate(
                paragraph, parsed_h1, paragraphs, idx, body_start, last_h1_chapter
            ):
                entries.append(TocEntry(text, f"kpfu_toc_{len(entries) + 1}", idx))
                last_h1_chapter = parsed_h1["chapter_num"]
                last_h2_in_chapter = None
                continue

        parsed_h2 = parse_heading2(text)
        if parsed_h2:
            if _is_heading2_style(paragraph):
                entries.append(TocEntry(text, f"kpfu_toc_{len(entries) + 1}", idx))
                if (
                    last_h1_chapter is not None
                    and parsed_h2["chapter_num"] == last_h1_chapter
                ):
                    last_h2_in_chapter = parsed_h2["paragraph_num"]
                continue
            if _is_fallback_heading2_candidate(
                paragraph,
                parsed_h2,
                paragraphs,
                idx,
                body_start,
                last_h1_chapter,
                last_h2_in_chapter,
            ):
                entries.append(TocEntry(text, f"kpfu_toc_{len(entries) + 1}", idx))
                last_h2_in_chapter = parsed_h2["paragraph_num"]
                continue

    return entries


def _find_existing_contents_start(document: Document, body_start: int) -> int | None:
    paragraphs = document.paragraphs[:body_start]
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text or ""
        if not _is_contents_heading(text):
            continue
        if not _is_safe_to_remove_pre_body_block(document, text, idx, body_start):
            continue
        return idx
    return None


def _remove_body_children_range(document: Document, start_elem, end_elem) -> None:
    body = document.element.body
    children = list(body)
    start_idx = children.index(start_elem)
    end_idx = children.index(end_elem)
    for child in children[start_idx:end_idx]:
        body.remove(child)


def _is_toc_sdt_block(sdt_elem) -> bool:
    """
    Identify a `<w:sdt>` block that wraps a Word-managed Table of Contents.

    Three accepted markers (any one is sufficient):
      * `<w:docPartGallery w:val="Table of Contents"/>` in sdtPr (Word's
        canonical TOC SDT — observed in real student docs).
      * any nested `<w:fldChar>` referencing a `TOC` field (legacy field-code
        TOC wrapped in an SDT).
      * a nested `<w:fldSimple w:instr="TOC ...">` element.
    """
    properties = sdt_elem.find(qn("w:sdtPr"))
    if properties is not None:
        doc_part_obj = properties.find(qn("w:docPartObj"))
        if doc_part_obj is not None:
            gallery = doc_part_obj.find(qn("w:docPartGallery"))
            if gallery is not None:
                value = (gallery.get(qn("w:val")) or "").strip().lower()
                if value == "table of contents":
                    return True
    # Field-code based TOCs nested inside SDT
    for instr_text in sdt_elem.findall(".//" + qn("w:instrText")):
        if "TOC" in (instr_text.text or "").upper():
            return True
    for fld_simple in sdt_elem.findall(".//" + qn("w:fldSimple")):
        instr = (fld_simple.get(qn("w:instr")) or "").upper()
        if "TOC" in instr:
            return True
    return False


def _remove_word_managed_toc_blocks(document: Document) -> int:
    """
    Remove every Word-managed TOC `<w:sdt>` block from the document body.

    `python-docx`'s `document.paragraphs` reports only top-level `<w:p>` body
    children; paragraphs nested inside `<w:sdt>` are invisible to the
    contents detection / removal helpers above. As a result a Word-managed
    Table of Contents SDT survives the canonical TOC rebuild and the
    rendered file contains two TOCs.

    This helper drops the entire SDT element (the whole TOC container,
    not its inner paragraphs) for every SDT block that satisfies
    `_is_toc_sdt_block`. Other SDT blocks (form fields, content controls
    that are NOT a TOC) are untouched.

    Returns the number of SDT blocks removed.
    """
    body = document.element.body
    removed = 0
    for sdt_elem in list(body.findall(qn("w:sdt"))):
        if not _is_toc_sdt_block(sdt_elem):
            continue
        body.remove(sdt_elem)
        removed += 1
    return removed


def _set_run_font(run, *, bold: bool) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.bold = bold
    run.italic = False
    run.underline = False
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")


def _run_pr_xml(*, bold: bool = False) -> OxmlElement:
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Times New Roman")
    r_fonts.set(qn("w:hAnsi"), "Times New Roman")
    r_fonts.set(qn("w:cs"), "Times New Roman")
    r_pr.append(r_fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "28")
    r_pr.append(size)

    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "28")
    r_pr.append(size_cs)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    r_pr.append(underline)

    if bold:
        r_pr.append(OxmlElement("w:b"))
        r_pr.append(OxmlElement("w:bCs"))

    return r_pr


def _append_hyperlink_run(hyperlink: OxmlElement, text: str | None = None, *, tab: bool = False) -> None:
    run = OxmlElement("w:r")
    run.append(_run_pr_xml(bold=False))
    if tab:
        run.append(OxmlElement("w:tab"))
    else:
        text_el = OxmlElement("w:t")
        text_el.text = text or ""
        if text and (text.startswith(" ") or text.endswith(" ")):
            text_el.set(qn("xml:space"), "preserve")
        run.append(text_el)
    hyperlink.append(run)


def _set_internal_hyperlink_text(paragraph, title: str, page: str, anchor: str) -> None:
    p = paragraph._element
    p_pr = p.get_or_add_pPr()
    for child in list(p):
        if child is not p_pr:
            p.remove(child)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    _append_hyperlink_run(hyperlink, title)
    _append_hyperlink_run(hyperlink, tab=True)
    _append_hyperlink_run(hyperlink, page)
    p.append(hyperlink)


def _clear_tab_stops(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    for tabs in list(p_pr.findall(qn("w:tabs"))):
        p_pr.remove(tabs)


def _set_toc_entry_tab_stop(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    _clear_tab_stops(paragraph)
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(Cm(16).twips))
    tabs.append(tab)
    p_pr.append(tabs)


def _apply_zero_indent_xml(paragraph) -> None:
    """
    Wipe inherited w:ind and write a fresh element with all-zero attributes so
    left/right/firstLine/hanging cannot drift in via styles or python-docx
    defaults.
    """
    p_pr = paragraph._element.get_or_add_pPr()
    for old in list(p_pr.findall(qn("w:ind"))):
        p_pr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:hanging"), "0")
    p_pr.append(ind)


def _format_contents_heading(paragraph) -> None:
    paragraph.text = "СОДЕРЖАНИЕ"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.page_break_before = False
    _clear_tab_stops(paragraph)
    _apply_zero_indent_xml(paragraph)
    for run in paragraph.runs:
        _set_run_font(run, bold=True)


def _format_toc_blank_paragraph(paragraph) -> None:
    """
    The methodical layout (Приложение 3) requires exactly one blank paragraph
    between the СОДЕРЖАНИЕ title and the first TOC entry. Match entry-level
    line spacing so the blank reads as a single empty TOC line.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.page_break_before = False
    _clear_tab_stops(paragraph)
    _apply_zero_indent_xml(paragraph)
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        _set_run_font(run, bold=False)


def _format_toc_entry(paragraph, title: str, page: str) -> None:
    paragraph.text = f"{title}\t{page}"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.page_break_before = False
    _set_toc_entry_tab_stop(paragraph)
    _apply_zero_indent_xml(paragraph)
    for run in paragraph.runs:
        _set_run_font(run, bold=False)


def _format_toc_entry_link(paragraph, entry: TocEntry, page: str) -> None:
    _format_toc_entry(paragraph, entry.title, page)
    _set_internal_hyperlink_text(paragraph, entry.title, page, entry.bookmark_name)


def _move_paragraphs_before(document: Document, paragraphs: list, reference_elem) -> None:
    body = document.element.body
    insert_idx = list(body).index(reference_elem)
    for paragraph in paragraphs:
        elem = paragraph._element
        body.remove(elem)
        body.insert(insert_idx, elem)
        insert_idx += 1


def _insert_contents_block(document: Document, entries: list[TocEntry], pages: dict[str, int | str]) -> int:
    body_start = _find_body_start_index_for_contents(document)
    if body_start is None:
        raise ValueError("real ВВЕДЕНИЕ heading not found")

    # Step 1: drop every Word-managed TOC `<w:sdt>` block from the body.
    # `document.paragraphs` is blind to SDT-nested content, so without this
    # step the previous Word TOC survives the rebuild and the rendered file
    # contains two TOCs side-by-side. Done BEFORE recomputing the paragraph
    # list because removing body children renumbers indices.
    sdt_removed = _remove_word_managed_toc_blocks(document)

    paragraphs = document.paragraphs
    if sdt_removed:
        # Re-locate body_start: paragraph indices shift when SDT-wrapped
        # children disappear. Re-running the detector keeps subsequent helpers
        # consistent.
        body_start = _find_body_start_index_for_contents(document)
        if body_start is None:
            raise ValueError("real ВВЕДЕНИЕ heading not found after SDT cleanup")
        paragraphs = document.paragraphs

    intro_elem = paragraphs[body_start]._element
    contents_start = _find_existing_contents_start(document, body_start)
    removed_count = 0
    if contents_start is not None:
        removed_count = sum(
            1
            for paragraph in paragraphs[contents_start:body_start]
            if _is_contents_heading(paragraph.text)
        )
        _remove_body_children_range(document, paragraphs[contents_start]._element, intro_elem)
    if sdt_removed:
        logger.info(
            "contents_rebuild_sdt_toc_removed path=<doc> count=%d", sdt_removed
        )

    new_paragraphs = [document.add_paragraph()]
    _format_contents_heading(new_paragraphs[0])

    blank = document.add_paragraph()
    _format_toc_blank_paragraph(blank)
    new_paragraphs.append(blank)

    for entry in entries:
        page = str(pages.get(entry.bookmark_name, _PAGE_PLACEHOLDER))
        p = document.add_paragraph()
        _format_toc_entry_link(p, entry, page)
        new_paragraphs.append(p)

    _move_paragraphs_before(document, new_paragraphs, intro_elem)
    return removed_count


def _line_norm(text: str) -> str:
    return clean_spaces(text).lower()


def _match_norm(text: str) -> str:
    text = _line_norm(text).replace("ё", "е")
    return re.sub(r"[^0-9a-zа-я]+", " ", text).strip()


def _is_heading_window_match(target: str, rendered: str) -> bool:
    target_norm = _match_norm(target)
    rendered_norm = _match_norm(rendered)
    if not target_norm or not rendered_norm:
        return False
    if rendered_norm == target_norm:
        return True
    if target_norm.startswith(rendered_norm) and len(rendered_norm) >= min(24, len(target_norm)):
        return True
    if rendered_norm.startswith(target_norm) and len(target_norm) >= 12:
        return True
    target_tokens = target_norm.split()
    rendered_tokens = rendered_norm.split()
    if len(target_tokens) >= 4 and len(rendered_tokens) >= 4:
        return rendered_tokens[:4] == target_tokens[:4] and len(rendered_norm) >= min(24, len(target_norm))
    return False


def _find_rendered_page(title: str, lines, *, min_page: int) -> int | None:
    target = _line_norm(title)
    if not target:
        return None
    for line in lines:
        if line.page_num < min_page:
            continue
        if _line_norm(line.text) == target:
            return line.page_num

    eligible = [line for line in lines if line.page_num >= min_page]
    for idx, line in enumerate(eligible):
        if _is_heading_window_match(title, line.text):
            return line.page_num
        joined = line.text
        for next_line in eligible[idx + 1:idx + 3]:
            if next_line.page_num != line.page_num:
                break
            joined = f"{joined} {next_line.text}"
            if _is_heading_window_match(title, joined):
                return line.page_num
    return None


def _validate_resolved_pages(entries: list[TocEntry], pages: dict[str, int]) -> None:
    ordered = [pages[entry.bookmark_name] for entry in entries]
    if not ordered:
        raise ValueError("no resolved TOC pages")

    if ordered[0] != 3:
        raise ValueError(f"resolved intro page must be 3, got {ordered[0]}")

    for prev, current in zip(ordered, ordered[1:]):
        if current < prev:
            raise ValueError(f"resolved TOC pages are not nondecreasing: {ordered}")

    if len(ordered) > 2 and len(set(ordered)) == 1:
        raise ValueError(f"degenerate resolved TOC pages: {ordered}")


def _resolve_display_pages(entries: list[TocEntry], lines) -> dict[str, int]:
    intro_pages = [
        line.page_num
        for line in lines
        if _line_norm(line.text) == "введение"
    ]
    if not intro_pages:
        raise ValueError("rendered ВВЕДЕНИЕ page not found")

    rendered_intro_page = max(intro_pages)
    offset = 3 - rendered_intro_page
    resolved: dict[str, int] = {}
    unresolved: list[str] = []
    min_page = rendered_intro_page
    for entry in entries:
        if _norm(entry.title) == "введение":
            page = rendered_intro_page
        else:
            page = _find_rendered_page(entry.title, lines, min_page=min_page)
        if page is None:
            unresolved.append(entry.title)
            continue
        resolved[entry.bookmark_name] = page + offset
        min_page = page
    if unresolved:
        logger.warning("contents_rebuild_unresolved headings=%s", unresolved)
        raise ValueError(f"rendered page not found for TOC entries: {unresolved}")
    _validate_resolved_pages(entries, resolved)
    return resolved


def _reapply_front_matter_layout(document: Document) -> None:
    from .safe_formatter import ensure_appendices_section_layout, ensure_front_matter_layout, normalize_sections

    body_start = _find_body_start_index_for_contents(document)
    if body_start is None:
        raise ValueError("real ВВЕДЕНИЕ heading not found")
    normalize_sections(document)
    ensure_front_matter_layout(document, body_start)
    apply_page_breaks(document, body_start)
    ensure_appendices_section_layout(document, body_start)
    apply_page_numbering_policy(document)


def _replace_contents_entry_pages(document: Document, entries: list[TocEntry], pages: dict[str, int]) -> None:
    body_start = _find_body_start_index_for_contents(document)
    if body_start is None:
        raise ValueError("real ВВЕДЕНИЕ heading not found")

    contents_start = _find_existing_contents_start(document, body_start)
    if contents_start is None:
        raise ValueError("draft contents block not found")

    # The draft block is [title, blank, entry, entry, ...]. Skip blanks so the
    # second-pass page rewrite stays aligned with `entries` even when the
    # methodical-required blank paragraph is present after the title.
    entry_paragraphs = [
        paragraph
        for paragraph in document.paragraphs[contents_start + 1:body_start]
        if clean_spaces(paragraph.text)
    ]
    if len(entry_paragraphs) != len(entries):
        raise ValueError("draft contents entry count changed")

    for paragraph, entry in zip(entry_paragraphs, entries):
        _format_toc_entry_link(paragraph, entry, str(pages[entry.bookmark_name]))


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    p = paragraph._element
    for existing in p.findall(qn("w:bookmarkStart")):
        if existing.get(qn("w:name")) == name:
            return

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    insert_at = 1 if p.pPr is not None else 0
    p.insert(insert_at, start)
    p.append(end)


def _add_body_bookmarks(document: Document, entries: list[TocEntry]) -> None:
    paragraphs = document.paragraphs
    for idx, entry in enumerate(entries, start=1000):
        if 0 <= entry.paragraph_index < len(paragraphs):
            _add_bookmark(paragraphs[entry.paragraph_index], entry.bookmark_name, idx)


def strip_obsolete_toc_blocks_inplace(docx_path: str | Path) -> dict:
    """
    Strip obsolete TOC artifacts from a DOCX on disk and save back to the same
    path. DOCX-only operation: no LibreOffice / PDF rendering involved.

    Always-safe pre-pass that complements `rebuild_static_contents_page`.
    The rebuild pipeline can fail mid-way (PDF render unavailable, degenerate
    page mapping, etc.) and currently leaves the source untouched; without
    this pre-pass that fail-safe meant the user saw the original TOC even
    after our patch removed it inside the rebuild. By stripping obsolete
    TOC artifacts FIRST and saving directly, we guarantee that the user
    never gets two TOCs: at worst (rebuild also fails) they get a document
    without a canonical TOC, which is recoverable.

    Removes:
      * Every Word-managed Table-of-Contents `<w:sdt>` block in the body
        (`_is_toc_sdt_block` — three accepted markers).
      * The plain-text old TOC block between a standalone Содержание /
        Оглавление paragraph and the real body intro, when the block is
        `_is_safe_to_remove_pre_body_block`.

    Returns a small report dict so callers can log. Source DOCX is touched
    only when at least one obsolete artifact was found.
    """
    source_path = Path(docx_path)
    document = Document(str(source_path))

    sdt_removed = _remove_word_managed_toc_blocks(document)

    plain_removed = 0
    body_start = _find_body_start_index_for_contents(document)
    if body_start is not None:
        contents_start = _find_existing_contents_start(document, body_start)
        if contents_start is not None:
            paragraphs = document.paragraphs
            plain_removed = body_start - contents_start
            _remove_body_children_range(
                document,
                paragraphs[contents_start]._element,
                paragraphs[body_start]._element,
            )

    if sdt_removed or plain_removed:
        document.save(str(source_path))
        logger.info(
            "obsolete_toc_stripped path=%s sdt=%d plain_paragraphs=%d",
            source_path,
            sdt_removed,
            plain_removed,
        )

    return {"sdt_removed": sdt_removed, "plain_toc_removed": plain_removed}


def rebuild_static_contents_page(docx_path: str | Path) -> bool:
    """
    Rebuild static KFU contents page. The source DOCX is replaced only after all
    TOC page numbers are resolved from a rendered draft; failures leave it intact.
    """
    source_path = Path(docx_path)
    workdir = Path(tempfile.mkdtemp(prefix="kpfu_contents_"))
    pdf_path: Path | None = None
    try:
        work_path = workdir / source_path.name
        shutil.copy2(source_path, work_path)

        document = Document(str(work_path))
        body_start = _find_body_start_index_for_contents(document)
        if body_start is None:
            logger.warning("contents_rebuild_skipped reason=intro_not_found path=%s", source_path)
            return False

        entries = _collect_body_entries(document, body_start)
        if not entries:
            logger.warning("contents_rebuild_skipped reason=no_entries path=%s", source_path)
            return False
        logger.info("contents_rebuild_entries_collected path=%s count=%d", source_path, len(entries))

        _add_body_bookmarks(document, entries)
        removed_count = _insert_contents_block(
            document,
            entries,
            {entry.bookmark_name: _PAGE_PLACEHOLDER for entry in entries},
        )
        logger.info("contents_rebuild_old_toc_removed path=%s count=%d", source_path, removed_count)
        _reapply_front_matter_layout(document)
        document.save(str(work_path))

        pdf_path = render_docx_to_pdf(work_path)
        lines = analyze_pdf_lines(pdf_path)
        pages = _resolve_display_pages(entries, lines)
        logger.info("contents_rebuild_render_resolved path=%s count=%d", source_path, len(pages))

        document = Document(str(work_path))
        _replace_contents_entry_pages(document, entries, pages)
        _reapply_front_matter_layout(document)
        document.save(str(work_path))

        shutil.copy2(work_path, source_path)
        logger.info("contents_rebuild_applied path=%s entries=%s", source_path, len(entries))
        return True
    except Exception as exc:
        logger.warning("contents_rebuild_skipped reason=%s path=%s", exc, source_path)
        return False
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)
        shutil.rmtree(workdir, ignore_errors=True)
