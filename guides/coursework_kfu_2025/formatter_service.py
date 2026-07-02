import logging
import os
import shutil
import tempfile
from pathlib import Path

from docx import Document

from .safe_formatter import process_document
from .pagination_rules import apply_pagination_rules
from .table_continuation import (
    apply_table_merging,
    apply_table_continuation,
    apply_rendered_table_continuation,
    apply_rule3_table_orphan,
    apply_rule4_empty_first_lines,
    apply_rule6_figure_orphan,
    remove_empty_before_figure_captions,
    restore_docx_if_same_page_continuation_markers,
    remove_same_page_continuation_markers_inplace,
    warn_same_page_continuation_marker_violations,
    repair_manual_chain_overflow_before_marker,
    normalize_exact_grid_same_page_repeated_fragments_inplace,
    normalize_compatible_grid_same_page_repeated_fragments_inplace,
    cleanup_same_page_incompatible_chains_inplace,
    cleanup_same_page_continuation_blockers_inplace,
    cleanup_cross_page_without_marker_blockers_inplace,
    normalize_continuation_semantic_header_inplace,
    normalize_fragment_grid_widths_inplace,
    apply_rendered_table_start_orphan_guard,
)
from .contents_builder import rebuild_static_contents_page, strip_obsolete_toc_blocks_inplace
from .docx_utils import FormattingReport
from .layout_render import render_docx_to_pdf
from .pdf_layout_analyzer import analyze_pdf_lines
from .rendered_table_validation import (
    RenderedContinuationViolation,
    TableLayoutBlocker,
    build_rendered_table_identities,
    evaluate_table_layout_acceptance,
    validate_rendered_continuations,
)
from .document_structure_validation import (
    evaluate_document_structure,
    source_has_appendix,
    source_has_toc,
)

logger = logging.getLogger(__name__)

_TRUTHY = {"1", "true", "yes", "on"}


def _flag_value(name: str) -> str:
    return os.getenv(name, "<unset>")


def _rendered_table_continuation_enabled() -> bool:
    """
    Stage 0 conservative table mode.

    Default (env unset / falsey): the risky rendered table-continuation
    insertion path and the same-page fragment merge passes are skipped, so the
    formatter never creates a same-page ``Продолжение таблицы`` split or a
    synthetic numeric row — it prefers leaving a table whole. Set
    ``KPFU_RENDERED_TABLE_CONTINUATION=1`` to re-enable the experimental path
    (still subject to the rendered acceptance gate).
    """
    return os.getenv("KPFU_RENDERED_TABLE_CONTINUATION", "").strip().lower() in _TRUTHY


def _rendered_continuation_violations_for_docx(
    docx_path: Path,
    *,
    source_table_identities: list | None = None,
) -> list[RenderedContinuationViolation]:
    pdf_path: Path | None = None
    try:
        pdf_path = render_docx_to_pdf(Path(docx_path))
        pdf_lines = analyze_pdf_lines(pdf_path)
        doc = Document(str(docx_path))
        identities = build_rendered_table_identities(doc)
        return validate_rendered_continuations(
            pdf_lines,
            identities,
            source_table_identities=source_table_identities,
        )
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)


def _rendered_continuation_warning(
    violation: RenderedContinuationViolation,
) -> str | None:
    table_num = violation.table_num or f"#{violation.table_index}"
    if (
        violation.violation_type == "missing_continuation_marker"
        and violation.confidence == "high"
    ):
        return (
            f"Проверьте перенос таблицы {table_num}: "
            f"стр. {violation.page} без маркера продолжения."
        )
    if (
        violation.violation_type == "suspected_missing_continuation_marker"
        and violation.confidence == "medium"
    ):
        return (
            f"Проверьте возможный перенос таблицы {table_num}: "
            f"стр. {violation.page} без маркера продолжения."
        )
    if (
        violation.violation_type == "late_continuation_marker"
        and violation.confidence == "high"
    ):
        return (
            f"Проверьте перенос таблицы {table_num}: "
            f"на стр. {violation.page} строки идут до маркера продолжения."
        )
    if (
        violation.violation_type == "same_page_repeated_fragment"
        and violation.confidence == "high"
    ):
        return (
            f"Проверьте таблицу {table_num}: "
            f"на стр. {violation.page} повторный фрагмент виден на той же странице."
        )
    if (
        violation.violation_type == "same_page_adjacent_fragment"
        and violation.confidence == "high"
    ):
        return (
            f"Проверьте таблицу {table_num}: "
            f"на стр. {violation.page} соседний фрагмент выглядит как часть той же таблицы."
        )
    if violation.violation_type == "missing_or_late_continuation_marker":
        return (
            f"Проверьте перенос таблицы {table_num}: "
            f"стр. {violation.page} без корректного маркера продолжения."
        )
    if violation.violation_type == "ambiguous_adjacent_tables":
        return (
            f"Проверьте таблицу {table_num}: "
            "соседние таблицы похожи на фрагменты, но связь не доказана."
        )
    if violation.violation_type == "source_bad_duplicated_content_rows":
        return (
            f"Проверьте таблицу {table_num}: "
            "в исходном файле есть повторяющиеся содержательные строки."
        )
    return None


def _append_rendered_continuation_warnings(
    report: FormattingReport,
    violations: list[RenderedContinuationViolation],
) -> None:
    for violation in violations:
        message = _rendered_continuation_warning(violation)
        if message is None:
            continue
        report.warn(message)
        logger.warning(
            "rendered_continuation_violation table_num=%s table_index=%s page=%s type=%s confidence=%s evidence=%s",
            violation.table_num,
            violation.table_index,
            violation.page,
            violation.violation_type,
            violation.confidence,
            violation.evidence,
        )


_LAYOUT_BLOCKER_MESSAGES = {
    "same_page_continuation": "таблица {num}: продолжение таблицы оказалось на той же странице (стр. {page}).",
    "single_table_crosses_pages_without_marker": "таблица {num}: переходит на следующую страницу без маркера «Продолжение таблицы» (стр. {page}).",
    "orphaned_header_row": "таблица {num}: шапка таблицы осталась без строк данных (стр. {page}).",
    "fragment_grid_mismatch": "таблица {num}: фрагменты одной таблицы имеют разную сетку столбцов.",
    "same_page_repeated_header": "таблица {num}: шапка таблицы повторяется на одной странице (стр. {page}).",
    "appendix_label_not_on_new_page": "приложение начинается не с новой страницы (стр. {page}).",
    "cell_text_overflow_or_illegible_squeeze": "таблица {num}: столбцы выглядят сжатыми, текст переносится по буквам (стр. {page}).",
}


def _emit_table_layout_acceptance_warnings(
    output_path: Path,
    report: FormattingReport,
    source_identities: list | None = None,
    source_text: str | None = None,
) -> list[TableLayoutBlocker]:
    """
    Render the final document once and surface rendered table-layout blockers AND
    document-structure regressions (missing TOC / required sections / appendices).

    Visible defects become structured, severity-bearing blockers that are logged
    and added to the user-facing report. ``fail``-level blockers mean the
    rendered output is NO-GO; the smoke/deploy decision reads them via
    ``evaluate_table_layout_acceptance`` / ``evaluate_document_structure``.
    ``format_docx`` still returns the file.
    """
    pdf_path: Path | None = None
    structure_issues = []
    try:
        pdf_path = render_docx_to_pdf(output_path)
        pdf_lines = analyze_pdf_lines(pdf_path)
        doc = Document(str(output_path))
        identities = build_rendered_table_identities(doc)
        blockers = evaluate_table_layout_acceptance(
            pdf_lines, identities, doc=doc, source_identities=source_identities
        )
        if source_text is not None:
            structure_issues = evaluate_document_structure(
                pdf_lines,
                expect_toc=source_has_toc(source_text),
                expect_appendix=source_has_appendix(source_text),
            )
    except Exception:
        logger.exception("format_docx: table/structure acceptance gate failed to evaluate")
        return []
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)

    for issue in structure_issues:
        if issue.severity == "fail":
            report.warn(
                f"Проверьте структуру документа — нарушение «{issue.issue_type}»"
                + (f" (стр. {issue.page})" if issue.page else "")
            )
        logger.warning(
            "document_structure_issue type=%s severity=%s page=%s evidence=%s",
            issue.issue_type, issue.severity, issue.page, issue.evidence,
        )

    fail_count = sum(1 for b in blockers if b.severity == "fail")
    review_count = sum(1 for b in blockers if b.severity == "needs_human_review")
    for blocker in blockers:
        template = _LAYOUT_BLOCKER_MESSAGES.get(blocker.blocker_type)
        if template is not None:
            detail = template.format(num=blocker.table_num or "?", page=blocker.page or "?")
            prefix = "Проверьте" if blocker.severity == "fail" else "Возможно, проверьте"
            report.warn(f"{prefix} вёрстку таблиц — {detail}")
        logger.warning(
            "table_layout_blocker type=%s severity=%s table_num=%s page=%s evidence=%s",
            blocker.blocker_type,
            blocker.severity,
            blocker.table_num,
            blocker.page,
            blocker.evidence,
        )
    logger.warning(
        "format_docx: table layout acceptance gate fail=%d needs_review=%d",
        fail_count,
        review_count,
    )
    return blockers


def format_docx(input_path: str, output_path: str) -> tuple[str, list[str]]:
    """
    Format *input_path* and write the result to *output_path*.

    Returns:
        (output_path_str, warnings) where *warnings* is a (possibly empty)
        list of short Russian strings describing issues the user should
        check manually (e.g. tables that could not be auto-split).
    """
    input_path = Path(input_path)
    output_path = Path(output_path)

    if not input_path.exists():
        raise FileNotFoundError(f"Файл не найден: {input_path}")

    if input_path.suffix.lower() != ".docx":
        raise ValueError("Поддерживаются только .docx файлы")

    report = FormattingReport()
    source_text: str | None = None
    try:
        source_doc = Document(str(input_path))
        source_table_identities = build_rendered_table_identities(source_doc)
        source_text = "\n".join((p.text or "") for p in source_doc.paragraphs)
    except Exception:
        source_table_identities = None

    # Phase 1: structural formatting
    process_document(input_path, output_path)

    if not output_path.exists():
        raise RuntimeError("Файл не был создан после Phase 1")

    # Phase 2: pagination rules (keep_with_next flags)
    try:
        doc = Document(str(output_path))
        apply_pagination_rules(doc)
        doc.save(str(output_path))
        logger.info("format_docx: phase2 pagination rules applied")
    except Exception:
        logger.exception("format_docx: phase2 failed, skipping (phase1 result preserved)")

    # Phase 3: DOCX-only cleanup/normalisation.
    try:
        doc = Document(str(output_path))
        logger.info(
            "format_docx: phase3_start output_path=%s tables=%d marker_split_enabled=%s marker_split_apply=%s",
            output_path,
            len(doc.tables),
            _flag_value("KPFU_ENABLE_MARKER_SPLIT"),
            _flag_value("KPFU_APPLY_MARKER_SPLIT"),
        )
        n_merged  = apply_table_merging(doc)
        n_tables  = apply_table_continuation(doc, report=report)
        n_rule3   = apply_rule3_table_orphan(doc)
        n_rule4   = apply_rule4_empty_first_lines(doc)
        n_rule6   = apply_rule6_figure_orphan(doc)
        n_gap     = remove_empty_before_figure_captions(doc)
        if n_merged > 0 or n_tables > 0 or n_rule3 > 0 or n_rule4 > 0 or n_rule6 > 0 or n_gap > 0:
            doc.save(str(output_path))
            logger.info(
                "format_docx: phase3 merged=%d tables=%d rule3=%d rule4=%d rule6=%d gap=%d",
                n_merged, n_tables, n_rule3, n_rule4, n_rule6, n_gap,
            )
        else:
            logger.info("format_docx: phase3 no changes")
    except Exception:
        logger.exception("format_docx: phase3 failed, skipping (phase2 result preserved)")

    # Strip obsolete TOC artifacts (Word SDT TOC and plain-text old TOC block)
    # BEFORE the rendered-continuation backup.  By stripping first the backup
    # becomes: DOCX-only continuation markers + canonical TOC (rebuilt below).
    # If the rendered-continuation gate fires and restores this backup the user
    # always sees exactly one canonical СОДЕРЖАНИЕ — never a stale hand-made
    # copy that was present in the source or from a prior formatting run.
    try:
        report_strip = strip_obsolete_toc_blocks_inplace(output_path)
        if report_strip["sdt_removed"] or report_strip["plain_toc_removed"]:
            logger.info(
                "format_docx: pre-backup TOC strip sdt=%d plain_paragraphs=%d",
                report_strip["sdt_removed"],
                report_strip["plain_toc_removed"],
            )
    except Exception:
        logger.exception("format_docx: pre-backup TOC strip failed, continuing")

    # Rebuild the canonical contents page BEFORE taking the rendered-continuation
    # backup.  The backup therefore contains a freshly resolved СОДЕРЖАНИЕ, so
    # gate restoration gives the user a clean, correct document instead of one
    # that either lacks a TOC or retains the old hand-made copy.
    try:
        if rebuild_static_contents_page(output_path):
            logger.info("format_docx: pre-backup canonical TOC rebuilt")
    except Exception:
        logger.exception("format_docx: pre-backup canonical TOC rebuild failed, continuing")

    # Stage 0 conservative table mode: only the risky rendered continuation
    # *insertion* path is gated off by default (it can create same-page
    # «Продолжение таблицы» splits / synthesized numeric rows). The same-page
    # merge-back normalizers below are CLEANUP that *removes* same-page splits
    # (e.g. unnecessary student manual chains that now fit on one page), so they
    # run unconditionally — with their own rendered rollback.
    if not _rendered_table_continuation_enabled():
        logger.info(
            "format_docx: conservative table mode — rendered continuation insertion skipped "
            "(KPFU_RENDERED_TABLE_CONTINUATION unset); same-page merge-back still runs"
        )
    else:
        # Rendered table continuation entry.  The backup is taken from the
        # stripped + rebuilt state so that any gate restoration returns a
        # document that already has exactly one canonical СОДЕРЖАНИЕ and zero
        # same-page continuation violations.
        table_gate_backup_dir: Path | None = None
        table_gate_backup_path: Path | None = None
        try:
            table_gate_backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_format_table_gate_"))
            table_gate_backup_path = table_gate_backup_dir / output_path.name
            shutil.copy2(output_path, table_gate_backup_path)
            n_rendered = apply_rendered_table_continuation(output_path, report=report)
            if n_rendered:
                logger.info("format_docx: rendered table continuation splits=%d", n_rendered)
        except Exception:
            logger.exception("format_docx: rendered table continuation failed")

        # Safety gate: if the rendered-continuation markers land on the same page
        # as the surrounding table segments, revert to the pre-rendered backup
        # (which already has the canonical TOC and zero same-page violations).
        _gate_restored = False
        if table_gate_backup_path is not None:
            try:
                if restore_docx_if_same_page_continuation_markers(
                    output_path,
                    table_gate_backup_path,
                    report=report,
                    context="format_docx_final",
                ):
                    _gate_restored = True
                    logger.warning(
                        "format_docx: final same-page continuation marker gate restored pre-rendered table state"
                    )
            except Exception:
                logger.exception("format_docx: final same-page continuation marker validation failed")
            finally:
                if table_gate_backup_dir is not None:
                    shutil.rmtree(table_gate_backup_dir, ignore_errors=True)

        # When the gate restored the canonical backup, DOCX-only markers that
        # were calibrated for the old student TOC layout may now be same-page in
        # the canonical layout (old TOC was typically larger by 1 page).  Remove
        # them: these markers are stale artefacts of the old layout — the table
        # actually fits on one page with the canonical TOC, so no continuation
        # header is needed.
        if _gate_restored:
            try:
                n_removed = remove_same_page_continuation_markers_inplace(
                    output_path,
                    report=report,
                )
                if n_removed:
                    logger.info(
                        "format_docx: removed %d same-page DOCX-only markers from canonical backup after gate",
                        n_removed,
                    )
            except Exception:
                logger.exception(
                    "format_docx: failed to remove same-page markers from canonical backup"
                )

    # Same-page merge-back (CLEANUP). Runs in conservative mode too: it merges
    # same-page repeated fragments (e.g. an unnecessary student manual chain
    # whose two halves now fit on one page) back into a single table and drops
    # the stray marker. Each merge has its own rendered rollback, so it never
    # worsens pagination.
    try:
        n_same_page_exact = normalize_exact_grid_same_page_repeated_fragments_inplace(
            output_path,
            source_docx_path=input_path,
            report=report,
        )
        if n_same_page_exact:
            logger.info(
                "format_docx: normalized %d exact-grid same-page table fragment(s)",
                n_same_page_exact,
            )
    except Exception:
        logger.exception("format_docx: exact-grid same-page fragment normalization failed")

    try:
        n_same_page_compatible = normalize_compatible_grid_same_page_repeated_fragments_inplace(
            output_path,
            source_docx_path=input_path,
            report=report,
        )
        if n_same_page_compatible:
            logger.info(
                "format_docx: normalized %d compatible-grid same-page table fragment(s)",
                n_same_page_compatible,
            )
    except Exception:
        logger.exception("format_docx: compatible-grid same-page fragment normalization failed")

    # Fallback cleanup for grid-incompatible same-page chains the mergers refuse:
    # drop the same-page marker + the second fragment's duplicate header/numeric,
    # keep both physical tables. Cleanup only; never reshapes grids or data rows.
    #
    # Gated to the experimental path. Verified (this round, Bondarev 3×) that it
    # does NOT fire for the remaining incompatible student chains (1.3.1/2.1.4/
    # 2.1.5) — their internal render does not emit a `same_page_repeated_fragment`
    # candidate — so enabling it by default only adds a render with no benefit.
    # A deterministic DOCX-level trigger is the Stage D follow-up.
    if _rendered_table_continuation_enabled():
        try:
            n_incompatible = cleanup_same_page_incompatible_chains_inplace(
                output_path,
                source_docx_path=input_path,
                report=report,
            )
            if n_incompatible:
                logger.info(
                    "format_docx: cleaned %d grid-incompatible same-page chain(s)",
                    n_incompatible,
                )
        except Exception:
            logger.exception("format_docx: incompatible-grid same-page cleanup failed")

    # Deterministic same-page manual continuation cleanup, driven by the
    # acceptance-gate `same_page_continuation` blocker (reliable table+page) — not
    # by the narrower `same_page_repeated_fragment` signal. Merges compatible-grid
    # chains, else drops the same-page marker + duplicate header keeping both
    # tables. Every applied cleanup is re-render verified (removes the fail, adds
    # no new fail, preserves all content) or rolled back.
    try:
        n_sp_cont = cleanup_same_page_continuation_blockers_inplace(
            output_path,
            source_docx_path=input_path,
            report=report,
        )
        if n_sp_cont:
            logger.info(
                "format_docx: cleaned %d same-page continuation chain(s)",
                n_sp_cont,
            )
    except Exception:
        logger.exception("format_docx: same-page continuation cleanup failed")

    # Canonical continuation rule on EXISTING manual chains: a `Продолжение
    # таблицы N` fragment must start with the numeric column row, not a repeated
    # semantic header. Strip the duplicate header from each continuation fragment
    # (deterministic, content-safe, rolled back on any new fail blocker).
    try:
        n_norm = normalize_continuation_semantic_header_inplace(
            output_path,
            source_docx_path=input_path,
            report=report,
        )
        if n_norm:
            logger.info(
                "format_docx: normalized %d continuation fragment header(s) to numeric row",
                n_norm,
            )
    except Exception:
        logger.exception("format_docx: continuation semantic-header normalization failed")

    try:
        n_final_orphan_moves = apply_rendered_table_start_orphan_guard(
            output_path,
            report=report,
        )
        if n_final_orphan_moves:
            logger.info(
                "format_docx: final table-start orphan guard moved %d table(s)",
                n_final_orphan_moves,
            )
    except Exception:
        logger.exception("format_docx: final table-start orphan guard failed")

    try:
        rendered_violations = _rendered_continuation_violations_for_docx(
            output_path,
            source_table_identities=source_table_identities,
        )
        if rendered_violations:
            n_repaired = repair_manual_chain_overflow_before_marker(
                output_path,
                rendered_violations,
                report=report,
            )
            if n_repaired:
                logger.info(
                    "format_docx: repaired %d manual-chain overflow continuation(s)",
                    n_repaired,
                )
                try:
                    remove_same_page_continuation_markers_inplace(output_path, report=None)
                except Exception:
                    logger.exception(
                        "format_docx: same-page marker cleanup after manual-chain repair failed"
                    )
                rendered_violations = _rendered_continuation_violations_for_docx(
                    output_path,
                    source_table_identities=source_table_identities,
                )
        if rendered_violations:
            _append_rendered_continuation_warnings(report, rendered_violations)
            hard_count = sum(
                1
                for violation in rendered_violations
                if violation.violation_type == "missing_continuation_marker"
                and violation.confidence == "high"
            )
            suspected_count = sum(
                1
                for violation in rendered_violations
                if violation.violation_type == "suspected_missing_continuation_marker"
                and violation.confidence == "medium"
            )
            logger.warning(
                "format_docx: final rendered continuation validation review_needed hard=%d suspected=%d",
                hard_count,
                suspected_count,
            )
    except Exception:
        logger.exception("format_docx: final rendered continuation validation failed")
        report.warn(
            "Автопроверка переносов таблиц по PDF не выполнена. Проверьте таблицы вручную."
        )

    # Deterministic marker-less long-table split (acceptance blocker
    # `single_table_crosses_pages_without_marker`). Runs LAST among the mutating
    # table stages so each split's re-render verify reflects the FINAL pagination:
    # a split that would push a neighbour table across a page boundary (cascade)
    # fails its own verify here and is rolled back. The table is split at the
    # rendered boundary (reliable per-row instrumentation when text matching is
    # ambiguous), a page-broken `Продолжение таблицы N` marker inserted, and ONLY
    # the numeric column row repeated on the continuation fragment (never the
    # semantic header). Content-preserving + budget-capped.
    try:
        n_cross = cleanup_cross_page_without_marker_blockers_inplace(
            output_path,
            source_docx_path=input_path,
            report=report,
        )
        if n_cross:
            logger.info(
                "format_docx: split %d marker-less cross-page table(s)",
                n_cross,
            )
    except Exception:
        logger.exception("format_docx: cross-page marker-less table split failed")

    # Normalize continuation-fragment column widths to the first fragment's grid
    # (acceptance review `fragment_grid_mismatch`). Runs after the cross-page
    # split so it sees the FINAL fragment set. Deterministic + content-safe
    # (only widths change); whole-doc rollback on any content regression or new
    # fail blocker (e.g. a squeeze).
    try:
        n_grid = normalize_fragment_grid_widths_inplace(
            output_path,
            source_docx_path=input_path,
            report=report,
        )
        if n_grid:
            logger.info(
                "format_docx: normalized %d continuation fragment grid(s) to first-fragment widths",
                n_grid,
            )
    except Exception:
        logger.exception("format_docx: fragment grid width normalization failed")

    try:
        n_same_page_marker_warnings = warn_same_page_continuation_marker_violations(
            output_path,
            report=report,
        )
        if n_same_page_marker_warnings:
            logger.warning(
                "format_docx: final same-page marker validation review_needed=%d",
                n_same_page_marker_warnings,
            )
    except Exception:
        logger.exception("format_docx: final same-page marker warning validation failed")

    # Stage A: rendered table layout acceptance gate. Surfaces visible table
    # layout defects (same-page continuation, orphaned header, grid mismatch,
    # appendix not on a new page, severe squeeze) as structured blockers in the
    # report + logs. fail-level blockers mean the rendered layout is NO-GO.
    _emit_table_layout_acceptance_warnings(
        output_path, report, source_identities=source_table_identities, source_text=source_text
    )

    return str(output_path), report.warnings
