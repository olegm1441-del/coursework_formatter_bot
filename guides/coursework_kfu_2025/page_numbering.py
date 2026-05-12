from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


TITLE_FOOTER_TEXT = "Казань – 2026 г."


def _is_contents_heading(text: str) -> bool:
    t = (text or "").strip().upper()
    return ("СОДЕРЖАН" in t) or ("ОГЛАВЛЕН" in t)


def _is_intro_heading(text: str) -> bool:
    return (text or "").strip().upper().rstrip(".") == "ВВЕДЕНИЕ"


def _is_appendices_heading(text: str) -> bool:
    return (text or "").strip().upper().rstrip(".") == "ПРИЛОЖЕНИЯ"


def _clear_xml_children(el):
    for child in list(el):
        el.remove(child)


def _reset_footer_part(footer):
    try:
        footer.is_linked_to_previous = False
    except Exception:
        pass

    ftr = footer._element
    _clear_xml_children(ftr)

    p = OxmlElement("w:p")
    ftr.append(p)


def _get_single_footer_paragraph(footer):
    _reset_footer_part(footer)
    return footer.paragraphs[0]


def _set_run_font(run):
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)

    rPr = run._element.get_or_add_rPr()

    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # 14 pt = 28 (half-points)
    sz = rPr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rPr.append(sz)
    sz.set(qn("w:val"), "28")

    szCs = rPr.find(qn("w:szCs"))
    if szCs is None:
        szCs = OxmlElement("w:szCs")
        rPr.append(szCs)
    szCs.set(qn("w:val"), "28")

def _clear_paragraph(paragraph):
    p = paragraph._element
    _clear_xml_children(p)


def _add_text_to_paragraph(paragraph, text):
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    _set_run_font(run)


def _add_page_field_to_paragraph(paragraph):
    _clear_paragraph(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    _set_run_font(run)

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._element.append(fld_char_begin)
    run._element.append(instr_text)
    run._element.append(fld_char_end)


def _remove_pg_num_type(section):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is not None:
        sectPr.remove(pgNumType)


def _set_page_number_start(section, start_value):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn("w:pgNumType"))
    if pgNumType is None:
        pgNumType = OxmlElement("w:pgNumType")
        sectPr.append(pgNumType)
    pgNumType.set(qn("w:start"), str(start_value))


def _reset_section_footer_state(section):
    section.different_first_page_header_footer = True

    _remove_pg_num_type(section)

    _reset_footer_part(section.first_page_footer)
    _reset_footer_part(section.footer)

    try:
        _reset_footer_part(section.even_page_footer)
    except Exception:
        pass


def _reset_all_footer_state(document):
    try:
        document.settings.odd_and_even_pages_header_footer = False
    except Exception:
        pass

    for section in document.sections:
        _reset_section_footer_state(section)


def _blank_section(section):
    _reset_section_footer_state(section)

    p1 = _get_single_footer_paragraph(section.first_page_footer)
    _clear_paragraph(p1)

    p2 = _get_single_footer_paragraph(section.footer)
    _clear_paragraph(p2)

    try:
        p3 = _get_single_footer_paragraph(section.even_page_footer)
        _clear_paragraph(p3)
    except Exception:
        pass

    _remove_pg_num_type(section)


def _title_section(section):
    _reset_section_footer_state(section)

    p1 = _get_single_footer_paragraph(section.first_page_footer)
    _add_text_to_paragraph(p1, TITLE_FOOTER_TEXT)

    p2 = _get_single_footer_paragraph(section.footer)
    _clear_paragraph(p2)

    try:
        p3 = _get_single_footer_paragraph(section.even_page_footer)
        _clear_paragraph(p3)
    except Exception:
        pass

    _remove_pg_num_type(section)


def _number_section(section, start_value=None):
    _reset_section_footer_state(section)

    # Для нумеруемых секций не нужен особый футер первой страницы.
    # Иначе Word может показывать на первой странице секции номер
    # не с того значения, а старт (например, 3) визуально "съезжает"
    # на следующую страницу.
    section.different_first_page_header_footer = False

    if start_value is not None:
        _set_page_number_start(section, start_value)

    dp = _get_single_footer_paragraph(section.footer)
    _add_page_field_to_paragraph(dp)

    try:
        ep = _get_single_footer_paragraph(section.even_page_footer)
        _add_page_field_to_paragraph(ep)
    except Exception:
        pass


def _appendix_section(section):
    _reset_section_footer_state(section)

    p1 = _get_single_footer_paragraph(section.first_page_footer)
    _add_page_field_to_paragraph(p1)

    p2 = _get_single_footer_paragraph(section.footer)
    _clear_paragraph(p2)

    try:
        p3 = _get_single_footer_paragraph(section.even_page_footer)
        _clear_paragraph(p3)
    except Exception:
        pass

    _remove_pg_num_type(section)


def _paragraph_section_index(document, paragraph_index):
    section_index = 0
    for idx, paragraph in enumerate(document.paragraphs):
        if idx >= paragraph_index:
            break
        p_pr = paragraph._element.pPr
        if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
            section_index += 1
    return section_index


def _find_paragraph_index(document, predicate):
    for idx, paragraph in enumerate(document.paragraphs):
        if predicate(paragraph.text):
            return idx
    return None


def apply_page_numbering_policy(document):
    sections = list(document.sections)
    if not sections:
        return

    _reset_all_footer_state(document)

    intro_idx = _find_paragraph_index(document, _is_intro_heading)
    appendices_idx = _find_paragraph_index(document, _is_appendices_heading)

    if intro_idx is None:
        _title_section(sections[0])
        for section in sections[1:]:
            _blank_section(section)
        return

    intro_section_idx = min(_paragraph_section_index(document, intro_idx), len(sections) - 1)
    appendices_section_idx = None
    if appendices_idx is not None:
        appendices_section_idx = min(_paragraph_section_index(document, appendices_idx), len(sections) - 1)

    for idx, section in enumerate(sections):
        if idx < intro_section_idx:
            if idx == 0:
                _title_section(section)
            else:
                _blank_section(section)
            continue

        if appendices_section_idx is not None and idx == appendices_section_idx:
            _appendix_section(section)
            continue

        if appendices_section_idx is not None and idx > appendices_section_idx:
            _blank_section(section)
            continue

        if idx == intro_section_idx:
            _number_section(section, start_value=3)
        else:
            _number_section(section)
