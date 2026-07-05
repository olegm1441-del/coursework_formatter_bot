from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx import Document
from docx.oxml.ns import qn

from .pdf_layout_analyzer import PdfLine


_STRICT_MARKER_RE = re.compile(
    r"^\s*Продолжение\s+таблицы\s+([0-9]+(?:\.[0-9]+)*)\s*$",
    re.IGNORECASE,
)
_ANY_MARKER_RE = re.compile(
    r"Продолжение\s+таблицы\s+([0-9]+(?:\.[0-9]+)*)",
    re.IGNORECASE,
)
_CAPTION_RE = re.compile(
    r"^\s*Таблица\s+([0-9]+(?:\.[0-9]+)*)\b",
    re.IGNORECASE,
)
_NUMERIC_ROW_RE = re.compile(r"^(?:\d+\s+){1,}\d+$")


@dataclass(frozen=True)
class RenderedTableIdentity:
    table_index: int
    body_order_index: int
    caption_num: str | None
    preceding_marker: str | None
    following_marker: str | None
    header_fingerprint: tuple[str, ...]
    numeric_row_fingerprint: str | None
    row_fingerprints: tuple[str, ...]
    preceding_inter_table_texts: tuple[str, ...] = ()


@dataclass(frozen=True)
class RenderedContinuationViolation:
    table_num: str | None
    table_index: int
    page: int
    violation_type: str
    confidence: str
    evidence: dict[str, object] = field(default_factory=dict)


@dataclass(frozen=True)
class ContinuationMarkerLine:
    text: str
    table_num: str | None
    marker_kind: str


def _norm_text(text: str) -> str:
    return " ".join((text or "").replace("\xa0", " ").split()).lower()


def _line_text(line: PdfLine) -> str:
    return " ".join((line.text or "").split())


def classify_continuation_marker_line(text: str) -> ContinuationMarkerLine:
    cleaned = " ".join((text or "").split())
    strict = _STRICT_MARKER_RE.match(cleaned)
    if strict:
        return ContinuationMarkerLine(
            text=cleaned,
            table_num=strict.group(1),
            marker_kind="strict_marker",
        )
    inline = _ANY_MARKER_RE.search(cleaned)
    if inline:
        return ContinuationMarkerLine(
            text=cleaned,
            table_num=inline.group(1),
            marker_kind="source_inline_marker_text",
        )
    return ContinuationMarkerLine(text=cleaned, table_num=None, marker_kind="none")


def _row_cell_texts(row_xml) -> list[str]:
    out: list[str] = []
    for cell in row_xml.findall(qn("w:tc")):
        text = " ".join(
            (node.text or "")
            for node in cell.findall(".//" + qn("w:t"))
        )
        out.append(" ".join(text.split()))
    return out


def _row_fingerprint(row_xml) -> str | None:
    fragments = [_norm_text(text) for text in _row_cell_texts(row_xml) if _norm_text(text)]
    if not fragments:
        return None
    return " ".join(fragments)


def _is_numeric_cells(cells: list[str]) -> bool:
    values = [" ".join((cell or "").split()) for cell in cells]
    if len(values) < 2:
        return False
    return values == [str(idx) for idx in range(1, len(values) + 1)]


def _caption_num(text: str) -> str | None:
    match = _CAPTION_RE.match(" ".join((text or "").split()))
    return match.group(1) if match else None


def _strict_marker_text(text: str) -> str | None:
    marker = classify_continuation_marker_line(text)
    if marker.marker_kind == "strict_marker":
        return marker.text
    return None


def _table_body_order(doc: Document) -> dict[object, int]:
    order: dict[object, int] = {}
    idx = 0
    for child in doc.element.body:
        if child.tag == qn("w:tbl"):
            order[child] = idx
            idx += 1
    return order


def build_rendered_table_identities(doc: Document) -> list[RenderedTableIdentity]:
    body_children = list(doc.element.body)
    para_by_xml = {p._element: p for p in doc.paragraphs}
    body_order = _table_body_order(doc)
    out: list[RenderedTableIdentity] = []

    for table_index, table in enumerate(doc.tables):
        tbl_xml = table._tbl
        try:
            body_idx = body_children.index(tbl_xml)
        except ValueError:
            body_idx = -1

        caption_num: str | None = None
        preceding_marker: str | None = None
        for node in reversed(body_children[:body_idx] if body_idx >= 0 else []):
            if node.tag == qn("w:tbl"):
                break
            if node.tag != qn("w:p"):
                continue
            para = para_by_xml.get(node)
            text = " ".join(((para.text if para is not None else "") or "").split())
            if not text:
                continue
            marker = _strict_marker_text(text)
            if marker and preceding_marker is None:
                preceding_marker = marker
                continue
            caption_num = _caption_num(text)
            if caption_num:
                break

        following_marker: str | None = None
        if body_idx >= 0:
            for node in body_children[body_idx + 1:]:
                if node.tag == qn("w:tbl"):
                    break
                if node.tag != qn("w:p"):
                    continue
                para = para_by_xml.get(node)
                text = " ".join(((para.text if para is not None else "") or "").split())
                marker = _strict_marker_text(text)
                if marker:
                    following_marker = marker
                    break
                if text:
                    break

        preceding_inter_table_texts: list[str] = []
        if body_idx >= 0:
            for node in reversed(body_children[:body_idx]):
                if node.tag == qn("w:tbl"):
                    break
                if node.tag != qn("w:p"):
                    continue
                para = para_by_xml.get(node)
                text = " ".join(((para.text if para is not None else "") or "").split())
                if text:
                    preceding_inter_table_texts.append(text)
            preceding_inter_table_texts.reverse()

        header_rows: list[str] = []
        numeric_row: str | None = None
        all_rows: list[str] = []
        for row_xml in tbl_xml.findall(qn("w:tr")):
            row_fp = _row_fingerprint(row_xml)
            if row_fp:
                all_rows.append(row_fp)
            cells = _row_cell_texts(row_xml)
            if _is_numeric_cells(cells):
                numeric_row = " ".join(cells)
                continue
            if row_fp and numeric_row is None and len(header_rows) < 3:
                header_rows.append(row_fp)

        out.append(
            RenderedTableIdentity(
                table_index=table_index,
                body_order_index=body_order.get(tbl_xml, table_index),
                caption_num=caption_num,
                preceding_marker=preceding_marker,
                following_marker=following_marker,
                header_fingerprint=tuple(header_rows),
                numeric_row_fingerprint=numeric_row,
                row_fingerprints=tuple(all_rows),
                preceding_inter_table_texts=tuple(preceding_inter_table_texts),
            )
        )
    return out


def _caption_lines(pdf_lines: list[PdfLine], table_num: str) -> list[PdfLine]:
    needle = re.compile(rf"^\s*Таблица\s+{re.escape(table_num)}\b", re.IGNORECASE)
    return [line for line in pdf_lines if needle.match(_line_text(line))]


def _strict_marker_lines(pdf_lines: list[PdfLine], table_num: str) -> list[PdfLine]:
    out: list[PdfLine] = []
    for line in pdf_lines:
        marker = classify_continuation_marker_line(_line_text(line))
        if marker.marker_kind == "strict_marker" and marker.table_num == table_num:
            out.append(line)
    return out


def _tokens(text: str) -> set[str]:
    return {token for token in re.findall(r"[0-9a-zа-яё]+", _norm_text(text)) if len(token) > 1}


def _contains_fingerprint(page_text: str, fingerprint: str, *, min_overlap: float = 0.55) -> bool:
    fp_tokens = _tokens(fingerprint)
    if not fp_tokens:
        return False
    page_tokens = _tokens(page_text)
    return len(fp_tokens & page_tokens) / len(fp_tokens) >= min_overlap


def _snippet(text: str, limit: int = 90) -> str:
    cleaned = " ".join(text.split())
    return cleaned[:limit]


def _page_lines(pdf_lines: list[PdfLine], page_num: int) -> list[PdfLine]:
    return sorted(
        [line for line in pdf_lines if line.page_num == page_num],
        key=lambda line: (line.top, line.text),
    )


def _page_text(pdf_lines: list[PdfLine], page_num: int) -> str:
    return " ".join(_norm_text(line.text) for line in _page_lines(pdf_lines, page_num))


def _fingerprint_overlap(left: str, right: str) -> float:
    left_tokens = _tokens(left)
    right_tokens = _tokens(right)
    if not left_tokens or not right_tokens:
        return 0.0
    return len(left_tokens & right_tokens) / max(len(left_tokens), len(right_tokens))


def _line_matches_numeric(line: PdfLine, numeric_row_fingerprint: str) -> bool:
    normalized = _norm_text(line.text)
    return normalized == numeric_row_fingerprint or (
        bool(_NUMERIC_ROW_RE.match(normalized)) and normalized == numeric_row_fingerprint
    )


def _numeric_row_count(lines: list[PdfLine], numeric_row_fingerprint: str | None) -> int:
    if not numeric_row_fingerprint:
        return 0
    return sum(1 for line in lines if _line_matches_numeric(line, numeric_row_fingerprint))


def _meaningful_rows(identity: RenderedTableIdentity) -> list[str]:
    rows: list[str] = []
    structural_headers = set(identity.header_fingerprint[:1])
    for row_fp in identity.row_fingerprints:
        if row_fp == identity.numeric_row_fingerprint:
            continue
        if row_fp in structural_headers:
            continue
        if len(_tokens(row_fp)) < 3:
            continue
        rows.append(row_fp)
    return rows


def _first_meaningful_row_on_page(
    identity: RenderedTableIdentity,
    page_text: str,
    *,
    min_overlap: float = 0.7,
) -> str | None:
    for row_fp in _meaningful_rows(identity):
        if _contains_fingerprint(page_text, row_fp, min_overlap=min_overlap):
            return row_fp
    return None


def _compatible_adjacent_fragment(
    current: RenderedTableIdentity,
    following: RenderedTableIdentity,
) -> bool:
    if following.caption_num is not None:
        return False
    if following.preceding_marker or following.following_marker:
        return False
    if following.preceding_inter_table_texts:
        return False
    if not current.header_fingerprint or not following.header_fingerprint:
        return False
    if current.numeric_row_fingerprint and following.numeric_row_fingerprint:
        if current.numeric_row_fingerprint != following.numeric_row_fingerprint:
            return False
    return _fingerprint_overlap(current.header_fingerprint[0], following.header_fingerprint[0]) >= 0.65


def _adjacent_identity_pairs(
    identities: list[RenderedTableIdentity],
) -> list[tuple[RenderedTableIdentity, RenderedTableIdentity]]:
    ordered = sorted(identities, key=lambda identity: identity.body_order_index)
    pairs: list[tuple[RenderedTableIdentity, RenderedTableIdentity]] = []
    for current, following in zip(ordered, ordered[1:]):
        if following.body_order_index != current.body_order_index + 1:
            continue
        if not current.caption_num:
            continue
        if _compatible_adjacent_fragment(current, following):
            pairs.append((current, following))
    return pairs


def _source_bad_duplicate_rows(
    identity: RenderedTableIdentity,
    source_by_num: dict[str, RenderedTableIdentity],
) -> list[str]:
    if not identity.caption_num:
        return []
    source = source_by_num.get(identity.caption_num)
    if source is None:
        return []

    source_counts: dict[str, int] = {}
    for row_fp in _meaningful_rows(source):
        source_counts[row_fp] = source_counts.get(row_fp, 0) + 1

    output_counts: dict[str, int] = {}
    for row_fp in _meaningful_rows(identity):
        output_counts[row_fp] = output_counts.get(row_fp, 0) + 1

    duplicates = [
        row_fp
        for row_fp, count in output_counts.items()
        if count > 1 and source_counts.get(row_fp, 0) > 1
    ]
    return duplicates


def _valid_marker_before_fragment(marker_lines: list[PdfLine], page_num: int, evidence_top: float) -> bool:
    return any(line.page_num == page_num and line.top < evidence_top - 1.0 for line in marker_lines)


def _nearest_marker_page(marker_lines: list[PdfLine], page_num: int, evidence_top: float) -> int | None:
    later = [
        line for line in marker_lines
        if line.page_num > page_num or (line.page_num == page_num and line.top > evidence_top + 1.0)
    ]
    if later:
        return sorted(later, key=lambda line: (line.page_num, line.top))[0].page_num
    earlier = [line for line in marker_lines if line.page_num < page_num]
    if earlier:
        return sorted(earlier, key=lambda line: (line.page_num, line.top))[-1].page_num
    return None


def validate_rendered_continuations(
    pdf_lines: list[PdfLine],
    table_identities: list[RenderedTableIdentity],
    source_table_identities: list[RenderedTableIdentity] | None = None,
) -> list[RenderedContinuationViolation]:
    violations: list[RenderedContinuationViolation] = []
    pages = sorted({line.page_num for line in pdf_lines})
    all_caption_pages = sorted(
        line.page_num
        for line in pdf_lines
        if _CAPTION_RE.match(_line_text(line))
    )

    for identity in table_identities:
        table_num = identity.caption_num
        if not table_num:
            continue

        captions = _caption_lines(pdf_lines, table_num)
        if not captions:
            continue
        first_caption_page = min(line.page_num for line in captions)
        next_caption_page = next(
            (page for page in all_caption_pages if page > first_caption_page),
            None,
        )
        markers = _strict_marker_lines(pdf_lines, table_num)

        for page_num in pages:
            if page_num <= first_caption_page:
                continue
            if next_caption_page is not None and page_num >= next_caption_page:
                continue

            lines = _page_lines(pdf_lines, page_num)
            if not lines:
                continue

            marker_on_page = sorted(
                [line for line in markers if line.page_num == page_num],
                key=lambda line: line.top,
            )
            marker_top = marker_on_page[0].top if marker_on_page else None
            window_limit = min(180.0, marker_top - 1.0) if marker_top is not None else 180.0
            window_lines = [line for line in lines if line.top <= window_limit]
            if not window_lines:
                continue

            window_text = " ".join(_norm_text(line.text) for line in window_lines)
            repeated_header = False
            header_snippet = None
            for header in identity.header_fingerprint:
                if _contains_fingerprint(window_text, header):
                    repeated_header = True
                    header_snippet = _snippet(header)
                    break

            repeated_numeric = False
            numeric_snippet = None
            if identity.numeric_row_fingerprint:
                for line in window_lines:
                    normalized = _norm_text(line.text)
                    if normalized == identity.numeric_row_fingerprint or (
                        _NUMERIC_ROW_RE.match(normalized)
                        and normalized == identity.numeric_row_fingerprint
                    ):
                        repeated_numeric = True
                        numeric_snippet = identity.numeric_row_fingerprint
                        break

            repeated_row = False
            row_snippet = None
            for row_fp in identity.row_fingerprints:
                if row_fp == identity.numeric_row_fingerprint:
                    continue
                if row_fp in identity.header_fingerprint:
                    continue
                if _contains_fingerprint(window_text, row_fp, min_overlap=0.7):
                    repeated_row = True
                    row_snippet = _snippet(row_fp)
                    break

            if not (repeated_header or repeated_numeric or repeated_row):
                continue

            evidence_top = min(line.top for line in window_lines)
            if _valid_marker_before_fragment(markers, page_num, evidence_top):
                continue
            marker_page = _nearest_marker_page(markers, page_num, evidence_top)

            if repeated_header and repeated_numeric and marker_on_page:
                violation_type = "late_continuation_marker"
                confidence = "high"
            elif repeated_header and repeated_numeric:
                violation_type = "missing_continuation_marker"
                confidence = "high"
            elif repeated_numeric and markers and marker_on_page:
                violation_type = "late_continuation_marker"
                confidence = "high"
            elif repeated_numeric and markers:
                violation_type = "missing_continuation_marker"
                confidence = "high"
            elif repeated_row and marker_page is not None:
                violation_type = "suspected_missing_continuation_marker"
                confidence = "medium"
            else:
                # Header-only or row-only matches without marker context are too
                # noisy in rendered PDFs; keep them out of hard diagnostics.
                continue

            violations.append(
                RenderedContinuationViolation(
                    table_num=table_num,
                    table_index=identity.table_index,
                    page=page_num,
                    violation_type=violation_type,
                    confidence=confidence,
                    evidence={
                        "repeated_header": repeated_header,
                        "repeated_numeric_row": repeated_numeric,
                        "repeated_row": repeated_row,
                        "previous_caption_page": first_caption_page,
                        "marker_page": marker_page,
                        "fragment_page": page_num,
                        "header_fingerprint": header_snippet,
                        "numeric_row_fingerprint": numeric_snippet,
                        "row_fingerprint": row_snippet,
                    },
                )
            )

    caption_page_by_num: dict[str, int] = {}
    for identity in table_identities:
        if not identity.caption_num:
            continue
        captions = _caption_lines(pdf_lines, identity.caption_num)
        if captions:
            caption_page_by_num[identity.caption_num] = min(line.page_num for line in captions)

    for current, following in _adjacent_identity_pairs(table_identities):
        table_num = current.caption_num
        if not table_num:
            continue
        first_page = caption_page_by_num.get(table_num)
        if first_page is None:
            continue
        same_page_lines = _page_lines(pdf_lines, first_page)
        same_page_text = " ".join(_norm_text(line.text) for line in same_page_lines)
        next_row = _first_meaningful_row_on_page(following, same_page_text)
        repeated_numeric_count = _numeric_row_count(
            same_page_lines,
            current.numeric_row_fingerprint,
        )
        repeated_header = any(
            _contains_fingerprint(same_page_text, header)
            for header in following.header_fingerprint
        )
        same_page_proven = bool(next_row)
        if same_page_proven:
            repeated_artifact = repeated_numeric_count >= 2 or repeated_header
            violations.append(
                RenderedContinuationViolation(
                    table_num=table_num,
                    table_index=current.table_index,
                    page=first_page,
                    violation_type=(
                        "same_page_repeated_fragment"
                        if repeated_artifact
                        else "same_page_adjacent_fragment"
                    ),
                    confidence="high",
                    evidence={
                        "following_table_index": following.table_index,
                        "fragment_page": first_page,
                        "repeated_header": repeated_header,
                        "repeated_numeric_row_count": repeated_numeric_count,
                        "row_fingerprint": _snippet(next_row),
                        "adjacent_fragment_proof": "following_row_on_caption_page",
                    },
                )
            )
            continue

        following_pages = [
            page
            for page in pages
            if _first_meaningful_row_on_page(following, _page_text(pdf_lines, page)) is not None
        ]
        if following_pages:
            first_following_page = min(following_pages)
            if (
                first_following_page > first_page
                and current.numeric_row_fingerprint
                and following.numeric_row_fingerprint
            ):
                markers = _strict_marker_lines(pdf_lines, table_num)
                marker_before = any(
                    line.page_num < first_following_page
                    or (
                        line.page_num == first_following_page
                        and line.top < min(l.top for l in _page_lines(pdf_lines, first_following_page))
                    )
                    for line in markers
                )
                violations.append(
                    RenderedContinuationViolation(
                        table_num=table_num,
                        table_index=current.table_index,
                        page=first_following_page,
                        violation_type=(
                            "ambiguous_adjacent_tables"
                            if marker_before
                            else "missing_or_late_continuation_marker"
                        ),
                        confidence="medium",
                        evidence={
                            "following_table_index": following.table_index,
                            "previous_caption_page": first_page,
                            "fragment_page": first_following_page,
                            "proof": "adjacent_fragment_row_on_later_page",
                        },
                    )
                )
                continue

        following_meaningful_rows = _meaningful_rows(following)
        if len(following_meaningful_rows) >= 2:
            violations.append(
                RenderedContinuationViolation(
                    table_num=table_num,
                    table_index=current.table_index,
                    page=first_page,
                    violation_type="ambiguous_adjacent_tables",
                    confidence="medium",
                    evidence={
                        "following_table_index": following.table_index,
                        "previous_caption_page": first_page,
                        "following_meaningful_rows": len(following_meaningful_rows),
                        "proof": "compatible_adjacent_docx_tables_render_not_proven",
                    },
                )
            )

    if source_table_identities:
        source_by_num = {
            identity.caption_num: identity
            for identity in source_table_identities
            if identity.caption_num
        }
        for identity in table_identities:
            duplicates = _source_bad_duplicate_rows(identity, source_by_num)
            if not duplicates:
                continue
            violations.append(
                RenderedContinuationViolation(
                    table_num=identity.caption_num,
                    table_index=identity.table_index,
                    page=caption_page_by_num.get(identity.caption_num or "", 0),
                    violation_type="source_bad_duplicated_content_rows",
                    confidence="high",
                    evidence={
                        "duplicate_row_count": len(duplicates),
                        "row_fingerprint": _snippet(duplicates[0]),
                        "source_proven": True,
                    },
                )
            )

    return violations


# --------------------------------------------------------------------------- #
# Stage A — rendered table layout acceptance gate
#
# These detectors turn *visible* rendered table defects into structured,
# severity-bearing blockers so the smoke / deploy decision can no longer report
# GO when the PDF layout is broken. The function is pure (no rendering, no I/O):
# callers render once, then pass the extracted lines + DOCX identities here.
#
#   severity == "fail"               -> NO-GO; a human would reject this layout.
#   severity == "needs_human_review" -> not clean; uncertain, must be reviewed.
# --------------------------------------------------------------------------- #

_APPENDIX_LABEL_RE = re.compile(r"^\s*приложение\s+[а-яёa-z]\b", re.IGNORECASE)


@dataclass(frozen=True)
class TableLayoutBlocker:
    blocker_type: str
    severity: str  # "fail" | "needs_human_review"
    table_num: str | None
    page: int | None
    evidence: dict[str, object] = field(default_factory=dict)


def _all_caption_pages(pdf_lines: list[PdfLine]) -> list[int]:
    return sorted({line.page_num for line in pdf_lines if _CAPTION_RE.match(_line_text(line))})


def _table_page_span(
    pdf_lines: list[PdfLine],
    identity: RenderedTableIdentity,
) -> list[int]:
    """Pages from this table's caption up to (but excluding) the next caption."""
    caption_pages = [line.page_num for line in _caption_lines(pdf_lines, identity.caption_num or "")]
    if not caption_pages:
        return []
    first = min(caption_pages)
    next_caption = next((p for p in _all_caption_pages(pdf_lines) if p > first), None)
    pages = sorted({line.page_num for line in pdf_lines})
    return [p for p in pages if p >= first and (next_caption is None or p < next_caption)]


def _same_page_continuation_blockers(
    pdf_lines: list[PdfLine],
    table_identities: list[RenderedTableIdentity],
) -> list[TableLayoutBlocker]:
    out: list[TableLayoutBlocker] = []
    for identity in table_identities:
        num = identity.caption_num
        if not num:
            continue
        caption_pages = {line.page_num for line in _caption_lines(pdf_lines, num)}
        flagged: set[int] = set()
        for marker in _strict_marker_lines(pdf_lines, num):
            page = marker.page_num
            if page in flagged:
                continue
            same_page = page in caption_pages
            if not same_page:
                above_text = " ".join(
                    _norm_text(line.text)
                    for line in pdf_lines
                    if line.page_num == page and line.top < marker.top - 1.0
                )
                if _first_meaningful_row_on_page(identity, above_text, min_overlap=0.7):
                    same_page = True
            if same_page:
                flagged.add(page)
                out.append(
                    TableLayoutBlocker(
                        blocker_type="same_page_continuation",
                        severity="fail",
                        table_num=num,
                        page=page,
                        evidence={"marker_text": _line_text(marker)},
                    )
                )
    return out


def _same_page_numeric_continuation_blockers(
    pdf_lines: list[PdfLine],
    table_identities: list[RenderedTableIdentity],
) -> list[TableLayoutBlocker]:
    """Marker-less same-page numeric continuation defect.

    A table pre-split in the source into numeric-led fragments whose fragments
    happen to fit on ONE page renders with its numeric label row ("1 2 … N")
    repeated mid-table, with no ``Продолжение таблицы`` marker between the
    fragments. ``_same_page_continuation_blockers`` misses it because it keys off
    strict markers, so this visible defect was previously undetected. Detect it
    structurally: the table's numeric_row_fingerprint appearing 2+ times on a
    single page WITHIN the table's own caption-to-next-caption span (the span
    excludes other tables' captions, so both occurrences belong to this table).
    """
    out: list[TableLayoutBlocker] = []
    for identity in table_identities:
        num = identity.caption_num
        if not num or not identity.numeric_row_fingerprint:
            continue
        for page in _table_page_span(pdf_lines, identity):
            count = _numeric_row_count(
                _page_lines(pdf_lines, page), identity.numeric_row_fingerprint
            )
            if count >= 2:
                out.append(
                    TableLayoutBlocker(
                        blocker_type="same_page_numeric_continuation",
                        severity="fail",
                        table_num=num,
                        page=page,
                        evidence={
                            "reason": "marker_less_numeric_repeat_same_page",
                            "numeric_row_count": count,
                        },
                    )
                )
                break
    return out


def _orphaned_header_blockers(
    pdf_lines: list[PdfLine],
    table_identities: list[RenderedTableIdentity],
) -> list[TableLayoutBlocker]:
    out: list[TableLayoutBlocker] = []
    for identity in table_identities:
        num = identity.caption_num
        if not num or not identity.header_fingerprint:
            continue
        span = _table_page_span(pdf_lines, identity)
        if not span or not _meaningful_rows(identity):
            continue
        matched = sorted(
            p for p in span
            if _first_meaningful_row_on_page(identity, _page_text(pdf_lines, p), min_overlap=0.7)
        )
        # Only the CONTIGUOUS run of data pages from the first belongs to this
        # physical table; a distant page that merely reuses the row text (a summary
        # / appendix table repeating the same figures) is a DIFFERENT table and
        # must not create a phantom "data page after" that fakes an orphan. (The
        # span runs to end-of-doc for the last captioned table.)
        data_pages: set[int] = set()
        prev: int | None = None
        for p in matched:
            if prev is None or p == prev + 1:
                data_pages.add(p)
                prev = p
            else:
                break
        for page in span:
            page_text = _page_text(pdf_lines, page)
            header_present = any(
                _contains_fingerprint(page_text, header, min_overlap=0.7)
                for header in identity.header_fingerprint
            )
            numeric_present = bool(
                identity.numeric_row_fingerprint
                and _numeric_row_count(_page_lines(pdf_lines, page), identity.numeric_row_fingerprint)
            )
            if (header_present or numeric_present) and page not in data_pages:
                if any(dp > page for dp in data_pages):
                    out.append(
                        TableLayoutBlocker(
                            blocker_type="orphaned_header_row",
                            severity="fail",
                            table_num=num,
                            page=page,
                            evidence={"header": _snippet(identity.header_fingerprint[0])},
                        )
                    )
    return out


def _appendix_label_blockers(pdf_lines: list[PdfLine]) -> list[TableLayoutBlocker]:
    out: list[TableLayoutBlocker] = []
    for line in pdf_lines:
        if not _APPENDIX_LABEL_RE.match(_line_text(line)):
            continue
        page = line.page_num
        substantial_above = [
            other
            for other in pdf_lines
            if other.page_num == page
            and other.top < line.top - 1.0
            and _tokens(other.text)
            and not _norm_text(other.text).startswith("приложени")
        ]
        if substantial_above:
            out.append(
                TableLayoutBlocker(
                    blocker_type="appendix_label_not_on_new_page",
                    severity="fail",
                    table_num=None,
                    page=page,
                    evidence={
                        "label": _line_text(line),
                        "content_above": _snippet(substantial_above[-1].text),
                    },
                )
            )
    return out


def _grid_col_widths(tbl_xml) -> list[int] | None:
    grid = tbl_xml.find(qn("w:tblGrid"))
    if grid is None:
        return None
    return [int(col.get(qn("w:w"), "0") or 0) for col in grid.findall(qn("w:gridCol"))]


def _is_blank_or_marker_paragraph(node) -> bool:
    if node.tag != qn("w:p"):
        return False
    text = " ".join((t.text or "") for t in node.findall(".//" + qn("w:t"))).strip()
    if not text:
        return True
    return classify_continuation_marker_line(text).table_num is not None


def _attribute_table(
    table_index: int,
    identities_by_index: dict[int, RenderedTableIdentity],
    pdf_lines: list[PdfLine],
) -> tuple[str | None, int | None]:
    """Best-effort (caption_num, first_page) for the table at doc.tables[index].
    Walks back to the nearest captioned fragment for unlabelled continuations."""
    num = None
    for idx in range(table_index, -1, -1):
        identity = identities_by_index.get(idx)
        if identity is not None and identity.caption_num:
            num = identity.caption_num
            break
    if not num:
        return None, None
    caption_pages = [line.page_num for line in _caption_lines(pdf_lines, num)]
    return num, (min(caption_pages) if caption_pages else None)


def _fragment_grid_mismatch_blockers(
    doc: Document,
    table_identities: list[RenderedTableIdentity],
    pdf_lines: list[PdfLine],
) -> list[TableLayoutBlocker]:
    out: list[TableLayoutBlocker] = []
    identities_by_index = {it.table_index: it for it in table_identities}
    body = list(doc.element.body)
    tbl_positions = [i for i, node in enumerate(body) if node.tag == qn("w:tbl")]
    for table_index, (a, b) in enumerate(zip(tbl_positions, tbl_positions[1:])):
        between = body[a + 1:b]
        if any(not _is_blank_or_marker_paragraph(node) for node in between):
            continue
        sig_a = _grid_col_widths(body[a])
        sig_b = _grid_col_widths(body[b])
        if not sig_a or not sig_b:
            continue
        # the first table of the pair is the table_index-th w:tbl in the body
        num, page = _attribute_table(table_index, identities_by_index, pdf_lines)
        if len(sig_a) != len(sig_b):
            out.append(
                TableLayoutBlocker(
                    blocker_type="fragment_grid_mismatch",
                    severity="fail",
                    table_num=num,
                    page=page,
                    evidence={"cols_first": len(sig_a), "cols_second": len(sig_b)},
                )
            )
            continue
        total_a = sum(sig_a) or 1
        total_b = sum(sig_b) or 1
        max_dev = max(
            abs(wa / total_a - wb / total_b) for wa, wb in zip(sig_a, sig_b)
        )
        # A correct split preserves column widths byte-equivalently (invariants
        # Rule 8). Rybakov's valid splits drift by 0.0; the small tolerance only
        # absorbs twip rounding noise.
        if max_dev > 0.02:
            out.append(
                TableLayoutBlocker(
                    blocker_type="fragment_grid_mismatch",
                    severity="needs_human_review",
                    table_num=num,
                    page=page,
                    evidence={"max_col_width_deviation": round(max_dev, 3)},
                )
            )
    return out


def _same_page_repeated_header_blockers(
    pdf_lines: list[PdfLine],
    table_identities: list[RenderedTableIdentity],
) -> list[TableLayoutBlocker]:
    """A table's semantic header rendered two+ times on the SAME page (a repeated
    header inside a same-page fragment), even without a continuation marker. A
    header repeated on a *different* (continuation) page is allowed."""
    out: list[TableLayoutBlocker] = []
    for identity in table_identities:
        num = identity.caption_num
        if not num or not identity.header_fingerprint:
            continue
        header = identity.header_fingerprint[0]
        if len(_tokens(header)) < 3:
            continue
        flagged: set[int] = set()
        for page in _table_page_span(pdf_lines, identity):
            page_lines = _page_lines(pdf_lines, page)
            # lines on this page that essentially reproduce the whole header row
            hits = [
                line for line in page_lines
                if _contains_fingerprint(_line_text(line), header, min_overlap=0.85)
            ]
            # A DIFFERENT table with a SIMILAR header (e.g. otchet's financial
            # tables all start "Показатель … 2022 г. 2023 г. 2024 г. …") renders its
            # own header under its own "Таблица M" caption on the same page. Only
            # count header matches within THIS table's block — above the next
            # other-table caption — so a neighbour's header is not counted as a
            # repeat of this one.
            page_caps = [
                (m.group(1), line.top)
                for line in page_lines
                for m in (_CAPTION_RE.match(_line_text(line)),)
                if m
            ]
            this_caps = [top for cnum, top in page_caps if cnum == num]
            if this_caps:
                this_top = min(this_caps)
                other_tops = [top for cnum, top in page_caps if cnum != num and top > this_top]
                if other_tops:
                    boundary = min(other_tops)
                    hits = [line for line in hits if line.top < boundary]
            # collapse wrapped header lines: count occurrences separated by a gap
            occurrences = 0
            last_top = None
            for line in sorted(hits, key=lambda l: l.top):
                if last_top is None or line.top - last_top > 30.0:
                    occurrences += 1
                last_top = line.top
            if occurrences >= 2 and page not in flagged:
                flagged.add(page)
                out.append(
                    TableLayoutBlocker(
                        blocker_type="same_page_repeated_header",
                        severity="fail",
                        table_num=num,
                        page=page,
                        evidence={"header": _snippet(header), "occurrences": occurrences},
                    )
                )
    return out


def _source_bad_caption_nums(
    table_identities: list[RenderedTableIdentity],
    source_identities: list[RenderedTableIdentity] | None,
) -> set[str]:
    """Caption numbers whose meaningful-row duplication is proven by the source
    (the formatter must not auto-delete these — they are manual/source-bad)."""
    if not source_identities:
        return set()
    source_by_num = {it.caption_num: it for it in source_identities if it.caption_num}
    bad: set[str] = set()
    for identity in table_identities:
        if identity.caption_num and _source_bad_duplicate_rows(identity, source_by_num):
            bad.add(identity.caption_num)
    return bad


def _cross_page_without_marker_blockers(
    pdf_lines: list[PdfLine],
    table_identities: list[RenderedTableIdentity],
) -> list[TableLayoutBlocker]:
    """A single logical table whose data rows render across >1 page with no
    valid ``Продолжение таблицы N`` continuation marker. KFU requires the marker
    on every continuation page; a marked split (Rybakov) is accepted."""
    out: list[TableLayoutBlocker] = []
    for identity in table_identities:
        num = identity.caption_num
        if not num:
            continue
        span = _table_page_span(pdf_lines, identity)
        if len(span) < 2:
            continue
        data_pages = sorted(
            p for p in span
            if _first_meaningful_row_on_page(identity, _page_text(pdf_lines, p), min_overlap=0.7)
        )
        if len(data_pages) < 2:
            continue
        # A single physical table renders as ONE contiguous block, so a genuine
        # cross flows onto the IMMEDIATELY following page. A non-adjacent match is a
        # DIFFERENT later table (or prose) that merely reuses this table's row text
        # — e.g. a summary/appendix table repeating the same figures — not a
        # continuation of THIS table. (`_table_page_span` runs to the next caption,
        # so for the LAST captioned table it reaches end-of-doc.) Require two
        # consecutive data pages and bound first/last to that contiguous cross.
        data_set = set(data_pages)
        cross_starts = [p for p in data_pages if (p + 1) in data_set]
        if not cross_starts:
            continue
        first = cross_starts[0]
        last = max(p + 1 for p in cross_starts)
        marker_pages = {line.page_num for line in _strict_marker_lines(pdf_lines, num)}
        has_marker = any(first < mp <= last for mp in marker_pages)
        if not has_marker:
            out.append(
                TableLayoutBlocker(
                    blocker_type="single_table_crosses_pages_without_marker",
                    severity="fail",
                    table_num=num,
                    page=first,
                    evidence={"data_pages": [p for p in data_pages if first <= p <= last]},
                )
            )
    return out


def _squeeze_blockers(
    pdf_lines: list[PdfLine],
    table_identities: list[RenderedTableIdentity],
) -> list[TableLayoutBlocker]:
    out: list[TableLayoutBlocker] = []
    for identity in table_identities:
        if not identity.caption_num:
            continue
        span = _table_page_span(pdf_lines, identity)
        if not span:
            continue
        # Count lone word-fragment lines (a single non-numeric token of <=2
        # chars, e.g. a Cyrillic "в" left over from a broken "месяцев"). Numeric
        # lines are excluded so page numbers / footers never count. A genuine
        # squeeze concentrates many such fragments on ONE page; require that
        # concentration so clean multi-page tables (Rybakov) do not false-fire.
        worst_page: int | None = None
        worst = 0
        for page in span:
            count = 0
            for line in _page_lines(pdf_lines, page):
                text = _line_text(line)
                if _CAPTION_RE.match(text) or _STRICT_MARKER_RE.match(text):
                    continue
                tokens = text.split()
                if (
                    len(tokens) == 1
                    and len(text) <= 2
                    and not text.isdigit()
                    and re.search(r"[a-zа-яё]", text, re.IGNORECASE)
                ):
                    count += 1
            if count > worst:
                worst, worst_page = count, page
        if worst >= 5:
            out.append(
                TableLayoutBlocker(
                    blocker_type="cell_text_overflow_or_illegible_squeeze",
                    severity="needs_human_review",
                    table_num=identity.caption_num,
                    page=worst_page,
                    evidence={"short_fragment_lines_on_page": worst},
                )
            )
    return out


def _semantic_header_on_continuation_blockers(
    table_identities: list[RenderedTableIdentity],
) -> list[TableLayoutBlocker]:
    """Canonical KFU continuation rule: a `Продолжение таблицы N` fragment must
    repeat ONLY the numeric column row (`1 2 3 ... N`), never the semantic header
    row(s) or the caption/title. ``header_fingerprint`` collects the non-numeric
    rows that appear BEFORE the first numeric row, so a continuation fragment
    (``preceding_marker`` set) with a NON-empty ``header_fingerprint`` still
    carries a semantic header above its numeric row — a fail. Rybakov-style valid
    continuations start with the numeric row, so their ``header_fingerprint`` is
    empty and they never fire."""
    out: list[TableLayoutBlocker] = []
    for identity in table_identities:
        if not identity.preceding_marker or identity.caption_num:
            continue
        if not identity.header_fingerprint:
            continue  # starts with the numeric row (or data) — canonical
        m = _ANY_MARKER_RE.search(identity.preceding_marker)
        num = m.group(1) if m else None
        out.append(
            TableLayoutBlocker(
                blocker_type="semantic_header_repeated_on_continuation",
                severity="fail",
                table_num=num,
                page=0,
                evidence={
                    "marker": _snippet(identity.preceding_marker),
                    "repeated_header": _snippet(" | ".join(identity.header_fingerprint)),
                },
            )
        )
    return out


def evaluate_table_layout_acceptance(
    pdf_lines: list[PdfLine],
    table_identities: list[RenderedTableIdentity],
    *,
    doc: Document | None = None,
    source_identities: list[RenderedTableIdentity] | None = None,
) -> list[TableLayoutBlocker]:
    """
    Evaluate rendered table layout and return structured blockers.

    A non-empty result with any ``severity == "fail"`` blocker means the
    rendered table layout is NO-GO. ``needs_human_review`` blockers mean the
    output is not provably clean and a human must look. Pure function: callers
    render the PDF once and pass extracted ``pdf_lines`` + DOCX identities (and
    optionally the ``doc`` for DOCX-grid checks and ``source_identities`` to
    recognise source-proven row duplication).

    Source-bad classification (D): a table whose meaningful-row duplication is
    proven by the source cannot be repaired without deleting source content
    (forbidden), so its layout *fails* are downgraded to ``needs_human_review``
    (still visible, never silently passed).
    """
    blockers: list[TableLayoutBlocker] = []
    same_page = _same_page_continuation_blockers(pdf_lines, table_identities)
    blockers.extend(same_page)
    # A repeated header on a page already flagged as a same-page continuation is
    # the same defect — don't double-count it.
    sp_keys = {(b.table_num, b.page) for b in same_page}
    blockers.extend(
        b for b in _same_page_repeated_header_blockers(pdf_lines, table_identities)
        if (b.table_num, b.page) not in sp_keys
    )
    # Marker-less same-page numeric continuation (numeric label row repeated
    # mid-table on one page, no strict marker) — the same defect class, deduped
    # against the marker-based same-page flags above.
    numeric_sp = _same_page_numeric_continuation_blockers(pdf_lines, table_identities)
    blockers.extend(b for b in numeric_sp if (b.table_num, b.page) not in sp_keys)
    sp_keys |= {(b.table_num, b.page) for b in numeric_sp}
    blockers.extend(_cross_page_without_marker_blockers(pdf_lines, table_identities))
    blockers.extend(_semantic_header_on_continuation_blockers(table_identities))
    blockers.extend(_orphaned_header_blockers(pdf_lines, table_identities))
    blockers.extend(_appendix_label_blockers(pdf_lines))
    blockers.extend(_squeeze_blockers(pdf_lines, table_identities))
    if doc is not None:
        blockers.extend(_fragment_grid_mismatch_blockers(doc, table_identities, pdf_lines))

    source_bad = _source_bad_caption_nums(table_identities, source_identities)
    if source_bad:
        downgraded: list[TableLayoutBlocker] = []
        for b in blockers:
            if b.severity == "fail" and b.table_num in source_bad:
                downgraded.append(
                    TableLayoutBlocker(
                        blocker_type=b.blocker_type,
                        severity="needs_human_review",
                        table_num=b.table_num,
                        page=b.page,
                        evidence={**b.evidence, "source_bad": True,
                                  "reason": "source_proven_duplicated_rows"},
                    )
                )
            else:
                downgraded.append(b)
        blockers = downgraded
    return blockers
