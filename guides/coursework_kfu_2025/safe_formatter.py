
import re

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt
FORMULA_NUMBER_RE = re.compile(r"\((\d+\.\d+\.\d+|\d+\.\d+)\)\s*$")
FORMULA_REFERENCE_RE = re.compile(
    r"\bформул[аеуы]\s+(\d+\.\d+(?:\.\d+)?)\b",
    re.IGNORECASE,
)
FORMULA_EXPLANATION_RE = re.compile(r"^\s*где\b", re.IGNORECASE)

MATH_TOKEN_RE = re.compile(r"[=+\-*/×÷^(){}\[\]<>]|[A-Za-zА-Яа-яЁё]\s*=")


def is_formula_paragraph_text(text: str) -> bool:
    t = clean_spaces(text)
    if not t:
        return False

    if not FORMULA_NUMBER_RE.search(t):
        return False

    # До номера должен быть не обычный текст, а выражение
    left = FORMULA_NUMBER_RE.sub("", t).strip()
    if len(left) > 120:
        return False

    # Формула должна содержать математический маркер
    return bool(MATH_TOKEN_RE.search(left))


def is_unnumbered_formula_paragraph_text(text: str) -> bool:
    t = clean_spaces(text)
    if not t or FORMULA_NUMBER_RE.search(t):
        return False
    if len(t) > 120:
        return False
    return bool(MATH_TOKEN_RE.search(t))


def is_formula_explanation_start(text: str) -> bool:
    return bool(FORMULA_EXPLANATION_RE.match(clean_spaces(text)))


def is_formula_explanation_continuation(text: str) -> bool:
    t = clean_spaces(text)
    if not t:
        return False
    if is_formula_explanation_start(t):
        return True
    # строка расшифровки символов: "V - ...", "R – ..."
    return bool(re.match(r"^[A-Za-zА-Яа-яЁё][A-Za-zА-Яа-яЁё0-9]*\s*[-–—=]\s*.+$", t))

def is_formula_block_paragraph_text(text: str) -> bool:
    """
    Возвращает True для любой строки, принадлежащей формульному блоку:
    - сама формула с номером справа: C = V*R (1.1.1)
    - первая строка пояснения: где C - ...
    - продолжение пояснений: V - ..., R – ...
    """
    t = clean_spaces(text)
    if not t:
        return False

    if is_formula_paragraph_text(t):
        return True

    if is_formula_explanation_start(t):
        return True

    if is_formula_explanation_continuation(t):
        return True

    return False

def normalize_formula_explanation_text(text: str, is_first=False) -> str:
    t = clean_spaces(text)
    t = re.sub(r"\s*,?\s*_\s*$", "", t)

    if is_first:
        t = re.sub(r"^\s*где\s*:\s*", "где ", t, flags=re.IGNORECASE)
        t = re.sub(r"^\s*где\s+", "где ", t, flags=re.IGNORECASE)
        t = re.sub(
            r"^где\s+([A-Za-zА-Яа-яЁё][A-Za-zА-Яа-яЁё0-9]*)\s*[-–—=]\s*",
            r"где \1 — ",
            t,
            flags=re.IGNORECASE,
        )

    # Нормализуем пробелы вокруг дефиса/тире после обозначения символа:
    # V- -> V – ; R –цена -> R – цена
    t = re.sub(r"^([A-Za-zА-Яа-яЁё][A-Za-zА-Яа-яЁё0-9]*)\s*[-–—=]\s*", r"\1 — ", t)

    return t


def _normalize_formula_expression(text: str) -> str:
    expr = clean_spaces(text)
    expr = re.sub(r"\s*([=+\-*/×÷^])\s*", r" \1 ", expr)
    expr = re.sub(r"\s+", " ", expr).strip()
    return expr


def _normalize_formula_explanation_punctuation(text: str, *, is_last: bool) -> str:
    cleaned = clean_spaces(text)
    cleaned = re.sub(r"\s*[,;._]\s*$", "", cleaned)
    return cleaned + ("." if is_last else ";")


def _find_formula_number_from_preceding_prose(paragraphs, idx: int, body_start: int) -> str | None:
    seen = 0
    j = idx - 1
    while j >= body_start and seen < 4:
        text = clean_spaces(paragraphs[j].text)
        if text:
            seen += 1
            match = FORMULA_REFERENCE_RE.search(text)
            if match:
                return match.group(1)
        j -= 1
    return None


def _next_nonempty_paragraph_starts_formula_explanation(paragraphs, idx: int) -> bool:
    j = idx + 1
    while j < len(paragraphs):
        text = clean_spaces(paragraphs[j].text)
        if text:
            return is_formula_explanation_start(text)
        j += 1
    return False


def split_formula_explanations_in_paragraph(paragraph, is_first=False):
    """
    Если в одном абзаце склеены несколько расшифровок формулы через ';',
    разбивает их на отдельные абзацы и сохраняет финальные ';' там, где они были.
    """
    text = normalize_formula_explanation_text(paragraph.text, is_first=is_first)
    if not text:
        replace_paragraph_text(paragraph, "")
        return []

    has_trailing_semicolon = text.rstrip().endswith(";")
    raw_parts = [clean_spaces(x) for x in text.split(";") if clean_spaces(x)]
    if not raw_parts:
        replace_paragraph_text(paragraph, text)
        return []

    rebuilt = []
    for i, part in enumerate(raw_parts):
        is_last = i == len(raw_parts) - 1
        if not is_last:
            rebuilt.append(part + ";")
        else:
            rebuilt.append(part + (";" if has_trailing_semicolon else ""))

    replace_paragraph_text(paragraph, rebuilt[0])
    inserted = []

    prev = paragraph
    for extra in rebuilt[1:]:
        new_p = insert_paragraph_after(prev, extra)
        inserted.append(new_p)
        prev = new_p

    return inserted

def format_formula_paragraph(paragraph):
    text = clean_spaces(paragraph.text)
    m = FORMULA_NUMBER_RE.search(text)
    if not m:
        return

    number = m.group(0)
    expr = _normalize_formula_expression(text[:m.start()].rstrip())

    replace_paragraph_text(paragraph, f"\t{expr}\t{number}")

    hard_reset_paragraph_format(paragraph, first_line_indent_cm=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)

    tabs = paragraph.paragraph_format.tab_stops
    tabs.clear_all()
    tabs.add_tab_stop(Cm(8), WD_TAB_ALIGNMENT.CENTER)
    tabs.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False)

def format_formula_explanation_paragraph(paragraph, is_first=False):
    hard_reset_paragraph_format(paragraph, first_line_indent_cm=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Cm(0)

    text = normalize_formula_explanation_text(paragraph.text, is_first=is_first)
    replace_paragraph_text(paragraph, text)

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False)

def normalize_formula_blocks(document, body_start):
    changed = False
    paragraphs = document.paragraphs
    idx = max(body_start, 0)

    while idx < len(paragraphs):
        paragraphs = document.paragraphs
        p = paragraphs[idx]
        text = clean_spaces(p.text)

        if not is_formula_paragraph_text(text):
            if (
                is_unnumbered_formula_paragraph_text(text)
                and _next_nonempty_paragraph_starts_formula_explanation(paragraphs, idx)
            ):
                formula_number = _find_formula_number_from_preceding_prose(paragraphs, idx, body_start)
                if formula_number:
                    replace_paragraph_text(p, f"{text} ({formula_number})")
                    text = clean_spaces(p.text)
                    changed = True
                else:
                    idx += 1
                    continue
            else:
                idx += 1
                continue

        if not is_formula_paragraph_text(text):
            idx += 1
            continue

        # 1. Форматируем строку формулы
        format_formula_paragraph(p)

        # 2. Перед формулой должна быть ровно одна пустая строка
        if idx > body_start:
            prev_idx = idx - 1
            if not is_empty_paragraph(paragraphs[prev_idx]):
                new_p = insert_paragraph_after(paragraphs[prev_idx], "")
                format_empty_paragraph(new_p)
                changed = True
                paragraphs = document.paragraphs
                idx += 1
                p = paragraphs[idx]
            else:
                while prev_idx - 1 >= body_start and is_empty_paragraph(paragraphs[prev_idx - 1]):
                    remove_paragraph(paragraphs[prev_idx - 1])
                    changed = True
                    paragraphs = document.paragraphs
                    idx -= 1
                    prev_idx -= 1
                format_empty_paragraph(paragraphs[prev_idx])

        # 3. Форматируем блок "где ..."
        paragraphs = document.paragraphs
        j = idx + 1
        first_expl = True
        explanation_paragraphs = []

        while j < len(paragraphs):
            t = clean_spaces(paragraphs[j].text)

            if not t:
                break

            if first_expl and is_formula_explanation_start(t):
                inserted = split_formula_explanations_in_paragraph(paragraphs[j], is_first=True)
                explanation_paragraphs.append((paragraphs[j], True))
                for new_p in inserted:
                    explanation_paragraphs.append((new_p, False))

                if inserted:
                    changed = True
                    paragraphs = document.paragraphs

                first_expl = False
                j += 1 + len(inserted)
                continue

            if not first_expl and is_formula_explanation_continuation(t):
                inserted = split_formula_explanations_in_paragraph(paragraphs[j], is_first=False)
                explanation_paragraphs.append((paragraphs[j], False))
                for new_p in inserted:
                    explanation_paragraphs.append((new_p, False))

                if inserted:
                    changed = True
                    paragraphs = document.paragraphs

                j += 1 + len(inserted)
                continue

            break

        for pos, (expl_p, is_first) in enumerate(explanation_paragraphs):
            format_formula_explanation_paragraph(expl_p, is_first=is_first)
            normalized = _normalize_formula_explanation_punctuation(
                expl_p.text,
                is_last=pos == len(explanation_paragraphs) - 1,
            )
            replace_paragraph_text(expl_p, normalized)
            for run in expl_p.runs:
                set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False)

        # 4. После формулы/пояснений должна быть ровно одна пустая строка
        tail_idx = j - 1 if j > idx + 1 else idx
        paragraphs = document.paragraphs

        if tail_idx + 1 >= len(paragraphs):
            new_p = insert_paragraph_after(paragraphs[tail_idx], "")
            format_empty_paragraph(new_p)
            changed = True
            paragraphs = document.paragraphs
        elif not is_empty_paragraph(paragraphs[tail_idx + 1]):
            new_p = insert_paragraph_after(paragraphs[tail_idx], "")
            format_empty_paragraph(new_p)
            changed = True
            paragraphs = document.paragraphs
        else:
            format_empty_paragraph(paragraphs[tail_idx + 1])
            while tail_idx + 2 < len(paragraphs) and is_empty_paragraph(paragraphs[tail_idx + 2]):
                remove_paragraph(paragraphs[tail_idx + 2])
                changed = True
                paragraphs = document.paragraphs

        idx = tail_idx + 2

    return changed
    
def clear_paragraph_outline_level(paragraph):
    try:
        pPr = paragraph._element.get_or_add_pPr()
        outline = pPr.find(qn("w:outlineLvl"))
        if outline is not None:
            pPr.remove(outline)
    except Exception:
        pass



def set_paragraph_style_safe(paragraph, *style_names):
    for name in style_names:
        try:
            paragraph.style = name
            return True
        except Exception:
            pass
    return False




# ===== STRUCTURAL SPACING FIX =====

STRUCTURAL_HEADINGS = {
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
}

def enforce_structural_spacing(doc):

    paragraphs = doc.paragraphs
    i = 0

    while i < len(paragraphs):

        p = paragraphs[i]
        text = p.text.strip().upper()

        if text in STRUCTURAL_HEADINGS:

            j = i + 1
            blank_count = 0

            while j < len(paragraphs) and not paragraphs[j].text.strip():
                blank_count += 1
                j += 1

            if blank_count == 0:
                new = insert_paragraph_after(p, "")
                new.paragraph_format.space_before = 0
                new.paragraph_format.space_after = 0
                i += 2
                paragraphs = doc.paragraphs
                continue

            if blank_count > 1:
                for k in range(i + 2, i + 1 + blank_count):
                    remove_paragraph(paragraphs[i + 2])

        i += 1

# ===== END STRUCTURAL SPACING FIX =====




# ===== AUTO PATCH: robust heading2 detection =====

def auto_detect_heading2(paragraph, current_chapter_num, next_paragraph_num, prev_kind=None):
    if current_chapter_num is None or next_paragraph_num is None:
        return False

    text = clean_spaces(paragraph.text)
    if not text:
        return False

    low = text.lower()

    forbidden_prefixes = (
        "таблица ",
        "рисунок ",
        "рис. ",
        "продолжение таблицы",
        "продолжение табл.",
        "источник:",
        "составлено по:",
        "рассчитано по:",
        "примечание:",
    )
    if low.startswith(forbidden_prefixes):
        return False

    if parse_heading1(text) or parse_heading2(text) or parse_broken_heading2(text):
        return False

    if is_table_continuation_text(text):
        return False

    if not looks_like_heading2_title(text):
        return False

    if is_heading2_promotion_safe(paragraph):
        return True

    return False
def auto_detect_numbered_heading1(paragraph, current_chapter_num=None, next_paragraph=None):
    text = clean_spaces(paragraph.text)
    if not text:
        return False

    low = text.lower()

    # Уже распознанный heading1 не трогаем
    if parse_heading1(text):
        return False

    # Не трогаем подписи таблиц/рисунков и служебные строки
    forbidden_prefixes = (
        "таблица",
        "табл.",
        "рисунок",
        "рис.",
        "источник:",
        "составлено по:",
        "рассчитано по:",
        "примечание:",
        "продолжение таблицы",
        "продолжение табл.",
    )
    if low.startswith(forbidden_prefixes):
        return False

    # Нужна именно Word-автонумерация / numbering
    if not paragraph_has_numbering(paragraph):
        return False

    has_structural_heading_signal = paragraph_has_heading_style_or_outline(paragraph)
    has_visual_heading_signal = is_probable_center_bold_heading(paragraph)
    if not (has_structural_heading_signal or has_visual_heading_signal):
        return False

    # Если это уже похоже на heading2, не считаем heading1
    if parse_heading2(text) or parse_broken_heading2(text):
        return False

    # Запрещённые финальные знаки
    if text.endswith((":", ";", "?", "!")):
        return False

    words = text.split()
    word_limit = 12 if "." in text else 15
    if len(words) < 1 or len(words) > word_limit:
        return False

    # Если следующий абзац тоже numbered и тоже короткий,
    # это больше похоже на список, а не на heading1
    if next_paragraph is not None and not has_structural_heading_signal:
        next_text = clean_spaces(next_paragraph.text)
        if next_text and paragraph_has_numbering(next_paragraph):
            if not parse_heading1(next_text) and not parse_heading2(next_text):
                next_words = next_text.split()
                next_limit = 12 if "." in next_text else 15
                if 1 <= len(next_words) <= next_limit and not next_text.endswith((":", ";", "?", "!")):
                    return False

    return True

def is_structural_heading_paragraph(paragraph):
    t = clean_spaces(paragraph.text).upper()
    return t in STRUCTURAL_HEADING_TEXTS

def enforce_single_blank_after_structural_headings(doc, body_start_idx=0):
    paragraphs = doc.paragraphs
    i = max(body_start_idx, 0)

    while i < len(paragraphs):
        p = paragraphs[i]
        if not is_structural_heading_paragraph(p):
            i += 1
            continue

        j = i + 1
        blank_idxs = []

        while j < len(paragraphs) and not clean_spaces(paragraphs[j].text):
            blank_idxs.append(j)
            j += 1

        if not blank_idxs:
            insert_paragraph_after(p, "")
            paragraphs = doc.paragraphs
            i += 2
            continue

        first_blank_idx = blank_idxs[0]
        for idx in reversed(blank_idxs[1:]):
            remove_paragraph(paragraphs[idx])

        paragraphs = doc.paragraphs
        format_empty_spacing_paragraph(paragraphs[first_blank_idx])
        i = first_blank_idx + 1

# ===== END FINAL PATCH =====



from pathlib import Path
import re
from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, Mm, RGBColor, RGBColor
from docx.text.paragraph import Paragraph

from .rules import (
    FONT_NAME,
    BODY_FONT_SIZE_PT,
    TABLE_FONT_SIZE_PT,
    LINE_SPACING_BODY,
    LINE_SPACING_TABLE,
    FIRST_LINE_INDENT_CM,
    LEFT_MARGIN_MM,
    RIGHT_MARGIN_MM,
    TOP_MARGIN_MM,
    BOTTOM_MARGIN_MM,
)
from .classifier import (
    find_body_start_index,
    classify_paragraph,
    clean_spaces,
    is_intro_heading_text,
    paragraph_text,
    parse_heading1,
    parse_heading2,
    parse_broken_heading2,
    caption_tail_is_reference_prose,
)
from .page_numbering import apply_page_numbering_policy
from .page_breaks import apply_page_breaks

MAX_NORMALIZATION_PASSES = 35

def run_with_pass_limit(step_name, func, document, body_start):
    """
    Re-run a normalization step until it stabilizes, but avoid full-text snapshots
    of the whole document on every pass.

    Preferred contract: a step may return:
      - True / positive int  -> document changed, run another pass
      - False / 0 / None     -> no changes, step is stable

    Backward compatibility: if a step returns None, we fall back to a cheap
    structural signature based on paragraph count and lengths.
    """
    previous_signature = None

    for _ in range(MAX_NORMALIZATION_PASSES):
        paragraphs = document.paragraphs
        before_signature = (
            len(paragraphs),
            sum(len(p.text) for p in paragraphs),
        )

        result = func(document, body_start)

        if isinstance(result, bool):
            if not result:
                return
            previous_signature = None
            continue

        if isinstance(result, int):
            if result <= 0:
                return
            previous_signature = None
            continue

        paragraphs_after = document.paragraphs
        after_signature = (
            len(paragraphs_after),
            sum(len(p.text) for p in paragraphs_after),
        )

        if after_signature == before_signature:
            return

        if after_signature == previous_signature:
            raise RuntimeError(f"Formatter step stuck: {step_name}")

        previous_signature = after_signature

    raise RuntimeError(f"Formatter step stuck: {step_name}")


TABLE_NUM_RE = re.compile(r"^\s*таблица\s*(\d+(?:\.\d+){0,2})\.?\s*(.*?)\s*$", re.IGNORECASE)
DASH_LINE_RE = re.compile(r"^\s*[—–\-•]\s*.+$")
FIG_RE = re.compile(
    r"^\s*(рисунок|рис\.)\s*(\d+(?:\.\d+){0,2})(?:\s*[.\-—–]?\s*(.+?))?\s*$",
    re.IGNORECASE,
)
FIG_SERVICE_LINE_RE = re.compile(
    r"^\s*(источник|составлено по|рассчитано по|примечание)\s*:",
    re.IGNORECASE,
)
HEADING2_ARTIFACT_RE = re.compile(r"^\s*[•·▪■◆►→\-–—]*\s*(\d+\.\d+\.?)\s*[•·▪■◆►→\-–—]*\s*(.+?)\s*$")

TABLE_CONTINUATION_RE = re.compile(r"^\s*продолжение\s+табл(?:ицы)?\.?\s*\d+(?:\.\d+){1,2}\.?\s*$", re.IGNORECASE)

def is_table_continuation_text(text: str) -> bool:
    t = clean_spaces(text)
    if not t:
        return False

    if TABLE_CONTINUATION_RE.match(t):
        return True

    # Защита от уже испорченных вариантов: "1.2. Продолжение таблицы 1.1.1"
    t2 = re.sub(r'^\s*\d+\.\d+\.?\s*', '', t, count=1)
    return bool(TABLE_CONTINUATION_RE.match(t2))


REFERENCE_SUBHEADINGS_CANON = {
    "официальные материалы": "Официальные материалы",
    "статистические материалы": "Статистические материалы",
    "справочные и архивные материалы": "Справочные и архивные материалы",
    "книги, монографии и диссертации": "Книги, монографии и диссертации",
    "книги, монографии, диссертации": "Книги, монографии, диссертации",
    "научные статьи": "Научные статьи",
    "монографии и статьи": "Монографии и статьи",
    "монографии и учебники": "Монографии и учебники",
    "учебники и учебные пособия": "Учебники и учебные пособия",
    "учебники, учебные пособия и материалы": "Учебники, учебные пособия и материалы",
    "электронные ресурсы": "Электронные ресурсы",
    "электронные источники": "Электронные источники",
    "интернет-ресурсы": "Интернет-ресурсы",
    "материалы на иностранных языках": "Материалы на иностранных языках",
    "нормативные правовые акты": "Нормативные правовые акты",
    "монографии, учебники": "Монографии, учебники",
    "монографии, учебники, учебные пособия": "Монографии, учебники, учебные пособия",
    "статьи": "Статьи",
    "статьи в периодических изданиях и сборниках": "Статьи в периодических изданиях и сборниках",
    "диссертации": "Диссертации",
    "диссертации, авторефераты диссертаций": "Диссертации, авторефераты диссертаций",
    "материалы интернет-сайтов": "Материалы интернет-сайтов",
}


REFERENCE_BLOCK_HEADING_MAX_CHARS = 64
REFERENCE_STYLED_HEADING_MAX_CHARS = 48
REFERENCE_HEADING_NUMBER_RE = re.compile(r"^\s*\d{1,3}[\.)]\s+(.+)$")
REFERENCE_SOURCE_LIKE_RE = re.compile(
    r"(https?://|www\.|doi\b|isbn\b|//|\b(?:19|20)\d{2}\b|"
    r"\b[А-ЯЁA-Z]\.\s*[А-ЯЁA-Z]\.|"
    r"\b(?:федеральн\w*\s+закон|кодекс|гост|№|статья\s+\d|"
    r"от\s+\d{2}\.\d{2}\.\d{4}|фз)\b)",
    re.IGNORECASE,
)
REFERENCE_HEADING_KEYWORDS = (
    "материал",
    "ресурс",
    "источник",
    "стать",
    "книг",
    "монограф",
    "диссертац",
    "учебник",
    "пособ",
    "акт",
)


def _reference_heading_key(text: str) -> str:
    t = clean_spaces(text).lower().replace("ё", "е").rstrip(".")
    t = re.sub(r"\s*,\s*", ", ", t)
    return clean_spaces(t)


def _edit_distance_at_most_two(a: str, b: str) -> bool:
    if abs(len(a) - len(b)) > 2:
        return False

    previous = list(range(len(b) + 1))
    for i, ca in enumerate(a, 1):
        current = [i]
        row_min = i
        for j, cb in enumerate(b, 1):
            cost = 0 if ca == cb else 1
            current.append(min(
                previous[j] + 1,
                current[j - 1] + 1,
                previous[j - 1] + cost,
            ))
            row_min = min(row_min, current[-1])
        if row_min > 2:
            return False
        previous = current
    return previous[-1] <= 2


def _looks_like_reference_source_entry(text: str) -> bool:
    t = clean_spaces(text)
    if len(t) > REFERENCE_BLOCK_HEADING_MAX_CHARS:
        return True
    if REFERENCE_SOURCE_LIKE_RE.search(t):
        return True
    return t.count(".") >= 2


def _canonical_reference_heading_candidate(text: str):
    t = clean_spaces(text).rstrip(".")
    if not t:
        return None

    key = _reference_heading_key(t)
    exact = REFERENCE_SUBHEADINGS_CANON.get(key)
    if exact:
        return exact

    if _looks_like_reference_source_entry(t):
        return None

    for known_key, canonical in REFERENCE_SUBHEADINGS_CANON.items():
        if _edit_distance_at_most_two(key, _reference_heading_key(known_key)):
            return canonical
    return None


def _looks_like_short_styled_reference_heading(text: str) -> bool:
    t = clean_spaces(text).rstrip(".")
    if not t or len(t) > REFERENCE_STYLED_HEADING_MAX_CHARS:
        return False
    if _looks_like_reference_source_entry(t):
        return False

    low = t.lower()
    return any(keyword in low for keyword in REFERENCE_HEADING_KEYWORDS)


def _canonicalize_styled_reference_heading(text: str) -> str:
    t = clean_spaces(text).rstrip(".")
    if not t:
        return t
    if t.isupper() or t.islower():
        low = t.lower()
        return low[:1].upper() + low[1:]
    return t


def canonical_reference_block_heading_paragraph(paragraph):
    canonical = canonical_reference_block_heading_text(paragraph.text)
    if canonical:
        return canonical

    text = clean_spaces(paragraph.text)
    if not _looks_like_short_styled_reference_heading(text):
        return None

    is_centered = paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER
    is_bold = any(run.bold is True for run in paragraph.runs)
    if is_centered and is_bold:
        return _canonicalize_styled_reference_heading(text)
    return None


def insert_paragraph_after(paragraph, text=""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def remove_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def replace_paragraph_text(paragraph, new_text: str):
    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)
    paragraph.add_run(new_text)


def paragraph_has_drawing(paragraph) -> bool:
    p = paragraph._element
    return bool(
        p.xpath(
            ".//*[local-name()='drawing' or local-name()='pict' or local-name()='object']"
        )
    )

def center_image_paragraphs(document, body_start):
    """
    Центрирует абзацы, содержащие рисунки (drawing/pict/object).

    Не трогает:
    - подписи рисунков
    - таблицы
    - текст

    Возвращает True только если реально изменил выравнивание.
    Это важно, чтобы run_with_pass_limit(...) не зацикливался.
    """
    changed = False

    paragraphs = document.paragraphs

    for idx, paragraph in enumerate(paragraphs):
        if idx < body_start:
            continue

        if not paragraph_has_drawing(paragraph):
            continue

        if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            changed = True

    return changed

def is_empty_paragraph(paragraph):
    return clean_spaces(paragraph.text) == "" and not paragraph_has_drawing(paragraph)


def is_reference_spacing_paragraph(paragraph):
    return (
        clean_spaces(paragraph.text) == ""
        and not paragraph._element.findall(".//" + qn("w:drawing"))
    )

def ensure_empty_run(paragraph):
    if not paragraph.runs:
        paragraph.add_run("")
    return paragraph.runs[0]


def force_paragraph_xml_spacing(paragraph, line_rule="auto"):
    pPr = paragraph._element.get_or_add_pPr()

    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)

    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:beforeAutospacing"), "0")
    spacing.set(qn("w:afterAutospacing"), "0")

    if line_rule == "auto":
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:line"), "360")
    elif line_rule == "exact":
        spacing.set(qn("w:lineRule"), "exact")
    elif line_rule == "atLeast":
        spacing.set(qn("w:lineRule"), "atLeast")

    snap = pPr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        pPr.append(snap)
    snap.set(qn("w:val"), "0")


def hard_reset_paragraph_format(paragraph, first_line_indent_cm=None):
    # Explicitly nuke any existing w:ind to prevent stale hanging/left attributes
    # from paragraph styles (e.g. List Paragraph) overriding our settings.
    _pPr = paragraph._element.find(qn("w:pPr"))
    if _pPr is not None:
        for _old_ind in list(_pPr.findall(qn("w:ind"))):
            _pPr.remove(_old_ind)

    force_paragraph_xml_spacing(paragraph, line_rule="auto")
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = LINE_SPACING_BODY

    # Write w:ind directly to guarantee w:left="0" is in XML (not just "default").
    # We include w:right="0" here instead of using fmt.right_indent to avoid
    # python-docx creating a second w:ind element.
    _pPr2 = paragraph._element.get_or_add_pPr()
    _ind = OxmlElement("w:ind")
    _ind.set(qn("w:left"), "0")
    _ind.set(qn("w:right"), "0")
    if first_line_indent_cm:
        # 1.25 cm = 709 twips (1 inch = 2.54 cm = 1440 twips → 1 cm = 566.9 twips)
        _twips = round(first_line_indent_cm * 1440 / 2.54)
        _ind.set(qn("w:firstLine"), str(_twips))
    _pPr2.append(_ind)

    fmt.keep_together = False
    fmt.keep_with_next = False
    fmt.page_break_before = False
    fmt.widow_control = False


def set_run_font(run, font_name=FONT_NAME, size_pt=BODY_FONT_SIZE_PT, bold=None, italic=False, all_caps=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)

    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.rFonts
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)

    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if all_caps is not None:
        run.font.all_caps = all_caps

    try:
        run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass

    color = rPr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rPr.append(color)

    color.set(qn("w:val"), "000000")

    for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        try:
            if color.get(qn(attr)) is not None:
                del color.attrib[qn(attr)]
        except Exception:
            pass

def set_section_margins(document):
    # Глобально отключаем зеркальные поля на уровне settings.xml
    try:
        settings_el = document.settings._element
        mirror = settings_el.find(qn("w:mirrorMargins"))
        if mirror is not None:
            settings_el.remove(mirror)
    except Exception:
        pass

    for section in document.sections:
        section.left_margin = Mm(LEFT_MARGIN_MM)
        section.right_margin = Mm(RIGHT_MARGIN_MM)
        section.top_margin = Mm(TOP_MARGIN_MM)
        section.bottom_margin = Mm(BOTTOM_MARGIN_MM)

        # На уровне секции убираем gutter/переплёт и следы зеркалинга
        try:
            sectPr = section._sectPr

            pgMar = sectPr.find(qn("w:pgMar"))
            if pgMar is not None:
                pgMar.set(qn("w:left"), str(Mm(LEFT_MARGIN_MM)._emu))
                pgMar.set(qn("w:right"), str(Mm(RIGHT_MARGIN_MM)._emu))
                pgMar.set(qn("w:top"), str(Mm(TOP_MARGIN_MM)._emu))
                pgMar.set(qn("w:bottom"), str(Mm(BOTTOM_MARGIN_MM)._emu))
                pgMar.set(qn("w:gutter"), "0")

            gutter = sectPr.find(qn("w:gutter"))
            if gutter is not None:
                sectPr.remove(gutter)
        except Exception:
            pass
def normalize_simple_paragraph_spaces(paragraph):
    if len(paragraph.runs) == 1 and "\n" not in paragraph.runs[0].text and "\v" not in paragraph.runs[0].text:
        old = paragraph.runs[0].text
        new = clean_spaces(old)
        if new != old:
            paragraph.runs[0].text = new

QUOTE_CHARS_DOUBLE = {
    '"',      # ASCII
    '“', '”', # curly double
    '„', '‟', # low/high double
    '«', '»', # уже правильные, но учитываем в общем потоке
    '″', '‟', '〝', '〞', '＂',
}

def _normalize_quotes_in_text_fragment(text: str, quote_state: dict) -> str:
    """
    Меняет все двойные кавычки на «» по принципу открытия/закрытия.
    Не трогает одинарные апострофы и штрихи — это сознательное ограничение
    ради безопасности обычных курсовых.
    """
    if not text:
        return text

    out = []
    for ch in text:
        if ch in QUOTE_CHARS_DOUBLE:
            if quote_state["open"]:
                out.append("«")
            else:
                out.append("»")
            quote_state["open"] = not quote_state["open"]
        else:
            out.append(ch)
    return "".join(out)


def normalize_quotes_in_paragraph_runs(paragraph, quote_state: dict):
    """
    Нормализует кавычки в run-ах абзаца без пересборки абзаца,
    чтобы не ломать гиперссылки, разметку и прочую структуру Word.
    """
    for run in paragraph.runs:
        old = run.text
        new = _normalize_quotes_in_text_fragment(old, quote_state)
        if new != old:
            run.text = new


def normalize_quotes_in_document(document, body_start=0):
    """
    Проходит по рабочей части документа последовательно сверху вниз.
    Состояние открытия/закрытия кавычек сохраняется между абзацами.
    """
    quote_state = {"open": True}

    for idx, paragraph in enumerate(document.paragraphs):
        if idx < body_start:
            continue
        normalize_quotes_in_paragraph_runs(paragraph, quote_state)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    normalize_quotes_in_paragraph_runs(paragraph, quote_state)

def normalize_dashes_in_runs(paragraph):
    for run in paragraph.runs:
        if '\u2014' in run.text:
            run.text = run.text.replace('\u2014', '\u2013')


def normalize_dashes_in_document(document, body_start):
    in_references = False
    for idx, paragraph in enumerate(document.paragraphs):
        if idx < body_start:
            continue
        text = clean_spaces(paragraph.text)
        if is_references_heading_text(text):
            in_references = True
            continue
        if in_references and is_appendix_heading_text(text):
            in_references = False
        if in_references:
            continue
        normalize_dashes_in_runs(paragraph)


_WORD_RE = re.compile(r'[А-ЯЁа-яё]+')


def normalize_yo_in_text(text: str) -> str:
    """Replace lowercase ё with е. Capital Ё is preserved (Python replace is case-sensitive)."""
    return text.replace('ё', 'е')


def normalize_yo_in_runs(paragraph):
    for run in paragraph.runs:
        if 'ё' in run.text:
            run.text = normalize_yo_in_text(run.text)


def normalize_yo_in_document(document, body_start):
    in_references = False
    for idx, paragraph in enumerate(document.paragraphs):
        if idx < body_start:
            continue
        text = clean_spaces(paragraph.text)
        if is_references_heading_text(text):
            in_references = True
            continue
        if in_references and is_appendix_heading_text(text):
            in_references = False
        if in_references:
            continue
        normalize_yo_in_runs(paragraph)


def normalize_semicolons_in_document(document, body_start):
    in_references = False
    for idx, paragraph in enumerate(document.paragraphs):
        if idx < body_start:
            continue
        text = clean_spaces(paragraph.text)
        if is_references_heading_text(text):
            in_references = True
            continue
        if in_references and is_appendix_heading_text(text):
            in_references = False
        if in_references:
            continue
        for run in paragraph.runs:
            if ';' in run.text:
                run.text = run.text.replace(';', ',')


# ── Source citation bracket splitting ────────────────────────────────────────
_CITATION_ENTRY_RE = re.compile(
    r'(\d+)(?:\s*[,;]\s*(?:с\.?|p\.?)\s*(\d+(?:\s*[–\-—]\s*\d+)?))?',
    re.IGNORECASE
)
_BRACKET_GROUP_RE = re.compile(r'\[([^\]]+)\]')


def _parse_citation_content(content: str) -> list[str]:
    """
    Parse inner content of a citation bracket into individual source strings.
    Handles both с. and p. page markers, semicolon separators.
    """
    # Normalise semicolons to commas
    content = content.replace(';', ',')
    entries = []
    pos = 0
    s = content.strip()

    while pos < len(s):
        skip = re.match(r'[\s,]+', s[pos:])
        if skip:
            pos += skip.end()
        if pos >= len(s):
            break

        m = _CITATION_ENTRY_RE.match(s, pos)
        if not m:
            return [s.strip()]  # unrecognised — return as single entry

        num = m.group(1)
        pages = m.group(2)
        if pages:
            # Normalise hyphen/em-dash to en-dash
            pages_norm = re.sub(r'\s*[\-—]\s*', '–', pages.strip())
            entries.append(f"{num}, с. {pages_norm}")
        else:
            entries.append(num)
        pos = m.end()

    return entries if entries else [s.strip()]


def _split_citation_brackets_in_text(text: str) -> str:
    """
    Find all [...] citation groups and split multi-source ones.
    Also normalises hyphens to en-dashes in page ranges within single-source citations.
    """
    def replace_bracket(m):
        content = m.group(1)
        entries = _parse_citation_content(content)
        if len(entries) <= 1:
            # Single source — just normalise hyphen in page range if present
            if entries:
                return f"[{entries[0]}]"
            return m.group(0)
        return ', '.join(f'[{e}]' for e in entries)

    return _BRACKET_GROUP_RE.sub(replace_bracket, text)


def normalize_citations_in_paragraph_runs(paragraph):
    """Apply citation bracket splitting to all runs in a paragraph."""
    for run in paragraph.runs:
        if '[' in run.text:
            new_text = _split_citation_brackets_in_text(run.text)
            if new_text != run.text:
                run.text = new_text


def normalize_citations_in_document(document, body_start):
    """
    Split multi-source citation brackets in all body paragraphs.
    Skips the references block.
    """
    in_references = False
    for idx, paragraph in enumerate(document.paragraphs):
        if idx < body_start:
            continue
        text = clean_spaces(paragraph.text)
        if is_references_heading_text(text):
            in_references = True
            continue
        if in_references and is_appendix_heading_text(text):
            in_references = False
        if in_references:
            continue
        if '[' in (paragraph.text or ''):
            normalize_citations_in_paragraph_runs(paragraph)


# ── List formatting (Level 1: а/б/в, Level 2: 1/2/3) ─────────────────────────

_CYRILLIC_LIST_ALPHA = 'абвгдежзиклмнопрстуфхцчшщэюя'
_CYRILLIC_LETTER_LIST_RE = re.compile(r'^([а-яё])\)\s+(.+)$', re.DOTALL)
_NUMERIC_PAREN_LIST_RE   = re.compile(r'^(\d+)\)\s+(.+)$',   re.DOTALL)
_NUMERIC_DOT_LIST_RE     = re.compile(r'^(\d+)\.\s+(.+)$',   re.DOTALL)
# Manual L1 dash/bullet markers (ascii hyphen, en/em dash, common black bullet
# glyphs). Methodical first-level list marker is en-dash `–`; everything in
# this class normalises to `– text`.
_DASH_BULLET_LIST_RE = re.compile(
    r'^([-–—•·●▪◆■►◦○])\s+(.+)$',
    re.DOTALL,
)

# Custom XML attribute key for marking KFU-converted list items so that a
# second pass of _normalize_plain_list_paragraphs can re-apply the correct
# formatting (e.g. restore {} ind for colon items, preserve – for en-dash items)
# without re-normalising them as free-standing dash blocks.
_KFU_LIST_TYPE_KEY = "{http://kfu.ru/formatter}listType"


def _mark_kfu_list_type(paragraph, list_type: str) -> None:
    """Stamp paragraph._element with the KFU list type for idempotent re-runs."""
    paragraph._element.set(_KFU_LIST_TYPE_KEY, list_type)


def _get_kfu_list_type(paragraph) -> str:
    """Return the KFU list type marker, or '' if not set."""
    return paragraph._element.get(_KFU_LIST_TYPE_KEY, "")


def _is_level1_list_text(text: str) -> bool:
    t = clean_spaces(text)
    return bool(
        _CYRILLIC_LETTER_LIST_RE.match(t)
        or _NUMERIC_PAREN_LIST_RE.match(t)
        or _NUMERIC_DOT_LIST_RE.match(t)
    )


def _is_dash_or_bullet_list_text(text: str) -> bool:
    return bool(_DASH_BULLET_LIST_RE.match(clean_spaces(text)))


def _is_letter_list_text(text: str) -> bool:
    return bool(_CYRILLIC_LETTER_LIST_RE.match(clean_spaces(text)))


def _lowercase_first(text: str) -> str:
    """Lowercase the first character of list body text.

    Only acts when text[0] is uppercase AND text[1] is lowercase (i.e. a
    capitalised word like "Раскрыть", not an acronym like "ПАО" or "IBM"
    whose second character is also uppercase).
    """
    if not text or not text[0].isupper():
        return text
    if len(text) > 1 and text[1].isupper():
        return text  # acronym — leave untouched
    return text[0].lower() + text[1:]


def _apply_list_indent_xml(paragraph, left_twips: int, hanging_twips: int):
    """Set list hanging indent directly via XML."""
    pPr = paragraph._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:ind"))):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),    str(left_twips))
    ind.set(qn("w:hanging"), str(hanging_twips))
    ind.set(qn("w:right"),   "0")
    pPr.append(ind)


def _apply_list_firstline_xml(paragraph, first_line_twips: int = 708):
    """Set firstLine-only indent directly via XML (KFU L1 standard: 1.25 cm ≈ 708 twips).

    Uses w:firstLine so the first line is indented 1.25 cm and continuation
    (wrapped) lines start at the left margin (0 cm), matching KFU list style.
    """
    pPr = paragraph._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:ind"))):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), str(first_line_twips))
    pPr.append(ind)


def _clear_paragraph_ind(paragraph):
    """Remove any explicit w:ind element from the paragraph's pPr."""
    pPr = paragraph._element.get_or_add_pPr()
    for old in list(pPr.findall(qn("w:ind"))):
        pPr.remove(old)


# ── KFU internal Word dash-list (real autonumbering) ─────────────────────────
# The KFU internal dash-list abstractNum is identified by a custom NSID so that
# repeated formatter runs find and reuse the same definition rather than
# appending a new one each time.
_KFU_DASH_NSID = "FFFEFFFF"


def _ensure_kfu_dash_list_numid(paragraph) -> "str | None":
    """
    Return the w:numId string for the KFU internal dash-list numbering definition,
    creating it in the document's numbering part if absent.

    The definition is identified by the NSID marker ``FFFEFFFF`` on the
    abstractNum element, so repeated calls always reuse the same numId.
    Returns ``None`` if the numbering part is not accessible (safe fallback).
    """
    try:
        numbering_part = paragraph.part.numbering_part
    except Exception:
        return None
    if numbering_part is None:
        return None

    numbering_el = numbering_part._element

    # ── 1. Look for existing KFU dash abstractNum by NSID ──────────────────
    kfu_abstract_num_id: "str | None" = None
    for ab in numbering_el.findall(qn("w:abstractNum")):
        nsid_el = ab.find(qn("w:nsid"))
        if nsid_el is not None and nsid_el.get(qn("w:val")) == _KFU_DASH_NSID:
            kfu_abstract_num_id = ab.get(qn("w:abstractNumId"))
            break

    if kfu_abstract_num_id is not None:
        # Find the w:num that references it
        for num_el in numbering_el.findall(qn("w:num")):
            ref = num_el.find(qn("w:abstractNumId"))
            if ref is not None and ref.get(qn("w:val")) == kfu_abstract_num_id:
                return num_el.get(qn("w:numId"))
        # abstractNum exists but no num points to it — fall through to create num
    else:
        # ── 2. Create new abstractNum ───────────────────────────────────────
        existing_ab_ids = [
            int(n.get(qn("w:abstractNumId"), "-1"))
            for n in numbering_el.findall(qn("w:abstractNum"))
        ]
        new_ab_int = max(existing_ab_ids, default=-1) + 1
        kfu_abstract_num_id = str(new_ab_int)

        ab_el = OxmlElement("w:abstractNum")
        ab_el.set(qn("w:abstractNumId"), kfu_abstract_num_id)

        nsid_mk = OxmlElement("w:nsid")
        nsid_mk.set(qn("w:val"), _KFU_DASH_NSID)
        ab_el.append(nsid_mk)

        ml_el = OxmlElement("w:multiLevelType")
        ml_el.set(qn("w:val"), "singleLevel")
        ab_el.append(ml_el)

        lvl_el = OxmlElement("w:lvl")
        lvl_el.set(qn("w:ilvl"), "0")

        s_el = OxmlElement("w:start"); s_el.set(qn("w:val"), "1")
        lvl_el.append(s_el)

        nf_el = OxmlElement("w:numFmt"); nf_el.set(qn("w:val"), "bullet")
        lvl_el.append(nf_el)

        lt_el = OxmlElement("w:lvlText"); lt_el.set(qn("w:val"), "–")  # en-dash
        lvl_el.append(lt_el)

        jc_el = OxmlElement("w:lvlJc"); jc_el.set(qn("w:val"), "left")
        lvl_el.append(jc_el)

        pPr_lvl = OxmlElement("w:pPr")
        ind_lvl = OxmlElement("w:ind")
        ind_lvl.set(qn("w:left"), "906")
        ind_lvl.set(qn("w:hanging"), "198")
        pPr_lvl.append(ind_lvl)
        lvl_el.append(pPr_lvl)

        rPr_lvl = OxmlElement("w:rPr")
        rFonts_lvl = OxmlElement("w:rFonts")
        rFonts_lvl.set(qn("w:ascii"), FONT_NAME)
        rFonts_lvl.set(qn("w:hAnsi"), FONT_NAME)
        rFonts_lvl.set(qn("w:cs"), FONT_NAME)
        rPr_lvl.append(rFonts_lvl)
        sz_lvl = OxmlElement("w:sz"); sz_lvl.set(qn("w:val"), "28")
        rPr_lvl.append(sz_lvl)
        szCs_lvl = OxmlElement("w:szCs"); szCs_lvl.set(qn("w:val"), "28")
        rPr_lvl.append(szCs_lvl)
        lvl_el.append(rPr_lvl)

        ab_el.append(lvl_el)

        # Insert abstractNum before first w:num to maintain schema order
        first_num = numbering_el.find(qn("w:num"))
        if first_num is not None:
            first_num.addprevious(ab_el)
        else:
            numbering_el.append(ab_el)

    # ── 3. Create w:num referencing kfu_abstract_num_id ────────────────────
    existing_num_ids = [
        int(n.get(qn("w:numId"), "0"))
        for n in numbering_el.findall(qn("w:num"))
    ]
    new_num_id = str(max(existing_num_ids, default=0) + 1)

    new_num_el = OxmlElement("w:num")
    new_num_el.set(qn("w:numId"), new_num_id)
    ab_ref = OxmlElement("w:abstractNumId")
    ab_ref.set(qn("w:val"), kfu_abstract_num_id)
    new_num_el.append(ab_ref)
    numbering_el.append(new_num_el)

    return new_num_id


def _format_word_dash_list_item(paragraph, body_text: str, marker: str = "-") -> None:
    """
    Format *paragraph* as a KFU plain-text dash-list item.

    * Text = '{marker} {body_text}' (literal marker embedded in run text, no numPr)
    * firstLine indent = 708 twips (≈ 1.25 cm, KFU L1 standard)
    * Times New Roman 14 pt · 1.5 line spacing · justified

    ``marker`` defaults to '-' (hyphen, uniform KFU standard for all list items).
    """
    body_text = _lowercase_first(body_text)
    remove_paragraph_numbering(paragraph)
    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)

    replace_paragraph_text(paragraph, f"{marker} {body_text}")

    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = LINE_SPACING_BODY
    fmt.keep_together = False
    fmt.keep_with_next = False
    fmt.page_break_before = False
    fmt.widow_control = False
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _apply_list_firstline_xml(paragraph, 708)
    force_paragraph_xml_spacing(paragraph, line_rule="auto")

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False, all_caps=False)

# ── End KFU internal Word dash-list ──────────────────────────────────────────


def _format_cyrillic_list_item(paragraph, letter: str, body_text: str):
    """Format a paragraph as level-1 Cyrillic list item (firstLine=708 twips)."""
    body_text = _lowercase_first(body_text)
    remove_paragraph_numbering(paragraph)
    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)

    replace_paragraph_text(paragraph, f"{letter}) {body_text}")

    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = LINE_SPACING_BODY
    fmt.keep_together = False
    fmt.keep_with_next = False
    fmt.page_break_before = False
    fmt.widow_control = False
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _apply_list_firstline_xml(paragraph, 708)
    force_paragraph_xml_spacing(paragraph, line_rule="auto")

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False, all_caps=False)


def _format_level2_list_item(paragraph, number: int, body_text: str):
    """Format a paragraph as level-2 (1)/2)/3)) list item (left=1200 hanging=198)."""
    body_text = _lowercase_first(body_text)
    remove_paragraph_numbering(paragraph)
    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)

    replace_paragraph_text(paragraph, f"{number}) {body_text}")

    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = LINE_SPACING_BODY
    fmt.keep_together = False
    fmt.keep_with_next = False
    fmt.page_break_before = False
    fmt.widow_control = False
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    _apply_list_indent_xml(paragraph, left_twips=1200, hanging_twips=198)
    force_paragraph_xml_spacing(paragraph, line_rule="auto")

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False, all_caps=False)


def _format_dash_list_item(paragraph, body_text: str):
    """Format a paragraph as a KFU dash-list item with hyphen '-' marker.
    Used for manual/free-standing dash blocks and colon-triggered numeric-dot.
    Applies firstLine=708 indent.
    """
    _format_word_dash_list_item(paragraph, body_text, marker="-")


def _format_endash_list_item(paragraph, body_text: str):
    """Format a paragraph as a KFU dash-list item with hyphen '-' marker.
    Previously used en-dash '–' for Word-decimal blocks; now unified to
    hyphen for consistency with all other list types.
    Applies firstLine=708 indent.
    """
    _format_word_dash_list_item(paragraph, body_text, marker="-")
    # Mark for idempotent second-pass recognition
    _mark_kfu_list_type(paragraph, "endash")


def _format_colon_dash_item(paragraph, body_text: str):
    """Format a paragraph as a colon-triggered dash item with hyphen '-' marker.
    No explicit firstLine indent (inherits style default). Used when a colon
    lead-in opens a numeric-dot task list (e.g. '1. задача' after 'задачи:').
    """
    body_text = _lowercase_first(body_text)
    remove_paragraph_numbering(paragraph)
    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)

    replace_paragraph_text(paragraph, f"- {body_text}")

    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = LINE_SPACING_BODY
    fmt.keep_together = False
    fmt.keep_with_next = False
    fmt.page_break_before = False
    fmt.widow_control = False
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Explicitly set firstLine=708 so items get 1.25 cm paragraph indent
    # (not cleared/inherited — template Normal style may have no indent at all)
    _apply_list_firstline_xml(paragraph, 708)
    force_paragraph_xml_spacing(paragraph, line_rule="auto")

    # Mark for idempotent second-pass recognition
    _mark_kfu_list_type(paragraph, "colon")

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False, all_caps=False)


def _paragraph_blocks_list_conversion(paragraph, text: str) -> bool:
    """
    Hard veto for free-standing list conversion. Paragraphs that look like
    structural headings, table/figure captions, table continuation labels,
    source/note service lines, formulas or formula explanations must never be
    rewritten as list items, regardless of how they start.
    """
    try:
        style_name = (paragraph.style.name or "").strip().lower()
    except Exception:
        style_name = ""
    if style_name in {
        "heading 1", "heading 2", "heading 3",
        "заголовок 1", "заголовок 2", "заголовок 3",
    }:
        return True
    pPr = paragraph._element.pPr
    if pPr is not None and pPr.find(qn("w:outlineLvl")) is not None:
        return True
    if TABLE_NUM_RE.match(text) or FIG_RE.match(text) or FIG_SERVICE_LINE_RE.match(text):
        return True
    if is_table_continuation_text(text):
        return True
    if is_formula_paragraph_text(text):
        return True
    if is_formula_explanation_start(text):
        return True
    return False


def _normalize_plain_list_paragraphs(paragraphs: list):
    """
    Detect and reformat plain-text list items in a sequence of paragraphs.

    Modes supported:

    1) Colon-trigger (existing behaviour, unchanged):
       A paragraph ending with ':' (that is not itself a list item) opens a
       list context. Subsequent numeric / letter markers become methodical
       L1 letters (`а)/б)/в)…`); a nested `1)` directly after the first L1
       item becomes L2 numeric (`1) 2) 3)…`).

    2) Free-standing dash / bullet (existing):
       Two or more consecutive paragraphs starting with `-`, `–`, `—`, `•`,
       `·`, `●`, etc. become methodical L1 dash items (`– text`) with
       hanging indent. Singletons stay untouched.

    3) Free-standing letter (existing — extended by G4):
       Two or more consecutive `а) б) в)` paragraphs get L1 letter layout.
       Letters preserved as-is from the source. G4 EXTENSION: letter items
       whose body ends with `:` are still valid in this block (e.g.
       `а) первый пункт первого уровня:` / `б) второй пункт первого уровня:`
       / `в) третий пункт:`). The trailing `:` is preserved in the output.

    4) Letter-colon + nested numeric L2 (G1, new):
       A letter L1 item ending with `:` may be followed by `1) ...`,
       `2) ...` children, which become L2 numeric (`left=963 hanging=198`).
       This works whether the parent is the first letter of a block (still
       buffered) — in which case the first numeric child confirms the
       buffer — or already inside a confirmed letter block.

    Aborts: blank, guarded paragraph, or any non-matching paragraph resets
    every buffer / context so singletons and mixed blocks remain untouched.
    """
    in_list = False
    level1_counter = 0
    level2_counter = 0
    prev_was_level1 = False
    prev_was_level2 = False

    dash_buffer = None        # (paragraph, body_text) pending dash candidate
    in_dash_block = False
    letter_buffer = None      # (paragraph, letter, body_text) pending letter candidate
    in_letter_block = False
    # G1: nested L2 numeric context opened by a confirmed letter L1 item
    # whose body ends with ':'. Counter restarts at 1 for each new letter.
    in_letter_colon_l2_context = False
    letter_colon_l2_counter = 0

    def _reset_free_blocks() -> None:
        nonlocal dash_buffer, in_dash_block, letter_buffer, in_letter_block
        nonlocal in_letter_colon_l2_context, letter_colon_l2_counter
        dash_buffer = None
        in_dash_block = False
        letter_buffer = None
        in_letter_block = False
        in_letter_colon_l2_context = False
        letter_colon_l2_counter = 0

    for p in paragraphs:
        text = clean_spaces(p.text)
        if not text:
            in_list = False
            level1_counter = 0
            level2_counter = 0
            prev_was_level1 = False
            prev_was_level2 = False
            _reset_free_blocks()
            continue

        # ── Idempotency fast-path ────────────────────────────────────────────────
        # Paragraphs already converted by a prior pass carry a kfuListType marker.
        # Re-apply the correct formatter directly so the second pass (line ~6475)
        # doesn't over-write their special geometry.
        _kfu_type = _get_kfu_list_type(p)
        if _kfu_type == "colon":
            # Restore {} ind (format_body resets it to firstLine=709)
            body = text[2:].strip() if text.startswith('- ') else text
            _format_colon_dash_item(p, body)
            # Keep in_list state intact — colon items extend the colon sequence
            continue
        if _kfu_type == "endash":
            # Restore firstLine=708; marker is now '-' (uniform hyphen)
            body = text[2:].strip() if text.startswith('- ') else text
            _format_endash_list_item(p, body)
            # End any colon/free-block context
            in_list = False
            level1_counter = 0
            level2_counter = 0
            prev_was_level1 = False
            prev_was_level2 = False
            _reset_free_blocks()
            continue
        # ── End fast-path ────────────────────────────────────────────────────────

        if text.endswith(':') and not _is_level1_list_text(text):
            in_list = True
            level1_counter = 0
            level2_counter = 0
            prev_was_level1 = False
            prev_was_level2 = False
            _reset_free_blocks()
            continue

        if in_list:
            m_cyr       = _CYRILLIC_LETTER_LIST_RE.match(text)
            m_num_paren = _NUMERIC_PAREN_LIST_RE.match(text)
            m_num_dot   = _NUMERIC_DOT_LIST_RE.match(text)

            if m_cyr:
                body = m_cyr.group(2).strip()
                letter = _CYRILLIC_LIST_ALPHA[level1_counter] if level1_counter < len(_CYRILLIC_LIST_ALPHA) else m_cyr.group(1)
                _format_cyrillic_list_item(p, letter, body)
                level1_counter += 1
                level2_counter = 0
                prev_was_level1 = True
                prev_was_level2 = False
                continue

            if m_num_dot:
                # Numeric-dot (1. 2. 3.) after a colon lead-in → hyphen dash items
                # (no firstLine indent; inherit from style — matches KFU task-list style)
                body = m_num_dot.group(2).strip()
                _format_colon_dash_item(p, body)
                level1_counter += 1
                prev_was_level1 = False
                prev_was_level2 = False
                continue

            if m_num_paren:
                body = m_num_paren.group(2).strip()
                num = int(m_num_paren.group(1))

                if prev_was_level1 and num == 1:
                    level2_counter = 1
                    _format_level2_list_item(p, level2_counter, body)
                    prev_was_level1 = False
                    prev_was_level2 = True
                elif prev_was_level2:
                    level2_counter += 1
                    _format_level2_list_item(p, level2_counter, body)
                else:
                    letter_idx = level1_counter
                    letter = _CYRILLIC_LIST_ALPHA[letter_idx] if letter_idx < len(_CYRILLIC_LIST_ALPHA) else str(level1_counter + 1)
                    _format_cyrillic_list_item(p, letter, body)
                    level1_counter += 1
                    level2_counter = 0
                    prev_was_level1 = True
                    prev_was_level2 = False
                continue

            # Colon-mode saw a non-numeric/non-letter paragraph: close the
            # colon context and re-evaluate this same paragraph against
            # free-standing detection. This is what lets `- text` items
            # right after a `задачи:` lead-in be recognised by the dash
            # path instead of being silently consumed by the reset.
            in_list = False
            level1_counter = 0
            level2_counter = 0
            prev_was_level1 = False
            prev_was_level2 = False
            # fall through to free-standing detection

        # Free-standing modes — apply hard guards first so headings, captions,
        # source lines, formulas, and figure-service lines can never be
        # rewritten as list items.
        if _paragraph_blocks_list_conversion(p, text):
            _reset_free_blocks()
            continue

        m_dash_free = _DASH_BULLET_LIST_RE.match(text)
        if m_dash_free:
            body = m_dash_free.group(2).strip()
            # Different family pending? abort it.
            letter_buffer = None
            in_letter_block = False
            if in_dash_block:
                _format_dash_list_item(p, body)
            elif dash_buffer is None:
                dash_buffer = (p, body)
            else:
                first_p, first_body = dash_buffer
                _format_dash_list_item(first_p, first_body)
                _format_dash_list_item(p, body)
                dash_buffer = None
                in_dash_block = True
            continue

        m_letter_free = _CYRILLIC_LETTER_LIST_RE.match(text)
        if m_letter_free:
            letter = m_letter_free.group(1)
            body = m_letter_free.group(2).strip()
            dash_buffer = None
            in_dash_block = False
            if in_letter_block:
                _format_cyrillic_list_item(p, letter, body)
                # G1: a new letter sibling restarts the L2 numeric context
                # only when this letter's body also ends with `:` — otherwise
                # any pending L2 context is closed.
                if body.endswith(":"):
                    in_letter_colon_l2_context = True
                    letter_colon_l2_counter = 0
                else:
                    in_letter_colon_l2_context = False
                    letter_colon_l2_counter = 0
            elif letter_buffer is None:
                letter_buffer = (p, letter, body)
            else:
                first_p, first_letter, first_body = letter_buffer
                _format_cyrillic_list_item(first_p, first_letter, first_body)
                _format_cyrillic_list_item(p, letter, body)
                letter_buffer = None
                in_letter_block = True
                # If THIS confirming letter ends with `:`, the nested L2
                # numeric context opens for the next paragraph. Note that
                # the FIRST (buffered) letter may also have ended with `:`,
                # but that does not retroactively open L2 — by the time we
                # confirm here the next paragraph follows the second letter
                # and its `:` decides.
                if body.endswith(":"):
                    in_letter_colon_l2_context = True
                    letter_colon_l2_counter = 0
                else:
                    in_letter_colon_l2_context = False
                    letter_colon_l2_counter = 0
            continue

        # G1: numeric-paren paragraph (`1) ...`, `2) ...`) following a
        # letter L1 item whose body ended with `:`. Two activation paths:
        #   * letter_buffer is pending AND its body ends with `:` — the
        #     numeric child confirms the buffered letter as L1, becomes
        #     numeric L2 child 1, and opens an L2 context for further
        #     children.
        #   * in_letter_colon_l2_context already True — current paragraph
        #     becomes the next numeric L2 child.
        m_numeric_l2 = _NUMERIC_PAREN_LIST_RE.match(text)
        if m_numeric_l2 and (
            (letter_buffer is not None and letter_buffer[2].endswith(":"))
            or in_letter_colon_l2_context
        ):
            child_body = m_numeric_l2.group(2).strip()
            if letter_buffer is not None and letter_buffer[2].endswith(":"):
                first_p, first_letter, first_body = letter_buffer
                _format_cyrillic_list_item(first_p, first_letter, first_body)
                letter_buffer = None
                in_letter_block = True
                in_letter_colon_l2_context = True
                letter_colon_l2_counter = 0
            letter_colon_l2_counter += 1
            _format_level2_list_item(p, letter_colon_l2_counter, child_body)
            continue

        # Anything else breaks any pending free-standing block.
        _reset_free_blocks()


def normalize_plain_lists_in_document(document, body_start):
    """
    Normalise plain-text list items in the document body.

    Skips the bibliography block (between `СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ`
    and `ПРИЛОЖЕНИЯ`) and the appendices block (everything from `ПРИЛОЖЕНИЯ`
    to the end of the document). Numbered entries in references and
    appendix-local paragraphs must remain plain text.
    """
    in_ref = False
    in_appendix = False
    body_paras: list = []
    for idx, p in enumerate(document.paragraphs):
        if idx < (body_start or 0):
            continue
        t = clean_spaces(p.text)
        if is_appendix_heading_text(t):
            in_appendix = True
        if is_references_heading_text(t):
            in_ref = True
        if in_ref and is_appendix_heading_text(t):
            in_ref = False
        if in_ref or in_appendix:
            continue
        body_paras.append(p)

    _normalize_plain_list_paragraphs(body_paras)


# ── Word-decimal autonumbering normalization ──────────────────────────────────

def _is_convertible_word_decimal_item(paragraph) -> bool:
    """
    True when paragraph carries Word decimal-numbered list formatting at ilvl=0
    and is safe to convert to a methodical dash list item.
    """
    if not paragraph_has_numbering(paragraph):
        return False
    num_fmt = _get_num_fmt_for_paragraph(paragraph)
    if num_fmt not in NUMERIC_NUM_FMTS:
        return False
    pPr = paragraph._element.pPr
    if pPr is None:
        return False
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        return False
    ilvl_el = numPr.find(qn("w:ilvl"))
    ilvl = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
    if ilvl != 0:
        return False
    if paragraph_has_heading_style_or_outline(paragraph):
        return False
    return True


def _normalize_word_numbered_list_paragraphs(paragraphs: list) -> None:
    """
    Convert contiguous blocks of Word decimal numPr (ilvl=0) paragraphs to
    methodical dash items (– text, left=906 hanging=198).
    Singletons (block size == 1) are left unchanged.
    """
    block: list = []  # [(paragraph, body_text), ...]

    def _flush() -> None:
        if len(block) >= 2:
            for p, body_text in block:
                _format_endash_list_item(p, body_text)
        block.clear()

    for p in paragraphs:
        text = clean_spaces(p.text)
        if not text:
            _flush()
            continue
        if _paragraph_blocks_list_conversion(p, text):
            _flush()
            continue
        if not _is_convertible_word_decimal_item(p):
            _flush()
            continue
        body_text = BULLET_CHARS_RE.sub("", text).strip() or text
        block.append((p, body_text))

    _flush()


def normalize_word_numbered_lists_in_document(document, body_start) -> None:
    """
    Normalize Word decimal-numbered list blocks in the document body.
    Skips the bibliography block and appendices (same contract as
    normalize_plain_lists_in_document).
    """
    in_ref = False
    in_appendix = False
    body_paras: list = []
    for idx, p in enumerate(document.paragraphs):
        if idx < (body_start or 0):
            continue
        t = clean_spaces(p.text)
        if is_appendix_heading_text(t):
            in_appendix = True
        if is_references_heading_text(t):
            in_ref = True
        if in_ref and is_appendix_heading_text(t):
            in_ref = False
        if in_ref or in_appendix:
            continue
        body_paras.append(p)
    _normalize_word_numbered_list_paragraphs(body_paras)

# ── End Word-decimal autonumbering normalization ──────────────────────────────


def canonical_reference_subheading_text(text: str):
    t = clean_spaces(text)
    if not t:
        return None

    return REFERENCE_SUBHEADINGS_CANON.get(t.lower())


def canonical_numbered_reference_subheading_text(text: str):
    t = clean_spaces(text)
    match = REFERENCE_HEADING_NUMBER_RE.match(t)
    if not match:
        return None

    return _canonical_reference_heading_candidate(match.group(1))


def canonical_reference_block_heading_text(text: str):
    return canonical_numbered_reference_subheading_text(text) or _canonical_reference_heading_candidate(text)



# ===== Reference list case normalization =====
_REF_URL_RE = re.compile(r'https?://\S+', re.IGNORECASE)
_PLAIN_URL_RE = re.compile(r'https?://[^\s<>"\']+', re.IGNORECASE)
_REF_TOKEN_RE = re.compile(r'([A-Za-zА-ЯЁа-яё]+(?:[-–—][A-Za-zА-ЯЁа-яё]+)*)')
_REF_ACRONYM_KEEP = {
    'ФНС', 'РФ', 'РБК', 'ТТС', 'ЭДО', 'СЭД', 'СМК', 'ГОСТ', 'ИСО', 'ЕС', 'АО', 'ООО', 'ПАО',
    'ISO', 'IEC', 'IEEE', 'OECD', 'EU', 'USA', 'UK', 'UN', 'PDF', 'HTML', 'URL', 'DOI', 'ISBN',
    'CRM', 'ERP', 'API', 'XML', 'JSON', 'UPD', 'B2B', 'B2G', 'B2C', 'ID', 'IT', 'AI', 'FTS',
}
_REF_CANONICAL_TOKEN_MAP = {
    'EIDAS': 'eIDAS',
    'BUSINESSTAT': 'BusinesStat',
    'CONSULTANTPLUS': 'КонсультантПлюс',
    'КОНСУЛЬТАНТПЛЮС': 'КонсультантПлюс',
}

def _looks_like_shouting_reference(text: str) -> bool:
    letters = [ch for ch in text if ch.isalpha()]
    if len(letters) < 12:
        return False
    uppers = sum(1 for ch in letters if ch.isupper())
    return (uppers / len(letters)) >= 0.65

def _normalize_reference_token(token: str) -> str:
    if not token:
        return token

    upper = token.upper()
    if upper in _REF_CANONICAL_TOKEN_MAP:
        return _REF_CANONICAL_TOKEN_MAP[upper]

    # Сохраняем общеупотребимые аббревиатуры и короткие токены с цифрами
    if upper in _REF_ACRONYM_KEEP:
        return upper
    if any(ch.isdigit() for ch in token):
        return token
    if len(token) <= 3 and token.isupper():
        return upper

    # Полностью верхний регистр -> нормальный Title Case
    if token.isupper():
        if '-' in token or '–' in token or '—' in token:
            parts = re.split(r'([-–—])', token)
            return ''.join(_normalize_reference_token(part) if part not in '-–—' else part for part in parts)
        low = token.lower()
        return low[:1].upper() + low[1:]

    return token

def _normalize_reference_case_fragment(fragment: str) -> str:
    return _REF_TOKEN_RE.sub(lambda m: _normalize_reference_token(m.group(0)), fragment)

def smart_normalize_reference_line_case(text: str) -> str:
    clean = clean_spaces(text)
    if not clean:
        return clean

    m = re.match(r'^(\d+\.\s+)(.+)$', clean)
    prefix = ''
    body = clean
    if m:
        prefix, body = m.group(1), m.group(2)

    if not _looks_like_shouting_reference(body):
        return clean

    urls = []
    def _url_repl(match):
        urls.append(match.group(0).lower())
        return f'__REFURL{len(urls)-1}__'

    body = _REF_URL_RE.sub(_url_repl, body)
    body = _normalize_reference_case_fragment(body)

    for i, url in enumerate(urls):
        body = body.replace(f'__REFURL{i}__', url)

    return f'{prefix}{body}' if prefix else body


def normalize_reference_url_spacing(text: str) -> str:
    return re.sub(
        r"(https?://[^\s]+?)\s*\((дата\s+обращения)",
        r"\1 (\2",
        text,
        flags=re.IGNORECASE,
    )

def strip_leading_heading_garbage(text: str) -> str:
    t = clean_spaces(text)
    if not t:
        return t

    # Убираем только явные маркеры-мусор Word/копипасты,
    # но НЕ трогаем обычные -, –, —, потому что они могут быть
    # реальными маркерами списка в основном тексте.
    t = re.sub(r'^\s*[•·▪■◆►→◦●○]+\s*', '', t)

    t = clean_spaces(t)
    return t

def is_probable_body_list_item(paragraph, prev_paragraph=None, prev_kind=None):
    """
    Обычный список в тексте, который надо сохранить как список, а не
    превращать в heading и не лишать Word-автонумерации.

    Целевой кейс:
    - перед списком абзац заканчивается двоеточием;
    - сам пункт списка начинается со строчной буквы;
    - у абзаца есть Word-numbering (numPr).
    Также поддерживаем продолжение такого списка:
    - если предыдущий непустой абзац уже был body_list_item.
    """
    text = clean_spaces(paragraph.text)
    if not text:
        return False

    if not paragraph_has_numbering(paragraph):
        return False

    if paragraph_has_heading_style_or_outline(paragraph):
        return False

    low = text.lower()

    forbidden_prefixes = (
        "таблица",
        "табл.",
        "рисунок",
        "рис.",
        "источник:",
        "составлено по:",
        "рассчитано по:",
        "примечание:",
        "продолжение таблицы",
        "продолжение табл.",
    )
    if low.startswith(forbidden_prefixes):
        return False

    if parse_heading1(text) or parse_heading2(text) or parse_broken_heading2(text):
        return False

    if TABLE_NUM_RE.match(text) or FIG_RE.match(text) or is_table_continuation_text(text):
        return False

    if not is_probable_center_bold_heading(paragraph):
        return True

    # Берём первую буквенную букву, если она есть
    first_alpha = next((ch for ch in text if ch.isalpha()), "")
    if not first_alpha or not first_alpha.islower():
        return False

    # Продолжение уже начавшегося текстового списка
    if prev_kind == "body_list_item":
        return True

    # Первый элемент списка сразу после двоеточия
    if prev_paragraph is not None:
        prev_text = clean_spaces(prev_paragraph.text)
        if prev_text.endswith(":"):
            return True

    return False
    
def normalize_heading2_artifacts(paragraph):
    text = clean_spaces(paragraph.text)
    if not text:
        return

    m = HEADING2_ARTIFACT_RE.match(text)
    if not m:
        return

    num = m.group(1)
    title = clean_spaces(m.group(2))
    if not parse_heading2(f"{num} {title}"):
        return

    if not num.endswith("."):
        num += "."
    replace_paragraph_text(paragraph, f"{num} {title}")


def is_probable_center_bold_heading(paragraph):
    if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return False
    if not paragraph.runs:
        return False

    non_empty_runs = [r for r in paragraph.runs if r.text and r.text.strip()]
    if not non_empty_runs:
        return False

    bold_runs = sum(1 for r in non_empty_runs if r.bold)
    return bold_runs >= max(1, len(non_empty_runs) // 2)


def paragraph_has_heading_style_or_outline(paragraph):
    try:
        style_name = (paragraph.style.name or "").strip().lower()
    except Exception:
        style_name = ""

    if style_name in {"heading 1", "heading 2", "заголовок 1", "заголовок 2"}:
        return True

    pPr = paragraph._element.pPr
    if pPr is None:
        return False

    return pPr.find(qn("w:outlineLvl")) is not None


def paragraph_has_numbering(paragraph):
    pPr = paragraph._element.pPr
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def remove_paragraph_numbering(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is not None:
        pPr.remove(numPr)


_HEADING_STYLE_NUMBERING_TOKENS = {
    "heading1",
    "heading2",
    "heading3",
    "заголовок1",
    "заголовок2",
    "заголовок3",
}


def _normalise_style_token(value: str | None) -> str:
    return re.sub(r"[\s_\-]+", "", (value or "").strip().lower())


def _is_heading_style_for_numbering_cleanup(style) -> bool:
    style_id = _normalise_style_token(getattr(style, "style_id", ""))
    style_name = _normalise_style_token(getattr(style, "name", ""))
    return (
        style_id in _HEADING_STYLE_NUMBERING_TOKENS
        or style_name in _HEADING_STYLE_NUMBERING_TOKENS
    )


def clear_heading_style_numbering(document) -> int:
    """
    Some input DOCX files carry numbering in Heading style definitions.
    Removing only paragraph-level w:numPr is not enough for those documents.
    """
    cleared = 0
    for style in document.styles:
        if not _is_heading_style_for_numbering_cleanup(style):
            continue

        style_element = getattr(style, "element", None)
        if style_element is None:
            style_element = getattr(style, "_element", None)
        if style_element is None:
            continue

        pPr = style_element.find(qn("w:pPr"))
        if pPr is None:
            continue

        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            continue

        pPr.remove(numPr)
        cleared += 1

    return cleared


def remove_page_break_artifacts_from_paragraph(paragraph):
    paragraph.paragraph_format.page_break_before = False
    paragraph.paragraph_format.keep_with_next = False
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.widow_control = False

    for run in paragraph.runs:
        r = run._element

        # Удаляем явные разрывы страницы внутри runs
        for br in list(r.findall(qn("w:br"))):
            br_type = br.get(qn("w:type"))
            if br_type in (None, "page"):
                r.remove(br)

        # На всякий случай убираем lastRenderedPageBreak
        for lrp in list(r.findall(qn("w:lastRenderedPageBreak"))):
            r.remove(lrp)

def is_references_heading_text(text: str) -> bool:
    low = clean_spaces(text).lower()
    return low in {
        "список использованных источников",
        "список использованной литературы",
    }


def is_appendix_heading_text(text: str) -> bool:
    low = clean_spaces(text).lower()
    return low in {"приложение", "приложения"}


APPENDIX_START_LABEL_RE = re.compile(
    r"^\s*приложение\s*(?P<num>\d{1,3}|[A-Za-zА-ЯЁ])\s*$",
    re.IGNORECASE,
)


def normalize_appendix_start_label_text(text: str) -> str | None:
    match = APPENDIX_START_LABEL_RE.match(clean_spaces(text))
    if not match:
        return None
    return f"ПРИЛОЖЕНИЕ {match.group('num').upper()}"


def is_appendix_start_label_like(text: str) -> bool:
    t = clean_spaces(text)
    return t.lower() != "приложения" and bool(re.match(r"^\s*приложение\s*\S+\s*$", t, re.IGNORECASE))


def _clear_page_break_before(paragraph):
    paragraph.paragraph_format.page_break_before = False
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return
    for page_break in list(p_pr.findall(qn("w:pageBreakBefore"))):
        p_pr.remove(page_break)


def format_appendix_start_label(paragraph, *, start_new_page=True):
    normalized = normalize_appendix_start_label_text(paragraph.text)
    if normalized and clean_spaces(paragraph.text) != normalized:
        replace_paragraph_text(paragraph, normalized)

    remove_page_break_artifacts_from_paragraph(paragraph)
    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)
    remove_paragraph_numbering(paragraph)
    force_paragraph_xml_spacing(paragraph, line_rule="auto")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = LINE_SPACING_BODY
    if start_new_page:
        paragraph.paragraph_format.page_break_before = True
    else:
        _clear_page_break_before(paragraph)
    paragraph.paragraph_format.keep_with_next = True

    p_pr = paragraph._element.get_or_add_pPr()
    for old_ind in list(p_pr.findall(qn("w:ind"))):
        p_pr.remove(old_ind)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "0")
    p_pr.append(ind)

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False, all_caps=False)


def is_appendix_continuation_label_text(text: str) -> bool:
    return bool(re.match(r"^\s*продолжение\s+приложения?\s+\S+\s*$", clean_spaces(text), re.IGNORECASE))


def is_table_like_appendix_title_text(text: str) -> bool:
    return bool(re.match(r"^\s*таблица\s+\S+", clean_spaces(text), re.IGNORECASE))


def is_obvious_appendix_body_paragraph(text: str) -> bool:
    t = clean_spaces(text)
    if len(t) > 100:
        return True
    if re.search(r"https?://|www\.|\[[^\]]*\d", t, re.IGNORECASE):
        return True
    if len(t.split()) > 14 and re.search(r"[.!?]\s+\S", t):
        return True
    return False


def is_appendix_title_candidate_text(text: str, *, has_later_content: bool) -> bool:
    t = clean_spaces(text)
    if not t or not has_later_content:
        return False
    if len(t) > 100:
        return False
    if normalize_appendix_start_label_text(t) or is_appendix_continuation_label_text(t):
        return False
    if is_table_like_appendix_title_text(t):
        return True
    return not is_obvious_appendix_body_paragraph(t)


def format_appendix_title(paragraph):
    text = strip_single_terminal_period(paragraph.text)
    if text != clean_spaces(paragraph.text):
        replace_paragraph_text(paragraph, text)

    remove_page_break_artifacts_from_paragraph(paragraph)
    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)
    remove_paragraph_numbering(paragraph)
    hard_reset_paragraph_format(paragraph, first_line_indent_cm=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    _clear_page_break_before(paragraph)

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False, all_caps=False)


def _has_later_appendix_content(children, start_idx: int) -> bool:
    idx = start_idx
    while idx < len(children):
        child = children[idx]
        if child.tag == qn("w:tbl"):
            return True
        if child.tag == qn("w:p"):
            if not _is_empty_paragraph_xml(child):
                return True
            idx += 1
            continue
        return False
    return False


def _appendix_label_followed_by_title_candidate(children, label_idx: int, paragraph_lookup, body_start) -> bool:
    next_idx = label_idx + 1
    while next_idx < len(children) and children[next_idx].tag == qn("w:p") and _is_empty_paragraph_xml(children[next_idx]):
        next_idx += 1
    if next_idx >= len(children) or children[next_idx].tag != qn("w:p"):
        return False

    paragraph_info = paragraph_lookup.get(children[next_idx])
    if paragraph_info is None:
        return False
    paragraph_idx, paragraph = paragraph_info
    if paragraph_idx < body_start:
        return False
    return is_appendix_title_candidate_text(
        paragraph.text,
        has_later_content=_has_later_appendix_content(children, next_idx + 1),
    )


def _merge_split_table_like_appendix_title(document, paragraph):
    if not is_table_like_appendix_title_text(paragraph.text):
        return

    body = document.element.body
    children = list(body)
    try:
        idx = children.index(paragraph._p)
    except ValueError:
        return

    blanks = []
    next_idx = idx + 1
    while next_idx < len(children) and children[next_idx].tag == qn("w:p") and _is_empty_paragraph_xml(children[next_idx]):
        blanks.append(children[next_idx])
        next_idx += 1

    if next_idx >= len(children) or children[next_idx].tag != qn("w:p"):
        return

    next_text = _paragraph_text_from_xml(children[next_idx])
    if (
        not next_text
        or normalize_appendix_start_label_text(next_text)
        or is_appendix_continuation_label_text(next_text)
        or is_table_like_appendix_title_text(next_text)
    ):
        return

    merged = clean_spaces(f"{clean_spaces(paragraph.text)} {strip_single_terminal_period(next_text)}")
    if len(merged) > 100:
        return

    replace_paragraph_text(paragraph, merged)
    for blank in blanks:
        body.remove(blank)
    body.remove(children[next_idx])


def _ensure_single_blank_after_paragraph(document, paragraph):
    body = document.element.body
    children = list(body)
    try:
        idx = children.index(paragraph._p)
    except ValueError:
        return

    blanks = []
    next_idx = idx + 1
    while next_idx < len(children) and children[next_idx].tag == qn("w:p") and _is_empty_paragraph_xml(children[next_idx]):
        blanks.append(children[next_idx])
        next_idx += 1

    if not blanks:
        blank = insert_paragraph_after(paragraph, "")
        format_empty_paragraph(blank)
        return

    paragraph_lookup = _paragraph_lookup(document)
    first_blank = paragraph_lookup.get(blanks[0])
    if first_blank is not None:
        _idx, blank_paragraph = first_blank
        format_empty_paragraph(blank_paragraph)

    for blank in blanks[1:]:
        body.remove(blank)


def normalize_appendix_titles(document, body_start):
    in_appendices = False
    children = list(document.element.body)
    paragraph_lookup = _paragraph_lookup(document)
    idx = 0

    while idx < len(children):
        child = children[idx]
        if child.tag != qn("w:p"):
            idx += 1
            continue

        paragraph_info = paragraph_lookup.get(child)
        if paragraph_info is None:
            idx += 1
            continue
        paragraph_idx, paragraph = paragraph_info
        if paragraph_idx < body_start:
            idx += 1
            continue

        text = _paragraph_text_from_xml(child)
        if text.lower() == "приложения":
            in_appendices = True
            idx += 1
            continue
        if not in_appendices or not normalize_appendix_start_label_text(text):
            idx += 1
            continue

        next_idx = idx + 1
        blanks_before_title = []
        while next_idx < len(children) and children[next_idx].tag == qn("w:p") and _is_empty_paragraph_xml(children[next_idx]):
            blanks_before_title.append(children[next_idx])
            next_idx += 1

        if (
            next_idx < len(children)
            and children[next_idx].tag == qn("w:p")
            and not _is_empty_paragraph_xml(children[next_idx])
        ):
            title_info = paragraph_lookup.get(children[next_idx])
            if title_info is not None:
                title_idx, title_paragraph = title_info
                if title_idx >= body_start and is_appendix_title_candidate_text(
                    title_paragraph.text,
                    has_later_content=_has_later_appendix_content(children, next_idx + 1),
                ):
                    for blank in blanks_before_title:
                        document.element.body.remove(blank)
                    _merge_split_table_like_appendix_title(document, title_paragraph)
                    format_appendix_title(title_paragraph)
                    _ensure_single_blank_after_paragraph(document, title_paragraph)
                    children = list(document.element.body)
                    paragraph_lookup = _paragraph_lookup(document)
                    idx = next_idx + 1
                    continue

        idx += 1


def _paragraph_text_from_xml(p_xml) -> str:
    return clean_spaces("".join(t.text or "" for t in p_xml.findall(".//" + qn("w:t"))))


def _is_empty_paragraph_xml(p_xml) -> bool:
    return not _paragraph_text_from_xml(p_xml) and not p_xml.findall(".//" + qn("w:drawing"))


def _paragraph_lookup(document):
    return {paragraph._p: (idx, paragraph) for idx, paragraph in enumerate(document.paragraphs)}


def normalize_appendix_start_labels(document, body_start):
    in_appendices = False
    first_appendix_label_seen = False
    for idx, paragraph in enumerate(document.paragraphs):
        if idx < body_start:
            continue
        text = clean_spaces(paragraph.text)
        if not text:
            continue
        if text.lower() == "приложения":
            in_appendices = True
            continue
        if in_appendices and normalize_appendix_start_label_text(text):
            format_appendix_start_label(
                paragraph,
                start_new_page=first_appendix_label_seen,
            )
            first_appendix_label_seen = True


def normalize_appendix_local_table_titles(document, body_start):
    paragraph_lookup = _paragraph_lookup(document)
    children = list(document.element.body)
    in_appendices = False

    for idx, child in enumerate(children):
        if child.tag == qn("w:p"):
            paragraph_info = paragraph_lookup.get(child)
            if paragraph_info is None:
                continue
            paragraph_idx, _paragraph = paragraph_info
            if paragraph_idx < body_start:
                continue
            if _paragraph_text_from_xml(child).lower() == "приложения":
                in_appendices = True
            continue

        if child.tag != qn("w:tbl") or not in_appendices:
            continue

        for prev in reversed(children[:idx]):
            if prev.tag != qn("w:p"):
                continue
            if _is_empty_paragraph_xml(prev):
                continue

            paragraph_info = paragraph_lookup.get(prev)
            if paragraph_info is None:
                break
            paragraph_idx, paragraph = paragraph_info
            if paragraph_idx < body_start:
                break

            text = _paragraph_text_from_xml(prev)
            if (
                text.lower() == "приложения"
                or is_appendix_start_label_like(text)
                or TABLE_NUM_RE.match(text)
                or is_table_continuation_text(text)
                or FIG_RE.match(text)
                or re.match(r"^\s*(источник|составлено по|рассчитано по|примечание)\s*:", text, re.IGNORECASE)
            ):
                break

            format_table_title(paragraph)
            _clear_page_break_before(paragraph)
            paragraph.paragraph_format.keep_with_next = True
            break


def remove_empty_paragraphs_between_appendices_heading_and_first_label(document, body_start):
    paragraph_lookup = _paragraph_lookup(document)
    children = list(document.element.body)

    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue

        paragraph_info = paragraph_lookup.get(child)
        if paragraph_info is None:
            continue
        paragraph_idx, _paragraph = paragraph_info
        if paragraph_idx < body_start:
            continue
        if _paragraph_text_from_xml(child).lower() != "приложения":
            continue

        blanks = []
        next_idx = idx + 1
        while next_idx < len(children) and children[next_idx].tag == qn("w:p") and _is_empty_paragraph_xml(children[next_idx]):
            blanks.append(children[next_idx])
            next_idx += 1

        if (
            blanks
            and next_idx < len(children)
            and children[next_idx].tag == qn("w:p")
            and is_appendix_start_label_like(_paragraph_text_from_xml(children[next_idx]))
        ):
            body = document.element.body
            for blank in blanks:
                body.remove(blank)
        return


def remove_empty_paragraphs_after_appendix_labels(document, body_start):
    paragraph_lookup = _paragraph_lookup(document)
    children = list(document.element.body)
    in_appendices = False

    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue

        paragraph_info = paragraph_lookup.get(child)
        if paragraph_info is None:
            continue
        paragraph_idx, _paragraph = paragraph_info
        if paragraph_idx < body_start:
            continue

        text = _paragraph_text_from_xml(child)
        if text.lower() == "приложения":
            in_appendices = True
            continue
        if not in_appendices or not is_appendix_start_label_like(text):
            continue
        if _appendix_label_followed_by_title_candidate(children, idx, paragraph_lookup, body_start):
            continue

        next_idx = idx + 1
        if next_idx < len(children) and children[next_idx].tag == qn("w:p") and _is_empty_paragraph_xml(children[next_idx]):
            continue

        insert_paragraph_after(_paragraph, "")


def format_empty_paragraphs_in_body(document, body_start):
    for idx, paragraph in enumerate(document.paragraphs):
        if idx < body_start:
            continue
        if is_empty_paragraph(paragraph):
            format_empty_paragraph(paragraph)

def format_body(paragraph, preserve_numbering=False):
    if preserve_numbering:
        # Для реальных Word-списков нельзя сбрасывать стиль в Normal
        # и нельзя трогать numbering/layout списка.
        remove_page_break_artifacts_from_paragraph(paragraph)
        force_paragraph_xml_spacing(paragraph, line_rule="auto")

        fmt = paragraph.paragraph_format
        fmt.space_before = Pt(0)
        fmt.space_after = Pt(0)
        fmt.line_spacing = LINE_SPACING_BODY
        fmt.keep_together = False
        fmt.keep_with_next = False
        fmt.page_break_before = False
        fmt.widow_control = False

        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Even numbered paragraphs need correct indent (body indent, not list hanging)
        _pPr_n = paragraph._element.find(qn("w:pPr"))
        if _pPr_n is not None:
            for _oi in list(_pPr_n.findall(qn("w:ind"))):
                _pPr_n.remove(_oi)
        _pPr_n2 = paragraph._element.get_or_add_pPr()
        _ind_n = OxmlElement("w:ind")
        _ind_n.set(qn("w:left"), "0")
        _ind_n.set(qn("w:firstLine"), "709")
        _pPr_n2.append(_ind_n)

        for run in paragraph.runs:
            set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, all_caps=False)
        return

    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)
    remove_paragraph_numbering(paragraph)

    hard_reset_paragraph_format(paragraph, first_line_indent_cm=FIRST_LINE_INDENT_CM)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, all_caps=False)


def _append_hyperlink_run(paragraph, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rPr.append(rFonts)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rPr.append(underline)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(BODY_FONT_SIZE_PT * 2)))
    rPr.append(size)

    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(int(BODY_FONT_SIZE_PT * 2)))
    rPr.append(size_cs)

    run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = url
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_plain_url_hyperlinks(paragraph) -> bool:
    text = paragraph.text
    matches = list(_PLAIN_URL_RE.finditer(text))
    if not matches:
        return False

    p = paragraph._element
    for child in list(p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p.remove(child)

    pos = 0
    for match in matches:
        url = match.group(0)
        trailing = ""
        while url and url[-1] in ".,;":
            trailing = url[-1] + trailing
            url = url[:-1]

        start = match.start()
        end = match.start() + len(url)
        if start > pos:
            paragraph.add_run(text[pos:start])

        _append_hyperlink_run(paragraph, url)
        if trailing:
            paragraph.add_run(trailing)
        pos = match.end()

    if pos < len(text):
        paragraph.add_run(text[pos:])

    return True


def format_reference_entry(paragraph) -> None:
    text = normalize_reference_url_spacing(clean_spaces(paragraph.text))
    if text != paragraph.text:
        replace_paragraph_text(paragraph, text)

    format_body(paragraph)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.first_line_indent = Cm(FIRST_LINE_INDENT_CM)

    pPr = paragraph._element.get_or_add_pPr()
    for old_tabs in list(pPr.findall(qn("w:tabs"))):
        pPr.remove(old_tabs)
    for old_ind in list(pPr.findall(qn("w:ind"))):
        pPr.remove(old_ind)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:firstLine"), "709")
    pPr.append(ind)

    add_plain_url_hyperlinks(paragraph)
    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, all_caps=False)


NUMERIC_NUM_FMTS = {"decimal", "decimalZero", "ordinal", "decimalEnclosedParen", "decimalEnclosedCircle"}

# numFmt values for letter / roman lists — preserved via
# format_body(preserve_numbering=True) rather than converted to KFU dash.
LETTER_NUM_FMTS = {"lowerLetter", "upperLetter", "lowerRoman", "upperRoman",
                   "aiueo", "iroha", "arabicAbjad", "arabicAlpha"}

BULLET_CHARS_RE = re.compile(r'^[•·▪■◆►→◦●○\u2013\u2014\-]+\s*')


def _get_num_fmt_for_paragraph(paragraph):
    """Returns numFmt val string like 'decimal', 'bullet', 'lowerLetter', etc. or None."""
    try:
        pPr = paragraph._element.pPr
        if pPr is None:
            return None
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return None
        numId_el = numPr.find(qn("w:numId"))
        ilvl_el = numPr.find(qn("w:ilvl"))
        if numId_el is None:
            return None
        num_id_val = numId_el.get(qn("w:val"))
        ilvl_val = ilvl_el.get(qn("w:val"), "0") if ilvl_el is not None else "0"

        numbering_part = paragraph.part.numbering_part
        if numbering_part is None:
            return None

        numbering_el = numbering_part._element

        num_el = None
        for n in numbering_el.findall(qn("w:num")):
            if n.get(qn("w:numId")) == num_id_val:
                num_el = n
                break
        if num_el is None:
            return None

        abstract_num_id_el = num_el.find(qn("w:abstractNumId"))
        if abstract_num_id_el is None:
            return None
        abstract_num_id_val = abstract_num_id_el.get(qn("w:val"))

        abstract_num_el = None
        for an in numbering_el.findall(qn("w:abstractNum")):
            if an.get(qn("w:abstractNumId")) == abstract_num_id_val:
                abstract_num_el = an
                break
        if abstract_num_el is None:
            return None

        for lvl in abstract_num_el.findall(qn("w:lvl")):
            if lvl.get(qn("w:ilvl")) == ilvl_val:
                num_fmt_el = lvl.find(qn("w:numFmt"))
                if num_fmt_el is not None:
                    return num_fmt_el.get(qn("w:val"))
                break
        return None
    except Exception:
        return None


def format_body_list_item(paragraph):
    """
    Format a list-item paragraph.

    * Numeric Word lists (decimal, ordinal, ...): preserved via
      format_body(preserve_numbering=True).
    * Letter / roman Word lists (lowerLetter, upperLetter, lowerRoman, ...):
      also preserved via format_body(preserve_numbering=True) so that
      Cyrillic-letter autonumbering is not stripped.
    * Everything else (bullet numPr or no numPr): converted to plain-text
      '– body_text' with firstLine=708 via _format_word_dash_list_item.
    """
    num_fmt = _get_num_fmt_for_paragraph(paragraph)
    is_numeric = num_fmt in NUMERIC_NUM_FMTS
    is_letter  = num_fmt in LETTER_NUM_FMTS

    if is_numeric or is_letter:
        format_body(paragraph, preserve_numbering=True)
        return

    # Bullet numPr (or no numPr) -> KFU real Word dash-list
    remove_page_break_artifacts_from_paragraph(paragraph)
    text = clean_spaces(paragraph.text)
    body_text = BULLET_CHARS_RE.sub('', text)
    body_text = clean_spaces(body_text) or text
    _format_word_dash_list_item(paragraph, body_text)


def format_heading1(paragraph):
    remove_page_break_artifacts_from_paragraph(paragraph)
    remove_paragraph_numbering(paragraph)

    text = clean_spaces(paragraph.text)
    if text:
        replace_paragraph_text(paragraph, text.upper())

    set_paragraph_style_safe(paragraph, "Heading 1", "Заголовок 1")
    clear_paragraph_outline_level(paragraph)
    hard_reset_paragraph_format(paragraph, first_line_indent_cm=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=True, italic=False, all_caps=False)

def format_heading2(paragraph):
    remove_page_break_artifacts_from_paragraph(paragraph)
    remove_paragraph_numbering(paragraph)

    set_paragraph_style_safe(paragraph, "Heading 2", "Заголовок 2")
    hard_reset_paragraph_format(paragraph, first_line_indent_cm=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=True, all_caps=False)

def format_table_caption(paragraph):
    text = clean_spaces(paragraph.text)
    m = TABLE_NUM_RE.match(text)
    if m:
        number = m.group(1)
        replace_paragraph_text(paragraph, f"Таблица {number}")

    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)
    remove_paragraph_numbering(paragraph)

    hard_reset_paragraph_format(paragraph, first_line_indent_cm=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False, all_caps=False)


def strip_single_terminal_period(text: str) -> str:
    text = clean_spaces(text)
    if text.endswith(".") and not text.endswith(".."):
        return text[:-1]
    return text


def format_table_title(paragraph):
    text = strip_single_terminal_period(paragraph.text)
    if text != clean_spaces(paragraph.text):
        replace_paragraph_text(paragraph, text)

    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)
    remove_paragraph_numbering(paragraph)

    hard_reset_paragraph_format(paragraph, first_line_indent_cm=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, all_caps=False)


def format_source_line(paragraph):
    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)
    remove_paragraph_numbering(paragraph)

    hard_reset_paragraph_format(paragraph, first_line_indent_cm=FIRST_LINE_INDENT_CM)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for run in paragraph.runs:
        set_run_font(run, size_pt=TABLE_FONT_SIZE_PT, bold=False, italic=False, all_caps=False)

def format_reference_subheading(paragraph):
    # Обязательно делаем обычным абзацем, а не заголовком
    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)
    remove_paragraph_numbering(paragraph)

    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0)
    paragraph.paragraph_format.page_break_before = False
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = False
    paragraph.paragraph_format.widow_control = False

    # Remove any firstLine/hanging indent via direct XML
    _pPr = paragraph._element.get_or_add_pPr()
    for _oi in list(_pPr.findall(qn("w:ind"))):
        _pPr.remove(_oi)
    _ind = OxmlElement("w:ind")
    _ind.set(qn("w:left"), "0")
    _ind.set(qn("w:firstLine"), "0")
    _pPr.append(_ind)

    # 1.5x line spacing
    force_paragraph_xml_spacing(paragraph, line_rule="auto")

    for run in paragraph.runs:
        set_run_font(
            run,
            size_pt=BODY_FONT_SIZE_PT,
            bold=True,
            italic=False,
            all_caps=False,
        )

def format_figure_caption(paragraph):
    hard_reset_paragraph_format(paragraph, first_line_indent_cm=FIRST_LINE_INDENT_CM)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.keep_together = True
    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, all_caps=False)


def set_cell_border(cell, color="000000", size="4", space="0"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = tcBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tcBorders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), space)
        element.set(qn("w:color"), color)

def clear_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is not None:
        tcPr.remove(tcBorders)


TABLE_BORDER_EDGES = ("top", "left", "bottom", "right", "insideH", "insideV")


def _get_or_add_tbl_pr(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    return tblPr


def _set_border_element_single(element, *, color="000000", size="4", space="0") -> None:
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), size)
    element.set(qn("w:space"), space)
    element.set(qn("w:color"), color)


def _ensure_tbl_borders_single(table, *, color="000000", size="4", space="0") -> None:
    tblPr = _get_or_add_tbl_pr(table)
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for edge in TABLE_BORDER_EDGES:
        element = tblBorders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tblBorders.append(element)
        _set_border_element_single(element, color=color, size=size, space=space)


def _normalize_existing_cell_borders_single(table, *, color="000000", size="4", space="0") -> None:
    for tc in table._tbl.findall(".//" + qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            continue
        tcBorders = tcPr.find(qn("w:tcBorders"))
        if tcBorders is None:
            continue
        for edge in TABLE_BORDER_EDGES:
            element = tcBorders.find(qn(f"w:{edge}"))
            if element is not None:
                _set_border_element_single(element, color=color, size=size, space=space)


def _normalize_table_borders_preserve_geometry(table, *, color="000000", size="4", space="0") -> None:
    """
    Border-only cleanup for preserve-mode tables.

    It removes cell spacing and normalizes visible border style, but intentionally
    avoids width/layout/topology nodes such as tblGrid, gridCol, tcW, gridSpan,
    vMerge, row heights, and tblCellMar.
    """
    tblPr = _get_or_add_tbl_pr(table)

    tblCellSpacing = tblPr.find(qn("w:tblCellSpacing"))
    if tblCellSpacing is not None:
        tblPr.remove(tblCellSpacing)

    for row in table.rows:
        trPr = row._tr.trPr
        if trPr is None:
            continue
        trCellSpacing = trPr.find(qn("w:tblCellSpacing"))
        if trCellSpacing is not None:
            trPr.remove(trCellSpacing)

    _ensure_tbl_borders_single(table, color=color, size=size, space=space)
    _normalize_existing_cell_borders_single(table, color=color, size=size, space=space)


def _safe_formatter_table_geometry_policy(table) -> str:
    """
    Classify table geometry before Phase 1 width/border normalization.

    Phase 1 fails closed here: if classification cannot be evaluated, existing
    geometry is preserved instead of rewritten.
    """
    try:
        from .table_continuation import classify_table_geometry_policy

        policy = classify_table_geometry_policy(table._tbl)
    except Exception:
        return "preserve_geometry"

    if policy == "simple":
        return "simple"
    if policy in {"preserve_geometry", "unsafe_no_split"}:
        return policy
    return "preserve_geometry"


def _safe_formatter_preserve_table_geometry(table) -> bool:
    return _safe_formatter_table_geometry_policy(table) != "simple"


def force_table_outer_borders_single(table, color="000000", size="4", space="0"):
    """
    Жестко задает одинарные границы таблицы и удаляет те table-level XML-узлы,
    которые в Word for Mac могут давать визуальный эффект двойного контура.
    """
    if _safe_formatter_preserve_table_geometry(table):
        return

    tblPr = _get_or_add_tbl_pr(table)

    # Убираем style/look-метаданные таблицы
    for tag in (
        "w:tblStyle",
        "w:tblLook",
        "w:tblStyleRowBandSize",
        "w:tblStyleColBandSize",
        "w:tblInd",
    ):
        node = tblPr.find(qn(tag))
        if node is not None:
            tblPr.remove(node)

    # tblCellSpacing вообще не должен оставаться в XML.
    # Даже при w:w="0" Word может продолжать рендерить контур так,
    # будто между ячейками/внешней рамкой есть зазор.
    tblCellSpacing = tblPr.find(qn("w:tblCellSpacing"))
    if tblCellSpacing is not None:
        tblPr.remove(tblCellSpacing)
    # Жестко фиксируем layout таблицы.
    # Без этого Word может автоподбирать ширины столбцов и визуально
    # давать "двойные" линии из-за дробной геометрии рендера.
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)

    tblLayout.set(qn("w:type"), "fixed")
    # И дополнительно убираем tblCellMar,
    # чтобы не было лишнего визуального "внутреннего отступа контура" в Word.
    node = tblPr.find(qn("w:tblCellMar"))
    if node is not None:
        tblPr.remove(node)

    _ensure_tbl_borders_single(table, color=color, size=size, space=space)

    # Убираем row-level overrides и row-level spacing, если они приехали из исходника.
    for row in table.rows:
        trPr = row._tr.trPr
        if trPr is not None:
            tblPrEx = trPr.find(qn("w:tblPrEx"))
            if tblPrEx is not None:
                trPr.remove(tblPrEx)

            trCellSpacing = trPr.find(qn("w:tblCellSpacing"))
            if trCellSpacing is not None:
                trPr.remove(trCellSpacing)

    # У ячеек оставляем без собственных borders/margins,
    # чтобы источник истины был только один — tblBorders.
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            tcBorders = tcPr.find(qn("w:tcBorders"))
            if tcBorders is not None:
                tcPr.remove(tcBorders)



def apply_table_borders(table):
    if _safe_formatter_preserve_table_geometry(table):
        _normalize_table_borders_preserve_geometry(table, size="4")
        return

    # Один источник истины для рамок — tblBorders на уровне таблицы.
    force_table_outer_borders_single(table, size="4")

    for row in table.rows:
        for cell in row.cells:
            clear_cell_borders(cell)
            
def force_zero_indent_in_table_paragraph(paragraph):
    """
    Жестко сбрасывает любые абзацные отступы внутри таблицы:
    первая строка, левый/правый отступ, а также XML-атрибуты w:ind.
    Это узкая функция только для абзацев внутри ячеек таблицы.
    """
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)

    pPr = paragraph._element.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)

    # Полностью обнуляем ключевые виды отступов Word
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:right"), "0")
    ind.set(qn("w:start"), "0")
    ind.set(qn("w:end"), "0")
    ind.set(qn("w:hanging"), "0")


def force_table_run_plain(run):
    """
    Жестко убирает жирность у run внутри таблицы.
    Обычного run.bold = False иногда недостаточно, поэтому
    дополнительно прибиваем XML-свойства жирности.
    """
    set_run_font(
        run,
        size_pt=TABLE_FONT_SIZE_PT,
        bold=False,
        italic=False,
        all_caps=False,
    )

    run.bold = False
    run.font.bold = False

    rPr = run._element.get_or_add_rPr()

    for tag in ("w:b", "w:bCs"):
        node = rPr.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            rPr.append(node)
        node.set(qn("w:val"), "0")

TABLE_PURE_NUMBER_RE = re.compile(r"^\s*\d+(?:[.,]\d+)?\s*$")


def _normalize_table_numeric_text(text: str) -> str | None:
    """
    Возвращает нормализованное число как строку или None, если это не "чисто число".

    Поддерживает:
    - 10
    - 12,3
    - 12.3
    - 16 000
    - 16 000,5

    Не считает числом:
    - 8-10
    - 7–12
    - 8—15
    - 12%
    - 12 тыс.
    - текст с буквами
    """
    t = clean_spaces(text)
    if not t:
        return None

    # Убираем обычные пробелы между разрядами: 16 000 -> 16000
    t = re.sub(r"\s+", "", t)

    # Диапазоны / интервалы / "не чистое число" сразу отсекаем
    if any(ch in t for ch in ("-", "–", "—")):
        return None

    if not re.fullmatch(r"\d+(?:[.,]\d+)?", t):
        return None

    # Для унификации внутри кода храним с точкой
    return t.replace(",", ".")


def _get_table_cell_text(cell) -> str:
    parts = []
    for p in cell.paragraphs:
        txt = clean_spaces(p.text)
        if txt:
            parts.append(txt)
    return " ".join(parts).strip()


def _table_cell_is_pure_number(cell) -> bool:
    return _normalize_table_numeric_text(_get_table_cell_text(cell)) is not None


def _get_table_numeric_column_scales(table) -> dict[int, int]:
    """
    Для каждого столбца определяет минимально нужное число знаков после запятой,
    которое надо применить к целым числам, если в столбце есть дробные значения.

    Логика:
    - если в столбце только целые -> scale = 0
    - если есть 12,3 -> scale = 1
    - если есть 12,34 -> scale = 2
    Берём максимум фактически встреченных дробных знаков в "чистых числах".
    """
    scales: dict[int, int] = {}

    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            raw = _normalize_table_numeric_text(_get_table_cell_text(cell))
            if raw is None:
                continue

            if "." in raw:
                frac_len = len(raw.split(".", 1)[1])
            else:
                frac_len = 0

            scales[col_idx] = max(scales.get(col_idx, 0), frac_len)

    return scales


def _format_table_number_for_column(text: str, scale: int) -> str | None:
    """
    Приводит число к нужному виду для конкретного столбца.
    Возвращает строку с запятой как десятичным разделителем.
    """
    raw = _normalize_table_numeric_text(text)
    if raw is None:
        return None

    if "." in raw:
        int_part, frac_part = raw.split(".", 1)
    else:
        int_part, frac_part = raw, ""

    if scale <= 0:
        # если столбец целочисленный — убираем дробную часть только если она нулевая/отсутствует
        if frac_part and any(ch != "0" for ch in frac_part):
            return f"{int_part},{frac_part}"
        return int_part

    # scale > 0
    frac_part = frac_part[:scale].ljust(scale, "0")
    return f"{int_part},{frac_part}"


def _replace_cell_text(cell, new_text: str) -> None:
    """
    Безопасная замена текста в ячейке.
    Не трогаем структуру таблицы, только содержимое абзацев.
    """
    if not cell.paragraphs:
        return

    first = cell.paragraphs[0]
    replace_paragraph_text(first, new_text)

    for p in cell.paragraphs[1:]:
        replace_paragraph_text(p, "")


def _set_table_paragraph_alignment(paragraph, alignment) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)


def _set_table_fixed_widths_from_grid(table, *, preserve_geometry: bool | None = None):
    """
    Жестко переносит ширины из tblGrid в tblW и tcW,
    чтобы Word не пересчитывал геометрию таблицы как auto.
    Это снижает риск визуально "двойных" линий.
    """
    if preserve_geometry is None:
        preserve_geometry = _safe_formatter_preserve_table_geometry(table)
    if preserve_geometry:
        return

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        return

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        return

    grid_cols = grid.findall(qn("w:gridCol"))
    if not grid_cols:
        return

    grid_widths = []
    for gc in grid_cols:
        w = gc.get(qn("w:w"))
        try:
            grid_widths.append(int(w))
        except Exception:
            return

    total_width = sum(grid_widths)

    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.insert(0, tblW)

    tblW.set(qn("w:w"), str(total_width))
    tblW.set(qn("w:type"), "dxa")

    for row in table.rows:
        logical_col_idx = 0

        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            grid_span = 1
            gridSpan = tcPr.find(qn("w:gridSpan"))
            if gridSpan is not None:
                try:
                    grid_span = int(gridSpan.get(qn("w:val")))
                except Exception:
                    grid_span = 1

            span_width = sum(grid_widths[logical_col_idx: logical_col_idx + grid_span])

            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)

            tcW.set(qn("w:w"), str(span_width))
            tcW.set(qn("w:type"), "dxa")

            logical_col_idx += grid_span
            
def ensure_all_table_rows_cant_split(document):
    """Set <w:cantSplit/> on every <w:tr> in the document so rows are never split
    across page boundaries. Idempotent — reuses the helper from
    table_split_prototype which checks for the existing element before adding."""
    from .table_split_prototype import _ensure_table_rows_cant_split
    for table in document.tables:
        _ensure_table_rows_cant_split(table._element)


def format_tables(document):
    for table in document.tables:
        preserve_geometry = _safe_formatter_preserve_table_geometry(table)
        apply_table_borders(table)

        if not preserve_geometry:
            try:
                table.autofit = False
            except Exception:
                pass

            _set_table_fixed_widths_from_grid(table, preserve_geometry=preserve_geometry)

        column_scales = _get_table_numeric_column_scales(table)

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = _get_table_cell_text(cell)
                normalized_number = _format_table_number_for_column(
                    cell_text,
                    column_scales.get(col_idx, 0),
                )

                # Если это числовая ячейка не в первой строке — нормализуем текст
                if row_idx != 0 and normalized_number is not None:
                    _replace_cell_text(cell, normalized_number)

                for paragraph in cell.paragraphs:
                    force_paragraph_xml_spacing(paragraph, line_rule="auto")
                    fmt = paragraph.paragraph_format
                    fmt.first_line_indent = Cm(0)
                    fmt.left_indent = Cm(0)
                    fmt.right_indent = Cm(0)
                    fmt.line_spacing = LINE_SPACING_TABLE
                    fmt.space_before = Pt(0)
                    fmt.space_after = Pt(0)
                    fmt.keep_together = False
                    fmt.keep_with_next = False
                    fmt.page_break_before = False
                    fmt.widow_control = False

                    if row_idx == 0:
                        _set_table_paragraph_alignment(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
                    elif normalized_number is not None:
                        _set_table_paragraph_alignment(paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
                    else:
                        _set_table_paragraph_alignment(paragraph, WD_ALIGN_PARAGRAPH.JUSTIFY)

                    for run in paragraph.runs:
                        set_run_font(
                            run,
                            size_pt=TABLE_FONT_SIZE_PT,
                            bold=False,
                            italic=False,
                            all_caps=False,
                        )
                        
def smart_repair_heading1(paragraph, text: str):
    cleaned = strip_leading_heading_garbage(text)
    parsed = parse_heading1(cleaned)
    if not parsed:
        return False

    if parsed["kind"] == "heading1_chapter":
        chapter_num = parsed["chapter_num"]
        title = parsed["title"].lstrip(".: ").upper()
        replace_paragraph_text(paragraph, f"{chapter_num}. {title}")
        remove_paragraph_numbering(paragraph)
        format_heading1(paragraph)
        return True

    if parsed["kind"] == "heading1_exact":
        replace_paragraph_text(paragraph, cleaned.upper())
        remove_paragraph_numbering(paragraph)
        format_heading1(paragraph)
        return True

    return False


def smart_repair_broken_heading2(paragraph, current_chapter_num, next_paragraph_num):
    if current_chapter_num is None or next_paragraph_num is None:
        return None

    text = clean_spaces(paragraph.text)
    parsed = parse_broken_heading2(text)
    if not parsed:
        return None

    if not is_probable_center_bold_heading(paragraph):
        return None

    title = parsed["title"].lstrip(". ").strip()
    new_text = f"{current_chapter_num}.{next_paragraph_num}. {title}"
    replace_paragraph_text(paragraph, new_text)
    remove_paragraph_numbering(paragraph)
    format_heading2(paragraph)
    return new_text


def looks_like_heading2_title(text: str) -> bool:
    t = clean_spaces(text)
    if not t:
        return False

    low = t.lower()

    if low.startswith("таблица "):
        return False
    if low.startswith("рисунок "):
        return False
    if low.startswith("рис. "):
        return False
    if low.startswith("продолжение таблицы"):
        return False
    if low.startswith("продолжение табл."):
        return False

    if is_table_continuation_text(t):
        return False

    if low in REFERENCE_SUBHEADINGS_CANON:
        return False
    if parse_heading1(t) or parse_heading2(t) or parse_broken_heading2(t):
        return False
    if TABLE_NUM_RE.match(t) or FIG_RE.match(t) or DASH_LINE_RE.match(t):
        return False
    if re.match(r"^\s*(источник|составлено по|рассчитано по|примечание)\s*:", t, re.IGNORECASE):
        return False
    if t.endswith((".", ":", ";", "?", "!")):
        return False
    if len(t) > 220:
        return False

    return True


def is_likely_numbered_heading2_candidate(paragraph, current_chapter_num, next_paragraph_num, prev_kind=None):
    if current_chapter_num is None or next_paragraph_num is None:
        return False
    if not paragraph_has_numbering(paragraph):
        return False

    text = clean_spaces(paragraph.text)
    if not looks_like_heading2_title(text):
        return False

    return is_heading2_promotion_safe(paragraph)


    
def normalize_heading2_numbering(paragraph, current_chapter_num, next_paragraph_num):
    if current_chapter_num is None or next_paragraph_num is None:
        return None

    text = strip_leading_heading_garbage(paragraph.text)
    text = clean_spaces(text)
    if not text:
        return None

    has_num = paragraph_has_numbering(paragraph)
    parsed = parse_heading2(text)

    if parsed:
        normalized = f"{parsed['chapter_num']}.{parsed['paragraph_num']}. {parsed['title']}"
        if text != normalized:
            replace_paragraph_text(paragraph, normalized)
        if has_num:
            remove_paragraph_numbering(paragraph)
        return normalized

    if has_num and looks_like_heading2_title(text) and is_heading2_promotion_safe(paragraph):
        title = text.lstrip(". ").strip()
        new_text = f"{current_chapter_num}.{next_paragraph_num}. {title}"
        replace_paragraph_text(paragraph, new_text)
        remove_paragraph_numbering(paragraph)
        format_heading2(paragraph)
        return new_text

    return None


def is_heading2_promotion_safe(paragraph, parsed_h2=None, toc_text=None):
    if toc_text:
        return True

    if parsed_h2 and paragraph_has_heading_style_or_outline(paragraph):
        return True

    if paragraph_has_heading_style_or_outline(paragraph):
        return True

    if is_probable_center_bold_heading(paragraph):
        return True

    return False


def is_heading1_promotion_safe(paragraph, parsed_h1, toc_text=None):
    if not parsed_h1:
        return False

    if parsed_h1["kind"] == "heading1_exact":
        return True

    # toc_text alone is NOT sufficient. A body sentence shaped like `N. Foo.`
    # may match parse_heading1 and find a chapter-N entry inside toc_h1_map
    # built from an old / fake TOC; relying on toc_text in that case promoted
    # the body sentence to Heading 1 and triggered uppercase + duplicate-H1
    # contamination (file 198 idx 80 "1. Маркетинговый подход. Данный
    # подход" → "1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ…"). Require an independent
    # structural signal — Heading style / outline level or a centred-bold
    # visual heading — for the toc_text shortcut to fire.
    has_structural_signal = (
        paragraph_has_heading_style_or_outline(paragraph)
        or is_probable_center_bold_heading(paragraph)
    )

    if toc_text and has_structural_signal:
        return True

    if has_structural_signal:
        return True

    title = clean_spaces(parsed_h1.get("title") or "")
    if re.search(r"\.\s+\S", title):
        return False

    return True

def normalize_table_continuation_text(paragraph):
    text = clean_spaces(paragraph.text)
    low = text.lower()

    if "продол" in low and "таб" in low:
        m = re.search(r"(\d+(?:\.\d+){1,2})", text)
        if m:
            replace_paragraph_text(paragraph, f"Продолжение таблицы {m.group(1)}")


def normalize_figure_caption_text(paragraph):
    text = clean_spaces(paragraph.text)
    if not text:
        return

    m = FIG_RE.match(text)
    if not m:
        return

    number = m.group(2)
    title = clean_spaces(m.group(3) or "")

    if title:
        normalized = f"Рис. {number}. {title}"
    else:
        normalized = f"Рис. {number}"

    if text != normalized:
        replace_paragraph_text(paragraph, normalized)


def normalize_toc_line(text: str) -> str:
    t = clean_spaces(text.replace("\t", " "))

    # Убираем хвосты содержания:
    # ..... 12
    # ……… 12
    # . . . 12
    # смешанные лидеры и пробелы перед номером страницы
    t = re.sub(r'[\s\.\u2024\u2025\u2026·•]+(\d+)\s*$', "", t).strip()

    # Дополнительно убираем хвосты вида "………………" без номера,
    # если Word уже отдельно разорвал страницу/табуляцию
    t = re.sub(r'[\s\.\u2024\u2025\u2026·•]+$', "", t).strip()

    return t


TOC_ENTRY_PAGE_RE = re.compile(
    r"^\s*(?P<title>.+?)(?:\t+|[\s\.\u2024\u2025\u2026·•]+)(?P<page>\d{1,4})\s*$"
)


def _contents_heading_text(text: str) -> str | None:
    t = normalize_toc_line(text).strip().upper().rstrip(".")
    if "СОДЕРЖАН" in t:
        return "СОДЕРЖАНИЕ"
    if "ОГЛАВЛЕН" in t:
        return "ОГЛАВЛЕНИЕ"
    return None


def _split_toc_entry_text(text: str) -> tuple[str, str] | None:
    raw = (text or "").replace("\u00A0", " ").replace("\u202F", " ")
    match = TOC_ENTRY_PAGE_RE.match(raw)
    if not match:
        return None

    title = clean_spaces(match.group("title").replace("\t", " "))
    title = re.sub(r"[\s\.\u2024\u2025\u2026·•]+$", "", title).strip()
    page = match.group("page")
    if not title:
        return None
    return title, page


def _build_toc_chapter_page_map(paragraphs, start_idx: int, end_idx: int) -> dict[str, str]:
    page_by_chapter: dict[str, str] = {}
    for paragraph in paragraphs[start_idx:end_idx]:
        entry = _split_toc_entry_text(paragraph.text)
        if not entry:
            continue

        title, page = entry
        parsed_h1 = parse_heading1(title)
        if parsed_h1 and parsed_h1["kind"] == "heading1_chapter":
            page_by_chapter.setdefault(parsed_h1["chapter_num"], page)
            continue

        parsed_h2 = parse_heading2(title)
        if parsed_h2:
            page_by_chapter.setdefault(parsed_h2["chapter_num"], page)

    return page_by_chapter


def _infer_toc_chapter_entry_text(text: str, page_by_chapter: dict[str, str]) -> tuple[str, str] | None:
    title = normalize_toc_line(text)
    if not title or _split_toc_entry_text(text):
        return None

    parsed_h1 = parse_heading1(title)
    if not parsed_h1 or parsed_h1["kind"] != "heading1_chapter":
        return None

    page = page_by_chapter.get(parsed_h1["chapter_num"])
    if not page:
        return None

    return f'{parsed_h1["chapter_num"]}. {parsed_h1["title"]}', page


def _clear_paragraph_tab_stops(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    for old_tabs in list(p_pr.findall(qn("w:tabs"))):
        p_pr.remove(old_tabs)


def _set_toc_tab_stop(paragraph) -> None:
    p_pr = paragraph._element.get_or_add_pPr()
    _clear_paragraph_tab_stops(paragraph)

    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(Cm(16).twips))
    tabs.append(tab)
    p_pr.append(tabs)


def _format_contents_heading(paragraph, text: str) -> None:
    replace_paragraph_text(paragraph, text)
    _clear_paragraph_tab_stops(paragraph)
    paragraph.paragraph_format.page_break_before = False
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=True, all_caps=False)


def _format_toc_entry(paragraph, title: str, page: str) -> None:
    replace_paragraph_text(paragraph, f"{title}\t{page}")
    _set_toc_tab_stop(paragraph)
    paragraph.paragraph_format.page_break_before = False
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.left_indent = Cm(0)
    paragraph.paragraph_format.right_indent = Cm(0.75)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, all_caps=False)


def normalize_contents_layout(document, body_start):
    # Product rule for B2.2: until TOC recovery is implemented as a separate
    # patch, front matter before the real ВВЕДЕНИЕ is text-frozen.
    return body_start


def build_toc_heading_maps(document, body_start):
    h1_map = {}
    h2_map = {}

    if body_start is None:
        return h1_map, h2_map

    for idx, p in enumerate(document.paragraphs):
        if idx >= body_start:
            break

        text = normalize_toc_line(p.text)
        if not text:
            continue

        parsed_h1 = parse_heading1(text)
        if parsed_h1 and parsed_h1["kind"] == "heading1_chapter":
            h1_map[parsed_h1["chapter_num"]] = f'{parsed_h1["chapter_num"]}. {parsed_h1["title"]}'
            continue

        parsed_h2 = parse_heading2(text)
        if parsed_h2:
            key = (parsed_h2["chapter_num"], parsed_h2["paragraph_num"])
            h2_map[key] = f'{parsed_h2["chapter_num"]}.{parsed_h2["paragraph_num"]}. {parsed_h2["title"]}'

    return h1_map, h2_map

def detect_kind_from_paragraph_object(paragraph, text: str, prev_kind=None) -> str:
    t = clean_spaces(text)
    low = t.lower()

    parsed_h1 = parse_heading1(t)
    if parsed_h1:
        if parsed_h1["kind"] == "heading1_exact" and low == "содержание":
            return "toc_heading"
        return "heading1"

    if parse_heading2(t):
        return "heading2"

    if parse_broken_heading2(t):
        return "broken_heading2"

    m_tab = TABLE_NUM_RE.match(t)
    if m_tab and not caption_tail_is_reference_prose(m_tab.group(2) or ""):
        return "table_caption"

    if is_table_continuation_text(t):
        return "table_continuation"

    m_fig = FIG_RE.match(t)
    if m_fig and not caption_tail_is_reference_prose(m_fig.group(3) or ""):
        return "figure_caption"

    if re.match(r"^\s*(источник|составлено по|рассчитано по|примечание)\s*:", t, re.IGNORECASE):
        return "source_line"

    style_name = ""
    try:
        style_name = (paragraph.style.name or "").strip().lower()
    except Exception:
        style_name = ""

    if style_name in {"heading 1", "заголовок 1"}:
        return "heading1"

    if style_name in {"heading 2", "заголовок 2"}:
        return "heading2"


    if prev_kind in {"table_caption", "table_continuation"}:
        return "table_title"

    return "body_text"
    
def split_manual_dash_lists(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            raw = p.text.replace("\r", "\n").replace("\v", "\n")
            if "\n" not in raw:
                continue

            parts = [x.strip() for x in re.split(r"[\n]+", raw) if x.strip()]
            if len(parts) < 2:
                continue

            _is_dash_split = all(DASH_LINE_RE.match(x) for x in parts[1:])
            _is_colon_numdot_split = (
                parts[0].endswith(":")
                and all(_NUMERIC_DOT_LIST_RE.match(x) for x in parts[1:])
            )
            if not (_is_dash_split or _is_colon_numdot_split):
                continue

            replace_paragraph_text(p, parts[0])
            prev = p
            for item in parts[1:]:
                prev = insert_paragraph_after(prev, item)

            changed = True
            break


def _nonempty_body_blocks_after(children, start_idx: int, *, max_blocks: int = 2) -> list:
    blocks = []
    idx = start_idx + 1
    while idx < len(children) and len(blocks) < max_blocks:
        child = children[idx]
        if child.tag == qn("w:p") and _is_empty_paragraph_xml(child):
            idx += 1
            continue
        blocks.append(child)
        idx += 1
    return blocks


def _is_table_title_candidate_block(child) -> bool:
    if child.tag != qn("w:p") or _is_empty_paragraph_xml(child):
        return False
    if child.findall(".//" + qn("w:drawing")):
        return False

    text = _paragraph_text_from_xml(child)
    if not text or len(text) > 120:
        return False

    low = text.lower()
    if (
        TABLE_NUM_RE.match(text)
        or is_table_continuation_text(text)
        or FIG_RE.match(text)
        or FIG_SERVICE_LINE_RE.match(text)
        or parse_heading1(text)
        or parse_heading2(text)
        or parse_broken_heading2(text)
        or low == "приложения"
        or normalize_appendix_start_label_text(text)
        or is_appendix_continuation_label_text(text)
    ):
        return False

    if re.match(r"^\s*в\s+таблиц[еуы]\b", text, re.IGNORECASE):
        return False

    if len(text.split()) > 14 and re.search(r"[.!?]\s+\S", text):
        return False

    return True


def _is_confirmed_table_caption_paragraph(document, paragraph, body_start) -> bool:
    text = clean_spaces(paragraph.text)
    m = TABLE_NUM_RE.match(text)
    if not m:
        return False

    # Reject reference-prose paragraphs like "Таблица 1.1.1 показывает ...":
    # these are body text referring to the table, not the caption itself.
    if caption_tail_is_reference_prose(m.group(2) or ""):
        return False

    paragraph_info = _paragraph_lookup(document).get(paragraph._p)
    if paragraph_info is None:
        return False
    paragraph_idx, _paragraph = paragraph_info
    if paragraph_idx < body_start:
        return False

    children = list(document.element.body)
    try:
        caption_child_idx = children.index(paragraph._p)
    except ValueError:
        return False

    blocks = _nonempty_body_blocks_after(children, caption_child_idx, max_blocks=2)
    if not blocks:
        return False

    has_inline_title = bool(clean_spaces(m.group(2)))
    if blocks[0].tag == qn("w:tbl"):
        return True

    if has_inline_title:
        return False

    return (
        len(blocks) >= 2
        and _is_table_title_candidate_block(blocks[0])
        and blocks[1].tag == qn("w:tbl")
    )


def _classify_paragraph_with_table_adjacency(document, paragraph, body_start, prev_kind=None) -> str:
    kind = classify_paragraph(clean_spaces(paragraph.text), prev_kind=prev_kind)
    if kind == "table_caption" and not _is_confirmed_table_caption_paragraph(document, paragraph, body_start):
        return "body_text"
    return kind


def _clean_inline_table_caption_title(text: str) -> str:
    title = clean_spaces(text)
    title = re.sub(r"^\s*[-–—]\s*", "", title, count=1)
    return clean_spaces(title)


def split_table_captions_prepass(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text)
            if not text or not text.lower().startswith("таблица"):
                continue

            m = TABLE_NUM_RE.match(text)
            if not m:
                continue
            if not _is_confirmed_table_caption_paragraph(document, p, body_start):
                continue

            number = m.group(1)
            raw_tail = clean_spaces(m.group(2))
            if not raw_tail:
                continue

            tail = _clean_inline_table_caption_title(raw_tail)
            if not tail:
                replace_paragraph_text(p, f"Таблица {number}")
                format_table_caption(p)
                changed = True
                break

            replace_paragraph_text(p, f"Таблица {number}")
            title_p = insert_paragraph_after(p, tail)

            format_table_caption(p)
            format_table_title(title_p)

            changed = True
            break


def _structural_soft_break_segment_kind(text: str) -> str:
    t = clean_spaces(text)
    if not t:
        return ""
    parsed_h1 = parse_heading1(t)
    if parsed_h1 and parsed_h1["kind"] == "heading1_chapter":
        return "heading1"
    if parse_heading2(t):
        return "heading2"
    return "body"


# P4 / DEFECT 4: source/note semantic-line prefix.
# Matches "Источник: …", "Примечание: …", "Составлено по: …", "Рассчитано по: …".
_SOURCE_NOTE_SEGMENT_RE = re.compile(
    r"^\s*(источник|составлено по|рассчитано по|примечание)\s*:",
    re.IGNORECASE,
)


def _is_source_note_segment(text: str) -> bool:
    return bool(_SOURCE_NOTE_SEGMENT_RE.match(clean_spaces(text)))


def _should_split_body_soft_break_segments(segments: list[str]) -> bool:
    if len(segments) < 2:
        return False
    first_kind = _structural_soft_break_segment_kind(segments[0])
    if first_kind in {"heading1", "heading2"}:
        second_kind = _structural_soft_break_segment_kind(segments[1])
        return second_kind in {"heading1", "heading2", "body"}
    # P4 / DEFECT 4: allow body→body split only when every segment is a
    # source/note semantic line (e.g. "Источник: … ↵ Примечание: …").
    # Avoids broadening generic body-soft-break splitting.
    if all(_is_source_note_segment(seg) for seg in segments):
        return True
    return False


def split_body_structural_soft_breaks(document, body_start):
    if body_start is None:
        return False

    paragraphs = document.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if idx < body_start:
            continue
        raw = paragraph.text
        if "\n" not in raw and "\v" not in raw:
            continue

        segments = [clean_spaces(part) for part in re.split(r"[\n\v]+", raw) if clean_spaces(part)]
        if not _should_split_body_soft_break_segments(segments):
            continue

        replace_paragraph_text(paragraph, segments[0])
        prev = paragraph
        for segment in segments[1:]:
            prev = insert_paragraph_after(prev, segment)
        return True

    return False


REFERENCE_LEADING_NUMBER_RE = re.compile(
    r"^\s*(?:\[\s*\d{1,3}\s*\]\s*|\d{1,3}[\.)](?:\s+|(?=[A-ZА-ЯЁ])))"
)


def strip_leading_reference_number(text: str) -> str:
    return clean_spaces(REFERENCE_LEADING_NUMBER_RE.sub("", clean_spaces(text), count=1))


def convert_reference_numbering_to_plain_text(document, body_start):
    """
    Приводит список источников к плоскому тексту и всегда строит
    СКВОЗНУЮ нумерацию заново по всем разделам списка.

    Разделы ("Официальные материалы", "Статистические материалы" и т.д.)
    не нумеруются, а обычные записи получают 1., 2., 3. ... без перезапуска.
    """
    in_references = False
    ref_counter = 1

    for idx, paragraph in enumerate(document.paragraphs):
        if idx < body_start:
            continue

        text = clean_spaces(paragraph.text)
        low = text.lower()
        canonical = canonical_reference_block_heading_paragraph(paragraph)

        if low in {
            "список использованных источников",
            "список использованной литературы",
        }:
            in_references = True
            ref_counter = 1
            continue

        if not in_references:
            continue

        if low in {"приложения", "приложение"}:
            in_references = False
            continue

        # Подзаголовки разделов внутри списка источников:
        # не нумеруем, просто нормализуем и форматируем.
        if canonical:
            replace_paragraph_text(paragraph, canonical)
            remove_paragraph_numbering(paragraph)
            remove_page_break_artifacts_from_paragraph(paragraph)
            format_reference_subheading(paragraph)
            continue

        if is_empty_paragraph(paragraph):
            continue

        # Любой обычный источник в блоке литературы:
        # полностью снимаем Word-numbering и видимую кривую нумерацию,
        # затем строим свою сквозную.
        remove_paragraph_numbering(paragraph)
        remove_page_break_artifacts_from_paragraph(paragraph)
        set_paragraph_style_safe(paragraph, "Normal", "Обычный")
        clear_paragraph_outline_level(paragraph)

        clean = clean_spaces(paragraph.text)

        # Снимаем видимый номер в начале строки, если он уже есть:
        # 1. ... / 1) ... / [1] ...
        clean = strip_leading_reference_number(clean)

        # Снимаем маркеры/буллеты, если они вдруг приехали из исходника
        clean = re.sub(r"^\s*[•·▪■◆►→\-–—]+\s*", "", clean)
        clean = clean_spaces(clean)

        if not clean:
            continue

        normalized = f"{ref_counter}. {clean}"
        ref_counter += 1

        normalized = smart_normalize_reference_line_case(normalize_reference_url_spacing(normalized))
        replace_paragraph_text(paragraph, normalized)
        format_reference_entry(paragraph)

def compact_references_block(document, body_start):
    changed = True

    while changed:
        changed = False
        in_references = False
        paragraphs = document.paragraphs

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text)
            low = text.lower()
            canonical = canonical_reference_block_heading_paragraph(p)

            if is_references_heading_text(text):
                in_references = True
                continue

            if not in_references:
                continue

            if is_appendix_heading_text(text):
                in_references = False
                continue

            # Полностью убираем пустые абзацы внутри блока литературы,
            # но сохраняем один пустой абзац прямо перед подзаголовком раздела.
            if is_empty_paragraph(p):
                # Peek ahead: keep blank if next non-empty para is a subheading
                next_nonempty = None
                for _np in paragraphs[idx + 1:]:
                    _nt = clean_spaces(_np.text)
                    if _nt:
                        next_nonempty = _np
                        break
                if next_nonempty and canonical_reference_block_heading_paragraph(next_nonempty):
                    pass  # keep the blank
                else:
                    remove_paragraph(p)
                    changed = True
                    break

            # Сначала снимаем весь мусор разрывов / списков / заголовков
            remove_page_break_artifacts_from_paragraph(p)
            remove_paragraph_numbering(p)
            set_paragraph_style_safe(p, "Normal", "Обычный")
            clear_paragraph_outline_level(p)

            # Подзаголовки разделов внутри литературы
            if canonical:
                replace_paragraph_text(p, canonical)
                format_reference_subheading(p)

                p.paragraph_format.page_break_before = False
                p.paragraph_format.keep_with_next = False
                p.paragraph_format.keep_together = False
                p.paragraph_format.widow_control = False

                continue

            # Обычный источник
            clean = clean_spaces(p.text)

            m = re.match(r"^\s*(\d+)\.\s+(.+)$", clean)
            if m:
                number = int(m.group(1))
                source_text = clean_spaces(m.group(2))
                normalized = f"{number}. {source_text}"
            else:
                normalized = clean

            normalized = smart_normalize_reference_line_case(normalize_reference_url_spacing(normalized))

            if clean != normalized:
                replace_paragraph_text(p, normalized)

            format_reference_entry(p)

            # Финальный добивающий reset именно после format_body
            p.paragraph_format.page_break_before = False
            p.paragraph_format.keep_with_next = False
            p.paragraph_format.keep_together = False
            p.paragraph_format.widow_control = False

def ensure_single_blank_after_references_heading(document, body_start):
    any_changes = False
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text)
            if not is_references_heading_text(text):
                continue

            # гарантируем 1 пустую строку после заголовка списка
            if idx + 1 >= len(paragraphs):
                new_p = insert_paragraph_after(p, "")
                format_empty_paragraph(new_p)
                changed = True
                any_changes = True
                break

            next_p = paragraphs[idx + 1]

            if not is_reference_spacing_paragraph(next_p):
                new_p = insert_paragraph_after(p, "")
                format_empty_paragraph(new_p)
                changed = True
                any_changes = True
                break

            # если пустых строк больше одной — удаляем лишние
            while idx + 2 < len(paragraphs) and is_reference_spacing_paragraph(paragraphs[idx + 2]):
                remove_paragraph(paragraphs[idx + 2])
                paragraphs = document.paragraphs
                changed = True
                any_changes = True

            format_empty_paragraph(next_p)
            break

    return any_changes


def ensure_blank_before_reference_subheadings(document, body_start):
    """
    Ensure exactly one blank paragraph appears immediately before each
    reference subheading (e.g. "Официальные материалы").
    """
    any_changes = False
    changed = True
    while changed:
        changed = False
        in_references = False
        paragraphs = document.paragraphs

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text)

            if is_references_heading_text(text):
                in_references = True
                continue

            if not in_references:
                continue

            if is_appendix_heading_text(text):
                in_references = False
                continue

            if not canonical_reference_block_heading_paragraph(p):
                continue

            if idx == 0:
                continue

            prev_p = paragraphs[idx - 1]
            if not is_reference_spacing_paragraph(prev_p):
                new_el = OxmlElement("w:p")
                p._element.addprevious(new_el)
                changed = True
                any_changes = True
                break

            while idx - 2 >= body_start and is_reference_spacing_paragraph(paragraphs[idx - 2]):
                remove_paragraph(paragraphs[idx - 2])
                changed = True
                any_changes = True
                break

            if changed:
                break

            format_empty_paragraph(prev_p)

    return any_changes


def ensure_single_blank_after_headings(document, body_start):
    paragraphs = document.paragraphs
    prev_kind = None
    changed = False
    in_references = False

    idx = max(body_start, 0)

    while idx < len(paragraphs):
        p = paragraphs[idx]
        text = clean_spaces(p.text)

        if is_references_heading_text(text):
            in_references = True
        elif in_references and is_appendix_heading_text(text):
            in_references = False

        kind = classify_paragraph(text, prev_kind=prev_kind)
        parsed_h1 = parse_heading1(text)

        need_blank_after = False

        # После параграфов 1.1 / 1.2 / 2.1 и т.д. нужна одна пустая строка.
        # Дополнительно поддерживаем "голый" номер вида "1.3" (без заголовка
        # в той же строке): в исходниках встречается такой артефакт, и после
        # него также должен быть один пустой абзац.
        is_heading2_number_only = bool(re.match(r"^\s*\d+\.\d+\.?\s*$", text))
        if kind == "heading2" or is_heading2_number_only:
            need_blank_after = True

        # После ВВЕДЕНИЯ / ЗАКЛЮЧЕНИЯ / СПИСКА ИСТОЧНИКОВ нужна одна пустая строка
        # После названий глав 1 / 2 / 3 пустая строка НЕ нужна
        elif parsed_h1:
            if parsed_h1["kind"] == "heading1_exact":
                need_blank_after = True
            elif parsed_h1["kind"] == "heading1_chapter":
                need_blank_after = False

        if not need_blank_after:
            prev_kind = kind
            idx += 1
            continue

        if idx + 1 >= len(paragraphs):
            new_p = insert_paragraph_after(p, "")
            format_empty_paragraph(new_p)
            changed = True
            break

        next_p = paragraphs[idx + 1]

        if not is_empty_paragraph(next_p):
            new_p = insert_paragraph_after(p, "")
            format_empty_paragraph(new_p)
            changed = True
            break

        while idx + 2 < len(paragraphs) and is_empty_paragraph(paragraphs[idx + 2]):
            remove_paragraph(paragraphs[idx + 2])
            paragraphs = document.paragraphs
            changed = True

        if idx + 1 < len(paragraphs) and is_empty_paragraph(paragraphs[idx + 1]):
            format_empty_paragraph(paragraphs[idx + 1])

        prev_kind = kind
        idx += 1

    return changed
    
def collapse_empty_paragraphs_in_body(paragraphs, body_start):
    empty_count = 0
    for idx, p in enumerate(list(paragraphs)):
        if idx < body_start:
            continue

        if is_empty_paragraph(p):
            empty_count += 1
            if empty_count > 1:
                remove_paragraph(p)
        else:
            empty_count = 0

def remove_single_empty_between_body_paragraphs(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            if not is_empty_paragraph(p):
                continue

            # Ищем ближайший непустой абзац слева
            prev_idx = idx - 1
            while prev_idx >= body_start and is_empty_paragraph(paragraphs[prev_idx]):
                prev_idx -= 1

            # Ищем ближайший непустой абзац справа
            next_idx = idx + 1
            while next_idx < len(paragraphs) and is_empty_paragraph(paragraphs[next_idx]):
                next_idx += 1

            if prev_idx < body_start or next_idx >= len(paragraphs):
                continue

            prev_text = clean_spaces(paragraphs[prev_idx].text)
            next_text = clean_spaces(paragraphs[next_idx].text)

            # ВАЖНО:
            # Не трогаем пустую строку рядом с формульным блоком.
            # Иначе сначала normalize_formula_blocks() вставит нужный отступ,
            # а потом этот проход его снесёт.
            if is_formula_block_paragraph_text(prev_text) or is_formula_block_paragraph_text(next_text):
                continue

            prev_prev_kind = None
            for j in range(body_start, prev_idx):
                t = clean_spaces(paragraphs[j].text)
                if not t:
                    continue
                prev_prev_kind = classify_paragraph(t, prev_kind=prev_prev_kind)

            prev_kind = classify_paragraph(prev_text, prev_kind=prev_prev_kind)
            next_kind = classify_paragraph(next_text, prev_kind=prev_kind)

            # Удаляем только случайную пустую строку между двумя обычными абзацами текста
            if prev_kind == "body_text" and next_kind == "body_text":
                remove_paragraph(p)
                changed = True
                break

    return changed
def ensure_empty_after_source_and_note(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs
        prev_kind = None
        in_references = False

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text)

            if is_references_heading_text(text):
                in_references = True
                prev_kind = "heading1"
                continue

            if in_references and is_appendix_heading_text(text):
                in_references = False

            # Внутри списка источников ничего не разрежаем
            if in_references:
                prev_kind = "body_text"
                continue

            kind = detect_kind_from_paragraph_object(p, text, prev_kind=prev_kind)
            is_note_line = bool(re.match(r"^\s*примечание\s*:", text, re.IGNORECASE))

            # ===== FIGURE CAPTION LOGIC =====
            # После подписи рисунка:
            # - если дальше идёт Источник:/Примечание: -> пустую строку НЕ вставляем
            # - иначе должна быть ровно одна пустая строка
            if kind == "figure_caption":
                if idx + 1 >= len(paragraphs):
                    new_p = insert_paragraph_after(p, "")
                    format_empty_paragraph(new_p)
                    changed = True
                    break

                next_text = clean_spaces(paragraphs[idx + 1].text)
                next_is_service = bool(FIG_SERVICE_LINE_RE.match(next_text))

                if next_is_service:
                    # Между подписью рисунка и Источником/Примечанием пустой строки быть не должно
                    if idx + 1 < len(paragraphs) and is_empty_paragraph(paragraphs[idx + 1]):
                        remove_paragraph(paragraphs[idx + 1])
                        changed = True
                        break

                    prev_kind = kind
                    continue

                # Иначе после подписи рисунка нужна ровно одна пустая строка
                if not is_empty_paragraph(paragraphs[idx + 1]):
                    new_p = insert_paragraph_after(p, "")
                    format_empty_paragraph(new_p)
                    changed = True
                    break

                if (
                    idx + 2 < len(paragraphs)
                    and is_empty_paragraph(paragraphs[idx + 1])
                    and is_empty_paragraph(paragraphs[idx + 2])
                ):
                    remove_paragraph(paragraphs[idx + 2])
                    changed = True
                    break

                format_empty_paragraph(paragraphs[idx + 1])
                prev_kind = kind
                continue

            # ===== SOURCE LINE LOGIC =====
            if kind == "source_line":

                # Если сразу после Источника идёт пустая строка,
                # а после неё Примечание: -> удалить эту пустую строку
                if (
                    idx + 2 < len(paragraphs)
                    and is_empty_paragraph(paragraphs[idx + 1])
                    and re.match(r"^\s*примечание\s*:", clean_spaces(paragraphs[idx + 2].text), re.IGNORECASE)
                ):
                    remove_paragraph(paragraphs[idx + 1])
                    changed = True
                    break

                # Если сразу после Источника идёт Примечание: -> ничего не вставляем
                if idx + 1 < len(paragraphs):
                    next_text = clean_spaces(paragraphs[idx + 1].text)
                    if re.match(r"^\s*примечание\s*:", next_text, re.IGNORECASE):
                        prev_kind = kind
                        continue

                # Во всех остальных случаях после Источника должна быть ровно одна пустая строка
                if idx + 1 >= len(paragraphs):
                    new_p = insert_paragraph_after(p, "")
                    format_empty_paragraph(new_p)
                    changed = True
                    break

                if not is_empty_paragraph(paragraphs[idx + 1]):
                    new_p = insert_paragraph_after(p, "")
                    format_empty_paragraph(new_p)
                    changed = True
                    break

                if (
                    idx + 2 < len(paragraphs)
                    and is_empty_paragraph(paragraphs[idx + 1])
                    and is_empty_paragraph(paragraphs[idx + 2])
                ):
                    remove_paragraph(paragraphs[idx + 2])
                    changed = True
                    break

                format_empty_paragraph(paragraphs[idx + 1])
                prev_kind = kind
                continue

            # ===== NOTE LINE LOGIC =====
            if is_note_line:
                # После Примечания должна быть ровно одна пустая строка
                if idx + 1 >= len(paragraphs):
                    new_p = insert_paragraph_after(p, "")
                    format_empty_paragraph(new_p)
                    changed = True
                    break

                if not is_empty_paragraph(paragraphs[idx + 1]):
                    new_p = insert_paragraph_after(p, "")
                    format_empty_paragraph(new_p)
                    changed = True
                    break

                if (
                    idx + 2 < len(paragraphs)
                    and is_empty_paragraph(paragraphs[idx + 1])
                    and is_empty_paragraph(paragraphs[idx + 2])
                ):
                    remove_paragraph(paragraphs[idx + 2])
                    changed = True
                    break

                format_empty_paragraph(paragraphs[idx + 1])
                prev_kind = "body_text"
                continue

            prev_kind = kind
            
def ensure_one_empty_after(paragraphs, index):
    """Ensure exactly one empty paragraph right after paragraphs[index]."""
    if index >= len(paragraphs):
        return False

    changed = False
    p = paragraphs[index]

    if index + 1 >= len(paragraphs):
        new_p = insert_paragraph_after(p, "")
        format_empty_paragraph(new_p)
        return True

    next_p = paragraphs[index + 1]
    if not is_empty_paragraph(next_p):
        new_p = insert_paragraph_after(p, "")
        format_empty_paragraph(new_p)
        return True

    format_empty_paragraph(next_p)

    while index + 2 < len(paragraphs) and is_empty_paragraph(paragraphs[index + 2]):
        remove_paragraph(paragraphs[index + 2])
        paragraphs = p._parent.paragraphs
        changed = True

    return changed


def ensure_single_blank_before_figure_captions(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs
        in_references = False

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text)

            if is_references_heading_text(text):
                in_references = True
                continue

            if in_references and is_appendix_heading_text(text):
                in_references = False

            if in_references or not FIG_RE.match(text) or idx <= body_start:
                continue

            prev_idx = idx - 1
            prev_p = paragraphs[prev_idx]

            if paragraph_has_drawing(prev_p):
                continue

            if is_empty_paragraph(prev_p):
                if prev_idx - 1 >= body_start and paragraph_has_drawing(paragraphs[prev_idx - 1]):
                    remove_paragraph(prev_p)
                    changed = True
                    break

                while prev_idx - 1 >= body_start and is_empty_paragraph(paragraphs[prev_idx - 1]):
                    remove_paragraph(paragraphs[prev_idx - 1])
                    changed = True
                    break

                if changed:
                    break

                format_empty_paragraph(prev_p)
                continue

            new_p = insert_paragraph_after(prev_p, "")
            format_empty_paragraph(new_p)
            changed = True
            break

    return changed


def ensure_single_blank_before_figure_blocks(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs
        in_references = False

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text)

            if is_references_heading_text(text):
                in_references = True
                continue

            if in_references and is_appendix_heading_text(text):
                in_references = False

            if in_references or not paragraph_has_drawing(p) or idx <= body_start:
                continue

            prev_idx = idx - 1
            prev_p = paragraphs[prev_idx]

            if is_empty_paragraph(prev_p):
                while prev_idx - 1 >= body_start and is_empty_paragraph(paragraphs[prev_idx - 1]):
                    remove_paragraph(paragraphs[prev_idx - 1])
                    changed = True
                    break

                if changed:
                    break

                format_empty_paragraph(prev_p)
                continue

            prev_text = clean_spaces(prev_p.text)
            prev_kind = detect_kind_from_paragraph_object(prev_p, prev_text, prev_kind=None)
            if prev_kind in {
                "figure_caption",
                "table_caption",
                "table_continuation",
                "table_title",
                "source_line",
                "heading1",
                "heading2",
                "reference_subheading",
                "toc_heading",
            }:
                continue

            new_p = insert_paragraph_after(prev_p, "")
            format_empty_paragraph(new_p)
            changed = True
            break

    return changed


def remove_empty_between_figure_caption_and_source(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs
        in_references = False

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text)

            if is_references_heading_text(text):
                in_references = True
                continue

            if in_references and is_appendix_heading_text(text):
                in_references = False

            if in_references or not FIG_RE.match(text):
                continue

            j = idx + 1
            empty_paragraphs = []
            while j < len(paragraphs) and is_empty_paragraph(paragraphs[j]):
                empty_paragraphs.append(paragraphs[j])
                j += 1

            if not empty_paragraphs or j >= len(paragraphs):
                continue

            next_text = clean_spaces(paragraphs[j].text)
            if not FIG_SERVICE_LINE_RE.match(next_text):
                continue

            for blank in reversed(empty_paragraphs):
                remove_paragraph(blank)
            changed = True
            break

    return changed


def remove_empty_between_figure_source_and_caption(document, body_start):
    """Remove blank paragraphs between Источник/Примечание and the real figure caption."""
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs
        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue
            text = clean_spaces(p.text)
            if not FIG_SERVICE_LINE_RE.match(text):
                continue
            j = idx + 1
            empty_paragraphs = []
            while j < len(paragraphs) and is_empty_paragraph(paragraphs[j]):
                empty_paragraphs.append(paragraphs[j])
                j += 1
            if not empty_paragraphs or j >= len(paragraphs):
                continue
            if not _is_figure_caption_text(clean_spaces(paragraphs[j].text)):
                continue
            for blank in reversed(empty_paragraphs):
                remove_paragraph(blank)
            changed = True
            break
    return changed


def ensure_figure_block_keep_with_next(document, body_start):
    """Keep IMG → optional source/note → real Рис. caption together across a page break.

    Walks from image paragraphs only; never starts from "Источник:" alone, so table
    sources are not affected. Caption itself stays keep_with_next=False (its keep_together
    flag, set by format_figure_caption, prevents the caption's own lines from splitting).
    """
    if body_start is None:
        return False

    changed = False
    paragraphs = document.paragraphs
    n = len(paragraphs)

    for idx, p in enumerate(paragraphs):
        if idx < body_start:
            continue
        if not paragraph_has_drawing(p):
            continue

        chain_keepnext = [p]
        j = idx + 1
        while j < n:
            nxt = paragraphs[j]
            if is_empty_paragraph(nxt):
                j += 1
                continue
            text = clean_spaces(nxt.text)
            if paragraph_has_drawing(nxt):
                break
            if _is_figure_service_text(text):
                chain_keepnext.append(nxt)
                j += 1
                continue
            if _is_figure_caption_text(text):
                # Confirmed figure block: apply keep_with_next to every link before the caption.
                for link in chain_keepnext:
                    if link.paragraph_format.keep_with_next is not True:
                        link.paragraph_format.keep_with_next = True
                        changed = True
                break
            break

    return changed


def _is_figure_caption_text(text: str) -> bool:
    match = FIG_RE.match(clean_spaces(text))
    if not match:
        return False
    title = clean_spaces(match.group(3) or "")
    if re.match(
        r"^(показыва|отража|содерж|представлен|представля|демонстрир|иллюстрир)\w*\b",
        title,
        re.IGNORECASE,
    ):
        return False
    return True


def _is_figure_service_text(text: str) -> bool:
    return bool(FIG_SERVICE_LINE_RE.match(clean_spaces(text)))


def _is_figure_note_text(text: str) -> bool:
    return bool(re.match(r"^\s*примечание\s*:", clean_spaces(text), re.IGNORECASE))


def _is_paragraph_xml_with_image(child) -> bool:
    return bool(
        child.xpath(
            ".//*[local-name()='drawing' or local-name()='pict' or local-name()='object']"
        )
    )


def _figure_metadata_segments(text: str) -> list[str]:
    segments = [clean_spaces(part) for part in re.split(r"[\n\v]+", text or "") if clean_spaces(part)]
    if len(segments) < 2:
        return []
    if all(_is_figure_service_text(segment) or _is_figure_caption_text(segment) for segment in segments):
        return segments
    return []


def _split_figure_metadata_soft_break_paragraph(paragraph) -> bool:
    segments = _figure_metadata_segments(paragraph.text)
    if not segments:
        return False

    replace_paragraph_text(paragraph, segments[0])
    prev = paragraph
    for segment in segments[1:]:
        prev = insert_paragraph_after(prev, segment)
    return True


def _is_figure_hard_boundary_paragraph(paragraph, text: str) -> bool:
    if paragraph_has_drawing(paragraph):
        return True
    if TABLE_NUM_RE.match(text):
        return True
    if is_table_continuation_text(text):
        return True
    if parse_heading1(text) or parse_heading2(text) or parse_broken_heading2(text):
        return True
    if is_references_heading_text(text):
        return True
    if text.lower() == "приложения":
        return True
    if normalize_appendix_start_label_text(text):
        return True
    if is_appendix_continuation_label_text(text):
        return True
    return False


def _figure_neighbor_paragraphs(document, image_child_idx: int, body_start: int, *, direction: int, max_blocks: int):
    children = list(document.element.body)
    paragraph_lookup = _paragraph_lookup(document)
    blocks = []
    idx = image_child_idx + direction

    while 0 <= idx < len(children) and len(blocks) < max_blocks:
        child = children[idx]

        if child.tag == qn("w:tbl"):
            return blocks, True

        if child.tag != qn("w:p"):
            return blocks, True

        if _is_empty_paragraph_xml(child):
            idx += direction
            continue

        info = paragraph_lookup.get(child)
        if info is None:
            return blocks, False

        paragraph_idx, paragraph = info
        if paragraph_idx < body_start:
            return blocks, False

        text = clean_spaces(paragraph.text)
        if _is_paragraph_xml_with_image(child):
            return blocks, False

        if _is_figure_service_text(text) or _is_figure_caption_text(text) or _figure_metadata_segments(paragraph.text):
            blocks.append(paragraph)
            idx += direction
            continue

        if _is_figure_hard_boundary_paragraph(paragraph, text):
            return blocks, True

        return blocks, True

    if len(blocks) >= max_blocks:
        peek_idx = idx
        while 0 <= peek_idx < len(children):
            child = children[peek_idx]
            if child.tag == qn("w:tbl"):
                break
            if child.tag != qn("w:p"):
                break
            if _is_empty_paragraph_xml(child):
                peek_idx += direction
                continue
            info = paragraph_lookup.get(child)
            if info is None:
                break
            paragraph_idx, paragraph = info
            if paragraph_idx < body_start:
                break
            text = clean_spaces(paragraph.text)
            if _is_paragraph_xml_with_image(child):
                return blocks, False
            if _is_figure_service_text(text) or _is_figure_caption_text(text) or _figure_metadata_segments(paragraph.text):
                return blocks, False
            break

    if direction < 0:
        blocks.reverse()
    return blocks, True


def _collect_figure_block_around_image(document, image_paragraph, body_start):
    paragraph_info = _paragraph_lookup(document).get(image_paragraph._p)
    if paragraph_info is None:
        return None

    paragraph_idx, _paragraph = paragraph_info
    if paragraph_idx < body_start:
        return None

    children = list(document.element.body)
    try:
        image_child_idx = children.index(image_paragraph._p)
    except ValueError:
        return None

    before, before_clear = _figure_neighbor_paragraphs(
        document,
        image_child_idx,
        body_start,
        direction=-1,
        max_blocks=2,
    )
    after, after_clear = _figure_neighbor_paragraphs(
        document,
        image_child_idx,
        body_start,
        direction=1,
        max_blocks=3,
    )
    if not before_clear or not after_clear:
        return None

    metadata = before + after
    if not metadata:
        return None

    captions = [p for p in metadata if _is_figure_caption_text(p.text)]
    services = [p for p in metadata if _is_figure_service_text(p.text)]
    split_candidates = [p for p in metadata if _figure_metadata_segments(p.text)]

    if split_candidates:
        if len(split_candidates) > 1:
            return None
        return {"split": split_candidates[0]}

    if len(captions) != 1:
        return None

    source_like = [p for p in services if not _is_figure_note_text(p.text)]
    notes = [p for p in services if _is_figure_note_text(p.text)]
    if len(source_like) > 1 or len(notes) > 1:
        return None

    desired = source_like + notes + captions
    current_after = [p for p in after if p in desired]
    already_canonical = not before and current_after == desired
    return {
        "split": None,
        "desired": desired,
        "already_canonical": already_canonical,
    }


def normalize_figure_blocks(document, body_start):
    if body_start is None:
        return False

    for paragraph in list(document.paragraphs):
        if not paragraph_has_drawing(paragraph):
            continue

        block = _collect_figure_block_around_image(document, paragraph, body_start)
        if not block:
            continue

        split_candidate = block.get("split")
        if split_candidate is not None:
            return _split_figure_metadata_soft_break_paragraph(split_candidate)

        if block.get("already_canonical"):
            for service in [p for p in block["desired"] if _is_figure_service_text(p.text)]:
                format_source_line(service)
            caption = block["desired"][-1]
            normalize_figure_caption_text(caption)
            format_figure_caption(caption)
            continue

        anchor = paragraph._p
        for metadata_paragraph in block["desired"]:
            anchor.addnext(metadata_paragraph._p)
            anchor = metadata_paragraph._p

        for service in [p for p in block["desired"] if _is_figure_service_text(p.text)]:
            format_source_line(service)
        caption = block["desired"][-1]
        normalize_figure_caption_text(caption)
        format_figure_caption(caption)
        return True

    return False


def reorder_figure_source_before_caption(document, body_start):
    return normalize_figure_blocks(document, body_start)


def ensure_empty_between_heading1_and_heading2(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs
        prev_kind = None

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            kind = classify_paragraph(clean_spaces(p.text), prev_kind=prev_kind)

            if kind == "heading1" and idx + 2 < len(paragraphs):
                next_p = paragraphs[idx + 1]
                next2_p = paragraphs[idx + 2]
                next2_kind = classify_paragraph(clean_spaces(next2_p.text), prev_kind="empty_paragraph")

                if is_empty_paragraph(next_p) and next2_kind == "heading2":
                    remove_paragraph(next_p)
                    changed = True
                    break

            prev_kind = kind


def ensure_compact_heading2_spacing(document, body_start):
    """
    Normalize spacing around heading2 in one pass.

    Rules:
      - exactly one empty paragraph immediately before heading2,
        except when the previous non-empty paragraph is heading1;
      - exactly one empty paragraph immediately after heading2.

    Returns True if any changes were made, otherwise False.
    """
    paragraphs = document.paragraphs
    prev_kind = None
    changed = False
    idx = max(body_start, 0)
    in_references = False

    while idx < len(paragraphs):
        p = paragraphs[idx]
        text = clean_spaces(p.text)

        if is_references_heading_text(text):
            in_references = True
            prev_kind = "heading1"
            idx += 1
            continue

        if in_references and is_appendix_heading_text(text):
            in_references = False

        # Внутри списка источников этот проход не должен ничего вставлять/удалять
        if in_references:
            prev_kind = "body_text"
            idx += 1
            continue

        kind = classify_paragraph(text, prev_kind=prev_kind)

        if kind != "heading2":
            prev_kind = kind
            idx += 1
            continue

        # Ищем предыдущий непустой абзац
        prev_nonempty_idx = idx - 1
        while prev_nonempty_idx >= body_start and is_empty_paragraph(paragraphs[prev_nonempty_idx]):
            prev_nonempty_idx -= 1

        prev_nonempty_kind = None
        if prev_nonempty_idx >= body_start:
            prev_nonempty_text = clean_spaces(paragraphs[prev_nonempty_idx].text)
            prev_nonempty_kind = classify_paragraph(prev_nonempty_text, prev_kind=None)

        # ПЕРЕД heading2:
        # - если сверху heading1 -> пустой строки быть не должно
        # - иначе должна быть ровно одна пустая строка
        if prev_nonempty_kind == "heading1":
            while idx - 1 >= body_start and is_empty_paragraph(paragraphs[idx - 1]):
                remove_paragraph(paragraphs[idx - 1])
                paragraphs = document.paragraphs
                idx -= 1
                changed = True
        else:
            if idx - 1 < body_start or not is_empty_paragraph(paragraphs[idx - 1]):
                new_p = OxmlElement("w:p")
                p._element.addprevious(new_p)
                format_heading2_spacing_paragraph(Paragraph(new_p, p._parent))
                paragraphs = document.paragraphs
                idx += 1
                changed = True

            while idx - 2 >= body_start and is_empty_paragraph(paragraphs[idx - 2]):
                remove_paragraph(paragraphs[idx - 2])
                paragraphs = document.paragraphs
                idx -= 1
                changed = True

            if idx - 1 >= body_start and is_empty_paragraph(paragraphs[idx - 1]):
                format_heading2_spacing_paragraph(paragraphs[idx - 1])

        # ПОСЛЕ heading2:
        # всегда ровно одна пустая строка
        if idx + 1 >= len(paragraphs) or not is_empty_paragraph(paragraphs[idx + 1]):
            new_p = OxmlElement("w:p")
            p._element.addnext(new_p)
            paragraphs = document.paragraphs
            changed = True

        while idx + 2 < len(paragraphs) and is_empty_paragraph(paragraphs[idx + 2]):
            remove_paragraph(paragraphs[idx + 2])
            paragraphs = document.paragraphs
            changed = True

        if idx + 1 < len(paragraphs) and is_empty_paragraph(paragraphs[idx + 1]):
            hard_reset_paragraph_format(paragraphs[idx + 1], first_line_indent_cm=None)

        prev_kind = kind
        idx += 1

    return changed



STRUCTURAL_HEADING_TEXTS_V2 = {
    "ВВЕДЕНИЕ",
    "ЗАКЛЮЧЕНИЕ",
}

def normalize_structural_heading_spacing_v2(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text).upper()
            if text not in STRUCTURAL_HEADING_TEXTS_V2:
                continue

            # Сразу после ВВЕДЕНИЕ / ЗАКЛЮЧЕНИЕ должна быть ровно одна пустая строка
            if idx + 1 >= len(paragraphs):
                new_p = insert_paragraph_after(p, "")
                hard_reset_paragraph_format(new_p, first_line_indent_cm=None)
                changed = True
                break

            next_p = paragraphs[idx + 1]

            if not is_empty_paragraph(next_p):
                new_p = insert_paragraph_after(p, "")
                hard_reset_paragraph_format(new_p, first_line_indent_cm=None)
                changed = True
                break

            # Если пустых строк больше одной — сжимаем до одной
            if idx + 2 < len(paragraphs) and is_empty_paragraph(paragraphs[idx + 2]):
                remove_paragraph(paragraphs[idx + 2])
                changed = True
                break

            # Нормализуем единственную пустую строку
            hard_reset_paragraph_format(next_p, first_line_indent_cm=None)


def ensure_empty_before_table_caption(document, body_start):
    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs
        prev_kind = None

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            kind = _classify_paragraph_with_table_adjacency(
                document,
                p,
                body_start,
                prev_kind=prev_kind,
            )

            if kind in {"table_caption", "table_continuation"}:
                if idx - 1 >= body_start:
                    prev_p = paragraphs[idx - 1]
                    if not is_empty_paragraph(prev_p):
                        new_p = OxmlElement("w:p")
                        prev_p._element.addnext(new_p)
                        changed = True
                        break
                    if idx - 2 >= body_start and is_empty_paragraph(paragraphs[idx - 2]):
                        remove_paragraph(prev_p)
                        changed = True
                        break

            prev_kind = kind


def remove_extra_empty_after_service_lines(document, body_start):
    target_kinds = {
        "table_caption",
        "table_title",
        "table_continuation",
        "reference_subheading",
    }

    changed = True
    while changed:
        changed = False
        paragraphs = document.paragraphs
        prev_kind = None

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            kind = _classify_paragraph_with_table_adjacency(
                document,
                p,
                body_start,
                prev_kind=prev_kind,
            )

            if kind in target_kinds:
                if idx + 1 < len(paragraphs) and is_empty_paragraph(paragraphs[idx + 1]):
                    remove_paragraph(paragraphs[idx + 1])
                    changed = True
                    break

            prev_kind = kind


def cleanup_reference_subheadings_layout(document, body_start):
    changed = True

    while changed:
        changed = False
        in_references = False
        paragraphs = document.paragraphs

        for idx, p in enumerate(paragraphs):
            if idx < body_start:
                continue

            text = clean_spaces(p.text)
            low = text.lower()

            if low in {
                "список использованных источников",
                "список использованной литературы",
            }:
                in_references = True
                continue

            if not in_references:
                continue

            if low in {"приложения", "приложение"}:
                in_references = False
                continue

            canonical = canonical_reference_block_heading_paragraph(p)
            if canonical:
                replace_paragraph_text(p, canonical)
                remove_paragraph_numbering(p)
                p.paragraph_format.page_break_before = False
                format_reference_subheading(p)

                if idx - 1 < body_start or not is_reference_spacing_paragraph(paragraphs[idx - 1]):
                    new_el = OxmlElement("w:p")
                    p._element.addprevious(new_el)
                    format_empty_paragraph(Paragraph(new_el, p._parent))
                    changed = True
                    break

                format_empty_paragraph(paragraphs[idx - 1])
                if idx - 2 >= body_start and is_reference_spacing_paragraph(paragraphs[idx - 2]):
                    remove_paragraph(paragraphs[idx - 2])
                    changed = True
                    break

                if idx + 1 < len(paragraphs) and is_empty_paragraph(paragraphs[idx + 1]):
                    remove_paragraph(paragraphs[idx + 1])
                    changed = True
                    break
                    
def format_empty_paragraph(paragraph):
    set_paragraph_style_safe(paragraph, "Normal", "Обычный")
    clear_paragraph_outline_level(paragraph)
    remove_paragraph_numbering(paragraph)

    hard_reset_paragraph_format(paragraph, first_line_indent_cm=None)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    ensure_empty_run(paragraph)
    for run in paragraph.runs:
        set_run_font(run, size_pt=BODY_FONT_SIZE_PT, bold=False, italic=False, all_caps=False)


def format_heading2_spacing_paragraph(paragraph):
    format_empty_paragraph(paragraph)
    # Phase 3 removes empty first lines at page top unless they are meaningful
    # spacers. A tiny space_before marks this as intentional while preserving
    # the visual one-blank-line layout before Heading 2.
    paragraph.paragraph_format.space_before = Pt(3)
    
def normalize_sections(document):
    """
    Удаляет секционные разрывы из абзацев внутри документа,
    чтобы потом можно было заново поставить один правильный
    разрыв секции перед ВВЕДЕНИЕМ.
    """
    for p in document.paragraphs:
        pPr = p._element.pPr
        if pPr is None:
            continue

        sectPr = pPr.find(qn("w:sectPr"))
        if sectPr is not None:
            pPr.remove(sectPr)

def ensure_section_break_before_introduction(document, body_start):
    """
    Ставит разрыв секции типа Next Page перед абзацем 'ВВЕДЕНИЕ'.
    Это нужно, чтобы:
    - 1-я страница имела свой футер,
    - 2-я страница была пустой,
    - с 3-й страницы можно было включить нумерацию отдельной секцией.
    """
    if body_start is None:
        return

    paragraphs = document.paragraphs
    if body_start <= 0 or body_start >= len(paragraphs):
        return

    intro_p = paragraphs[body_start]
    prev_p = paragraphs[body_start - 1]

    prev_pPr = prev_p._element.get_or_add_pPr()

    # Если секционный разрыв уже стоит — второй раз не добавляем
    existing_sectPr = prev_pPr.find(qn("w:sectPr"))
    if existing_sectPr is not None:
        return

    next_pPr = intro_p._element.pPr
    if next_pPr is not None and next_pPr.find(qn("w:sectPr")) is not None:
        return

    body = document._body._element
    body_sectPr = body.sectPr
    if body_sectPr is None:
        return

    new_sectPr = deepcopy(body_sectPr)

    # Делаем разрыв секции "со следующей страницы"
    type_el = new_sectPr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        new_sectPr.insert(0, type_el)
    type_el.set(qn("w:val"), "nextPage")

    prev_pPr.append(new_sectPr)

def _append_next_page_section_break_after(paragraph, body_sectpr):
    pPr = paragraph._element.get_or_add_pPr()

    old = pPr.find(qn("w:sectPr"))
    if old is not None:
        pPr.remove(old)

    new_sectPr = deepcopy(body_sectpr)

    # Не тащим старые ссылки на футеры/заголовки и старый старт нумерации
    for tag in ("w:pgNumType", "w:footerReference", "w:headerReference"):
        for el in list(new_sectPr.findall(qn(tag))):
            new_sectPr.remove(el)

    type_el = new_sectPr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        new_sectPr.insert(0, type_el)
    type_el.set(qn("w:val"), "nextPage")

    pPr.append(new_sectPr)


def ensure_front_matter_layout(document, body_start):
    """
    Целевая модель:
    если есть содержание:
        секция 1 = титул
        секция 2 = содержание
        секция 3 = введение и далее
    если содержания нет:
        секция 1 = титул
        секция 2 = введение и далее
    """
    if body_start is None or body_start <= 0:
        return

    paragraphs = document.paragraphs
    if body_start >= len(paragraphs):
        return

    body = document._body._element
    body_sectpr = body.sectPr
    if body_sectpr is None:
        return

    # Полная очистка page-break артефактов до введения
    for i in range(body_start):
        p = paragraphs[i]
        p.paragraph_format.page_break_before = False

        for run in p.runs:
            r = run._element
            for br in list(r.findall(qn("w:br"))):
                br_type = br.get(qn("w:type"))
                if br_type == "page":
                    r.remove(br)

    # Ищем содержание до введения
    contents_idx = None
    for i in range(body_start):
        t = clean_spaces(paragraphs[i].text).upper()
        if ("СОДЕРЖАН" in t) or ("ОГЛАВЛЕН" in t):
            contents_idx = i
            break

    # На самом абзаце введения обычный page break не нужен
    paragraphs[body_start].paragraph_format.page_break_before = False

    if contents_idx is not None and contents_idx > 0:
        # титул -> содержание
        _append_next_page_section_break_after(paragraphs[contents_idx - 1], body_sectpr)
        # содержание -> введение
        _append_next_page_section_break_after(paragraphs[body_start - 1], body_sectpr)
    else:
        # титул -> введение
        _append_next_page_section_break_after(paragraphs[body_start - 1], body_sectpr)


def ensure_appendices_section_layout(document, body_start):
    if body_start is None:
        return

    paragraphs = document.paragraphs
    body = document._body._element
    body_sectpr = body.sectPr
    if body_sectpr is None:
        return

    appendices_idx = None
    for idx, paragraph in enumerate(paragraphs):
        if idx < body_start:
            continue
        if clean_spaces(paragraph.text).lower() == "приложения":
            appendices_idx = idx
            break

    if appendices_idx is None or appendices_idx <= 0:
        return

    appendices_paragraph = paragraphs[appendices_idx]
    appendices_paragraph.paragraph_format.page_break_before = False
    _append_next_page_section_break_after(paragraphs[appendices_idx - 1], body_sectpr)

def remove_all_italic(doc):
    """
    Убирает курсив, highlight, цвет текста и XML-заливку из всего документа.
    """

    def clear_run(run):
        run.italic = False

        try:
            run.font.highlight_color = None
        except Exception:
            pass

        try:
            run.font.color.rgb = RGBColor(0, 0, 0)
        except Exception:
            pass

        rPr = run._element.get_or_add_rPr()

        for tag in ("w:highlight", "w:shd"):
            node = rPr.find(qn(tag))
            if node is not None:
                rPr.remove(node)

        color = rPr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rPr.append(color)
        color.set(qn("w:val"), "000000")

        for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
            qname = qn(attr)
            if qname in color.attrib:
                del color.attrib[qname]

    for p in doc.paragraphs:
        for r in p.runs:
            clear_run(r)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        clear_run(r)


# ── Footnote standardization ──────────────────────────────────────────────────

import lxml.etree as _lxml_etree

_FOOTNOTES_RTYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
)
# Separator footnotes (id ≤ 0) contain Word-internal separator lines — skip them.
_FOOTNOTE_SKIP_IDS = {"-1", "0"}


def _fn_ensure(parent, tag: str):
    """Return existing child with Clark-notation *tag*, or create and append one."""
    el = parent.find(tag)
    if el is None:
        el = _lxml_etree.SubElement(parent, tag)
    return el


def _format_footnote_para(p_elem) -> None:
    """
    Apply KFU footnote style to a single <w:p> XML element:
      • Times New Roman 10 pt, not bold
      • Single line spacing (240/240), no space_before/space_after
      • Zero indent (left, right, firstLine, hanging)

    Operates directly on the XML — safe to call on footnote paragraphs
    that are not exposed as python-docx Paragraph objects.
    Hyperlink runs (inside w:hyperlink) are formatted too.

    Uses lxml.etree.SubElement directly (instead of OxmlElement) because
    OxmlElement requires namespace-prefixed tags ("w:b") whereas qn()
    returns Clark notation ("{...}b").
    """
    # ── Paragraph-level properties ────────────────────────────────────────
    pPr = _fn_ensure(p_elem, qn("w:pPr"))
    # Keep pPr as first child
    if list(p_elem)[0] is not pPr:
        p_elem.remove(pPr)
        p_elem.insert(0, pPr)

    # Indent → 0
    ind = _fn_ensure(pPr, qn("w:ind"))
    for attr in (qn("w:left"), qn("w:right"), qn("w:firstLine"), qn("w:hanging")):
        ind.set(attr, "0")

    # Spacing → single (line=240, lineRule=auto), no before/after
    spacing = _fn_ensure(pPr, qn("w:spacing"))
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")

    # ── Run-level properties (all w:r descendants, including inside hyperlinks) ─
    for r_elem in p_elem.findall(".//" + qn("w:r")):
        rPr = _fn_ensure(r_elem, qn("w:rPr"))
        # Keep rPr as first child of run
        if list(r_elem)[0] is not rPr:
            r_elem.remove(rPr)
            r_elem.insert(0, rPr)

        # Font → Times New Roman
        rFonts = _fn_ensure(rPr, qn("w:rFonts"))
        for attr in (qn("w:ascii"), qn("w:hAnsi"), qn("w:cs")):
            rFonts.set(attr, "Times New Roman")

        # Size → 10 pt (20 half-points)
        for tag in (qn("w:sz"), qn("w:szCs")):
            el = _fn_ensure(rPr, tag)
            el.set(qn("w:val"), "20")

        # Bold → suppress
        for tag in (qn("w:b"), qn("w:bCs")):
            el = _fn_ensure(rPr, tag)
            el.set(qn("w:val"), "0")


def _format_footnotes(doc: Document) -> int:
    """
    Standardise all footnotes in the document to KFU footnote style:
    10 pt Times New Roman, single spacing, no bold, zero indent.

    Separator footnotes (id ≤ 0) are left untouched.
    Returns the number of footnote paragraphs formatted.

    Works with both XmlPart (has ._element) and plain Part (has ._blob),
    since python-docx may load the footnotes part as either depending on version.
    """
    try:
        footnotes_part = doc.part.part_related_by(_FOOTNOTES_RTYPE)
    except KeyError:
        return 0   # document has no footnotes

    # XmlPart path (python-docx >= 0.8.x typically)
    if hasattr(footnotes_part, "_element"):
        fn_root = footnotes_part._element
        need_serialize = False
    elif hasattr(footnotes_part, "_blob") and footnotes_part._blob:
        # Plain Part — parse the raw XML blob, modify, re-serialize
        fn_root = _lxml_etree.fromstring(footnotes_part._blob)
        need_serialize = True
    else:
        return 0

    count = 0
    for fn_elem in fn_root.findall(qn("w:footnote")):
        fn_id = fn_elem.get(qn("w:id"), "")
        if fn_id in _FOOTNOTE_SKIP_IDS:
            continue
        for p_elem in fn_elem.findall(".//" + qn("w:p")):
            _format_footnote_para(p_elem)
            count += 1

    if need_serialize:
        footnotes_part._blob = _lxml_etree.tostring(
            fn_root,
            xml_declaration=True,
            encoding="UTF-8",
            standalone=True,
        )

    return count


_PLAIN_CONTENTS_HEADING_RE = re.compile(
    r"^\s*(содержание|оглавление)\s*[.:;]?\s*[.․‥…·•\s]*\d{0,4}\s*$",
    re.IGNORECASE,
)
_TOC_ENTRY_PAGE_TAIL_RE = re.compile(r"[\s.․‥…·•]+\d{1,4}\s*$")
_TOC_ENTRY_APPENDIX_RE = re.compile(
    r"^приложение\s+(?:\d{1,3}|[a-zа-яё])\b", re.IGNORECASE
)


def _is_plain_contents_heading_paragraph(text: str) -> bool:
    cleaned = clean_spaces(text or "")
    if not cleaned:
        return False
    first = re.split(r"[\n\v]+", cleaned, maxsplit=1)[0].strip()
    return bool(_PLAIN_CONTENTS_HEADING_RE.match(first))


def _looks_like_toc_entry_text(text: str) -> bool:
    t = clean_spaces(text or "").strip()
    if not t:
        return False
    low_raw = t.lower().rstrip(".").strip()
    # Check appendix label BEFORE stripping page tail: "ПРИЛОЖЕНИЕ 1" would be
    # mistakenly reduced to "ПРИЛОЖЕНИЕ" by the page-tail RE (treating " 1" as
    # a page number), causing the appendix regex to reject it.
    if _TOC_ENTRY_APPENDIX_RE.match(low_raw):
        return True
    t = _TOC_ENTRY_PAGE_TAIL_RE.sub("", t).strip()
    if not t:
        return False
    low = t.lower().rstrip(".").strip()
    if low in {"введение", "заключение", "приложения"}:
        return True
    if low.startswith("список использованных") or low.startswith("список использованной"):
        return True
    if _TOC_ENTRY_APPENDIX_RE.match(low):
        return True
    if re.match(r"^\d+\.\d+\.?\s+\S", t):
        return True
    if re.match(r"^\d+\.\s+\S", t):
        return True
    return False


def _paragraph_is_heading_styled(paragraph) -> bool:
    try:
        style_name = (paragraph.style.name or "").strip().lower()
    except Exception:
        style_name = ""
    if style_name in {
        "heading 1", "heading 2", "heading 3",
        "заголовок 1", "заголовок 2", "заголовок 3",
    }:
        return True
    pPr = paragraph._element.pPr
    if pPr is not None and pPr.find(qn("w:outlineLvl")) is not None:
        return True
    return False


def find_real_body_start_index(document):
    """
    Find the real body ВВЕДЕНИЕ paragraph, skipping past entries of an old
    plain-text TOC where a 'ВВЕДЕНИЕ' line appears as a TOC entry rather
    than as the real body intro.

    Strategy (most specific first):

    1. Among all standalone 'ВВЕДЕНИЕ' paragraphs, prefer one that is
       styled as a Heading (Heading 1/2/3 / Заголовок 1/2/3) or carries
       a `w:outlineLvl`. A styled intro is an unambiguous real-body
       signal and overrides any plain-text TOC entries before it.

    2. Otherwise, if a standalone 'Содержание' / 'Оглавление' paragraph
       (with optional trailing punctuation / leaders / page number)
       exists earlier in the document AND at least one TOC-like entry
       sits between it and a later 'ВВЕДЕНИЕ', the later 'ВВЕДЕНИЕ' is
       the real body intro — return the last such candidate after the
       last contents heading.

    3. Otherwise fall back to `find_body_start_index` — the legacy
       first-match behaviour, preserved for documents without any TOC
       structure.
    """
    paragraphs = document.paragraphs
    intros: list[int] = [
        idx
        for idx, p in enumerate(paragraphs)
        if is_intro_heading_text(paragraph_text(p))
    ]
    if not intros:
        return find_body_start_index(document)

    styled_intros = [idx for idx in intros if _paragraph_is_heading_styled(paragraphs[idx])]
    if styled_intros:
        return styled_intros[-1]

    contents_indices = [
        idx
        for idx, p in enumerate(paragraphs)
        if _is_plain_contents_heading_paragraph(paragraph_text(p))
    ]
    if contents_indices:
        last_contents_idx = max(contents_indices)
        # Require at least one TOC-like entry between the contents heading and
        # a later ВВЕДЕНИЕ before we treat it as the real intro. Without that
        # evidence, a standalone Содержание that did not lead a real TOC must
        # not displace the legacy first-match rule.
        for j in range(last_contents_idx + 1, len(paragraphs)):
            text = paragraph_text(paragraphs[j])
            if not clean_spaces(text):
                continue
            if _looks_like_toc_entry_text(text):
                candidates_after = [i for i in intros if i > last_contents_idx]
                if candidates_after:
                    return candidates_after[-1]
                break

    return intros[0]


def process_document(input_path: Path, output_path: Path):
    doc = Document(str(input_path))

    # Сразу чистим визуальный мусор по всему документу
    remove_all_italic(doc)
    set_section_margins(doc)

    body_start = find_real_body_start_index(doc)
    if body_start is None:
        raise RuntimeError("Не найден заголовок 'Введение'; файл пропущен из соображений безопасности.")

    toc_h1_map, toc_h2_map = build_toc_heading_maps(doc, body_start)

    split_manual_dash_lists(doc, body_start)
    split_table_captions_prepass(doc, body_start)
    normalize_quotes_in_document(doc, body_start or 0)
    normalize_dashes_in_document(doc, body_start)
    normalize_yo_in_document(doc, body_start)
    normalize_semicolons_in_document(doc, body_start)
    normalize_citations_in_document(doc, body_start)
    normalize_plain_lists_in_document(doc, body_start)
    normalize_word_numbered_lists_in_document(doc, body_start)
    run_with_pass_limit(
        "split_body_structural_soft_breaks",
        split_body_structural_soft_breaks,
        doc,
        body_start,
    )
    # Преднормализация только тела работы; содержание не трогаем
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < body_start:
            continue
        normalize_simple_paragraph_spaces(paragraph)
        normalize_heading2_artifacts(paragraph)

    paragraphs = doc.paragraphs
    prev_kind = None
    current_chapter_num = None
    next_paragraph_num = None
    in_references = False

    # Основной проход по телу документа
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < body_start:
            continue

        text = clean_spaces(paragraph.text)
        if is_references_heading_text(text):
            in_references = True
        elif in_references and is_appendix_heading_text(text):
            in_references = False

        if not text:
            prev_kind = "empty_paragraph"
            continue

        if in_references and not is_references_heading_text(text):
            canonical = canonical_reference_block_heading_paragraph(paragraph)
            if canonical:
                replace_paragraph_text(paragraph, canonical)
                format_reference_subheading(paragraph)
                prev_kind = "reference_subheading"
            else:
                format_reference_entry(paragraph)
                prev_kind = "body_text"
            continue

        text = strip_leading_heading_garbage(text)
        if text != clean_spaces(paragraph.text):
            replace_paragraph_text(paragraph, text)

        kind = detect_kind_from_paragraph_object(paragraph, text, prev_kind=prev_kind)
        if kind == "table_caption" and not _is_confirmed_table_caption_paragraph(doc, paragraph, body_start):
            kind = "body_text"
        if kind == "figure_caption":
            m_fig_main = FIG_RE.match(text)
            if m_fig_main and caption_tail_is_reference_prose(m_fig_main.group(3) or ""):
                kind = "body_text"
        prev_paragraph_obj = doc.paragraphs[idx - 1] if idx - 1 >= body_start else None
        is_body_list_item = is_probable_body_list_item(
            paragraph,
            prev_paragraph=prev_paragraph_obj,
            prev_kind=prev_kind,
        )
        if is_body_list_item:
            kind = "body_list_item"

        parsed_h1 = parse_heading1(text)
        if parsed_h1:
            if parsed_h1["kind"] == "heading1_chapter":
                toc_text = toc_h1_map.get(parsed_h1["chapter_num"])
                if not is_heading1_promotion_safe(paragraph, parsed_h1, toc_text=toc_text):
                    parsed_h1 = None
                    if kind == "heading1":
                        kind = "body_text"
                else:
                    current_text = f'{parsed_h1["chapter_num"]}. {parsed_h1["title"]}'

                    if toc_text and len(current_text) < len(toc_text):
                        replace_paragraph_text(paragraph, toc_text)
                        text = clean_spaces(paragraph.text)
                        parsed_h1 = parse_heading1(text)

                    current_chapter_num = parsed_h1["chapter_num"]
                    next_paragraph_num = 1
                    smart_repair_heading1(paragraph, text)
                    kind = "heading1"

            elif parsed_h1["kind"] == "heading1_exact":
                current_chapter_num = None
                next_paragraph_num = None
                smart_repair_heading1(paragraph, text)
                kind = "heading1"

        parsed_h2_existing = parse_heading2(text)
        if parsed_h2_existing:
            toc_text = toc_h2_map.get(
                (parsed_h2_existing["chapter_num"], parsed_h2_existing["paragraph_num"])
            )
            current_text = (
                f'{parsed_h2_existing["chapter_num"]}.'
                f'{parsed_h2_existing["paragraph_num"]}. '
                f'{parsed_h2_existing["title"]}'
            )

            if toc_text and len(current_text) < len(toc_text):
                replace_paragraph_text(paragraph, toc_text)
                text = clean_spaces(paragraph.text)
                kind = "heading2"

        if kind == "broken_heading2":
            repaired = smart_repair_broken_heading2(
                paragraph,
                current_chapter_num,
                next_paragraph_num,
            )
            if repaired:
                text = clean_spaces(paragraph.text)
                kind = "heading2"

        if kind not in {"table_continuation", "body_list_item"} and (
            kind == "heading2"
            or auto_detect_heading2(
                paragraph,
                current_chapter_num,
                next_paragraph_num,
                prev_kind,
            )
            or is_likely_numbered_heading2_candidate(
                paragraph,
                current_chapter_num,
                next_paragraph_num,
                prev_kind=prev_kind,
            )
        ):
            normalized_text = normalize_heading2_numbering(
                paragraph,
                current_chapter_num,
                next_paragraph_num,
            )
            if normalized_text:
                kind = "heading2"
                parsed_h2 = parse_heading2(clean_spaces(paragraph.text))
                if parsed_h2:
                    current_chapter_num = parsed_h2["chapter_num"]
                    next_paragraph_num = parsed_h2["paragraph_num"] + 1

        if kind not in {
            "heading1",
            "heading2",
            "body_list_item",
            "table_caption",
            "table_continuation",
            "table_title",
            "figure_caption",
            "source_line",
            "reference_subheading",
        }:
            if auto_detect_numbered_heading1(
                paragraph,
                current_chapter_num=current_chapter_num,
                next_paragraph=doc.paragraphs[idx + 1] if idx + 1 < len(doc.paragraphs) else None,
            ):
                inferred_chapter_num = 1 if current_chapter_num is None else current_chapter_num + 1
                heading_text = clean_spaces(paragraph.text)
                replace_paragraph_text(paragraph, f"{inferred_chapter_num}. {heading_text}")
                kind = "heading1"
                current_chapter_num = inferred_chapter_num
                next_paragraph_num = 1

        if kind == "heading1" and paragraph_has_numbering(paragraph):
            text_without_number = clean_spaces(paragraph.text)
            if (
                text_without_number
                and not parse_heading1(text_without_number)
                and auto_detect_numbered_heading1(
                    paragraph,
                    current_chapter_num=current_chapter_num,
                    next_paragraph=doc.paragraphs[idx + 1] if idx + 1 < len(doc.paragraphs) else None,
                )
            ):
                inferred_chapter_num = 1 if current_chapter_num is None else current_chapter_num + 1
                replace_paragraph_text(paragraph, f"{inferred_chapter_num}. {text_without_number}")
                current_chapter_num = inferred_chapter_num
                next_paragraph_num = 1

        if kind == "table_continuation":
            normalize_table_continuation_text(paragraph)

        if kind == "figure_caption":
            normalize_figure_caption_text(paragraph)

        if kind == "heading1":
            format_heading1(paragraph)

        elif kind == "heading2":
            format_heading2(paragraph)

        elif kind == "table_caption":
            format_table_caption(paragraph)

        elif kind == "table_continuation":
            format_table_caption(paragraph)

        elif kind == "table_title":
            format_table_title(paragraph)

        elif kind == "figure_caption":
            format_figure_caption(paragraph)

        elif kind == "source_line":
            format_source_line(paragraph)
        elif kind == "body_list_item":
            format_body_list_item(paragraph)
        elif kind == "reference_subheading":
            canonical = canonical_reference_block_heading_paragraph(paragraph)
            if canonical:
                replace_paragraph_text(paragraph, canonical)
            format_reference_subheading(paragraph)

        else:
            format_body(paragraph)

        prev_kind = kind

    format_tables(doc)
    ensure_all_table_rows_cant_split(doc)
    run_with_pass_limit(
        "center_image_paragraphs",
        center_image_paragraphs,
        doc,
        body_start,
    )
    convert_reference_numbering_to_plain_text(doc, body_start)

    run_with_pass_limit(
        "normalize_formula_blocks",
        normalize_formula_blocks,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "compact_references_block",
        compact_references_block,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "ensure_blank_before_reference_subheadings",
        ensure_blank_before_reference_subheadings,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "ensure_single_blank_after_references_heading",
        ensure_single_blank_after_references_heading,
        doc,
        body_start,
    )

    collapse_empty_paragraphs_in_body(doc.paragraphs, body_start)

    run_with_pass_limit(
        "ensure_compact_heading2_spacing",
        ensure_compact_heading2_spacing,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "ensure_empty_before_table_caption",
        ensure_empty_before_table_caption,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "normalize_figure_blocks",
        normalize_figure_blocks,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "ensure_single_blank_before_figure_blocks",
        ensure_single_blank_before_figure_blocks,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "ensure_single_blank_before_figure_captions",
        ensure_single_blank_before_figure_captions,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "remove_empty_between_figure_caption_and_source",
        remove_empty_between_figure_caption_and_source,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "remove_extra_empty_after_service_lines",
        remove_extra_empty_after_service_lines,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "ensure_empty_after_source_and_note",
        ensure_empty_after_source_and_note,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "remove_empty_between_figure_source_and_caption",
        remove_empty_between_figure_source_and_caption,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "cleanup_reference_subheadings_layout",
        cleanup_reference_subheadings_layout,
        doc,
        body_start,
    )

    collapse_empty_paragraphs_in_body(doc.paragraphs, body_start)

    run_with_pass_limit(
        "ensure_single_blank_after_headings",
        ensure_single_blank_after_headings,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "normalize_structural_heading_spacing_v2",
        normalize_structural_heading_spacing_v2,
        doc,
        body_start,
    )
    run_with_pass_limit(
        "remove_single_empty_between_body_paragraphs",
        remove_single_empty_between_body_paragraphs,
        doc,
        body_start,
    )

    # Финальный жёсткий проход:
    # добиваем заголовки, таблицы и обычный текст уже после всех структурных вставок/удалений
    prev_nonempty_kind = None
    _final_in_references = False
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx < body_start:
            continue

        if paragraph_has_drawing(paragraph):
            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            prev_nonempty_kind = "body_text"
            continue
        text = clean_spaces(paragraph.text)

        if is_references_heading_text(text):
            _final_in_references = True
        elif is_appendix_heading_text(text):
            _final_in_references = False

        if not text:
            format_empty_paragraph(paragraph)
            continue

        if _final_in_references and not is_references_heading_text(text):
            canonical = canonical_reference_block_heading_paragraph(paragraph)
            if canonical:
                replace_paragraph_text(paragraph, canonical)
                format_reference_subheading(paragraph)
                prev_nonempty_kind = "reference_subheading"
            else:
                format_reference_entry(paragraph)
                prev_nonempty_kind = "body_text"
            continue

        text = strip_leading_heading_garbage(text)
        if text != clean_spaces(paragraph.text):
            replace_paragraph_text(paragraph, text)
        if is_formula_paragraph_text(text):
            format_formula_paragraph(paragraph)
            prev_nonempty_kind = "formula"
            continue

        if is_formula_explanation_start(text):
            format_formula_explanation_paragraph(paragraph, is_first=True)
            prev_nonempty_kind = "formula_explanation"
            continue

        if prev_nonempty_kind in {"formula", "formula_explanation"} and is_formula_explanation_continuation(text):
            format_formula_explanation_paragraph(paragraph, is_first=False)
            prev_nonempty_kind = "formula_explanation"
            continue
        prev_paragraph_obj = doc.paragraphs[idx - 1] if idx - 1 >= body_start else None
        if is_probable_body_list_item(
            paragraph,
            prev_paragraph=prev_paragraph_obj,
            prev_kind=prev_nonempty_kind,
        ):
            format_body_list_item(paragraph)
            prev_nonempty_kind = "body_list_item"
            continue
        parsed_h1_final = parse_heading1(text)
        if parsed_h1_final and is_heading1_promotion_safe(paragraph, parsed_h1_final, toc_text=(
            toc_h1_map.get(parsed_h1_final["chapter_num"])
            if parsed_h1_final["kind"] == "heading1_chapter"
            else None
        )):
            smart_repair_heading1(paragraph, text)
            format_heading1(paragraph)
            prev_nonempty_kind = "heading1"
            continue

        if parse_heading2(text):
            format_heading2(paragraph)
            prev_nonempty_kind = "heading2"
            continue

        if TABLE_NUM_RE.match(text) and _is_confirmed_table_caption_paragraph(doc, paragraph, body_start):
            format_table_caption(paragraph)
            prev_nonempty_kind = "table_caption"
            continue

        if is_table_continuation_text(text):
            normalize_table_continuation_text(paragraph)
            format_table_caption(paragraph)
            prev_nonempty_kind = "table_continuation"
            continue

        if prev_nonempty_kind in {"table_caption", "table_continuation"}:
            format_table_title(paragraph)
            prev_nonempty_kind = "table_title"
            continue

        m_fig_pass = FIG_RE.match(text)
        if m_fig_pass and not caption_tail_is_reference_prose(m_fig_pass.group(3) or ""):
            normalize_figure_caption_text(paragraph)
            format_figure_caption(paragraph)
            prev_nonempty_kind = "figure_caption"
            continue

        if re.match(r"^\s*(источник|составлено по|рассчитано по|примечание)\s*:", text, re.IGNORECASE):
            format_source_line(paragraph)
            prev_nonempty_kind = "source_line"
            continue

        canonical = canonical_reference_subheading_text(text)
        if canonical:
            replace_paragraph_text(paragraph, canonical)
            format_reference_subheading(paragraph)
            prev_nonempty_kind = "reference_subheading"
            continue

        format_body(paragraph)
        prev_nonempty_kind = "body_text"

    run_with_pass_limit(
        "ensure_compact_heading2_spacing_final",
        ensure_compact_heading2_spacing,
        doc,
        body_start,
    )

    body_start = normalize_contents_layout(doc, body_start)

    normalize_sections(doc)
    ensure_front_matter_layout(doc, body_start)
    apply_page_breaks(doc, body_start)
    ensure_appendices_section_layout(doc, body_start)
    apply_page_numbering_policy(doc)

    # И ещё раз дочищаем цвет / highlight в самом конце
    remove_all_italic(doc)

    # Стандартизация сносок: TNR 10pt, без полужирного, одинарный интервал
    _format_footnotes(doc)

    run_with_pass_limit(
        "ensure_single_blank_after_references_heading_final",
        ensure_single_blank_after_references_heading,
        doc,
        body_start,
    )

    run_with_pass_limit(
        "ensure_compact_heading2_spacing_ultimate",
        ensure_compact_heading2_spacing,
        doc,
        body_start,
    )

    normalize_appendix_start_labels(doc, body_start)
    remove_empty_paragraphs_between_appendices_heading_and_first_label(doc, body_start)
    normalize_appendix_titles(doc, body_start)
    normalize_appendix_local_table_titles(doc, body_start)
    remove_empty_paragraphs_after_appendix_labels(doc, body_start)

    clear_heading_style_numbering(doc)

    run_with_pass_limit(
        "ensure_figure_block_keep_with_next",
        ensure_figure_block_keep_with_next,
        doc,
        body_start,
    )

    # Re-apply plain-list normalisation as the final pass: `format_body` in the
    # main classification loop resets paragraph indents back to body defaults
    # (left=0 firstLine=709), which strips the methodical list hanging indent
    # written by the early Phase-1 pass. Running the same normaliser again
    # against the now-stable, heading-styled document re-attaches `left=906
    # hanging=198` to dash/letter list blocks while heading guards continue to
    # protect chapter and subchapter paragraphs.
    normalize_plain_lists_in_document(doc, body_start)

    doc.save(str(output_path))
