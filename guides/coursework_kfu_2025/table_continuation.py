"""
Phase 3 — Table formatting utilities.

### What works now (geometry-based, no LibreOffice required)
  - apply_table_merging:      stub — returns 0 (see FUTURE note below)
  - apply_table_continuation: stub — returns 0 (see FUTURE note below)
  - _optimize_table_col_widths: active — fixes oversized/phantom columns
  - apply_rule3_table_orphan: active — prevents table caption orphaned at page bottom
  - apply_rule4_empty_first_lines: active — removes empty paragraphs at page top
  - apply_rule6_figure_orphan: active — keeps image with its caption

### FUTURE: Table splitting via LibreOffice (Rule 1)
#
# The table-continuation system (merge pre-split tables → re-split at real page
# breaks → insert "Продолжение таблицы X.Y.Z" headers) requires knowing EXACTLY
# where page breaks fall after formatting.  Pure geometry estimation (without a
# rendering engine) is too unreliable for production use:
#
#   Problem A — w:lastRenderedPageBreak (LRPB) is stale.
#     Word writes LRPB markers when it saves.  After Phase 1 reformatting
#     (fonts, margins, spacing all change) the LRPBs reflect the OLD layout,
#     not the new one.  Fresh KFU-formatted documents have NO LRPB at all,
#     producing 9-12 spurious "check manually" warnings per document.
#
#   Problem B — Geometry estimator is approximate.
#     Font metrics, line-wrap, cell merges, images in cells, and Word's own
#     internal kerning all introduce errors that compound over many rows.
#     A 2% per-row error on a 50-row table → entire page off.
#
# Recommended future approach — LibreOffice headless PDF-info:
#   1. Run `soffice --headless --convert-to pdf <formatted.docx>` (separate
#      Railway service or sidecar, NOT inline — adds ~400 MB + 8-15 s startup).
#   2. Parse the PDF page-stream to find exact row → page mapping.
#   3. Split at real page breaks, insert "Продолжение таблицы X.Y.Z" headers.
#
# Required helper functions (written, now commented-out):
#   _FORMATTER_RSID         — unique rsidR stamp for formatter-inserted paragraphs
#   _make_continuation_para — builds <w:p> "Продолжение таблицы X.Y.Z"
#   _split_table            — splits tbl_xml after row N, inserts continuation para
#   _is_formatter_continuation — detects formatter-stamped continuation paras
#   _rows_match / _merge_tables — merges two table parts (undo student splits)
#   apply_table_merging     — pre-pass: detect & merge student-split table pairs
#   apply_table_continuation — main pass: split at real page breaks
#
# To re-enable: restore those functions from git history (commit before this one),
# replace the stubs below, and integrate with a LibreOffice rendering step.
"""

from __future__ import annotations

import logging
import math
import os
import re
import shutil
import tempfile
import time
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


from .docx_utils import xml_has_image, is_source_or_note_line, FormattingReport
from .layout_render import LibreOfficeNotFoundError, render_docx_to_pdf
from .pdf_layout_analyzer import PdfLine, analyze_pdf_lines
from .table_split_prototype import (
    apply_numbered_split_to_document,
    _build_appendix_continuation_paragraph,
    _find_preceding_appendix_number,
)

logger = logging.getLogger(__name__)


@dataclass
class _MarkerSplitStats:
    renders: int = 0
    candidates: int = 0
    applied: int = 0
    skipped_reasons: dict[str, int] = field(default_factory=dict)
    applied_table_indexes: list[int] = field(default_factory=list)
    applied_captions: list[str] = field(default_factory=list)


_ACTIVE_MARKER_STATS: _MarkerSplitStats | None = None


@dataclass(frozen=True)
class _SamePageContinuationMarkerViolation:
    marker_text: str
    marker_page: int
    previous_table_page: int
    following_table_page: int | None = None
    confidence: str = "low"


_STRICT_CONTINUATION_MARKER_RE = re.compile(
    r"^\s*Продолжение\s+таблицы\s+\d+(?:\.\d+)*\s*$",
    re.IGNORECASE,
)
_ANY_CONTINUATION_MARKER_RE = re.compile(
    r"продолжение\s+таблицы\s+\d+(?:\.\d+)*",
    re.IGNORECASE,
)

# ── Unit helpers ─────────────────────────────────────────────────────────────

EMU_PER_PT  = 12700   # 1 pt  = 12 700 EMU  (python-docx stores lengths in EMU)
TWIP_PER_PT = 20      # 1 pt  = 20 twips    (w:trHeight val is in twips)

def _emu_pt(v: int) -> float: return v / EMU_PER_PT
def _twip_pt(v: int) -> float: return v / TWIP_PER_PT


# ── w:lastRenderedPageBreak helpers ──────────────────────────────────────────

_LRPB_TAG = qn("w:lastRenderedPageBreak")


def _para_has_lrpb(p_elem) -> bool:
    """True if this paragraph contains w:lastRenderedPageBreak.
    Used by _lrpb_calibrate (Rule 4 geometry estimator).
    """
    return p_elem.find(".//" + _LRPB_TAG) is not None

# ── Page geometry ─────────────────────────────────────────────────────────────

# Safety margin subtracted from body height so we don't overfill a page.
# Accounts for rounding + minor rendering differences between LO and Word.
_PAGE_BUFFER_PT = 36

# Minimum column width (pt) for column-width optimisation.
# Columns narrower than this are "phantom" (invisible/accidental).
# Using 20 pt (variant C): only truly phantom columns are redistributed;
# legitimate narrow columns (e.g. 30 pt numbering column) are left as-is.
_MIN_COL_PT = 20.0   # ≈ 0.7 cm — only phantom columns


def _body_height_pt(doc: Document) -> float:
    s = doc.sections[0]
    return _emu_pt(s.page_height - s.top_margin - s.bottom_margin) - _PAGE_BUFFER_PT


def _body_width_pt(doc: Document) -> float:
    s = doc.sections[0]
    return _emu_pt(s.page_width - s.left_margin - s.right_margin)


# ── Height estimators ─────────────────────────────────────────────────────────

# KFU body: Times New Roman 14 pt, 1.5 line spacing → ~21 pt/line
_BODY_LINE_PT  = 14 * 1.5
# Table cells: Times New Roman 12 pt, 1.0 line spacing → ~12 pt/line
_TABLE_LINE_PT = 12 * 1.0
# Empirical chars-per-line for 14 pt TNR in a 17 cm body column.
# Lowered from 68 → 62 to avoid underestimating multi-line paragraphs
# (shorter effective measure due to first-line indent + word-wrap).
_BODY_CHARS_PER_LINE = 62

# Approx pt per char for 12 pt TNR (used to derive chars-per-column)
_PT_PER_CHAR_TABLE = 6.0

# Top+bottom cell padding in pt (default Word cell margins ≈ 2.25 pt each side)
_CELL_PADDING_PT = 4.5


def _estimate_para_height(p) -> float:
    """Estimated rendered height of a body paragraph in points."""
    text = (p.text or "").strip()
    n_lines = max(1, math.ceil(len(text) / _BODY_CHARS_PER_LINE)) if text else 1

    line_h = _BODY_LINE_PT
    try:
        pf = p.paragraph_format
        ls = pf.line_spacing
        if ls is not None:
            # python-docx may return:
            #  • a Length subclass (Emu, Twips, …) with .pt for exact/atLeast rules
            #  • a plain float multiplier (e.g. 1.5) for auto rule
            #  • a raw int in 240ths-of-a-line when rule is unset (older python-docx)
            # Detection order: .pt first (handles all Length objects correctly),
            # then float multiplier, then 240ths fallback.
            # NOTE: WD_LINE_SPACING.EXACTLY == 4 (not 1), so checking int(rule)==1
            #       was wrong — we now rely on type detection instead.
            if hasattr(ls, "pt"):
                # Length object: .pt converts to points regardless of sub-type
                line_h = float(ls.pt)
            elif isinstance(ls, float):
                # Pure Python float → line spacing multiplier (e.g. 1.5)
                line_h = 14 * ls
            elif isinstance(ls, int):
                ls_i = int(ls)
                if ls_i > 10:
                    # Raw 240ths-of-a-line value (240=single, 360=1.5×, 480=double)
                    line_h = 14 * (ls_i / 240)
                else:
                    # Small integer treated as a multiplier (rare)
                    line_h = 14 * ls_i
    except Exception:
        pass

    sb = sa = 0.0
    try:
        if p.paragraph_format.space_before:
            sb = p.paragraph_format.space_before.pt
        if p.paragraph_format.space_after:
            sa = p.paragraph_format.space_after.pt
    except Exception:
        pass

    return n_lines * line_h + sb + sa


def _tbl_col_widths_pt(tbl_elem) -> list[float]:
    """
    Read actual column widths (in pt) from w:tblGrid / w:gridCol w:w (twips).
    Returns an empty list if not present.
    """
    tblGrid = tbl_elem.find(qn("w:tblGrid"))
    if tblGrid is None:
        return []
    widths = []
    for gc in tblGrid.findall(qn("w:gridCol")):
        w_val = gc.get(qn("w:w"))
        if w_val and w_val.isdigit():
            widths.append(_twip_pt(int(w_val)))
    return widths


def _cell_margins_pt(cell_elem) -> float:
    """
    Return total vertical cell margin (top + bottom) in pt from w:tcPr/w:tcMar.
    Falls back to _CELL_PADDING_PT if not specified.
    """
    tcPr = cell_elem.find(qn("w:tcPr"))
    if tcPr is None:
        return _CELL_PADDING_PT
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        return _CELL_PADDING_PT
    total = 0.0
    found = False
    for side in ("w:top", "w:bottom"):
        el = tcMar.find(qn(side))
        if el is not None:
            w_type = el.get(qn("w:type"), "dxa")
            val = el.get(qn("w:w"), "0")
            if val.lstrip("-").isdigit():
                if w_type == "dxa":
                    total += _twip_pt(int(val))
                elif w_type == "nil":
                    pass   # zero
            found = True
    return total if found else _CELL_PADDING_PT


def _para_font_size_pt(p_elem) -> float:
    """
    Read font size (pt) from the paragraph's rPr or its first run's rPr.
    Checks paragraph-level rPr first (w:pPr/w:rPr), then first w:r/w:rPr.
    Falls back to _TABLE_LINE_PT.
    """
    # Paragraph-level run properties (pPr > rPr)
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is not None:
        rPr = pPr.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                val = sz.get(qn("w:val"))
                if val and val.isdigit():
                    return int(val) / 2

    # First run's rPr
    for r in p_elem.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            sz = rPr.find(qn("w:sz"))
            if sz is not None:
                val = sz.get(qn("w:val"))
                if val and val.isdigit():
                    return int(val) / 2

    return _TABLE_LINE_PT   # default: 12 pt


def _para_line_height_pt(p_elem, font_pt: float) -> float:
    """
    Resolve actual single-line rendered height (pt) for a paragraph,
    reading w:spacing w:line + w:lineRule from the paragraph's pPr.
    """
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return font_pt
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        return font_pt

    line_val = spacing.get(qn("w:line"))
    line_rule = spacing.get(qn("w:lineRule"), "auto")

    if line_val and line_val.lstrip("-").isdigit():
        lv = int(line_val)
        if line_rule == "exact":
            # Exact: value is in twips
            return _twip_pt(lv)
        elif line_rule == "atLeast":
            # At-least: value in twips, but could be taller
            return max(font_pt, _twip_pt(lv))
        else:
            # "auto" (default): value is in 240ths of a line
            # 240 = single spacing; 360 = 1.5x
            return font_pt * (lv / 240.0)

    return font_pt


def _para_spacing_pt(p_elem) -> tuple[float, float]:
    """Return (space_before_pt, space_after_pt) for a paragraph."""
    pPr = p_elem.find(qn("w:pPr"))
    if pPr is None:
        return 0.0, 0.0
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        return 0.0, 0.0
    sb = sa = 0.0
    before = spacing.get(qn("w:before"))
    after  = spacing.get(qn("w:after"))
    if before and before.lstrip("-").isdigit():
        sb = _twip_pt(int(before))
    if after and after.lstrip("-").isdigit():
        sa = _twip_pt(int(after))
    return sb, sa


def _estimate_cell_height(cell, col_w_pt: float) -> float:
    """
    Estimate total height of a table cell in points.

    Accounts for:
    - All paragraphs in the cell (not just concatenated text)
    - Per-paragraph font size, line spacing, space_before, space_after
    - Proportional TNR character width for line-wrap estimation
    - Cell top+bottom margins from w:tcMar
    """
    p_elems = cell._element.findall(qn("w:p"))
    if not p_elems:
        return _TABLE_LINE_PT + _CELL_PADDING_PT

    total_h = 0.0
    for p_elem in p_elems:
        font_pt = _para_font_size_pt(p_elem)
        line_h  = _para_line_height_pt(p_elem, font_pt)
        # If no explicit line spacing is set in the paragraph XML, the cell
        # inherits the document's Normal style (typically 1.5× in KFU docs).
        # Apply 1.5× as a conservative default to avoid underestimating row height.
        if abs(line_h - font_pt) < 0.5:   # "line_h == font_pt" means unset (single)
            line_h = font_pt * 1.5
        sb, sa  = _para_spacing_pt(p_elem)

        # TNR avg char width ≈ 0.50 × font size (conservative — Cyrillic glyphs are wider than Latin)
        pt_per_char  = font_pt * 0.50
        chars_per_line = max(4, int(col_w_pt / pt_per_char))

        # Gather text from all runs (preserves multi-run paragraphs)
        text = "".join(
            (r.find(qn("w:t")).text or "")
            for r in p_elem.findall(qn("w:r"))
            if r.find(qn("w:t")) is not None
        ).strip()

        n_lines = max(1, math.ceil(len(text) / chars_per_line)) if text else 1
        total_h += n_lines * line_h + sb + sa

    # Cell top+bottom margins
    cell_margin = _cell_margins_pt(cell._element)
    return total_h + cell_margin


def _estimate_row_height(row, body_width_pt: float, col_widths_pt: list[float] | None = None) -> float:
    """
    Estimated rendered height of a table row in points.

    Priority:
    1. Explicit w:trHeight (hRule=exact) → use as-is
    2. Explicit w:trHeight (hRule=atLeast) → use as minimum
    3. Estimate from cell content via _estimate_cell_height
    """
    tr = row._tr
    trPr = tr.find(qn("w:trPr"))
    explicit_min = 0.0
    if trPr is not None:
        trH = trPr.find(qn("w:trHeight"))
        if trH is not None:
            val = trH.get(qn("w:val"))
            h_rule = trH.get(qn("w:hRule"), "atLeast")
            if val and val.lstrip("-").isdigit():
                h = _twip_pt(int(val))
                if h > 2:
                    if h_rule == "exact":
                        return h   # exact → trust it completely
                    else:
                        explicit_min = h   # atLeast → use as lower bound

    cells = row.cells
    if not cells:
        return max(explicit_min, _TABLE_LINE_PT + _CELL_PADDING_PT)

    num_cols = len(cells)

    # Per-cell column width: actual XML widths preferred
    if col_widths_pt and len(col_widths_pt) >= num_cols:
        col_ws = col_widths_pt
    else:
        equal_w = max(20.0, body_width_pt / num_cols)
        col_ws = [equal_w] * num_cols

    max_h = 0.0
    seen: set[int] = set()
    col_idx = 0
    for cell in cells:
        cid = id(cell._element)
        if cid in seen:
            col_idx += 1
            continue
        seen.add(cid)

        col_w_pt = col_ws[col_idx] if col_idx < len(col_ws) else max(20.0, body_width_pt / num_cols)
        cell_h = _estimate_cell_height(cell, col_w_pt)
        max_h = max(max_h, cell_h)
        col_idx += 1

    return max(explicit_min, max_h)


# ── Body element iteration ────────────────────────────────────────────────────

def _iter_body(doc: Document):
    """
    Yield (kind, xml_element, py_object) for each direct child of <w:body>.
    kind ∈ {"paragraph", "table"}
    """
    body = doc.element.body
    para_map  = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    for child in body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "p" and child in para_map:
            yield "paragraph", child, para_map[child]
        elif local == "tbl" and child in table_map:
            yield "table", child, table_map[child]


# ── Table number extraction ───────────────────────────────────────────────────

_TBL_NUM_RE = re.compile(
    r"(?:таблица|table)\s+(\d+(?:\.\d+){0,2})",
    re.IGNORECASE,
)
_CONT_NUM_RE = re.compile(
    r"продолжение\s+таблицы\s+(\d+(?:\.\d+){0,2})",
    re.IGNORECASE,
)


def _extract_table_num(text: str) -> str | None:
    m = _TBL_NUM_RE.search(text.strip())
    return m.group(1) if m else None


# ── Table merging / continuation detection helpers ────────────────────────────
# (splitting/merging logic is stubbed — see module docstring for FUTURE plan)

_CONT_RE = re.compile(r"продолжени", re.IGNORECASE)
_TBL_WORD_RE = re.compile(r"таблиц", re.IGNORECASE)
_APPENDIX_CONTINUATION_RE = re.compile(
    r"^\s*продолжение\s+приложения\s+(?:\d{1,3}|[A-Za-zА-ЯЁ])\s*$",
    re.IGNORECASE,
)


def _is_student_continuation(text: str) -> bool:
    """
    True if paragraph text looks like a student-written standalone
    'Продолжение таблицы X.Y.Z' header.

    Guard: text must be short (≤30 chars) — long paragraphs are prose
    that merely happen to contain those words mid-sentence.
    30 chars covers table numbers up to e.g. "100.10.10" depth.
    """
    if len(text) > 30:
        return False
    return bool(_CONT_RE.search(text) and _TBL_WORD_RE.search(text))


def _norm_text(text: str) -> str:
    return " ".join((text or "").split())


def _is_any_continuation_marker(text: str) -> bool:
    t = _norm_text(text)
    return bool(t and t.lower().startswith("продолжение таблицы"))


def _table_col_count(tbl_xml) -> int:
    grid = tbl_xml.find(qn("w:tblGrid"))
    if grid is not None:
        cols = grid.findall(qn("w:gridCol"))
        if cols:
            return len(cols)

    first_row = tbl_xml.find(qn("w:tr"))
    if first_row is None:
        return 0
    count = 0
    for tc in first_row.findall(qn("w:tc")):
        tcPr = tc.find(qn("w:tcPr"))
        gs = tcPr.find(qn("w:gridSpan")) if tcPr is not None else None
        span = int(gs.get(qn("w:val"), 1)) if gs is not None else 1
        count += max(1, span)
    return count


def _row_cell_texts(tr_xml) -> list[str]:
    vals: list[str] = []
    for tc in tr_xml.findall(qn("w:tc")):
        txt = "".join(
            (t.text or "")
            for t in tc.findall(".//" + qn("w:t"))
            if t.text
        )
        vals.append(_norm_text(txt))
    return vals


def _row_is_simple_full_width(tr_xml, col_count: int) -> bool:
    if col_count <= 0:
        return False
    cells = tr_xml.findall(qn("w:tc"))
    if len(cells) != col_count:
        return False
    for tc in cells:
        tc_pr = tc.find(qn("w:tcPr"))
        if tc_pr is not None and (
            tc_pr.find(qn("w:gridSpan")) is not None
            or tc_pr.find(qn("w:vMerge")) is not None
        ):
            return False
    return True


def _row_is_exact_numeric_row(tr_xml, col_count: int) -> bool:
    if not _row_is_simple_full_width(tr_xml, col_count):
        return False
    return _row_cell_texts(tr_xml) == [str(i) for i in range(1, col_count + 1)]


def _row_looks_numeric_but_malformed(tr_xml, col_count: int) -> bool:
    if not _row_is_simple_full_width(tr_xml, col_count):
        return False
    values = _row_cell_texts(tr_xml)
    if not values:
        return False
    if not all(re.fullmatch(r"\d+", value or "") for value in values):
        return False
    return values != [str(i) for i in range(1, col_count + 1)]


def _manual_chain_rows_compatible(tbl1, tbl2) -> bool:
    col_count = _table_col_count(tbl1)
    if col_count <= 0 or _table_col_count(tbl2) != col_count:
        return False

    rows1 = tbl1.findall(qn("w:tr"))
    rows2 = tbl2.findall(qn("w:tr"))
    if not rows1 or not rows2 or not _tbl_has_at_least_two_rows(tbl2):
        return False

    if _rows_match(rows1[0], rows2[0]):
        return True

    # Manual KFU continuations often omit the semantic header in continuation
    # fragments and repeat only the numeric column row before continuation data.
    return (
        len(rows1) > 1
        and _row_is_exact_numeric_row(rows1[1], col_count)
        and _row_is_exact_numeric_row(rows2[0], col_count)
    )


def _clear_cell_content_preserving_properties(tc_xml) -> None:
    tc_pr = tc_xml.find(qn("w:tcPr"))
    for child in list(tc_xml):
        tc_xml.remove(child)
    if tc_pr is not None:
        tc_xml.append(tc_pr)


def _build_numeric_row_from_header(header_row_xml, col_count: int):
    if not _row_is_simple_full_width(header_row_xml, col_count):
        raise ValueError("numeric row synthesis requires simple full-width header")
    numeric_row = deepcopy(header_row_xml)
    for idx, tc in enumerate(numeric_row.findall(qn("w:tc")), start=1):
        _clear_cell_content_preserving_properties(tc)
        p = OxmlElement("w:p")
        p_pr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        p_pr.append(jc)
        p.append(p_pr)

        r = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:hAnsi"), "Times New Roman")
        fonts.set(qn("w:cs"), "Times New Roman")
        r_pr.append(fonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        r_pr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), "24")
        r_pr.append(sz_cs)
        r.append(r_pr)

        t = OxmlElement("w:t")
        t.text = str(idx)
        r.append(t)
        p.append(r)
        tc.append(p)
    return numeric_row


def _ensure_fragment_numeric_row(tbl_xml) -> tuple[int, str | None]:
    col_count = _table_col_count(tbl_xml)
    if col_count <= 0:
        return 0, "no_columns"
    rows = tbl_xml.findall(qn("w:tr"))
    if len(rows) < 2:
        return 0, "no_data_row"

    if _row_is_exact_numeric_row(rows[0], col_count):
        return 0, None
    if _row_looks_numeric_but_malformed(rows[0], col_count):
        return 0, "malformed_numeric_row"
    if _row_is_exact_numeric_row(rows[1], col_count):
        return 0, None
    if _row_looks_numeric_but_malformed(rows[1], col_count):
        return 0, "malformed_numeric_row"
    if not _row_is_simple_full_width(rows[0], col_count):
        return 0, "complex_header"

    rows[0].addnext(_build_numeric_row_from_header(rows[0], col_count))
    return 1, None


def _ensure_manual_chain_numeric_rows(tbl1, tbl2) -> tuple[int, list[str]]:
    repairs = 0
    reasons: list[str] = []
    for label, tbl_xml in (("first", tbl1), ("continuation", tbl2)):
        changed, reason = _ensure_fragment_numeric_row(tbl_xml)
        repairs += changed
        if reason:
            reasons.append(f"{label}:{reason}")
    return repairs, reasons


_NESTED_TABLE_HEADER_PREFIXES = (
    ("уровень", "формальные органы", "неформальные практики", "основные функции"),
)


def _ordinary_table_has_nested_header_row(doc: Document, table_index: int) -> bool:
    if table_index < 0 or table_index >= len(doc.tables):
        return False

    rows = doc.tables[table_index]._tbl.findall(qn("w:tr"))
    for row_index, row_xml in enumerate(rows):
        if row_index == 0:
            continue
        cells = tuple(text.lower() for text in _row_cell_texts(row_xml) if text)
        if not cells:
            continue
        for prefix in _NESTED_TABLE_HEADER_PREFIXES:
            if len(cells) >= len(prefix) and cells[: len(prefix)] == prefix:
                return True
    return False


def _rows_match(row1_xml, row2_xml) -> bool:
    return _row_cell_texts(row1_xml) == _row_cell_texts(row2_xml)


@dataclass(frozen=True)
class RowSignature:
    row_idx: int
    key: str
    fragments: tuple[str, ...]


@dataclass(frozen=True)
class TableSignature:
    table_idx: int
    tbl_xml: object
    rows: tuple[RowSignature, ...]


@dataclass(frozen=True)
class RenderedSplitCandidate:
    table_idx: int
    tbl_xml: object
    split_after: int


@dataclass(frozen=True)
class RenderedWholeTableMoveCandidate:
    table_idx: int
    tbl_xml: object
    caption_para_xml: object


_START_HAS_COMPLETE_DATA_ROW = "has_complete_data_row"
_START_NO_COMPLETE_DATA_ROW = "no_complete_data_row"
_START_AMBIGUOUS = "ambiguous"


def _tbl_has_at_least_two_rows(tbl_xml) -> bool:
    return len(tbl_xml.findall(qn("w:tr"))) >= 2


def _is_vmerge_continue(tc_xml) -> bool:
    tcPr = tc_xml.find(qn("w:tcPr"))
    if tcPr is None:
        return False
    vm = tcPr.find(qn("w:vMerge"))
    if vm is None:
        return False
    val = vm.get(qn("w:val"))
    # w:vMerge with no val is "continue" by spec.
    return val is None or val == "continue"


def _is_split_boundary_safe(rows_xml: list, split_after: int) -> bool:
    """
    split_after is index of the last row in part 1.
    Boundary is between rows[split_after] and rows[split_after+1].
    """
    if split_after < 0 or split_after + 1 >= len(rows_xml):
        return False
    next_row = rows_xml[split_after + 1]
    for tc in next_row.findall(qn("w:tc")):
        if _is_vmerge_continue(tc):
            return False
    return True


def _find_safe_split_after(rows_xml: list, candidate_after: int) -> int | None:
    """
    Move split boundary upward until it is safe and leaves at least
    header + 1 data row in part 1.
    """
    s = candidate_after
    while s >= 1:
        if _is_split_boundary_safe(rows_xml, s):
            return s
        s -= 1
    return None


_GEOMETRY_SIMPLE = "simple"
_GEOMETRY_PRESERVE = "preserve_geometry"
_GEOMETRY_UNSAFE = "unsafe_no_split"
_NARROW_GRID_COL_TWIPS = 120


def _int_attr(el, attr: str, default: int | None = None) -> int | None:
    value = el.get(qn(attr))
    if value is None:
        return default
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _append_unique(items: list[str], item: str) -> None:
    if item not in items:
        items.append(item)


def _is_default_tbl_look(tbl_look) -> bool:
    if tbl_look is None:
        return False
    return (
        tbl_look.get(qn("w:val")) == "04A0"
        and tbl_look.get(qn("w:firstColumn")) == "1"
        and tbl_look.get(qn("w:firstRow")) == "1"
        and tbl_look.get(qn("w:lastColumn")) == "0"
        and tbl_look.get(qn("w:lastRow")) == "0"
        and tbl_look.get(qn("w:noHBand")) == "0"
        and tbl_look.get(qn("w:noVBand")) == "1"
    )


def _is_default_auto_tbl_w(tbl_w) -> bool:
    if tbl_w is None:
        return False
    return tbl_w.get(qn("w:type")) == "auto" and tbl_w.get(qn("w:w")) in {None, "0"}


def _row_has_vmerge_continue(row_xml) -> bool:
    return any(_is_vmerge_continue(tc) for tc in row_xml.findall(qn("w:tc")))


def _table_has_adjacent_continuation_marker(
    body_children: list,
    para_by_xml: dict,
    tbl_xml,
) -> bool:
    try:
        idx = body_children.index(tbl_xml)
    except ValueError:
        return False

    for offset in (-1, 1):
        j = idx + offset
        if j < 0 or j >= len(body_children):
            continue
        node = body_children[j]
        if node.tag != qn("w:p"):
            continue
        para = para_by_xml.get(node)
        if para is not None and _is_any_continuation_marker(para.text or ""):
            return True
    return False


def _table_geometry_policy_details(
    tbl_xml,
    *,
    split_before_row: int | None = None,
    has_existing_continuation_marker: bool = False,
) -> tuple[str, list[str]]:
    """
    Classify whether table geometry may be safely width-optimized/split.

    Default python-docx table artifacts (`tblW auto w=0`, default `tblLook`)
    are intentionally not treated as sensitive on their own; otherwise every
    simple synthetic/default table would bypass the optimizer.
    """
    reasons: list[str] = []

    grid = tbl_xml.find(qn("w:tblGrid"))
    if grid is None:
        return _GEOMETRY_UNSAFE, ["malformed_grid:missing_tblGrid"]
    gridcols = grid.findall(qn("w:gridCol"))
    if not gridcols:
        return _GEOMETRY_UNSAFE, ["malformed_grid:empty_tblGrid"]

    col_count = len(gridcols)
    for grid_col in gridcols:
        width = _int_attr(grid_col, "w:w")
        if width is None or width <= 0:
            return _GEOMETRY_UNSAFE, ["malformed_grid:invalid_gridCol_width"]
        if width <= _NARROW_GRID_COL_TWIPS:
            _append_unique(reasons, "narrow_grid_col")

    tbl_pr = tbl_xml.find(qn("w:tblPr"))
    if tbl_pr is not None:
        if tbl_pr.find(qn("w:tblCellSpacing")) is not None:
            _append_unique(reasons, "tblCellSpacing")
        if tbl_pr.find(qn("w:tblCellMar")) is not None:
            _append_unique(reasons, "tblCellMar")

        tbl_look = tbl_pr.find(qn("w:tblLook"))
        if tbl_look is not None and not _is_default_tbl_look(tbl_look):
            _append_unique(reasons, "tblLook")

        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if (
            tbl_layout is not None
            and (tbl_layout.get(qn("w:type")) or "").lower() == "fixed"
        ):
            _append_unique(reasons, "tblLayout_fixed")

        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is not None and not _is_default_auto_tbl_w(tbl_w):
            tbl_w_type = (tbl_w.get(qn("w:type")) or "").lower()
            if tbl_w_type in {"auto", "pct"}:
                _append_unique(reasons, f"tblW_{tbl_w_type}")

    rows = tbl_xml.findall(qn("w:tr"))
    if split_before_row is not None:
        if split_before_row < 0 or split_before_row >= len(rows):
            return _GEOMETRY_UNSAFE, ["ambiguous_topology:split_out_of_range"]
        if _row_has_vmerge_continue(rows[split_before_row]):
            return _GEOMETRY_UNSAFE, ["split_crosses_vMerge_continuation"]

    for row in rows:
        tr_pr = row.find(qn("w:trPr"))
        if tr_pr is not None and tr_pr.find(qn("w:tblPrEx")) is not None:
            _append_unique(reasons, "tblPrEx")

        raw_cells = row.findall(qn("w:tc"))
        span_total = 0
        has_grid_span = False
        for tc in raw_cells:
            tc_pr = tc.find(qn("w:tcPr"))
            span = 1
            if tc_pr is not None:
                grid_span = tc_pr.find(qn("w:gridSpan"))
                if grid_span is not None:
                    has_grid_span = True
                    _append_unique(reasons, "gridSpan")
                    span = _int_attr(grid_span, "w:val", 1) or 1
                    if span < 1:
                        return _GEOMETRY_UNSAFE, ["ambiguous_topology:invalid_gridSpan"]

                vmerge = tc_pr.find(qn("w:vMerge"))
                if vmerge is not None:
                    _append_unique(reasons, "vMerge")

                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is not None and tc_w.get(qn("w:w")) == "0":
                    _append_unique(reasons, "tcW_zero")

            span_total += span

        if len(raw_cells) != col_count:
            _append_unique(reasons, "raw_tc_count_mismatch")
        if span_total > col_count:
            return _GEOMETRY_UNSAFE, ["ambiguous_topology:span_exceeds_grid"]
        if span_total < col_count and not has_grid_span:
            _append_unique(reasons, "raw_tc_count_mismatch")

    if has_existing_continuation_marker:
        _append_unique(reasons, "existing_continuation_marker")

    if reasons:
        return _GEOMETRY_PRESERVE, reasons
    return _GEOMETRY_SIMPLE, []


def classify_table_geometry_policy(
    tbl_xml,
    *,
    split_before_row: int | None = None,
    has_existing_continuation_marker: bool = False,
) -> str:
    """Public testable wrapper for geometry preservation policy."""
    policy, _ = _table_geometry_policy_details(
        tbl_xml,
        split_before_row=split_before_row,
        has_existing_continuation_marker=has_existing_continuation_marker,
    )
    return policy


def _format_geometry_reasons(reasons: list[str]) -> str:
    return ",".join(reasons) if reasons else "-"


def _split_geometry_is_safe(
    tbl_xml,
    *,
    table_index: int,
    split_before_row: int,
    log_prefix: str,
) -> bool:
    policy, reasons = _table_geometry_policy_details(
        tbl_xml,
        split_before_row=split_before_row,
    )
    if policy != _GEOMETRY_UNSAFE:
        return True
    logger.info(
        "%s table_index=%s reason=unsafe_geometry geometry_reasons=%s",
        log_prefix,
        table_index,
        _format_geometry_reasons(reasons),
    )
    return False


def _find_caption_number_before_table(doc: Document, tbl_xml) -> str | None:
    """
    Strict source of truth: caption paragraph before the table.
    Supports:
      - "Таблица X.X"
      - "Таблица X.X.X"
    and two-paragraph format (caption line + title line).
    """
    body = doc.element.body
    children = list(body)
    try:
        idx = children.index(tbl_xml)
    except ValueError:
        return None

    # Build a fast map of paragraph XML -> paragraph text
    para_text = {p._element: _norm_text(p.text) for p in doc.paragraphs}

    j = idx - 1
    nonempty_seen = 0
    while j >= 0 and nonempty_seen < 4:
        node = children[j]
        if node.tag == qn("w:p"):
            txt = para_text.get(node, "")
            if txt:
                nonempty_seen += 1
                m = _TBL_NUM_RE.match(txt)
                if m:
                    return m.group(1)
        elif node.tag == qn("w:tbl"):
            break
        j -= 1
    return None


def _find_caption_paragraph_before_table(doc: Document, tbl_xml):
    """
    Return the strict table caption paragraph XML and number before tbl_xml.
    The caption paragraph, not the title paragraph, is the only safe anchor for
    whole-table moves.
    """
    body = doc.element.body
    children = list(body)
    try:
        idx = children.index(tbl_xml)
    except ValueError:
        return None

    para_text = {p._element: _norm_text(p.text) for p in doc.paragraphs}

    j = idx - 1
    nonempty_seen = 0
    while j >= 0 and nonempty_seen < 4:
        node = children[j]
        if node.tag == qn("w:p"):
            txt = para_text.get(node, "")
            if txt:
                nonempty_seen += 1
                m = _TBL_NUM_RE.match(txt)
                if m:
                    return node, m.group(1)
        elif node.tag == qn("w:tbl"):
            break
        j -= 1
    return None


def _norm_match_text(text: str) -> str:
    return _norm_text(text).lower()


def _row_signature(tr_xml, row_idx: int) -> RowSignature | None:
    fragments = tuple(
        frag
        for frag in (_norm_match_text(t) for t in _row_cell_texts(tr_xml))
        if frag
    )
    if not fragments:
        return None
    return RowSignature(row_idx=row_idx, key=" || ".join(fragments), fragments=fragments)


def _collect_table_signatures(doc: Document) -> list[TableSignature]:
    out: list[TableSignature] = []
    for table_idx, table in enumerate(doc.tables):
        rows: list[RowSignature] = []
        for row_idx, tr in enumerate(table._tbl.findall(qn("w:tr"))):
            sig = _row_signature(tr, row_idx)
            if sig is not None:
                rows.append(sig)
        out.append(TableSignature(table_idx=table_idx, tbl_xml=table._tbl, rows=tuple(rows)))
    return out


def _valid_manual_continuation_table_indexes(doc: Document) -> set[int]:
    """
    Return stable table indexes that are already part of a valid continuation chain
    (manual student-authored OR auto-inserted by an earlier marker-split run).

    Accepted chain shapes (both must be preserved across re-runs):
      tbl → marker_p → tbl                  (manual; no blank between marker and table 2)
      tbl → marker_p → blank_p → tbl        (auto; one blank between marker and table 2)
    """
    body = doc.element.body
    children = list(body)
    skip: set[int] = set()
    para_by_xml = {p._element: p for p in doc.paragraphs}
    table_index_by_xml = {table._tbl: idx for idx, table in enumerate(doc.tables)}

    def _is_empty_p(elem) -> bool:
        if elem.tag != qn("w:p"):
            return False
        p_obj = para_by_xml.get(elem)
        return p_obj is not None and not _norm_text(p_obj.text or "")

    i = 1
    while i < len(children) - 1:
        prev_node = children[i - 1]
        node = children[i]

        if prev_node.tag != qn("w:tbl") or node.tag != qn("w:p"):
            i += 1
            continue

        p_obj = para_by_xml.get(node)
        marker_text = _norm_text(p_obj.text if p_obj is not None else "")
        if not _is_any_continuation_marker(marker_text):
            i += 1
            continue

        # Locate the second table: either immediately after the marker (manual)
        # or after exactly one blank paragraph (auto-inserted by marker split).
        next_idx = i + 1
        if next_idx < len(children) and _is_empty_p(children[next_idx]):
            next_idx += 1
        if next_idx >= len(children):
            i += 1
            continue
        next_node = children[next_idx]
        if next_node.tag != qn("w:tbl"):
            i += 1
            continue

        if (
            _is_valid_manual_continuation_chain(doc, prev_node, node, next_node)
            or _is_structurally_valid_student_chain(doc, prev_node, node, next_node)
        ):
            if prev_node in table_index_by_xml:
                skip.add(table_index_by_xml[prev_node])
            if next_node in table_index_by_xml:
                skip.add(table_index_by_xml[next_node])

        i += 1

    return skip


def _valid_manual_continuation_table_ids(doc: Document) -> set[int]:
    """Compatibility wrapper: manual-chain protection is now index-based."""
    return _valid_manual_continuation_table_indexes(doc)


def _repair_manual_continuation_numeric_rows(doc: Document) -> int:
    """Ensure numeric column rows inside already-recognized continuation chains.

    Conservative scope: this scans only `tbl -> Продолжение таблицы -> tbl`
    chains (with the existing optional one blank paragraph before table 2).
    Ordinary unsplit tables are not considered.
    """
    body = doc.element.body
    children = list(body)
    para_by_xml = {p._element: p for p in doc.paragraphs}

    def _is_empty_p(elem) -> bool:
        if elem.tag != qn("w:p"):
            return False
        p_obj = para_by_xml.get(elem)
        return p_obj is not None and not _norm_text(p_obj.text or "")

    repaired = 0
    i = 1
    while i < len(children) - 1:
        prev_node = children[i - 1]
        node = children[i]

        if prev_node.tag != qn("w:tbl") or node.tag != qn("w:p"):
            i += 1
            continue

        p_obj = para_by_xml.get(node)
        marker_text = _norm_text(p_obj.text if p_obj is not None else "")
        if not _is_any_continuation_marker(marker_text):
            i += 1
            continue

        next_idx = i + 1
        if next_idx < len(children) and _is_empty_p(children[next_idx]):
            next_idx += 1
        if next_idx >= len(children):
            i += 1
            continue
        next_node = children[next_idx]
        if next_node.tag != qn("w:tbl"):
            i += 1
            continue

        if not (
            _is_valid_manual_continuation_chain(doc, prev_node, node, next_node)
            or _is_structurally_valid_student_chain(doc, prev_node, node, next_node)
        ):
            i += 1
            continue

        changed, reasons = _ensure_manual_chain_numeric_rows(prev_node, next_node)
        if changed:
            repaired += changed
            logger.info(
                "manual_continuation_numeric_rows_repaired marker=%s rows=%s",
                marker_text,
                changed,
            )
        for reason in reasons:
            logger.info(
                "manual_continuation_numeric_rows_skipped marker=%s reason=%s",
                marker_text,
                reason,
            )
        i += 1

    return repaired


def _paragraph_text_from_xml(p_xml) -> str:
    return _norm_text("".join(t.text or "" for t in p_xml.findall(".//" + qn("w:t"))))


def _is_blank_service_paragraph(p_xml) -> bool:
    if _paragraph_text_from_xml(p_xml):
        return False
    if xml_has_image(p_xml):
        return False
    pPr = p_xml.find(qn("w:pPr"))
    return pPr is None or pPr.find(qn("w:sectPr")) is None


def _is_appendix_continuation_paragraph(p_xml) -> bool:
    return bool(_APPENDIX_CONTINUATION_RE.match(_paragraph_text_from_xml(p_xml)))


def _first_row_is_generated_numbered_row(tbl_xml) -> bool:
    col_count = _table_col_count(tbl_xml)
    if col_count <= 0:
        return False
    first_row = tbl_xml.find(qn("w:tr"))
    if first_row is None:
        return False
    return _row_cell_texts(first_row) == [str(i) for i in range(1, col_count + 1)]


def _previous_significant_body_child_is_table(doc: Document, tbl_xml) -> bool:
    children = list(doc.element.body)
    table_body_index = None
    for idx, child in enumerate(children):
        if child is tbl_xml:
            table_body_index = idx
            break
    if table_body_index is None:
        return False

    idx = table_body_index - 1
    while idx >= 0:
        child = children[idx]
        if child.tag == qn("w:p"):
            if _is_blank_service_paragraph(child) or _is_appendix_continuation_paragraph(child):
                idx -= 1
                continue
        return child.tag == qn("w:tbl")
    return False


def _is_generated_appendix_continuation_table(doc: Document, table_index: int) -> bool:
    if table_index < 0 or table_index >= len(doc.tables):
        return False
    tbl_xml = doc.tables[table_index]._tbl
    return (
        _first_row_is_generated_numbered_row(tbl_xml)
        and _previous_significant_body_child_is_table(doc, tbl_xml)
    )


def _paragraph_has_keep_next(p_xml) -> bool:
    pPr = p_xml.find(qn("w:pPr"))
    if pPr is None:
        return False
    keep = pPr.find(qn("w:keepNext"))
    if keep is None:
        return False
    return keep.get(qn("w:val")) not in {"0", "false", "False"}


def _paragraph_is_right_aligned(p_xml) -> bool:
    pPr = p_xml.find(qn("w:pPr"))
    if pPr is None:
        return False
    jc = pPr.find(qn("w:jc"))
    return bool(jc is not None and jc.get(qn("w:val")) == "right")


def _manual_marker_matches_caption(doc: Document, tbl_xml, marker_text: str) -> bool:
    caption_num = _find_caption_number_before_table(doc, tbl_xml)
    marker_match = _CONT_NUM_RE.search(marker_text)
    marker_num = marker_match.group(1) if marker_match else None
    if caption_num is None and marker_num is None:
        return True
    return caption_num == marker_num


def _is_valid_manual_continuation_chain(doc: Document, tbl1, marker_p, tbl2) -> bool:
    marker_text = ""
    for text_node in marker_p.findall(".//" + qn("w:t")):
        marker_text += text_node.text or ""
    marker_text = _norm_text(marker_text)
    if not _is_any_continuation_marker(marker_text):
        return False
    if not _manual_marker_matches_caption(doc, tbl1, marker_text):
        return False
    if not _paragraph_is_right_aligned(marker_p):
        return False
    if not _paragraph_has_keep_next(marker_p):
        return False

    return _manual_chain_rows_compatible(tbl1, tbl2)


def _is_structurally_valid_student_chain(doc: Document, tbl1, marker_p, tbl2) -> bool:
    # Looser merge-safety gate: structurally sound student-authored chain that
    # only lacks keepNext (a formatter-applied attribute). Used by
    # apply_table_merging to avoid destroying valid student splits.
    marker_text = ""
    for text_node in marker_p.findall(".//" + qn("w:t")):
        marker_text += text_node.text or ""
    marker_text = _norm_text(marker_text)
    if not _is_any_continuation_marker(marker_text):
        return False
    if not _manual_marker_matches_caption(doc, tbl1, marker_text):
        return False
    if not _paragraph_is_right_aligned(marker_p):
        return False

    return _manual_chain_rows_compatible(tbl1, tbl2)


def _marker_has_enabled_page_break(marker_p) -> bool:
    # True if marker paragraph has <w:pageBreakBefore/> enabled (no w:val
    # attribute, or w:val in {"1", "true", "on"}). Used to keep enable
    # operation idempotent across re-runs.
    pPr = marker_p.find(qn("w:pPr"))
    if pPr is None:
        return False
    pb = pPr.find(qn("w:pageBreakBefore"))
    if pb is None:
        return False
    val = pb.get(qn("w:val"))
    return val is None or val in {"1", "true", "True", "on"}


def _ensure_paragraph_bool_property_active(p_xml, prop_name: str, *, prepend: bool = False) -> bool:
    pPr = p_xml.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_xml.insert(0, pPr)

    props = pPr.findall(qn(prop_name))
    changed = False
    if props:
        for prop in props:
            if qn("w:val") in prop.attrib:
                del prop.attrib[qn("w:val")]
                changed = True
        return changed

    new_prop = OxmlElement(prop_name)
    if prepend:
        pPr.insert(0, new_prop)
    else:
        pPr.append(new_prop)
    return True


def _enable_marker_page_break_for_student_chain(marker_p) -> int:
    # P1-critical / DEFECT E: Phase 1 `hard_reset_paragraph_format` neutralizes
    # `<w:pageBreakBefore/>` and `<w:keepNext/>` on the marker paragraph by
    # setting w:val="0" (disabled). When the chain is preserved as a student
    # chain (not formatter-authored), the marker renders at the bottom of the
    # previous page instead of the top of the continuation page because both
    # properties stay disabled.
    #
    # Formatter-authored markers (built by `_build_continuation_para`) keep
    # these properties ENABLED with no w:val attribute and render correctly.
    # Replicate that XML shape on preserved student markers: strip the
    # disabling w:val from existing elements, or insert fresh enabled ones.
    # Idempotent. Returns how many OOXML properties were inserted or changed so
    # callers can persist marker-only repairs.
    changed = 0
    if _ensure_paragraph_bool_property_active(marker_p, "w:pageBreakBefore", prepend=True):
        changed += 1
    if _ensure_paragraph_bool_property_active(marker_p, "w:keepNext"):
        changed += 1
    return changed


def _normalise_ordinary_continuation_anchors(doc: Document) -> int:
    """Anchor ordinary table continuation marker chains.

    Scope is intentionally narrow: only body-level
    tbl -> "Продолжение таблицы N" -> [optional blank] -> tbl chains are
    normalized. Appendix continuation labels are not matched by
    `_is_any_continuation_marker` and are left untouched.
    """
    body = doc.element.body
    children = list(body)
    para_by_xml = {p._element: p for p in doc.paragraphs}
    changed = 0

    def _is_empty_p(elem) -> bool:
        if elem.tag != qn("w:p"):
            return False
        p_obj = para_by_xml.get(elem)
        return p_obj is not None and not _norm_text(p_obj.text or "")

    i = 0
    while i < len(children):
        marker_node = children[i]
        if marker_node.tag != qn("w:p"):
            i += 1
            continue

        marker_text = _paragraph_text_from_xml(marker_node)
        if not _is_any_continuation_marker(marker_text):
            i += 1
            continue

        prev_is_table = i > 0 and children[i - 1].tag == qn("w:tbl")
        next_idx = i + 1
        blank_node = None
        if next_idx < len(children) and _is_empty_p(children[next_idx]):
            blank_node = children[next_idx]
            next_idx += 1
        if next_idx >= len(children) or children[next_idx].tag != qn("w:tbl"):
            i += 1
            continue
        if not prev_is_table:
            i += 1
            continue

        marker_changed = 0
        if _ensure_paragraph_bool_property_active(marker_node, "w:pageBreakBefore", prepend=True):
            marker_changed += 1
        if _ensure_paragraph_bool_property_active(marker_node, "w:keepNext"):
            marker_changed += 1
        if marker_changed:
            changed += marker_changed
            logger.info(
                "ordinary_continuation_anchor_marker_normalized marker=%s changes=%s",
                marker_text,
                marker_changed,
            )

        if blank_node is not None:
            if _ensure_paragraph_bool_property_active(blank_node, "w:keepNext"):
                changed += 1
                logger.info(
                    "ordinary_continuation_anchor_blank_normalized marker=%s",
                    marker_text,
                )

        i += 1

    return changed


def _row_matches_line(sig: RowSignature, line_text: str) -> bool:
    pos = 0
    for fragment in sig.fragments:
        found = line_text.find(fragment, pos)
        if found < 0:
            return False
        pos = found + len(fragment)
    return True


def _match_row_pages(table_sig: TableSignature, pdf_lines: list[PdfLine]) -> dict[int, int] | None:
    data_rows = [sig for sig in table_sig.rows if sig.row_idx > 0]
    if len(data_rows) < 2:
        return None

    keys = [sig.key for sig in data_rows]
    if len(keys) != len(set(keys)):
        return None

    line_texts = [(_norm_match_text(line.text), line.page_num) for line in pdf_lines]
    row_pages: dict[int, int] = {}
    last_match_idx = -1

    for sig in data_rows:
        matches = [
            (idx, page_num)
            for idx, (line_text, page_num) in enumerate(line_texts)
            if idx > last_match_idx and _row_matches_line(sig, line_text)
        ]
        if len(matches) != 1:
            return None
        last_match_idx, page_num = matches[0]
        row_pages[sig.row_idx] = page_num

    return row_pages


def _match_row_pages_relaxed_for_appendix(
    table_sig: TableSignature,
    pdf_lines: list[PdfLine],
    *,
    min_matched_data_rows: int = 4,
    max_window_size: int = 3,
) -> dict[int, int] | None:
    """P2-a' fallback: appendix-only row→page matcher that is tolerant of
    (a) duplicate row signatures — uses sequential positional matching;
    (b) wrapped cell text — concatenates a small window of adjacent PDF
    lines before fragment matching.

    Returns the row_idx → page_num mapping only when:
    - at least `min_matched_data_rows` data rows mapped successfully;
    - mapped pages are non-decreasing in row order (monotonic);
    - at least 2 distinct pages observed (otherwise no multi-page span).

    Fail-open: returns None when confidence is low. This matcher is only
    used inside the P2-a appendix continuation collector — strict
    `_match_row_pages` is untouched and remains the primary path.
    """
    data_rows = [sig for sig in table_sig.rows if sig.row_idx > 0]
    if len(data_rows) < min_matched_data_rows:
        return None

    line_texts = [(_norm_match_text(line.text), line.page_num) for line in pdf_lines]
    if not line_texts:
        return None

    row_pages: dict[int, int] = {}
    last_match_idx = -1
    last_page = -1

    for sig in data_rows:
        match_idx = -1
        match_page: int | None = None
        # Sequential scan AFTER previous match — no global uniqueness needed.
        for idx in range(last_match_idx + 1, len(line_texts)):
            line_text, page_num = line_texts[idx]
            # Skip lines that would create a non-monotonic page jump.
            if last_page >= 0 and page_num < last_page:
                continue
            # 1) try single-line match
            if _row_matches_line(sig, line_text):
                match_idx, match_page = idx, page_num
                break
            # 2) try window aggregation (2..max_window_size adjacent lines)
            #    — concatenate normalized texts. PDF line-wrap may either
            #    insert a space at the boundary ("foo" + "bar") or break
            #    inside a hyphenated token ("foo-" + "bar"), so try both
            #    " " and "" separators.
            found_in_window = False
            for win in range(2, max_window_size + 1):
                if idx + win > len(line_texts):
                    break
                chunks = [t for t, _ in line_texts[idx : idx + win]]
                for sep in (" ", ""):
                    combined_text = sep.join(chunks)
                    if _row_matches_line(sig, combined_text):
                        match_idx = idx + win - 1
                        match_page = line_texts[match_idx][1]
                        found_in_window = True
                        break
                if found_in_window:
                    break
            if found_in_window:
                break
            # 3) tables with many narrow columns (e.g. Bondarev 5-col appendix)
            #    render with cell texts on SEPARATE PDF lines — no window of
            #    N adjacent lines can contain all N fragments in order.
            #    Fall back to a strict first-fragment match: each row is
            #    identified by its leftmost cell text only. Require the
            #    fragment to be ≥ 6 chars to stay distinctive.
            first_frag = sig.fragments[0] if sig.fragments else ""
            if len(first_frag) >= 6 and line_text.strip() == first_frag:
                match_idx, match_page = idx, page_num
                break

        if match_idx < 0:
            # Cannot find this row; continue (fail-open per-row).
            continue
        row_pages[sig.row_idx] = match_page  # type: ignore[assignment]
        last_match_idx = match_idx
        last_page = match_page  # type: ignore[assignment]

    if len(row_pages) < min_matched_data_rows:
        return None

    # Monotonic page check (sequential in row order).
    pages_in_row_order = [row_pages[r] for r in sorted(row_pages.keys())]
    if pages_in_row_order != sorted(pages_in_row_order):
        return None

    # Require at least 2 distinct pages for "multi-page span" verdict.
    if len(set(pages_in_row_order)) < 2:
        return None

    return row_pages


_TOKEN_RE = re.compile(r"[0-9A-Za-zА-Яа-яЁё]+")


def _distinctive_tokens(text: str) -> set[str]:
    tokens = {
        token.lower()
        for token in _TOKEN_RE.findall(_norm_text(text))
        if len(token) >= 4 and not token.isdigit()
    }
    return tokens


def _row_distinctive_tokens(sig: RowSignature) -> set[str]:
    out: set[str] = set()
    for fragment in sig.fragments:
        out.update(_distinctive_tokens(fragment))
    return out


def _unique_data_row_tokens(data_rows: list[RowSignature]) -> dict[int, set[str]]:
    all_tokens: dict[str, int] = {}
    row_tokens: dict[int, set[str]] = {}
    for sig in data_rows:
        tokens = _row_distinctive_tokens(sig)
        row_tokens[sig.row_idx] = tokens
        for token in tokens:
            all_tokens[token] = all_tokens.get(token, 0) + 1
    return {
        row_idx: {token for token in tokens if all_tokens[token] == 1}
        for row_idx, tokens in row_tokens.items()
    }


def _line_matches_caption_number(line_text: str, num: str) -> bool:
    m = _TBL_NUM_RE.match(_norm_text(line_text))
    return bool(m and m.group(1) == num)


_DISABLED_PAGE_BREAK_VALUES = {"0", "false", "False", "off"}


def _is_active_page_break_before(page_break_elem) -> bool:
    if page_break_elem is None:
        return False
    return page_break_elem.get(qn("w:val")) not in _DISABLED_PAGE_BREAK_VALUES


def _find_page_break_before(pPr):
    if pPr is None:
        return None
    return pPr.find(qn("w:pageBreakBefore"))


def _pdf_caption_match_count(caption_num: str, pdf_lines: list[PdfLine]) -> int:
    return sum(1 for line in pdf_lines if _line_matches_caption_number(line.text, caption_num))


def _row_has_any_token_in_text(sig: RowSignature, text: str) -> bool:
    tokens = _row_distinctive_tokens(sig)
    if not tokens:
        return False
    text_tokens = _distinctive_tokens(text)
    return bool(tokens & text_tokens)


def _row_fragment_token_hits(sig: RowSignature, text: str) -> int:
    text_tokens = _distinctive_tokens(text)
    if not text_tokens:
        return 0
    hits = 0
    for fragment in sig.fragments:
        fragment_tokens = _distinctive_tokens(fragment)
        if fragment_tokens and fragment_tokens & text_tokens:
            hits += 1
    return hits


def _row_complete_in_page_texts(sig: RowSignature, page_texts: list[str]) -> bool:
    joined = " ".join(page_texts)
    if _row_matches_line(sig, joined):
        return True
    nonempty_fragments = [
        fragment for fragment in sig.fragments
        if _distinctive_tokens(fragment)
    ]
    if not nonempty_fragments:
        return False
    return _row_fragment_token_hits(sig, joined) == len(nonempty_fragments)


def _page_has_complete_data_row(data_rows: list[RowSignature], page_texts: list[str]) -> bool:
    return any(_row_complete_in_page_texts(sig, page_texts) for sig in data_rows)


def _page_has_visible_data_row(data_rows: list[RowSignature], page_texts: list[str]) -> bool:
    joined = " ".join(page_texts)
    for sig in data_rows:
        if _row_complete_in_page_texts(sig, page_texts):
            return True
        required_hits = min(2, len([fragment for fragment in sig.fragments if _distinctive_tokens(fragment)]))
        if required_hits > 0 and _row_fragment_token_hits(sig, joined) >= required_hits:
            return True
    return False


def _tokens_in_text(tokens: set[str], text: str) -> bool:
    if not tokens:
        return False
    return bool(tokens & _distinctive_tokens(text))


def _has_complete_data_row_in_page_window(
    data_rows: list[RowSignature],
    unique_tokens: dict[int, set[str]],
    page_texts: list[str],
) -> bool:
    max_window = 4
    for sig in data_rows:
        tokens = unique_tokens[sig.row_idx]
        if len(tokens) < 2:
            continue
        for start in range(len(page_texts)):
            for end in range(start + 1, min(len(page_texts), start + max_window) + 1):
                window_text = " ".join(page_texts[start:end])
                if tokens <= _distinctive_tokens(window_text):
                    return True
    return False


def _header_line_indexes(header: RowSignature, page_texts: list[str]) -> list[int]:
    header_search_limit = min(len(page_texts), 12)
    return [
        idx
        for idx, line_text in enumerate(page_texts[:header_search_limit])
        if _row_matches_line(header, line_text) or _row_has_any_token_in_text(header, line_text)
    ]


def _first_data_row_spills_to_next_page(
    first_row: RowSignature,
    first_row_tokens: set[str],
    header: RowSignature,
    caption_idx: int,
    start_page: int,
    pdf_lines: list[PdfLine],
    data_page_texts: list[str],
) -> bool:
    if len(first_row_tokens) < 2:
        return False

    start_page_joined = " ".join(data_page_texts)
    start_page_distinctive = _distinctive_tokens(start_page_joined)
    start_tokens = first_row_tokens & start_page_distinctive
    if len(start_tokens) < 2:
        return False
    next_only_tokens = first_row_tokens - start_tokens
    if not next_only_tokens:
        return False

    next_page = start_page + 1
    next_page_lines = [
        line
        for idx, line in enumerate(pdf_lines)
        if idx > caption_idx and line.page_num == next_page
    ]
    if not next_page_lines:
        return False

    next_page_texts = [_norm_match_text(line.text) for line in next_page_lines]
    next_header_indexes = _header_line_indexes(header, next_page_texts)
    if not next_header_indexes:
        return False

    next_page_data_texts = next_page_texts[(max(next_header_indexes) + 1):]
    if not next_page_data_texts:
        return False

    # Conservative continuation evidence: the next page should expose a very
    # short residue of the same first row right after the repeated header, not
    # a long prose line that happens to reuse one token.
    for start in range(min(len(next_page_data_texts), 4)):
        for end in range(start + 1, min(len(next_page_data_texts), start + 2) + 1):
            window_tokens = _distinctive_tokens(" ".join(next_page_data_texts[start:end]))
            if len(window_tokens) <= 3 and (window_tokens & next_only_tokens):
                return True

    return False


def _classify_start_page_usability(
    table_sig: TableSignature,
    caption_num: str,
    pdf_lines: list[PdfLine],
) -> str:
    """
    Conservative Patch 2.1 detector.

    It does not reconstruct the table. It only answers whether the rendered
    caption page contains one clearly complete data row. Ambiguous evidence is
    intentionally treated as no-op by the caller.
    """
    caption_matches = [
        (idx, line)
        for idx, line in enumerate(pdf_lines)
        if _line_matches_caption_number(line.text, caption_num)
    ]
    if len(caption_matches) != 1:
        return _START_AMBIGUOUS

    caption_idx, caption_line = caption_matches[0]
    start_page = caption_line.page_num
    same_page_lines = [
        line
        for idx, line in enumerate(pdf_lines)
        if idx > caption_idx and line.page_num == start_page
    ]
    if not same_page_lines:
        return _START_AMBIGUOUS

    header = next((sig for sig in table_sig.rows if sig.row_idx == 0), None)
    data_rows = [
        sig
        for sig in table_sig.rows
        if sig.row_idx > 0 and not _is_docx_numeric_row(list(sig.fragments))
    ]
    if header is None or not data_rows:
        return _START_AMBIGUOUS

    same_page_texts = [_norm_match_text(line.text) for line in same_page_lines]
    same_page_joined = " ".join(same_page_texts)
    header_line_indexes = _header_line_indexes(header, same_page_texts)
    if not header_line_indexes and not _row_has_any_token_in_text(header, same_page_joined):
        return _START_AMBIGUOUS

    data_page_texts = same_page_texts[(max(header_line_indexes) + 1):] if header_line_indexes else same_page_texts
    if _page_has_complete_data_row(data_rows, data_page_texts):
        return _START_HAS_COMPLETE_DATA_ROW

    next_page = start_page + 1
    next_page_texts = [
        _norm_match_text(line.text)
        for idx, line in enumerate(pdf_lines)
        if idx > caption_idx and line.page_num == next_page
    ]
    if next_page_texts and _page_has_visible_data_row(data_rows, next_page_texts):
        return _START_NO_COMPLETE_DATA_ROW

    data_keys = [sig.key for sig in data_rows]
    if len(data_keys) != len(set(data_keys)):
        return _START_AMBIGUOUS
    unique_tokens = _unique_data_row_tokens(data_rows)
    if any(not unique_tokens.get(sig.row_idx) for sig in data_rows):
        return _START_AMBIGUOUS

    first_row = data_rows[0]
    if _first_data_row_spills_to_next_page(
        first_row=first_row,
        first_row_tokens=unique_tokens[first_row.row_idx],
        header=header,
        caption_idx=caption_idx,
        start_page=start_page,
        pdf_lines=pdf_lines,
        data_page_texts=data_page_texts,
    ):
        return _START_NO_COMPLETE_DATA_ROW

    data_page_joined = " ".join(data_page_texts)
    rows_with_start_page_tokens = [
        sig for sig in data_rows if _tokens_in_text(unique_tokens[sig.row_idx], data_page_joined)
    ]
    if not rows_with_start_page_tokens:
        return _START_NO_COMPLETE_DATA_ROW

    later_text = " ".join(
        _norm_match_text(line.text)
        for idx, line in enumerate(pdf_lines)
        if idx > caption_idx and line.page_num > start_page
    )
    split_like_rows = [
        sig for sig in rows_with_start_page_tokens
        if _tokens_in_text(unique_tokens[sig.row_idx], later_text)
    ]
    if len(split_like_rows) == 1:
        return _START_NO_COMPLETE_DATA_ROW

    return _START_AMBIGUOUS


def _find_rendered_whole_table_move_candidate(
    doc: Document,
    pdf_lines: list[PdfLine],
    diagnostics: dict[str, bool] | None = None,
) -> RenderedWholeTableMoveCandidate | None:
    manual_skip = _valid_manual_continuation_table_indexes(doc)
    inspected = 0

    for table_sig in _collect_table_signatures(doc):
        inspected += 1
        if table_sig.table_idx in manual_skip:
            logger.info(
                "rendered_whole_table_candidate table_idx=%s skip=valid_manual_continuation",
                table_sig.table_idx,
            )
            continue

        caption = _find_caption_paragraph_before_table(doc, table_sig.tbl_xml)
        if caption is None:
            logger.info(
                "rendered_whole_table_candidate table_idx=%s skip=caption_missing",
                table_sig.table_idx,
            )
            continue
        caption_para_xml, caption_num = caption
        pdf_caption_matches = _pdf_caption_match_count(caption_num, pdf_lines)
        caption_pPr = caption_para_xml.find(qn("w:pPr"))
        if _is_active_page_break_before(_find_page_break_before(caption_pPr)):
            logger.info(
                "rendered_whole_table_candidate table_idx=%s caption=%s pdf_caption_matches=%s strict_caption_found=%s skip=existing_active_page_break",
                table_sig.table_idx,
                caption_num,
                pdf_caption_matches,
                pdf_caption_matches == 1,
            )
            continue

        usability = _classify_start_page_usability(table_sig, caption_num, pdf_lines)
        logger.info(
            "rendered_whole_table_candidate table_idx=%s caption=%s pdf_caption_matches=%s strict_caption_found=%s start_page_usability=%s",
            table_sig.table_idx,
            caption_num,
            pdf_caption_matches,
            pdf_caption_matches == 1,
            usability,
        )
        if usability == _START_NO_COMPLETE_DATA_ROW:
            logger.info(
                "rendered_whole_table_candidate_selected table_idx=%s caption=%s reason=%s",
                table_sig.table_idx,
                caption_num,
                usability,
            )
            return RenderedWholeTableMoveCandidate(
                table_idx=table_sig.table_idx,
                tbl_xml=table_sig.tbl_xml,
                caption_para_xml=caption_para_xml,
            )

        if usability == _START_AMBIGUOUS and diagnostics is not None:
            diagnostics["ambiguous"] = True
        logger.info(
            "rendered_whole_table_candidate table_idx=%s caption=%s skip=%s",
            table_sig.table_idx,
            caption_num,
            "ambiguous" if usability == _START_AMBIGUOUS else "has_complete_data_row",
        )

    logger.info("rendered_whole_table_no_candidate inspected=%s", inspected)
    return None


def _ensure_page_break_before(para_elem) -> bool:
    pPr = para_elem.find(qn("w:pPr"))
    page_break = _find_page_break_before(pPr)
    if _is_active_page_break_before(page_break):
        return False
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para_elem.insert(0, pPr)
    if page_break is None:
        pPr.append(OxmlElement("w:pageBreakBefore"))
    else:
        page_break.attrib.pop(qn("w:val"), None)
    return True


def _is_blank_paragraph_xml(p_xml) -> bool:
    if p_xml is None or p_xml.tag != qn("w:p"):
        return False
    return not "".join((node.text or "") for node in p_xml.findall(".//" + qn("w:t"))).strip()


def _set_table_start_orphan_blank_spacing(p_xml) -> bool:
    if p_xml is None or p_xml.tag != qn("w:p"):
        return False

    changed = False
    p_pr = p_xml.find(qn("w:pPr"))
    if p_pr is None:
        p_pr = OxmlElement("w:pPr")
        p_xml.insert(0, p_pr)
        changed = True

    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
        changed = True

    desired = {
        qn("w:before"): "0",
        qn("w:after"): "0",
        qn("w:line"): "840",
        qn("w:lineRule"): "auto",
    }
    for attr, value in desired.items():
        if spacing.get(attr) != value:
            spacing.set(attr, value)
            changed = True

    if p_xml.find(qn("w:r")) is None:
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text.text = ""
        run.append(text)
        p_xml.append(run)
        changed = True

    return changed


def _count_blank_paragraphs_before(para_elem) -> int:
    parent = para_elem.getparent()
    if parent is None:
        return 0
    children = list(parent)
    try:
        idx = children.index(para_elem)
    except ValueError:
        return 0
    count = 0
    j = idx - 1
    while j >= 0 and _is_blank_paragraph_xml(children[j]):
        count += 1
        j -= 1
    return count


def _insert_table_start_orphan_blanks(para_elem, *, target_count: int = 2) -> bool:
    parent = para_elem.getparent()
    if parent is None:
        return False

    children = list(parent)
    try:
        idx = children.index(para_elem)
    except ValueError:
        return False

    blanks: list = []
    j = idx - 1
    while j >= 0 and _is_blank_paragraph_xml(children[j]):
        blanks.append(children[j])
        j -= 1

    changed = False
    while len(blanks) < target_count:
        blank = OxmlElement("w:p")
        _set_table_start_orphan_blank_spacing(blank)
        para_elem.addprevious(blank)
        blanks.insert(0, blank)
        changed = True

    for blank in blanks[:target_count]:
        if _set_table_start_orphan_blank_spacing(blank):
            changed = True

    return changed


def _same_table_start_orphan_remains(docx_path: Path, table_idx: int) -> bool:
    pdf_path: Path | None = None
    try:
        pdf_path = render_docx_to_pdf(docx_path)
        pdf_lines = analyze_pdf_lines(pdf_path)
        doc = Document(str(docx_path))
        candidate = _find_rendered_whole_table_move_candidate(doc, pdf_lines)
        return candidate is not None and candidate.table_idx == table_idx
    except Exception as exc:
        logger.warning(
            "table_start_orphan_validation_failed table_idx=%s error=%s",
            table_idx,
            exc,
        )
        return True
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)


def apply_rendered_table_start_orphan_guard(docx_path: str | Path, report=None) -> int:
    """
    Final rendered table-start orphan guard.

    This intentionally applies only the accepted whole-table-start move:
    exactly two blank paragraphs before the caption/table block, no page break,
    no continuation marker, and no table split. It is safe to run after later
    geometry-changing table passes such as same-page fragment normalization.
    """
    docx_path = Path(docx_path)
    pdf_path: Path | None = None
    try:
        pdf_path = render_docx_to_pdf(docx_path)
        pdf_lines = analyze_pdf_lines(pdf_path)
    except LibreOfficeNotFoundError as exc:
        _warn_rendered_split_unavailable(report, str(exc))
        logger.info(
            "table_start_orphan_guard_skip path=%s reason=libreoffice_unavailable",
            docx_path,
        )
        return 0
    except Exception as exc:
        _warn_rendered_split_unavailable(report, str(exc))
        logger.info(
            "table_start_orphan_guard_skip path=%s reason=render_or_pdf_analysis_failed error=%s",
            docx_path,
            exc,
        )
        return 0
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)

    try:
        doc = Document(str(docx_path))
    except Exception as exc:
        logger.info(
            "table_start_orphan_guard_skip path=%s reason=docx_load_failed error=%s",
            docx_path,
            exc,
        )
        return 0

    diagnostics: dict[str, bool] = {"ambiguous": False}
    move_candidate = _find_rendered_whole_table_move_candidate(doc, pdf_lines, diagnostics)
    if move_candidate is None:
        logger.info(
            "table_start_orphan_guard_skip path=%s reason=%s",
            docx_path,
            "ambiguous" if diagnostics.get("ambiguous") else "no_candidate",
        )
        return 0

    validation_backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_final_table_start_orphan_"))
    validation_backup_path = validation_backup_dir / docx_path.name
    try:
        shutil.copy2(docx_path, validation_backup_path)
        if not _insert_table_start_orphan_blanks(move_candidate.caption_para_xml, target_count=2):
            logger.info(
                "table_start_orphan_guard_skip path=%s table_idx=%s reason=already_has_target_blanks",
                docx_path,
                move_candidate.table_idx,
            )
            return 0

        doc.save(str(docx_path))
        if _same_table_start_orphan_remains(docx_path, move_candidate.table_idx):
            shutil.copy2(validation_backup_path, docx_path)
            logger.info(
                "table_start_orphan_guard_rollback path=%s table_idx=%s reason=orphan_remains",
                docx_path,
                move_candidate.table_idx,
            )
            return 0

        logger.info(
            "table_start_orphan_guard_applied path=%s table_idx=%s blanks=2",
            docx_path,
            move_candidate.table_idx,
        )
        return 1
    except Exception as exc:
        try:
            shutil.copy2(validation_backup_path, docx_path)
        except Exception:
            logger.exception(
                "table_start_orphan_guard_rollback_failed path=%s table_idx=%s",
                docx_path,
                move_candidate.table_idx,
            )
        logger.info(
            "table_start_orphan_guard_rollback path=%s table_idx=%s reason=exception error=%s",
            docx_path,
            move_candidate.table_idx,
            exc,
        )
        return 0
    finally:
        shutil.rmtree(validation_backup_dir, ignore_errors=True)


def _find_rendered_split_candidate(
    doc: Document,
    pdf_lines: list[PdfLine],
    diagnostics: dict[str, bool] | None = None,
) -> RenderedSplitCandidate | None:
    manual_skip = _valid_manual_continuation_table_indexes(doc)
    inspected = 0

    for table_sig in _collect_table_signatures(doc):
        inspected += 1
        if table_sig.table_idx in manual_skip:
            logger.info(
                "rendered_split_candidate table_idx=%s skip=valid_manual_continuation",
                table_sig.table_idx,
            )
            continue

        rows_xml = table_sig.tbl_xml.findall(qn("w:tr"))
        if len(rows_xml) < 3:
            logger.info(
                "rendered_split_candidate table_idx=%s rows=%s skip=too_few_rows",
                table_sig.table_idx,
                len(rows_xml),
            )
            continue

        row_pages = _match_row_pages(table_sig, pdf_lines)
        if row_pages is None:
            if diagnostics is not None:
                diagnostics["ambiguous"] = True
            logger.info(
                "rendered_split_candidate table_idx=%s rows=%s skip=row_mapping_ambiguous",
                table_sig.table_idx,
                len(rows_xml),
            )
            continue

        page_boundary_found = False
        for row_idx in sorted(row_pages):
            next_idx = row_idx + 1
            if next_idx not in row_pages:
                continue
            if row_pages[row_idx] < row_pages[next_idx]:
                page_boundary_found = True
                safe_after = _find_safe_split_after(rows_xml, row_idx)
                if safe_after is None or safe_after < 1:
                    if diagnostics is not None:
                        diagnostics["ambiguous"] = True
                    logger.info(
                        "rendered_split_candidate table_idx=%s row_idx=%s skip=merged_boundary_conflict",
                        table_sig.table_idx,
                        row_idx,
                    )
                    return None
                if len(rows_xml) - (safe_after + 1) < 1:
                    logger.info(
                        "rendered_split_candidate table_idx=%s row_idx=%s safe_after=%s skip=no_continuation_data_row",
                        table_sig.table_idx,
                        row_idx,
                        safe_after,
                    )
                    return None
                logger.info(
                    "rendered_split_candidate_selected table_idx=%s row_idx=%s split_after=%s",
                    table_sig.table_idx,
                    row_idx,
                    safe_after,
                )
                return RenderedSplitCandidate(
                    table_idx=table_sig.table_idx,
                    tbl_xml=table_sig.tbl_xml,
                    split_after=safe_after,
                )
        if not page_boundary_found:
            logger.info(
                "rendered_split_candidate table_idx=%s rows=%s skip=no_page_boundary",
                table_sig.table_idx,
                len(rows_xml),
            )

    logger.info("rendered_split_no_candidate inspected=%s", inspected)
    return None


def _build_continuation_para(text: str):
    """
    Create:
      - right align
      - Times New Roman 14 pt
      - no first-line indent
      - pageBreakBefore=True
      - keepWithNext=True
    """
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    p.append(pPr)

    page_break = OxmlElement("w:pageBreakBefore")
    pPr.append(page_break)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)

    ind = OxmlElement("w:ind")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:left"), "0")
    pPr.append(ind)

    keep_next = OxmlElement("w:keepNext")
    pPr.append(keep_next)

    r = OxmlElement("w:r")
    p.append(r)
    rPr = OxmlElement("w:rPr")
    r.append(rPr)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    fonts.set(qn("w:cs"), "Times New Roman")
    rPr.append(fonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "28")  # 14 pt
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), "28")
    rPr.append(szCs)

    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    return p


def _split_table_at(doc: Document, tbl_xml, split_after: int, continuation_text: str) -> bool:
    rows = tbl_xml.findall(qn("w:tr"))
    if len(rows) < 3:  # header + at least 2 data rows to split
        return False
    if split_after < 1 or split_after >= len(rows) - 1:
        return False

    header_row = deepcopy(rows[0])
    tail_rows = [deepcopy(r) for r in rows[split_after + 1:]]
    if not tail_rows:
        return False

    # part2 must have at least header + 1 data row
    if len(tail_rows) < 1:
        return False

    tbl2 = deepcopy(tbl_xml)
    for tr in list(tbl2.findall(qn("w:tr"))):
        tbl2.remove(tr)
    tbl2.append(header_row)
    for tr in tail_rows:
        tbl2.append(tr)

    # mark repeated header row
    trPr = header_row.find(qn("w:trPr"))
    if trPr is None:
        trPr = OxmlElement("w:trPr")
        header_row.insert(0, trPr)
    if trPr.find(qn("w:tblHeader")) is None:
        trPr.append(OxmlElement("w:tblHeader"))

    # trim part1
    for tr in rows[split_after + 1:]:
        tbl_xml.remove(tr)

    body = doc.element.body
    marker = _build_continuation_para(continuation_text)
    tbl_xml.addnext(marker)
    marker.addnext(tbl2)
    return True



_NUMERIC_CELL_RE = re.compile(r"^[\d\s\+\-−–,.%]+$")
_PT_PER_CHAR_NUMERIC = 6.0   # approx pt/char for 12pt TNR digits
_CELL_H_PADDING = 8.0        # left+right cell padding (pt) added to content width


def _compute_col_minimums(tbl_xml, n_cols: int) -> list[float]:
    """
    Compute per-column minimum widths (pt) in a single pass over all rows.

    For cells containing only numbers/symbols (no letters), the minimum is
    set to the width needed to render the longest value on one line:
        min_w = len(text) × _PT_PER_CHAR_NUMERIC + _CELL_H_PADDING

    For all other cells (header or text), the minimum falls back to _MIN_COL_PT.
    This protects numeric columns (e.g. "9 503 005") from being scaled so narrow
    that values wrap to multiple lines.

    Only single-column cells (gridSpan = 1) are considered.
    """
    minimums = [_MIN_COL_PT] * n_cols

    for tr in tbl_xml.findall(qn("w:tr")):
        col_idx = 0
        for tc in tr.findall(qn("w:tc")):
            if col_idx >= n_cols:
                break
            tcPr = tc.find(qn("w:tcPr"))
            gs = tcPr.find(qn("w:gridSpan")) if tcPr is not None else None
            span = int(gs.get(qn("w:val"), 1)) if gs is not None else 1
            span = max(1, min(span, n_cols - col_idx))

            if span == 1:
                for p_el in tc.findall(".//" + qn("w:p")):
                    cell_text = "".join(
                        (r.find(qn("w:t")).text or "")
                        for r in p_el.findall(qn("w:r"))
                        if r.find(qn("w:t")) is not None
                    ).strip()
                    if cell_text and _NUMERIC_CELL_RE.match(cell_text):
                        content_w = len(cell_text) * _PT_PER_CHAR_NUMERIC + _CELL_H_PADDING
                        if content_w > minimums[col_idx]:
                            minimums[col_idx] = content_w

            col_idx += span

    return minimums


def _optimize_table_col_widths(tbl_xml, body_width_pt: float) -> bool:
    """
    Ensure no column is narrower than its content minimum and total width ≤ body_width_pt.

    Algorithm:
      1. Scale all columns down proportionally if total > body_width_pt.
      2. Identify undersized columns (based on content-aware per-column minimums);
         redistribute deficit from wider donor columns.

    The per-column minimums are content-aware: numeric-only cells (digits, spaces,
    punctuation) set a minimum wide enough to display their content on one line.
    This prevents number columns from being scaled too narrow when proportionally
    shrinking a wide table.

    Updates both w:tblGrid/w:gridCol, w:tblPr/w:tblW, and each w:tc/w:tcPr/w:tcW
    (honouring w:gridSpan for merged cells).

    Returns True if any width was changed.
    """
    grid = tbl_xml.find(qn("w:tblGrid"))
    if grid is None:
        return False
    gridcols = grid.findall(qn("w:gridCol"))
    if not gridcols:
        return False

    widths = [int(c.get(qn("w:w"), 0)) / TWIP_PER_PT for c in gridcols]
    n = len(widths)
    total = sum(widths)
    if total < 1:
        return False

    # Content-aware per-column minimums (protects numeric columns from over-shrinking)
    col_mins = _compute_col_minimums(tbl_xml, n)

    changed = False

    # Step 1: scale down if total exceeds body width
    if total > body_width_pt + 0.5:
        scale = body_width_pt / total
        widths = [w * scale for w in widths]
        total = sum(widths)
        changed = True

    # Step 2: redistribute to fix undersized columns (up to n iterations).
    # Uses per-column minimums: numeric columns have higher minimums to keep
    # values on one line; other columns use the global _MIN_COL_PT floor.
    for _ in range(n):
        undersized = [(i, col_mins[i] - widths[i]) for i in range(n)
                      if widths[i] < col_mins[i] - 0.5]
        if not undersized:
            break
        donors = [i for i in range(n) if widths[i] > col_mins[i] + 0.5]
        if not donors:
            break
        total_deficit = sum(d for _, d in undersized)
        total_donor_excess = sum(widths[i] - col_mins[i] for i in donors)
        take_frac = min(1.0, total_donor_excess / total_deficit)

        for i, deficit in undersized:
            widths[i] += deficit * take_frac
        actual_taken = total_deficit * take_frac
        for i in donors:
            donor_excess = widths[i] - col_mins[i]
            widths[i] -= actual_taken * (donor_excess / total_donor_excess)
        changed = True

    if not changed:
        return False

    # Round to integer twips, keep total consistent
    twip_widths = [max(1, round(w * TWIP_PER_PT)) for w in widths]

    # Apply to grid
    for col_el, tw in zip(gridcols, twip_widths):
        col_el.set(qn("w:w"), str(tw))

    # Update w:tblPr/w:tblW to the new column total.
    # Without this, Word uses the original (too-wide) tblW as master table width
    # and ignores the corrected gridCol / tcW values.
    tblPr = tbl_xml.find(qn("w:tblPr"))
    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:w"), str(sum(twip_widths)))
        tblW.set(qn("w:type"), "dxa")

    # Apply to each row's cells (respecting gridSpan)
    for tr in tbl_xml.findall(qn("w:tr")):
        col_idx = 0
        for tc in tr.findall(qn("w:tc")):
            if col_idx >= n:
                break
            tcPr = tc.find(qn("w:tcPr"))
            gridSpan_el = tcPr.find(qn("w:gridSpan")) if tcPr is not None else None
            span = int(gridSpan_el.get(qn("w:val"), 1)) if gridSpan_el is not None else 1
            span = max(1, min(span, n - col_idx))

            cell_tw = sum(twip_widths[col_idx: col_idx + span])

            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(cell_tw))
            tcW.set(qn("w:type"), "dxa")

            col_idx += span

    return True


def apply_table_merging(doc: Document) -> int:
    """
    Phase 3 pre-pass — STUB (table splitting/merging disabled).

    Previously: detected student-split table pairs (table + "Продолжение
    таблицы X" paragraph + continuation table) and merged them back into one
    table so apply_table_continuation could re-split at the real page boundary.

    Disabled because reliable page-break detection requires a rendering engine
    (LibreOffice / Word) — pure geometry estimation was too unreliable.
    See module docstring for the FUTURE implementation plan.

    Returns 0 (no changes made).
    """
    body = doc.element.body
    children = list(body)
    merges = 0

    i = 1
    while i < len(children) - 1:
        prev_node = children[i - 1]
        node = children[i]
        next_node = children[i + 1]

        if prev_node.tag != qn("w:tbl") or node.tag != qn("w:p") or next_node.tag != qn("w:tbl"):
            i += 1
            continue

        p_obj = next((p for p in doc.paragraphs if p._element is node), None)
        marker_text = _norm_text(p_obj.text if p_obj is not None else "")
        if not _is_any_continuation_marker(marker_text):
            i += 1
            continue

        tbl1 = prev_node
        tbl2 = next_node

        rows1 = tbl1.findall(qn("w:tr"))
        rows2 = tbl2.findall(qn("w:tr"))
        headers_match = bool(rows1 and rows2 and _rows_match(rows1[0], rows2[0]))
        keep_manual_split = _is_valid_manual_continuation_chain(doc, tbl1, node, tbl2)
        keep_student_chain = _is_structurally_valid_student_chain(doc, tbl1, node, tbl2)

        if keep_manual_split or keep_student_chain:
            # P1-critical / DEFECT E: preserved student chains have
            # <w:pageBreakBefore w:val='0'/> + <w:keepNext w:val='0'/> applied
            # by Phase 1 hard_reset, so LibreOffice renders the marker at the
            # bottom of the previous page instead of at the top of the
            # continuation page. Enable both properties to match formatter-
            # authored markers (which render correctly). Formatter-authored
            # chains (keep_manual_split) already pass the strict validator
            # that requires keepNext enabled — they must NOT be touched.
            if keep_student_chain and not keep_manual_split:
                repairs = _enable_marker_page_break_for_student_chain(node)
                if repairs:
                    merges += repairs
                    logger.info(
                        "student_continuation_anchor_marker_normalized marker=%s changes=%s",
                        marker_text,
                        repairs,
                    )
            i += 1
            continue

        # Rebuild invalid split: merge tbl2 into tbl1, skipping duplicate header if present.
        start_idx = 1 if headers_match else 0
        for tr in rows2[start_idx:]:
            tbl1.append(deepcopy(tr))

        parent = node.getparent()
        if parent is not None:
            parent.remove(node)
        parent2 = tbl2.getparent()
        if parent2 is not None:
            parent2.remove(tbl2)
        merges += 1

        # refresh snapshot after mutations
        children = list(body)
        i = max(1, i - 1)

    return merges


# ── Main entry point ──────────────────────────────────────────────────────────

def apply_table_continuation(
    doc: Document,
    report: FormattingReport | None = None,
) -> int:
    """
    Phase 3 Rule 1 — STUB (table page-break splitting disabled).

    Still active: column-width optimisation (_optimize_table_col_widths) runs
    for simple tables. Geometry-sensitive/unsafe tables bypass this pass so
    authored widths and grid topology are not re-inferred before splitting.

    The splitting part is disabled because reliable page-break detection
    requires a rendering engine.  See module docstring for the FUTURE plan.

    Returns the number of width normalisations plus manual-chain numeric-row
    repairs and continuation-anchor normalisations. Does not split tables or
    insert continuation markers.
    """
    anchor_repairs = _normalise_ordinary_continuation_anchors(doc)
    numeric_repairs = _repair_manual_continuation_numeric_rows(doc)

    # ── Column-width optimisation (always active) ──────────────────────────
    body_w = _body_width_pt(doc)
    n_col_fixed = 0
    table_index = -1
    body_children = list(doc.element.body)
    para_by_xml = {p._element: p for p in doc.paragraphs}
    for kind, tbl_xml, _ in _iter_body(doc):
        if kind != "table":
            continue
        table_index += 1
        policy, reasons = _table_geometry_policy_details(
            tbl_xml,
            has_existing_continuation_marker=_table_has_adjacent_continuation_marker(
                body_children,
                para_by_xml,
                tbl_xml,
            ),
        )
        if policy != _GEOMETRY_SIMPLE:
            logger.info(
                "table_continuation: col-width optimizer skipped table_index=%s policy=%s reasons=%s",
                table_index,
                policy,
                _format_geometry_reasons(reasons),
            )
            continue
        if _optimize_table_col_widths(tbl_xml, body_w):
            n_col_fixed += 1
    if n_col_fixed:
        logger.info("table_continuation: col-width optimised %d table(s)", n_col_fixed)

    return n_col_fixed + numeric_repairs + anchor_repairs


def _warn_rendered_split_unavailable(
    report: FormattingReport | None,
    reason: str,
) -> None:
    logger.info("rendered table continuation skipped: %s", reason)
    if report is not None:
        report.warn("Автоперенос таблиц по PDF временно недоступен")


@dataclass(frozen=True)
class _MarkerSplitDecision:
    eligible: bool
    split_before_row: int | None
    skip_reason: str | None


def _marker_split_enabled() -> bool:
    return os.getenv("KPFU_ENABLE_MARKER_SPLIT", "").strip().lower() in {
        "1", "true", "yes", "on",
    }


def _marker_split_apply_enabled() -> bool:
    return os.getenv("KPFU_APPLY_MARKER_SPLIT", "").strip().lower() in {
        "1", "true", "yes", "on",
    }


_MARKER_SPLIT_MAX_RENDERS_DEFAULT = 20
_MARKER_SPLIT_HARD_TIMEOUT_DEFAULT = 300.0
_MARKER_SPLIT_MODE_DEFAULT = "candidate"
_MARKER_SPLIT_MODE_VALID = {"candidate", "global_skip"}


def _marker_split_max_renders() -> int:
    """
    Per-document cap on rendered marker-split *candidates*. Each candidate
    diagnose costs one LibreOffice render. Env override:
    KPFU_MARKER_SPLIT_MAX_RENDERS. Missing/invalid/negative values fall back
    to the default. Default raised to 20 for quality-first candidate mode.
    """
    raw = os.getenv("KPFU_MARKER_SPLIT_MAX_RENDERS", "")
    try:
        value = int(raw)
    except (TypeError, ValueError):
        return _MARKER_SPLIT_MAX_RENDERS_DEFAULT
    if value < 0:
        return _MARKER_SPLIT_MAX_RENDERS_DEFAULT
    return value


def _marker_split_mode() -> str:
    """
    Phase 3 marker split execution mode. Env override KPFU_MARKER_SPLIT_MODE.
      "candidate"   (default) — per-candidate diagnose loop with hard wall-time
                                 cap. User warning only if real candidates were
                                 skipped.
      "global_skip" (legacy)  — pre-E2 behaviour. Whole pass skipped when
                                 total table count exceeds the budget.
    Unknown / missing values fall back to the default.
    """
    raw = (os.getenv("KPFU_MARKER_SPLIT_MODE") or "").strip().lower()
    if raw not in _MARKER_SPLIT_MODE_VALID:
        return _MARKER_SPLIT_MODE_DEFAULT
    return raw


def _marker_split_hard_timeout_seconds() -> float:
    """
    Hard wall-time cap (seconds) for the per-candidate marker-split loop.
    Checked before starting each candidate's diagnose. Does NOT interrupt an
    in-flight diagnose. Env override KPFU_MARKER_SPLIT_HARD_TIMEOUT_SECONDS.
    Defaults to 300 s. Non-positive / invalid → default.
    """
    raw = os.getenv("KPFU_MARKER_SPLIT_HARD_TIMEOUT_SECONDS", "")
    if not raw:
        return _MARKER_SPLIT_HARD_TIMEOUT_DEFAULT
    try:
        value = float(raw)
    except (TypeError, ValueError):
        return _MARKER_SPLIT_HARD_TIMEOUT_DEFAULT
    if value <= 0:
        return _MARKER_SPLIT_HARD_TIMEOUT_DEFAULT
    return value


def _marker_split_num_row_compensation_enabled() -> bool:
    """
    E3: NUM-row compensation for marker-split first fragment.

    When apply_numbered_split_to_document inserts a synthesized "1, 2, 3, …"
    row at the top of the first fragment (because the original table doesn't
    have one), the fragment is one row taller than LO measured. With cantSplit
    on every row (TCF-A), the last data row can be pushed alone to the next
    page (orphan / near-empty page). Compensation: subtract 1 from
    split_before_row so the fragment matches LO's measured capacity.

    Env override KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION. Default ON. Values
    "0", "false", "off", "no" (case-insensitive) disable it; anything else
    enables it.
    """
    raw = (os.getenv("KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION") or "").strip().lower()
    if raw in {"0", "false", "off", "no"}:
        return False
    return True


def _render_block_page(block: dict) -> int:
    try:
        return int(block.get("page") or 0)
    except (TypeError, ValueError):
        return 0


def _render_block_float(block: dict | None, key: str) -> float:
    if block is None:
        return 0.0
    try:
        return float(block.get(key) or 0.0)
    except (TypeError, ValueError):
        return 0.0


def _nearest_previous_rendered_table(tables: list[dict], marker: dict) -> dict | None:
    marker_page = _render_block_page(marker)
    marker_y0 = _render_block_float(marker, "y0")
    candidates: list[dict] = []
    for table in tables:
        page = _render_block_page(table)
        if page < marker_page:
            candidates.append(table)
        elif page == marker_page and _render_block_float(table, "y1") <= marker_y0 + 2.0:
            candidates.append(table)
    if not candidates:
        return None
    return max(candidates, key=lambda item: (_render_block_page(item), _render_block_float(item, "y1")))


def _nearest_following_rendered_table(tables: list[dict], marker: dict) -> dict | None:
    marker_page = _render_block_page(marker)
    marker_y1 = _render_block_float(marker, "y1")
    candidates: list[dict] = []
    for table in tables:
        page = _render_block_page(table)
        if page > marker_page:
            candidates.append(table)
        elif page == marker_page and _render_block_float(table, "y0") >= marker_y1 - 2.0:
            candidates.append(table)
    if not candidates:
        return None
    return min(candidates, key=lambda item: (_render_block_page(item), _render_block_float(item, "y0")))


def _same_page_continuation_marker_violations_from_blocks(
    blocks: list[dict],
) -> list[_SamePageContinuationMarkerViolation]:
    tables = [
        block for block in blocks
        if (block.get("kind") or "").strip().lower() == "table"
        and _render_block_page(block) > 0
    ]
    marker_blocks = [
        block for block in blocks
        if (block.get("kind") or "").strip().lower() == "text"
        and _render_block_page(block) > 0
        and _ANY_CONTINUATION_MARKER_RE.search(str(block.get("text") or ""))
    ]

    violations: list[_SamePageContinuationMarkerViolation] = []
    for marker in sorted(marker_blocks, key=lambda item: (_render_block_page(item), _render_block_float(item, "y0"))):
        text = " ".join(str(marker.get("text") or "").split())
        if not _STRICT_CONTINUATION_MARKER_RE.match(text):
            continue
        previous = _nearest_previous_rendered_table(tables, marker)
        if previous is None:
            continue
        marker_page = _render_block_page(marker)
        previous_page = _render_block_page(previous)
        if previous_page != marker_page:
            continue
        following = _nearest_following_rendered_table(tables, marker)
        following_page = _render_block_page(following) if following is not None else None
        confidence = "high" if following_page == marker_page else "medium"
        violations.append(
            _SamePageContinuationMarkerViolation(
                marker_text=text,
                marker_page=marker_page,
                previous_table_page=previous_page,
                following_table_page=following_page,
                confidence=confidence,
            )
        )
    return violations


def _continuation_marker_line_blocks_from_pdf_page(page, *, page_number: int) -> list[dict]:
    words = page.extract_words(
        x_tolerance=3,
        y_tolerance=3,
        keep_blank_chars=False,
        use_text_flow=False,
    ) or []
    words = sorted(words, key=lambda item: (float(item.get("top", 0.0)), float(item.get("x0", 0.0))))
    lines: list[list[dict]] = []
    for word in words:
        top = float(word.get("top", 0.0))
        if not lines or abs(float(lines[-1][0].get("top", 0.0)) - top) > 3.0:
            lines.append([word])
        else:
            lines[-1].append(word)

    blocks: list[dict] = []
    for line in lines:
        text = " ".join(
            str(word.get("text") or "")
            for word in sorted(line, key=lambda item: float(item.get("x0", 0.0)))
        )
        text = " ".join(text.split())
        if not _ANY_CONTINUATION_MARKER_RE.search(text):
            continue
        blocks.append(
            {
                "kind": "text",
                "page": page_number,
                "text": text,
                "y0": min(float(word.get("top", 0.0)) for word in line),
                "y1": max(float(word.get("bottom", 0.0)) for word in line),
            }
        )
    return blocks


def _continuation_marker_render_blocks_from_pdf(pdf_path: Path) -> list[dict]:
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber is required for Phase 3 continuation marker validation. "
            "Install it: pip install pdfplumber"
        )

    blocks: list[dict] = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            try:
                tables = page.find_tables()
            except Exception:
                tables = []
            for table_idx, table in enumerate(tables):
                _x0, y0, _x1, y1 = table.bbox
                blocks.append(
                    {
                        "kind": "table",
                        "page": page_number,
                        "table_block_index": table_idx,
                        "y0": float(y0),
                        "y1": float(y1),
                    }
                )
            blocks.extend(_continuation_marker_line_blocks_from_pdf_page(page, page_number=page_number))
    return blocks


def _same_page_continuation_marker_violations_for_docx(
    docx_path: Path,
) -> list[_SamePageContinuationMarkerViolation]:
    pdf_path: Path | None = None
    try:
        pdf_path = render_docx_to_pdf(Path(docx_path))
        blocks = _continuation_marker_render_blocks_from_pdf(pdf_path)
        return _same_page_continuation_marker_violations_from_blocks(blocks)
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)


def _rendered_continuation_violations_for_docx(docx_path: Path):
    pdf_path: Path | None = None
    try:
        from .rendered_table_validation import (
            build_rendered_table_identities,
            validate_rendered_continuations,
        )

        pdf_path = render_docx_to_pdf(Path(docx_path))
        pdf_lines = analyze_pdf_lines(pdf_path)
        doc = Document(str(docx_path))
        identities = build_rendered_table_identities(doc)
        return validate_rendered_continuations(pdf_lines, identities)
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)


def _rendered_continuation_deletion_regressions(docx_path: Path) -> list:
    return [
        violation
        for violation in _rendered_continuation_violations_for_docx(Path(docx_path))
        if violation.violation_type in {
            "missing_continuation_marker",
            "suspected_missing_continuation_marker",
        }
    ]


def _remove_strict_continuation_marker_texts(docx_path: Path, marker_texts: set[str]) -> int:
    doc = Document(str(docx_path))
    body = doc.element.body

    removed = 0
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "p":
            continue
        raw = "".join(
            t.text or ""
            for t in child.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        ).strip()
        if raw in marker_texts and _STRICT_CONTINUATION_MARKER_RE.match(raw):
            body.remove(child)
            removed += 1

    if removed:
        doc.save(str(docx_path))
    return removed


def _format_rendered_deletion_regressions(violations: list) -> str:
    details = []
    for violation in violations:
        table_num = getattr(violation, "table_num", None) or f"#{getattr(violation, 'table_index', '?')}"
        page = getattr(violation, "page", "?")
        violation_type = getattr(violation, "violation_type", "?")
        confidence = getattr(violation, "confidence", "?")
        details.append(f"{table_num}@p{page}:{violation_type}:{confidence}")
    return ";".join(details)


def _warn_unsafe_same_page_marker_cleanup(
    report: FormattingReport | None,
    violations: list,
) -> None:
    if report is None:
        return
    seen: set[tuple[str, int, str]] = set()
    for violation in violations:
        table_num = getattr(violation, "table_num", None) or f"#{getattr(violation, 'table_index', '?')}"
        page = int(getattr(violation, "page", 0) or 0)
        violation_type = str(getattr(violation, "violation_type", ""))
        key = (table_num, page, violation_type)
        if key in seen:
            continue
        seen.add(key)
        if violation_type == "missing_continuation_marker":
            report.warn(
                f"Проверьте перенос таблицы {table_num}: "
                f"маркер продолжения сохранён, потому что его удаление оставляет стр. {page} без маркера."
            )
        elif violation_type == "suspected_missing_continuation_marker":
            report.warn(
                f"Проверьте возможный перенос таблицы {table_num}: "
                f"маркер продолжения сохранён, потому что его удаление может оставить стр. {page} без маркера."
            )


def _repair_row_fingerprint(row_xml) -> str:
    return " ".join(
        _norm_match_text(text)
        for text in _row_cell_texts(row_xml)
        if _norm_match_text(text)
    )


def _repair_fingerprint_key(text: str) -> str:
    cleaned = _norm_match_text(text)
    cleaned = re.sub(r"\s*[-–—]\s*", "-", cleaned)
    cleaned = re.sub(r"\s*/\s*", "/", cleaned)
    return " ".join(cleaned.split())


def _repair_text_matches_row(row_xml, row_fingerprint: str) -> bool:
    needle = _repair_fingerprint_key(row_fingerprint)
    if not needle:
        return False
    row_fp = _repair_fingerprint_key(_repair_row_fingerprint(row_xml))
    return bool(row_fp and (needle in row_fp or row_fp in needle))


def _repair_table_has_unsafe_manual_chain_markup(tbl_xml) -> bool:
    for tr in tbl_xml.findall(qn("w:tr")):
        trPr = tr.find(qn("w:trPr"))
        if trPr is not None and trPr.find(qn("w:trHeight")) is not None:
            return True
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                continue
            if tcPr.find(qn("w:gridSpan")) is not None:
                return True
            if tcPr.find(qn("w:vMerge")) is not None:
                return True
    return False


def _repair_paragraph_text(doc: Document, node) -> str:
    para_by_xml = {p._element: p for p in doc.paragraphs}
    paragraph = para_by_xml.get(node)
    if paragraph is not None:
        return _norm_text(paragraph.text)
    return _norm_text("".join(t.text or "" for t in node.findall(".//" + qn("w:t"))))


def _manual_chain_after_table(
    doc: Document,
    table_index: int,
    table_num: str,
) -> tuple[object, int] | None:
    if table_index < 0 or table_index >= len(doc.tables):
        return None
    first_tbl = doc.tables[table_index]._tbl
    children = list(doc.element.body)
    try:
        pos = children.index(first_tbl)
    except ValueError:
        return None

    expected_marker = f"Продолжение таблицы {table_num}"
    marker_seen = False
    next_table_index = table_index + 1
    for node in children[pos + 1:]:
        if node.tag == qn("w:p"):
            text = _repair_paragraph_text(doc, node)
            if not text:
                continue
            if not marker_seen:
                if text == expected_marker and _STRICT_CONTINUATION_MARKER_RE.match(text):
                    marker_seen = True
                    continue
                return None
            return None
        if node.tag == qn("w:tbl"):
            if marker_seen and next_table_index < len(doc.tables) and doc.tables[next_table_index]._tbl is node:
                return node, next_table_index
            return None
    return None


def _manual_chain_overflow_boundary(
    first_tbl,
    second_tbl,
    *,
    row_fingerprint: str,
) -> int | None:
    col_count = _table_col_count(first_tbl)
    if col_count <= 0 or _table_col_count(second_tbl) != col_count:
        return None
    first_rows = first_tbl.findall(qn("w:tr"))
    second_rows = second_tbl.findall(qn("w:tr"))
    if len(first_rows) < 4 or len(second_rows) < 3:
        return None
    if not _rows_match(first_rows[0], second_rows[0]):
        return None
    if not (_row_is_exact_numeric_row(first_rows[1], col_count) and _row_is_exact_numeric_row(second_rows[1], col_count)):
        return None
    if _repair_table_has_unsafe_manual_chain_markup(first_tbl):
        return None
    if _repair_table_has_unsafe_manual_chain_markup(second_tbl):
        return None
    for row_idx in range(2, len(first_rows)):
        if _repair_text_matches_row(first_rows[row_idx], row_fingerprint):
            if row_idx <= 2 or row_idx >= len(first_rows):
                return None
            return row_idx
    return None


def _move_manual_chain_rows(first_tbl, second_tbl, split_before_row: int) -> int:
    first_rows = first_tbl.findall(qn("w:tr"))
    second_rows = second_tbl.findall(qn("w:tr"))
    moving = list(first_rows[split_before_row:])
    if not moving or len(second_rows) < 2:
        return 0
    insert_after = second_rows[1]
    insert_pos = second_tbl.index(insert_after) + 1
    for row in moving:
        first_tbl.remove(row)
    for offset, row in enumerate(moving):
        second_tbl.insert(insert_pos + offset, row)
    return len(moving)


def _target_rendered_violations(violations: list, table_num: str) -> list:
    return [
        violation
        for violation in violations
        if getattr(violation, "table_num", None) == table_num
        and getattr(violation, "violation_type", None) in {
            "missing_continuation_marker",
            "suspected_missing_continuation_marker",
        }
    ]


def _target_same_page_violations(violations: list[_SamePageContinuationMarkerViolation], table_num: str) -> list:
    marker_text = f"Продолжение таблицы {table_num}"
    return [
        violation
        for violation in violations
        if violation.marker_text.strip() == marker_text
    ]


def _attempt_manual_chain_overflow_repair(
    docx_path: Path,
    violation,
    *,
    report: FormattingReport | None = None,
) -> bool:
    table_num = getattr(violation, "table_num", None)
    table_index = int(getattr(violation, "table_index", -1))
    violation_type = getattr(violation, "violation_type", None)
    confidence = getattr(violation, "confidence", None)
    evidence = getattr(violation, "evidence", {}) or {}
    row_fingerprint = str(evidence.get("row_fingerprint") or "")
    if (
        not table_num
        or violation_type != "missing_continuation_marker"
        or confidence != "high"
        or not row_fingerprint
    ):
        return False

    with tempfile.TemporaryDirectory(prefix="kpfu_manual_chain_repair_") as tmp:
        candidate_path = Path(tmp) / Path(docx_path).name
        shutil.copy2(docx_path, candidate_path)
        doc = Document(str(candidate_path))
        if table_index < 0 or table_index >= len(doc.tables):
            return False
        first_tbl = doc.tables[table_index]._tbl
        caption_num = _find_caption_number_before_table(doc, first_tbl)
        if caption_num != table_num:
            return False
        chain = _manual_chain_after_table(doc, table_index, table_num)
        if chain is None:
            return False
        second_tbl, second_table_index = chain
        boundary = _manual_chain_overflow_boundary(
            first_tbl,
            second_tbl,
            row_fingerprint=row_fingerprint,
        )
        if boundary is None:
            return False
        if not _split_geometry_is_safe(
            first_tbl,
            table_index=table_index,
            split_before_row=boundary,
            log_prefix="manual_chain_overflow_repair_skip",
        ):
            return False

        moved = _move_manual_chain_rows(first_tbl, second_tbl, boundary)
        if moved <= 0:
            return False
        doc.save(str(candidate_path))

        try:
            rendered_after = _rendered_continuation_violations_for_docx(candidate_path)
        except Exception as exc:
            logger.info(
                "manual_chain_overflow_repair_rollback table_num=%s table_index=%s reason=rendered_validation_failed error=%s",
                table_num, table_index, exc,
            )
            if report is not None:
                report.warn(
                    f"Проверьте перенос таблицы {table_num}: автоматическая коррекция отменена, PDF-проверка не выполнена."
                )
            return False
        if _target_rendered_violations(rendered_after, table_num):
            logger.info(
                "manual_chain_overflow_repair_rollback table_num=%s table_index=%s reason=rendered_violation_remains",
                table_num, table_index,
            )
            if report is not None:
                report.warn(
                    f"Проверьте перенос таблицы {table_num}: автоматическая коррекция отменена, перенос всё ещё требует проверки."
                )
            return False
        try:
            same_page_after = _same_page_continuation_marker_violations_for_docx(candidate_path)
        except Exception as exc:
            logger.info(
                "manual_chain_overflow_repair_rollback table_num=%s table_index=%s reason=same_page_validation_failed error=%s",
                table_num, table_index, exc,
            )
            if report is not None:
                report.warn(
                    f"Проверьте перенос таблицы {table_num}: автоматическая коррекция отменена, проверка маркеров не выполнена."
                )
            return False
        if _target_same_page_violations(same_page_after, table_num):
            logger.info(
                "manual_chain_overflow_repair_rollback table_num=%s table_index=%s reason=same_page_marker_remains",
                table_num, table_index,
            )
            if report is not None:
                report.warn(
                    f"Проверьте перенос таблицы {table_num}: автоматическая коррекция отменена, маркер остался на той же странице."
                )
            return False

        shutil.copy2(candidate_path, docx_path)
        logger.info(
            "manual_chain_overflow_repair_applied table_num=%s table_index=%s continuation_table_index=%s moved_rows=%s split_before_row=%s",
            table_num, table_index, second_table_index, moved, boundary,
        )
        return True


def repair_manual_chain_overflow_before_marker(
    docx_path: Path,
    violations: list,
    *,
    report: FormattingReport | None = None,
) -> int:
    """
    Repair one narrow high-confidence manual-chain failure:

    table A -> "Продолжение таблицы X" -> table B exists, but the rendered PDF
    shows rows from table A spilling onto the next page before the marker.  Move
    only the rows identified by the rendered validator's row fingerprint into
    table B, after its numeric row.  The candidate DOCX is accepted only when
    repeated PDF validation for the target table is clean.
    """
    repaired = 0
    seen: set[tuple[str, int, int]] = set()
    for violation in violations:
        table_num = getattr(violation, "table_num", None)
        table_index = int(getattr(violation, "table_index", -1))
        page = int(getattr(violation, "page", 0) or 0)
        key = (str(table_num or ""), table_index, page)
        if key in seen:
            continue
        seen.add(key)
        if _attempt_manual_chain_overflow_repair(
            Path(docx_path),
            violation,
            report=report,
        ):
            repaired += 1
    return repaired


def _paragraph_text_xml(p_xml) -> str:
    if p_xml is None or p_xml.tag != qn("w:p"):
        return ""
    return " ".join((node.text or "") for node in p_xml.findall(".//" + qn("w:t"))).strip()


def _table_grid_signature(table) -> tuple[str, ...]:
    grid = table._tbl.tblGrid
    if grid is None:
        return ()
    return tuple(col.get(qn("w:w")) or "" for col in grid.findall(qn("w:gridCol")))


def _table_width_signature(table) -> tuple[str | None, str | None] | None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW")) if tbl_pr is not None else None
    if tbl_w is None:
        return None
    return (tbl_w.get(qn("w:type")), tbl_w.get(qn("w:w")))


def _table_border_signature(table) -> tuple[tuple[str, str | None, str | None, str | None], ...] | None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders")) if tbl_pr is not None else None
    if borders is None:
        return None
    return tuple(
        (
            child.tag.rsplit("}", 1)[-1],
            child.get(qn("w:val")),
            child.get(qn("w:sz")),
            child.get(qn("w:color")),
        )
        for child in borders
    )


def _table_margin_signature(table) -> tuple[tuple[str, str | None, str | None], ...] | None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar")) if tbl_pr is not None else None
    if margins is None:
        return None
    return tuple(
        (
            child.tag.rsplit("}", 1)[-1],
            child.get(qn("w:w")),
            child.get(qn("w:type")),
        )
        for child in margins
    )


def _table_has_merged_cells_docx(table) -> bool:
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.tcPr
            if tc_pr is None:
                continue
            if tc_pr.find(qn("w:gridSpan")) is not None:
                return True
            if tc_pr.find(qn("w:vMerge")) is not None:
                return True
    return False


def _docx_row_cell_texts(row) -> list[str]:
    return [" ".join(cell.text.split()) for cell in row.cells]


def _is_docx_numeric_row(values: list[str]) -> bool:
    return len(values) >= 2 and values == [str(idx) for idx in range(1, len(values) + 1)]


def _docx_row_fingerprint_values(values: list[str]) -> str:
    return " ".join(" ".join(value.split()).lower() for value in values if " ".join(value.split()))


def _docx_row_fingerprint(row) -> str:
    return _docx_row_fingerprint_values(_docx_row_cell_texts(row))


def _docx_data_fingerprints(table) -> list[str]:
    out: list[str] = []
    for idx, row in enumerate(table.rows):
        values = _docx_row_cell_texts(row)
        if idx == 0:
            continue
        if _is_docx_numeric_row(values):
            continue
        fp = _docx_row_fingerprint_values(values)
        if fp:
            out.append(fp)
    return out


def _docx_table_col_count(table) -> int:
    return max((len(row.cells) for row in table.rows), default=0)


def _docx_table_has_numeric_row(table) -> bool:
    return any(_is_docx_numeric_row(_docx_row_cell_texts(row)) for row in table.rows)


def _strict_marker_table_num(text: str) -> str | None:
    match = _STRICT_CONTINUATION_MARKER_RE.match(" ".join((text or "").split()))
    if not match:
        return None
    number = re.search(r"\d+(?:\.\d+)*", match.group(0))
    return number.group(0) if number else None


def _caption_table_num(text: str) -> str | None:
    match = re.match(r"^\s*Таблица\s+(\d+(?:\.\d+)*)\b", " ".join((text or "").split()), re.IGNORECASE)
    return match.group(1) if match else None


def _source_or_note_text(text: str) -> bool:
    return bool(re.match(r"^\s*(Источник|Примечание)\s*:", text or "", re.IGNORECASE))


def _doc_table_xml_map(doc: Document) -> dict[object, int]:
    return {table._tbl: idx for idx, table in enumerate(doc.tables)}


def _find_strict_marker_paragraph(doc: Document, marker_text: str):
    wanted = " ".join((marker_text or "").split())
    for child in doc.element.body:
        if child.tag == qn("w:p") and " ".join(_paragraph_text_xml(child).split()) == wanted:
            return child
    return None


def _nearest_table_indexes_around_marker(doc: Document, marker_para) -> tuple[int, int] | None:
    body = list(doc.element.body)
    table_by_xml = _doc_table_xml_map(doc)
    try:
        marker_idx = body.index(marker_para)
    except ValueError:
        return None

    previous = None
    idx = marker_idx - 1
    while idx >= 0:
        child = body[idx]
        if child.tag == qn("w:tbl"):
            previous = child
            break
        if child.tag == qn("w:p") and _paragraph_text_xml(child):
            return None
        idx -= 1

    following = None
    idx = marker_idx + 1
    while idx < len(body):
        child = body[idx]
        if child.tag == qn("w:tbl"):
            following = child
            break
        if child.tag == qn("w:p") and _paragraph_text_xml(child):
            return None
        idx += 1

    if previous is None or following is None:
        return None
    if previous not in table_by_xml or following not in table_by_xml:
        return None
    return table_by_xml[previous], table_by_xml[following]


def _strict_marker_between_table_indexes(doc: Document, first_idx: int, second_idx: int, table_num: str):
    try:
        first_tbl = doc.tables[first_idx]._tbl
        second_tbl = doc.tables[second_idx]._tbl
    except IndexError:
        return None
    body = list(doc.element.body)
    try:
        first_pos = body.index(first_tbl)
        second_pos = body.index(second_tbl)
    except ValueError:
        return None
    if second_pos <= first_pos:
        return None
    expected = f"Продолжение таблицы {table_num}"
    marker = None
    for child in body[first_pos + 1:second_pos]:
        if child.tag != qn("w:p"):
            return None
        text = " ".join(_paragraph_text_xml(child).split())
        if not text:
            continue
        if text == expected and _STRICT_CONTINUATION_MARKER_RE.match(text):
            if marker is not None:
                return None
            marker = child
            continue
        return None
    return marker


def _caption_before_table_matches(doc: Document, table_idx: int, table_num: str) -> bool:
    body = list(doc.element.body)
    try:
        tbl_pos = body.index(doc.tables[table_idx]._tbl)
    except (IndexError, ValueError):
        return False
    for child in reversed(body[:tbl_pos]):
        if child.tag == qn("w:tbl"):
            return False
        if child.tag != qn("w:p"):
            continue
        text = _paragraph_text_xml(child)
        if not text:
            continue
        caption_num = _caption_table_num(text)
        if caption_num:
            return caption_num == table_num
    return False


def _source_table_indexes_for_num(doc: Document, table_num: str) -> list[int]:
    table_by_xml = _doc_table_xml_map(doc)
    current_num: str | None = None
    out: list[int] = []
    any_marker_re = re.compile(
        r"Продолжение\s+(?:таблицы|табл\.)\s+(\d+(?:\.\d+)*)",
        re.IGNORECASE,
    )
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            text = _paragraph_text_xml(child)
            caption_num = _caption_table_num(text)
            if caption_num:
                current_num = caption_num
                continue
            marker = any_marker_re.search(text)
            if marker:
                current_num = marker.group(1)
                continue
            if current_num == table_num and _source_or_note_text(text):
                current_num = None
                continue
        elif child.tag == qn("w:tbl") and current_num == table_num and child in table_by_xml:
            out.append(table_by_xml[child])
    return out


def _source_has_numeric_row_for_table(source_docx_path: Path | None, table_num: str) -> bool | None:
    if source_docx_path is None:
        return None
    try:
        doc = Document(str(source_docx_path))
    except Exception:
        return None
    table_indexes = _source_table_indexes_for_num(doc, table_num)
    if not table_indexes:
        return None
    for table_idx in table_indexes:
        for row in doc.tables[table_idx].rows:
            if _is_docx_numeric_row(_docx_row_cell_texts(row)):
                return True
    return False


def _source_has_meaningful_duplicate_for_table(source_docx_path: Path | None, table_num: str) -> bool | None:
    if source_docx_path is None:
        return None
    try:
        doc = Document(str(source_docx_path))
    except Exception:
        return None
    table_indexes = _source_table_indexes_for_num(doc, table_num)
    if not table_indexes:
        return None
    counts: dict[str, int] = {}
    for table_idx in table_indexes:
        for fp in _docx_data_fingerprints(doc.tables[table_idx]):
            counts[fp] = counts.get(fp, 0) + 1
    return any(count > 1 for count in counts.values())


def _tables_have_exact_same_layout(left, right) -> bool:
    if _docx_table_col_count(left) != _docx_table_col_count(right):
        return False
    if _table_grid_signature(left) != _table_grid_signature(right):
        return False
    if _table_width_signature(left) != _table_width_signature(right):
        return False
    if _table_margin_signature(left) != _table_margin_signature(right):
        return False
    if _table_border_signature(left) != _table_border_signature(right):
        return False
    return True


def _tables_have_compatible_same_page_layout(left, right) -> bool:
    if _docx_table_col_count(left) != _docx_table_col_count(right):
        return False
    if _table_width_signature(left) != _table_width_signature(right):
        return False
    if _table_margin_signature(left) != _table_margin_signature(right):
        return False
    if _table_border_signature(left) != _table_border_signature(right):
        return False
    return True


def _exact_grid_same_page_candidate(
    doc: Document,
    violation: _SamePageContinuationMarkerViolation,
    *,
    source_docx_path: Path | None,
) -> tuple[str, int, int, object] | None:
    table_num = _strict_marker_table_num(violation.marker_text)
    if not table_num or violation.confidence != "high":
        return None
    if violation.previous_table_page != violation.marker_page:
        return None
    if violation.following_table_page != violation.marker_page:
        return None

    marker_para = _find_strict_marker_paragraph(doc, violation.marker_text)
    if marker_para is None:
        return None
    table_pair = _nearest_table_indexes_around_marker(doc, marker_para)
    if table_pair is None:
        return None
    first_idx, second_idx = table_pair
    if not _caption_before_table_matches(doc, first_idx, table_num):
        return None

    first = doc.tables[first_idx]
    second = doc.tables[second_idx]
    if _table_has_merged_cells_docx(first) or _table_has_merged_cells_docx(second):
        return None
    if not _tables_have_exact_same_layout(first, second):
        return None
    if len(first.rows) < 2 or len(second.rows) < 2:
        return None
    if _docx_row_fingerprint(first.rows[0]) != _docx_row_fingerprint(second.rows[0]):
        return None

    has_numeric = any(_is_docx_numeric_row(_docx_row_cell_texts(row)) for row in first.rows)
    has_numeric = has_numeric or any(_is_docx_numeric_row(_docx_row_cell_texts(row)) for row in second.rows)
    if has_numeric and _source_has_numeric_row_for_table(source_docx_path, table_num) is not False:
        return None
    if _source_has_meaningful_duplicate_for_table(source_docx_path, table_num) is not False:
        return None

    first_data = set(_docx_data_fingerprints(first))
    second_data = set(_docx_data_fingerprints(second))
    if not second_data or first_data & second_data:
        return None
    return table_num, first_idx, second_idx, marker_para


def _exact_grid_same_page_candidate_from_rendered(
    doc: Document,
    violation,
    *,
    source_docx_path: Path | None,
) -> tuple[str, int, int, object] | None:
    table_num = getattr(violation, "table_num", None)
    violation_type = getattr(violation, "violation_type", None)
    confidence = getattr(violation, "confidence", None)
    evidence = getattr(violation, "evidence", {}) or {}
    if not table_num or violation_type != "same_page_repeated_fragment" or confidence != "high":
        return None
    try:
        first_idx = int(getattr(violation, "table_index"))
        second_idx = int(evidence.get("following_table_index"))
    except (TypeError, ValueError):
        return None
    if second_idx != first_idx + 1:
        return None
    marker_para = _strict_marker_between_table_indexes(doc, first_idx, second_idx, str(table_num))
    if not _caption_before_table_matches(doc, first_idx, str(table_num)):
        return None

    first = doc.tables[first_idx]
    second = doc.tables[second_idx]
    if _table_has_merged_cells_docx(first) or _table_has_merged_cells_docx(second):
        return None
    if not _tables_have_exact_same_layout(first, second):
        return None
    if len(first.rows) < 2 or len(second.rows) < 2:
        return None
    if _docx_row_fingerprint(first.rows[0]) != _docx_row_fingerprint(second.rows[0]):
        return None
    has_numeric = any(_is_docx_numeric_row(_docx_row_cell_texts(row)) for row in first.rows)
    has_numeric = has_numeric or any(_is_docx_numeric_row(_docx_row_cell_texts(row)) for row in second.rows)
    if has_numeric and _source_has_numeric_row_for_table(source_docx_path, str(table_num)) is not False:
        return None
    if _source_has_meaningful_duplicate_for_table(source_docx_path, str(table_num)) is not False:
        return None
    first_data = set(_docx_data_fingerprints(first))
    second_data = set(_docx_data_fingerprints(second))
    if not second_data or first_data & second_data:
        return None
    return str(table_num), first_idx, second_idx, marker_para


def _same_page_candidate_common_checks(
    doc: Document,
    table_num: str,
    first_idx: int,
    second_idx: int,
    *,
    source_docx_path: Path | None,
    require_exact_layout: bool,
    require_numeric_rows_in_both: bool = False,
) -> bool:
    if second_idx != first_idx + 1:
        return False
    if not _caption_before_table_matches(doc, first_idx, table_num):
        return False
    try:
        first = doc.tables[first_idx]
        second = doc.tables[second_idx]
    except IndexError:
        return False
    if _table_has_merged_cells_docx(first) or _table_has_merged_cells_docx(second):
        return False
    if require_exact_layout:
        if not _tables_have_exact_same_layout(first, second):
            return False
    elif not _tables_have_compatible_same_page_layout(first, second):
        return False
    if len(first.rows) < 2 or len(second.rows) < 2:
        return False
    if _docx_row_fingerprint(first.rows[0]) != _docx_row_fingerprint(second.rows[0]):
        return False
    first_has_numeric = _docx_table_has_numeric_row(first)
    second_has_numeric = _docx_table_has_numeric_row(second)
    if require_numeric_rows_in_both and not (first_has_numeric and second_has_numeric):
        return False
    has_numeric = first_has_numeric or second_has_numeric
    if has_numeric and _source_has_numeric_row_for_table(source_docx_path, table_num) is not False:
        return False
    if _source_has_meaningful_duplicate_for_table(source_docx_path, table_num) is not False:
        return False
    first_data = set(_docx_data_fingerprints(first))
    second_data = set(_docx_data_fingerprints(second))
    if not second_data or first_data & second_data:
        return False
    return True


def _compatible_grid_same_page_candidate_from_rendered(
    doc: Document,
    violation,
    *,
    source_docx_path: Path | None,
) -> tuple[str, int, int, object] | None:
    table_num = getattr(violation, "table_num", None)
    violation_type = getattr(violation, "violation_type", None)
    confidence = getattr(violation, "confidence", None)
    evidence = getattr(violation, "evidence", {}) or {}
    if not table_num or violation_type != "same_page_repeated_fragment" or confidence != "high":
        return None
    try:
        first_idx = int(getattr(violation, "table_index"))
        second_idx = int(evidence.get("following_table_index"))
    except (TypeError, ValueError):
        return None
    marker_para = _strict_marker_between_table_indexes(doc, first_idx, second_idx, str(table_num))
    if not _same_page_candidate_common_checks(
        doc,
        str(table_num),
        first_idx,
        second_idx,
        source_docx_path=source_docx_path,
        require_exact_layout=False,
        require_numeric_rows_in_both=True,
    ):
        return None
    return str(table_num), first_idx, second_idx, marker_para


def _no_numeric_same_page_header_cleanup_candidate_from_rendered(
    doc: Document,
    violation,
    *,
    source_docx_path: Path | None,
) -> tuple[str, int, int] | None:
    table_num = getattr(violation, "table_num", None)
    violation_type = getattr(violation, "violation_type", None)
    confidence = getattr(violation, "confidence", None)
    evidence = getattr(violation, "evidence", {}) or {}
    if not table_num or violation_type != "same_page_repeated_fragment" or confidence != "high":
        return None
    if evidence.get("repeated_numeric_row_count") != 0:
        return None
    try:
        first_idx = int(getattr(violation, "table_index"))
        second_idx = int(evidence.get("following_table_index"))
    except (TypeError, ValueError):
        return None
    if _strict_marker_between_table_indexes(doc, first_idx, second_idx, str(table_num)) is not None:
        return None
    if _source_has_numeric_row_for_table(source_docx_path, str(table_num)) is True:
        return None
    if not _same_page_candidate_common_checks(
        doc,
        str(table_num),
        first_idx,
        second_idx,
        source_docx_path=source_docx_path,
        require_exact_layout=False,
        require_numeric_rows_in_both=False,
    ):
        return None
    first = doc.tables[first_idx]
    second = doc.tables[second_idx]
    if _docx_table_has_numeric_row(first) or _docx_table_has_numeric_row(second):
        return None
    return str(table_num), first_idx, second_idx


def _row_cell_widths(row) -> list[tuple[str | None, str | None]]:
    widths: list[tuple[str | None, str | None]] = []
    for cell in row.cells:
        tc_pr = cell._tc.tcPr
        tc_w = tc_pr.find(qn("w:tcW")) if tc_pr is not None else None
        widths.append(
            (
                tc_w.get(qn("w:w")) if tc_w is not None else None,
                tc_w.get(qn("w:type")) if tc_w is not None else None,
            )
        )
    return widths


def _apply_row_cell_widths(row, widths: list[tuple[str | None, str | None]]) -> None:
    for idx, cell in enumerate(row.cells):
        if idx >= len(widths):
            break
        width, width_type = widths[idx]
        if width is None and width_type is None:
            continue
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_w = tc_pr.find(qn("w:tcW"))
        if tc_w is None:
            tc_w = OxmlElement("w:tcW")
            tc_pr.insert(0, tc_w)
        if width is not None:
            tc_w.set(qn("w:w"), width)
        if width_type is not None:
            tc_w.set(qn("w:type"), width_type)


def _append_second_fragment_data_rows(first, second, *, normalize_to_first_grid: bool = False) -> int:
    appended = 0
    survivor_widths = _row_cell_widths(first.rows[0]) if normalize_to_first_grid and first.rows else []
    for idx, row in enumerate(list(second.rows)):
        values = _docx_row_cell_texts(row)
        if idx == 0:
            continue
        if _is_docx_numeric_row(values):
            continue
        first._tbl.append(deepcopy(row._tr))
        if normalize_to_first_grid and survivor_widths:
            _apply_row_cell_widths(first.rows[-1], survivor_widths)
        appended += 1
    return appended


def _remove_duplicate_second_fragment_header(second) -> bool:
    if not second.rows:
        return False
    _remove_xml_node(second.rows[0]._tr)
    return True


def _clear_same_page_merge_repeat_metadata(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.trPr
        if tr_pr is not None:
            for tbl_header in list(tr_pr.findall(qn("w:tblHeader"))):
                tr_pr.remove(tbl_header)
        for paragraph in row._tr.findall(".//" + qn("w:p")):
            p_pr = paragraph.find(qn("w:pPr"))
            if p_pr is None:
                continue
            for tag in ("w:pageBreakBefore", "w:keepNext"):
                for node in list(p_pr.findall(qn(tag))):
                    p_pr.remove(node)


def _remove_xml_node(node) -> None:
    if node is None:
        return
    parent = node.getparent()
    if parent is not None:
        parent.remove(node)


def _same_page_marker_text_remains(docx_path: Path, marker_text: str) -> bool:
    wanted = " ".join((marker_text or "").split())
    for violation in _same_page_continuation_marker_violations_for_docx(docx_path):
        if " ".join(violation.marker_text.split()) == wanted:
            return True
    return False


def _same_page_rendered_target_remains(docx_path: Path, table_num: str) -> bool:
    for violation in _rendered_continuation_violations_for_docx(docx_path):
        if (
            getattr(violation, "table_num", None) == table_num
            and getattr(violation, "violation_type", None) in {
                "same_page_repeated_fragment",
                "same_page_adjacent_fragment",
            }
        ):
            return True
    return False


def normalize_exact_grid_same_page_repeated_fragments_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """
    Merge only exact-grid same-page repeated fragments.

    This is intentionally narrower than the future continuation engine: it
    requires rendered same-page marker evidence, exact DOCX grid/layout match,
    repeated semantic header, source-proven generated numeric rows, and no
    source-bad data duplicates. Grid-mismatch cases remain warnings.
    """
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    try:
        marker_violations = _same_page_continuation_marker_violations_for_docx(docx_path)
    except Exception as exc:
        logger.info(
            "same_page_exact_grid_normalize_marker_probe_skip path=%s reason=render_failed error=%s",
            docx_path, exc,
        )
        marker_violations = []
    try:
        rendered_violations = _rendered_continuation_violations_for_docx(docx_path)
    except Exception as exc:
        logger.info(
            "same_page_exact_grid_normalize_render_probe_skip path=%s reason=render_failed error=%s",
            docx_path, exc,
        )
        rendered_violations = []
    if not marker_violations and not rendered_violations:
        return 0

    repaired = 0
    candidates: list[tuple[str, object]] = [
        ("rendered", violation)
        for violation in rendered_violations
        if getattr(violation, "violation_type", None) == "same_page_repeated_fragment"
    ]
    candidates.extend(("marker", violation) for violation in marker_violations)
    seen_tables: set[str] = set()
    for source_kind, violation in candidates:
        backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_same_page_exact_grid_"))
        backup_path = backup_dir / docx_path.name
        try:
            shutil.copy2(docx_path, backup_path)
            doc = Document(str(docx_path))
            if source_kind == "rendered":
                candidate = _exact_grid_same_page_candidate_from_rendered(
                    doc,
                    violation,
                    source_docx_path=source_docx_path,
                )
            else:
                candidate = _exact_grid_same_page_candidate(
                    doc,
                    violation,
                    source_docx_path=source_docx_path,
                )
            if candidate is None:
                continue
            table_num, first_idx, second_idx, marker_para = candidate
            if table_num in seen_tables:
                continue
            first = doc.tables[first_idx]
            second = doc.tables[second_idx]
            appended = _append_second_fragment_data_rows(first, second)
            if appended <= 0:
                continue
            _clear_same_page_merge_repeat_metadata(first)
            _remove_xml_node(marker_para)
            _remove_xml_node(second._tbl)
            doc.save(str(docx_path))

            marker_text = f"Продолжение таблицы {table_num}"
            if _same_page_marker_text_remains(docx_path, marker_text):
                shutil.copy2(backup_path, docx_path)
                logger.info(
                    "same_page_exact_grid_normalize_rollback table_num=%s reason=marker_remains",
                    table_num,
                )
                continue
            if _same_page_rendered_target_remains(docx_path, table_num):
                shutil.copy2(backup_path, docx_path)
                logger.info(
                    "same_page_exact_grid_normalize_rollback table_num=%s reason=same_page_rendered_target_remains",
                    table_num,
                )
                continue
            regressions = _rendered_continuation_deletion_regressions(docx_path)
            if regressions:
                shutil.copy2(backup_path, docx_path)
                logger.info(
                    "same_page_exact_grid_normalize_rollback table_num=%s reason=rendered_regression violations=%s",
                    table_num,
                    _format_rendered_deletion_regressions(regressions),
                )
                continue

            repaired += 1
            seen_tables.add(table_num)
            logger.info(
                "same_page_exact_grid_normalize_applied table_num=%s first_table=%s second_table=%s appended_rows=%s",
                table_num, first_idx, second_idx, appended,
            )
        except Exception as exc:
            shutil.copy2(backup_path, docx_path)
            logger.info(
                "same_page_exact_grid_normalize_rollback marker=%r reason=exception error=%s",
                getattr(violation, "marker_text", ""), exc,
            )
        finally:
            shutil.rmtree(backup_dir, ignore_errors=True)
    return repaired


def normalize_compatible_grid_same_page_repeated_fragments_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """
    Merge compatible-grid same-page repeated fragments.

    This deliberately handles only the next bounded class after exact-grid:
    rendered high-confidence same-page fragments with equal column count,
    compatible width/margins/borders, no merged cells, repeated semantic header,
    formatter-generated numeric rows, and distinct source data rows.  The first
    table's grid survives; appended rows are adapted to that layout.  A narrower
    no-numeric variant keeps both physical fragments and removes only the
    duplicate header from the second fragment when merging would be unsafe.
    """
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    repaired = 0
    seen_tables: set[str] = set()
    for _pass in range(20):
        try:
            rendered_violations = _rendered_continuation_violations_for_docx(docx_path)
        except Exception as exc:
            logger.info(
                "same_page_compatible_grid_normalize_render_probe_skip path=%s reason=render_failed error=%s",
                docx_path, exc,
            )
            break

        candidates = [
            violation
            for violation in rendered_violations
            if getattr(violation, "violation_type", None) == "same_page_repeated_fragment"
        ]
        if not candidates:
            break

        made_progress = False
        for violation in candidates:
            backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_same_page_compatible_grid_"))
            backup_path = backup_dir / docx_path.name
            try:
                shutil.copy2(docx_path, backup_path)
                doc = Document(str(docx_path))
                candidate = _compatible_grid_same_page_candidate_from_rendered(
                    doc,
                    violation,
                    source_docx_path=source_docx_path,
                )
                if candidate is None:
                    header_cleanup_candidate = _no_numeric_same_page_header_cleanup_candidate_from_rendered(
                        doc,
                        violation,
                        source_docx_path=source_docx_path,
                    )
                    if header_cleanup_candidate is None:
                        continue
                    table_num, first_idx, second_idx = header_cleanup_candidate
                    if table_num in seen_tables:
                        continue
                    first = doc.tables[first_idx]
                    second = doc.tables[second_idx]
                    if not _remove_duplicate_second_fragment_header(second):
                        continue
                    _clear_same_page_merge_repeat_metadata(first)
                    _clear_same_page_merge_repeat_metadata(second)
                    doc.save(str(docx_path))

                    if _same_page_rendered_target_remains(docx_path, table_num):
                        shutil.copy2(backup_path, docx_path)
                        logger.info(
                            "same_page_no_numeric_header_cleanup_rollback table_num=%s reason=same_page_rendered_target_remains",
                            table_num,
                        )
                        continue
                    if _same_table_start_orphan_remains(docx_path, first_idx):
                        shutil.copy2(backup_path, docx_path)
                        logger.info(
                            "same_page_no_numeric_header_cleanup_rollback table_num=%s reason=table_start_orphan",
                            table_num,
                        )
                        continue
                    regressions = _rendered_continuation_deletion_regressions(docx_path)
                    if regressions:
                        shutil.copy2(backup_path, docx_path)
                        logger.info(
                            "same_page_no_numeric_header_cleanup_rollback table_num=%s reason=rendered_regression violations=%s",
                            table_num,
                            _format_rendered_deletion_regressions(regressions),
                        )
                        continue

                    repaired += 1
                    seen_tables.add(table_num)
                    made_progress = True
                    logger.info(
                        "same_page_no_numeric_header_cleanup_applied table_num=%s first_table=%s second_table=%s",
                        table_num, first_idx, second_idx,
                    )
                    break
                else:
                    table_num, first_idx, second_idx, marker_para = candidate
                    if table_num in seen_tables:
                        continue
                    first = doc.tables[first_idx]
                    second = doc.tables[second_idx]
                    appended = _append_second_fragment_data_rows(
                        first,
                        second,
                        normalize_to_first_grid=True,
                    )
                    if appended <= 0:
                        continue
                    _clear_same_page_merge_repeat_metadata(first)
                    _remove_xml_node(marker_para)
                    _remove_xml_node(second._tbl)
                    doc.save(str(docx_path))

                    marker_text = f"Продолжение таблицы {table_num}"
                    if _same_page_marker_text_remains(docx_path, marker_text):
                        shutil.copy2(backup_path, docx_path)
                        logger.info(
                            "same_page_compatible_grid_normalize_rollback table_num=%s reason=marker_remains",
                            table_num,
                        )
                        continue
                    if _same_page_rendered_target_remains(docx_path, table_num):
                        shutil.copy2(backup_path, docx_path)
                        logger.info(
                            "same_page_compatible_grid_normalize_rollback table_num=%s reason=same_page_rendered_target_remains",
                            table_num,
                        )
                        continue
                    regressions = _rendered_continuation_deletion_regressions(docx_path)
                    if regressions:
                        shutil.copy2(backup_path, docx_path)
                        logger.info(
                            "same_page_compatible_grid_normalize_rollback table_num=%s reason=rendered_regression violations=%s",
                            table_num,
                            _format_rendered_deletion_regressions(regressions),
                        )
                        continue

                    repaired += 1
                    seen_tables.add(table_num)
                    made_progress = True
                    logger.info(
                        "same_page_compatible_grid_normalize_applied table_num=%s first_table=%s second_table=%s appended_rows=%s",
                        table_num, first_idx, second_idx, appended,
                    )
                    break
            except Exception as exc:
                shutil.copy2(backup_path, docx_path)
                logger.info(
                    "same_page_compatible_grid_normalize_rollback table_num=%s reason=exception error=%s",
                    getattr(violation, "table_num", ""), exc,
                )
            finally:
                shutil.rmtree(backup_dir, ignore_errors=True)

        if not made_progress:
            break
    return repaired


def _same_page_meaningful_row_fps(first, second, header_fp: str) -> list[str]:
    """Sorted multiset of meaningful (non-header, non-numeric) row fingerprints
    across both fragments. Invariant to header/numeric-row removal, so it can
    verify that a cleanup preserved every real data row exactly once."""
    out: list[str] = []
    for table in (first, second):
        for row in table.rows:
            values = _docx_row_cell_texts(row)
            if _is_docx_numeric_row(values):
                continue
            fp = _docx_row_fingerprint(row)
            if not fp or fp == header_fp:
                continue
            out.append(fp)
    return sorted(out)


def _remove_second_fragment_duplicate_leading_rows(first, second) -> int:
    """Remove the second fragment's LEADING duplicate SEMANTIC HEADER row(s) so it
    starts with the numeric column row (canonical KFU continuation: numeric row,
    not semantic header). The numeric row is KEPT. Stops at the first numeric or
    data row so meaningful data is never removed, and never empties the table."""
    if not first.rows or not second.rows:
        return 0
    first_header_fp = _docx_row_fingerprint(first.rows[0])
    removed = 0
    while len(second.rows) > 1:
        row = second.rows[0]
        values = _docx_row_cell_texts(row)
        if _is_docx_numeric_row(values):
            break  # numeric row is the canonical continuation lead — keep it
        if _docx_row_fingerprint(row) != first_header_fp:
            break  # reached a data row
        _remove_xml_node(row._tr)
        removed += 1
    return removed


def _incompatible_grid_same_page_cleanup_candidate_from_rendered(
    doc: Document,
    violation,
    *,
    source_docx_path: Path | None,
) -> tuple[str, int, int, object] | None:
    """Same-page repeated fragment that the exact/compatible mergers refused
    (typically because the two fragments have incompatible grids). Safe to clean
    when the second fragment merely REPEATS the header and carries distinct data:
    we drop the duplicate header/numeric + the same-page marker and keep BOTH
    physical tables (no merge, no grid reshape). Source-proven duplicates are
    left untouched (classified manual/source-bad elsewhere)."""
    table_num = getattr(violation, "table_num", None)
    if (
        not table_num
        or getattr(violation, "violation_type", None) != "same_page_repeated_fragment"
        or getattr(violation, "confidence", None) != "high"
    ):
        return None
    evidence = getattr(violation, "evidence", {}) or {}
    try:
        first_idx = int(getattr(violation, "table_index"))
        second_idx = int(evidence.get("following_table_index"))
    except (TypeError, ValueError):
        return None
    if second_idx != first_idx + 1:
        return None
    if not _caption_before_table_matches(doc, first_idx, str(table_num)):
        return None
    try:
        first = doc.tables[first_idx]
        second = doc.tables[second_idx]
    except IndexError:
        return None
    if _table_has_merged_cells_docx(first) or _table_has_merged_cells_docx(second):
        return None
    if len(first.rows) < 2 or len(second.rows) < 2:
        return None
    # second fragment must START with a proven duplicate of the first header
    if _docx_row_fingerprint(second.rows[0]) != _docx_row_fingerprint(first.rows[0]):
        return None
    # never touch source-proven duplicated content (manual / source-bad)
    if _source_has_meaningful_duplicate_for_table(source_docx_path, str(table_num)) is not False:
        return None
    # data rows must be distinct and present in the second fragment
    first_data = set(_docx_data_fingerprints(first))
    second_data = set(_docx_data_fingerprints(second))
    if not second_data or (first_data & second_data):
        return None
    marker_para = _strict_marker_between_table_indexes(doc, first_idx, second_idx, str(table_num))
    return str(table_num), first_idx, second_idx, marker_para


def cleanup_same_page_incompatible_chains_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """Fallback after exact/compatible merge: for grid-incompatible same-page
    repeated fragments, remove the same-page continuation marker and the second
    fragment's duplicate header/numeric rows, keeping both physical tables. Never
    reshapes grids, never alters meaningful data rows. Rolls back if a same-page
    target/marker remains, a data row is lost/duplicated/reordered, a table-start
    orphan appears, or any rendered deletion regression is detected."""
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    repaired = 0
    seen_tables: set[str] = set()
    for _pass in range(20):
        try:
            rendered_violations = _rendered_continuation_violations_for_docx(docx_path)
        except Exception as exc:
            logger.info(
                "same_page_incompatible_cleanup_render_probe_skip path=%s reason=render_failed error=%s",
                docx_path, exc,
            )
            break
        candidates = [
            v for v in rendered_violations
            if getattr(v, "violation_type", None) == "same_page_repeated_fragment"
        ]
        if not candidates:
            break
        made_progress = False
        for violation in candidates:
            backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_same_page_incompatible_"))
            backup_path = backup_dir / docx_path.name
            try:
                shutil.copy2(docx_path, backup_path)
                doc = Document(str(docx_path))
                candidate = _incompatible_grid_same_page_cleanup_candidate_from_rendered(
                    doc, violation, source_docx_path=source_docx_path,
                )
                if candidate is None:
                    continue
                table_num, first_idx, second_idx, marker_para = candidate
                if table_num in seen_tables:
                    continue
                first = doc.tables[first_idx]
                second = doc.tables[second_idx]
                header_fp = _docx_row_fingerprint(first.rows[0])
                before_data = _same_page_meaningful_row_fps(first, second, header_fp)
                removed = _remove_second_fragment_duplicate_leading_rows(first, second)
                if removed <= 0:
                    continue
                after_data = _same_page_meaningful_row_fps(first, second, header_fp)
                if after_data != before_data:
                    # a meaningful data row would change — refuse
                    continue
                if marker_para is not None:
                    _remove_xml_node(marker_para)
                _clear_same_page_merge_repeat_metadata(first)
                _clear_same_page_merge_repeat_metadata(second)
                doc.save(str(docx_path))

                marker_text = f"Продолжение таблицы {table_num}"
                if marker_para is not None and _same_page_marker_text_remains(docx_path, marker_text):
                    shutil.copy2(backup_path, docx_path)
                    logger.info("same_page_incompatible_cleanup_rollback table_num=%s reason=marker_remains", table_num)
                    continue
                if _same_page_rendered_target_remains(docx_path, table_num):
                    shutil.copy2(backup_path, docx_path)
                    logger.info("same_page_incompatible_cleanup_rollback table_num=%s reason=same_page_rendered_target_remains", table_num)
                    continue
                if _same_table_start_orphan_remains(docx_path, first_idx):
                    shutil.copy2(backup_path, docx_path)
                    logger.info("same_page_incompatible_cleanup_rollback table_num=%s reason=table_start_orphan", table_num)
                    continue
                regressions = _rendered_continuation_deletion_regressions(docx_path)
                if regressions:
                    shutil.copy2(backup_path, docx_path)
                    logger.info(
                        "same_page_incompatible_cleanup_rollback table_num=%s reason=rendered_regression violations=%s",
                        table_num, _format_rendered_deletion_regressions(regressions),
                    )
                    continue

                repaired += 1
                seen_tables.add(table_num)
                made_progress = True
                logger.info(
                    "same_page_incompatible_cleanup_applied table_num=%s first_table=%s second_table=%s removed_rows=%s marker=%s",
                    table_num, first_idx, second_idx, removed, marker_para is not None,
                )
                break
            except Exception as exc:
                shutil.copy2(backup_path, docx_path)
                logger.info(
                    "same_page_incompatible_cleanup_rollback table_num=%s reason=exception error=%s",
                    getattr(violation, "table_num", ""), exc,
                )
            finally:
                shutil.rmtree(backup_dir, ignore_errors=True)
        if not made_progress:
            break
    return repaired


def _same_page_continuation_fail_keys(docx_path: Path, source_docx_path: Path | None):
    """Render once and return the acceptance-gate FAIL blocker keys plus the set
    of table numbers flagged `same_page_continuation`. Driven by the acceptance
    gate (which reliably attributes table+page), not `same_page_repeated_fragment`."""
    from .rendered_table_validation import (
        build_rendered_table_identities,
        evaluate_table_layout_acceptance,
    )

    pdf_path: Path | None = None
    try:
        pdf_path = render_docx_to_pdf(Path(docx_path))
        pdf_lines = analyze_pdf_lines(pdf_path)
        doc = Document(str(docx_path))
        identities = build_rendered_table_identities(doc)
        source_ids = None
        if source_docx_path is not None:
            try:
                source_ids = build_rendered_table_identities(Document(str(source_docx_path)))
            except Exception:
                source_ids = None
        blockers = evaluate_table_layout_acceptance(
            pdf_lines, identities, doc=doc, source_identities=source_ids
        )
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)

    fail_keys = {(b.blocker_type, b.table_num) for b in blockers if b.severity == "fail"}
    same_page = {
        b.table_num for b in blockers
        if b.severity == "fail" and b.blocker_type == "same_page_continuation" and b.table_num
    }
    return fail_keys, same_page


def _content_regressed(source_docx_path: Path | None, docx_path: Path) -> bool:
    if source_docx_path is None:
        return False
    try:
        from .content_preservation import evaluate_content_preservation
        _rep, issues = evaluate_content_preservation(
            Document(str(source_docx_path)), Document(str(docx_path))
        )
        return any(i.severity == "fail" for i in issues)
    except Exception:
        return False  # never block cleanup on a measurement error


def _force_marker_page_break(marker_para) -> None:
    """Force the marker paragraph onto a new page: pageBreakBefore + keepNext set
    ACTIVE (val='true'), overriding any earlier val='0' that an upstream pass left
    (that disabled flag is exactly why these markers still render same-page). No
    rows/grid touched — the repeated header becomes a valid continuation header."""
    pPr = marker_para.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        marker_para.insert(0, pPr)
    for tag in ("w:pageBreakBefore", "w:keepNext"):
        el = pPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            pPr.append(el)
        el.set(qn("w:val"), "true")


def _same_page_continuation_cleanup_candidate(
    doc: Document,
    table_num: str,
    *,
    source_docx_path: Path | None,
) -> tuple[int, int, object] | None:
    """Locate the manual continuation chain for a `same_page_continuation` blocker
    via its strict marker, and validate it is safe to clean: caption matches, the
    following fragment STARTS with a proven duplicate header, data rows are
    distinct, no merged cells, and the table is not source-bad duplicated."""
    marker_text = f"Продолжение таблицы {table_num}"
    marker_para = _find_strict_marker_paragraph(doc, marker_text)
    if marker_para is None:
        return None
    idxs = _nearest_table_indexes_around_marker(doc, marker_para)
    if idxs is None:
        return None
    first_idx, second_idx = idxs
    if not _caption_before_table_matches(doc, first_idx, str(table_num)):
        return None
    try:
        first = doc.tables[first_idx]
        second = doc.tables[second_idx]
    except IndexError:
        return None
    if _table_has_merged_cells_docx(first) or _table_has_merged_cells_docx(second):
        return None
    if len(first.rows) < 2 or len(second.rows) < 2:
        return None
    if _docx_row_fingerprint(second.rows[0]) != _docx_row_fingerprint(first.rows[0]):
        return None
    if _source_has_meaningful_duplicate_for_table(source_docx_path, str(table_num)) is not False:
        return None
    first_data = set(_docx_data_fingerprints(first))
    second_data = set(_docx_data_fingerprints(second))
    if not second_data or (first_data & second_data):
        return None
    return first_idx, second_idx, marker_para


def cleanup_same_page_continuation_blockers_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """Deterministically clean same-page manual continuation chains that the
    acceptance gate flags as `same_page_continuation`.

    Strategy per chain (after locating it from the gate blocker + its marker):
      1. compatible grid -> merge fragments (drop marker + duplicate header/
         numeric, append the second fragment's data rows to the first);
      2. incompatible grid (same shape but differing widths) -> drop the marker +
         the second fragment's duplicate header/numeric, keep both tables;
      3. otherwise -> no mutation.

    Every applied cleanup is verified by re-rendering: it must remove that
    `same_page_continuation` fail, introduce NO new fail blocker, preserve all
    content (content-preservation gate), and not orphan the table start —
    otherwise it is rolled back."""
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    repaired = 0
    seen: set[str] = set()
    for _pass in range(10):
        try:
            baseline_keys, same_page = _same_page_continuation_fail_keys(docx_path, source_docx_path)
        except Exception as exc:
            logger.info(
                "same_page_continuation_cleanup_probe_skip path=%s reason=render_failed error=%s",
                docx_path, exc,
            )
            break
        todo = sorted(n for n in same_page if n not in seen)
        if not todo:
            break
        made_progress = False
        for table_num in todo:
            backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_sp_continuation_cleanup_"))
            backup_path = backup_dir / docx_path.name
            try:
                shutil.copy2(docx_path, backup_path)
                probe = Document(str(docx_path))
                candidate = _same_page_continuation_cleanup_candidate(
                    probe, table_num, source_docx_path=source_docx_path
                )
                if candidate is None:
                    seen.add(table_num)
                    continue
                pf, ps, _ = candidate
                compatible = _tables_have_compatible_same_page_layout(
                    probe.tables[pf], probe.tables[ps]
                )
                # Strategy order (reference fixture): prefer the gentle "merge"
                # for compatible grids; else "keep_both" (drop marker + duplicate
                # header, keep both tables); else the "page_break" fallback — force
                # the EXISTING marker to a new page (pageBreakBefore+keepNext) so
                # the marker + continuation fragment move together to the next
                # page, leaving the repeated header valid. The fallback changes no
                # grid and deletes no rows.
                strategies = (["merge"] if compatible else []) + ["keep_both", "page_break"]

                applied = False
                for strategy in strategies:
                    shutil.copy2(backup_path, docx_path)  # always start from a clean copy
                    doc = Document(str(docx_path))
                    cand = _same_page_continuation_cleanup_candidate(
                        doc, table_num, source_docx_path=source_docx_path
                    )
                    if cand is None:
                        continue
                    first_idx, second_idx, marker_para = cand
                    first = doc.tables[first_idx]
                    second = doc.tables[second_idx]
                    if strategy == "merge":
                        appended = _append_second_fragment_data_rows(
                            first, second, normalize_to_first_grid=True
                        )
                        if appended <= 0:
                            continue
                        _clear_same_page_merge_repeat_metadata(first)
                        _remove_xml_node(marker_para)
                        _remove_xml_node(second._tbl)
                    elif strategy == "keep_both":
                        removed = _remove_second_fragment_duplicate_leading_rows(first, second)
                        if removed <= 0:
                            continue
                        _remove_xml_node(marker_para)
                        _clear_same_page_merge_repeat_metadata(first)
                        _clear_same_page_merge_repeat_metadata(second)
                    else:  # page_break: keep marker + both fragments + repeated header
                        _force_marker_page_break(marker_para)
                    doc.save(str(docx_path))

                    # Verification with a single gate render (the acceptance gate
                    # already detects orphaned_header_row, neighbour same-page, and
                    # cross-page-without-marker, so one render covers marker-remains,
                    # orphan, neighbour-flip and new cross-page) + a docx-only
                    # content check. Keeps the cleanup render budget bounded.
                    reason = None
                    if _content_regressed(source_docx_path, docx_path):
                        reason = "content_regression"
                    else:
                        try:
                            after_keys, _asp = _same_page_continuation_fail_keys(docx_path, source_docx_path)
                        except Exception:
                            reason = "post_render_failed"
                        else:
                            if ("same_page_continuation", table_num) in after_keys:
                                reason = "same_page_continuation_remains"
                            elif after_keys - baseline_keys:
                                reason = f"new_fail_blocker={sorted(after_keys - baseline_keys)}"

                    if reason is None:
                        repaired += 1
                        made_progress = True
                        applied = True
                        logger.info(
                            "same_page_continuation_cleanup_applied table_num=%s strategy=%s first=%s second=%s",
                            table_num, strategy, first_idx, second_idx,
                        )
                        break
                    logger.info(
                        "same_page_continuation_cleanup_rollback table_num=%s strategy=%s reason=%s",
                        table_num, strategy, reason,
                    )

                if not applied:
                    shutil.copy2(backup_path, docx_path)
                seen.add(table_num)
                if applied:
                    break
            except Exception as exc:
                shutil.copy2(backup_path, docx_path)
                seen.add(table_num)
                logger.info(
                    "same_page_continuation_cleanup_rollback table_num=%s reason=exception error=%s",
                    table_num, exc,
                )
            finally:
                shutil.rmtree(backup_dir, ignore_errors=True)
        if not made_progress:
            break
    if repaired and report is not None:
        logger.info("same_page_continuation_cleanup total_repaired=%d", repaired)
    return repaired


# ── Cross-page (marker-less) long-table split ────────────────────────────────
#
# A DIFFERENT defect class from same-page manual markers: ONE physical table
# whose data rows render across pages with NO ``Продолжение таблицы N`` marker
# (acceptance blocker ``single_table_crosses_pages_without_marker``, e.g. Demo
# 1.1.3). The KFU-valid repair (kfu_long_table_split_reference): split the
# physical table at the rendered page boundary, insert the marker ONLY on the
# continuation page, repeat the header (and the numeric column-index row when
# present), and keep Источник:/Примечание: after the FINAL fragment.
#
# This is acceptance-gate driven + re-render verified + rolled back on any new
# defect. It is NOT the old globally-gated rendered-continuation path: scope is
# limited to two-page tables (single boundary) whose rows map unambiguously to
# pages, so a single split makes every fragment single-page (a 3+ page table is
# skipped honestly rather than leaving a caption-less continuation fragment that
# still crosses pages — which the caption-based gate cannot see).


def _match_data_row_pages_by_lead(
    table,
    pdf_lines: list[PdfLine],
    skip_idxs: set[int],
) -> dict[int, int] | None:
    """Map each DATA row (table-row index, excluding the header row 0 and any
    ``skip_idxs`` such as a numeric column-index row) to its rendered page using
    the leading first-cell text. Tolerant of rows that wrap across several PDF
    lines (the strict whole-row matcher fails on those). Sequential + monotonic;
    all-or-nothing — returns None if any data row cannot be located."""
    lines = [(_norm_match_text(line.text), line.page_num) for line in pdf_lines]
    if not lines:
        return None
    result: dict[int, int] = {}
    last_idx = -1
    last_page = -1
    for ri, row in enumerate(table.rows):
        if ri == 0 or ri in skip_idxs:
            continue
        cells = _docx_row_cell_texts(row)
        first = _norm_match_text(cells[0]) if cells else ""
        if len(first) < 3:
            return None  # leading cell too short to anchor reliably
        matched_page: int | None = None
        matched_idx = -1
        for j in range(last_idx + 1, len(lines)):
            text, page = lines[j]
            if last_page >= 0 and page < last_page:
                continue
            if text == first or text.startswith(first + " "):
                matched_page, matched_idx = page, j
                break
        if matched_page is None:
            return None
        result[ri] = matched_page
        last_idx = matched_idx
        last_page = matched_page
    return result or None


def _table_pdf_window(pdf_lines: list[PdfLine], table_num: str) -> list[PdfLine]:
    """Restrict ``pdf_lines`` to the rendered region of table ``table_num``: from
    its ``Таблица N`` caption to the line before the next table caption /
    continuation marker / ``Источник:``/``Примечание:``. This keeps row→page
    matching TABLE-LOCAL so a row's first-cell text that also appears in body
    prose / citations elsewhere cannot mis-anchor the mapping (Stage 4C). Falls
    back to the full list when the caption is not found."""
    cap_re = re.compile(rf"^\s*Таблица\s+{re.escape(table_num)}\b", re.IGNORECASE)
    any_cap_re = re.compile(r"^\s*Таблица\s+\d", re.IGNORECASE)
    stop_re = re.compile(r"^\s*(Источник|Примечание)\s*[:.]|^\s*Продолжение\s+табл", re.IGNORECASE)
    start = None
    for i, line in enumerate(pdf_lines):
        if cap_re.match(" ".join((line.text or "").split())):
            start = i
            break
    if start is None:
        return pdf_lines
    end = len(pdf_lines)
    for j in range(start + 1, len(pdf_lines)):
        t = " ".join((pdf_lines[j].text or "").split())
        if stop_re.match(t):
            end = j + 1  # keep the stop line's page so the last data row is covered
            break
        if any_cap_re.match(t) and not cap_re.match(t):
            end = j
            break
    return pdf_lines[start:end]


def _cross_page_instrumentation_enabled() -> bool:
    """Per-row marker instrumentation reliably maps tables whose text the strict /
    lead matchers can't anchor, but each call renders the WHOLE document (1–2×)
    plus a full PDF parse — ~10–30s on a long doc. That cost is unacceptable on
    the live bot (one format per request), so it is OFF by default and enabled
    only for eval/smoke via KPFU_CROSS_PAGE_INSTRUMENT=1. With it off the
    cross-page split still runs the cheap text-matcher path."""
    return os.environ.get("KPFU_CROSS_PAGE_INSTRUMENT", "").strip().lower() in {"1", "true", "yes", "on"}


def _instrumented_data_row_pages(docx_path: Path | None, table_idx: int) -> dict[int, int] | None:
    """Reliable DATA-row → page map via unique per-row marker instrumentation
    (``table_markers.map_table_rows_to_pages``) — structural, not text-matching,
    so it works on tables whose cell text overlaps body prose/citations. The
    header row (idx 0) marker is unreliable across the 1pt/2pt re-render, so it is
    ignored; every DATA row (idx>0, non-numeric) must be found exactly once and
    the pages must be monotonic. Returns {table_row_idx: page} or None."""
    if docx_path is None:
        return None
    try:
        from .table_markers import map_table_rows_to_pages
        res = map_table_rows_to_pages(Path(docx_path), table_idx)
        table = Document(str(docx_path)).tables[table_idx]
    except Exception:
        return None
    dup = set(res.duplicate_rows or {})
    out: dict[int, int] = {}
    for ri in range(1, len(table.rows)):
        if _is_docx_numeric_row(_docx_row_cell_texts(table.rows[ri])):
            continue
        if ri not in res.row_pages or ri in dup:
            return None  # a real data row could not be located reliably
        out[ri] = res.row_pages[ri]
    if len(out) < 2:
        return None
    ordered = [out[ri] for ri in sorted(out)]
    if ordered != sorted(ordered):
        return None  # non-monotonic → unreliable
    return out


def _cross_page_data_row_pages(
    doc: Document,
    table,
    table_idx: int,
    pdf_lines: list[PdfLine],
    numeric_row_idx: int | None,
    table_num: str | None = None,
    docx_path: Path | None = None,
) -> dict[int, int] | None:
    """Row→page map for the split: strict whole-row matcher first (rows that
    render on a single distinctive line), then the leading first-cell matcher
    (wrapped rows, table-local window), then — when both are ambiguous — the
    reliable per-row instrumentation mapper. Returns table-row-index → page, or
    None when no method maps it unambiguously."""
    window = _table_pdf_window(pdf_lines, table_num) if table_num else pdf_lines
    sig = next(
        (s for s in _collect_table_signatures(doc) if s.table_idx == table_idx),
        None,
    )
    if sig is not None:
        strict = _match_row_pages(sig, window)
        if strict is not None and len(set(strict.values())) >= 2:
            return strict
    skip = {numeric_row_idx} if numeric_row_idx is not None else set()
    lead = _match_data_row_pages_by_lead(table, window, skip)
    if lead is not None and len(set(lead.values())) >= 2:
        return lead
    # Both text matchers ambiguous → structural instrumentation fallback (the
    # expensive, eval-only path — see _cross_page_instrumentation_enabled).
    if not _cross_page_instrumentation_enabled():
        return None
    return _instrumented_data_row_pages(docx_path, table_idx)


def _cross_page_split_candidate(
    doc: Document,
    table_num: str,
    pdf_lines: list[PdfLine],
    *,
    source_docx_path: Path | None,
    docx_path: Path | None = None,
):
    """Locate the physical marker-less table for a cross-page blocker and a SAFE
    split point. Returns ``(table_idx, split_after, numeric_row_idx)`` or None.

    Eligible only if (reference spec + Stage 3): not source-bad; no existing
    adjacent continuation marker; rows map unambiguously to EXACTLY two pages;
    ≥1 data row stays in BOTH fragments; safe (non-merged) split boundary; safe
    table geometry."""
    # Source-bad duplicated content must never be auto-split/restructured.
    if _source_has_meaningful_duplicate_for_table(source_docx_path, str(table_num)) is not False:
        return None

    body_children = list(doc.element.body)
    para_by_xml = {p._element: p for p in doc.paragraphs}

    for table_idx, table in enumerate(doc.tables):
        if not _caption_before_table_matches(doc, table_idx, str(table_num)):
            continue
        tbl_xml = table._tbl
        if _table_has_adjacent_continuation_marker(body_children, para_by_xml, tbl_xml):
            continue
        rows_xml = tbl_xml.findall(qn("w:tr"))
        if len(rows_xml) < 3:
            continue

        numeric_row_idx = None
        if len(table.rows) > 1 and _is_docx_numeric_row(_docx_row_cell_texts(table.rows[1])):
            numeric_row_idx = 1

        row_pages = _cross_page_data_row_pages(
            doc, table, table_idx, pdf_lines, numeric_row_idx,
            table_num=str(table_num), docx_path=docx_path)
        if row_pages is None:
            continue
        distinct_pages = sorted(set(row_pages.values()))
        if len(distinct_pages) != 2:
            # only single-boundary (two-page) tables are split here; a 3+ page
            # table would leave a caption-less continuation fragment still
            # crossing pages that the gate cannot detect — skip honestly.
            continue

        first_page = distinct_pages[0]
        last_on_first = max(ri for ri, pg in row_pages.items() if pg == first_page)
        safe_after = _find_safe_split_after(rows_xml, last_on_first)
        if safe_after is None or safe_after < 1:
            continue
        min_after = 2 if numeric_row_idx is not None else 1
        if safe_after < min_after:
            continue  # would leave no real data row in the first fragment
        if len(rows_xml) - (safe_after + 1) < 1:
            continue  # would leave no real data row in the continuation fragment

        if not _split_geometry_is_safe(
            tbl_xml,
            table_index=table_idx,
            split_before_row=safe_after + 1,
            log_prefix="cross_page_split action=skip",
        ):
            continue

        return table_idx, safe_after, numeric_row_idx, min_after
    return None


def _set_tc_number_text(tc_xml, text: str) -> None:
    """Replace a table cell's text with ``text``, preserving the first run's
    properties (font) when present. Used to build a synthesized numeric column
    row from a cloned template row."""
    paras = tc_xml.findall(qn("w:p"))
    first = paras[0] if paras else None
    if first is None:
        first = OxmlElement("w:p")
        tc_xml.append(first)
    for extra in paras[1:]:
        tc_xml.remove(extra)
    runs = first.findall(qn("w:r"))
    if runs:
        keep = runs[0]
        for extra in runs[1:]:
            first.remove(extra)
        for t in keep.findall(qn("w:t")):
            keep.remove(t)
        t = OxmlElement("w:t")
        t.text = text
        keep.append(t)
    else:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        first.append(r)


def _synthesize_numeric_row_xml(template_tr):
    """Clone a real row (keeping cell structure / widths / formatting) and set its
    cells to the KFU numeric column index ``1 2 ... N``."""
    new = deepcopy(template_tr)
    for i, tc in enumerate(new.findall(qn("w:tc"))):
        _set_tc_number_text(tc, str(i + 1))
    return new


def _split_cross_page_table_with_marker(
    doc: Document,
    table_idx: int,
    split_after: int,
    table_num: str,
    *,
    numeric_row_idx: int | None,
) -> bool:
    """Split ``doc.tables[table_idx]`` after row ``split_after`` into two physical
    fragments per the canonical KFU rule. The continuation fragment repeats ONLY
    the numeric column row (``1 2 ... N``) followed by data — NEVER the semantic
    header, caption or title. The first fragment keeps caption/header + numeric
    row + data; when the source has no numeric row one is synthesized into BOTH
    fragments. Grid widths are inherited verbatim. Источник:/Примечание: that
    followed the original table stay after the continuation (final) fragment."""
    tbl_xml = doc.tables[table_idx]._tbl
    rows = tbl_xml.findall(qn("w:tr"))
    if split_after < 1 or split_after >= len(rows) - 1:
        return False

    tail_rows = [deepcopy(r) for r in rows[split_after + 1:]]
    if not tail_rows:
        return False

    if numeric_row_idx is not None and 0 < numeric_row_idx <= split_after:
        numeric_for_continuation = deepcopy(rows[numeric_row_idx])
        numeric_for_first = None  # already present in the first fragment
    else:
        # no numeric row in the source — synthesize from a data-row template
        # (normal formatting) for BOTH fragments.
        numeric_for_continuation = _synthesize_numeric_row_xml(rows[split_after])
        numeric_for_first = _synthesize_numeric_row_xml(rows[split_after])

    # continuation fragment: numeric row first, then data rows (NO semantic header)
    tbl2 = deepcopy(tbl_xml)
    for tr in list(tbl2.findall(qn("w:tr"))):
        tbl2.remove(tr)
    tbl2.append(numeric_for_continuation)
    for tr in tail_rows:
        tbl2.append(tr)

    # trim the first fragment
    for tr in rows[split_after + 1:]:
        tbl_xml.remove(tr)

    # ensure the first fragment carries the numeric row (right after the header)
    if numeric_for_first is not None:
        first_rows = tbl_xml.findall(qn("w:tr"))
        first_rows[0].addnext(numeric_for_first)

    marker = _build_continuation_para(f"Продолжение таблицы {table_num}")
    tbl_xml.addnext(marker)
    marker.addnext(tbl2)
    return True


def _cross_page_without_marker_probe(docx_path: Path, source_docx_path: Path | None):
    """Render once; return ``(fail_keys, cross_page_table_nums, pdf_lines)``. The
    pdf_lines are reused for split-boundary detection so a pass renders once for
    both the gate verdict and the row→page mapping."""
    from .rendered_table_validation import (
        build_rendered_table_identities,
        evaluate_table_layout_acceptance,
    )

    pdf_path: Path | None = None
    try:
        pdf_path = render_docx_to_pdf(Path(docx_path))
        pdf_lines = analyze_pdf_lines(pdf_path)
        doc = Document(str(docx_path))
        identities = build_rendered_table_identities(doc)
        source_ids = None
        if source_docx_path is not None:
            try:
                source_ids = build_rendered_table_identities(Document(str(source_docx_path)))
            except Exception:
                source_ids = None
        blockers = evaluate_table_layout_acceptance(
            pdf_lines, identities, doc=doc, source_identities=source_ids
        )
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)

    fail_keys = {(b.blocker_type, b.table_num) for b in blockers if b.severity == "fail"}
    cross = {
        b.table_num for b in blockers
        if b.severity == "fail"
        and b.blocker_type == "single_table_crosses_pages_without_marker"
        and b.table_num
    }
    return fail_keys, cross, pdf_lines


def _push_first_fragment_to_next_page(docx_path: Path, table_num: str) -> bool:
    """Insert the two-blank table-start-orphan guard before the caption of the
    table numbered ``table_num`` so its first fragment (caption + header) moves
    off a page bottom. Needed because after the split the first fragment is a
    continuation-chain head, which the standalone rendered orphan guard skips."""
    doc = Document(str(docx_path))
    for idx, table in enumerate(doc.tables):
        if not _caption_before_table_matches(doc, idx, str(table_num)):
            continue
        caption = _find_caption_paragraph_before_table(doc, table._tbl)
        if caption is None:
            return False
        caption_para_xml, _num = caption
        if _insert_table_start_orphan_blanks(caption_para_xml, target_count=2):
            doc.save(str(docx_path))
            return True
        return False
    return False


def _cross_page_cleanup_budget_seconds() -> float:
    """Per-document wall-clock cap for the cross-page cleanup (instrumentation +
    adaptive split attempts each render). Override with KPFU_CROSS_PAGE_BUDGET_S."""
    try:
        return max(30.0, float(os.environ.get("KPFU_CROSS_PAGE_BUDGET_S", "90")))
    except (TypeError, ValueError):
        return 90.0


def _verify_cross_page_split(
    docx_path: Path,
    source_docx_path: Path | None,
    table_num: str,
    table_idx: int,
    baseline_keys: set,
) -> str | None:
    """Re-render verify of an applied cross-page split. Returns None to ACCEPT, or
    a rollback reason. Clears a first-fragment orphan via the two-blank guard
    (the split makes the first fragment a continuation head the downstream orphan
    guard skips); requires the target cross-page fail gone, no new fail blocker,
    content preserved, and the caption-less continuation fragment single-page."""
    cross_key = ("single_table_crosses_pages_without_marker", table_num)
    orphan_key = ("orphaned_header_row", table_num)
    if _content_regressed(source_docx_path, docx_path):
        return "content_regression"
    try:
        after_keys, _ac, _apl = _cross_page_without_marker_probe(docx_path, source_docx_path)
    except Exception:
        return "post_render_failed"
    if orphan_key in after_keys and _push_first_fragment_to_next_page(docx_path, table_num):
        if _content_regressed(source_docx_path, docx_path):
            return "content_regression_after_orphan_fix"
        try:
            after_keys, _ac, _apl = _cross_page_without_marker_probe(docx_path, source_docx_path)
        except Exception:
            return "post_render_failed"
    if cross_key in after_keys:
        return "cross_page_remains"
    if orphan_key in after_keys:
        return "first_fragment_orphan_unfixable"
    new_fails = after_keys - baseline_keys
    if new_fails:
        return f"new_fail_blocker={sorted(new_fails)}"
    cont = _instrumented_data_row_pages(docx_path, table_idx + 1)
    if cont is not None and len(set(cont.values())) > 1:
        return "continuation_fragment_crosses_pages"
    return None


def cleanup_cross_page_without_marker_blockers_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """Insert valid continuation markers for marker-less tables that truly cross
    pages (acceptance blocker ``single_table_crosses_pages_without_marker``).

    Driven by the rendered acceptance gate (reliable table+page) — NOT the old
    globally-gated rendered-continuation path. For each flagged table the
    physical DOCX table is split at the rendered page boundary, a page-broken
    ``Продолжение таблицы N`` marker is inserted, and the header (+ numeric row)
    is repeated on the continuation fragment. Every split is re-render verified:
    it must clear that cross-page fail, add NO new fail blocker (same-page
    marker, orphaned header, neighbour flip, grid mismatch …) and preserve all
    content, else it is rolled back. At most one accepted split per pass, then
    re-probe because pagination shifts."""
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    repaired = 0
    seen: set[str] = set()
    entry_keys: set | None = None
    # Whole-document rollback snapshot: a per-split verify only sees the state at
    # that split, but accepted splits add a marker + numeric row + page break that
    # can SHIFT pagination and push OTHER tables across boundaries (a cascade that
    # surfaces only later). At exit we re-probe and, if ANY new fail blocker
    # appeared vs entry, roll the WHOLE cleanup back — it must never grow the
    # corpus fail set.
    entry_backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_cross_page_entry_"))
    entry_backup = entry_backup_dir / docx_path.name
    shutil.copy2(docx_path, entry_backup)
    # Hard wall-clock budget per document: instrumentation + adaptive split
    # attempts each render, so cap total time to keep format_docx bounded and
    # never runaway. Honoured between tables/passes (a single in-flight verify
    # always completes).
    deadline = time.monotonic() + _cross_page_cleanup_budget_seconds()
    for _pass in range(12):
        if time.monotonic() > deadline:
            logger.info("cross_page_split_budget_exhausted path=%s repaired=%s", docx_path, repaired)
            break
        try:
            baseline_keys, cross, pdf_lines = _cross_page_without_marker_probe(
                docx_path, source_docx_path
            )
        except Exception as exc:
            logger.info(
                "cross_page_split_probe_skip path=%s reason=render_failed error=%s",
                docx_path, exc,
            )
            break
        if entry_keys is None:
            entry_keys = set(baseline_keys)
        todo = sorted(n for n in cross if n not in seen)
        if not todo:
            break
        made_progress = False
        for table_num in todo:
            if time.monotonic() > deadline:
                logger.info("cross_page_split_budget_exhausted path=%s repaired=%s", docx_path, repaired)
                break
            backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_cross_page_split_"))
            backup_path = backup_dir / docx_path.name
            try:
                shutil.copy2(docx_path, backup_path)
                doc = Document(str(docx_path))
                cand = _cross_page_split_candidate(
                    doc, table_num, pdf_lines, source_docx_path=source_docx_path,
                    docx_path=docx_path,
                )
                if cand is None:
                    seen.add(table_num)
                    continue
                table_idx, split_after, numeric_row_idx, min_after = cand
                # Adaptive boundary: the rendered page boundary is measured on the
                # UNSPLIT table; adding the continuation marker + numeric row pushes
                # the first fragment down, so the natural boundary can overflow.
                # Try it, then progressively EARLIER boundaries (≥ min_after) until
                # one verifies — bounded to a few attempts.
                applied = False
                last_reason = "no_attempt"
                for sa in range(split_after, max(min_after, split_after - 3) - 1, -1):
                    if time.monotonic() > deadline:
                        last_reason = "budget_exhausted"
                        break
                    shutil.copy2(backup_path, docx_path)  # clean slate per attempt
                    doc = Document(str(docx_path))
                    if len(doc.tables[table_idx]._tbl.findall(qn("w:tr"))) - (sa + 1) < 1:
                        continue  # no data row left in the continuation fragment
                    if not _split_cross_page_table_with_marker(
                        doc, table_idx, sa, table_num, numeric_row_idx=numeric_row_idx
                    ):
                        continue
                    doc.save(str(docx_path))
                    last_reason = _verify_cross_page_split(
                        docx_path, source_docx_path, table_num, table_idx, baseline_keys
                    )
                    if last_reason is None:
                        repaired += 1
                        made_progress = True
                        applied = True
                        logger.info(
                            "cross_page_split_applied table_num=%s table_idx=%s split_after=%s numeric_row=%s",
                            table_num, table_idx, sa, numeric_row_idx,
                        )
                        break
                    logger.info(
                        "cross_page_split_attempt_rollback table_num=%s split_after=%s reason=%s",
                        table_num, sa, last_reason,
                    )
                seen.add(table_num)
                if applied:
                    break  # re-probe — pagination shifted
                shutil.copy2(backup_path, docx_path)
                logger.info(
                    "cross_page_split_rollback table_num=%s reason=%s",
                    table_num, last_reason,
                )
            except Exception as exc:
                shutil.copy2(backup_path, docx_path)
                seen.add(table_num)
                logger.info(
                    "cross_page_split_rollback table_num=%s reason=exception error=%s",
                    table_num, exc,
                )
            finally:
                shutil.rmtree(backup_dir, ignore_errors=True)
        if not made_progress:
            break
    # Whole-document cascade guard: if the accepted splits net-introduced any new
    # fail blocker vs entry (e.g. shifted a neighbour table across a page), roll
    # the entire cleanup back so the corpus fail set never grows.
    if repaired and entry_keys is not None:
        try:
            final_keys, _fc, _fl = _cross_page_without_marker_probe(docx_path, source_docx_path)
        except Exception:
            final_keys = None
        if final_keys is not None and (final_keys - entry_keys):
            shutil.copy2(entry_backup, docx_path)
            logger.info(
                "cross_page_split_doc_rollback path=%s reason=cascade_new_fail=%s reverted_splits=%d",
                docx_path, sorted(final_keys - entry_keys), repaired,
            )
            repaired = 0
    shutil.rmtree(entry_backup_dir, ignore_errors=True)
    if repaired:
        logger.info("cross_page_split total_repaired=%d", repaired)
    return repaired


def _find_table_index_by_caption(doc: Document, table_num: str) -> int | None:
    for i in range(len(doc.tables)):
        if _caption_before_table_matches(doc, i, str(table_num)):
            return i
    return None


def _cross_page_index_search_budget_seconds() -> float:
    """Per-document wall-clock cap for the DOCX-index cross-page split search.
    Override with KPFU_CROSS_INDEX_BUDGET_S."""
    try:
        return max(30.0, float(os.environ.get("KPFU_CROSS_INDEX_BUDGET_S", "150")))
    except (TypeError, ValueError):
        return 150.0


def cleanup_cross_page_by_index_search_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """Fallback marker-less cross-page splitter that does NOT depend on row→page
    mapping.

    The primary cleanup (`cleanup_cross_page_without_marker_blockers_inplace`)
    needs to map DOCX rows to rendered pages (strict/lead/instrumentation) and
    SKIPS a table when that mapping is ambiguous — leaving winnable tables
    unsplit. This fallback instead enumerates the split point by DOCX ROW INDEX:
    for each table still failing `single_table_crosses_pages_without_marker`, it
    builds the canonical 2-fragment layout (first = header + numeric + rows[:k];
    continuation = page-broken `Продолжение таблицы N` + numeric + rows[k:]) for
    candidate k (largest first fragment first), RENDERS, and accepts the first
    candidate that clears that table's cross fail with NO new fail key and
    preserved content. Otherwise every candidate is rolled back. PDF is used only
    to VERIFY, never to choose k. Bounded: per-table candidate cap + per-doc
    wall-clock budget; re-probes after each applied split (pagination shifts)."""
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    deadline = time.monotonic() + _cross_page_index_search_budget_seconds()
    repaired = 0
    seen: set[str] = set()
    for _pass in range(8):
        if time.monotonic() > deadline:
            break
        try:
            baseline_fail, cross_nums, _lines = _cross_page_without_marker_probe(
                docx_path, source_docx_path)
        except Exception as exc:
            logger.info("cross_index_search_probe_skip reason=render_failed error=%s", exc)
            break
        todo = sorted(n for n in cross_nums if n not in seen)
        if not todo:
            break
        made_progress = False
        for num in todo:
            if time.monotonic() > deadline:
                break
            seen.add(num)
            probe = Document(str(docx_path))
            tidx = _find_table_index_by_caption(probe, str(num))
            if tidx is None:
                continue
            rows_xml = probe.tables[tidx]._tbl.findall(qn("w:tr"))
            n = len(rows_xml)
            if n < 3:
                continue
            numeric_idx = 1 if (
                len(probe.tables[tidx].rows) > 1
                and _is_docx_numeric_row(_docx_row_cell_texts(probe.tables[tidx].rows[1]))
            ) else None
            min_after = 2 if numeric_idx is not None else 1
            # candidate split points: largest first fragment first (canonical),
            # capped so a stubborn table cannot exhaust the budget.
            candidates = [k for k in range(n - 2, min_after - 1, -1)][:8]
            applied = False
            for k in candidates:
                if time.monotonic() > deadline:
                    break
                backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_cross_index_"))
                backup_path = backup_dir / docx_path.name
                try:
                    shutil.copy2(docx_path, backup_path)
                    doc = Document(str(docx_path))
                    if not _split_cross_page_table_with_marker(
                        doc, tidx, k, str(num), numeric_row_idx=numeric_idx
                    ):
                        shutil.copy2(backup_path, docx_path)
                        continue
                    doc.save(str(docx_path))
                    reason = None
                    if _content_regressed(source_docx_path, docx_path):
                        reason = "content_regression"
                    else:
                        try:
                            after_fail, _c, _l = _cross_page_without_marker_probe(
                                docx_path, source_docx_path)
                        except Exception:
                            reason = "post_render_failed"
                        else:
                            if ("single_table_crosses_pages_without_marker", num) in after_fail:
                                reason = "cross_remains"
                            elif after_fail - baseline_fail:
                                reason = f"new_fail={sorted(after_fail - baseline_fail)}"
                    if reason is None:
                        repaired += 1
                        made_progress = True
                        applied = True
                        logger.info(
                            "cross_index_split_applied table_num=%s k=%s of_rows=%s", num, k, n
                        )
                        break
                    shutil.copy2(backup_path, docx_path)
                    logger.info(
                        "cross_index_split_reject table_num=%s k=%s reason=%s", num, k, reason
                    )
                except Exception as exc:
                    shutil.copy2(backup_path, docx_path)
                    logger.info(
                        "cross_index_split_reject table_num=%s k=%s reason=exception error=%s",
                        num, k, exc,
                    )
                finally:
                    shutil.rmtree(backup_dir, ignore_errors=True)
            if applied:
                break  # re-probe from the top: pagination shifted
        if not made_progress:
            break
    if repaired:
        logger.info("cross_index_search total_repaired=%d", repaired)
    return repaired


def _move_table_block_to_next_page(doc: Document, table_num: str) -> bool:
    """Set ``pageBreakBefore`` on the caption paragraph of table ``table_num`` so
    the whole block (caption + title + table) starts on the next page — real Word
    mechanics, no blank paragraphs. Returns True when applied (skips if the caption
    already carries a page break, so a table is never moved twice)."""
    tidx = _find_table_index_by_caption(doc, str(table_num))
    if tidx is None:
        return False
    cap = _find_caption_paragraph_before_table(doc, doc.tables[tidx]._tbl)
    if cap is None:
        return False
    cap_para = cap[0]
    pPr = cap_para.find(qn("w:pPr"))
    if pPr is not None:
        pbb = pPr.find(qn("w:pageBreakBefore"))
        if pbb is not None:
            val = pbb.get(qn("w:val"))
            # skip only when the break is already ENABLED; a disabled break
            # (val='0', left by an upstream pass) is exactly what we re-enable.
            if val is None or val.strip().lower() in ("true", "1", "on"):
                return False
    _force_marker_page_break(cap_para)
    return True


def cleanup_cross_page_by_block_move_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """Whole-block move (with cascade) for cross-page tables that are small/medium
    and cross only because they START near a page bottom — moving the whole block
    to the next page makes it fit, whereas a split would orphan the header.

    For each table still failing `single_table_crosses_pages_without_marker`, set a
    page break before its caption. If the move shifts a NEIGHBOUR table into a new
    cross / same-page fail, add that neighbour to the SAME batch and move it too
    (bounded cascade depth). The batch is accepted only if the rendered fail set
    STRICTLY SHRINKS with NO new fail key and content is preserved; otherwise the
    whole batch is rolled back. Runs after the split-based cleanups, so it only
    handles residual tables those could not split cleanly. Budget-capped."""
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    deadline = time.monotonic() + _cross_page_index_search_budget_seconds()
    repaired = 0
    seen: set[str] = set()
    for _pass in range(8):
        if time.monotonic() > deadline:
            break
        try:
            baseline_fail, cross_nums, _l = _cross_page_without_marker_probe(
                docx_path, source_docx_path)
        except Exception as exc:
            logger.info("cross_block_move_probe_skip reason=render_failed error=%s", exc)
            break
        todo = sorted(n for n in cross_nums if n not in seen)
        if not todo:
            break
        made_progress = False
        for num in todo:
            if time.monotonic() > deadline:
                break
            seen.add(num)
            backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_block_move_"))
            backup_path = backup_dir / docx_path.name
            try:
                shutil.copy2(docx_path, backup_path)
                batch: list[str] = []
                cur = num
                ok_batch = False
                final_fail: set | None = None
                for _depth in range(5):  # cascade depth cap
                    if time.monotonic() > deadline:
                        break
                    doc = Document(str(docx_path))
                    if not _move_table_block_to_next_page(doc, cur):
                        break
                    doc.save(str(docx_path))
                    batch.append(cur)
                    if _content_regressed(source_docx_path, docx_path):
                        break
                    try:
                        after_fail, _c2, _l2 = _cross_page_without_marker_probe(
                            docx_path, source_docx_path)
                    except Exception:
                        break
                    new = after_fail - baseline_fail
                    if not new and len(after_fail) < len(baseline_fail):
                        ok_batch = True
                        final_fail = after_fail
                        break
                    nxt = None
                    for (bt, tn) in sorted(new):
                        if bt in (
                            "single_table_crosses_pages_without_marker",
                            "same_page_continuation",
                        ) and tn not in batch:
                            nxt = tn
                            break
                    if nxt is None:
                        break  # dead end — a non-movable new fail
                    cur = nxt
                if ok_batch and final_fail is not None:
                    repaired += len(baseline_fail) - len(final_fail)
                    made_progress = True
                    for b in batch:
                        seen.discard(b)
                    logger.info("cross_block_move_batch_applied batch=%s", batch)
                    break  # re-probe from the top: pagination shifted
                shutil.copy2(backup_path, docx_path)
                logger.info(
                    "cross_block_move_batch_rollback start=%s batch=%s", num, batch
                )
            except Exception as exc:
                shutil.copy2(backup_path, docx_path)
                logger.info(
                    "cross_block_move_batch_rollback start=%s reason=exception error=%s",
                    num, exc,
                )
            finally:
                shutil.rmtree(backup_dir, ignore_errors=True)
        if not made_progress:
            break
    if repaired:
        logger.info("cross_block_move total_repaired=%d", repaired)
    return repaired


# ── Continuation semantic-header normalizer (canonical KFU rule) ──────────────
#
# A `Продолжение таблицы N` fragment must repeat ONLY the numeric column row
# (`1 2 ... N`), never the semantic header / caption / title (Rybakov gold).
# Existing manual chains (student-authored) often duplicate the semantic header
# above the numeric row — strip that leading duplicate header so the fragment
# starts with the numeric row. Content-safe (a duplicate header row is excluded
# from the content gate); rolled back if a re-render adds any new fail blocker.


def _iter_continuation_chains(doc: Document):
    """Yield (first_table_idx, marker_para_xml, second_table_idx, num) for each
    ``tbl -> 'Продолжение таблицы N' -> tbl`` chain, in body order."""
    body = list(doc.element.body)
    para_by_xml = {p._element: p for p in doc.paragraphs}
    tbl_index = {t._tbl: i for i, t in enumerate(doc.tables)}
    for i, ch in enumerate(body):
        if ch.tag != qn("w:p"):
            continue
        para = para_by_xml.get(ch)
        num = _strict_marker_table_num(para.text if para is not None else "")
        if not num:
            continue
        first_idx = next(
            (tbl_index[body[k]] for k in range(i - 1, -1, -1)
             if body[k].tag == qn("w:tbl") and body[k] in tbl_index),
            None,
        )
        second_idx = next(
            (tbl_index[body[k]] for k in range(i + 1, min(i + 3, len(body)))
             if body[k].tag == qn("w:tbl") and body[k] in tbl_index),
            None,
        )
        if first_idx is None or second_idx is None:
            continue
        yield first_idx, ch, second_idx, num


def _continuation_starts_with_semantic_header(first, second) -> bool:
    if not first.rows or not second.rows:
        return False
    v0 = _docx_row_cell_texts(second.rows[0])
    if _is_docx_numeric_row(v0):
        return False  # already canonical
    return _docx_row_fingerprint(second.rows[0]) == _docx_row_fingerprint(first.rows[0])


def _strip_continuation_semantic_headers(doc: Document, only_num: str | None = None) -> int:
    """Remove the leading duplicate semantic header row(s) from continuation
    fragments so each starts with the numeric row; if no numeric row follows,
    insert the first fragment's numeric row (or a synthesized one). Never removes
    a data row; never empties a table; skips merged-cell fragments. When
    ``only_num`` is given, only that chain is touched."""
    fixed = 0
    for first_idx, _marker, second_idx, num in _iter_continuation_chains(doc):
        if only_num is not None and num != only_num:
            continue
        try:
            first = doc.tables[first_idx]
            second = doc.tables[second_idx]
        except IndexError:
            continue
        if _table_has_merged_cells_docx(second):
            continue
        if not _continuation_starts_with_semantic_header(first, second):
            continue
        header_fp = _docx_row_fingerprint(first.rows[0])
        removed = 0
        while len(second.rows) > 1:
            r0 = second.rows[0]
            if _is_docx_numeric_row(_docx_row_cell_texts(r0)):
                break
            if _docx_row_fingerprint(r0) != header_fp:
                break
            _remove_xml_node(r0._tr)
            removed += 1
        if removed <= 0:
            continue
        # ensure a numeric row leads the continuation fragment
        if not (second.rows and _is_docx_numeric_row(_docx_row_cell_texts(second.rows[0]))):
            num_tr = None
            for r in first.rows:
                if _is_docx_numeric_row(_docx_row_cell_texts(r)):
                    num_tr = deepcopy(r._tr)
                    break
            if num_tr is None and second.rows:
                num_tr = _synthesize_numeric_row_xml(second.rows[0]._tr)
            if num_tr is not None and second.rows:
                second.rows[0]._tr.addprevious(num_tr)
        fixed += 1
    return fixed


def normalize_continuation_semantic_header_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """Enforce the canonical KFU continuation rule on existing manual chains:
    strip the duplicate semantic header from each `Продолжение таблицы N` fragment
    so it starts with the numeric column row. Deterministic + content-safe, and
    verified PER CHAIN: a chain is accepted only if its
    `semantic_header_repeated_on_continuation` fail clears, content is preserved,
    and the re-render adds NO new fail blocker (e.g. a pagination flip) — else
    that chain is rolled back. At most one accepted fix per pass, then re-probe
    because removing a row shifts pagination for the remaining chains."""
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    repaired = 0
    seen: set[str] = set()
    for _pass in range(12):
        try:
            baseline_keys, _c, _l = _cross_page_without_marker_probe(docx_path, source_docx_path)
        except Exception as exc:
            logger.info("continuation_header_normalize_skip reason=render_failed error=%s", exc)
            break
        try:
            probe = Document(str(docx_path))
        except Exception:
            break
        targets = sorted({
            num for f, _m, s, num in _iter_continuation_chains(probe)
            if num not in seen
            and _continuation_starts_with_semantic_header(probe.tables[f], probe.tables[s])
        })
        if not targets:
            break
        made_progress = False
        for num in targets:
            backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_cont_header_norm_"))
            backup_path = backup_dir / docx_path.name
            try:
                shutil.copy2(docx_path, backup_path)
                doc = Document(str(docx_path))
                if _strip_continuation_semantic_headers(doc, only_num=num) <= 0:
                    seen.add(num)
                    continue
                doc.save(str(docx_path))

                reason = None
                if _content_regressed(source_docx_path, docx_path):
                    reason = "content_regression"
                else:
                    try:
                        after_keys, _c2, _l2 = _cross_page_without_marker_probe(
                            docx_path, source_docx_path
                        )
                    except Exception:
                        reason = "post_render_failed"
                    else:
                        if ("semantic_header_repeated_on_continuation", num) in after_keys:
                            reason = "header_remains"
                        elif after_keys - baseline_keys:
                            reason = f"new_fail_blocker={sorted(after_keys - baseline_keys)}"

                if reason is None:
                    repaired += 1
                    made_progress = True
                    seen.add(num)
                    logger.info("continuation_header_normalize_applied table_num=%s", num)
                    break  # re-probe — pagination shifted
                shutil.copy2(backup_path, docx_path)
                seen.add(num)
                logger.info(
                    "continuation_header_normalize_rollback table_num=%s reason=%s", num, reason
                )
            except Exception as exc:
                shutil.copy2(backup_path, docx_path)
                seen.add(num)
                logger.info(
                    "continuation_header_normalize_rollback table_num=%s reason=exception error=%s",
                    num, exc,
                )
            finally:
                shutil.rmtree(backup_dir, ignore_errors=True)
        if not made_progress:
            break
    if repaired:
        logger.info("continuation_header_normalize total_repaired=%d", repaired)
    return repaired


def _table_grid_proportions(tbl_xml) -> list[float] | None:
    grid = tbl_xml.find(qn("w:tblGrid"))
    if grid is None:
        return None
    cols = [int(c.get(qn("w:w"), "0") or 0) for c in grid.findall(qn("w:gridCol"))]
    total = sum(cols)
    if not cols or total <= 0:
        return None
    return [c / total for c in cols]


def _fragment_grid_drifts(first, second) -> bool:
    """True when two fragments share a column count but their grid column-width
    proportions drift beyond the detector's twip-rounding tolerance (0.02)."""
    pa = _table_grid_proportions(first._tbl)
    pb = _table_grid_proportions(second._tbl)
    if not pa or not pb or len(pa) != len(pb):
        return False
    return max(abs(a - b) for a, b in zip(pa, pb)) > 0.02


def _copy_fragment_grid_first_to_second(first, second) -> bool:
    """Make the continuation fragment ``second`` render with the first fragment's
    column grid: copy tblGrid column widths, the master tblW, and per-cell tcW.
    Content-safe — touches only widths, never a row or its text. Skips merged-cell
    fragments and mismatched column counts. Returns True when widths changed."""
    fgrid = first._tbl.find(qn("w:tblGrid"))
    sgrid = second._tbl.find(qn("w:tblGrid"))
    if fgrid is None or sgrid is None:
        return False
    fcols = fgrid.findall(qn("w:gridCol"))
    scols = sgrid.findall(qn("w:gridCol"))
    if not fcols or len(fcols) != len(scols):
        return False
    if _table_has_merged_cells_docx(first) or _table_has_merged_cells_docx(second):
        return False
    changed = False
    for fc, sc in zip(fcols, scols):
        w = fc.get(qn("w:w"))
        if w is not None and sc.get(qn("w:w")) != w:
            sc.set(qn("w:w"), w)
            changed = True
    fpr = first._tbl.find(qn("w:tblPr"))
    spr = second._tbl.find(qn("w:tblPr"))
    if fpr is not None and spr is not None:
        fw = fpr.find(qn("w:tblW"))
        sw = spr.find(qn("w:tblW"))
        if fw is not None and sw is not None:
            for attr in (qn("w:w"), qn("w:type")):
                v = fw.get(attr)
                if v is not None and sw.get(attr) != v:
                    sw.set(attr, v)
                    changed = True
    widths = _row_cell_widths(first.rows[0]) if first.rows else []
    if widths:
        for row in second.rows:
            _apply_row_cell_widths(row, widths)
        changed = True
    return changed


def _iter_grid_fragment_pairs(doc: Document):
    """Yield (first_idx, second_idx) for adjacent table pairs separated only by
    blank/marker paragraphs — the exact pairs the fragment_grid_mismatch detector
    inspects, so normalizing them matches what the gate measures."""
    from .rendered_table_validation import _is_blank_or_marker_paragraph
    body = list(doc.element.body)
    tbl_by_xml = {t._tbl: i for i, t in enumerate(doc.tables)}
    tbl_positions = [i for i, node in enumerate(body) if node.tag == qn("w:tbl")]
    for a, b in zip(tbl_positions, tbl_positions[1:]):
        between = body[a + 1:b]
        if any(not _is_blank_or_marker_paragraph(node) for node in between):
            continue
        fi = tbl_by_xml.get(body[a])
        si = tbl_by_xml.get(body[b])
        if fi is None or si is None:
            continue
        yield fi, si


def normalize_fragment_grid_widths_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """Normalize continuation-fragment column widths to the first fragment's grid
    so ``fragment_grid_mismatch`` (needs_human_review) clears. Deterministic and
    content-safe: only column widths change; never a row or its text. All drifting
    same-column-count fragment pairs are normalized in one shot, then verified
    whole-doc — accepted only if content is preserved and NO new fail blocker
    (e.g. a squeeze) appears; otherwise the whole doc is rolled back. Bounded: at
    most two renders (baseline + verify)."""
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    try:
        baseline_fail, _cross, _lines = _cross_page_without_marker_probe(
            docx_path, source_docx_path)
    except Exception as exc:
        logger.info("fragment_grid_normalize_skip reason=baseline_render_failed error=%s", exc)
        return 0
    try:
        doc = Document(str(docx_path))
    except Exception:
        return 0
    changed = 0
    for fi, si in _iter_grid_fragment_pairs(doc):
        try:
            first = doc.tables[fi]
            second = doc.tables[si]
        except IndexError:
            continue
        if not _fragment_grid_drifts(first, second):
            continue
        if _copy_fragment_grid_first_to_second(first, second):
            changed += 1
    if changed == 0:
        return 0
    backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_grid_norm_"))
    backup_path = backup_dir / docx_path.name
    try:
        shutil.copy2(docx_path, backup_path)
        doc.save(str(docx_path))
        reason = None
        if _content_regressed(source_docx_path, docx_path):
            reason = "content_regression"
        else:
            try:
                after_fail, _c2, _l2 = _cross_page_without_marker_probe(
                    docx_path, source_docx_path)
            except Exception:
                reason = "post_render_failed"
            else:
                if after_fail - baseline_fail:
                    reason = f"new_fail_blocker={sorted(after_fail - baseline_fail)}"
        if reason is None:
            logger.info("fragment_grid_normalize_applied pairs=%d", changed)
            return changed
        shutil.copy2(backup_path, docx_path)
        logger.info("fragment_grid_normalize_rollback reason=%s", reason)
        return 0
    except Exception as exc:
        shutil.copy2(backup_path, docx_path)
        logger.info("fragment_grid_normalize_rollback reason=exception error=%s", exc)
        return 0
    finally:
        shutil.rmtree(backup_dir, ignore_errors=True)


def cleanup_entangled_same_page_group_inplace(
    docx_path: Path,
    *,
    source_docx_path: Path | None = None,
    report: FormattingReport | None = None,
) -> int:
    """Repair ENTANGLED same-page continuation chains that the per-table cleanup
    cannot fix alone.

    When two or more adjacent `same_page_continuation` /
    `semantic_header_repeated_on_continuation` chains sit close together, fixing
    one shifts pagination and pushes a fail onto its neighbour, so
    `cleanup_same_page_continuation_blockers_inplace` and
    `normalize_continuation_semantic_header_inplace` roll each candidate back
    individually. Fix the whole group ATOMICALLY instead: force a page break
    before every remaining such marker AND strip its continuation's semantic
    header (numeric-led canonical), in a single mutation, then verify whole-doc.

    Accept only if content is preserved and the rendered fail set STRICTLY SHRINKS
    with NO new fail key (a page break that merely trades `same_page_continuation`
    for `single_table_crosses_pages_without_marker` — a multi-page first fragment —
    is a new key and is rejected). Otherwise the whole group is rolled back.
    Content-safe: page-break + header-strip never delete a data row. Bounded: one
    mutation + one verify render."""
    docx_path = Path(docx_path)
    source_docx_path = Path(source_docx_path) if source_docx_path is not None else None
    try:
        baseline_fail, _c, _l = _cross_page_without_marker_probe(docx_path, source_docx_path)
    except Exception as exc:
        logger.info("entangled_same_page_group_skip reason=baseline_render_failed error=%s", exc)
        return 0
    group = sorted({
        num for (btype, num) in baseline_fail
        if btype in ("same_page_continuation", "semantic_header_repeated_on_continuation")
        and num
    })
    if len(group) < 2:
        # a single residual marker is already handled (and rolled back) by the
        # per-table cleanup; grouping adds value only for entangled neighbours.
        return 0
    backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_entangled_group_"))
    backup_path = backup_dir / docx_path.name
    try:
        shutil.copy2(docx_path, backup_path)
        doc = Document(str(docx_path))
        touched: set[str] = set()
        for _first_idx, marker_xml, _second_idx, num in list(_iter_continuation_chains(doc)):
            if num in group:
                _force_marker_page_break(marker_xml)
                touched.add(num)
        for num in group:
            _strip_continuation_semantic_headers(doc, only_num=num)
        if not touched:
            shutil.copy2(backup_path, docx_path)
            return 0
        doc.save(str(docx_path))

        reason = None
        if _content_regressed(source_docx_path, docx_path):
            reason = "content_regression"
        else:
            try:
                after_fail, _c2, _l2 = _cross_page_without_marker_probe(docx_path, source_docx_path)
            except Exception:
                reason = "post_render_failed"
            else:
                new = after_fail - baseline_fail
                if new:
                    reason = f"new_fail_blocker={sorted(new)}"
                elif len(after_fail) >= len(baseline_fail):
                    reason = "no_net_reduction"
        if reason is None:
            removed = len(baseline_fail) - len(after_fail)
            logger.info(
                "entangled_same_page_group_applied group=%s fails_removed=%d", group, removed
            )
            return removed
        shutil.copy2(backup_path, docx_path)
        logger.info("entangled_same_page_group_rollback group=%s reason=%s", group, reason)
        return 0
    except Exception as exc:
        shutil.copy2(backup_path, docx_path)
        logger.info("entangled_same_page_group_rollback reason=exception error=%s", exc)
        return 0
    finally:
        shutil.rmtree(backup_dir, ignore_errors=True)


def restore_docx_if_same_page_continuation_markers(
    docx_path: Path,
    backup_docx_path: Path,
    *,
    report: FormattingReport | None = None,
    context: str = "post_render",
) -> bool:
    """Restore backup when final render contains fake same-page continuations."""
    try:
        violations = _same_page_continuation_marker_violations_for_docx(Path(docx_path))
    except Exception as exc:
        logger.info(
            "post_render_same_page_marker_validation_skip context=%s reason=validation_failed error=%s",
            context, exc,
        )
        return False

    if not violations:
        return False

    shutil.copy2(backup_docx_path, docx_path)
    details = [
        f"{v.marker_text}@p{v.marker_page}:prev={v.previous_table_page}:next={v.following_table_page}:confidence={v.confidence}"
        for v in violations
    ]
    logger.warning(
        "post_render_same_page_marker_violation context=%s restored_backup=1 violations=%s",
        context, ";".join(details),
    )
    try:
        remaining = _same_page_continuation_marker_violations_for_docx(Path(docx_path))
    except Exception as exc:
        logger.info(
            "post_render_same_page_marker_recheck_skip context=%s reason=validation_failed error=%s",
            context, exc,
        )
        return True
    if remaining:
        logger.warning(
            "post_render_same_page_marker_violation_unresolved context=%s restored_backup=1 violations=%s",
            context,
            ";".join(f"{v.marker_text}@p{v.marker_page}" for v in remaining),
        )
    return True


def warn_same_page_continuation_marker_violations(
    docx_path: Path,
    *,
    report: FormattingReport | None = None,
) -> int:
    if report is None:
        return 0
    try:
        violations = _same_page_continuation_marker_violations_for_docx(Path(docx_path))
    except Exception as exc:
        logger.info(
            "same_page_marker_warning_skip path=%s reason=validation_failed error=%s",
            docx_path,
            exc,
        )
        return 0
    if not violations:
        return 0

    existing = set(report.warnings)
    emitted = 0
    seen: set[tuple[str, int]] = set()
    for violation in violations:
        table_num = _strict_marker_table_num(violation.marker_text) or "?"
        page = int(violation.marker_page or 0)
        key = (table_num, page)
        if key in seen:
            continue
        seen.add(key)
        warning = (
            f"Проверьте таблицу {table_num}: "
            f"на стр. {page} повторный фрагмент виден на той же странице."
        )
        if warning in existing:
            continue
        report.warn(warning)
        existing.add(warning)
        emitted += 1
    return emitted


def remove_same_page_continuation_markers_inplace(
    docx_path: Path,
    *,
    report: FormattingReport | None = None,
) -> int:
    """
    Render *docx_path*, find any "Продолжение таблицы" markers that land on the
    same page as both their preceding and following table segments, and remove
    those marker paragraphs from the document.

    Such markers arise when a DOCX-only split (from `apply_table_continuation`)
    is no longer valid after the TOC size changes: the surrounding table now fits
    on one page and the marker is a stale artefact.  Removing it is correct —
    the table no longer straddles a page boundary, so no continuation header is
    needed.

    Returns the number of markers removed.  The document is saved only when at
    least one marker is removed.  If rendering fails the function logs and returns 0.
    """
    try:
        violations = _same_page_continuation_marker_violations_for_docx(Path(docx_path))
    except Exception as exc:
        logger.info(
            "remove_same_page_markers_skip path=%s reason=render_failed error=%s",
            docx_path, exc,
        )
        return 0

    if not violations:
        return 0

    violation_texts = {v.marker_text.strip() for v in violations}
    with tempfile.TemporaryDirectory(prefix="kpfu_same_page_marker_cleanup_") as tmp:
        candidate_path = Path(tmp) / Path(docx_path).name
        shutil.copy2(docx_path, candidate_path)
        removed = _remove_strict_continuation_marker_texts(candidate_path, violation_texts)
        if not removed:
            return 0
        try:
            regressions = _rendered_continuation_deletion_regressions(candidate_path)
        except Exception as exc:
            logger.info(
                "remove_same_page_markers_preserved path=%s reason=rendered_validation_failed error=%s",
                docx_path, exc,
            )
            return 0
        if regressions:
            logger.warning(
                "remove_same_page_markers_preserved path=%s reason=rendered_continuation_regression violations=%s",
                docx_path,
                _format_rendered_deletion_regressions(regressions),
            )
            _warn_unsafe_same_page_marker_cleanup(report, regressions)
            return 0

        shutil.copy2(candidate_path, docx_path)
        for raw in sorted(violation_texts):
            logger.info("remove_same_page_markers: removed %r from %s", raw, docx_path)
        logger.info(
            "remove_same_page_markers_done path=%s removed=%d",
            docx_path, removed,
        )
        return removed


# E1 — Phase 3 marker-split candidate classification (logging-only, no behavior change).
# Cheap predicate that selects which tables WOULD be diagnosed under future E2.
# Does not render. Does not mutate. Does not call LibreOffice or any PDF tool.
_MIN_ROWS_FOR_SPLIT_CANDIDACY = 6


def _classify_marker_split_candidates(doc: Document) -> dict:
    """Return cheap candidate classification for Phase 3 marker split.

    Output keys:
      total_tables                 -- int, len(doc.tables)
      manual_continuation_skipped  -- list[int], indices in an already-valid manual chain
      no_caption_skipped           -- list[int], no standard 'Таблица N' caption above
      tiny_table_skipped           -- list[int], rows < _MIN_ROWS_FOR_SPLIT_CANDIDACY
      candidate_tables             -- list[int], priority-ordered
      candidate_priority           -- list[(table_idx, priority_score)], priority-ordered
      candidate_budget             -- int, current _marker_split_max_renders()
      would_process                -- list[int], first N candidates within budget
      would_skip_for_budget        -- list[int], candidates beyond budget
    """
    from . import table_markers

    total_tables = len(doc.tables)
    manual_skip_ids = _valid_manual_continuation_table_indexes(doc)

    try:
        contexts = table_markers._iter_body_tables_with_context(doc)
    except Exception:
        contexts = []

    manual_continuation_skipped: list[int] = []
    no_caption_skipped: list[int] = []
    tiny_table_skipped: list[int] = []
    candidate_pairs: list[tuple[int, int]] = []

    for idx, table in enumerate(doc.tables):
        if idx in manual_skip_ids:
            manual_continuation_skipped.append(idx)
            continue

        if idx < len(contexts):
            ctx = contexts[idx]
            has_standard = bool(ctx.get("has_standard_table_caption"))
            is_appendix = bool(ctx.get("appendix_table"))
        else:
            has_standard = False
            is_appendix = False

        if not has_standard:
            no_caption_skipped.append(idx)
            continue

        rows_count = len(table.rows)
        if rows_count < _MIN_ROWS_FOR_SPLIT_CANDIDACY:
            tiny_table_skipped.append(idx)
            continue

        priority = rows_count + (1 if is_appendix else 0)
        candidate_pairs.append((idx, priority))

    candidate_priority = sorted(candidate_pairs, key=lambda x: (-x[1], x[0]))
    candidate_tables = [idx for idx, _ in candidate_priority]
    budget = _marker_split_max_renders()

    return {
        "total_tables": total_tables,
        "manual_continuation_skipped": manual_continuation_skipped,
        "no_caption_skipped": no_caption_skipped,
        "tiny_table_skipped": tiny_table_skipped,
        "candidate_tables": candidate_tables,
        "candidate_priority": candidate_priority,
        "candidate_budget": budget,
        "would_process": candidate_tables[:budget],
        "would_skip_for_budget": candidate_tables[budget:],
    }


def _classify_marker_duplicate_rows(
    diagnostic,
    *,
    header_rows: int = 1,
) -> tuple[str, list[int]]:
    if not diagnostic.duplicate_rows:
        return "none", []

    data_duplicate_rows = sorted(
        row_index
        for row_index in diagnostic.duplicate_rows
        if row_index >= header_rows
    )
    if data_duplicate_rows:
        return "data_rows", data_duplicate_rows

    return "header_only", sorted(diagnostic.duplicate_rows)


def _non_header_rows_are_clean(
    diagnostic,
    *,
    header_rows: int = 1,
) -> bool:
    if any(row_index >= header_rows for row_index in diagnostic.missing_rows):
        return False
    if any(row_index >= header_rows for row_index in diagnostic.duplicate_rows):
        return False

    row_pages = {
        row_index: page_num
        for row_index, page_num in sorted(diagnostic.row_pages.items())
        if row_index >= header_rows
    }
    expected_rows = list(range(header_rows, diagnostic.rows_count))
    return list(row_pages) == expected_rows


def _is_header_only_duplicate_safe(
    diagnostic,
    *,
    header_rows: int = 1,
) -> bool:
    duplicate_classification, _ = _classify_marker_duplicate_rows(
        diagnostic,
        header_rows=header_rows,
    )
    return (
        duplicate_classification == "header_only"
        and _non_header_rows_are_clean(diagnostic, header_rows=header_rows)
    )


def _evaluate_marker_split_diagnostic(
    diagnostic,
    *,
    header_rows: int = 1,
) -> _MarkerSplitDecision:
    if diagnostic.error_message:
        return _MarkerSplitDecision(False, None, "mapping_error")
    if len(diagnostic.pages_detected) != 2:
        return _MarkerSplitDecision(False, None, "not_2_pages")

    duplicate_classification, _ = _classify_marker_duplicate_rows(
        diagnostic,
        header_rows=header_rows,
    )
    if duplicate_classification == "data_rows":
        return _MarkerSplitDecision(False, None, "duplicate_rows")
    if diagnostic.missing_rows not in ([], [0]):
        return _MarkerSplitDecision(False, None, "missing_rows_outside_header")
    if (
        duplicate_classification == "header_only"
        and not _is_header_only_duplicate_safe(diagnostic, header_rows=header_rows)
    ):
        return _MarkerSplitDecision(False, None, "duplicate_rows")

    row_pages = {
        row_index: page_num
        for row_index, page_num in sorted(diagnostic.row_pages.items())
        if row_index >= header_rows
    }
    expected_rows = list(range(header_rows, diagnostic.rows_count))
    if list(row_pages) != expected_rows:
        return _MarkerSplitDecision(False, None, "no_boundary")

    first_page = None
    second_page = None
    split_before_row = None
    expected_first_page, expected_second_page = diagnostic.pages_detected

    for row_index, page_num in row_pages.items():
        if page_num not in diagnostic.pages_detected:
            return _MarkerSplitDecision(False, None, "non_monotonic_pages")

        if first_page is None:
            if page_num != expected_first_page:
                return _MarkerSplitDecision(False, None, "non_monotonic_pages")
            first_page = page_num
            continue

        if second_page is None:
            if page_num == first_page:
                continue
            if page_num != expected_second_page:
                return _MarkerSplitDecision(False, None, "non_monotonic_pages")
            second_page = page_num
            split_before_row = row_index
            continue

        if page_num != second_page:
            return _MarkerSplitDecision(False, None, "non_monotonic_pages")

    if split_before_row is None:
        return _MarkerSplitDecision(False, None, "no_boundary")
    return _MarkerSplitDecision(True, split_before_row, None)


def _map_marker_split_apply_error(exc: Exception) -> str:
    text = str(exc).lower()
    if "standard table caption" in text:
        return "ordinary_without_standard_caption"
    if "tblgrid" in text or "grid" in text:
        return "unsupported_grid"
    if "complex merged header" in text:
        return "complex_merged_header"
    if "malformed" in text:
        return "malformed_numbered_row"
    if "header_rows" in text:
        return "unsupported_header_rows"
    return "mutation_error"


def _request_id_from_docx_path(docx_path: Path) -> str:
    match = re.match(r"^(\d+)_", docx_path.name)
    return match.group(1) if match else "-"


def _count_marker_renders(diagnostics) -> int:
    # table_markers may retry with 2pt markers; marker_font_size_pt reflects
    # whether the 1pt attempt was enough or a second render was needed.
    return sum(2 if diagnostic.marker_font_size_pt == 2 else 1 for diagnostic in diagnostics)


def _record_marker_skip(reason: str | None) -> None:
    stats = _ACTIVE_MARKER_STATS
    if stats is None:
        return
    key = reason or "unknown"
    stats.skipped_reasons[key] = stats.skipped_reasons.get(key, 0) + 1


def _record_marker_applied(docx_path: Path, table_index: int) -> None:
    stats = _ACTIVE_MARKER_STATS
    if stats is None:
        return
    stats.applied += 1
    stats.applied_table_indexes.append(table_index)

    caption = "-"
    try:
        doc = Document(str(docx_path))
        if 0 <= table_index < len(doc.tables):
            caption = _find_caption_number_before_table(doc, doc.tables[table_index]._tbl) or "-"
    except Exception:
        caption = "-"
    stats.applied_captions.append(caption)


def _format_marker_skip_reasons(stats: _MarkerSplitStats) -> str:
    if not stats.skipped_reasons:
        return "-"
    return ",".join(
        f"{reason}:{count}"
        for reason, count in sorted(stats.skipped_reasons.items())
    )


def _log_phase3_marker_summary(
    docx_path: Path,
    stats: _MarkerSplitStats,
    started_at: float,
) -> None:
    logger.info(
        "phase3_marker_summary request_id=%s renders=%s candidates=%s applied=%s skipped_reasons=%s elapsed_seconds=%.3f",
        _request_id_from_docx_path(docx_path),
        stats.renders,
        stats.candidates,
        stats.applied,
        _format_marker_skip_reasons(stats),
        time.monotonic() - started_at,
    )
    if stats.applied:
        logger.info(
            "phase3_marker_applied_summary request_id=%s table_indexes=%s captions=%s",
            _request_id_from_docx_path(docx_path),
            ",".join(str(index) for index in stats.applied_table_indexes) or "-",
            ",".join(stats.applied_captions) or "-",
        )


def _effective_marker_split_before_row(
    diagnostic,
    decision: _MarkerSplitDecision,
    *,
    header_rows: int = 1,
    docx_path=None,
) -> int | None:
    split_before_row = decision.split_before_row
    if split_before_row is None:
        return None

    if diagnostic.appendix_table and split_before_row > header_rows + 1:
        # Appendix continuations have no visible "Продолжение таблицы" marker.
        # Leave one fewer data row in the first fragment so Word does not push
        # a repeated textual header above the generated numbered row.
        return split_before_row - 1

    # E3: NUM-row compensation for ordinary body tables.
    # apply_numbered_split_to_document will insert a synthesized "1, 2, 3, ..."
    # row at index 1 of the first fragment if the original table doesn't
    # already have one. That extra row was NOT present when LO marker render
    # measured split_before_row, so post-split the fragment is one row taller
    # than measured — and TCF-A cantSplit can push the last data row alone
    # to a new page (orphan). Compensate by returning K - 1, but only when:
    #   (a) feature flag enabled
    #   (b) ordinary body table (appendix branch above handles its own case)
    #   (c) we have a docx_path to inspect the original table
    #   (d) original row 1 is NOT already an exact numbered row
    #   (e) K - 1 still leaves ≥ 1 data row in the first fragment
    #   (f) fragment 2 will have ≥ 2 data rows after the bump
    # Condition (f) avoids the Case-A worsening: when K is already at the
    # natural maximum (rows_count - 1), a -1 would push a row unnecessarily.
    if (
        not diagnostic.appendix_table
        and docx_path is not None
        and _marker_split_num_row_compensation_enabled()
        and split_before_row - 1 >= header_rows + 1
        and diagnostic.rows_count - split_before_row >= 2
    ):
        try:
            from .table_split_prototype import (
                _row_is_exact_numbered_row,
                _tbl_grid_column_count,
            )
            _doc = Document(str(docx_path))
            if 0 <= diagnostic.table_index < len(_doc.tables):
                tbl_xml = _doc.tables[diagnostic.table_index]._tbl
                cols = _tbl_grid_column_count(tbl_xml)
                trs = tbl_xml.findall(qn("w:tr"))
                if (
                    cols
                    and len(trs) > 1
                    and not _row_is_exact_numbered_row(trs[1], cols)
                ):
                    logger.info(
                        "marker_split_num_row_compensation table_index=%s K=%s -> %s rows_count=%s",
                        diagnostic.table_index,
                        split_before_row,
                        split_before_row - 1,
                        diagnostic.rows_count,
                    )
                    return split_before_row - 1
        except Exception as exc:
            logger.warning(
                "marker_split_num_row_compensation_failed table_index=%s error=%s",
                diagnostic.table_index, exc,
            )

    return split_before_row


def _ensure_blank_between_marker_and_second_table(doc, first_idx: int) -> bool:
    """Insert exactly one blank <w:p> between the continuation marker and the
    second table of a just-applied auto-split, if not already present.
    Idempotent. No-op if the chain shape is unexpected.
    """
    tables = doc.tables
    second_idx = first_idx + 1
    if first_idx < 0 or second_idx >= len(tables):
        return False
    tbl1 = tables[first_idx]._element
    tbl2 = tables[second_idx]._element
    body = tbl1.getparent()
    if body is None:
        return False
    children = list(body)
    try:
        i1 = children.index(tbl1)
        i2 = children.index(tbl2)
    except ValueError:
        return False
    # Expect tbl1 → marker_p → tbl2 (auto-split just inserted).
    if i2 != i1 + 2:
        # Either no marker between (rare appendix shape with no continuation_p)
        # or already shaped tbl1 → marker → blank → tbl2 → nothing to do.
        return False
    marker_node = children[i1 + 1]
    if marker_node.tag != qn("w:p"):
        return False
    blank = OxmlElement("w:p")
    _ensure_paragraph_bool_property_active(blank, "w:keepNext")
    marker_node.addnext(blank)
    return True


def _apply_marker_split_candidate(
    docx_path: Path,
    diagnostic,
    decision: _MarkerSplitDecision,
):
    doc = Document(str(docx_path))
    manual_skip = _valid_manual_continuation_table_indexes(doc)
    if diagnostic.table_index in manual_skip:
        return None, "valid_manual_continuation"
    if diagnostic.appendix_table and _is_generated_appendix_continuation_table(
        doc,
        diagnostic.table_index,
    ):
        return None, "generated_appendix_continuation"
    if not diagnostic.appendix_table and not diagnostic.has_standard_table_caption:
        return None, "ordinary_without_standard_caption"
    if (
        not diagnostic.appendix_table
        and _ordinary_table_has_nested_header_row(doc, diagnostic.table_index)
    ):
        return None, "body_contains_nested_table_header"

    split_before_row = _effective_marker_split_before_row(
        diagnostic, decision, docx_path=docx_path,
    )
    if split_before_row is None:
        return None, "no_boundary"
    if diagnostic.table_index < 0 or diagnostic.table_index >= len(doc.tables):
        return None, "invalid_table_index"
    if not _split_geometry_is_safe(
        doc.tables[diagnostic.table_index]._tbl,
        table_index=diagnostic.table_index,
        split_before_row=split_before_row,
        log_prefix="marker_split_skipped",
    ):
        return None, "unsafe_geometry"

    try:
        result = apply_numbered_split_to_document(
            doc,
            diagnostic.table_index,
            split_before_row,
            header_rows=1,
            numbered_header=True,
            appendix_table=diagnostic.appendix_table,
            continuation_paragraph_builder=_build_continuation_para,
        )
    except Exception as exc:
        return None, _map_marker_split_apply_error(exc)

    if result.source_note_after_second is False:
        return None, "source_note_ordering_failed"

    # Insert exactly one blank paragraph between the body-table continuation
    # marker "Продолжение таблицы X.Y.Z" and the second table fragment.
    # KFU norm: caption / EMPTY LINE / continued fragment.
    # Appendix continuation labels ("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ N") use a different
    # layout — leave them adjacent so existing tests + visual checks remain
    # correct. Idempotent — no-op when the blank is already there.
    if not diagnostic.appendix_table:
        _ensure_blank_between_marker_and_second_table(doc, diagnostic.table_index)

    _normalise_ordinary_continuation_anchors(doc)
    doc.save(str(docx_path))
    return result, None


def _run_marker_split_detection_pass(docx_path: Path, *, apply_split: bool = False) -> int:
    from . import table_markers

    def _format_row_pages(row_pages: dict[int, int]) -> str:
        if not row_pages:
            return "-"
        return ",".join(
            f"{row_index}:{page_num}"
            for row_index, page_num in sorted(row_pages.items())
        )

    def _format_duplicate_rows(duplicate_rows: dict[int, list[int]]) -> str:
        if not duplicate_rows:
            return "-"
        return ",".join(
            f"{row_index}:{'/'.join(str(page) for page in pages)}"
            for row_index, pages in sorted(duplicate_rows.items())
        )

    def _format_page_spans(page_spans) -> str:
        if not page_spans:
            return "-"
        return ",".join(
            f"{span.start_row}-{span.end_row}:{span.page_num}"
            for span in page_spans
        )

    eligible_count = 0
    eligible_apply_candidates = []
    try:
        diagnostics = table_markers.diagnose_all_tables(docx_path, keep_temp=False)
    except Exception as exc:
        logger.info("marker_split_skipped reason=mapping_error error=%s", exc)
        _record_marker_skip("mapping_error")
        return 0

    stats = _ACTIVE_MARKER_STATS
    if stats is not None:
        stats.renders += _count_marker_renders(diagnostics)
        stats.candidates += len(diagnostics)

    for diagnostic in diagnostics:
        logger.info(
            "marker_split_candidate table_index=%s rows=%s pages=%s row_pages=%s page_spans=%s missing_rows=%s duplicate_rows=%s",
            diagnostic.table_index,
            diagnostic.rows_count,
            diagnostic.pages_detected,
            _format_row_pages(diagnostic.row_pages),
            _format_page_spans(diagnostic.page_spans),
            diagnostic.missing_rows,
            _format_duplicate_rows(diagnostic.duplicate_rows),
        )
        duplicate_classification, duplicate_rows = _classify_marker_duplicate_rows(
            diagnostic,
            header_rows=1,
        )
        if duplicate_classification != "none":
            logger.info(
                "marker_split_duplicate_rows_classified table_index=%s classification=%s rows=%s missing_rows=%s duplicate_rows=%s page_spans=%s",
                diagnostic.table_index,
                duplicate_classification,
                duplicate_rows,
                diagnostic.missing_rows,
                _format_duplicate_rows(diagnostic.duplicate_rows),
                _format_page_spans(diagnostic.page_spans),
            )
        decision = _evaluate_marker_split_diagnostic(diagnostic, header_rows=1)
        if decision.eligible:
            if duplicate_classification == "header_only":
                logger.info(
                    "marker_split_header_duplicate_allowed table_index=%s missing_rows=%s duplicate_rows=%s page_spans=%s",
                    diagnostic.table_index,
                    diagnostic.missing_rows,
                    _format_duplicate_rows(diagnostic.duplicate_rows),
                    _format_page_spans(diagnostic.page_spans),
                )
            logger.info(
                "marker_split_boundary table_index=%s split_before_row=%s",
                diagnostic.table_index,
                decision.split_before_row,
            )
            logger.info(
                "marker_split_decision=ELIGIBLE table_index=%s",
                diagnostic.table_index,
            )
            eligible_count += 1
            if apply_split:
                eligible_apply_candidates.append((diagnostic, decision))
            continue

        logger.info(
            "marker_split_skipped table_index=%s reason=%s missing_rows=%s duplicate_rows=%s page_spans=%s",
            diagnostic.table_index,
            decision.skip_reason,
            diagnostic.missing_rows,
            _format_duplicate_rows(diagnostic.duplicate_rows),
            _format_page_spans(diagnostic.page_spans),
        )
        _record_marker_skip(decision.skip_reason)

    if not apply_split:
        return eligible_count

    applied_count = 0
    for diagnostic, decision in sorted(
        eligible_apply_candidates,
        key=lambda item: item[0].table_index,
        reverse=True,
    ):
        result, skip_reason = _apply_marker_split_candidate(
            docx_path,
            diagnostic,
            decision,
        )
        if result is not None:
            applied_count += 1
            _record_marker_applied(docx_path, diagnostic.table_index)
            logger.info(
                "marker_split_applied table_index=%s split_before_row=%s first_rows=%s second_rows=%s appendix=%s continuation=%s",
                diagnostic.table_index,
                result.split_before_row,
                result.first_table_rows_count,
                result.second_table_rows_count,
                diagnostic.appendix_table,
                result.continuation_paragraph_inserted,
            )
            continue
        logger.info(
            "marker_split_skipped table_index=%s reason=%s missing_rows=%s duplicate_rows=%s page_spans=%s",
            diagnostic.table_index,
            skip_reason,
            diagnostic.missing_rows,
            _format_duplicate_rows(diagnostic.duplicate_rows),
            _format_page_spans(diagnostic.page_spans),
        )
        _record_marker_skip(skip_reason)

    return applied_count


def _run_marker_split_for_candidates(
    docx_path: Path,
    candidates: list[int],
    overflow: list[int],
    hard_timeout_seconds: float,
    apply_split: bool,
    report,
) -> int:
    """Quality-first per-candidate marker-split pass (E2).

    Diagnoses each candidate table individually (one LibreOffice render per
    call). Per-candidate try/except so a single failure does not abort the
    rest. Wall-time check before starting each diagnose. Overflow indices
    (filtered out by candidate budget) are logged individually.

    Emits a user warning only when at least one *candidate* was skipped due
    to budget / wall-time / diagnose error — never just because the document
    has many tables.
    """
    from . import table_markers

    stats = _ACTIVE_MARKER_STATS
    started_at = time.monotonic()

    eligible_apply_candidates: list = []
    diagnose_failed: list[int] = []
    timed_out: list[int] = []

    for ci, idx in enumerate(candidates):
        elapsed = time.monotonic() - started_at
        if elapsed > hard_timeout_seconds:
            for r in candidates[ci:]:
                logger.warning(
                    "marker_split_skipped table_index=%s reason=hard_timeout_exceeded elapsed=%.1fs limit=%.1fs",
                    r, elapsed, hard_timeout_seconds,
                )
                _record_marker_skip("hard_timeout_exceeded")
                timed_out.append(r)
            break

        try:
            diagnostic = table_markers.diagnose_table(docx_path, idx, keep_temp=False)
        except Exception as exc:
            logger.warning(
                "marker_split_skipped table_index=%s reason=diagnose_error error=%s",
                idx, exc,
            )
            _record_marker_skip("diagnose_error")
            diagnose_failed.append(idx)
            continue

        if stats is not None:
            stats.renders += _count_marker_renders([diagnostic])
            stats.candidates += 1

        logger.info(
            "marker_split_candidate table_index=%s rows=%s pages=%s",
            diagnostic.table_index,
            diagnostic.rows_count,
            diagnostic.pages_detected,
        )

        decision = _evaluate_marker_split_diagnostic(diagnostic, header_rows=1)
        if decision.eligible:
            logger.info(
                "marker_split_decision=ELIGIBLE table_index=%s split_before_row=%s",
                diagnostic.table_index, decision.split_before_row,
            )
            if apply_split:
                eligible_apply_candidates.append((diagnostic, decision))
        else:
            logger.info(
                "marker_split_skipped table_index=%s reason=%s",
                diagnostic.table_index, decision.skip_reason,
            )
            _record_marker_skip(decision.skip_reason)

    for idx in overflow:
        logger.warning(
            "marker_split_skipped table_index=%s reason=candidate_budget_exhausted",
            idx,
        )
        _record_marker_skip("candidate_budget_exhausted")

    applied_count = 0
    if apply_split:
        for diagnostic, decision in sorted(
            eligible_apply_candidates,
            key=lambda item: item[0].table_index,
            reverse=True,
        ):
            result, skip_reason = _apply_marker_split_candidate(
                docx_path, diagnostic, decision,
            )
            if result is not None:
                applied_count += 1
                _record_marker_applied(docx_path, diagnostic.table_index)
                logger.info(
                    "marker_split_applied table_index=%s split_before_row=%s appendix=%s",
                    diagnostic.table_index,
                    result.split_before_row,
                    diagnostic.appendix_table,
                )
                continue
            logger.info(
                "marker_split_skipped table_index=%s reason=%s",
                diagnostic.table_index, skip_reason or "unknown",
            )
            _record_marker_skip(skip_reason or "apply_failed")

    real_skips = list(overflow) + list(timed_out) + list(diagnose_failed)
    if report is not None and real_skips:
        total_attempted = len(candidates) + len(overflow)
        processed = total_attempted - len(real_skips)
        report.warn(
            "Автоматическое разделение длинных таблиц выполнено частично: "
            f"обработано {processed} из {total_attempted} длинных таблиц. "
            "Остальные проверьте вручную."
        )

    return applied_count


# ── P1-c / DEFECT B — fallback split-last-row for detached source/note ───────

@dataclass(frozen=True)
class _SourceNoteDetachmentCandidate:
    table_index: int
    caption_num: str
    last_data_row_page: int
    source_note_first_page: int
    rows_count: int


def _paragraph_text_for_match(p_xml) -> str:
    return _norm_match_text(
        "".join(t.text or "" for t in p_xml.findall(".//" + qn("w:t")))
    )


def _find_paragraph_pdf_page(p_xml, pdf_lines: list[PdfLine]) -> int | None:
    """Best-effort page lookup: find which rendered PDF page contains the
    paragraph's text. Uses substring match on normalized text. Returns None
    if no reliable match (text too short, image-only, or absent)."""
    norm = _paragraph_text_for_match(p_xml)
    if not norm or len(norm) < 8:
        return None
    for line in pdf_lines:
        if norm in _norm_match_text(line.text):
            return line.page_num
    # Fallback: try first 30 chars (tolerate line wrapping / truncation).
    head = norm[:30]
    if len(head) >= 8:
        for line in pdf_lines:
            ln = _norm_match_text(line.text)
            if head in ln or ln.startswith(head):
                return line.page_num
    return None


def _detect_table_following_source_note_paragraphs(
    table_xml, para_by_xml: dict, body_children: list,
) -> list:
    """Return ordered list of <w:p> elements immediately following `table_xml`
    that constitute one or more source/note paragraphs (Источник:/Примечание:
    /Составлено по/Рассчитано по). At most one blank line is allowed between
    the table and the first source/note paragraph — matches the structural
    pattern recognized by `_apply_rule_source_note`."""
    try:
        tbl_idx = body_children.index(table_xml)
    except ValueError:
        return []
    result: list = []
    j = tbl_idx + 1
    seen_nonempty = False
    while j < len(body_children):
        el = body_children[j]
        if el.tag != qn("w:p"):
            break
        p_obj = para_by_xml.get(el)
        if p_obj is None:
            break
        text = (p_obj.text or "").strip()
        if not text:
            if seen_nonempty:
                break
            j += 1
            continue
        if is_source_or_note_line(text):
            result.append(el)
            seen_nonempty = True
            j += 1
            continue
        break
    return result


def _collect_source_note_detachment_candidates(
    doc: Document,
    pdf_lines: list[PdfLine],
    manual_chain_ids: set[int],
) -> tuple[list[_SourceNoteDetachmentCandidate], list[tuple[int, str]]]:
    """Detect tables whose immediately-following source/note paragraphs render
    on a strictly later PDF page than the table's last data row. Pure logic:
    no document mutation, no rendering. Returns (candidates, skips) where
    `skips` is a list of (table_index, skip_reason) for telemetry."""
    body_children = list(doc.element.body)
    para_by_xml = {p._element: p for p in doc.paragraphs}
    signatures = _collect_table_signatures(doc)
    candidates: list[_SourceNoteDetachmentCandidate] = []
    skips: list[tuple[int, str]] = []

    for table_index, table in enumerate(doc.tables):
        tbl_xml = table._tbl

        if table_index in manual_chain_ids or id(tbl_xml) in manual_chain_ids:
            skips.append((table_index, "already_in_manual_chain"))
            continue

        sn_paras = _detect_table_following_source_note_paragraphs(
            tbl_xml, para_by_xml, body_children,
        )
        if not sn_paras:
            skips.append((table_index, "no_source_note"))
            continue

        caption_num = _find_caption_number_before_table(doc, tbl_xml)
        if not caption_num:
            skips.append((table_index, "no_caption"))
            continue

        rows_xml = tbl_xml.findall(qn("w:tr"))
        # Need at least 3 rows: 1 header + 2 data rows. After split-last-row,
        # fragment-1 has header + ≥1 data row, fragment-2 has header + 1 data.
        if len(rows_xml) < 3:
            skips.append((table_index, "no_safe_data_row"))
            continue

        if _ordinary_table_has_nested_header_row(doc, table_index):
            skips.append((table_index, "body_contains_nested_table_header"))
            continue

        # Render probe: rows → pages.
        sig = signatures[table_index] if table_index < len(signatures) else None
        if sig is None:
            skips.append((table_index, "render_probe_unreliable"))
            continue
        row_pages = _match_row_pages(sig, pdf_lines)
        if not row_pages:
            skips.append((table_index, "render_probe_unreliable"))
            continue
        data_row_pages = [p for idx, p in row_pages.items() if idx > 0]
        if not data_row_pages:
            skips.append((table_index, "render_probe_unreliable"))
            continue
        last_data_row_page = max(data_row_pages)

        sn_pages = []
        for sn_xml in sn_paras:
            pg = _find_paragraph_pdf_page(sn_xml, pdf_lines)
            if pg is not None:
                sn_pages.append(pg)
        if not sn_pages:
            skips.append((table_index, "render_probe_unreliable"))
            continue
        sn_first_page = min(sn_pages)

        if sn_first_page <= last_data_row_page:
            skips.append((table_index, "not_detached"))
            continue

        candidates.append(
            _SourceNoteDetachmentCandidate(
                table_index=table_index,
                caption_num=caption_num,
                last_data_row_page=last_data_row_page,
                source_note_first_page=sn_first_page,
                rows_count=len(rows_xml),
            )
        )

    return candidates, skips


def _apply_source_note_detachment_split(
    docx_path: Path,
    report: FormattingReport | None = None,
) -> int:
    """P1-c: split off the last data row into a continuation fragment for
    every table whose Источник:/Примечание: paragraphs would otherwise land
    on a different page (detached). Reuses the existing split engine
    (`apply_numbered_split_to_document`) with `_build_continuation_para`,
    so the new fragment chain matches the formatter-authored continuation
    invariants (right-aligned marker, pageBreakBefore, keepNext, numbered row
    repeated).

    Per-table try/except: a single failure does not abort the document.
    Idempotent: tables already in a manual continuation chain are skipped via
    `_valid_manual_continuation_table_ids`, so a second run is a no-op.
    """
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))
    if not doc.tables:
        return 0

    pdf_path: Path | None = None
    try:
        pdf_path = render_docx_to_pdf(docx_path)
        pdf_lines = analyze_pdf_lines(pdf_path)
    except LibreOfficeNotFoundError as exc:
        logger.info(
            "p1c_source_note_split_skip reason=libreoffice_unavailable error=%s", exc,
        )
        return 0
    except Exception as exc:
        logger.info(
            "p1c_source_note_split_skip reason=render_failed error=%s", exc,
        )
        return 0
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)

    manual_chain_ids = _valid_manual_continuation_table_indexes(doc)
    candidates, skips = _collect_source_note_detachment_candidates(
        doc, pdf_lines, manual_chain_ids,
    )
    for table_index, reason in skips:
        logger.info(
            "p1c_source_note_split_skip table_index=%s reason=%s",
            table_index, reason,
        )

    if not candidates:
        return 0

    applied = 0
    # Iterate from highest table_index to lowest so each split inserts a new
    # tbl at table_index+1 without shifting any not-yet-processed candidate.
    for cand in sorted(candidates, key=lambda c: -c.table_index):
        split_before_row = cand.rows_count - 1  # split off last data row
        if cand.table_index < 0 or cand.table_index >= len(doc.tables):
            logger.info(
                "p1c_source_note_split_skip table_index=%s reason=invalid_table_index",
                cand.table_index,
            )
            continue
        if not _split_geometry_is_safe(
            doc.tables[cand.table_index]._tbl,
            table_index=cand.table_index,
            split_before_row=split_before_row,
            log_prefix="p1c_source_note_split_skip",
        ):
            continue
        try:
            result = apply_numbered_split_to_document(
                doc,
                cand.table_index,
                split_before_row,
                header_rows=1,
                numbered_header=True,
                appendix_table=False,
                continuation_paragraph_builder=_build_continuation_para,
            )
        except Exception as exc:
            logger.info(
                "p1c_source_note_split_skip table_index=%s reason=apply_split_failed error=%s",
                cand.table_index, exc,
            )
            continue

        if result.source_note_after_second is False:
            logger.warning(
                "p1c_source_note_split_warn table_index=%s caption=%s source_note_after_second=False",
                cand.table_index, cand.caption_num,
            )

        _ensure_blank_between_marker_and_second_table(doc, cand.table_index)
        applied += 1
        logger.info(
            "p1c_source_note_split_applied table_index=%s caption=%s split_before_row=%s last_row_page=%s sn_page=%s",
            cand.table_index, cand.caption_num, split_before_row,
            cand.last_data_row_page, cand.source_note_first_page,
        )

    if applied:
        _normalise_ordinary_continuation_anchors(doc)
        doc.save(str(docx_path))
        logger.info("p1c_source_note_split_summary applied=%s", applied)
    return applied


# ── P2-a — appendix continuation for table-based appendices ──────────────────

@dataclass(frozen=True)
class _AppendixTableContinuationCandidate:
    table_index: int
    appendix_num: str
    first_data_row_page: int
    last_data_row_page: int
    rows_count: int
    split_before_row: int


def _next_nonempty_body_paragraph_after(
    body_children: list, table_xml, para_by_xml: dict,
):
    """Walk forward from `table_xml`, skip at most one empty paragraph, and
    return the next non-empty `<w:p>` element (or None). Used to detect a
    pre-existing «ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ N» marker directly after the table."""
    try:
        tbl_idx = body_children.index(table_xml)
    except ValueError:
        return None
    j = tbl_idx + 1
    seen_blank = False
    while j < len(body_children):
        el = body_children[j]
        if el.tag != qn("w:p"):
            return None
        p_obj = para_by_xml.get(el)
        text = (p_obj.text or "").strip() if p_obj is not None else ""
        if not text:
            if seen_blank:
                return None
            seen_blank = True
            j += 1
            continue
        return el
    return None


def _collect_appendix_table_continuation_candidates(
    doc: Document,
    pdf_lines: list[PdfLine],
    manual_chain_ids: set[int],
) -> tuple[list[_AppendixTableContinuationCandidate], list[tuple[int, str]]]:
    """Detect appendix tables whose data rows render across > 1 PDF page.
    Pure logic — no document mutation, no rendering. Returns
    (candidates, skips) where `skips` is a list of (table_index, skip_reason).

    Eligibility (all must hold):
    - Table is flagged `appendix_table=True` by
      `_iter_body_tables_with_context` (i.e. preceded by `ПРИЛОЖЕНИЕ …`
      heading somewhere upstream).
    - `_find_preceding_appendix_number` returns a non-None number (required
      by `apply_numbered_split_to_document` with `appendix_table=True`).
    - Table has ≥ 3 rows (1 header + ≥ 2 data rows; split-last-row leaves
      ≥ 1 data row on each fragment).
    - Header row is not a nested-table header.
    - Table is NOT in a manual continuation chain protected by P0-α.
    - Table is NOT itself a generated appendix continuation fragment
      (`_is_generated_appendix_continuation_table`).
    - Table is NOT already followed by a «ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ N» marker
      (idempotency: skip if some prior pass / source authoring already
      inserted one).
    - Render probe maps rows to pages reliably AND the data rows actually
      span ≥ 2 pages.
    """
    from .table_markers import _iter_body_tables_with_context
    contexts = _iter_body_tables_with_context(doc)
    signatures = _collect_table_signatures(doc)
    body_children = list(doc.element.body)
    para_by_xml = {p._element: p for p in doc.paragraphs}
    candidates: list[_AppendixTableContinuationCandidate] = []
    skips: list[tuple[int, str]] = []

    for table_index, ctx in enumerate(contexts):
        table_obj = ctx.get("table_obj")
        if table_obj is None:
            skips.append((table_index, "table_not_in_body"))
            continue
        tbl_xml = table_obj._tbl

        if not ctx.get("appendix_table"):
            skips.append((table_index, "not_appendix"))
            continue

        if table_index in manual_chain_ids or id(tbl_xml) in manual_chain_ids:
            skips.append((table_index, "already_in_manual_chain"))
            continue

        if _is_generated_appendix_continuation_table(doc, table_index):
            skips.append((table_index, "generated_appendix_continuation"))
            continue

        next_para = _next_nonempty_body_paragraph_after(
            body_children, tbl_xml, para_by_xml,
        )
        if next_para is not None and _is_appendix_continuation_paragraph(next_para):
            skips.append((table_index, "already_followed_by_continuation"))
            continue

        try:
            tbl_body_idx = body_children.index(tbl_xml)
        except ValueError:
            skips.append((table_index, "table_not_in_body"))
            continue
        appendix_num = _find_preceding_appendix_number(body_children, tbl_body_idx)
        if appendix_num is None:
            skips.append((table_index, "no_appendix_number"))
            continue

        rows_xml = tbl_xml.findall(qn("w:tr"))
        if len(rows_xml) < 3:
            skips.append((table_index, "no_safe_data_row"))
            continue

        if _ordinary_table_has_nested_header_row(doc, table_index):
            skips.append((table_index, "body_contains_nested_table_header"))
            continue

        sig = signatures[table_index] if table_index < len(signatures) else None
        if sig is None:
            skips.append((table_index, "render_probe_unreliable"))
            continue
        row_pages = _match_row_pages(sig, pdf_lines)
        if not row_pages:
            # P2-a' relaxed fallback — appendix tables only. Tolerates
            # duplicate row signatures and wrapped cell text via window
            # aggregation. Strict matcher remains untouched.
            row_pages = _match_row_pages_relaxed_for_appendix(sig, pdf_lines)
        if not row_pages:
            skips.append((table_index, "render_probe_unreliable"))
            continue
        data_row_pages = {idx: p for idx, p in row_pages.items() if idx > 0}
        if not data_row_pages:
            skips.append((table_index, "render_probe_unreliable"))
            continue
        first_data_row_page = min(data_row_pages.values())
        last_data_row_page = max(data_row_pages.values())
        if last_data_row_page <= first_data_row_page:
            skips.append((table_index, "single_page"))
            continue

        # split_before_row = first row index whose page is strictly later
        # than the first data row page.
        split_before_row = None
        for idx, page in sorted(data_row_pages.items()):
            if page > first_data_row_page:
                split_before_row = idx
                break
        if split_before_row is None or split_before_row <= 1:
            skips.append((table_index, "no_safe_split_boundary"))
            continue

        candidates.append(
            _AppendixTableContinuationCandidate(
                table_index=table_index,
                appendix_num=appendix_num,
                first_data_row_page=first_data_row_page,
                last_data_row_page=last_data_row_page,
                rows_count=len(rows_xml),
                split_before_row=split_before_row,
            )
        )

    return candidates, skips


def _apply_appendix_table_continuation_split(
    docx_path: Path,
    report: FormattingReport | None = None,
) -> int:
    """P2-a: insert «ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ N» on the continuation page of
    every appendix table whose data rows span > 1 PDF page. Reuses the
    existing split engine (`apply_numbered_split_to_document` with
    `appendix_table=True`) and the existing
    `_build_appendix_continuation_paragraph` marker builder.

    Per-table try/except: a single failure does not abort the document.
    Idempotent: tables already followed by a continuation marker, generated
    appendix continuation fragments, and manually-chained tables are
    skipped — so a second run is a no-op.
    """
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))
    if not doc.tables:
        return 0

    pdf_path: Path | None = None
    try:
        pdf_path = render_docx_to_pdf(docx_path)
        pdf_lines = analyze_pdf_lines(pdf_path)
    except LibreOfficeNotFoundError as exc:
        logger.info(
            "p2a_appendix_continuation_skip reason=libreoffice_unavailable error=%s", exc,
        )
        return 0
    except Exception as exc:
        logger.info(
            "p2a_appendix_continuation_skip reason=render_failed error=%s", exc,
        )
        return 0
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)

    manual_chain_ids = _valid_manual_continuation_table_indexes(doc)
    candidates, skips = _collect_appendix_table_continuation_candidates(
        doc, pdf_lines, manual_chain_ids,
    )
    for table_index, reason in skips:
        logger.info(
            "p2a_appendix_continuation_skip table_index=%s reason=%s",
            table_index, reason,
        )

    if not candidates:
        return 0

    applied = 0
    # Reverse table_index order — each split inserts a new <w:tbl> at
    # `table_index + 1` and would shift any not-yet-processed candidate.
    for cand in sorted(candidates, key=lambda c: -c.table_index):
        if cand.table_index < 0 or cand.table_index >= len(doc.tables):
            logger.info(
                "p2a_appendix_continuation_skip table_index=%s reason=invalid_table_index",
                cand.table_index,
            )
            continue
        if not _split_geometry_is_safe(
            doc.tables[cand.table_index]._tbl,
            table_index=cand.table_index,
            split_before_row=cand.split_before_row,
            log_prefix="p2a_appendix_continuation_skip",
        ):
            continue
        try:
            apply_numbered_split_to_document(
                doc,
                cand.table_index,
                cand.split_before_row,
                header_rows=1,
                numbered_header=True,
                appendix_table=True,
                continuation_paragraph_builder=_build_appendix_continuation_paragraph,
            )
        except Exception as exc:
            logger.info(
                "p2a_appendix_continuation_skip table_index=%s reason=apply_split_failed error=%s",
                cand.table_index, exc,
            )
            continue
        applied += 1
        logger.info(
            "p2a_appendix_continuation_applied table_index=%s appendix_num=%s split_before_row=%s first_row_page=%s last_row_page=%s",
            cand.table_index, cand.appendix_num, cand.split_before_row,
            cand.first_data_row_page, cand.last_data_row_page,
        )

    if applied:
        doc.save(str(docx_path))
        logger.info("p2a_appendix_continuation_summary applied=%s", applied)
    return applied


def apply_rendered_table_continuation(
    docx_path: Path,
    report: FormattingReport | None = None,
    max_passes: int = 1,
) -> int:
    docx_path = Path(docx_path)
    backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_rendered_table_gate_"))
    backup_path = backup_dir / "pre_rendered_table_continuation.docx"
    try:
        shutil.copy2(docx_path, backup_path)
        applied = _apply_rendered_table_continuation_impl(
            docx_path,
            report=report,
            max_passes=max_passes,
        )
        if applied and restore_docx_if_same_page_continuation_markers(
            docx_path,
            backup_path,
            report=report,
            context="apply_rendered_table_continuation",
        ):
            logger.info(
                "rendered_final_decision action=rendered_no_action reason=post_render_same_page_marker_violation"
            )
            return 0
        return applied
    finally:
        shutil.rmtree(backup_dir, ignore_errors=True)


def _apply_rendered_table_continuation_impl(
    docx_path: Path,
    report: FormattingReport | None = None,
    max_passes: int = 1,
) -> int:
    """
    Phase 3 rendered table continuation entry point.

    Patch 1 only wires LibreOffice/PDF availability checks and disables the
    previous heuristic splitter. Actual rendered row-to-page splitting is added
    in a later patch.
    """
    docx_path = Path(docx_path)
    marker_stats = _MarkerSplitStats()
    marker_started_at = time.monotonic()
    marker_summary_logged = False

    def emit_marker_summary() -> None:
        nonlocal marker_summary_logged
        if marker_summary_logged:
            return
        marker_summary_logged = True
        _log_phase3_marker_summary(docx_path, marker_stats, marker_started_at)

    doc = Document(str(docx_path))
    if not doc.tables:
        emit_marker_summary()
        logger.info("rendered_table_continuation_enter tables=0 pdf_lines=0 max_passes=%s", max_passes)
        logger.info("rendered_final_decision action=rendered_no_action reason=no_tables")
        return 0

    logger.info(
        "rendered_table_continuation_start tables=%s max_passes=%s",
        len(doc.tables),
        max_passes,
    )

    # E1: log future-E2 candidate classification BEFORE any existing skip/run logic.
    # Strictly observe-only — no behavior change.
    if _marker_split_enabled():
        try:
            _classification = _classify_marker_split_candidates(doc)
            logger.info(
                "phase3_candidate_classification total=%s manual_chain=%s no_caption=%s tiny=%s candidates=%s budget=%s would_process=%s would_skip_for_budget=%s",
                _classification["total_tables"],
                len(_classification["manual_continuation_skipped"]),
                len(_classification["no_caption_skipped"]),
                len(_classification["tiny_table_skipped"]),
                len(_classification["candidate_tables"]),
                _classification["candidate_budget"],
                _classification["would_process"],
                _classification["would_skip_for_budget"],
            )
            try:
                from . import table_markers as _tm_e1
                _ctxs = _tm_e1._iter_body_tables_with_context(doc)
            except Exception:
                _ctxs = []
            for _idx, _priority in _classification["candidate_priority"]:
                _rows = len(doc.tables[_idx].rows) if _idx < len(doc.tables) else 0
                _appendix = bool(_ctxs[_idx].get("appendix_table")) if _idx < len(_ctxs) else False
                logger.info(
                    "marker_split_candidate_preview table_index=%s priority=%s rows=%s appendix=%s",
                    _idx, _priority, _rows, _appendix,
                )
        except Exception as _exc:
            logger.warning("phase3_candidate_classification_failed error=%s", _exc)

    # E2: branch on KPFU_MARKER_SPLIT_MODE.
    # candidate (default) — per-candidate diagnose, hard wall-time cap, warning
    #                       fires only if real candidates were skipped.
    # global_skip (legacy) — pre-E2 behaviour: skip wholesale when total > budget.
    _split_mode = _marker_split_mode()

    if _marker_split_enabled() and _split_mode == "global_skip":
        budget = _marker_split_max_renders()
        table_count = len(doc.tables)
        if table_count > budget:
            emit_marker_summary()
            logger.warning(
                "marker_split_skipped reason=render_budget_exceeded count=%s budget=%s",
                table_count,
                budget,
            )
            if report is not None:
                report.warn(
                    "Автоматическое разделение длинных таблиц пропущено: "
                    "в документе много таблиц. Проверьте переносы таблиц вручную."
                )
            logger.info(
                "rendered_final_decision action=rendered_no_action reason=render_budget_exceeded"
            )
            return 0

    if _marker_split_enabled():
        apply_marker_split = _marker_split_apply_enabled()
        global _ACTIVE_MARKER_STATS
        previous_marker_stats = _ACTIVE_MARKER_STATS
        _ACTIVE_MARKER_STATS = marker_stats
        try:
            try:
                if _split_mode == "candidate":
                    marker_total = _run_marker_split_for_candidates(
                        docx_path,
                        candidates=_classification["would_process"],
                        overflow=_classification["would_skip_for_budget"],
                        hard_timeout_seconds=_marker_split_hard_timeout_seconds(),
                        apply_split=apply_marker_split,
                        report=report,
                    )
                elif apply_marker_split:
                    marker_total = _run_marker_split_detection_pass(
                        docx_path,
                        apply_split=True,
                    )
                else:
                    _run_marker_split_detection_pass(
                        docx_path,
                        apply_split=False,
                    )
                    marker_total = 0
                if marker_total and apply_marker_split:
                    emit_marker_summary()
                    logger.info(
                        "rendered_final_decision action=marker_split_applied count=%s",
                        marker_total,
                    )
                    # P1-c: scan remaining tables for detached source/note even
                    # after marker-split succeeded on some tables. Manual chains
                    # produced by marker-split are protected (see
                    # `_valid_manual_continuation_table_ids`), so we will not
                    # re-split anything marker-split already touched.
                    p1c_total = _apply_source_note_detachment_split(
                        docx_path, report=report,
                    )
                    # P2-a: appendix continuation for table-based appendices.
                    # Runs AFTER P1-c so any P1-c-created chains are already
                    # protected by `_valid_manual_continuation_table_ids` or
                    # the `already_followed_by_continuation` skip.
                    p2a_total = _apply_appendix_table_continuation_split(
                        docx_path, report=report,
                    )
                    return marker_total + p1c_total + p2a_total
            except Exception:
                emit_marker_summary()
                raise
        finally:
            _ACTIVE_MARKER_STATS = previous_marker_stats

    emit_marker_summary()

    # P1-c: detached source/note pass — runs even when marker-split did not
    # apply (table fits on one page geometrically but Word/LO cannot keep
    # source/note attached due to last-row overflow). Returns early if any
    # splits were applied so the legacy rendered-split fallback below does
    # not double-act on the now-restructured document.
    p1c_total_pre_legacy = _apply_source_note_detachment_split(
        docx_path, report=report,
    )
    # P2-a: appendix continuation for table-based appendices. Runs after P1-c
    # so its detection sees the doc shape that P1-c left.
    p2a_total_pre_legacy = _apply_appendix_table_continuation_split(
        docx_path, report=report,
    )
    combined_pre_legacy = p1c_total_pre_legacy + p2a_total_pre_legacy
    if combined_pre_legacy > 0:
        logger.info(
            "rendered_final_decision action=p1c_p2a_split p1c=%s p2a=%s",
            p1c_total_pre_legacy, p2a_total_pre_legacy,
        )
        return combined_pre_legacy

    pdf_path: Path | None = None
    try:
        pdf_path = render_docx_to_pdf(docx_path)
        pdf_lines = analyze_pdf_lines(pdf_path)
    except LibreOfficeNotFoundError as exc:
        _warn_rendered_split_unavailable(report, str(exc))
        logger.info("rendered_final_decision action=rendered_no_action reason=libreoffice_unavailable")
        return 0
    except Exception as exc:
        _warn_rendered_split_unavailable(report, str(exc))
        logger.info("rendered_final_decision action=rendered_no_action reason=render_or_pdf_analysis_failed")
        return 0
    finally:
        if pdf_path is not None:
            shutil.rmtree(pdf_path.parent, ignore_errors=True)

    logger.info(
        "rendered_table_continuation_enter tables=%s pdf_lines=%s max_passes=%s",
        len(doc.tables),
        len(pdf_lines),
        max_passes,
    )

    diagnostics: dict[str, bool] = {"ambiguous": False}
    move_candidate = _find_rendered_whole_table_move_candidate(doc, pdf_lines, diagnostics)
    if move_candidate is not None:
        validation_backup_dir = Path(tempfile.mkdtemp(prefix="kpfu_table_start_orphan_"))
        validation_backup_path = validation_backup_dir / docx_path.name
        shutil.copy2(docx_path, validation_backup_path)

        if not _insert_table_start_orphan_blanks(move_candidate.caption_para_xml, target_count=2):
            shutil.rmtree(validation_backup_dir, ignore_errors=True)
            logger.info(
                "rendered_final_decision action=rendered_no_action reason=table_start_orphan_already_has_blanks table_idx=%s",
                move_candidate.table_idx,
            )
            return 0
        doc.save(str(docx_path))
        if _same_table_start_orphan_remains(docx_path, move_candidate.table_idx):
            shutil.copy2(validation_backup_path, docx_path)
            shutil.rmtree(validation_backup_dir, ignore_errors=True)
            logger.info(
                "rendered_final_decision action=rendered_no_action reason=table_start_orphan_validation_failed table_idx=%s",
                move_candidate.table_idx,
            )
            return 0
        shutil.rmtree(validation_backup_dir, ignore_errors=True)
        logger.info(
            "rendered_final_decision action=table_start_orphan_move table_idx=%s blanks=2",
            move_candidate.table_idx,
        )
        return 1

    candidate = _find_rendered_split_candidate(doc, pdf_lines, diagnostics)
    if candidate is None:
        if diagnostics["ambiguous"]:
            logger.info("rendered_final_decision action=rendered_skip_ambiguous reason=no_safe_rendered_candidate")
        else:
            logger.info("rendered_final_decision action=rendered_no_action reason=no_rendered_candidate")
        return 0

    num = _find_caption_number_before_table(doc, candidate.tbl_xml)
    continuation_text = f"Продолжение таблицы {num}" if num else "Продолжение таблицы"
    if not _split_geometry_is_safe(
        candidate.tbl_xml,
        table_index=candidate.table_idx,
        split_before_row=candidate.split_after + 1,
        log_prefix="rendered_final_decision action=rendered_no_action",
    ):
        return 0
    if not _split_table_at(doc, candidate.tbl_xml, candidate.split_after, continuation_text):
        logger.info(
            "rendered_final_decision action=rendered_no_action reason=split_mutation_failed table_idx=%s split_after=%s",
            candidate.table_idx,
            candidate.split_after,
        )
        return 0

    _normalise_ordinary_continuation_anchors(doc)
    doc.save(str(docx_path))
    logger.info(
        "rendered_final_decision action=rendered_split table_idx=%s split_after=%s",
        candidate.table_idx,
        candidate.split_after,
    )
    return 1


# ── Remove empty paragraphs between image and figure caption ─────────────────

def remove_empty_before_figure_captions(doc: Document) -> int:
    """
    Remove empty paragraphs that appear immediately between an image paragraph
    and a figure caption ("Рис. X.Y.Z — ...").

    Students often insert a blank line between the figure and its caption.
    This leaves a visual gap in the formatted output.  We remove such blanks
    only when the paragraph immediately before the empty run contains a drawing.

    Returns the number of paragraphs removed.
    """
    paragraphs = doc.paragraphs
    n = len(paragraphs)
    to_remove: list = []

    i = 0
    while i < n:
        text = (paragraphs[i].text or "").strip()
        if _FIGURE_CAP_RE_GEOM.match(text):
            # Collect preceding empty paragraphs
            j = i - 1
            empty_elems: list = []
            while j >= 0:
                prev_text = (paragraphs[j].text or "").strip()
                if not prev_text and not _para_has_image(paragraphs[j]._element):
                    empty_elems.append(paragraphs[j]._element)
                    j -= 1
                else:
                    break
            # Only remove if the paragraph immediately before the run is an image
            if empty_elems and j >= 0 and _para_has_image(paragraphs[j]._element):
                to_remove.extend(empty_elems)
        i += 1

    removed = 0
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            removed += 1

    if removed:
        logger.info("remove_empty_before_figure_captions: removed %d gap paragraph(s)", removed)
    return removed


# Only trust a w:lastRenderedPageBreak calibration signal when we have already
# accumulated at least this fraction of the page.  LRPB markers that fire at
# very low cumulative heights are stale artefacts from the ORIGINAL layout that
# no longer reflect page boundaries in the MODIFIED document (e.g. a paragraph
# that was the first on a page in the source but is now mid-page after a table
# split was inserted above it).
_LRPB_TRUST_RATIO = 0.25   # 25 % of body height ≈ ~178 pt for a KFU page


def _lrpb_calibrate(xml_elem, current_h: float, body_h: float) -> float:
    """
    Return the new current_h after applying an optional LRPB calibration.

    Resets to 0.0 only when:
      1. The paragraph contains a w:lastRenderedPageBreak, AND
      2. current_h >= body_h * _LRPB_TRUST_RATIO
         (enough content has been seen that the LRPB is likely genuine).
    """
    if _para_has_lrpb(xml_elem) and current_h >= body_h * _LRPB_TRUST_RATIO:
        return 0.0
    return current_h


# ── Helpers for geometry-based page-break rules ───────────────────────────────

_TABLE_CAP_RE_GEOM = re.compile(
    r"^\s*(таблица|table)\s+\d+(?:\.\d+){0,2}",
    re.IGNORECASE,
)
_FIGURE_CAP_RE_GEOM = re.compile(
    r"^\s*(рис\.|рисунок|figure|fig\.)\s*\d+",
    re.IGNORECASE,
)


def _para_has_image(p_elem) -> bool:
    """True if the paragraph XML element contains an inline drawing or picture."""
    return xml_has_image(p_elem)


def _get_image_height_pt(p_elem) -> float | None:
    """
    Return the rendered height (pt) of the first drawing in a paragraph by
    reading the wp:extent cy attribute (in EMU).

    Word stores drawing dimensions in EMU (English Metric Units):
        1 pt = 12 700 EMU  (EMU_PER_PT constant)

    Returns None if no wp:extent element is found.
    """
    for drawing in p_elem.findall(".//" + qn("w:drawing")):
        for container_tag in (qn("wp:inline"), qn("wp:anchor")):
            container = drawing.find(container_tag)
            if container is not None:
                extent = container.find(qn("wp:extent"))
                if extent is not None:
                    cy = extent.get("cy")
                    if cy and cy.lstrip("-").isdigit():
                        return int(cy) / EMU_PER_PT
    return None


def _set_page_break_before(para_elem) -> None:
    """Add w:pageBreakBefore to a paragraph's pPr (idempotent)."""
    pPr = para_elem.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        para_elem.insert(0, pPr)
    if pPr.find(qn("w:pageBreakBefore")) is None:
        pb = OxmlElement("w:pageBreakBefore")
        pPr.append(pb)


# ── Rule 4: no empty first line of page ──────────────────────────────────────

def _apply_rule4_pass(doc: Document) -> int:
    """
    Single pass of Rule 4 — remove empty paragraphs at the very top of a page.

    Conservative: only removes paragraphs with no text AND no meaningful
    spacing (space_before ≤ 2 pt). This avoids deleting intentional
    visual separators.

    Returns the number of paragraphs removed in this pass.
    """
    body_h = _body_height_pt(doc)
    body_w = _body_width_pt(doc)

    body_elems = list(_iter_body(doc))
    current_h = 0.0
    to_remove: list = []
    prev_nonempty_kind: str | None = None

    for kind, xml_elem, py_obj in body_elems:
        if kind == "paragraph":
            # LRPB calibration — only trust when enough page content was seen
            current_h = _lrpb_calibrate(xml_elem, current_h, body_h)

            h = _estimate_para_height(py_obj)

            page_overflow = (current_h + h > body_h)
            if page_overflow:
                current_h = 0.0   # new page starts

            text = (py_obj.text or "").strip()
            # A paragraph with an image but no text must never be treated as
            # "empty" — removing it would delete the figure from the document.
            is_empty = not text and not xml_has_image(xml_elem)

            if page_overflow and is_empty:
                pPr = xml_elem.find(qn("w:pPr"))
                if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                    current_h += h
                    continue

                # Preserve intentional blank lines that must remain after
                # headings and table note blocks ("Источник:" / "Примечание:").
                if prev_nonempty_kind in {"heading", "source_or_note"}:
                    current_h += h
                    continue

                # Check it's not a meaningful spacer (large space_before)
                try:
                    sb = py_obj.paragraph_format.space_before
                    if sb and sb.pt > 2:
                        current_h += h
                        continue
                except Exception:
                    pass
                to_remove.append(xml_elem)
                # current_h stays 0 — next element is still first on page
            else:
                current_h += h
                if not is_empty:
                    if _looks_like_heading(text):
                        prev_nonempty_kind = "heading"
                    elif is_source_or_note_line(text):
                        prev_nonempty_kind = "source_or_note"
                    else:
                        prev_nonempty_kind = "text"

        elif kind == "table":
            prev_nonempty_kind = "table"
            rows = py_obj.rows
            col_widths = _tbl_col_widths_pt(xml_elem)
            for rh in (_estimate_row_height(r, body_w, col_widths) for r in rows):
                current_h += rh
                if current_h > body_h:
                    current_h = rh

    removed = 0
    for elem in reversed(to_remove):
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            removed += 1

    return removed


def apply_rule4_empty_first_lines(doc: Document) -> int:
    """
    Rule 4 — Remove empty paragraphs that land at the very top of a page.

    Runs iteratively until convergence: each removal can shift subsequent
    page boundaries, potentially exposing new violations that the first
    pass missed (stale LRPB calibration + cascading removals).

    Returns total number of paragraphs removed across all passes.
    """
    total = 0
    for _ in range(5):   # cap at 5 iterations to prevent infinite loops
        n = _apply_rule4_pass(doc)
        total += n
        if n == 0:
            break
    logger.info("rule4: removed %d empty first-line paragraph(s) total", total)
    return total


# ── Rule 3: no orphan table caption at page bottom ────────────────────────────

def apply_rule3_table_orphan(doc: Document) -> int:
    """
    Rule 3 (geometry) — Prevent table caption from hanging alone at page bottom.

    If a table_caption paragraph (optionally followed by a short title line)
    fits on the current page but the table's first data row does not,
    set w:pageBreakBefore on the caption so the caption and table land
    together on the next page.

    Returns the number of captions given a pageBreakBefore.
    """
    body_h = _body_height_pt(doc)
    body_w = _body_width_pt(doc)
    body_elems = list(_iter_body(doc))
    n = len(body_elems)
    current_h = 0.0
    count = 0

    i = 0
    while i < n:
        kind, xml_elem, py_obj = body_elems[i]

        if kind == "paragraph":
            current_h = _lrpb_calibrate(xml_elem, current_h, body_h)

            text = (py_obj.text or "").strip()
            h = _estimate_para_height(py_obj)

            if not _TABLE_CAP_RE_GEOM.match(text):
                if current_h + h > body_h:
                    current_h = h
                else:
                    current_h += h
                i += 1
                continue

            # Found a table caption — collect caption + possible title lines
            cap_start_h = current_h
            cap_items: list[tuple] = [(xml_elem, h)]   # (xml_elem, height)

            j = i + 1
            while j < n:
                k2, xe2, po2 = body_elems[j]
                if k2 != "paragraph":
                    break
                t2 = (po2.text or "").strip()
                # Stop at: empty para, very long line (body text), another caption
                if not t2 or len(t2) > 200 or _TABLE_CAP_RE_GEOM.match(t2):
                    break
                cap_items.append((xe2, _estimate_para_height(po2)))
                j += 1

            cap_total_h = sum(h2 for _, h2 in cap_items)

            # j should point to the table element
            if j < n and body_elems[j][0] == "table":
                tbl_py = body_elems[j][2]
                tbl_xml = body_elems[j][1]
                rows = tbl_py.rows
                if rows:
                    col_widths = _tbl_col_widths_pt(tbl_xml)
                    first_row_h = _estimate_row_height(rows[0], body_w, col_widths)

                    caption_fits     = (cap_start_h + cap_total_h <= body_h)
                    first_row_orphan = (cap_start_h + cap_total_h + first_row_h > body_h)
                    fits_fresh       = (cap_total_h + first_row_h <= body_h)

                    if caption_fits and first_row_orphan and fits_fresh:
                        _set_page_break_before(cap_items[0][0])
                        count += 1
                        logger.info(
                            "rule3: pageBreakBefore on table caption [%s]",
                            text[:50],
                        )
                        current_h = cap_total_h
                        i = j      # resume from the table element
                        continue

            # No action — advance geometry past caption + title
            current_h = cap_start_h + cap_total_h
            if current_h > body_h:
                current_h = cap_items[-1][1]
            i = j
            continue

        elif kind == "table":
            rows = py_obj.rows
            col_widths = _tbl_col_widths_pt(xml_elem)
            for row in rows:
                rh = _estimate_row_height(row, body_w, col_widths)
                if current_h + rh > body_h:
                    current_h = rh
                else:
                    current_h += rh

        i += 1

    logger.info("rule3: %d table caption(s) given pageBreakBefore", count)
    return count


# ── Rule 6: figure must stay with its caption ─────────────────────────────────

def apply_rule6_figure_orphan(doc: Document) -> int:
    """
    Rule 6 (geometry) — Prevent figure caption from being stranded at the
    top of the next page while the figure itself is on the current page.

    If an image paragraph fits on the current page but the immediately
    following figure_caption does not, set w:pageBreakBefore on the image
    so both the image and caption land on the next page together.

    Returns the number of images given a pageBreakBefore.
    """
    body_h = _body_height_pt(doc)
    body_w = _body_width_pt(doc)
    body_elems = list(_iter_body(doc))
    n = len(body_elems)
    current_h = 0.0
    count = 0

    i = 0
    while i < n:
        kind, xml_elem, py_obj = body_elems[i]

        if kind == "paragraph":
            current_h = _lrpb_calibrate(xml_elem, current_h, body_h)

            h = _estimate_para_height(py_obj)

            if not _para_has_image(xml_elem):
                if current_h + h > body_h:
                    current_h = h
                else:
                    current_h += h
                i += 1
                continue

            # Image paragraph — use actual rendered height from wp:extent cy if
            # available; fall back to the generic paragraph height estimate.
            # The generic estimate returns ~21 pt (1 empty line) for image-only
            # paragraphs, massively underestimating real figure heights.
            img_h = _get_image_height_pt(xml_elem) or h

            # Image paragraph — check if the next paragraph is a figure caption.
            # Skip past any empty paragraphs between image and caption first.
            j = i + 1
            while j < n and body_elems[j][0] == "paragraph":
                nk, nxe, npo = body_elems[j]
                if (npo.text or "").strip():
                    break
                j += 1

            if j < n:
                nk, nxe, npo = body_elems[j]
                if nk == "paragraph":
                    next_text = (npo.text or "").strip()
                    if _FIGURE_CAP_RE_GEOM.match(next_text):
                        caption_h   = _estimate_para_height(npo)
                        img_fits    = (current_h + img_h <= body_h)
                        cap_orphan  = (current_h + img_h + caption_h > body_h)
                        fits_fresh  = (img_h + caption_h <= body_h)

                        if img_fits and cap_orphan and fits_fresh:
                            _set_page_break_before(xml_elem)
                            count += 1
                            logger.info(
                                "rule6: pageBreakBefore on image before [%s]",
                                next_text[:50],
                            )
                            # Both now start fresh on next page
                            current_h = img_h + caption_h
                            i = j + 1
                            continue

            # Normal advance (use img_h for accurate geometry tracking)
            if current_h + img_h > body_h:
                current_h = img_h
            else:
                current_h += img_h

        elif kind == "table":
            rows = py_obj.rows
            col_widths = _tbl_col_widths_pt(xml_elem)
            for row in rows:
                rh = _estimate_row_height(row, body_w, col_widths)
                if current_h + rh > body_h:
                    current_h = rh
                else:
                    current_h += rh

        i += 1

    logger.info("rule6: %d figure(s) given pageBreakBefore", count)
    return count


# ── Rule 2: no trailing empty lines at page bottom before a heading ───────────

_HEADING_RE = re.compile(
    r"^\s*\d+(?:\.\d+)*\.?\s",   # matches "1. …" / "1.1. …" / "1.1.1. …"
)


def _looks_like_heading(text: str) -> bool:
    return bool(_HEADING_RE.match(text))


def apply_rule2_trailing_empties(doc: Document) -> int:
    """
    Rule 2 — Remove empty paragraphs that sit at the very bottom of a page
    when the next non-empty element is a heading (heading1 / heading2).

    These ghost lines appear because the geometry estimator places them
    mid-page, but Word's real line-breaking pushes them to page bottom,
    so Rule 4 (which only catches first-on-page empties) never fires.

    Strategy:
      Walk body elements in order.  Collect runs of consecutive empty
      paragraphs.  When the run is followed by a heading-like paragraph
      AND the geometry says the run straddles or is near the page
      boundary (within _BOTTOM_TOLERANCE_PT), mark the empties for removal.

    Conservative: requires the very next non-empty paragraph to be a heading
    so we don't accidentally eat intentional visual separators between sections.

    Returns the number of paragraphs removed.
    """
    body_h = _body_height_pt(doc)
    body_w = _body_width_pt(doc)
    _BOTTOM_TOLERANCE_PT = _BODY_LINE_PT * 3   # empty lines within last ~3 lines

    body_elems = list(_iter_body(doc))
    n = len(body_elems)

    current_h = 0.0
    to_remove: list = []

    i = 0
    while i < n:
        kind, xml_elem, py_obj = body_elems[i]

        if kind == "paragraph":
            text = (py_obj.text or "").strip()
            h = _estimate_para_height(py_obj)

            if current_h + h > body_h:
                current_h = h   # new page

            if not text:
                # Start of a potential empty-paragraph run
                run_start = i
                run_elems = [(xml_elem, h)]
                run_h_start = current_h   # height at start of run

                j = i + 1
                while j < n:
                    k2, xe2, po2 = body_elems[j]
                    if k2 != "paragraph":
                        break
                    t2 = (po2.text or "").strip()
                    if t2:
                        break
                    run_elems.append((xe2, _estimate_para_height(po2)))
                    j += 1

                # j now points to the first non-empty element after the run
                next_is_heading = False
                if j < n:
                    k_next, _, po_next = body_elems[j]
                    if k_next == "paragraph":
                        t_next = (po_next.text or "").strip()
                        next_is_heading = _looks_like_heading(t_next)

                if next_is_heading:
                    run_total_h = sum(rh for _, rh in run_elems)
                    heading_h = _estimate_para_height(po_next)

                    # Only remove if the heading lands on the SAME page.
                    # If the empty run already pushes past body_h → heading is on
                    # the next page → the empties are harmless bottom-of-page padding,
                    # leave them alone (user confirmed this is acceptable).
                    heading_on_next_page = (
                        run_h_start + run_total_h + heading_h > body_h
                    )
                    if not heading_on_next_page:
                        for xe, _ in run_elems:
                            to_remove.append(xe)
                        current_h = run_h_start   # pretend the run wasn't there
                        i = j
                        continue

                # Otherwise just advance normally through the run
                for _, rh in run_elems:
                    current_h += rh
                    if current_h > body_h:
                        current_h = rh
                i = j
                continue

            else:
                current_h += h

        elif kind == "table":
            rows = py_obj.rows
            for rh in (_estimate_row_height(r, body_w) for r in rows):
                current_h += rh
                if current_h > body_h:
                    current_h = rh

        i += 1

    removed = 0
    for elem in reversed(to_remove):
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
            removed += 1

    logger.info("rule2: removed %d trailing empty paragraph(s) before headings", removed)
    return removed
