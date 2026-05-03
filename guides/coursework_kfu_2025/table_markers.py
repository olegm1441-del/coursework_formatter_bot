from __future__ import annotations

import re
import shutil
import tempfile
import unicodedata
import uuid
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.shared import Pt

from .classifier import TABLE_CAPTION_RE, clean_spaces
from .layout_render import render_docx_to_pdf
from .pdf_layout_analyzer import PdfLine, analyze_pdf_lines


_MARKER_RE = re.compile(
    r"KPFU_TMARK_(?P<salt>[A-F0-9]{6})_T(?P<table>\d{2})_R(?P<row>\d{3})"
)
_APPENDIX_HEADING_RE = re.compile(r"^\s*приложени(?:е|я)\b", re.IGNORECASE)


@dataclass(frozen=True)
class TableMarkerInstrumentation:
    instrumented_docx_path: Path
    marker_salt: str
    table_index: int
    total_rows: int
    row_markers: dict[int, str]


@dataclass(frozen=True)
class TableMarkerResult:
    row_pages: dict[int, int]
    found_rows: list[int]
    missing_rows: list[int]
    duplicate_rows: dict[int, list[int]]
    instrumented_docx_path: Path | None = None
    pdf_path: Path | None = None
    marker_font_size_pt: int = 1


@dataclass(frozen=True)
class TablePageSpan:
    start_row: int
    end_row: int
    page_num: int


@dataclass(frozen=True)
class TableMarkerDiagnostic:
    table_index: int
    rows_count: int
    pages_detected: list[int]
    row_pages: dict[int, int]
    found_rows: list[int]
    missing_rows: list[int]
    duplicate_rows: dict[int, list[int]]
    candidate_for_split: bool
    page_spans: list[TablePageSpan]
    appendix_table: bool
    caption_detected: bool
    has_standard_table_caption: bool
    preceding_paragraph_text: str | None = None
    instrumented_docx_path: Path | None = None
    pdf_path: Path | None = None
    marker_font_size_pt: int = 1
    error_message: str | None = None


def _build_marker(marker_salt: str, table_index: int, row_index: int) -> str:
    return f"KPFU_TMARK_{marker_salt}_T{table_index:02d}_R{row_index:03d}"


def _pick_marker_paragraph(row) -> object:
    best_paragraph = None
    best_len = None

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            text_len = len((paragraph.text or "").strip())
            if text_len == 0:
                continue
            if best_len is None or text_len < best_len:
                best_paragraph = paragraph
                best_len = text_len

    if best_paragraph is not None:
        return best_paragraph

    return row.cells[0].paragraphs[0]


def _resolve_existing_docx_path(docx_path) -> Path:
    candidate = Path(docx_path)
    if candidate.exists():
        return candidate

    parent = candidate.parent if str(candidate.parent) else Path(".")
    target_name = unicodedata.normalize("NFC", candidate.name)
    if parent.exists():
        siblings = list(parent.iterdir())
        for sibling in siblings:
            if unicodedata.normalize("NFC", sibling.name) == target_name:
                return sibling

        target_stem = unicodedata.normalize("NFC", candidate.stem)
        fuzzy_matches = [
            sibling
            for sibling in siblings
            if sibling.suffix.lower() == candidate.suffix.lower()
            and unicodedata.normalize("NFC", sibling.stem).startswith(target_stem)
        ]
        if len(fuzzy_matches) == 1:
            return fuzzy_matches[0]

    raise FileNotFoundError(f"DOCX not found: {candidate}")


def instrument_table_rows_copy(
    docx_path: Path | str,
    table_index: int,
    *,
    workdir: Path | str | None = None,
    marker_font_size_pt: int = 1,
) -> TableMarkerInstrumentation:
    source_path = _resolve_existing_docx_path(docx_path)

    if workdir is None:
        workdir_path = Path(tempfile.mkdtemp(prefix="table_markers_"))
    else:
        workdir_path = Path(workdir)
        workdir_path.mkdir(parents=True, exist_ok=True)

    instrumented_docx_path = workdir_path / f"{source_path.stem}_markers_{marker_font_size_pt}pt.docx"
    shutil.copy2(source_path, instrumented_docx_path)

    doc = Document(str(instrumented_docx_path))
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(
            f"table_index out of range: {table_index} (tables={len(doc.tables)})"
        )

    marker_salt = uuid.uuid4().hex[:6].upper()
    table = doc.tables[table_index]
    row_markers: dict[int, str] = {}

    for row_index, row in enumerate(table.rows):
        marker = _build_marker(marker_salt, table_index, row_index)
        paragraph = _pick_marker_paragraph(row)
        run = paragraph.add_run(marker)
        run.bold = False
        run.italic = False
        run.font.size = Pt(marker_font_size_pt)
        row_markers[row_index] = marker

    doc.save(str(instrumented_docx_path))
    return TableMarkerInstrumentation(
        instrumented_docx_path=instrumented_docx_path,
        marker_salt=marker_salt,
        table_index=table_index,
        total_rows=len(table.rows),
        row_markers=row_markers,
    )


def extract_row_pages_from_pdf_lines(
    pdf_lines: list[PdfLine],
    marker_salt: str,
    table_index: int,
    total_rows: int,
) -> TableMarkerResult:
    row_page_hits: dict[int, set[int]] = {}

    for line in pdf_lines:
        text = line.text or ""
        for match in _MARKER_RE.finditer(text):
            if match.group("salt") != marker_salt:
                continue
            if int(match.group("table")) != table_index:
                continue

            row_index = int(match.group("row"))
            row_page_hits.setdefault(row_index, set()).add(line.page_num)

    row_pages: dict[int, int] = {}
    duplicate_rows: dict[int, list[int]] = {}

    for row_index in range(total_rows):
        pages = sorted(row_page_hits.get(row_index, set()))
        if len(pages) == 1:
            row_pages[row_index] = pages[0]
        elif len(pages) > 1:
            duplicate_rows[row_index] = pages

    found_rows = sorted(row_pages)
    missing_rows = [
        row_index
        for row_index in range(total_rows)
        if row_index not in row_pages and row_index not in duplicate_rows
    ]

    return TableMarkerResult(
        row_pages=row_pages,
        found_rows=found_rows,
        missing_rows=missing_rows,
        duplicate_rows=duplicate_rows,
    )


def _cleanup_attempt(workdir_path: Path | None, pdf_path: Path | None) -> None:
    if pdf_path is not None:
        shutil.rmtree(pdf_path.parent, ignore_errors=True)
    if workdir_path is not None:
        shutil.rmtree(workdir_path, ignore_errors=True)


def _should_preserve_artifacts(
    result: TableMarkerResult,
    keep_temp: bool,
) -> bool:
    return keep_temp or bool(result.missing_rows or result.duplicate_rows)


def _iter_body_tables_with_context(doc: Document) -> list[dict]:
    body = doc.element.body
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    in_appendix = False
    last_nonempty_paragraph_text: str | None = None
    out: list[dict] = []

    for child in body:
        local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if local == "p" and child in para_map:
            text = clean_spaces(para_map[child].text or "")
            if text:
                if _APPENDIX_HEADING_RE.match(text):
                    in_appendix = True
                last_nonempty_paragraph_text = text
            continue

        if local == "tbl" and child in table_map:
            out.append({
                "table_obj": table_map[child],
                "appendix_table": in_appendix,
                "preceding_paragraph_text": last_nonempty_paragraph_text,
                "caption_detected": bool(last_nonempty_paragraph_text),
                "has_standard_table_caption": bool(
                    last_nonempty_paragraph_text
                    and TABLE_CAPTION_RE.match(last_nonempty_paragraph_text)
                ),
            })

    return out


def summarize_row_page_spans(row_pages: dict[int, int]) -> list[TablePageSpan]:
    if not row_pages:
        return []

    spans: list[TablePageSpan] = []
    sorted_rows = sorted(row_pages.items())
    start_row, current_page = sorted_rows[0]
    end_row = start_row

    for row_index, page_num in sorted_rows[1:]:
        if row_index == end_row + 1 and page_num == current_page:
            end_row = row_index
            continue

        spans.append(TablePageSpan(
            start_row=start_row,
            end_row=end_row,
            page_num=current_page,
        ))
        start_row = row_index
        end_row = row_index
        current_page = page_num

    spans.append(TablePageSpan(
        start_row=start_row,
        end_row=end_row,
        page_num=current_page,
    ))
    return spans


def map_table_rows_to_pages(
    docx_path: Path | str,
    table_index: int,
    *,
    keep_temp: bool = False,
) -> TableMarkerResult:
    source_path = _resolve_existing_docx_path(docx_path)
    last_result: TableMarkerResult | None = None
    last_workdir: Path | None = None
    last_pdf_path: Path | None = None

    for marker_font_size_pt in (1, 2):
        workdir_path = Path(tempfile.mkdtemp(prefix="table_markers_"))
        instrumentation = instrument_table_rows_copy(
            source_path,
            table_index,
            workdir=workdir_path,
            marker_font_size_pt=marker_font_size_pt,
        )
        pdf_path = render_docx_to_pdf(instrumentation.instrumented_docx_path)
        pdf_lines = analyze_pdf_lines(pdf_path)
        result = extract_row_pages_from_pdf_lines(
            pdf_lines,
            instrumentation.marker_salt,
            instrumentation.table_index,
            instrumentation.total_rows,
        )
        preserve = _should_preserve_artifacts(result, keep_temp)
        result = TableMarkerResult(
            row_pages=result.row_pages,
            found_rows=result.found_rows,
            missing_rows=result.missing_rows,
            duplicate_rows=result.duplicate_rows,
            instrumented_docx_path=instrumentation.instrumented_docx_path if preserve else None,
            pdf_path=pdf_path if preserve else None,
            marker_font_size_pt=marker_font_size_pt,
        )

        is_reliable = not result.missing_rows and not result.duplicate_rows
        if is_reliable:
            if not preserve:
                _cleanup_attempt(workdir_path, pdf_path)
            if last_result is not None:
                _cleanup_attempt(last_workdir, last_pdf_path)
            return result

        if marker_font_size_pt == 1 and not keep_temp:
            _cleanup_attempt(workdir_path, pdf_path)
            continue

        if last_result is not None:
            _cleanup_attempt(last_workdir, last_pdf_path)
        last_result = result
        last_workdir = workdir_path
        last_pdf_path = pdf_path

        if marker_font_size_pt == 2:
            return result

    if last_result is None:
        raise RuntimeError("Failed to map table rows to PDF pages")
    return last_result


def diagnose_table(
    docx_path: Path | str,
    table_index: int,
    *,
    keep_temp: bool = False,
) -> TableMarkerDiagnostic:
    doc = Document(str(docx_path))
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(
            f"table_index out of range: {table_index} (tables={len(doc.tables)})"
        )
    table_contexts = _iter_body_tables_with_context(doc)
    table_context = table_contexts[table_index]

    try:
        result = map_table_rows_to_pages(docx_path, table_index, keep_temp=keep_temp)
    except Exception as exc:
        return TableMarkerDiagnostic(
            table_index=table_index,
            rows_count=len(doc.tables[table_index].rows),
            pages_detected=[],
            row_pages={},
            found_rows=[],
            missing_rows=[],
            duplicate_rows={},
            candidate_for_split=False,
            page_spans=[],
            appendix_table=table_context["appendix_table"],
            caption_detected=table_context["caption_detected"],
            has_standard_table_caption=table_context["has_standard_table_caption"],
            preceding_paragraph_text=table_context["preceding_paragraph_text"],
            error_message=str(exc),
        )

    pages_detected = sorted(set(result.row_pages.values()))
    page_spans = summarize_row_page_spans(result.row_pages)
    candidate_for_split = (
        len(pages_detected) >= 2
        and not result.missing_rows
        and not result.duplicate_rows
    )
    return TableMarkerDiagnostic(
        table_index=table_index,
        rows_count=len(doc.tables[table_index].rows),
        pages_detected=pages_detected,
        row_pages=result.row_pages,
        found_rows=result.found_rows,
        missing_rows=result.missing_rows,
        duplicate_rows=result.duplicate_rows,
        candidate_for_split=candidate_for_split,
        page_spans=page_spans,
        appendix_table=table_context["appendix_table"],
        caption_detected=table_context["caption_detected"],
        has_standard_table_caption=table_context["has_standard_table_caption"],
        preceding_paragraph_text=table_context["preceding_paragraph_text"],
        instrumented_docx_path=result.instrumented_docx_path,
        pdf_path=result.pdf_path,
        marker_font_size_pt=result.marker_font_size_pt,
        error_message=None,
    )


def diagnose_all_tables(
    docx_path: Path | str,
    *,
    keep_temp: bool = False,
) -> list[TableMarkerDiagnostic]:
    doc = Document(str(docx_path))
    diagnostics: list[TableMarkerDiagnostic] = []
    for table_index in range(len(doc.tables)):
        diagnostics.append(
            diagnose_table(docx_path, table_index, keep_temp=keep_temp)
        )
    return diagnostics
