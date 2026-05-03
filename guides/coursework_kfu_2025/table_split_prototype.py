from __future__ import annotations

import re
import shutil
import tempfile
import unicodedata
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .docx_utils import is_source_or_note_line


_STANDARD_TABLE_NUMBER_RE = re.compile(
    r"^\s*(?:таблица|table)\s+(?P<num>\d+(?:\.\d+){0,2})\.?\s*(?:[-—–].*)?$",
    re.IGNORECASE,
)


@dataclass(frozen=True)
class PrototypeSplitResult:
    output_docx_path: Path | None
    workdir_path: Path | None
    table_index: int
    second_table_index: int
    total_tables_before: int
    total_tables_after: int
    original_rows_count: int
    first_table_rows_count: int
    second_table_rows_count: int
    split_before_row: int
    header_rows: int
    source_note_after_second: bool | None
    source_note_text: str | None
    continuation_paragraph_inserted: bool
    continuation_text: str | None
    numbered_header_enabled: bool
    numbered_row_reused: bool | None
    column_count: int | None
    diagnostics: list[str] = field(default_factory=list)


def _body_children(doc: Document) -> list:
    return list(doc.element.body)


def _local_name(elem) -> str:
    return elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag


def _find_table_body_index(doc: Document, table_index: int) -> int:
    target_tbl = doc.tables[table_index]._tbl
    for idx, child in enumerate(_body_children(doc)):
        if child is target_tbl:
            return idx
    raise RuntimeError(f"Table body element not found for table_index={table_index}")


def _paragraph_text_from_elem(p_elem) -> str:
    texts = p_elem.findall(".//" + qn("w:t"))
    return "".join(t.text or "" for t in texts).strip()


def _find_following_source_note(children: list, table_body_index: int) -> str | None:
    seen_empty = 0
    for child in children[table_body_index + 1:]:
        if _local_name(child) != "p":
            break
        text = _paragraph_text_from_elem(child)
        if not text:
            seen_empty += 1
            if seen_empty > 1:
                break
            continue
        if is_source_or_note_line(text):
            return text
        break
    return None


def _find_preceding_nonempty_paragraph_text(children: list, table_body_index: int) -> str | None:
    for child in reversed(children[:table_body_index]):
        if _local_name(child) != "p":
            break
        text = _paragraph_text_from_elem(child)
        if text:
            return text
    return None


def _resolve_existing_docx_path(docx_path) -> Path:
    candidate = Path(docx_path)
    if candidate.exists():
        return candidate

    parent = candidate.parent if str(candidate.parent) else Path(".")
    target_name = unicodedata.normalize("NFC", candidate.name)
    if parent.exists():
        normalized_siblings = list(parent.iterdir())
        for sibling in parent.iterdir():
            if unicodedata.normalize("NFC", sibling.name) == target_name:
                return sibling

        target_stem = unicodedata.normalize("NFC", candidate.stem)
        fuzzy_matches = [
            sibling
            for sibling in normalized_siblings
            if sibling.suffix.lower() == candidate.suffix.lower()
            and unicodedata.normalize("NFC", sibling.stem).startswith(target_stem)
        ]
        if len(fuzzy_matches) == 1:
            return fuzzy_matches[0]

    raise FileNotFoundError(f"DOCX not found: {candidate}")


def _ensure_tbl_header(tr_elem) -> None:
    tr_pr = tr_elem.find(qn("w:trPr"))
    if tr_pr is None:
        tr_pr = OxmlElement("w:trPr")
        tr_elem.insert(0, tr_pr)
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def _extract_standard_table_number(text: str | None) -> str | None:
    if not text:
        return None
    match = _STANDARD_TABLE_NUMBER_RE.match(text.strip())
    if not match:
        return None
    return match.group("num")


def _tbl_grid_column_count(tbl_xml) -> int:
    tbl_grid = tbl_xml.find(qn("w:tblGrid"))
    if tbl_grid is None:
        raise ValueError("table has no tblGrid; cannot build numbered header safely")
    grid_cols = tbl_grid.findall(qn("w:gridCol"))
    if not grid_cols:
        raise ValueError("table tblGrid is empty; cannot build numbered header safely")
    return len(grid_cols)


def _iter_row_cells(tr_elem) -> list:
    return tr_elem.findall(qn("w:tc"))


def _cell_grid_span(tc_elem) -> int:
    tc_pr = tc_elem.find(qn("w:tcPr"))
    if tc_pr is None:
        return 1
    grid_span = tc_pr.find(qn("w:gridSpan"))
    if grid_span is None:
        return 1
    val = grid_span.get(qn("w:val")) or grid_span.get("w:val") or grid_span.get("val")
    try:
        return max(1, int(val))
    except Exception:
        return 1


def _has_vmerge(tc_elem) -> bool:
    tc_pr = tc_elem.find(qn("w:tcPr"))
    return tc_pr is not None and tc_pr.find(qn("w:vMerge")) is not None


def _cell_text(tc_elem) -> str:
    texts = tc_elem.findall(".//" + qn("w:t"))
    return "".join(t.text or "" for t in texts).strip()


def _row_is_simple_full_width(tr_elem, column_count: int) -> bool:
    cells = _iter_row_cells(tr_elem)
    if not cells:
        return False
    total_span = 0
    for tc in cells:
        if _has_vmerge(tc):
            return False
        span = _cell_grid_span(tc)
        if span != 1:
            return False
        total_span += span
    return total_span == column_count and len(cells) == column_count


def _row_is_exact_numbered_row(tr_elem, column_count: int) -> bool:
    if not _row_is_simple_full_width(tr_elem, column_count):
        return False
    values = [_cell_text(tc) for tc in _iter_row_cells(tr_elem)]
    return values == [str(i) for i in range(1, column_count + 1)]


def _row_looks_numbered_but_malformed(tr_elem, column_count: int) -> bool:
    if not _row_is_simple_full_width(tr_elem, column_count):
        return False
    values = [_cell_text(tc) for tc in _iter_row_cells(tr_elem)]
    if not values:
        return False
    if not all(re.fullmatch(r"\d+", value or "") for value in values):
        return False
    return values != [str(i) for i in range(1, column_count + 1)]


def _row_has_numpr(tr_elem) -> bool:
    for p in tr_elem.findall(".//" + qn("w:p")):
        p_pr = p.find(qn("w:pPr"))
        if p_pr is not None and p_pr.find(qn("w:numPr")) is not None:
            return True
    return False


def _remove_all_children(elem) -> None:
    for child in list(elem):
        elem.remove(child)


def _build_numbered_row_from_header(header_tr_elem, column_count: int):
    if not _row_is_simple_full_width(header_tr_elem, column_count):
        raise ValueError("complex merged header is not supported for numbered prototype split")
    numbered_tr = deepcopy(header_tr_elem)
    _ensure_tbl_header(numbered_tr)
    for idx, tc in enumerate(_iter_row_cells(numbered_tr), start=1):
        tc_pr = tc.find(qn("w:tcPr"))
        _remove_all_children(tc)
        if tc_pr is not None:
            tc.append(tc_pr)
        p = OxmlElement("w:p")
        p_pr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        p_pr.append(jc)
        p.append(p_pr)
        r = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Times New Roman")
        r_fonts.set(qn("w:hAnsi"), "Times New Roman")
        r_fonts.set(qn("w:cs"), "Times New Roman")
        r_pr.append(r_fonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        r_pr.append(sz)
        sz_cs = OxmlElement("w:szCs")
        sz_cs.set(qn("w:val"), "24")
        r_pr.append(sz_cs)
        r.append(r_pr)
        t = OxmlElement("w:t")
        t.text = str(idx)
        r.append(t)
        p.append(r)
        tc.append(p)
    return numbered_tr


def _build_plain_paragraph(text: str):
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def apply_numbered_split_to_document(
    doc: Document,
    table_index: int,
    split_before_row: int,
    *,
    header_rows: int = 1,
    numbered_header: bool = False,
    appendix_table: bool = False,
    continuation_paragraph_builder=None,
) -> PrototypeSplitResult:
    total_tables_before = len(doc.tables)
    if table_index < 0 or table_index >= total_tables_before:
        raise ValueError(
            f"table_index out of range: {table_index} (tables={total_tables_before})"
        )

    table = doc.tables[table_index]
    rows_xml = table._tbl.findall(qn("w:tr"))
    original_rows_count = len(rows_xml)
    if header_rows < 0 or header_rows >= original_rows_count:
        raise ValueError(
            f"header_rows out of range: {header_rows} (rows={original_rows_count})"
        )
    if header_rows != 1:
        raise ValueError("prototype split v1 supports only header_rows=1")
    if split_before_row <= 0 or split_before_row >= original_rows_count:
        raise ValueError(
            f"split_before_row out of range: {split_before_row} (rows={original_rows_count})"
        )
    if split_before_row < header_rows:
        raise ValueError(
            f"split_before_row must be >= header_rows: split_before_row={split_before_row}, header_rows={header_rows}"
        )

    children_before = _body_children(doc)
    table_body_index = _find_table_body_index(doc, table_index)
    source_note_before = _find_following_source_note(children_before, table_body_index)
    preceding_paragraph_text = _find_preceding_nonempty_paragraph_text(children_before, table_body_index)

    tbl_xml = table._tbl
    second_tbl = deepcopy(tbl_xml)
    for tr in list(second_tbl.findall(qn("w:tr"))):
        second_tbl.remove(tr)

    column_count = None
    numbered_row_reused = None
    continuation_text = None
    continuation_inserted = False
    numbered_row_for_second = None

    if numbered_header:
        column_count = _tbl_grid_column_count(tbl_xml)
        header_row_xml = rows_xml[0]
        if not _row_is_simple_full_width(header_row_xml, column_count):
            raise ValueError("complex merged header is not supported for numbered prototype split")

        if len(rows_xml) > 1 and _row_is_exact_numbered_row(rows_xml[1], column_count):
            if _row_has_numpr(rows_xml[1]):
                raise ValueError("existing numbered row uses paragraph numbering")
            numbered_row_reused = True
            numbered_row_for_second = deepcopy(rows_xml[1])
        elif len(rows_xml) > 1 and _row_looks_numbered_but_malformed(rows_xml[1], column_count):
            raise ValueError("existing numbered row is malformed")
        else:
            numbered_row_reused = False
            numbered_row_for_second = _build_numbered_row_from_header(header_row_xml, column_count)
            header_row_xml.addnext(deepcopy(numbered_row_for_second))

        if not appendix_table:
            table_number = _extract_standard_table_number(preceding_paragraph_text)
            if table_number is None:
                raise ValueError("ordinary numbered split requires a standard table caption")
            continuation_text = f"Продолжение таблицы {table_number}"
            continuation_inserted = True

    tail_rows_xml = [deepcopy(r) for r in rows_xml[split_before_row:]]
    if not tail_rows_xml:
        raise ValueError("split would create an empty second table")

    if numbered_header:
        second_tbl.append(deepcopy(numbered_row_for_second))
    else:
        header_rows_xml = [deepcopy(r) for r in rows_xml[:header_rows]]
        for tr in header_rows_xml:
            _ensure_tbl_header(tr)
            second_tbl.append(tr)
    for tr in tail_rows_xml:
        second_tbl.append(tr)

    for tr in rows_xml[split_before_row:]:
        tbl_xml.remove(tr)

    if continuation_inserted and continuation_text is not None:
        builder = continuation_paragraph_builder or _build_plain_paragraph
        continuation_p = builder(continuation_text)
        tbl_xml.addnext(continuation_p)
        continuation_p.addnext(second_tbl)
    else:
        tbl_xml.addnext(second_tbl)

    total_tables_after = len(doc.tables)
    second_table_index = table_index + 1
    if second_table_index >= total_tables_after:
        raise RuntimeError("Second table missing after split")

    first_table_rows_count = len(doc.tables[table_index].rows)
    second_table_rows_count = len(doc.tables[second_table_index].rows)

    source_note_after = None
    if source_note_before is not None:
        children_after = _body_children(doc)
        second_table_body_index = _find_table_body_index(doc, second_table_index)
        source_note_after = _find_following_source_note(children_after, second_table_body_index)

    return PrototypeSplitResult(
        output_docx_path=None,
        workdir_path=None,
        table_index=table_index,
        second_table_index=second_table_index,
        total_tables_before=total_tables_before,
        total_tables_after=total_tables_after,
        original_rows_count=original_rows_count,
        first_table_rows_count=first_table_rows_count,
        second_table_rows_count=second_table_rows_count,
        split_before_row=split_before_row,
        header_rows=header_rows,
        source_note_after_second=(
            source_note_after == source_note_before if source_note_before is not None else None
        ),
        source_note_text=source_note_before,
        continuation_paragraph_inserted=continuation_inserted,
        continuation_text=continuation_text,
        numbered_header_enabled=numbered_header,
        numbered_row_reused=numbered_row_reused,
        column_count=column_count,
        diagnostics=[
            f"first_table_rows={first_table_rows_count}",
            f"second_table_rows={second_table_rows_count}",
            f"total_tables_after={total_tables_after}",
            f"numbered_header={numbered_header}",
            f"numbered_row_reused={numbered_row_reused}",
            f"continuation_inserted={continuation_inserted}",
        ],
    )


def prototype_split_table_copy(
    docx_path,
    table_index,
    split_before_row,
    header_rows=1,
    *,
    numbered_header=False,
    appendix_table=False,
    keep_temp=False,
) -> PrototypeSplitResult:
    source_path = _resolve_existing_docx_path(docx_path)

    workdir_path = Path(tempfile.mkdtemp(prefix="table_split_proto_"))
    output_docx_path = workdir_path / f"{source_path.stem}_split_prototype.docx"
    shutil.copy2(source_path, output_docx_path)

    try:
        doc = Document(str(output_docx_path))
        result = apply_numbered_split_to_document(
            doc,
            table_index,
            split_before_row,
            header_rows=header_rows,
            numbered_header=numbered_header,
            appendix_table=appendix_table,
            continuation_paragraph_builder=_build_plain_paragraph,
        )
        doc.save(str(output_docx_path))
        result = PrototypeSplitResult(
            output_docx_path=output_docx_path if keep_temp else None,
            workdir_path=workdir_path if keep_temp else None,
            table_index=result.table_index,
            second_table_index=result.second_table_index,
            total_tables_before=result.total_tables_before,
            total_tables_after=result.total_tables_after,
            original_rows_count=result.original_rows_count,
            first_table_rows_count=result.first_table_rows_count,
            second_table_rows_count=result.second_table_rows_count,
            split_before_row=result.split_before_row,
            header_rows=result.header_rows,
            source_note_after_second=result.source_note_after_second,
            source_note_text=result.source_note_text,
            continuation_paragraph_inserted=result.continuation_paragraph_inserted,
            continuation_text=result.continuation_text,
            numbered_header_enabled=result.numbered_header_enabled,
            numbered_row_reused=result.numbered_row_reused,
            column_count=result.column_count,
            diagnostics=result.diagnostics,
        )
        if not keep_temp:
            shutil.rmtree(workdir_path, ignore_errors=True)
        return result

    except Exception:
        if not keep_temp:
            shutil.rmtree(workdir_path, ignore_errors=True)
        raise
