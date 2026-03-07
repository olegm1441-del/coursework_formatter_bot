#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import re
import sys
from pathlib import Path
from statistics import median

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


FONT_NAME_TARGET = "Times New Roman"
BODY_FONT_SIZE_TARGET = 14
TABLE_FONT_SIZE_TARGET = 12
FIRST_LINE_INDENT_TARGET_CM = 1.25
LINE_SPACING_TARGET = 1.5

INTRO_HEADING = "введение"

H1_EXACT = {
    "содержание",
    "введение",
    "заключение",
    "список использованных источников",
    "список использованной литературы",
    "приложения",
}

CHAPTER_RE = re.compile(r"^\s*глава\s+\d+\b.*$", re.IGNORECASE)
NORMALIZED_H1_RE = re.compile(r"^\s*\d+\.\s+\S.*$")
H2_RE = re.compile(r"^\s*\d+\.\d+\.?\s+\S.*$")
TABLE_CAPTION_RE = re.compile(r"^\s*таблица\s+\d+(\.\d+){1,2}\.?\s*$", re.IGNORECASE)
FIGURE_CAPTION_RE = re.compile(r"^\s*(рис\.|рисунок)\s*\d+(\.\d+){1,2}\.?\s+.+$", re.IGNORECASE)
SOURCE_LINE_RE = re.compile(
    r"^\s*(источник|составлено по|рассчитано по|примечание)\s*:\s*.+$",
    re.IGNORECASE
)


def clean_spaces(text: str) -> str:
    if text is None:
        return ""
    text = text.replace("\u00A0", " ")
    text = text.replace("\u2007", " ")
    text = text.replace("\u202F", " ")
    text = text.replace("\t", " ")
    text = re.sub(r"[ ]{2,}", " ", text)
    text = re.sub(r"\s+([,.;:!?])", r"\1", text)
    return text.strip()


def alignment_to_str(alignment):
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        None: "default",
    }
    return mapping.get(alignment, str(alignment))


def pt_to_float(value):
    if value is None:
        return None
    try:
        return round(value.pt, 2)
    except Exception:
        return None


def indent_to_cm(value):
    if value is None:
        return None
    try:
        return round(value.cm, 2)
    except Exception:
        return None


def normalize_font_name(name):
    if not name:
        return None
    return str(name).strip()


def median_or_none(values):
    vals = [v for v in values if v is not None]
    if not vals:
        return None
    return round(float(median(vals)), 2)


def paragraph_text(paragraph):
    return clean_spaces(paragraph.text)


def is_empty_paragraph(paragraph):
    return paragraph_text(paragraph) == ""


def is_probable_heading1(text: str) -> bool:
    t = clean_spaces(text)
    if not t:
        return False
    low = t.lower()
    return low in H1_EXACT or bool(CHAPTER_RE.match(t)) or bool(NORMALIZED_H1_RE.match(t))


def is_probable_heading2(text: str) -> bool:
    t = clean_spaces(text)
    if not t:
        return False
    if len(t) > 180:
        return False
    if t.endswith("."):
        return False
    return bool(H2_RE.match(t))


def classify_paragraph(text: str):
    t = clean_spaces(text)
    low = t.lower()
    if not t:
        return "empty_paragraph"
    if low == "содержание":
        return "toc_heading_candidate"
    if TABLE_CAPTION_RE.match(t):
        return "table_caption_candidate"
    if FIGURE_CAPTION_RE.match(t):
        return "figure_caption_candidate"
    if SOURCE_LINE_RE.match(t):
        return "source_line_candidate"
    if is_probable_heading1(t):
        return "heading1_candidate"
    if is_probable_heading2(t):
        return "heading2_candidate"
    return "body_text"


def collect_run_stats(paragraph):
    font_names = []
    font_sizes = []
    bold_count = 0
    italic_count = 0
    run_count = 0

    for run in paragraph.runs:
        run_count += 1
        fname = normalize_font_name(run.font.name)
        fsize = pt_to_float(run.font.size)

        if fname:
            font_names.append(fname)
        if fsize is not None:
            font_sizes.append(fsize)
        if run.bold:
            bold_count += 1
        if run.italic:
            italic_count += 1

    return {
        "run_count": run_count,
        "font_names": sorted(set(font_names)),
        "median_font_size_pt": median_or_none(font_sizes),
        "bold_runs": bold_count,
        "italic_runs": italic_count,
    }


def collect_paragraph_metrics(paragraph):
    fmt = paragraph.paragraph_format
    return {
        "alignment": alignment_to_str(paragraph.alignment),
        "first_line_indent_cm": indent_to_cm(fmt.first_line_indent),
        "left_indent_cm": indent_to_cm(fmt.left_indent),
        "right_indent_cm": indent_to_cm(fmt.right_indent),
        "line_spacing": fmt.line_spacing if isinstance(fmt.line_spacing, (int, float)) else None,
        "space_before_pt": pt_to_float(fmt.space_before),
        "space_after_pt": pt_to_float(fmt.space_after),
    }


def paragraph_flags(info):
    flags = []
    kind = info["kind"]

    if kind == "body_text":
        if info["metrics"]["alignment"] not in {"justify", "default"}:
            flags.append("body_not_justify")
        if info["run_stats"]["median_font_size_pt"] not in {None, 14.0}:
            flags.append("body_font_size_not_14")
        if info["metrics"]["first_line_indent_cm"] not in {None, 1.25}:
            flags.append("body_first_line_indent_not_1_25")
        if FONT_NAME_TARGET not in info["run_stats"]["font_names"] and info["run_stats"]["font_names"]:
            flags.append("body_font_not_tnr")

    if kind == "toc_heading_candidate":
        if info["metrics"]["alignment"] != "center":
            flags.append("toc_heading_not_center")

    if kind == "heading1_candidate":
        if info["metrics"]["alignment"] != "center":
            flags.append("h1_not_center")
        if info["run_stats"]["median_font_size_pt"] not in {None, 14.0}:
            flags.append("h1_font_size_not_14")

    if kind == "heading2_candidate":
        if info["metrics"]["alignment"] != "center":
            flags.append("h2_not_center")
        if info["run_stats"]["median_font_size_pt"] not in {None, 14.0}:
            flags.append("h2_font_size_not_14")

    return flags


def analyze_tables(doc):
    tables_report = []

    for table_idx, table in enumerate(doc.tables, start=1):
        para_count = 0
        font_names = []
        font_sizes = []

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    para_count += 1
                    for run in paragraph.runs:
                        fname = normalize_font_name(run.font.name)
                        fsize = pt_to_float(run.font.size)
                        if fname:
                            font_names.append(fname)
                        if fsize is not None:
                            font_sizes.append(fsize)

        tables_report.append({
            "table_index": table_idx,
            "paragraphs_in_table": para_count,
            "font_names": sorted(set(font_names)),
            "median_font_size_pt": median_or_none(font_sizes),
            "font_size_target_pt": TABLE_FONT_SIZE_TARGET,
        })

    return tables_report


def find_body_start_index(doc):
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph_text(paragraph).lower() == INTRO_HEADING:
            return idx
    return None


def analyze_document(path: Path):
    doc = Document(str(path))
    body_start = find_body_start_index(doc)

    paragraphs_report = []
    if body_start is not None:
        for idx, paragraph in enumerate(doc.paragraphs):
            if idx < body_start:
                continue

            info = {
                "index": idx,
                "text": paragraph_text(paragraph),
                "kind": classify_paragraph(paragraph.text),
                "metrics": collect_paragraph_metrics(paragraph),
                "run_stats": collect_run_stats(paragraph),
            }
            info["flags"] = paragraph_flags(info)
            paragraphs_report.append(info)

    return {
        "file": str(path),
        "body_start_index": body_start,
        "paragraphs": paragraphs_report,
        "tables": analyze_tables(doc),
    }


def main():
    if len(sys.argv) != 3:
        print("Использование:")
        print("  python3 inspector.py input.docx report.json")
        sys.exit(1)

    input_path = Path(sys.argv[1]).expanduser().resolve()
    output_path = Path(sys.argv[2]).expanduser().resolve()

    report = analyze_document(input_path)
    output_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"[OK] Отчёт сохранён: {output_path}")


if __name__ == "__main__":
    main()
