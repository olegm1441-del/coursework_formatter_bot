"""Each appendix START label (ПРИЛОЖЕНИЕ N) must begin on a new page — not only
the first appendix after the references block. Pins the page_breaks fix.

Run: python3 tests/test_appendix_page_break.py
"""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
from guides.coursework_kfu_2025.page_breaks import apply_page_breaks, _needs_page_break_before  # noqa: E402


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def test_appendix_start_label_needs_page_break() -> tuple[bool, str]:
    for label in ("ПРИЛОЖЕНИЕ Б", "Приложение А", "ПРИЛОЖЕНИЕ 2"):
        if not _needs_page_break_before(label):
            return _result(False, f"{label!r} should require a page break")
    return _result(True, "appendix start labels require a page break")


def test_appendix_body_phrase_does_not_break() -> tuple[bool, str]:
    # standalone plural section heading is handled elsewhere; body phrases that
    # merely mention приложение must NOT force a page break
    for t in ("приложение к договору поставки", "Приложение А содержит данные по"):
        if _needs_page_break_before(t):
            return _result(False, f"body phrase wrongly flagged for page break: {t!r}")
    return _result(True, "appendix body phrases do not force a page break")


def test_second_appendix_label_gets_page_break_in_doc() -> tuple[bool, str]:
    doc = Document()
    for t in [
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "1. Иванов И.И. Что-то. — М., 2020.",
        "ПРИЛОЖЕНИЯ",
        "ПРИЛОЖЕНИЕ А",
        "Таблица А Данные",
        "строка данных один два три",
        "ПРИЛОЖЕНИЕ Б",
        "Съел дедал",
    ]:
        doc.add_paragraph(t)
    apply_page_breaks(doc, 0)
    by_text = {p.text: p for p in doc.paragraphs}
    if not by_text["ПРИЛОЖЕНИЕ Б"].paragraph_format.page_break_before:
        return _result(False, "ПРИЛОЖЕНИЕ Б must start on a new page")
    if not by_text["ПРИЛОЖЕНИЕ А"].paragraph_format.page_break_before:
        return _result(False, "ПРИЛОЖЕНИЕ А must start on a new page")
    # a normal data row in between must NOT get a page break
    if by_text["строка данных один два три"].paragraph_format.page_break_before:
        return _result(False, "ordinary appendix body row must not get a page break")
    return _result(True, "each appendix label starts on a new page")


def main() -> int:
    tests = [
        ("appendix start label needs break", test_appendix_start_label_needs_page_break),
        ("appendix body phrase no break", test_appendix_body_phrase_does_not_break),
        ("second appendix label breaks in doc", test_second_appendix_label_gets_page_break_in_doc),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        print(f"[{'PASS' if ok else 'FAIL'}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
