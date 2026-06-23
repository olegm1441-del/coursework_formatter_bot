"""Content-preservation gate tests: number/case normalization must not cause
false positives, but real cell loss/dup, lost source-notes, and lost references
must be hard fails. Pins `content_preservation`.

Run: python3 tests/test_content_preservation.py
"""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
import guides.coursework_kfu_2025.content_preservation as cp  # noqa: E402


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def _doc(tables=(), paras=()):
    d = Document()
    for p in paras:
        d.add_paragraph(p)
    for rows in tables:
        t = d.add_table(rows=len(rows), cols=len(rows[0]))
        for r, vals in zip(t.rows, rows):
            for c, v in zip(r.cells, vals):
                c.text = v
    return d


def test_normalize_cell_numbers_and_case() -> tuple[bool, str]:
    if cp.normalize_cell("31,75") != cp.normalize_cell("31.75"):
        return _result(False, "decimal comma/dot not unified")
    if cp.normalize_cell("13,00") != cp.normalize_cell("13"):
        return _result(False, "trailing decimal zeros not stripped")
    if cp.normalize_cell("Лидер совета") != cp.normalize_cell("лидер совета"):
        return _result(False, "case not folded")
    return _result(True, "number/case normalization canonicalises accepted variants")


def test_decimal_and_capitalization_not_flagged() -> tuple[bool, str]:
    src = _doc(tables=[[["Орган", "31,75", "13,00"], ["правление", "по плану", "2,5"]]])
    out = _doc(tables=[[["Орган", "31.75", "13"], ["Правление", "по плану", "2.50"]]])
    report, issues = cp.evaluate_content_preservation(src, out)
    if any(i.severity == "fail" for i in issues):
        return _result(False, f"accepted normalization wrongly flagged: {[i.issue_type for i in issues]}")
    return _result(True, "decimal/capitalization normalization is not a content failure")


def test_lost_cell_is_fail() -> tuple[bool, str]:
    src = _doc(tables=[[["альфа данные", "бета данные"], ["гамма данные", "дельта данные"]]])
    out = _doc(tables=[[["альфа данные", "бета данные"]]])  # lost a row of cells
    report, issues = cp.evaluate_content_preservation(src, out)
    if not any(i.issue_type == "lost_table_cell_content" for i in issues):
        return _result(False, f"lost cell not flagged: {[i.issue_type for i in issues]}")
    return _result(True, "lost table cell content is a hard fail")


def test_duplicated_data_row_is_fail() -> tuple[bool, str]:
    # header + one data row in source; output repeats the DATA row (not the header)
    src = _doc(tables=[[["Заголовок A", "Заголовок B"],
                        ["уникальные данные один", "значение один"]]])
    out = _doc(tables=[[["Заголовок A", "Заголовок B"],
                        ["уникальные данные один", "значение один"],
                        ["уникальные данные один", "значение один"]]])
    report, issues = cp.evaluate_content_preservation(src, out)
    if not any(i.issue_type == "duplicated_table_data_row" for i in issues):
        return _result(False, f"duplicated data row not flagged: {[i.issue_type for i in issues]}")
    return _result(True, "duplicated table data row is a hard fail")


def test_repeated_header_on_continuation_not_flagged() -> tuple[bool, str]:
    # two fragments of one logical table each begin with the SAME header — the
    # header repetition must NOT register as duplicated content
    src = _doc(tables=[[["Орган", "Функция"], ["совет", "надзор"], ["правление", "управление"]]])
    out = _doc(tables=[[["Орган", "Функция"], ["совет", "надзор"]],
                       [["Орган", "Функция"], ["правление", "управление"]]])
    report, issues = cp.evaluate_content_preservation(src, out)
    if any(i.severity == "fail" for i in issues):
        return _result(False, f"header repetition wrongly flagged: {[i.issue_type for i in issues]}")
    return _result(True, "repeated header across fragments is not a content failure")


def test_lost_source_note_is_fail() -> tuple[bool, str]:
    src = _doc(paras=["Источник: составлено автором.", "Примечание: важное."])
    out = _doc(paras=["Источник: составлено автором."])  # lost the примечание
    report, issues = cp.evaluate_content_preservation(src, out)
    if not any(i.issue_type == "lost_source_note_line" for i in issues):
        return _result(False, f"lost source/note not flagged: {[i.issue_type for i in issues]}")
    return _result(True, "lost Источник:/Примечание: line is a hard fail")


def test_reference_renumbering_not_a_content_fail() -> tuple[bool, str]:
    # references are reformatted/renumbered by the B1 subsystem; that must not be
    # a content fail (reference SECTION presence is covered by the structure gate)
    src = _doc(paras=["СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "Иванов И. Книга. 2020.", "Петров П. Статья. 2021."])
    out = _doc(paras=["СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "1. Иванов И. Книга. 2020.", "2. Петров П. Статья. 2021."])
    report, issues = cp.evaluate_content_preservation(src, out)
    if any(i.severity == "fail" for i in issues):
        return _result(False, f"reference renumbering wrongly flagged: {[i.issue_type for i in issues]}")
    return _result(True, "reference renumbering is not a content fail")


def main() -> int:
    tests = [
        ("normalize numbers/case", test_normalize_cell_numbers_and_case),
        ("decimal/capitalization not flagged", test_decimal_and_capitalization_not_flagged),
        ("lost cell is fail", test_lost_cell_is_fail),
        ("duplicated data row is fail", test_duplicated_data_row_is_fail),
        ("repeated header not flagged", test_repeated_header_on_continuation_not_flagged),
        ("lost source/note is fail", test_lost_source_note_is_fail),
        ("reference renumbering not a fail", test_reference_renumbering_not_a_content_fail),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        print(f"[{'PASS' if ok else 'FAIL'}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
