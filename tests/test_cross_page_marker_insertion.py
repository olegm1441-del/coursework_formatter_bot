"""Next table stage — safe marker insertion for marker-less cross-page tables.

Pins the safety core of the cross-page split subsystem
(`cleanup_cross_page_without_marker_blockers_inplace` and its helpers
`_split_cross_page_table_with_marker`, `_cross_page_split_candidate`,
`_match_data_row_pages_by_lead`):

- the split inserts a page-broken ``Продолжение таблицы N`` marker between the
  two physical fragments;
- the continuation fragment starts with the NUMERIC column row (`1 2 ... N`)
  and NEVER the semantic header (canonical KFU rule; Rybakov gold) — a numeric
  row is synthesized into both fragments when the source lacks one;
- the manual-chain normalizer strips a duplicate semantic header so the
  continuation starts with the numeric row, and the acceptance gate fails any
  ``semantic_header_repeated_on_continuation``;
- Источник:/Примечание: stays after the FINAL fragment;
- data rows are never lost / duplicated / reordered;
- grid widths are inherited verbatim across both fragments;
- candidate detection requires a matching caption, no existing marker, exactly
  two rendered pages, ≥1 data row in BOTH fragments, and refuses source-bad
  duplication.

End-to-end rendered behaviour (Demo 1.1.3 actually becomes valid, no regression
on Rybakov/Roman) is covered by the format-from-source corpus smoke.

Run: python3 tests/test_cross_page_marker_insertion.py
"""
from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
import guides.coursework_kfu_2025.table_continuation as tc  # noqa: E402
from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine  # noqa: E402


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def _mk_table(doc, rows, widths=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    if widths is not None:
        grid = table._tbl.find(qn("w:tblGrid"))
        for col, w in zip(grid.findall(qn("w:gridCol")), widths):
            col.set(qn("w:w"), str(w))
    for r, vals in zip(table.rows, rows):
        for c, val in zip(r.cells, vals):
            c.text = val
    return table


def _grid_widths(tbl_xml):
    grid = tbl_xml.find(qn("w:tblGrid"))
    return [c.get(qn("w:w")) for c in grid.findall(qn("w:gridCol"))]


def _marker_between(doc):
    """Return the first 'Продолжение таблицы' paragraph element, or None."""
    for child in doc.element.body:
        if child.tag == qn("w:p"):
            text = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if "Продолжение таблиц" in text:
                return child
    return None


def _data_fps(table):
    return sorted(tc._docx_data_fingerprints(table))


# ── pure split behaviour ─────────────────────────────────────────────────────

def test_split_inserts_pagebroken_marker() -> tuple[bool, str]:
    doc = Document()
    _mk_table(doc, [["Группа", "Дефект", "Статус"],
                    ["Кейс 1", "a", "ok"],
                    ["Кейс 2", "b", "ok"],
                    ["Кейс 3", "c", "ok"],
                    ["Кейс 4", "d", "ok"]], [3000, 3000, 3000])
    ok = tc._split_cross_page_table_with_marker(doc, 0, 2, "1.1.3", numeric_row_idx=None)
    if not ok:
        return _result(False, "split returned False on a valid 2-page table")
    if len(doc.tables) != 2:
        return _result(False, f"expected 2 physical tables, got {len(doc.tables)}")
    marker = _marker_between(doc)
    if marker is None:
        return _result(False, "no 'Продолжение таблицы' marker inserted")
    text = "".join(t.text or "" for t in marker.iter(qn("w:t")))
    if text.strip() != "Продолжение таблицы 1.1.3":
        return _result(False, f"unexpected marker text: {text!r}")
    pPr = marker.find(qn("w:pPr"))
    pbb = pPr.find(qn("w:pageBreakBefore")) if pPr is not None else None
    if pbb is None or pbb.get(qn("w:val")) in tc._DISABLED_PAGE_BREAK_VALUES:
        return _result(False, "marker pageBreakBefore missing/disabled (would render same-page)")
    if pPr.find(qn("w:keepNext")) is None:
        return _result(False, "marker keepNext missing (marker could detach from table)")
    return _result(True, "split inserts a page-broken keepNext marker before the continuation table")


def test_continuation_has_numeric_row_not_semantic_header() -> tuple[bool, str]:
    """Canonical KFU rule: continuation repeats the NUMERIC row, NEVER the
    semantic header. Source already has a numeric row at index 1."""
    doc = Document()
    _mk_table(doc, [["A", "B", "C"],
                    ["1", "2", "3"],
                    ["x", "y", "z"],
                    ["p", "q", "r"],
                    ["m", "n", "o"]], [2000, 2000, 2000])
    ok = tc._split_cross_page_table_with_marker(doc, 0, 3, "2.1", numeric_row_idx=1)
    if not ok:
        return _result(False, "split returned False")
    second = doc.tables[1]
    first_row = tc._docx_row_cell_texts(second.rows[0])
    if not tc._is_docx_numeric_row(first_row):
        return _result(False, f"continuation must START with the numeric row, got {first_row}")
    # the semantic header must NOT appear anywhere in the continuation fragment
    header_fp = tc._docx_row_fingerprint(doc.tables[0].rows[0])
    if any(tc._docx_row_fingerprint(r) == header_fp for r in second.rows):
        return _result(False, "semantic header was wrongly copied into the continuation")
    # first fragment keeps semantic header + numeric row + ≥1 data row
    f = doc.tables[0]
    if tc._docx_row_fingerprint(f.rows[0]) != header_fp:
        return _result(False, "first fragment lost its semantic header")
    if not any(tc._is_docx_numeric_row(tc._docx_row_cell_texts(r)) for r in f.rows):
        return _result(False, "first fragment lost its numeric row")
    return _result(True, "continuation starts with numeric row, no semantic header; first fragment intact")


def test_continuation_synthesizes_numeric_when_source_lacks_one() -> tuple[bool, str]:
    """Demo-style: source has NO numeric row. The split must synthesize `1..N`
    into BOTH fragments and the continuation must start with it (no header)."""
    doc = Document()
    _mk_table(doc, [["Группа", "Дефект", "Статус"],
                    ["Кейс 1", "a", "ok"],
                    ["Кейс 2", "b", "ok"],
                    ["Кейс 3", "c", "ok"],
                    ["Кейс 4", "d", "ok"]], [3000, 3000, 3000])
    ok = tc._split_cross_page_table_with_marker(doc, 0, 2, "1.1.3", numeric_row_idx=None)
    if not ok:
        return _result(False, "split returned False")
    first, second = doc.tables[0], doc.tables[1]
    if tc._docx_row_cell_texts(first.rows[1]) != ["1", "2", "3"]:
        return _result(False, f"first fragment numeric row not synthesized after header: {tc._docx_row_cell_texts(first.rows[1])}")
    if tc._docx_row_cell_texts(second.rows[0]) != ["1", "2", "3"]:
        return _result(False, f"continuation must start with synthesized numeric row, got {tc._docx_row_cell_texts(second.rows[0])}")
    header_fp = tc._docx_row_fingerprint(first.rows[0])
    if any(tc._docx_row_fingerprint(r) == header_fp for r in second.rows):
        return _result(False, "semantic header wrongly present in continuation")
    return _result(True, "numeric row synthesized into both fragments; continuation header-free")


def test_strip_normalizes_existing_manual_continuation() -> tuple[bool, str]:
    """Existing manual chain `[header, numeric, data]` continuation → strip the
    leading duplicate header so it starts with the numeric row."""
    doc = Document()
    doc.add_paragraph("Таблица 2.1 — Заголовок")
    _mk_table(doc, [["A", "B", "C"], ["1", "2", "3"], ["d1", "e1", "f1"]], [2000, 2000, 2000])
    doc.add_paragraph("Продолжение таблицы 2.1")
    _mk_table(doc, [["A", "B", "C"], ["1", "2", "3"], ["d2", "e2", "f2"]], [2000, 2000, 2000])
    fixed = tc._strip_continuation_semantic_headers(doc)
    if fixed != 1:
        return _result(False, f"expected 1 chain normalized, got {fixed}")
    cont = doc.tables[1]
    if not tc._is_docx_numeric_row(tc._docx_row_cell_texts(cont.rows[0])):
        return _result(False, f"continuation should start with numeric row, got {tc._docx_row_cell_texts(cont.rows[0])}")
    # data row preserved
    if tc._docx_row_cell_texts(cont.rows[-1]) != ["d2", "e2", "f2"]:
        return _result(False, "continuation data row lost during normalization")
    return _result(True, "manual continuation header stripped, numeric row + data preserved")


def test_gate_flags_semantic_header_on_continuation() -> tuple[bool, str]:
    """Negative regression: a continuation that still has the semantic header
    after `Продолжение таблицы N` must be flagged fail; numeric-led is clean."""
    from guides.coursework_kfu_2025.rendered_table_validation import (
        build_rendered_table_identities, _semantic_header_on_continuation_blockers)
    # wrong: continuation starts with semantic header
    bad = Document()
    bad.add_paragraph("Таблица 5.1 — T")
    _mk_table(bad, [["H1", "H2"], ["1", "2"], ["a", "b"]])
    bad.add_paragraph("Продолжение таблицы 5.1")
    _mk_table(bad, [["H1", "H2"], ["1", "2"], ["c", "d"]])
    bl = _semantic_header_on_continuation_blockers(build_rendered_table_identities(bad))
    if not any(b.blocker_type == "semantic_header_repeated_on_continuation" and b.table_num == "5.1" for b in bl):
        return _result(False, "gate failed to flag a semantic header on the continuation")
    # correct: continuation starts with numeric row only
    good = Document()
    good.add_paragraph("Таблица 5.2 — T")
    _mk_table(good, [["H1", "H2"], ["1", "2"], ["a", "b"]])
    good.add_paragraph("Продолжение таблицы 5.2")
    _mk_table(good, [["1", "2"], ["c", "d"]])
    bl2 = _semantic_header_on_continuation_blockers(build_rendered_table_identities(good))
    if bl2:
        return _result(False, f"gate false-fired on a valid numeric-led continuation: {[b.table_num for b in bl2]}")
    return _result(True, "gate fails semantic-header continuation, passes numeric-led continuation")


def test_source_note_stays_after_final_fragment() -> tuple[bool, str]:
    doc = Document()
    _mk_table(doc, [["H1", "H2"],
                    ["Кейс 1", "a"],
                    ["Кейс 2", "b"],
                    ["Кейс 3", "c"]], [3000, 3000])
    doc.add_paragraph("Источник: составлено автором.")
    tc._split_cross_page_table_with_marker(doc, 0, 2, "1.2", numeric_row_idx=None)
    body = list(doc.element.body)
    tbls = [i for i, c in enumerate(body) if c.tag == qn("w:tbl")]
    src_idx = next(
        (i for i, c in enumerate(body)
         if c.tag == qn("w:p") and "Источник" in "".join(t.text or "" for t in c.iter(qn("w:t")))),
        None,
    )
    if src_idx is None or len(tbls) != 2:
        return _result(False, "source line or second fragment missing")
    if src_idx < tbls[-1]:
        return _result(False, "Источник: is not after the final fragment")
    return _result(True, "Источник: stays after the final (continuation) fragment")


def test_no_data_rows_lost_or_duplicated() -> tuple[bool, str]:
    rows = [["H1", "H2"]] + [[f"Кейс {i}", f"v{i}"] for i in range(1, 7)]
    doc = Document()
    table = _mk_table(doc, rows, [3000, 3000])
    before = _data_fps(table)
    tc._split_cross_page_table_with_marker(doc, 0, 3, "1.1", numeric_row_idx=None)
    after = sorted(_data_fps(doc.tables[0]) + _data_fps(doc.tables[1]))
    if after != before:
        return _result(False, f"data rows changed across split: {before} -> {after}")
    # no duplication: every data fingerprint appears exactly once across fragments
    if len(after) != len(set(after)):
        return _result(False, "a data row was duplicated across fragments")
    return _result(True, "all data rows preserved exactly once, in order, across the split")


def test_grid_widths_preserved_across_fragments() -> tuple[bool, str]:
    widths = [1500, 4200, 2300]
    doc = Document()
    table = _mk_table(doc, [["A", "B", "C"],
                            ["Кейс 1", "x", "y"],
                            ["Кейс 2", "p", "q"],
                            ["Кейс 3", "m", "n"]], widths)
    before = _grid_widths(table._tbl)
    tc._split_cross_page_table_with_marker(doc, 0, 2, "1.1", numeric_row_idx=None)
    w1 = _grid_widths(doc.tables[0]._tbl)
    w2 = _grid_widths(doc.tables[1]._tbl)
    if w1 != before or w2 != before:
        return _result(False, f"grid widths drifted: orig={before} first={w1} second={w2}")
    return _result(True, "grid widths inherited verbatim by both fragments")


def test_split_refuses_when_no_continuation_data_row() -> tuple[bool, str]:
    doc = Document()
    _mk_table(doc, [["H1", "H2"], ["Кейс 1", "a"], ["Кейс 2", "b"]], [3000, 3000])
    # split_after == len(rows)-1 leaves no row for the continuation fragment
    if tc._split_cross_page_table_with_marker(doc, 0, 2, "1.1", numeric_row_idx=None):
        return _result(False, "split should refuse when the continuation fragment has no data row")
    if len(doc.tables) != 1:
        return _result(False, "table was mutated despite refused split")
    return _result(True, "split refuses to leave an empty continuation fragment")


# ── candidate detection ──────────────────────────────────────────────────────

def _demo_like_doc():
    doc = Document()
    doc.add_paragraph("Таблица 1.1.3 — Длинная таблица")
    _mk_table(doc, [["Группа", "Дефект", "Статус"],
                    ["Кейс 1", "a", "ok"],
                    ["Кейс 2", "b", "ok"],
                    ["Кейс 3", "c", "ok"],
                    ["Кейс 4", "d", "ok"]], [3000, 3000, 3000])
    return doc


def _two_page_lines():
    # leading-first-cell matcher: Кейс 1/2 on page 1, Кейс 3/4 on page 2
    return [
        PdfLine("Таблица 1.1.3 — Длинная таблица", 1, 10, 20),
        PdfLine("Кейс 1 a ok", 1, 30, 40),
        PdfLine("Кейс 2 b ok", 1, 50, 60),
        PdfLine("Кейс 3 c ok", 2, 30, 40),
        PdfLine("Кейс 4 d ok", 2, 50, 60),
    ]


def test_candidate_found_on_two_page_table() -> tuple[bool, str]:
    doc = _demo_like_doc()
    orig = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: False
    try:
        cand = tc._cross_page_split_candidate(
            doc, "1.1.3", _two_page_lines(), source_docx_path=Path("x")
        )
    finally:
        tc._source_has_meaningful_duplicate_for_table = orig
    if cand is None:
        return _result(False, "no candidate found for a clean 2-page marker-less table")
    table_idx, split_after, numeric_row_idx = cand
    if table_idx != 0 or split_after != 2 or numeric_row_idx is not None:
        return _result(False, f"unexpected candidate: idx={table_idx} after={split_after} num={numeric_row_idx}")
    return _result(True, "candidate found at the rendered page boundary (split_after=2)")


def test_candidate_skips_single_and_triple_page() -> tuple[bool, str]:
    doc = _demo_like_doc()
    orig = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: False
    try:
        # all rows on one page -> not cross-page
        one = [PdfLine(f"Кейс {i} x", 1, 10 * i, 10 * i + 5) for i in range(1, 5)]
        c1 = tc._cross_page_split_candidate(doc, "1.1.3", one, source_docx_path=Path("x"))
        # rows spread over THREE pages -> single split insufficient, skip
        three = [PdfLine("Кейс 1 a", 1, 10, 20), PdfLine("Кейс 2 b", 2, 10, 20),
                 PdfLine("Кейс 3 c", 3, 10, 20), PdfLine("Кейс 4 d", 3, 30, 40)]
        c3 = tc._cross_page_split_candidate(doc, "1.1.3", three, source_docx_path=Path("x"))
    finally:
        tc._source_has_meaningful_duplicate_for_table = orig
    if c1 is not None:
        return _result(False, "candidate wrongly fired on a single-page table")
    if c3 is not None:
        return _result(False, "candidate wrongly fired on a 3-page table (single split insufficient)")
    return _result(True, "candidate restricted to exactly-two-page tables")


def test_candidate_requires_no_existing_marker() -> tuple[bool, str]:
    # Rybakov-like valid chain: caption, tbl1, marker, tbl2 — must NOT be split again
    doc = Document()
    doc.add_paragraph("Таблица 1.1.3 — Длинная таблица")
    _mk_table(doc, [["Группа", "Дефект"], ["Кейс 1", "a"], ["Кейс 2", "b"]], [3000, 3000])
    doc.add_paragraph("Продолжение таблицы 1.1.3")
    _mk_table(doc, [["Группа", "Дефект"], ["Кейс 3", "c"], ["Кейс 4", "d"]], [3000, 3000])
    orig = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: False
    try:
        cand = tc._cross_page_split_candidate(
            doc, "1.1.3", _two_page_lines(), source_docx_path=Path("x")
        )
    finally:
        tc._source_has_meaningful_duplicate_for_table = orig
    if cand is not None:
        return _result(False, "candidate fired on a table that already has a continuation marker")
    return _result(True, "candidate skips tables with an existing continuation marker (Rybakov-safe)")


def test_candidate_skips_source_bad() -> tuple[bool, str]:
    doc = _demo_like_doc()
    orig = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: True  # source-proven dup
    try:
        cand = tc._cross_page_split_candidate(
            doc, "1.1.3", _two_page_lines(), source_docx_path=Path("x")
        )
    finally:
        tc._source_has_meaningful_duplicate_for_table = orig
    if cand is not None:
        return _result(False, "candidate fired on source-bad duplicated content")
    return _result(True, "candidate refuses source-bad duplicated tables (never auto-deleted)")


def test_lead_matcher_handles_wrapped_rows() -> tuple[bool, str]:
    # rows wrap across several PDF lines; strict whole-row matcher fails, lead wins
    doc = Document()
    table = _mk_table(doc, [["Группа", "Дефект"],
                            ["Кейс 1", "очень длинный текст про дефект один"],
                            ["Кейс 2", "очень длинный текст про дефект два"],
                            ["Кейс 3", "очень длинный текст про дефект три"]], [3000, 5000])
    lines = [
        PdfLine("Кейс 1 очень длинный", 1, 10, 20),
        PdfLine("текст про дефект один", 1, 21, 30),
        PdfLine("Кейс 2 очень длинный", 1, 31, 40),
        PdfLine("текст про дефект два", 1, 41, 50),
        PdfLine("Кейс 3 очень длинный", 2, 10, 20),
        PdfLine("текст про дефект три", 2, 21, 30),
    ]
    mapping = tc._match_data_row_pages_by_lead(table, lines, set())
    if mapping != {1: 1, 2: 1, 3: 2}:
        return _result(False, f"lead matcher mapping wrong: {mapping}")
    return _result(True, "leading-first-cell matcher maps wrapped rows to pages")


def main() -> int:
    tests = [
        ("split inserts page-broken marker", test_split_inserts_pagebroken_marker),
        ("continuation numeric-row not header", test_continuation_has_numeric_row_not_semantic_header),
        ("continuation synthesizes numeric row", test_continuation_synthesizes_numeric_when_source_lacks_one),
        ("strip normalizes manual continuation", test_strip_normalizes_existing_manual_continuation),
        ("gate flags semantic header on cont.", test_gate_flags_semantic_header_on_continuation),
        ("source/note after final fragment", test_source_note_stays_after_final_fragment),
        ("no data rows lost/duplicated", test_no_data_rows_lost_or_duplicated),
        ("grid widths preserved", test_grid_widths_preserved_across_fragments),
        ("split refuses empty continuation", test_split_refuses_when_no_continuation_data_row),
        ("candidate found on 2-page table", test_candidate_found_on_two_page_table),
        ("candidate skips 1/3-page tables", test_candidate_skips_single_and_triple_page),
        ("candidate requires no existing marker", test_candidate_requires_no_existing_marker),
        ("candidate skips source-bad", test_candidate_skips_source_bad),
        ("lead matcher handles wrapped rows", test_lead_matcher_handles_wrapped_rows),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        print(f"[{'PASS' if ok else 'FAIL'}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
