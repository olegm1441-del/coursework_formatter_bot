from __future__ import annotations

import inspect
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml.ns import qn

from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine
from guides.coursework_kfu_2025 import formatter_service
import guides.coursework_kfu_2025.table_continuation as tc


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def _add_orphan_candidate_doc() -> Document:
    doc = Document()
    doc.add_paragraph("Перед таблицей текст.")
    doc.add_paragraph("Таблица 1.3.1")
    doc.add_paragraph("Сравнение моделей корпоративного управления")
    table = doc.add_table(rows=3, cols=4)
    values = [
        ["Страна/модель", "Роль государства", "Роль операторов", "Выводы"],
        ["1", "2", "3", "4"],
        ["Россия", "регулирование", "платформы", "координация"],
    ]
    for row, row_values in zip(table.rows, values):
        for cell, value in zip(row.cells, row_values):
            cell.text = value
    doc.add_paragraph("Источник: составлено автором.")
    return doc


def _blank_count_before_caption(doc: Document, caption_text: str) -> int:
    children = list(doc.element.body)
    caption_idx = None
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t"))).strip()
        if text == caption_text:
            caption_idx = idx
            break
    if caption_idx is None:
        raise AssertionError(f"caption not found: {caption_text!r}")

    count = 0
    for child in reversed(children[:caption_idx]):
        if child.tag != qn("w:p"):
            break
        text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t"))).strip()
        if text:
            break
        count += 1
    return count


def _caption_has_page_break(doc: Document, caption_text: str) -> bool:
    para = next(p for p in doc.paragraphs if p.text == caption_text)
    p_pr = para._element.find(qn("w:pPr"))
    if p_pr is None:
        return False
    page_break = p_pr.find(qn("w:pageBreakBefore"))
    if page_break is None:
        return False
    return page_break.get(qn("w:val")) not in {"0", "false", "False", "off"}


def _source_note_after_table(doc: Document) -> bool:
    children = list(doc.element.body)
    table_idx = next((idx for idx, child in enumerate(children) if child.tag == qn("w:tbl")), None)
    source_idx = None
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t"))).strip()
        if text.startswith("Источник:"):
            source_idx = idx
            break
    return table_idx is not None and source_idx is not None and source_idx > table_idx


def _orphan_lines() -> list[PdfLine]:
    return [
        PdfLine("Таблица 1.3.1", 21, 646.0, 658.0),
        PdfLine("Сравнение моделей корпоративного управления", 21, 668.0, 680.0),
        PdfLine("Страна/модель Роль государства Роль операторов Выводы", 21, 692.0, 704.0),
        PdfLine("1 2 3 4", 21, 716.0, 728.0),
        PdfLine("Россия регулирование платформы координация", 22, 60.0, 72.0),
    ]


def _repaired_lines() -> list[PdfLine]:
    return [
        PdfLine("Таблица 1.3.1", 22, 58.0, 70.0),
        PdfLine("Сравнение моделей корпоративного управления", 22, 82.0, 94.0),
        PdfLine("Страна/модель Роль государства Роль операторов Выводы", 22, 106.0, 118.0),
        PdfLine("1 2 3 4", 22, 124.0, 136.0),
        PdfLine("Россия регулирование платформы координация", 22, 148.0, 160.0),
    ]


def _run_final_guard(lines: list[PdfLine], after_lines: list[PdfLine]):
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "candidate.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "candidate.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        _add_orphan_candidate_doc().save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            calls = {"n": 0}

            def fake_analyze(_path):
                calls["n"] += 1
                return lines if calls["n"] == 1 else after_lines

            tc.analyze_pdf_lines = fake_analyze
            changed = tc.apply_rendered_table_start_orphan_guard(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        return changed, Document(str(path))


def test_post_padding_orphan_guard_moves_table_with_two_blanks() -> tuple[bool, str]:
    changed, reread = _run_final_guard(_orphan_lines(), _repaired_lines())
    blanks = _blank_count_before_caption(reread, "Таблица 1.3.1")
    return _result(
        changed == 1
        and blanks == 2
        and not _caption_has_page_break(reread, "Таблица 1.3.1")
        and _source_note_after_table(reread)
        and len(reread.tables) == 1,
        f"changed={changed}, blanks={blanks}, page_break={_caption_has_page_break(reread, 'Таблица 1.3.1')}, "
        f"tables={len(reread.tables)}, source_after={_source_note_after_table(reread)}",
    )


def test_post_padding_orphan_guard_rolls_back_when_orphan_remains() -> tuple[bool, str]:
    changed, reread = _run_final_guard(_orphan_lines(), _orphan_lines())
    blanks = _blank_count_before_caption(reread, "Таблица 1.3.1")
    return _result(
        changed == 0 and blanks == 0,
        f"changed={changed}, blanks={blanks}",
    )


def test_formatter_runs_final_orphan_guard_after_same_page_normalizers() -> tuple[bool, str]:
    source = inspect.getsource(formatter_service.format_docx)
    compatible_idx = source.find("normalize_compatible_grid_same_page_repeated_fragments_inplace")
    guard_idx = source.find("apply_rendered_table_start_orphan_guard")
    final_validation_idx = source.find("rendered_violations = _rendered_continuation_violations_for_docx")
    return _result(
        compatible_idx != -1
        and guard_idx != -1
        and final_validation_idx != -1
        and compatible_idx < guard_idx < final_validation_idx,
        f"compatible_idx={compatible_idx}, guard_idx={guard_idx}, final_validation_idx={final_validation_idx}",
    )


def main() -> int:
    tests = [
        ("post-padding orphan guard two blanks", test_post_padding_orphan_guard_moves_table_with_two_blanks),
        ("post-padding orphan guard rollback", test_post_padding_orphan_guard_rolls_back_when_orphan_remains),
        ("formatter guard ordering", test_formatter_runs_final_orphan_guard_after_same_page_normalizers),
    ]
    failed = 0
    for name, test in tests:
        ok, message = test()
        if ok:
            print(f"[PASS] {name} — {message}")
        else:
            failed += 1
            print(f"[FAIL] {name} — {message}")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
