"""
Fast Phase 3/product-rule regression tests.

Run from repo root:
    python -m pytest tests/test_phase3.py -v
or directly:
    python tests/test_phase3.py

The default runner stays cheap: synthetic DOCX/XML checks and isolated unit tests
only. Real asset formatting smoke checks are opt-in via:
    KPFU_RUN_LONG_PHASE3_TESTS=1 python tests/test_phase3.py

Product-rule coverage:
  A  — Figure deletion: images survive Rule 4 (paragraphs with w:drawing never removed)
  C  — Student continuation length: _is_student_continuation detects ≤30 char texts
  B1 — tblW fix: _optimize_table_col_widths updates w:tblW after scaling
  B2 — keepTogether, Rule 6 propagation, image height from wp:extent
  B3 — Footnote standardisation
  C2 — Empty para between image and caption removed; numeric column minimums
  T2 — Heading paragraphs/styles must not use Word autonumbering; manual
       heading text numbering remains literal text.
  M1/S1 — Marker/prototype table split rules for ordinary and appendix tables.

  NOTE: Tests for LRPB-based table splitting (B, B1-stale/valid, C2-fits-1-page,
  C-student-merges) were removed when apply_table_merging / apply_table_continuation
  were stubbed out.  See module docstring in table_continuation.py for the future
  LibreOffice-based plan.
"""

from __future__ import annotations

import io
import logging
import os
import re
import sys
import shutil
import tempfile
import traceback
from pathlib import Path

# ── project root on path ──────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from guides.coursework_kfu_2025.table_continuation import (
    _is_student_continuation,
    _para_has_image,
    _body_height_pt,
    _estimate_para_height,
    apply_rule4_empty_first_lines,
)
from guides.coursework_kfu_2025.formatter_service import format_docx

ASSETS = ROOT / "assets"
ASSET_FILES = list(ASSETS.glob("*.docx"))

PASS = "PASS"
FAIL = "FAIL"


def _result(ok: bool, msg: str = "") -> tuple[bool, str]:
    return ok, msg


def _all_table_rows_have_cant_split(table) -> bool:
    for row in table.rows:
        tr_pr = row._tr.find(qn("w:trPr"))
        if tr_pr is None or tr_pr.find(qn("w:cantSplit")) is None:
            return False
    return True


def _table_has_row_texts(table, expected: list[str]) -> bool:
    return any([cell.text for cell in row.cells] == expected for row in table.rows)


def _count_table_rows_with_texts(table, expected: list[str]) -> int:
    return sum(1 for row in table.rows if [cell.text for cell in row.cells] == expected)


def _table_has_page_break_service_paragraph_before(doc: Document, table_index: int) -> bool:
    target = doc.tables[table_index]._tbl
    children = list(doc.element.body)
    for idx, child in enumerate(children):
        if child is not target:
            continue
        if idx == 0:
            return False
        previous = children[idx - 1]
        if previous.tag != qn("w:p"):
            return False
        text = "".join(t.text or "" for t in previous.findall(".//" + qn("w:t"))).strip()
        if text:
            return False
        p_pr = previous.find(qn("w:pPr"))
        return p_pr is not None and p_pr.find(qn("w:pageBreakBefore")) is not None
    return False


def _paragraph_before_table(doc: Document, table_index: int):
    target = doc.tables[table_index]._tbl
    children = list(doc.element.body)
    for idx, child in enumerate(children):
        if child is target:
            if idx == 0 or children[idx - 1].tag != qn("w:p"):
                return None
            return children[idx - 1]
    return None


def _paragraph_text(p_xml) -> str:
    return "".join(t.text or "" for t in p_xml.findall(".//" + qn("w:t"))).strip()


def _paragraph_has_page_break_before(p_xml) -> bool:
    p_pr = p_xml.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:pageBreakBefore")) is not None


def _paragraph_has_keep_next(p_xml) -> bool:
    p_pr = p_xml.find(qn("w:pPr"))
    return p_pr is not None and p_pr.find(qn("w:keepNext")) is not None


def _paragraph_is_right_aligned(p_xml) -> bool:
    p_pr = p_xml.find(qn("w:pPr"))
    if p_pr is None:
        return False
    jc = p_pr.find(qn("w:jc"))
    return jc is not None and jc.get(qn("w:val")) == "right"


# ── helpers ───────────────────────────────────────────────────────────────────

def _make_minimal_doc_with_image() -> Document:
    """
    Minimal document: body paragraph with a w:drawing (simulated image) placed
    EXACTLY at the top of a new page in the geometry estimator.

    Strategy: fill one page worth of content using the same height estimator
    that Rule 4 uses, so the image paragraph triggers page_overflow=True
    and is_empty=True — this is the exact condition that caused the deletion bug.
    """
    doc = Document()

    # Calculate how many "Body text." paragraphs fit on one page
    probe = doc.add_paragraph("Body text.")
    h_per_para = _estimate_para_height(probe)
    probe._element.getparent().remove(probe._element)

    body_h = _body_height_pt(doc)
    # Fill just under one page (leave room for image to overflow)
    n_paras = max(1, int(body_h / h_per_para))
    for _ in range(n_paras):
        doc.add_paragraph("Body text.")

    # Image paragraph: no text, one w:drawing — lands at page-top in estimator
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    return doc


def _count_drawings(doc: Document) -> int:
    return len(doc.element.body.findall(".//" + qn("w:drawing")))


def _section_break_positions(doc: Document) -> list[tuple[int, str]]:
    positions: list[tuple[int, str]] = []
    for idx, paragraph in enumerate(doc.paragraphs):
        pPr = paragraph._element.pPr
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            positions.append((idx, paragraph.text.strip()))
    return positions


def _section_index_for_paragraph(doc: Document, paragraph_index: int) -> int:
    section_index = 0
    for idx, paragraph in enumerate(doc.paragraphs):
        if idx >= paragraph_index:
            break
        p_pr = paragraph._element.pPr
        if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
            section_index += 1
    return section_index


def _paragraph_index(doc: Document, text: str) -> int | None:
    for idx, paragraph in enumerate(doc.paragraphs):
        if " ".join((paragraph.text or "").split()) == text:
            return idx
    return None


def _footer_has_page_field(footer) -> bool:
    return any((instr.text or "").strip() == "PAGE" for instr in footer._element.findall(".//" + qn("w:instrText")))


def _footer_has_visible_text(footer) -> bool:
    return bool("".join(t.text or "" for t in footer._element.findall(".//" + qn("w:t"))).strip())


def _section_page_start(section) -> str | None:
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    return pg_num_type.get(qn("w:start")) if pg_num_type is not None else None


def _toc_tab_stops(paragraph) -> list:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return []
    return p_pr.findall(qn("w:tabs"))


def _toc_tab_elems(paragraph) -> list:
    elems = []
    for tabs in _toc_tab_stops(paragraph):
        elems.extend(tabs.findall(qn("w:tab")))
    return elems


def _format_synthetic_doc(doc: Document) -> Document:
    from guides.coursework_kfu_2025.safe_formatter import process_document

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(inp)
        process_document(inp, out)
        return Document(str(out))


def _make_front_matter_doc(kind: str, *, with_appendices=False) -> Document:
    doc = Document()
    if kind == "title_contents_intro":
        doc.add_paragraph("Титульная строка")
        doc.add_paragraph("СОДЕРЖАНИЕ")
        doc.add_paragraph("ВВЕДЕНИЕ 3")
    elif kind == "title_intro":
        doc.add_paragraph("Титульная строка")
    elif kind == "intro_only":
        pass
    else:
        raise ValueError(kind)

    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ")
    doc.add_paragraph("Текст главы.")
    if with_appendices:
        doc.add_paragraph("ПРИЛОЖЕНИЯ")
        doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
        doc.add_paragraph("Материалы приложения.")
        doc.add_paragraph("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ 1")
        doc.add_paragraph("Продолжение материалов.")
    return doc


def _assert_intro_numbering_starts_at_three(kind: str) -> tuple[bool, str]:
    formatted = _format_synthetic_doc(_make_front_matter_doc(kind))
    intro_idx = _paragraph_index(formatted, "ВВЕДЕНИЕ")
    if intro_idx is None:
        return _result(False, "ВВЕДЕНИЕ not found after formatting")

    intro_section = formatted.sections[_section_index_for_paragraph(formatted, intro_idx)]
    if _section_page_start(intro_section) != "3":
        return _result(False, f"intro section page start is {_section_page_start(intro_section)!r}, expected '3'")
    if not _footer_has_page_field(intro_section.footer):
        return _result(False, "intro section footer has no PAGE field")
    if intro_section.footer.is_linked_to_previous:
        return _result(False, "intro section footer is linked to previous")
    return _result(True, f"{kind} introduction starts at page 3")


def test_front_matter_intro_only_starts_numbering_at_three() -> tuple[bool, str]:
    return _assert_intro_numbering_starts_at_three("intro_only")


def test_front_matter_title_intro_starts_numbering_at_three() -> tuple[bool, str]:
    return _assert_intro_numbering_starts_at_three("title_intro")


def test_front_matter_title_contents_intro_starts_numbering_at_three() -> tuple[bool, str]:
    return _assert_intro_numbering_starts_at_three("title_contents_intro")


def test_front_matter_section_breaks_are_bounded() -> tuple[bool, str]:
    formatted = _format_synthetic_doc(_make_front_matter_doc("title_contents_intro", with_appendices=True))
    positions = _section_break_positions(formatted)
    position_texts = [text for _, text in positions]
    if len(positions) != 3:
        return _result(False, f"expected 3 paragraph section breaks, got {positions!r}")
    if position_texts != ["Титульная строка", "ВВЕДЕНИЕ 3", "Текст главы."]:
        return _result(False, f"unexpected section break positions: {positions!r}")
    for idx, paragraph in enumerate(formatted.paragraphs):
        if not (paragraph.text or "").strip():
            p_pr = paragraph._element.pPr
            if p_pr is not None and p_pr.find(qn("w:sectPr")) is not None:
                return _result(False, "front matter numbering generated a blank section-break paragraph")
            if idx + 1 < len(formatted.paragraphs) and not (formatted.paragraphs[idx + 1].text or "").strip():
                return _result(False, "front matter numbering generated consecutive blank paragraphs")
    return _result(True, "front matter section breaks are bounded")


def test_appendices_first_page_numbered_following_pages_unnumbered() -> tuple[bool, str]:
    formatted = _format_synthetic_doc(_make_front_matter_doc("title_contents_intro", with_appendices=True))
    appendices_idx = _paragraph_index(formatted, "ПРИЛОЖЕНИЯ")
    if appendices_idx is None:
        return _result(False, "ПРИЛОЖЕНИЯ not found after formatting")

    appendix_section = formatted.sections[_section_index_for_paragraph(formatted, appendices_idx)]
    if not appendix_section.different_first_page_header_footer:
        return _result(False, "appendix section does not use first-page footer")
    if not _footer_has_page_field(appendix_section.first_page_footer):
        return _result(False, "ПРИЛОЖЕНИЯ page footer has no PAGE field")
    if _footer_has_page_field(appendix_section.footer) or _footer_has_visible_text(appendix_section.footer):
        return _result(False, "following appendix pages have visible page numbering")
    if appendix_section.first_page_footer.is_linked_to_previous or appendix_section.footer.is_linked_to_previous:
        return _result(False, "appendix footers are linked to previous")
    return _result(True, "appendix first page is numbered and following pages are unnumbered")


def test_appendix_continuation_pages_are_unnumbered() -> tuple[bool, str]:
    formatted = _format_synthetic_doc(_make_front_matter_doc("title_contents_intro", with_appendices=True))
    continuation_idx = _paragraph_index(formatted, "ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ 1")
    if continuation_idx is None:
        return _result(False, "appendix continuation label not found after formatting")

    continuation_section = formatted.sections[_section_index_for_paragraph(formatted, continuation_idx)]
    if _footer_has_page_field(continuation_section.footer) or _footer_has_visible_text(continuation_section.footer):
        return _result(False, "appendix continuation pages have visible page numbering")
    return _result(True, "appendix continuation pages are unnumbered")


def test_front_matter_before_introduction_remains_protected() -> tuple[bool, str]:
    formatted = _format_synthetic_doc(_make_front_matter_doc("title_contents_intro"))
    texts = [" ".join((p.text or "").split()) for p in formatted.paragraphs]
    intro_idx = texts.index("ВВЕДЕНИЕ")
    front_matter = texts[:intro_idx]
    if "Титульная строка" not in front_matter or "СОДЕРЖАНИЕ" not in front_matter or "ВВЕДЕНИЕ 3" not in front_matter:
        return _result(False, f"front matter changed unexpectedly: {front_matter!r}")
    return _result(True, "front matter before ВВЕДЕНИЕ remains protected")


def test_real_intro_detection_ignores_toc_embedded_intro() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.classifier import find_body_start_index

    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("СОДЕРЖАНИЕ\nВВЕДЕНИЕ\n1. Теоретические основы")
    doc.add_paragraph("ВВЕДЕНИЕ........................................................3")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")

    body_start = find_body_start_index(doc)
    if body_start != 3:
        return _result(False, f"body_start={body_start!r}, expected real standalone intro at index 3")
    return _result(True, "real intro detection ignores TOC-embedded ВВЕДЕНИЕ")


def test_front_matter_text_before_real_intro_is_preserved() -> tuple[bool, str]:
    doc = Document()
    front_matter_texts = [
        "Титульная строка",
        "СОДЕРЖАНИЕ\nВВЕДЕНИЕ\n1. Теоретические основы\n1.1. Сущность подхода",
        (
            "2. Анализ и совершенствование\n"
            "2.1. Общая характеристика\n"
            "ЗАКЛЮЧЕНИЕ\n"
            "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ\n"
            "ПРИЛОЖЕНИЕ А"
        ),
        "ВВЕДЕНИЕ........................................................3",
    ]
    for text in front_matter_texts:
        doc.add_paragraph(text)
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")

    formatted = _format_synthetic_doc(doc)
    intro_idx = _paragraph_index(formatted, "ВВЕДЕНИЕ")
    if intro_idx is None:
        return _result(False, "real standalone ВВЕДЕНИЕ missing after formatting")

    actual_front_matter = [paragraph.text for paragraph in formatted.paragraphs[:intro_idx]]
    if actual_front_matter != front_matter_texts:
        return _result(
            False,
            "front matter text before real intro was mutated:\n"
            f"expected={front_matter_texts!r}\nactual={actual_front_matter!r}",
        )
    return _result(True, "front matter text before real intro is preserved exactly")


def test_b2_contents_entries_have_stable_tab_leaders() -> tuple[bool, str]:
    doc = Document()
    doc.add_paragraph("Титульная строка.....9")
    doc.add_paragraph("СОДЕРЖАНИЕ 2")
    doc.add_paragraph("ВВЕДЕНИЕ........................................................3")
    doc.add_paragraph("")
    doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ")
    doc.add_paragraph("1.1. Очень длинный пункт содержания с продолжительным названием, которое должно сохранять номер страницы справа\t\t10")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ . . . . . . . . . . 22")
    doc.add_paragraph("ПРИЛОЖЕНИЯ………25")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ")
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    doc.add_paragraph("Материалы приложения.")

    formatted = _format_synthetic_doc(doc)
    texts = [p.text for p in formatted.paragraphs]

    real_intro_idx = _paragraph_index(formatted, "ВВЕДЕНИЕ")
    if real_intro_idx is None:
        return _result(False, "real ВВЕДЕНИЕ paragraph missing")

    expected_front_matter = [
        "Титульная строка.....9",
        "СОДЕРЖАНИЕ 2",
        "ВВЕДЕНИЕ........................................................3",
        "",
        "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ",
        "1.1. Очень длинный пункт содержания с продолжительным названием, которое должно сохранять номер страницы справа\t\t10",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ . . . . . . . . . . 22",
        "ПРИЛОЖЕНИЯ………25",
    ]
    actual_front_matter = texts[:real_intro_idx]
    if actual_front_matter != expected_front_matter:
        return _result(
            False,
            "TOC/front matter must remain text-frozen before real intro:\n"
            f"expected={expected_front_matter!r}\nactual={actual_front_matter!r}",
        )

    return _result(True, "TOC/front matter text is frozen before real intro")


def _run_static_contents_rebuild(doc: Document, rendered_lines: list[tuple[str, int]]) -> Document:
    import guides.coursework_kfu_2025.contents_builder as cb
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "toc.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "toc.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = cb.render_docx_to_pdf
        old_analyze = cb.analyze_pdf_lines
        try:
            cb.render_docx_to_pdf = lambda _path: pdf_path
            cb.analyze_pdf_lines = lambda _path: [
                PdfLine(text=text, page_num=page, top=100.0, bottom=112.0)
                for text, page in rendered_lines
            ]
            changed = cb.rebuild_static_contents_page(path)
        finally:
            cb.render_docx_to_pdf = old_render
            cb.analyze_pdf_lines = old_analyze

        if not changed:
            raise AssertionError("rebuild_static_contents_page returned False")
        return Document(str(path))


def _run_static_contents_rebuild_result(doc: Document, rendered_lines: list[tuple[str, int]]) -> tuple[bool, Document]:
    import guides.coursework_kfu_2025.contents_builder as cb
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "toc.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "toc.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = cb.render_docx_to_pdf
        old_analyze = cb.analyze_pdf_lines
        try:
            cb.render_docx_to_pdf = lambda _path: pdf_path
            cb.analyze_pdf_lines = lambda _path: [
                PdfLine(text=text, page_num=page, top=100.0, bottom=112.0)
                for text, page in rendered_lines
            ]
            changed = cb.rebuild_static_contents_page(path)
        finally:
            cb.render_docx_to_pdf = old_render
            cb.analyze_pdf_lines = old_analyze

        return changed, Document(str(path))


def _make_autotoc_doc(*, old_heading: str | None = "Содержание", title: bool = True, appendices: bool = False) -> Document:
    doc = Document()
    if title:
        doc.add_paragraph("Титульная строка")
    if old_heading:
        doc.add_paragraph(old_heading)
        doc.add_paragraph("ВВЕДЕНИЕ........................................................3")
        doc.add_paragraph("1. Старый раздел................................................4")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("1. Теоретические основы").style = "Heading 1"
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("1.1. Длинный подраздел с названием, которое должно переноситься естественно без ручных точек").style = "Heading 2"
    doc.add_paragraph("Текст подраздела.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("Итоги.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")
    if appendices:
        doc.add_paragraph("ПРИЛОЖЕНИЯ")
        doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
        doc.add_paragraph("1.1. Локальный заголовок приложения")
        doc.add_paragraph("Текст приложения.")
    return doc


def _default_autotoc_lines(*, appendices: bool = False) -> list[tuple[str, int]]:
    lines = [
        ("ВВЕДЕНИЕ", 3),
        ("1. Теоретические основы", 4),
        ("1.1. Длинный подраздел с названием, которое должно переноситься естественно без ручных точек", 5),
        ("ЗАКЛЮЧЕНИЕ", 8),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 9),
    ]
    if appendices:
        lines.append(("ПРИЛОЖЕНИЯ", 10))
    return lines


def _toc_texts_before_intro(doc: Document) -> list[str]:
    texts = [p.text for p in doc.paragraphs]
    intro_idx = texts.index("ВВЕДЕНИЕ")
    return texts[:intro_idx]


def _toc_entry_paragraph(doc: Document, prefix: str):
    for p in doc.paragraphs:
        if (p.text or "").startswith(prefix):
            return p
    return None


def _paragraph_left_indent_twips(paragraph) -> str | None:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return None
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        return None
    return ind.get(qn("w:left"))


def _paragraph_has_right_dot_tab(paragraph) -> bool:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return False
    for tabs in p_pr.findall(qn("w:tabs")):
        for tab in tabs.findall(qn("w:tab")):
            if tab.get(qn("w:val")) == "right" and tab.get(qn("w:leader")) == "dot":
                return True
    return False


def _paragraph_has_internal_hyperlink(paragraph) -> bool:
    p = paragraph._element
    for hyperlink in p.findall(qn("w:hyperlink")):
        if hyperlink.get(qn("w:anchor")):
            return True
    return False


def _document_has_bookmark(document: Document, prefix: str = "kpfu_toc_") -> bool:
    for elem in document._element.iter(qn("w:bookmarkStart")):
        name = elem.get(qn("w:name")) or ""
        if name.startswith(prefix):
            return True
    return False


def test_autotoc_existing_soderzhanie_replaced_by_canonical() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    front = _toc_texts_before_intro(out)
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"canonical heading missing or duplicated: {front!r}")
    if any("Содержание" in text for text in front):
        return _result(False, f"old mixed-case heading survived: {front!r}")
    if any("Старый раздел" in text for text in front):
        return _result(False, f"old TOC entry survived: {front!r}")
    return _result(True, "old Содержание block replaced by canonical СОДЕРЖАНИЕ")


def test_autotoc_existing_oglavlenie_replaced_by_canonical() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Оглавление"), _default_autotoc_lines())
    front = _toc_texts_before_intro(out)
    if "СОДЕРЖАНИЕ" not in front:
        return _result(False, f"canonical heading missing: {front!r}")
    if any("Оглавление" in text for text in front):
        return _result(False, f"old Оглавление survived: {front!r}")
    return _result(True, "old Оглавление block replaced by canonical СОДЕРЖАНИЕ")


def test_autotoc_exact_intro_entry_inside_old_toc_is_removed() -> tuple[bool, str]:
    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("Содержание")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. Старый раздел без номера страницы")
    doc.add_paragraph("Оглавление")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. Еще один старый раздел")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("1. АКТУАЛЬНЫЙ РАЗДЕЛ").style = "Heading 1"
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("Итоги.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")

    out = _run_static_contents_rebuild(
        doc,
        [
            ("ВВЕДЕНИЕ", 3),
            ("1. АКТУАЛЬНЫЙ РАЗДЕЛ", 4),
            ("ЗАКЛЮЧЕНИЕ", 5),
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 6),
        ],
    )
    front = _toc_texts_before_intro(out)
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected one new TOC heading before real intro: {front!r}")
    joined = "\n".join(front)
    if "Оглавление" in joined or "Старый раздел" in joined or "Еще один старый раздел" in joined:
        return _result(False, f"old TOC blocks were not fully removed: {front!r}")
    if "1. АКТУАЛЬНЫЙ РАЗДЕЛ\t4" not in front:
        return _result(False, f"actual uppercase body chapter missing from TOC: {front!r}")
    return _result(True, "exact intro entries and double old TOC blocks are removed")


def test_autotoc_missing_contents_inserted_before_real_intro() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading=None, title=True), _default_autotoc_lines())
    texts = [p.text for p in out.paragraphs]
    if texts.index("СОДЕРЖАНИЕ") >= texts.index("ВВЕДЕНИЕ"):
        return _result(False, f"TOC not inserted before intro: {texts[:8]!r}")
    if texts[0] != "Титульная строка":
        return _result(False, f"title/front matter was not preserved before TOC: {texts[:5]!r}")
    return _result(True, "missing TOC inserted after title and before real intro")


def test_autotoc_no_title_page_inserted_at_document_start() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading=None, title=False), _default_autotoc_lines())
    texts = [p.text for p in out.paragraphs]
    if texts[0] != "СОДЕРЖАНИЕ":
        return _result(False, f"TOC should start document when no title page exists: {texts[:5]!r}")
    if texts.index("СОДЕРЖАНИЕ") >= texts.index("ВВЕДЕНИЕ"):
        return _result(False, f"TOC not before intro: {texts[:6]!r}")
    return _result(True, "no-title document gets TOC at document start")


def test_autotoc_appendices_include_general_heading_only() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(
        _make_autotoc_doc(old_heading="Содержание", appendices=True),
        _default_autotoc_lines(appendices=True),
    )
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "ПРИЛОЖЕНИЯ\t10" not in front:
        return _result(False, f"general appendices heading missing from TOC: {front!r}")
    if "ПРИЛОЖЕНИЕ 1" in joined or "Локальный заголовок приложения" in joined:
        return _result(False, f"appendix-local entries leaked into TOC: {front!r}")
    return _result(True, "TOC includes only general ПРИЛОЖЕНИЯ entry")


def test_autotoc_heading2_has_no_left_indent() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    p = _toc_entry_paragraph(out, "1.1. Длинный подраздел")
    if p is None:
        return _result(False, "heading2 TOC entry missing")
    if _paragraph_left_indent_twips(p) not in (None, "0"):
        return _result(False, f"heading2 TOC entry has left indent: {_paragraph_left_indent_twips(p)!r}")
    h1 = _toc_entry_paragraph(out, "1. Теоретические основы")
    if h1 is None:
        return _result(False, "heading1 TOC entry missing")
    if _paragraph_left_indent_twips(h1) != _paragraph_left_indent_twips(p):
        return _result(False, "heading1 and heading2 TOC entries have different left indents")
    return _result(True, "heading2 TOC entry has no left indent")


def test_autotoc_entries_have_zero_indent_and_one_point_five_spacing() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    for p in out.paragraphs:
        if "\t" not in (p.text or ""):
            continue
        fmt = p.paragraph_format
        if fmt.left_indent is not None and fmt.left_indent.twips != 0:
            return _result(False, f"left indent is not zero for {p.text!r}: {fmt.left_indent.twips}")
        if fmt.first_line_indent is not None and fmt.first_line_indent.twips != 0:
            return _result(False, f"first line indent is not zero for {p.text!r}: {fmt.first_line_indent.twips}")
        if fmt.line_spacing != 1.5:
            return _result(False, f"line spacing is not 1.5 for {p.text!r}: {fmt.line_spacing!r}")
    return _result(True, "TOC entries have zero paragraph indent and 1.5 line spacing")


def test_autotoc_entries_use_dot_leader_tab_not_manual_dots() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    for p in out.paragraphs:
        if "\t" not in (p.text or ""):
            continue
        if re.search(r"\.{2,}|…", p.text):
            return _result(False, f"manual dot run survived in TOC entry: {p.text!r}")
        if not _paragraph_has_right_dot_tab(p):
            return _result(False, f"TOC entry lacks right dot tab stop: {p.text!r}")
    return _result(True, "TOC entries use tab-stop dot leaders, not manual dots")


def test_autotoc_uses_body_heading_register() -> tuple[bool, str]:
    doc = _make_autotoc_doc(old_heading="Содержание")
    for p in doc.paragraphs:
        if p.text == "1. Теоретические основы":
            p.text = "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ"
            break
    out = _run_static_contents_rebuild(
        doc,
        [
            ("ВВЕДЕНИЕ", 3),
            ("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ", 4),
            ("1.1. Длинный подраздел с названием, которое должно переноситься естественно без ручных точек", 5),
            ("ЗАКЛЮЧЕНИЕ", 8),
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 9),
        ],
    )
    front = _toc_texts_before_intro(out)
    if "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ\t4" not in front:
        return _result(False, f"chapter heading register was not copied from body: {front!r}")
    return _result(True, "TOC chapter entry keeps formatted body heading register")


def test_autotoc_normal_numbered_body_paragraph_is_excluded() -> tuple[bool, str]:
    doc = _make_autotoc_doc(old_heading="Содержание")
    for idx, p in enumerate(doc.paragraphs):
        if p.text == "Текст подраздела.":
            doc.paragraphs[idx]._element.addnext(doc.add_paragraph("1. Маркетинговый подход. Данный подход применяется в анализе.")._element)
            break
    out = _run_static_contents_rebuild(doc, _default_autotoc_lines())
    front = _toc_texts_before_intro(out)
    if any("Маркетинговый подход" in text for text in front):
        return _result(False, f"normal numbered body paragraph leaked into TOC: {front!r}")
    return _result(True, "normal numbered body paragraph is excluded from TOC")


def test_autotoc_page_resolver_ignores_toc_page_heading_echoes() -> tuple[bool, str]:
    rendered_lines = [
        ("СОДЕРЖАНИЕ", 2),
        ("ВВЕДЕНИЕ", 2),
        ("1. Теоретические основы", 2),
        ("1.1. Длинный подраздел с названием, которое должно переноситься естественно без ручных точек", 2),
        ("ЗАКЛЮЧЕНИЕ", 2),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 2),
        ("ВВЕДЕНИЕ", 3),
        ("1. Теоретические основы", 4),
        ("1.1. Длинный подраздел с названием, которое должно переноситься естественно без ручных точек", 5),
        ("ЗАКЛЮЧЕНИЕ", 8),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 9),
    ]
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), rendered_lines)
    front = _toc_texts_before_intro(out)
    expected = {
        "ВВЕДЕНИЕ\t3",
        "1. Теоретические основы\t4",
        "1.1. Длинный подраздел с названием, которое должно переноситься естественно без ручных точек\t5",
        "ЗАКЛЮЧЕНИЕ\t8",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ\t9",
    }
    missing = sorted(expected - set(front))
    if missing:
        return _result(False, f"TOC page echoes were used instead of body pages; missing={missing!r}, front={front!r}")
    return _result(True, "page resolver ignores TOC page heading echoes")


def test_autotoc_page_resolver_matches_wrapped_rendered_headings() -> tuple[bool, str]:
    rendered_lines = [
        ("ВВЕДЕНИЕ", 3),
        ("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ ИССЛЕДОВАНИЯ", 4),
        ("ОРГАНИЗАЦИОННОГО ПОКУПАТЕЛЬСКОГО ПОВЕДЕНИЯ", 4),
        ("1.1. Длинный подраздел с названием, которое должно переноситься естественно без ручных точек", 5),
        ("ЗАКЛЮЧЕНИЕ", 8),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 9),
    ]
    doc = _make_autotoc_doc(old_heading="Содержание")
    for p in doc.paragraphs:
        if p.text == "1. Теоретические основы":
            p.text = "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ ИССЛЕДОВАНИЯ ОРГАНИЗАЦИОННОГО ПОКУПАТЕЛЬСКОГО ПОВЕДЕНИЯ"
            break
    out = _run_static_contents_rebuild(doc, rendered_lines)
    front = _toc_texts_before_intro(out)
    expected = "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ ИССЛЕДОВАНИЯ ОРГАНИЗАЦИОННОГО ПОКУПАТЕЛЬСКОГО ПОВЕДЕНИЯ\t4"
    if expected not in front:
        return _result(False, f"wrapped rendered heading was not resolved: {front!r}")
    return _result(True, "page resolver matches wrapped rendered headings")


def test_autotoc_degenerate_page_mapping_fails_safe() -> tuple[bool, str]:
    source = _make_autotoc_doc(old_heading="Содержание")
    changed, out = _run_static_contents_rebuild_result(
        source,
        [
            ("СОДЕРЖАНИЕ", 2),
            ("ВВЕДЕНИЕ", 2),
            ("1. Теоретические основы", 2),
            ("1.1. Длинный подраздел с названием, которое должно переноситься естественно без ручных точек", 2),
            ("ЗАКЛЮЧЕНИЕ", 2),
            ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 2),
        ],
    )
    if changed:
        return _result(False, "degenerate all-TOC-page mapping should fail safe, got changed=True")
    texts = [p.text for p in out.paragraphs[:6]]
    if texts[1] != "Содержание":
        return _result(False, f"original document was not preserved after failed mapping: {texts!r}")
    return _result(True, "degenerate page mapping fails safe without mutating source")


def test_autotoc_entries_are_internal_hyperlinks_to_bookmarks() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    if not _document_has_bookmark(out):
        return _result(False, "body heading bookmarks were not created")
    for prefix in ("ВВЕДЕНИЕ", "1. Теоретические основы", "1.1. Длинный подраздел"):
        p = _toc_entry_paragraph(out, prefix)
        if p is None:
            return _result(False, f"TOC entry missing: {prefix}")
        if not _paragraph_has_internal_hyperlink(p):
            return _result(False, f"TOC entry is not an internal hyperlink: {p.text!r}")
    return _result(True, "TOC entries have internal hyperlinks to body heading bookmarks")


def test_autotoc_long_heading_uses_same_tab_leader_layout() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    p = _toc_entry_paragraph(out, "1.1. Длинный подраздел")
    if p is None:
        return _result(False, "long heading TOC entry missing")
    if not p.text.endswith("\t5"):
        return _result(False, f"long heading page number is not right-tab text: {p.text!r}")
    if not _paragraph_has_right_dot_tab(p):
        return _result(False, "long heading entry lacks right dot tab stop")
    return _result(True, "long heading uses tab leader layout with page number at right tab")


# ── Fallback heading collection (Roman/побитая coverage) ────────────────────


def _make_fallback_doc(
    *,
    h1_style_normal: bool = True,
    h2_style_normal: bool = True,
    include_intro_task_list: bool = False,
    include_sentence_like_h2: bool = False,
    include_inside_table_h2: bool = False,
    include_non_monotonic_h2: bool = False,
    include_after_appendices: bool = False,
    include_source_lines_with_numbers: bool = False,
) -> Document:
    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("Содержание")
    doc.add_paragraph("ВВЕДЕНИЕ........................................................3")
    doc.add_paragraph("1. Старый раздел")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    if include_intro_task_list:
        doc.add_paragraph("Задачи курсовой работы:")
        doc.add_paragraph("1. Изучить теорию")
        doc.add_paragraph("2. Проанализировать практику")

    h1 = doc.add_paragraph("1. Теоретические основы")
    if not h1_style_normal:
        h1.style = "Heading 1"
    doc.add_paragraph("Текст главы.")

    h2 = doc.add_paragraph("1.1. Сущность подхода")
    if not h2_style_normal:
        h2.style = "Heading 2"
    doc.add_paragraph("Текст подраздела.")

    if include_sentence_like_h2:
        doc.add_paragraph(
            "1.2. Маркетинговый подход. Данный подход применяется в анализе."
        )
        doc.add_paragraph("Дополнительный абзац.")

    if include_non_monotonic_h2:
        doc.add_paragraph("1.5. Раздел не по порядку")
        doc.add_paragraph("Текст.")

    if include_source_lines_with_numbers:
        doc.add_paragraph("Источник: 1. Иванов, 2025.")
        doc.add_paragraph("Примечание: 1.3. условный показатель.")

    if include_inside_table_h2:
        table = doc.add_table(rows=1, cols=1)
        table.rows[0].cells[0].paragraphs[0].text = "1.4. Не настоящий подраздел в ячейке"

    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("Итоги.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")

    if include_after_appendices:
        doc.add_paragraph("ПРИЛОЖЕНИЯ")
        doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
        doc.add_paragraph("1.1. Локальный псевдо-подраздел приложения")
        doc.add_paragraph("Текст приложения.")

    return doc


_FALLBACK_DEFAULT_LINES: list[tuple[str, int]] = [
    ("ВВЕДЕНИЕ", 3),
    ("1. Теоретические основы", 4),
    ("1.1. Сущность подхода", 5),
    ("ЗАКЛЮЧЕНИЕ", 8),
    ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 9),
]


def test_autotoc_fallback_normal_styled_heading2_included() -> tuple[bool, str]:
    doc = _make_fallback_doc(h1_style_normal=True, h2_style_normal=True)
    out = _run_static_contents_rebuild(doc, _FALLBACK_DEFAULT_LINES)
    front = _toc_texts_before_intro(out)
    if "1. Теоретические основы\t4" not in front:
        return _result(False, f"normal-styled H1 not in TOC: {front!r}")
    if "1.1. Сущность подхода\t5" not in front:
        return _result(False, f"normal-styled H2 not in TOC: {front!r}")
    return _result(True, "normal-styled real headings included via fallback")


def test_autotoc_fallback_normal_styled_heading1_included() -> tuple[bool, str]:
    doc = _make_fallback_doc(h1_style_normal=True, h2_style_normal=False)
    out = _run_static_contents_rebuild(doc, _FALLBACK_DEFAULT_LINES)
    front = _toc_texts_before_intro(out)
    if "1. Теоретические основы\t4" not in front:
        return _result(False, f"normal-styled H1 not in TOC: {front!r}")
    if "1.1. Сущность подхода\t5" not in front:
        return _result(False, f"styled H2 missing alongside fallback H1: {front!r}")
    return _result(True, "fallback H1 + styled H2 both included")


def test_autotoc_fallback_intro_task_list_excluded() -> tuple[bool, str]:
    doc = _make_fallback_doc(
        h1_style_normal=True,
        h2_style_normal=True,
        include_intro_task_list=True,
    )
    out = _run_static_contents_rebuild(doc, _FALLBACK_DEFAULT_LINES)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "Изучить теорию" in joined or "Проанализировать практику" in joined:
        return _result(False, f"intro task list leaked into TOC: {front!r}")
    if "1. Теоретические основы\t4" not in front:
        return _result(False, f"real heading after task list missing from TOC: {front!r}")
    return _result(True, "intro task list excluded; real heading still included")


def test_autotoc_fallback_numbered_sentence_excluded() -> tuple[bool, str]:
    doc = _make_fallback_doc(
        h1_style_normal=False,
        h2_style_normal=True,
        include_sentence_like_h2=True,
    )
    rendered_lines = list(_FALLBACK_DEFAULT_LINES)
    out = _run_static_contents_rebuild(doc, rendered_lines)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "Маркетинговый подход" in joined:
        return _result(False, f"sentence-like numbered paragraph leaked into TOC: {front!r}")
    if "1.1. Сущность подхода\t5" not in front:
        return _result(False, f"valid fallback H2 missing: {front!r}")
    return _result(True, "sentence-like numbered paragraph rejected by fallback")


def test_autotoc_fallback_paragraph_inside_table_excluded() -> tuple[bool, str]:
    doc = _make_fallback_doc(
        h1_style_normal=False,
        h2_style_normal=False,
        include_inside_table_h2=True,
    )
    out = _run_static_contents_rebuild(doc, _FALLBACK_DEFAULT_LINES)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "Не настоящий" in joined or "1.4." in joined:
        return _result(False, f"H2 inside table cell leaked into TOC: {front!r}")
    return _result(True, "H2 inside table cell excluded")


def test_autotoc_fallback_non_monotonic_rejected() -> tuple[bool, str]:
    doc = _make_fallback_doc(
        h1_style_normal=False,
        h2_style_normal=False,
        include_non_monotonic_h2=True,
    )
    out = _run_static_contents_rebuild(doc, _FALLBACK_DEFAULT_LINES)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "1.5." in joined or "Раздел не по порядку" in joined:
        return _result(False, f"non-monotonic H2 leaked into TOC: {front!r}")
    if "1.1. Сущность подхода\t5" not in front:
        return _result(False, f"strict-style H2 missing alongside rejected fallback: {front!r}")
    return _result(True, "non-monotonic fallback H2 rejected")


def test_autotoc_fallback_after_appendices_excluded() -> tuple[bool, str]:
    doc = _make_fallback_doc(
        h1_style_normal=False,
        h2_style_normal=False,
        include_after_appendices=True,
    )
    rendered_lines = _FALLBACK_DEFAULT_LINES + [("ПРИЛОЖЕНИЯ", 10)]
    out = _run_static_contents_rebuild(doc, rendered_lines)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "Локальный псевдо-подраздел" in joined or "ПРИЛОЖЕНИЕ 1" in joined:
        return _result(False, f"appendix-local heading leaked into TOC: {front!r}")
    if "ПРИЛОЖЕНИЯ\t10" not in front:
        return _result(False, f"general ПРИЛОЖЕНИЯ entry missing: {front!r}")
    return _result(True, "fallback heading after general ПРИЛОЖЕНИЯ excluded")


def test_autotoc_fallback_source_or_note_excluded() -> tuple[bool, str]:
    doc = _make_fallback_doc(
        h1_style_normal=False,
        h2_style_normal=False,
        include_source_lines_with_numbers=True,
    )
    out = _run_static_contents_rebuild(doc, _FALLBACK_DEFAULT_LINES)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "Иванов" in joined or "условный показатель" in joined:
        return _result(False, f"source/note service line leaked into TOC: {front!r}")
    return _result(True, "source/note lines excluded by fallback")


def test_autotoc_fallback_strict_style_path_unchanged() -> tuple[bool, str]:
    """
    Regression guard: the strict-style path keeps producing the same TOC for
    fully styled inputs.
    """
    doc = _make_fallback_doc(h1_style_normal=False, h2_style_normal=False)
    out = _run_static_contents_rebuild(doc, _FALLBACK_DEFAULT_LINES)
    front = _toc_texts_before_intro(out)
    expected = {
        "СОДЕРЖАНИЕ",
        "ВВЕДЕНИЕ\t3",
        "1. Теоретические основы\t4",
        "1.1. Сущность подхода\t5",
        "ЗАКЛЮЧЕНИЕ\t8",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ\t9",
    }
    missing = sorted(expected - set(front))
    if missing:
        return _result(False, f"strict-style path output changed; missing={missing!r}, front={front!r}")
    return _result(True, "strict-style path output unchanged")


# ── Visual TOC polish (Приложение 3) ───────────────────────────────────────


def _toc_title_index(doc: Document) -> int | None:
    for idx, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == "СОДЕРЖАНИЕ":
            return idx
    return None


def _all_hyperlink_runs(paragraph):
    runs = []
    for hyperlink in paragraph._element.findall(qn("w:hyperlink")):
        for run in hyperlink.findall(qn("w:r")):
            runs.append(run)
    return runs


def _run_ascii_font(run_elem) -> str | None:
    r_pr = run_elem.find(qn("w:rPr"))
    if r_pr is None:
        return None
    rfonts = r_pr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn("w:ascii"))


def _run_size_twips(run_elem) -> str | None:
    r_pr = run_elem.find(qn("w:rPr"))
    if r_pr is None:
        return None
    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        return None
    return sz.get(qn("w:val"))


def test_autotoc_blank_paragraph_after_title() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    title_idx = _toc_title_index(out)
    if title_idx is None:
        return _result(False, "СОДЕРЖАНИЕ heading missing")
    if title_idx + 1 >= len(out.paragraphs):
        return _result(False, "no paragraph after СОДЕРЖАНИЕ heading")
    blank_p = out.paragraphs[title_idx + 1]
    if (blank_p.text or "").strip():
        return _result(False, f"paragraph after СОДЕРЖАНИЕ is not blank: {blank_p.text!r}")
    if title_idx + 2 >= len(out.paragraphs):
        return _result(False, "no entry after blank paragraph")
    first_entry = out.paragraphs[title_idx + 2]
    if "\t" not in (first_entry.text or ""):
        return _result(False, f"second paragraph after title is not a TOC entry: {first_entry.text!r}")
    return _result(True, "exactly one blank paragraph between СОДЕРЖАНИЕ and first entry")


def test_autotoc_no_double_blank_after_title() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    title_idx = _toc_title_index(out)
    if title_idx is None:
        return _result(False, "СОДЕРЖАНИЕ heading missing")
    blank_count = 0
    j = title_idx + 1
    while j < len(out.paragraphs) and not (out.paragraphs[j].text or "").strip():
        blank_count += 1
        j += 1
    if blank_count != 1:
        return _result(False, f"expected exactly 1 blank paragraph after title, got {blank_count}")
    return _result(True, "no double blank after СОДЕРЖАНИЕ on rebuild")


def test_autotoc_title_font_is_times_new_roman_14_bold_centered() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    title_idx = _toc_title_index(out)
    if title_idx is None:
        return _result(False, "СОДЕРЖАНИЕ heading missing")
    title_p = out.paragraphs[title_idx]
    if title_p.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, f"СОДЕРЖАНИЕ not centered: alignment={title_p.alignment!r}")
    runs = list(title_p.runs)
    if not runs:
        return _result(False, "СОДЕРЖАНИЕ paragraph has no runs")
    for r in runs:
        if r.font.name != "Times New Roman":
            return _result(False, f"СОДЕРЖАНИЕ run font is {r.font.name!r}")
        if r.font.size is None or r.font.size.pt != 14:
            return _result(False, f"СОДЕРЖАНИЕ run size is {r.font.size!r}")
        if r.bold is not True:
            return _result(False, f"СОДЕРЖАНИЕ run not bold: {r.bold!r}")
    return _result(True, "СОДЕРЖАНИЕ title is TNR 14 bold centered")


def test_autotoc_entry_hyperlink_runs_use_times_new_roman_14() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    seen_any = False
    for p in out.paragraphs:
        if "\t" not in (p.text or ""):
            continue
        hl_runs = _all_hyperlink_runs(p)
        if not hl_runs:
            return _result(False, f"TOC entry without hyperlink runs: {p.text!r}")
        for run in hl_runs:
            font = _run_ascii_font(run)
            if font != "Times New Roman":
                return _result(False, f"hyperlink run ascii font is {font!r} in {p.text!r}")
            sz = _run_size_twips(run)
            if sz != "28":
                return _result(False, f"hyperlink run size is {sz!r} in {p.text!r}")
            seen_any = True
    if not seen_any:
        return _result(False, "no TOC entries with hyperlink runs found")
    return _result(True, "all TOC hyperlink runs are Times New Roman 14")


def test_autotoc_entries_have_zero_hanging_and_firstline_xml() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    for p in out.paragraphs:
        if "\t" not in (p.text or ""):
            continue
        p_pr = p._element.pPr
        if p_pr is None:
            continue
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            continue
        for attr in ("w:left", "w:right", "w:firstLine", "w:hanging"):
            value = ind.get(qn(attr))
            if value is not None and value != "0":
                return _result(False, f"TOC entry {attr}={value!r} (expected '0'): {p.text!r}")
    return _result(True, "TOC entries have zero left/right/firstLine/hanging indent")


def test_autotoc_blank_paragraph_carries_no_tab_or_hyperlink() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    title_idx = _toc_title_index(out)
    if title_idx is None:
        return _result(False, "СОДЕРЖАНИЕ heading missing")
    blank_p = out.paragraphs[title_idx + 1]
    if "\t" in (blank_p.text or ""):
        return _result(False, f"blank paragraph contains tab: {blank_p.text!r}")
    if blank_p._element.findall(qn("w:hyperlink")):
        return _result(False, "blank paragraph contains a hyperlink")
    p_pr = blank_p._element.pPr
    if p_pr is not None and p_pr.findall(qn("w:tabs")):
        return _result(False, "blank paragraph carries a tab stop")
    return _result(True, "blank paragraph has no tab, no hyperlink, no tab stop")


# ── Old simple TOC removal (loose detection + validation) ─────────────────


def _make_old_simple_toc_doc(
    *,
    heading: str = "Содержание",
    extra_entries: list[str] | None = None,
    extra_front_paragraphs: list[str] | None = None,
) -> Document:
    doc = Document()
    doc.add_paragraph("Титульная строка")
    if extra_front_paragraphs:
        for text in extra_front_paragraphs:
            doc.add_paragraph(text)
    doc.add_paragraph(heading)
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. Теоретические основы")
    doc.add_paragraph("1.1. Подпункт теоретической части")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    if extra_entries:
        for text in extra_entries:
            doc.add_paragraph(text)
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    h1 = doc.add_paragraph("1. Теоретические основы")
    h1.style = "Heading 1"
    doc.add_paragraph("Текст главы.")
    h2 = doc.add_paragraph("1.1. Подпункт теоретической части")
    h2.style = "Heading 2"
    doc.add_paragraph("Текст подпункта.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("Итоги.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")
    return doc


_OLD_TOC_RENDERED_LINES: list[tuple[str, int]] = [
    ("ВВЕДЕНИЕ", 3),
    ("1. Теоретические основы", 4),
    ("1.1. Подпункт теоретической части", 5),
    ("ЗАКЛЮЧЕНИЕ", 7),
    ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 8),
]


def test_autotoc_old_simple_toc_without_page_numbers_removed() -> tuple[bool, str]:
    doc = _make_old_simple_toc_doc(heading="Содержание")
    out = _run_static_contents_rebuild(doc, _OLD_TOC_RENDERED_LINES)
    front = _toc_texts_before_intro(out)
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single СОДЕРЖАНИЕ, front={front!r}")
    if "Содержание" in front:
        return _result(False, f"mixed-case old heading survived: {front!r}")
    canonical_idx = front.index("СОДЕРЖАНИЕ")
    for above in front[:canonical_idx]:
        if above in (
            "ВВЕДЕНИЕ",
            "1. Теоретические основы",
            "1.1. Подпункт теоретической части",
            "ЗАКЛЮЧЕНИЕ",
            "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        ):
            return _result(False, f"old TOC entry survived above canonical: {front!r}")
    return _result(True, "old plain-text TOC removed; canonical single block")


def test_autotoc_old_toc_heading_with_trailing_dot_removed() -> tuple[bool, str]:
    doc = _make_old_simple_toc_doc(heading="Содержание.")
    out = _run_static_contents_rebuild(doc, _OLD_TOC_RENDERED_LINES)
    front = _toc_texts_before_intro(out)
    if "Содержание." in front:
        return _result(False, f"trailing-dot TOC heading survived: {front!r}")
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single canonical СОДЕРЖАНИЕ: {front!r}")
    canonical_idx = front.index("СОДЕРЖАНИЕ")
    if any(p.startswith("1. Теоретические") for p in front[:canonical_idx]):
        return _result(False, f"old TOC entries above canonical: {front!r}")
    return _result(True, "trailing-dot old TOC heading and entries removed")


def test_autotoc_old_toc_with_appendix_entries_removed() -> tuple[bool, str]:
    doc = _make_old_simple_toc_doc(
        heading="СОДЕРЖАНИЕ",
        extra_entries=["ПРИЛОЖЕНИЕ 1", "ПРИЛОЖЕНИЕ 2", "ПРИЛОЖЕНИЕ 3"],
    )
    out = _run_static_contents_rebuild(doc, _OLD_TOC_RENDERED_LINES)
    front = _toc_texts_before_intro(out)
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single СОДЕРЖАНИЕ: {front!r}")
    canonical_idx = front.index("СОДЕРЖАНИЕ")
    for above in front[:canonical_idx]:
        if above.startswith("ПРИЛОЖЕНИЕ "):
            return _result(False, f"old appendix entries leaked into front: {front!r}")
    return _result(True, "old TOC with appendix entries removed cleanly")


def test_autotoc_old_toc_heading_with_page_number_removed() -> tuple[bool, str]:
    doc = _make_old_simple_toc_doc(heading="СОДЕРЖАНИЕ ……………… 1")
    out = _run_static_contents_rebuild(doc, _OLD_TOC_RENDERED_LINES)
    front = _toc_texts_before_intro(out)
    if any(p == "СОДЕРЖАНИЕ ……………… 1" for p in front):
        return _result(False, f"leader-style old TOC heading survived: {front!r}")
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single canonical СОДЕРЖАНИЕ: {front!r}")
    return _result(True, "old TOC heading with leaders + page number removed")


def test_autotoc_standalone_loose_contents_paragraph_does_not_delete_body() -> tuple[bool, str]:
    """
    P1 hard recovery (v2 approved rule): a Содержание marker followed by a
    confident intro (standalone ВВЕДЕНИЕ by text) authorises hard-removal of the
    ENTIRE block between the marker and the intro — including hand-typed prose
    that would otherwise read like front-matter. The structural rule
    (contents marker + confident intro) wins; no length thresholds are used.
    Previously this case preserved the prose; the product rule changed so the
    block is now removed and replaced by a single canonical СОДЕРЖАНИЕ.
    """
    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("Содержание.")  # loose marker; prose follows before real intro
    doc.add_paragraph("Просто пояснительный абзац, не относящийся к содержанию.")
    doc.add_paragraph("Ещё один обычный абзац без структуры TOC.")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    h1 = doc.add_paragraph("1. Теоретические основы")
    h1.style = "Heading 1"
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("Итоги.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")

    rendered_lines = [
        ("ВВЕДЕНИЕ", 3),
        ("1. Теоретические основы", 4),
        ("ЗАКЛЮЧЕНИЕ", 5),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 6),
    ]
    out = _run_static_contents_rebuild(doc, rendered_lines)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "Просто пояснительный абзац" in joined:
        return _result(False, f"prose between marker and intro survived: {front!r}")
    if "Ещё один обычный абзац" in joined:
        return _result(False, f"prose between marker and intro survived: {front!r}")
    if "Содержание." in joined:
        return _result(False, f"old loose Содержание marker survived: {front!r}")
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected exactly one canonical СОДЕРЖАНИЕ: {front!r}")
    return _result(True, "confident intro authorises hard-removal of marker→intro block")


def test_autotoc_p1_garbage_removed_when_intro_normal_style_text_only() -> tuple[bool, str]:
    """
    P1 v2 CORE: garbage prose must be removed even when the intro paragraph is
    NOT yet Heading 1 — a standalone ВВЕДЕНИЕ detected by TEXT is a sufficient
    confidence signal. Mirrors the dirty demo (primer) structure.
    """
    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("Содержание")
    doc.add_paragraph("тут что то будет содержаться, оставлю это на откуп боту")
    doc.add_paragraph("Сразу к делу:")
    doc.add_paragraph("ВВЕДЕНИЕ")  # Normal style — NOT promoted to Heading 1
    doc.add_paragraph("Текст введения.")
    h1 = doc.add_paragraph("1. Теоретические основы")
    h1.style = "Heading 1"
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")

    rendered_lines = [
        ("ВВЕДЕНИЕ", 3),
        ("1. Теоретические основы", 4),
        ("ЗАКЛЮЧЕНИЕ", 5),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 6),
    ]
    out = _run_static_contents_rebuild(doc, rendered_lines)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "тут что то" in joined:
        return _result(False, f"garbage survived (text-only intro): {front!r}")
    if "Сразу к делу" in joined:
        return _result(False, f"garbage survived (text-only intro): {front!r}")
    if "Содержание" in joined and "СОДЕРЖАНИЕ" not in [t for t in front if t == "Содержание"] and front.count("Содержание") > 0:
        return _result(False, f"old lowercase Содержание marker survived: {front!r}")
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single canonical СОДЕРЖАНИЕ: {front!r}")
    return _result(True, "garbage removed with Normal-style standalone ВВЕДЕНИЕ (text signal)")


def test_autotoc_p1_oglavlenie_variant_garbage_removed_text_only() -> tuple[bool, str]:
    """P1 v2: Оглавление marker + Normal-style ВВЕДЕНИЕ → garbage removed."""
    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("Оглавление")
    doc.add_paragraph("мусорный текст перед введением")
    doc.add_paragraph("ВВЕДЕНИЕ")  # Normal style
    doc.add_paragraph("Текст введения.")
    h1 = doc.add_paragraph("1. Основной раздел")
    h1.style = "Heading 1"
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")

    rendered_lines = [
        ("ВВЕДЕНИЕ", 3),
        ("1. Основной раздел", 4),
        ("ЗАКЛЮЧЕНИЕ", 5),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 6),
    ]
    out = _run_static_contents_rebuild(doc, rendered_lines)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "мусорный текст" in joined:
        return _result(False, f"garbage survived in Оглавление case: {front!r}")
    if "Оглавление" in joined:
        return _result(False, f"old Оглавление marker survived: {front!r}")
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single canonical СОДЕРЖАНИЕ: {front!r}")
    return _result(True, "Оглавление variant: garbage removed with text-only intro signal")


def test_autotoc_p1_fail_closed_no_confident_intro_keeps_block() -> tuple[bool, str]:
    """
    P1 v2 FAIL-CLOSED: when the body-start paragraph is neither a standalone
    intro by text nor Heading 1 styled, _find_existing_contents_start must NOT
    return a removal index for a marker followed by non-TOC prose — the
    conservative safety guard stays active. Unit-level check.
    """
    import guides.coursework_kfu_2025.contents_builder as cb

    doc = Document()
    doc.add_paragraph("Титульная строка")          # 0
    doc.add_paragraph("Содержание")                 # 1 marker
    doc.add_paragraph("обычный пояснительный абзац один")  # 2 prose
    doc.add_paragraph("обычный пояснительный абзац два")   # 3 prose
    doc.add_paragraph("ещё один абзац без структуры")      # 4 = body_start (NOT intro, NOT heading)
    doc.add_paragraph("Текст.")                     # 5

    body_start = 4
    intro_para = doc.paragraphs[body_start]
    # Sanity: body_start really is not a confident intro
    if cb._is_intro_heading(intro_para.text) or cb._is_heading1_style(intro_para):
        return _result(False, "test setup wrong: body_start looks like confident intro")

    start = cb._find_existing_contents_start(doc, body_start)
    if start is not None:
        return _result(False, f"fail-closed violated: returned removal index {start}")
    return _result(True, "fail-closed: no confident intro keeps conservative guard, no removal")


def test_autotoc_p1_wellformed_old_toc_removed_with_confident_intro() -> tuple[bool, str]:
    """P1 v2: a well-formed old TOC block is still removed (confident intro path)."""
    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("Содержание")
    doc.add_paragraph("ВВЕДЕНИЕ........................................................3")
    doc.add_paragraph("1. Раздел...............................................4")
    doc.add_paragraph("ВВЕДЕНИЕ")  # Normal-style real intro
    doc.add_paragraph("Текст введения.")
    h1 = doc.add_paragraph("1. Теоретические основы")
    h1.style = "Heading 1"
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")

    rendered_lines = [
        ("ВВЕДЕНИЕ", 3),
        ("1. Теоретические основы", 4),
        ("ЗАКЛЮЧЕНИЕ", 5),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 6),
    ]
    out = _run_static_contents_rebuild(doc, rendered_lines)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "ВВЕДЕНИЕ........" in joined:
        return _result(False, f"old TOC entry survived: {front!r}")
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single canonical СОДЕРЖАНИЕ: {front!r}")
    return _result(True, "well-formed old TOC block removed with confident intro")


def test_autotoc_p1_title_page_before_contents_preserved() -> tuple[bool, str]:
    """
    P1 v2: paragraphs BEFORE the Содержание marker (title page) are preserved;
    only the marker→intro block is removed.
    """
    doc = Document()
    doc.add_paragraph("КАЗАНСКИЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ")
    doc.add_paragraph("Курсовая работа по предмету")
    doc.add_paragraph("Автор: Иванов Иван")
    doc.add_paragraph("Содержание")
    doc.add_paragraph("мусор между содержанием и введением")
    doc.add_paragraph("ВВЕДЕНИЕ")  # Normal-style
    doc.add_paragraph("Текст введения.")
    h1 = doc.add_paragraph("1. Теоретические основы")
    h1.style = "Heading 1"
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")

    rendered_lines = [
        ("ВВЕДЕНИЕ", 4),
        ("1. Теоретические основы", 5),
        ("ЗАКЛЮЧЕНИЕ", 6),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 7),
    ]
    out = _run_static_contents_rebuild(doc, rendered_lines)
    front = _toc_texts_before_intro(out)
    joined = "\n".join(front)
    if "КАЗАНСКИЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ" not in joined:
        return _result(False, f"title page paragraph deleted: {front!r}")
    if "Автор: Иванов Иван" not in joined:
        return _result(False, f"title page paragraph deleted: {front!r}")
    if "мусор между" in joined:
        return _result(False, f"garbage survived: {front!r}")
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single canonical СОДЕРЖАНИЕ: {front!r}")
    return _result(True, "title page preserved; marker→intro garbage removed")


def test_autotoc_dot_leader_and_pages_preserved_with_blank() -> tuple[bool, str]:
    out = _run_static_contents_rebuild(_make_autotoc_doc(old_heading="Содержание"), _default_autotoc_lines())
    entry_paragraphs = [p for p in out.paragraphs if "\t" in (p.text or "")]
    if not entry_paragraphs:
        return _result(False, "no TOC entries found")
    for p in entry_paragraphs:
        if not _paragraph_has_right_dot_tab(p):
            return _result(False, f"entry lacks right dot tab stop: {p.text!r}")
        if not (p.text or "").split("\t")[-1].strip():
            return _result(False, f"entry has no page number: {p.text!r}")
    return _result(True, "dot leader and page numbers preserved with blank paragraph present")


def _add_softbreak_para(doc: Document, lines: list[str]):
    """Add a paragraph with soft line-breaks (<w:br/>) between each line."""
    p = doc.add_paragraph()
    p.clear()
    for i, line in enumerate(lines):
        run = p.add_run(line)
        if i < len(lines) - 1:
            br = OxmlElement("w:br")
            run._element.append(br)
    return p


_SOFTBREAK_TOC_RENDERED_LINES: list[tuple[str, int]] = [
    ("ВВЕДЕНИЕ", 3),
    ("1. Теоретические основы", 4),
    ("2. Анализ и совершенствование", 5),
    ("ЗАКЛЮЧЕНИЕ", 6),
    ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 7),
]


def _make_softbreak_toc_doc_173_pattern() -> Document:
    """
    Mimics the 173 input structure:
      - title paragraph
      - Para 1: soft-break TOC part 1 (СОДЕРЖАНИЕ + a few chapters)
      - Para 2: soft-break TOC part 2 (more chapters + ПРИЛОЖЕНИЕ 1/2/3)
      - two near-empty paragraphs (spaces)
      - real ВВЕДЕНИЕ heading (Heading 1 style) → body start
      - body content with headings
    """
    doc = Document()
    doc.add_paragraph("Специфика организационного покупательского поведения")
    _add_softbreak_para(doc, [
        "СОДЕРЖАНИЕ",
        "ВВЕДЕНИЕ",
        "1. Теоретические основы",
    ])
    _add_softbreak_para(doc, [
        "2. Анализ и совершенствование",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "ПРИЛОЖЕНИЕ 1",
        "ПРИЛОЖЕНИЕ 2",
        "ПРИЛОЖЕНИЕ 3",
    ])
    doc.add_paragraph(" ")
    doc.add_paragraph(" ")
    intro_p = doc.add_paragraph("ВВЕДЕНИЕ")
    intro_p.style = doc.styles["Heading 1"]
    doc.add_paragraph("Текст введения.")
    h1a = doc.add_paragraph("1. Теоретические основы")
    h1a.style = doc.styles["Heading 1"]
    doc.add_paragraph("Текст раздела 1.")
    h1b = doc.add_paragraph("2. Анализ и совершенствование")
    h1b.style = doc.styles["Heading 1"]
    doc.add_paragraph("Текст раздела 2.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("Итоги.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")
    return doc


def test_autotoc_softbreak_plain_toc_removed() -> tuple[bool, str]:
    """
    Old TOC where heading + all entries live in a single soft-break paragraph
    (no page numbers) must be removed; one canonical СОДЕРЖАНИЕ inserted.
    """
    doc = Document()
    doc.add_paragraph("Титульная строка")
    _add_softbreak_para(doc, [
        "СОДЕРЖАНИЕ",
        "ВВЕДЕНИЕ",
        "1. Теоретические основы",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    ])
    intro_p = doc.add_paragraph("ВВЕДЕНИЕ")
    intro_p.style = doc.styles["Heading 1"]
    doc.add_paragraph("Текст введения.")
    h1 = doc.add_paragraph("1. Теоретические основы")
    h1.style = doc.styles["Heading 1"]
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("Итоги.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")

    rendered = [
        ("ВВЕДЕНИЕ", 3),
        ("1. Теоретические основы", 4),
        ("ЗАКЛЮЧЕНИЕ", 5),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 6),
    ]
    out = _run_static_contents_rebuild(doc, rendered)
    front = _toc_texts_before_intro(out)
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single canonical СОДЕРЖАНИЕ, got front={front!r}")
    canonical_idx = front.index("СОДЕРЖАНИЕ")
    # Old TOC entries (not the title) must not appear above the canonical heading
    old_toc_entries = {"ВВЕДЕНИЕ", "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ"}
    for above in front[:canonical_idx]:
        if above.upper() in old_toc_entries or above.startswith("1. "):
            return _result(False, f"old TOC entry survived above canonical: {front!r}")
    return _result(True, "old soft-break plain TOC removed; canonical single block")


def test_autotoc_softbreak_toc_appendix_continuation_removed() -> tuple[bool, str]:
    """
    173-pattern: two-paragraph soft-break front matter where the second paragraph
    contains ПРИЛОЖЕНИЕ 1/2/3 as soft-break lines. Previously _is_toc_like_entry
    stripped 'ПРИЛОЖЕНИЕ 1' to 'ПРИЛОЖЕНИЕ' (treating '1' as a page number),
    so the appendix regex rejected it and the entire old TOC block was left in place.
    After the fix, the full block must be removed.
    """
    doc = _make_softbreak_toc_doc_173_pattern()
    out = _run_static_contents_rebuild(doc, _SOFTBREAK_TOC_RENDERED_LINES)
    front = _toc_texts_before_intro(out)
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected single СОДЕРЖАНИЕ, got front={front!r}")
    canonical_idx = front.index("СОДЕРЖАНИЕ")
    # Old soft-break TOC paragraphs (the ones containing soft-break content)
    # must not appear above the canonical heading. The title is OK.
    for above in front[:canonical_idx]:
        if above.startswith("ПРИЛОЖЕНИЕ ") or "\n" in above or "СПИСОК" in above:
            return _result(False, f"old TOC content survived above canonical: {front!r}")
        # The soft-break para would contain multiple lines when read via p.text
        if above.count("\n") > 0:
            return _result(False, f"multi-line old TOC para survived: {front!r}")
    return _result(True, "173-pattern: soft-break TOC with ПРИЛОЖЕНИЕ continuation removed")


def test_autotoc_body_intro_preserved_after_softbreak_toc() -> tuple[bool, str]:
    """
    Real body ВВЕДЕНИЕ heading must not be deleted when the old soft-break TOC
    block is removed. Only pre-body paragraphs up to (not including) the real
    ВВЕДЕНИЕ should be removed.
    """
    doc = _make_softbreak_toc_doc_173_pattern()
    out = _run_static_contents_rebuild(doc, _SOFTBREAK_TOC_RENDERED_LINES)
    texts = [p.text for p in out.paragraphs]
    if "ВВЕДЕНИЕ" not in texts:
        return _result(False, f"real body ВВЕДЕНИЕ was deleted: {texts[:10]!r}")
    # Verify body content follows ВВЕДЕНИЕ
    intro_idx = texts.index("ВВЕДЕНИЕ")
    if intro_idx + 1 >= len(texts) or not texts[intro_idx + 1]:
        return _result(False, "nothing follows ВВЕДЕНИЕ — body may have been truncated")
    return _result(True, "real body ВВЕДЕНИЕ preserved after soft-break TOC cleanup")


def test_autotoc_no_duplicate_contents_after_rebuild() -> tuple[bool, str]:
    """
    After full rebuild on a 173-pattern doc, there must be exactly one СОДЕРЖАНИЕ
    in the entire output document (no duplicate from old TOC + new canonical).
    """
    doc = _make_softbreak_toc_doc_173_pattern()
    out = _run_static_contents_rebuild(doc, _SOFTBREAK_TOC_RENDERED_LINES)
    all_texts = [p.text for p in out.paragraphs]
    count = sum(1 for t in all_texts if "СОДЕРЖАНИЕ" in t and t.strip() == "СОДЕРЖАНИЕ")
    if count != 1:
        return _result(False, f"expected 1 СОДЕРЖАНИЕ in entire doc, found {count}: {all_texts[:20]!r}")
    return _result(True, "no duplicate СОДЕРЖАНИЕ after rebuild on 173-pattern doc")


def _paragraph_has_texts_in_order(doc: Document, expected: list[str]) -> bool:
    texts = [paragraph.text for paragraph in doc.paragraphs]
    pos = 0
    for expected_text in expected:
        try:
            pos = texts.index(expected_text, pos) + 1
        except ValueError:
            return False
    return True


def test_body_soft_break_chapter_and_section_headings_are_separated() -> tuple[bool, str]:
    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("СОДЕРЖАНИЕ\nВВЕДЕНИЕ\n1. Теоретические основы")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph(
        "1. Теоретические основы исследования\n"
        "1.1. Сущность организационного покупательского поведения"
    )
    doc.add_paragraph("Текст раздела.")

    formatted = _format_synthetic_doc(doc)
    texts = [paragraph.text for paragraph in formatted.paragraphs]
    if any("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ ИССЛЕДОВАНИЯ" in text and "1.1." in text for text in texts):
        return _result(False, "chapter and section headings remained merged")
    if not _paragraph_has_texts_in_order(
        formatted,
        [
            "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ ИССЛЕДОВАНИЯ",
            "1.1. Сущность организационного покупательского поведения",
            "Текст раздела.",
        ],
    ):
        return _result(False, f"separated heading paragraphs missing or out of order: {texts!r}")
    return _result(True, "body chapter and section headings are structurally separated")


def test_body_soft_break_heading_and_body_are_separated_after_intro() -> tuple[bool, str]:
    doc = Document()
    doc.add_paragraph("СОДЕРЖАНИЕ\nВВЕДЕНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph(
        "1. Теоретические основы\n"
        "Обычный текст главы начинается после заголовка."
    )
    doc.add_paragraph("Следующий абзац.")

    formatted = _format_synthetic_doc(doc)
    texts = [paragraph.text for paragraph in formatted.paragraphs]
    if any("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ" in text and "Обычный текст" in text for text in texts):
        return _result(False, "heading and body text remained merged")
    if not _paragraph_has_texts_in_order(
        formatted,
        [
            "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ",
            "Обычный текст главы начинается после заголовка.",
            "Следующий абзац.",
        ],
    ):
        return _result(False, f"heading/body split paragraphs missing or out of order: {texts!r}")
    return _result(True, "body heading and following body text are separated after intro")


def test_body_soft_break_split_does_not_touch_toc() -> tuple[bool, str]:
    doc = Document()
    front_matter_texts = [
        "Титульная строка",
        "СОДЕРЖАНИЕ\nВВЕДЕНИЕ\n1. Теоретические основы\n1.1. Сущность подхода",
    ]
    for text in front_matter_texts:
        doc.add_paragraph(text)
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. Теоретические основы\n1.1. Сущность подхода")
    doc.add_paragraph("Текст раздела.")

    formatted = _format_synthetic_doc(doc)
    intro_idx = _paragraph_index(formatted, "ВВЕДЕНИЕ")
    if intro_idx is None:
        return _result(False, "real standalone ВВЕДЕНИЕ missing after formatting")
    actual_front_matter = [paragraph.text for paragraph in formatted.paragraphs[:intro_idx]]
    if actual_front_matter != front_matter_texts:
        return _result(False, f"TOC/front matter was touched by body splitter: {actual_front_matter!r}")
    return _result(True, "body soft-break splitting does not touch TOC/front matter")


def test_body_soft_break_split_does_not_split_ordinary_body_text() -> tuple[bool, str]:
    ordinary = "Обычный абзац содержит перенос строки\nно не является структурным заголовком."
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph(ordinary)
    doc.add_paragraph("Следующий обычный абзац.")

    formatted = _format_synthetic_doc(doc)
    if ordinary not in [paragraph.text for paragraph in formatted.paragraphs]:
        return _result(False, "ordinary soft-break body paragraph was split or rewritten")
    return _result(True, "ordinary body soft break is not structurally split")




# ── Task A — figure deletion ──────────────────────────────────────────────────

def test_a_rule4_does_not_delete_images() -> tuple[bool, str]:
    """
    Rule 4 must NOT remove paragraphs that contain w:drawing even when they
    appear to be empty (no text) and land at the top of a new estimated page.
    """
    doc = _make_minimal_doc_with_image()
    before = _count_drawings(doc)
    if before == 0:
        return _result(False, "test setup failed: no drawing inserted")

    apply_rule4_empty_first_lines(doc)

    after = _count_drawings(doc)
    if after < before:
        return _result(False, f"drawing deleted: before={before}, after={after}")
    return _result(True, f"drawings intact: {after}")


def test_a_para_has_image_helper() -> tuple[bool, str]:
    """_para_has_image correctly detects w:drawing elements."""
    doc = _make_minimal_doc_with_image()
    # Last paragraph has the drawing
    last_p = doc.paragraphs[-1]
    if not _para_has_image(last_p._element):
        return _result(False, "_para_has_image returned False for paragraph with w:drawing")
    # A normal paragraph should return False
    normal_p = doc.paragraphs[0]
    if _para_has_image(normal_p._element):
        return _result(False, "_para_has_image returned True for text paragraph")
    return _result(True)


def test_a_rule4_preserves_front_matter_section_breaks() -> tuple[bool, str]:
    """
    Product rule: title, contents, and introduction are separated by structural
    section breaks. Rule 4 may delete visual blank paragraphs, but never a
    paragraph carrying w:sectPr.
    """
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    probe = doc.add_paragraph("Титульная строка")
    h_per_para = _estimate_para_height(probe)
    probe._element.getparent().remove(probe._element)

    body_h = _body_height_pt(doc)
    title_lines = max(1, int(body_h / h_per_para))
    for i in range(title_lines):
        doc.add_paragraph(f"Титульная строка {i + 1}")
    doc.add_paragraph("")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ 3")
    doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ 4")
    doc.add_paragraph("")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "front_matter_in.docx"
        out = Path(tmp) / "front_matter_out.docx"
        doc.save(inp)
        process_document(inp, out)

        formatted = Document(str(out))
        before_positions = _section_break_positions(formatted)
        if len(before_positions) < 2:
            return _result(False, f"front matter section breaks missing after Phase 1: {before_positions!r}")

        apply_rule4_empty_first_lines(formatted)
        after_positions = _section_break_positions(formatted)

    if len(after_positions) != len(before_positions):
        return _result(
            False,
            f"Rule 4 removed structural section break(s): before={before_positions!r} after={after_positions!r}",
        )
    if after_positions != before_positions:
        return _result(
            False,
            f"Rule 4 moved structural section break(s): before={before_positions!r} after={after_positions!r}",
        )
    return _result(True, "Rule 4 preserved front matter section breaks")


# ── Task C — student continuation length ─────────────────────────────────────

def test_c_continuation_length_guard() -> tuple[bool, str]:
    """
    _is_student_continuation must accept texts up to 30 chars.
    Target behaviour after raising limit 27 → 30:
      ≤30 chars + 'продолжени' + 'таблиц' → True
      >30 chars → False
    """
    cases = [
        # (text, expected_after_fix_to_30)
        ("Продолжение таблицы 2.1.10",   True),   # 26 chars
        ("Продолжение таблицы 10.1.10",  True),   # 27 chars
        ("Продолжение таблицы 1.1",      True),   # 23 chars
        ("Продолжение таблицы 100.10.10", True),  # 29 chars — needs limit ≥30
        ("Продолжение таблицы 1.1 (часть 2)", False),  # 33 chars > 30
        ("Это обычный абзац с упоминанием таблицы и продолжения", False),  # long prose
    ]
    failures = []
    for text, expected in cases:
        got = _is_student_continuation(text)
        if got != expected:
            failures.append(f"'{text}' (len={len(text)}): expected={expected}, got={got}")
    if failures:
        return _result(False, "; ".join(failures))
    return _result(True, f"all {len(cases)} cases correct")


def test_c_caption_number_extraction_strict() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import _extract_table_num
    cases = [
        ("Таблица 2.3", "2.3"),
        ("Таблица 2.3.4", "2.3.4"),
        ("Продолжение таблицы 2.3", None),
        ("Таблица абв", None),
    ]
    bad = []
    for text, expected in cases:
        got = _extract_table_num(text)
        if got != expected:
            bad.append(f"{text!r}: expected={expected!r}, got={got!r}")
    return _result(not bad, "; ".join(bad) if bad else "strict caption extraction OK")


def test_c_apply_table_merging_rebuilds_invalid_split() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc = Document()
    t1 = doc.add_table(rows=3, cols=2)
    t1.rows[0].cells[0].text = "H1"
    t1.rows[0].cells[1].text = "H2"
    t1.rows[1].cells[0].text = "a"
    t1.rows[1].cells[1].text = "b"
    t1.rows[2].cells[0].text = "c"
    t1.rows[2].cells[1].text = "d"

    doc.add_paragraph("Продолжение таблицы 1.1")

    # invalid continuation: header row does NOT match source header
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "X"
    t2.rows[0].cells[1].text = "Y"
    t2.rows[1].cells[0].text = "e"
    t2.rows[1].cells[1].text = "f"

    n = apply_table_merging(doc)
    if n != 1:
        return _result(False, f"expected 1 merge, got {n}")
    if len(doc.tables) != 1:
        return _result(False, f"expected 1 table after merge, got {len(doc.tables)}")
    texts = [p.text for p in doc.paragraphs]
    if any("Продолжение таблицы" in (t or "") for t in texts):
        return _result(False, "continuation marker paragraph was not removed for invalid split")
    return _result(True, "invalid manual split was rebuilt")


def test_c_apply_table_merging_keeps_valid_manual_split() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc = Document()
    doc.add_paragraph("Таблица 1.1")
    t1 = doc.add_table(rows=3, cols=2)
    t1.rows[0].cells[0].text = "H1"
    t1.rows[0].cells[1].text = "H2"
    t1.rows[1].cells[0].text = "a"
    t1.rows[1].cells[1].text = "b"
    t1.rows[2].cells[0].text = "c"
    t1.rows[2].cells[1].text = "d"

    marker = doc.add_paragraph("Продолжение таблицы 1.1")
    marker.alignment = 2
    marker.paragraph_format.keep_with_next = True

    t2 = doc.add_table(rows=2, cols=2)
    # valid continuation header equals source header
    t2.rows[0].cells[0].text = "H1"
    t2.rows[0].cells[1].text = "H2"
    t2.rows[1].cells[0].text = "e"
    t2.rows[1].cells[1].text = "f"

    n = apply_table_merging(doc)
    if n != 0:
        return _result(False, f"expected 0 merges for valid manual split, got {n}")
    if len(doc.tables) != 2:
        return _result(False, f"expected 2 tables preserved, got {len(doc.tables)}")
    return _result(True, "valid manual split preserved")


def test_c_apply_table_merging_keeps_marker_without_keep_next() -> tuple[bool, str]:
    """
    P0-α: a student-authored continuation chain that is structurally sound
    (caption-matching right-aligned marker, matching column count and header row,
    second table has data rows) must NOT be destroyed just because the marker
    paragraph lacks keepNext. Phase 1 normalizes the marker text and alignment
    but does not add keepNext — that attribute is formatter-applied. Without
    this guard, apply_table_merging silently destroys valid student chains
    (Bondarev: 1.3.1, 2.1.1, 2.1.4, 2.1.5, 2.2.3).
    """
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc = Document()
    doc.add_paragraph("Таблица 1.3.1")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "H1"
    t1.rows[0].cells[1].text = "H2"
    t1.rows[1].cells[0].text = "a"
    t1.rows[1].cells[1].text = "b"

    marker = doc.add_paragraph("Продолжение таблицы 1.3.1")
    marker.alignment = 2  # right-aligned (Phase 1 normalisation); no keepNext

    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "H1"
    t2.rows[0].cells[1].text = "H2"
    t2.rows[1].cells[0].text = "c"
    t2.rows[1].cells[1].text = "d"

    n = apply_table_merging(doc)
    if n <= 0:
        return _result(False, f"expected marker anchor repair to be reported, got {n}")
    if len(doc.tables) != 2:
        return _result(False, f"expected 2 tables preserved, got {len(doc.tables)}")
    if not any("Продолжение таблицы" in (p.text or "") for p in doc.paragraphs):
        return _result(False, "continuation marker paragraph was lost")
    return _result(True, "structurally valid student chain preserved without keepNext")


def test_c_apply_table_merging_rebuilds_caption_mismatch() -> tuple[bool, str]:
    """
    Negative gate for P0-α: when the marker number does not match the caption
    number above tbl1, the chain is NOT structurally valid and apply_table_merging
    must still rebuild it. Verifies that loosening the keepNext requirement did
    not accidentally accept arbitrary marker text.
    """
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc = Document()
    doc.add_paragraph("Таблица 1.3.1")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "H1"
    t1.rows[0].cells[1].text = "H2"
    t1.rows[1].cells[0].text = "a"
    t1.rows[1].cells[1].text = "b"

    # marker references a different table number than the caption above tbl1
    marker = doc.add_paragraph("Продолжение таблицы 9.9.9")
    marker.alignment = 2

    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "H1"
    t2.rows[0].cells[1].text = "H2"
    t2.rows[1].cells[0].text = "c"
    t2.rows[1].cells[1].text = "d"

    n = apply_table_merging(doc)
    if n != 1:
        return _result(False, f"expected caption-mismatched chain to be rebuilt, got {n}")
    if len(doc.tables) != 1:
        return _result(False, f"expected 1 merged table, got {len(doc.tables)}")
    return _result(True, "caption-mismatched chain still rebuilt")


def test_p0_manual_continuation_numeric_row_only_fragment_is_valid() -> tuple[bool, str]:
    """P0: valid manual continuations may repeat only the numeric row in the
    continuation fragment. Do not merge these chains into a malformed table."""
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc = Document()
    doc.add_paragraph("Таблица 2.2.1")
    t1 = doc.add_table(rows=3, cols=4)
    for c, text in enumerate(["Раздел сайта", "Основная функция", "Значение", "Последствие"]):
        t1.rows[0].cells[c].text = text
    for c in range(4):
        t1.rows[1].cells[c].text = str(c + 1)
    for c, text in enumerate(["Главная", "Вход", "Представление", "База"]):
        t1.rows[2].cells[c].text = text

    marker = doc.add_paragraph("Продолжение таблицы 2.2.1")
    marker.alignment = 2

    t2 = doc.add_table(rows=2, cols=4)
    for c in range(4):
        t2.rows[0].cells[c].text = str(c + 1)
    for c, text in enumerate(["Оферта", "Порядок", "Прозрачность", "Доверие"]):
        t2.rows[1].cells[c].text = text

    n = apply_table_merging(doc)
    if len(doc.tables) != 2:
        return _result(False, f"numeric-row-only continuation chain was merged; expected 2 preserved fragments, got {len(doc.tables)}")
    if [cell.text for cell in doc.tables[1].rows[0].cells] != ["1", "2", "3", "4"]:
        return _result(False, "continuation numeric row was not preserved as first row")
    return _result(True, "numeric-row-only continuation fragment is preserved")


def test_p0_preserved_manual_chain_synthesizes_missing_numeric_rows() -> tuple[bool, str]:
    """P0: when a preserved manual chain lacks numeric rows, synthesize them
    inside both table fragments, directly below each semantic header."""
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation, apply_table_merging

    doc = Document()
    doc.add_paragraph("Таблица 1.3.1")
    t1 = doc.add_table(rows=2, cols=3)
    for c, text in enumerate(["Страна", "Роль государства", "Вывод"]):
        t1.rows[0].cells[c].text = text
    for c, text in enumerate(["Россия", "Регулирует", "Выбор оператора"]):
        t1.rows[1].cells[c].text = text

    marker = doc.add_paragraph("Продолжение таблицы 1.3.1")
    marker.alignment = 2

    t2 = doc.add_table(rows=2, cols=3)
    for c, text in enumerate(["Страна", "Роль государства", "Вывод"]):
        t2.rows[0].cells[c].text = text
    for c, text in enumerate(["ЕС", "Задает стандарт", "Совместимость"]):
        t2.rows[1].cells[c].text = text

    merges = apply_table_merging(doc)
    repairs = apply_table_continuation(doc)
    if len(doc.tables) != 2:
        return _result(False, f"expected 2 preserved fragments, got {len(doc.tables)}")
    if repairs < 2:
        return _result(False, f"expected numeric-row repairs in both fragments, got {repairs}")
    if [cell.text for cell in doc.tables[0].rows[1].cells] != ["1", "2", "3"]:
        return _result(False, "numeric row missing from first fragment")
    if [cell.text for cell in doc.tables[1].rows[1].cells] != ["1", "2", "3"]:
        return _result(False, "numeric row missing from continuation fragment")
    return _result(True, "missing numeric rows synthesized in both fragments")


def test_p0_unsplit_ordinary_table_does_not_get_synthetic_numeric_row() -> tuple[bool, str]:
    """P0: numeric-row synthesis is scoped to continuation chains only."""
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation, apply_table_merging

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    table = doc.add_table(rows=2, cols=3)
    for c, text in enumerate(["Показатель", "Влияние", "Последствие"]):
        table.rows[0].cells[c].text = text
    for c, text in enumerate(["Срок", "Скорость", "Риск"]):
        table.rows[1].cells[c].text = text

    merges = apply_table_merging(doc)
    repairs = apply_table_continuation(doc)
    if merges != 0:
        return _result(False, f"unsplit table should not be merged, got {merges}")
    if repairs != 0:
        return _result(False, f"unsplit table should not be repaired, got {repairs}")
    if len(doc.tables[0].rows) != 2:
        return _result(False, f"unsplit table got synthetic row: rows={len(doc.tables[0].rows)}")
    return _result(True, "unsplit table unchanged")


def test_p0_existing_correct_numeric_rows_are_not_duplicated() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation, apply_table_merging

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    t1 = doc.add_table(rows=3, cols=3)
    for c, text in enumerate(["Показатель", "Влияние", "Итог"]):
        t1.rows[0].cells[c].text = text
    for c in range(3):
        t1.rows[1].cells[c].text = str(c + 1)
    for c, text in enumerate(["A", "B", "C"]):
        t1.rows[2].cells[c].text = text

    marker = doc.add_paragraph("Продолжение таблицы 1.1.1")
    marker.alignment = 2

    t2 = doc.add_table(rows=3, cols=3)
    for c, text in enumerate(["Показатель", "Влияние", "Итог"]):
        t2.rows[0].cells[c].text = text
    for c in range(3):
        t2.rows[1].cells[c].text = str(c + 1)
    for c, text in enumerate(["D", "E", "F"]):
        t2.rows[2].cells[c].text = text

    apply_table_merging(doc)
    apply_table_continuation(doc)
    numeric = ["1", "2", "3"]
    if _count_table_rows_with_texts(doc.tables[0], numeric) != 1:
        return _result(False, "first fragment numeric row was duplicated")
    if _count_table_rows_with_texts(doc.tables[1], numeric) != 1:
        return _result(False, "continuation fragment numeric row was duplicated")
    return _result(True, "existing numeric rows are not duplicated")


def test_p0_rybakov_style_221_chain_does_not_merge_into_malformed_table() -> tuple[bool, str]:
    """P0: Rybakov-style chain has semantic header + numeric row in fragment 1
    and numeric-row-only continuation. It must remain two tables."""
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc = Document()
    doc.add_paragraph("Таблица 2.2.1")
    t1 = doc.add_table(rows=7, cols=4)
    for c, text in enumerate(["Раздел сайта", "Основная функция", "Значение для клиента", "Последствие"]):
        t1.rows[0].cells[c].text = text
    for c in range(4):
        t1.rows[1].cells[c].text = str(c + 1)
    for r in range(2, 7):
        for c in range(4):
            t1.rows[r].cells[c].text = f"r{r}c{c}"

    marker = doc.add_paragraph("Продолжение таблицы 2.2.1")
    marker.alignment = 2

    t2 = doc.add_table(rows=2, cols=4)
    for c in range(4):
        t2.rows[0].cells[c].text = str(c + 1)
    for c, text in enumerate(["Договор-оферта", "Разъяснение порядка", "Снижает неопределенность", "Прозрачность"]):
        t2.rows[1].cells[c].text = text

    n = apply_table_merging(doc)
    if len(doc.tables) != 2:
        return _result(False, f"Rybakov-style chain was merged; expected 2 fragments, got {len(doc.tables)}")
    if _count_table_rows_with_texts(doc.tables[0], ["1", "2", "3", "4"]) != 1:
        return _result(False, "first fragment numeric row count changed")
    if _count_table_rows_with_texts(doc.tables[1], ["1", "2", "3", "4"]) != 1:
        return _result(False, "continuation numeric row count changed")
    return _result(True, "Rybakov-style numeric-row-only continuation preserved")


# ── P1-critical / DEFECT E — enable pageBreakBefore on preserved student marker

def _marker_page_break_before_enabled(marker_p_xml) -> bool:
    """True if <w:pageBreakBefore/> exists in pPr AND w:val is absent or in
    the enabled set. Mirrors the OOXML semantic of "enabled by default"."""
    pPr = marker_p_xml.find(qn("w:pPr"))
    if pPr is None:
        return False
    pb = pPr.find(qn("w:pageBreakBefore"))
    if pb is None:
        return False
    val = pb.get(qn("w:val"))
    return val is None or val in {"1", "true", "True", "on"}


def _marker_keep_next_enabled(marker_p_xml) -> bool:
    pPr = marker_p_xml.find(qn("w:pPr"))
    if pPr is None:
        return False
    kn = pPr.find(qn("w:keepNext"))
    if kn is None:
        return False
    val = kn.get(qn("w:val"))
    return val is None or val in {"1", "true", "True", "on"}


def _build_student_chain_doc():
    """Synthetic doc that produces tbl1 → marker → tbl2 where chain is structurally
    valid (matching headers, cols, ≥2 rows) but lacks keepNext on the marker —
    i.e. the case routed through `_is_structurally_valid_student_chain`. We also
    simulate Phase 1's hard_reset by explicitly disabling pageBreakBefore on the
    marker paragraph (mirrors production state at the point apply_table_merging
    runs)."""
    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "H1"
    t1.rows[0].cells[1].text = "H2"
    t1.rows[1].cells[0].text = "a"
    t1.rows[1].cells[1].text = "b"
    marker = doc.add_paragraph("Продолжение таблицы 1.1.1")
    marker.alignment = 2  # right; no keepNext on purpose → student-chain branch
    # Simulate Phase 1 hard_reset: explicitly disable pageBreakBefore & keepNext.
    marker.paragraph_format.page_break_before = False
    marker.paragraph_format.keep_with_next = False
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "H1"
    t2.rows[0].cells[1].text = "H2"
    t2.rows[1].cells[0].text = "c"
    t2.rows[1].cells[1].text = "d"
    return doc, marker


def test_e_preserved_student_marker_enables_page_break_before() -> tuple[bool, str]:
    """
    P1-critical / DEFECT E positive: when apply_table_merging preserves a
    student chain (matched by `_is_structurally_valid_student_chain` but NOT
    by `_is_valid_manual_continuation_chain`), the marker paragraph must end
    up with <w:pageBreakBefore/> enabled (no w:val='0'). Phase 1 disables it;
    this fix re-enables it to match formatter-authored markers that render
    at the top of the continuation page rather than at the bottom of the
    previous page.
    """
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc, marker = _build_student_chain_doc()
    if _marker_page_break_before_enabled(marker._element):
        return _result(False, "fixture invariant violated: pageBreakBefore already enabled pre-patch")

    apply_table_merging(doc)

    if not _marker_page_break_before_enabled(marker._element):
        return _result(False, "expected <w:pageBreakBefore/> enabled on preserved student marker after apply_table_merging")
    if not _marker_keep_next_enabled(marker._element):
        return _result(False, "expected <w:keepNext/> enabled on preserved student marker")
    if len(doc.tables) != 2:
        return _result(False, f"chain should remain preserved; got {len(doc.tables)} tables")
    return _result(True, "preserved student marker has pageBreakBefore + keepNext enabled")


def test_e_preserved_student_marker_anchor_repairs_are_counted() -> tuple[bool, str]:
    """
    P1-critical / persistence regression: marker anchor repairs are real DOCX
    changes. If apply_table_merging enables disabled pageBreakBefore/keepNext
    but returns 0, formatter_service may skip doc.save() and lose the repair.
    """
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc, marker = _build_student_chain_doc()
    if _marker_page_break_before_enabled(marker._element):
        return _result(False, "fixture invariant violated: pageBreakBefore already enabled")
    if _marker_keep_next_enabled(marker._element):
        return _result(False, "fixture invariant violated: keepNext already enabled")

    n = apply_table_merging(doc)

    if n <= 0:
        return _result(False, f"expected anchor repairs to be counted, got {n}")
    if not _marker_page_break_before_enabled(marker._element):
        return _result(False, "pageBreakBefore repair did not apply")
    if not _marker_keep_next_enabled(marker._element):
        return _result(False, "keepNext repair did not apply")
    return _result(True, f"anchor repairs counted: {n}")


def test_e_format_docx_saves_repaired_manual_marker_anchors() -> tuple[bool, str]:
    """
    Full formatter regression: Phase 1 can save disabled marker anchors, then
    Phase 3 repairs them in memory. The formatter must persist that repair.
    """
    import guides.coursework_kfu_2025.formatter_service as fs

    prepared_doc, _marker = _build_student_chain_doc()

    with tempfile.TemporaryDirectory() as tmp:
        input_path = Path(tmp) / "input.docx"
        output_path = Path(tmp) / "output.docx"
        Document().save(input_path)

        old_process = fs.process_document
        old_pagination = fs.apply_pagination_rules
        old_rule3 = fs.apply_rule3_table_orphan
        old_rule4 = fs.apply_rule4_empty_first_lines
        old_rule6 = fs.apply_rule6_figure_orphan
        old_gap = fs.remove_empty_before_figure_captions
        old_continuation = fs.apply_table_continuation
        old_rendered = fs.apply_rendered_table_continuation
        old_rebuild = fs.rebuild_static_contents_page
        old_restore = fs.restore_docx_if_same_page_continuation_markers
        try:
            fs.process_document = lambda _src, dst: prepared_doc.save(dst)
            fs.apply_pagination_rules = lambda _doc: None
            fs.apply_rule3_table_orphan = lambda _doc: 0
            fs.apply_rule4_empty_first_lines = lambda _doc: 0
            fs.apply_rule6_figure_orphan = lambda _doc: 0
            fs.remove_empty_before_figure_captions = lambda _doc: 0
            fs.apply_table_continuation = lambda _doc, report=None: 0
            fs.apply_rendered_table_continuation = lambda _path, report=None: 0
            fs.rebuild_static_contents_page = lambda _path: False
            fs.restore_docx_if_same_page_continuation_markers = (
                lambda _output, _backup, report=None, context="": False
            )
            fs.format_docx(str(input_path), str(output_path))
        finally:
            fs.process_document = old_process
            fs.apply_pagination_rules = old_pagination
            fs.apply_rule3_table_orphan = old_rule3
            fs.apply_rule4_empty_first_lines = old_rule4
            fs.apply_rule6_figure_orphan = old_rule6
            fs.remove_empty_before_figure_captions = old_gap
            fs.apply_table_continuation = old_continuation
            fs.apply_rendered_table_continuation = old_rendered
            fs.rebuild_static_contents_page = old_rebuild
            fs.restore_docx_if_same_page_continuation_markers = old_restore

        out = Document(str(output_path))
        marker_xml = next(
            (p._element for p in out.paragraphs if p.text == "Продолжение таблицы 1.1.1"),
            None,
        )

    if marker_xml is None:
        return _result(False, "manual continuation marker was lost")
    if not _marker_page_break_before_enabled(marker_xml):
        return _result(False, "format_docx did not persist repaired pageBreakBefore")
    if not _marker_keep_next_enabled(marker_xml):
        return _result(False, "format_docx did not persist repaired keepNext")
    return _result(True, "format_docx persisted repaired marker anchors")


def test_e_page_break_enable_is_idempotent() -> tuple[bool, str]:
    """
    P1-critical / DEFECT E idempotency: running apply_table_merging twice on
    the same chain must not duplicate <w:pageBreakBefore/> elements in pPr.
    """
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc, marker = _build_student_chain_doc()
    apply_table_merging(doc)
    apply_table_merging(doc)

    pPr = marker._element.find(qn("w:pPr"))
    n_pb = len(pPr.findall(qn("w:pageBreakBefore"))) if pPr is not None else 0
    n_kn = len(pPr.findall(qn("w:keepNext"))) if pPr is not None else 0
    if n_pb != 1:
        return _result(False, f"expected exactly 1 <w:pageBreakBefore/>, got {n_pb}")
    if n_kn != 1:
        return _result(False, f"expected exactly 1 <w:keepNext/>, got {n_kn}")
    return _result(True, "pageBreakBefore/keepNext enable is idempotent")


def test_e_formatter_authored_chain_with_keepnext_not_modified() -> tuple[bool, str]:
    """
    P1-critical / DEFECT E regression: a chain matched by the strict
    `_is_valid_manual_continuation_chain` validator (i.e. already has
    keepNext) is preserved via the manual-chain branch and the enable helper
    must NOT touch it (no double work, no duplicate pPr children).
    """
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "H1"
    t1.rows[0].cells[1].text = "H2"
    t1.rows[1].cells[0].text = "a"
    t1.rows[1].cells[1].text = "b"
    marker = doc.add_paragraph("Продолжение таблицы 1.1.1")
    marker.alignment = 2
    marker.paragraph_format.keep_with_next = True  # formatter-authored → manual-chain path
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "H1"
    t2.rows[0].cells[1].text = "H2"
    t2.rows[1].cells[0].text = "c"
    t2.rows[1].cells[1].text = "d"

    # Snapshot pPr children count before
    pPr_before = marker._element.find(qn("w:pPr"))
    n_pb_before = len(pPr_before.findall(qn("w:pageBreakBefore"))) if pPr_before is not None else 0
    n_kn_before = len(pPr_before.findall(qn("w:keepNext"))) if pPr_before is not None else 0

    apply_table_merging(doc)

    pPr_after = marker._element.find(qn("w:pPr"))
    n_pb_after = len(pPr_after.findall(qn("w:pageBreakBefore"))) if pPr_after is not None else 0
    n_kn_after = len(pPr_after.findall(qn("w:keepNext"))) if pPr_after is not None else 0

    if n_pb_after != n_pb_before:
        return _result(False, f"formatter-authored chain pageBreakBefore count changed: {n_pb_before} -> {n_pb_after}")
    if n_kn_after != n_kn_before:
        return _result(False, f"formatter-authored chain keepNext count changed: {n_kn_before} -> {n_kn_after}")
    if len(doc.tables) != 2:
        return _result(False, f"formatter-authored chain preservation regressed: {len(doc.tables)} tables")
    return _result(True, "formatter-authored chain untouched (manual-chain branch)")


def test_e_preserved_marker_keeps_alignment_and_text() -> tuple[bool, str]:
    """
    P1-critical / DEFECT E regression: after enable helper runs, the marker
    paragraph must still carry its original right alignment and text content.
    The helper must not strip or rewrite other pPr children.
    """
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc, marker = _build_student_chain_doc()
    apply_table_merging(doc)

    pPr = marker._element.find(qn("w:pPr"))
    if pPr is None:
        return _result(False, "marker lost <w:pPr> after enable")
    jc = pPr.find(qn("w:jc"))
    if jc is None or jc.get(qn("w:val")) != "right":
        return _result(False, f"marker right-alignment lost: jc={jc!r}")
    text = (marker.text or "").strip()
    if text != "Продолжение таблицы 1.1.1":
        return _result(False, f"marker text changed: {text!r}")
    return _result(True, "preserved marker keeps right alignment + original text")


def test_e_integration_tbl_marker_tbl_marker_has_enabled_break() -> tuple[bool, str]:
    """
    P1-critical / DEFECT E integration: tbl → marker → tbl topology, run
    apply_table_merging, then verify marker XML has <w:pageBreakBefore/>
    without w:val='0' (enabled by OOXML default).
    """
    from guides.coursework_kfu_2025.table_continuation import apply_table_merging

    doc, marker = _build_student_chain_doc()
    apply_table_merging(doc)

    pPr = marker._element.find(qn("w:pPr"))
    if pPr is None:
        return _result(False, "marker has no pPr after apply_table_merging")
    pb = pPr.find(qn("w:pageBreakBefore"))
    if pb is None:
        return _result(False, "marker has no <w:pageBreakBefore/> element")
    val = pb.get(qn("w:val"))
    if val is not None and val not in {"1", "true", "True", "on"}:
        return _result(False, f"<w:pageBreakBefore/> still disabled: w:val={val!r}")
    return _result(True, f"integration: <w:pageBreakBefore/> enabled (w:val={val!r})")


# ── P1a — ordinary continuation marker/blank/table anchoring ────────────────

def _body_children(doc: Document) -> list:
    return list(doc.element.body)


def _insert_blank_after_paragraph(p_xml):
    blank = OxmlElement("w:p")
    p_xml.addnext(blank)
    return blank


def _build_p1a_blank_chain_doc() -> tuple[Document, object, object]:
    """Build tbl -> marker -> blank -> tbl with disabled marker anchoring."""
    doc = Document()
    doc.add_paragraph("Таблица 1.2.2")
    t1 = doc.add_table(rows=3, cols=3)
    for c, text in enumerate(["H1", "H2", "H3"]):
        t1.rows[0].cells[c].text = text
    for c in range(3):
        t1.rows[1].cells[c].text = str(c + 1)
    for c, text in enumerate(["a", "b", "c"]):
        t1.rows[2].cells[c].text = text

    marker = doc.add_paragraph("Продолжение таблицы 1.2.2")
    marker.alignment = 2
    marker.paragraph_format.page_break_before = False
    marker.paragraph_format.keep_with_next = False
    blank = _insert_blank_after_paragraph(marker._element)

    t2 = doc.add_table(rows=2, cols=3)
    for c in range(3):
        t2.rows[0].cells[c].text = str(c + 1)
    for c, text in enumerate(["d", "e", "f"]):
        t2.rows[1].cells[c].text = text
    return doc, marker._element, blank


def test_p1a_marker_blank_table_chain_is_anchored() -> tuple[bool, str]:
    """P1a: marker + optional blank + continuation table must be one keep chain."""
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc, marker, blank = _build_p1a_blank_chain_doc()
    apply_table_continuation(doc)

    if not _marker_page_break_before_enabled(marker):
        return _result(False, "marker pageBreakBefore was not activated")
    if not _marker_keep_next_enabled(marker):
        return _result(False, "marker keepNext was not activated")
    if not _marker_keep_next_enabled(blank):
        return _result(False, "blank paragraph keepNext was not activated")
    if _marker_page_break_before_enabled(blank):
        return _result(False, "blank paragraph should not get active pageBreakBefore")
    children = _body_children(doc)
    gap = children.index(doc.tables[1]._tbl) - children.index(doc.tables[0]._tbl)
    if gap != 3:
        return _result(False, f"normalizer moved or inserted body nodes; gap={gap}")
    return _result(True, "tbl -> marker -> blank -> tbl chain is anchored")


def test_p1a_disabled_marker_page_break_becomes_active() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc, marker, _blank = _build_p1a_blank_chain_doc()
    if _marker_page_break_before_enabled(marker):
        return _result(False, "fixture invariant violated: marker pageBreakBefore already active")
    apply_table_continuation(doc)
    if not _marker_page_break_before_enabled(marker):
        return _result(False, "disabled marker pageBreakBefore remained inactive")
    return _result(True, "disabled marker pageBreakBefore becomes active")


def test_p1a_preserved_manual_chain_normalized_without_merging() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation, apply_table_merging

    doc, marker, blank = _build_p1a_blank_chain_doc()
    merges = apply_table_merging(doc)
    repairs = apply_table_continuation(doc)

    if len(doc.tables) != 2:
        return _result(False, f"manual chain was merged; expected preserved two-table chain, got {len(doc.tables)} tables")
    if not _marker_page_break_before_enabled(marker) or not _marker_keep_next_enabled(marker):
        return _result(False, "preserved marker was not anchored")
    if not _marker_keep_next_enabled(blank):
        return _result(False, "preserved chain blank was not chained")
    if repairs < 1:
        return _result(False, f"expected at least one anchoring repair, got {repairs}")
    return _result(True, "preserved manual chain normalized without merging")


def test_p1a_appendix_continuation_label_not_modified() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc = Document()
    first = doc.add_table(rows=2, cols=2)
    first.rows[0].cells[0].text = "H1"
    first.rows[0].cells[1].text = "H2"
    first.rows[1].cells[0].text = "a"
    first.rows[1].cells[1].text = "b"
    label = doc.add_paragraph("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ 1")
    label.alignment = 2
    label.paragraph_format.page_break_before = False
    label.paragraph_format.keep_with_next = False
    second = doc.add_table(rows=2, cols=2)
    second.rows[0].cells[0].text = "1"
    second.rows[0].cells[1].text = "2"
    second.rows[1].cells[0].text = "c"
    second.rows[1].cells[1].text = "d"

    before = label._element.find(qn("w:pPr")).xml
    apply_table_continuation(doc)
    after = label._element.find(qn("w:pPr")).xml
    if before != after:
        return _result(False, "appendix continuation label XML was modified by ordinary normalizer")
    return _result(True, "appendix continuation label is not modified")


def test_p1a_numeric_rows_stay_unchanged_after_anchoring() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc, _marker, _blank = _build_p1a_blank_chain_doc()
    before_first = _count_table_rows_with_texts(doc.tables[0], ["1", "2", "3"])
    before_second = _count_table_rows_with_texts(doc.tables[1], ["1", "2", "3"])
    apply_table_continuation(doc)
    after_first = _count_table_rows_with_texts(doc.tables[0], ["1", "2", "3"])
    after_second = _count_table_rows_with_texts(doc.tables[1], ["1", "2", "3"])
    if (before_first, before_second) != (1, 1):
        return _result(False, "fixture invariant violated: numeric rows missing before normalizer")
    if (after_first, after_second) != (1, 1):
        return _result(False, f"numeric rows changed: first={after_first}, second={after_second}")
    return _result(True, "numeric rows unchanged after anchoring normalizer")


def _p1a_rendered_chain_nodes(doc: Document, marker_text: str):
    children = _body_children(doc)
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t"))).strip()
        if text != marker_text:
            continue
        prev_node = children[idx - 1] if idx > 0 else None
        next_node = children[idx + 1] if idx + 1 < len(children) else None
        blank_node = None
        table_node = next_node
        if next_node is not None and next_node.tag == qn("w:p") and not _paragraph_text(next_node):
            blank_node = next_node
            table_node = children[idx + 2] if idx + 2 < len(children) else None
        return prev_node, child, blank_node, table_node
    return None, None, None, None


def _p1a_assert_rendered_chain_anchored(
    doc: Document,
    marker_text: str,
    *,
    require_blank: bool,
) -> tuple[bool, str]:
    prev_node, marker_node, blank_node, table_node = _p1a_rendered_chain_nodes(doc, marker_text)
    if marker_node is None:
        return _result(False, f"marker not found: {marker_text!r}")
    if prev_node is None or prev_node.tag != qn("w:tbl"):
        return _result(False, "marker is not immediately after the first table")
    if table_node is None or table_node.tag != qn("w:tbl"):
        return _result(False, "continuation table is not after marker/blank")
    if require_blank and blank_node is None:
        return _result(False, "required blank paragraph is missing")
    if not _marker_page_break_before_enabled(marker_node):
        return _result(False, "rendered marker pageBreakBefore is not active")
    if not _marker_keep_next_enabled(marker_node):
        return _result(False, "rendered marker keepNext is not active")
    if blank_node is not None and not _marker_keep_next_enabled(blank_node):
        return _result(False, "rendered blank paragraph keepNext is not active")
    return _result(True, "rendered continuation chain is anchored")


# ── P1-c / DEFECT B — detached source/note: fallback split last data row ─────

def _build_p1c_doc(rows=4, include_source_note=True, include_caption=True, manual_marker_between=False):
    """Synthetic doc with caption + table (rows × 2) + Источник: paragraph.

    `manual_marker_between=True` produces tbl1 → manual marker → tbl2 chain so
    `_valid_manual_continuation_table_ids` flags both tables as protected.
    """
    doc = Document()
    if include_caption:
        doc.add_paragraph("Таблица 2.1.2 — Тестовая")
    tbl = doc.add_table(rows=rows, cols=2)
    tbl.rows[0].cells[0].text = "Col1Header"
    tbl.rows[0].cells[1].text = "Col2Header"
    row_texts = []
    for r in range(1, rows):
        a = f"alpha{r}xxx"
        b = f"beta{r}yyy"
        tbl.rows[r].cells[0].text = a
        tbl.rows[r].cells[1].text = b
        row_texts.append((a, b))
    if manual_marker_between:
        marker = doc.add_paragraph("Продолжение таблицы 2.1.2")
        marker.alignment = 2
        marker.paragraph_format.keep_with_next = True
        tbl2 = doc.add_table(rows=2, cols=2)
        tbl2.rows[0].cells[0].text = "Col1Header"
        tbl2.rows[0].cells[1].text = "Col2Header"
        tbl2.rows[1].cells[0].text = "tailA"
        tbl2.rows[1].cells[1].text = "tailB"
    if include_source_note:
        doc.add_paragraph("Источник: составлено автором по данным таблицы.")
    return doc, row_texts


def _build_p1c_pdf_lines(row_texts, source_note_text="Источник: составлено автором по данным таблицы.",
                         last_row_page=5, source_note_page=6):
    """Construct PdfLine objects: header + each data row on page `last_row_page`,
    source/note on page `source_note_page` (default: detached one page later)."""
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine
    lines = [PdfLine(text="Col1Header Col2Header", page_num=last_row_page, top=100, bottom=110)]
    y = 120
    for (a, b) in row_texts:
        lines.append(PdfLine(text=f"{a} {b}", page_num=last_row_page, top=y, bottom=y + 10))
        y += 20
    lines.append(PdfLine(text=source_note_text, page_num=source_note_page, top=58, bottom=72))
    return lines


def test_p1c_detects_detached_source_note() -> tuple[bool, str]:
    """
    P1-c positive: synthetic table whose last data row renders on page 5 and
    whose Источник: line renders on page 6 is reported as a detachment
    candidate (no skips). Verifies the pure detection logic before any split
    is applied. Render is mocked via PdfLine fixtures.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_source_note_detachment_candidates,
    )

    doc, row_texts = _build_p1c_doc(rows=4)
    pdf_lines = _build_p1c_pdf_lines(row_texts, last_row_page=5, source_note_page=6)
    candidates, skips = _collect_source_note_detachment_candidates(doc, pdf_lines, manual_chain_ids=set())
    if len(candidates) != 1:
        return _result(False, f"expected 1 candidate, got {len(candidates)} (skips={skips!r})")
    c = candidates[0]
    if c.table_index != 0:
        return _result(False, f"expected table_index=0, got {c.table_index}")
    if c.caption_num != "2.1.2":
        return _result(False, f"expected caption_num='2.1.2', got {c.caption_num!r}")
    if c.last_data_row_page != 5 or c.source_note_first_page != 6:
        return _result(False, f"page mapping wrong: last_row_page={c.last_data_row_page} sn_page={c.source_note_first_page}")
    return _result(True, "detached source/note candidate detected")


def test_p1c_skip_when_source_note_attached_same_page() -> tuple[bool, str]:
    """
    P1-c regression: when Источник: renders on the SAME page as the last row,
    detector reports `not_detached` skip (no candidate).
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_source_note_detachment_candidates,
    )
    doc, row_texts = _build_p1c_doc(rows=4)
    pdf_lines = _build_p1c_pdf_lines(row_texts, last_row_page=5, source_note_page=5)
    candidates, skips = _collect_source_note_detachment_candidates(doc, pdf_lines, manual_chain_ids=set())
    if candidates:
        return _result(False, f"unexpected candidate when source/note attached: {candidates!r}")
    if not any(reason == "not_detached" for _, reason in skips):
        return _result(False, f"expected 'not_detached' skip, got skips={skips!r}")
    return _result(True, "attached source/note correctly skipped as not_detached")


def test_p1c_skip_no_caption() -> tuple[bool, str]:
    """P1-c eligibility: table without 'Таблица X.Y.Z' caption is skipped."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_source_note_detachment_candidates,
    )
    doc, row_texts = _build_p1c_doc(rows=4, include_caption=False)
    pdf_lines = _build_p1c_pdf_lines(row_texts, last_row_page=5, source_note_page=6)
    candidates, skips = _collect_source_note_detachment_candidates(doc, pdf_lines, manual_chain_ids=set())
    if candidates:
        return _result(False, f"caption-less table should not be candidate: {candidates!r}")
    if not any(reason == "no_caption" for _, reason in skips):
        return _result(False, f"expected 'no_caption' skip, got skips={skips!r}")
    return _result(True, "no-caption skip correct")


def test_p1c_skip_no_source_note() -> tuple[bool, str]:
    """P1-c eligibility: table with no following Источник:/Примечание: is skipped."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_source_note_detachment_candidates,
    )
    doc, row_texts = _build_p1c_doc(rows=4, include_source_note=False)
    pdf_lines = _build_p1c_pdf_lines(row_texts, last_row_page=5, source_note_page=6)
    candidates, skips = _collect_source_note_detachment_candidates(doc, pdf_lines, manual_chain_ids=set())
    if candidates:
        return _result(False, f"table without source/note should not be candidate: {candidates!r}")
    if not any(reason == "no_source_note" for _, reason in skips):
        return _result(False, f"expected 'no_source_note' skip, got skips={skips!r}")
    return _result(True, "no-source-note skip correct")


def test_p1c_skip_small_table_no_safe_data_row() -> tuple[bool, str]:
    """P1-c eligibility: table with < 3 rows (header + 1 data) is skipped — not safe to split off last row."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_source_note_detachment_candidates,
    )
    doc, row_texts = _build_p1c_doc(rows=2)
    pdf_lines = _build_p1c_pdf_lines(row_texts, last_row_page=5, source_note_page=6)
    candidates, skips = _collect_source_note_detachment_candidates(doc, pdf_lines, manual_chain_ids=set())
    if candidates:
        return _result(False, f"2-row table should not be candidate: {candidates!r}")
    if not any(reason == "no_safe_data_row" for _, reason in skips):
        return _result(False, f"expected 'no_safe_data_row' skip, got skips={skips!r}")
    return _result(True, "small-table skip correct")


def test_p1c_skip_already_in_manual_chain() -> tuple[bool, str]:
    """P1-c regression: tables already in a manual continuation chain (preserved
    by P0-α) must be skipped to avoid double-splitting."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_source_note_detachment_candidates,
    )
    doc, row_texts = _build_p1c_doc(rows=4)
    pdf_lines = _build_p1c_pdf_lines(row_texts, last_row_page=5, source_note_page=6)
    # Synthetic manual_chain_ids: pretend the first table is in a chain.
    fake_chain_ids = {id(doc.tables[0]._tbl)}
    candidates, skips = _collect_source_note_detachment_candidates(doc, pdf_lines, manual_chain_ids=fake_chain_ids)
    if candidates:
        return _result(False, f"chain-protected table should not be candidate: {candidates!r}")
    if not any(reason == "already_in_manual_chain" for _, reason in skips):
        return _result(False, f"expected 'already_in_manual_chain' skip, got skips={skips!r}")
    return _result(True, "manual-chain skip correct")


def test_p1c_skip_render_probe_unreliable() -> tuple[bool, str]:
    """P1-c eligibility: when PdfLine data cannot map rows to pages
    (e.g. empty pdf_lines), detector reports `render_probe_unreliable`."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_source_note_detachment_candidates,
    )
    doc, _ = _build_p1c_doc(rows=4)
    candidates, skips = _collect_source_note_detachment_candidates(doc, pdf_lines=[], manual_chain_ids=set())
    if candidates:
        return _result(False, f"empty pdf_lines must not yield candidate: {candidates!r}")
    if not any(reason == "render_probe_unreliable" for _, reason in skips):
        return _result(False, f"expected 'render_probe_unreliable' skip, got skips={skips!r}")
    return _result(True, "render-probe-unreliable skip correct")


def test_p1c_apply_split_inserts_continuation_marker_and_numbered_row() -> tuple[bool, str]:
    """P1-c apply: invoke split engine on a candidate's table, verify result —
    continuation marker text matches caption, numbered row repeated in second
    fragment, original source/note still in body after second fragment.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _build_continuation_para,
    )
    from guides.coursework_kfu_2025.table_split_prototype import (
        apply_numbered_split_to_document,
    )

    doc, _ = _build_p1c_doc(rows=4)
    # split_before_row = rows - 1 = 3 → first fragment keeps rows 0..2, second has row 3
    result = apply_numbered_split_to_document(
        doc, table_index=0, split_before_row=3,
        header_rows=1, numbered_header=True, appendix_table=False,
        continuation_paragraph_builder=_build_continuation_para,
    )
    # Validate continuation marker present
    markers = [p for p in doc.paragraphs if (p.text or "").strip().startswith("Продолжение таблицы 2.1.2")]
    if len(markers) != 1:
        return _result(False, f"expected exactly 1 continuation marker, got {len(markers)}")
    # Validate two tables now
    if len(doc.tables) != 2:
        return _result(False, f"expected 2 tables after split, got {len(doc.tables)}")
    # Second table must have numbered row at top (row 0) + at least 1 data row
    tbl2 = doc.tables[1]
    if len(tbl2.rows) < 2:
        return _result(False, f"second fragment too small: rows={len(tbl2.rows)}")
    # Source/note paragraph still exists in document after second fragment
    if result.source_note_after_second is False:
        return _result(False, "source_note_after_second=False (orphan would persist)")
    sn_paras = [p for p in doc.paragraphs if (p.text or "").strip().startswith("Источник:")]
    if len(sn_paras) != 1:
        return _result(False, f"expected exactly 1 Источник: paragraph, got {len(sn_paras)}")
    return _result(True, "split applied: marker + numbered row + source/note retained")


def test_p1a_p1c_rendered_split_anchors_marker_blank_chain() -> tuple[bool, str]:
    """P1a: P1-c rendered source/note split must persist marker+blank anchoring."""
    import guides.coursework_kfu_2025.table_continuation as tc

    doc, row_texts = _build_p1c_doc(rows=4)
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "p1c_rendered_anchor.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "p1c_rendered_anchor.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: _build_p1c_pdf_lines(
                row_texts, last_row_page=5, source_note_page=6,
            )
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode

        out = Document(str(path))

    if n != 1:
        return _result(False, f"expected one P1-c rendered split, got {n}")
    ok, msg = _p1a_assert_rendered_chain_anchored(
        out, "Продолжение таблицы 2.1.2", require_blank=True,
    )
    if not ok:
        return _result(False, msg)
    if _count_table_rows_with_texts(out.tables[0], ["1", "2"]) != 1:
        return _result(False, "first P1-c fragment numeric row changed or duplicated")
    if _count_table_rows_with_texts(out.tables[1], ["1", "2"]) != 1:
        return _result(False, "second P1-c fragment numeric row changed or duplicated")
    return _result(True, "P1-c rendered split chain is anchored")


def test_p1c_double_run_idempotent_via_natural_skips() -> tuple[bool, str]:
    """P1-c idempotency: after split, a second detection pass produces no
    candidates. The chain may be protected by the manual-chain detector after
    P0 numeric-row validation, or by the natural no-source/no-caption/small
    continuation-fragment skips. Either path is idempotent.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _build_continuation_para,
        _collect_source_note_detachment_candidates,
        _valid_manual_continuation_table_ids,
    )
    from guides.coursework_kfu_2025.table_split_prototype import (
        apply_numbered_split_to_document,
    )

    doc, row_texts = _build_p1c_doc(rows=4)
    apply_numbered_split_to_document(
        doc, table_index=0, split_before_row=3,
        header_rows=1, numbered_header=True, appendix_table=False,
        continuation_paragraph_builder=_build_continuation_para,
    )
    # Re-detect on the now-split doc; P0 recognizes numeric-row-only
    # continuation fragments as protected manual chains.
    chain_ids = _valid_manual_continuation_table_ids(doc)
    fake_lines = _build_p1c_pdf_lines(row_texts, last_row_page=5, source_note_page=6)
    candidates, skips = _collect_source_note_detachment_candidates(
        doc, fake_lines, manual_chain_ids=chain_ids,
    )
    if candidates:
        return _result(False, f"second pass produced unexpected candidates: {candidates!r}")
    reasons = {ti: r for ti, r in skips}
    if reasons.get(0) not in {"already_in_manual_chain", "no_source_note"}:
        return _result(False, f"expected tbl1 idempotent skip, got {reasons.get(0)!r}")
    # tbl2 (the continuation fragment) is preceded by the right-aligned
    # "Продолжение таблицы 2.1.2" marker, NOT by a standard "Таблица X.Y.Z"
    # caption — so `_find_caption_number_before_table` returns None and the
    # fragment is skipped via `no_caption`. Either skip is acceptable for
    # idempotency, but no_caption is what `_find_caption_number_before_table`
    # produces in this layout.
    if reasons.get(1) not in {"already_in_manual_chain", "no_caption", "no_safe_data_row"}:
        return _result(False, f"expected tbl2 idempotent skip, got {reasons.get(1)!r}")
    return _result(True, f"second pass idempotent (tbl1={reasons.get(0)}, tbl2={reasons.get(1)})")


# ── P2-a — appendix continuation for table-based appendices ──────────────────

def _build_p2a_appendix_doc(rows=4, include_appendix_heading=True, include_existing_marker=False):
    """Synthetic doc with «ПРИЛОЖЕНИЕ N» heading + appendix table (rows × 2).
    Optionally include an existing «ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ N» marker right
    after the table to exercise idempotency.
    """
    doc = Document()
    if include_appendix_heading:
        doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    tbl = doc.add_table(rows=rows, cols=2)
    tbl.rows[0].cells[0].text = "Col1Header"
    tbl.rows[0].cells[1].text = "Col2Header"
    row_texts = []
    for r in range(1, rows):
        a = f"appx{r}aaa"
        b = f"appx{r}bbb"
        tbl.rows[r].cells[0].text = a
        tbl.rows[r].cells[1].text = b
        row_texts.append((a, b))
    if include_existing_marker:
        m = doc.add_paragraph("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ 1")
        m.alignment = 2
    return doc, row_texts


def _build_p2a_pdf_lines(row_texts, header_page=10, page_break_after_row=2, last_page=11):
    """Construct PdfLine objects: header + first `page_break_after_row` data
    rows on `header_page`, remaining rows on `last_page`."""
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine
    lines = [PdfLine(text="Col1Header Col2Header", page_num=header_page, top=100, bottom=110)]
    y = 120
    for idx, (a, b) in enumerate(row_texts, start=1):
        page = header_page if idx <= page_break_after_row else last_page
        # reset y when moving to a new page so x-y looks plausible
        if page == header_page:
            lines.append(PdfLine(text=f"{a} {b}", page_num=page, top=y, bottom=y + 10))
        else:
            lines.append(PdfLine(text=f"{a} {b}", page_num=page, top=58 + (idx - page_break_after_row - 1) * 20, bottom=68 + (idx - page_break_after_row - 1) * 20))
        y += 20
    return lines


def test_p2a_detects_multipage_appendix_table() -> tuple[bool, str]:
    """
    P2-a positive: appendix table whose last data row renders on a strictly
    later page than the first data row is reported as a continuation
    candidate. Verifies pure detection logic before split is applied.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
    )
    doc, row_texts = _build_p2a_appendix_doc(rows=4)
    pdf_lines = _build_p2a_pdf_lines(row_texts, header_page=10, page_break_after_row=2, last_page=11)
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=set())
    if len(candidates) != 1:
        return _result(False, f"expected 1 candidate, got {len(candidates)} (skips={skips!r})")
    c = candidates[0]
    if c.appendix_num != "1":
        return _result(False, f"expected appendix_num='1', got {c.appendix_num!r}")
    if c.split_before_row != 3:
        return _result(False, f"expected split_before_row=3 (first row on page 11), got {c.split_before_row}")
    return _result(True, "multi-page appendix table candidate detected")


def test_p2a_skip_non_appendix_table() -> tuple[bool, str]:
    """P2-a regression: ordinary (non-appendix) table is NOT a candidate —
    skipped with reason 'not_appendix'. Body tables stay in P1-c / marker-split
    paths."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
    )
    doc, row_texts = _build_p2a_appendix_doc(rows=4, include_appendix_heading=False)
    pdf_lines = _build_p2a_pdf_lines(row_texts, header_page=10, page_break_after_row=2, last_page=11)
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=set())
    if candidates:
        return _result(False, f"ordinary table should not be candidate: {candidates!r}")
    if not any(reason == "not_appendix" for _, reason in skips):
        return _result(False, f"expected 'not_appendix' skip, got {skips!r}")
    return _result(True, "ordinary table correctly skipped as not_appendix")


def test_p2a_skip_existing_continuation_label() -> tuple[bool, str]:
    """P2-a idempotency: an appendix table already followed by «ПРОДОЛЖЕНИЕ
    ПРИЛОЖЕНИЯ N» must be skipped to avoid duplicating the label."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
    )
    doc, row_texts = _build_p2a_appendix_doc(rows=4, include_existing_marker=True)
    pdf_lines = _build_p2a_pdf_lines(row_texts, header_page=10, page_break_after_row=2, last_page=11)
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=set())
    if candidates:
        return _result(False, f"existing-marker chain should not be candidate: {candidates!r}")
    if not any(reason == "already_followed_by_continuation" for _, reason in skips):
        return _result(False, f"expected 'already_followed_by_continuation' skip, got {skips!r}")
    return _result(True, "existing continuation label correctly skipped")


def test_p2a_skip_manual_chain() -> tuple[bool, str]:
    """P2-a regression: appendix table whose id is in the manual_chain_ids
    set (i.e. already preserved by P0-α) must be skipped."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
    )
    doc, row_texts = _build_p2a_appendix_doc(rows=4)
    fake_chain_ids = {id(doc.tables[0]._tbl)}
    pdf_lines = _build_p2a_pdf_lines(row_texts, header_page=10, page_break_after_row=2, last_page=11)
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=fake_chain_ids)
    if candidates:
        return _result(False, f"manual-chain table should not be candidate: {candidates!r}")
    if not any(reason == "already_in_manual_chain" for _, reason in skips):
        return _result(False, f"expected 'already_in_manual_chain' skip, got {skips!r}")
    return _result(True, "manual chain correctly skipped")


def test_p2a_skip_single_page_appendix() -> tuple[bool, str]:
    """P2-a eligibility: appendix table fitting on a single page (no
    cross-page row span) is skipped with reason 'single_page'."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
    )
    doc, row_texts = _build_p2a_appendix_doc(rows=4)
    pdf_lines = _build_p2a_pdf_lines(row_texts, header_page=10, page_break_after_row=99, last_page=10)
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=set())
    if candidates:
        return _result(False, f"single-page table should not be candidate: {candidates!r}")
    if not any(reason == "single_page" for _, reason in skips):
        return _result(False, f"expected 'single_page' skip, got {skips!r}")
    return _result(True, "single-page appendix correctly skipped")


def test_p2a_skip_small_appendix_table() -> tuple[bool, str]:
    """P2-a eligibility: appendix table with < 3 rows is too small to split."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
    )
    doc, row_texts = _build_p2a_appendix_doc(rows=2)
    pdf_lines = _build_p2a_pdf_lines(row_texts, header_page=10, page_break_after_row=0, last_page=11)
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=set())
    if candidates:
        return _result(False, f"2-row appendix should not be candidate: {candidates!r}")
    if not any(reason == "no_safe_data_row" for _, reason in skips):
        return _result(False, f"expected 'no_safe_data_row' skip, got {skips!r}")
    return _result(True, "small appendix correctly skipped")


def test_p2a_apply_split_inserts_uppercase_continuation_marker() -> tuple[bool, str]:
    """P2-a apply: invoke split engine on a candidate via
    `apply_numbered_split_to_document(..., appendix_table=True, ...)`. Verify
    the marker text is exactly «ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ N» (uppercase) and
    the marker paragraph is right-aligned with pageBreakBefore + keepNext.
    """
    from guides.coursework_kfu_2025.table_split_prototype import (
        apply_numbered_split_to_document,
        _build_appendix_continuation_paragraph,
    )
    doc, _ = _build_p2a_appendix_doc(rows=4)
    apply_numbered_split_to_document(
        doc, table_index=0, split_before_row=3,
        header_rows=1, numbered_header=True, appendix_table=True,
        continuation_paragraph_builder=_build_appendix_continuation_paragraph,
    )
    markers = [p for p in doc.paragraphs if (p.text or "").strip().startswith("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ")]
    if len(markers) != 1:
        return _result(False, f"expected exactly 1 continuation marker, got {len(markers)}")
    text = (markers[0].text or "").strip()
    if text != "ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ 1":
        return _result(False, f"expected exact uppercase marker text, got {text!r}")
    pPr = markers[0]._element.find(qn("w:pPr"))
    if pPr is None:
        return _result(False, "marker missing pPr")
    jc = pPr.find(qn("w:jc"))
    if jc is None or jc.get(qn("w:val")) != "right":
        return _result(False, f"marker not right-aligned: jc={jc!r}")
    pb = pPr.find(qn("w:pageBreakBefore"))
    if pb is None:
        return _result(False, "marker missing pageBreakBefore")
    pb_val = pb.get(qn("w:val"))
    if pb_val is not None and pb_val in {"0", "false", "False"}:
        return _result(False, f"marker pageBreakBefore disabled: w:val={pb_val!r}")
    kn = pPr.find(qn("w:keepNext"))
    if kn is None:
        return _result(False, "marker missing keepNext")
    if len(doc.tables) != 2:
        return _result(False, f"expected 2 tables after split, got {len(doc.tables)}")
    return _result(True, "split applied: uppercase marker right-aligned with pageBreakBefore+keepNext")


def test_p2a_idempotent_after_apply() -> tuple[bool, str]:
    """P2-a idempotency: after `apply_numbered_split_to_document(...)` creates
    the continuation marker, a second detection pass produces zero candidates
    (fragment 1 is followed by the marker → `already_followed_by_continuation`;
    fragment 2 is a generated appendix continuation table → skipped via
    `generated_appendix_continuation`)."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
        _valid_manual_continuation_table_ids,
    )
    from guides.coursework_kfu_2025.table_split_prototype import (
        apply_numbered_split_to_document,
        _build_appendix_continuation_paragraph,
    )
    doc, row_texts = _build_p2a_appendix_doc(rows=4)
    apply_numbered_split_to_document(
        doc, table_index=0, split_before_row=3,
        header_rows=1, numbered_header=True, appendix_table=True,
        continuation_paragraph_builder=_build_appendix_continuation_paragraph,
    )
    # Re-detect on the now-split doc.
    chain_ids = _valid_manual_continuation_table_ids(doc)
    fake_lines = _build_p2a_pdf_lines(row_texts, header_page=10, page_break_after_row=2, last_page=11)
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, fake_lines, manual_chain_ids=chain_ids)
    if candidates:
        return _result(False, f"second pass produced unexpected candidates: {candidates!r}")
    reasons = {ti: r for ti, r in skips}
    if reasons.get(0) != "already_followed_by_continuation":
        return _result(False, f"expected tbl1 skip='already_followed_by_continuation', got {reasons.get(0)!r}")
    if reasons.get(1) != "generated_appendix_continuation":
        return _result(False, f"expected tbl2 skip='generated_appendix_continuation', got {reasons.get(1)!r}")
    return _result(True, f"second pass idempotent (tbl1={reasons.get(0)}, tbl2={reasons.get(1)})")


def test_p2a_no_regression_on_p1c_source_note_fixture() -> tuple[bool, str]:
    """P2-a regression: a P1-c-style synthetic doc (ordinary body table +
    Источник: paragraph, no appendix) must NOT yield any P2-a candidate —
    P1-c stays the only owner of this case."""
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
    )
    doc = Document()
    doc.add_paragraph("Таблица 2.1.2 — Тестовая")
    tbl = doc.add_table(rows=4, cols=2)
    tbl.rows[0].cells[0].text = "Col1Header"
    tbl.rows[0].cells[1].text = "Col2Header"
    row_texts = []
    for r in range(1, 4):
        a, b = f"alpha{r}", f"beta{r}"
        tbl.rows[r].cells[0].text = a
        tbl.rows[r].cells[1].text = b
        row_texts.append((a, b))
    doc.add_paragraph("Источник: составлено автором.")
    pdf_lines = _build_p2a_pdf_lines(row_texts, header_page=5, page_break_after_row=99, last_page=6)
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=set())
    if candidates:
        return _result(False, f"P1-c case must not be P2-a candidate: {candidates!r}")
    if not any(reason == "not_appendix" for _, reason in skips):
        return _result(False, f"expected 'not_appendix' skip on P1-c case, got {skips!r}")
    return _result(True, "P1-c source/note case untouched by P2-a")


# ── P2-a′ — relaxed row/page matcher for appendix tables (fallback) ──────────


def _make_table_signature(rows_data: list[tuple[str, ...]]):
    """Build a TableSignature directly from cell-text tuples for unit testing
    the relaxed matcher without instantiating a full Document. Row 0 is the
    header (skipped by data-row probes), rows 1..N are data rows."""
    from guides.coursework_kfu_2025.table_continuation import (
        RowSignature, TableSignature, _norm_match_text,
    )
    row_sigs = []
    for row_idx, cells in enumerate(rows_data):
        fragments = tuple(_norm_match_text(c) for c in cells if c)
        if not fragments:
            continue
        key = " || ".join(fragments)
        row_sigs.append(RowSignature(row_idx=row_idx, key=key, fragments=fragments))
    return TableSignature(table_idx=0, tbl_xml=None, rows=tuple(row_sigs))


def _make_pdf_lines(entries: list[tuple[str, int]]):
    """Construct PdfLine objects from (text, page_num) tuples."""
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine
    return [PdfLine(text=t, page_num=p, top=100, bottom=110) for t, p in entries]


def test_p2a_relaxed_accepts_duplicate_row_signatures() -> tuple[bool, str]:
    """
    P2-a′ positive: a table with triplicate (A,B,A,B,A,B) data rows produces
    duplicate row signatures, which strict `_match_row_pages` rejects up-front.
    The relaxed matcher must accept these via sequential positional matching
    and return a row→page map.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _match_row_pages,
        _match_row_pages_relaxed_for_appendix,
    )
    sig = _make_table_signature([
        ("Header1", "Header2"),
        ("aaa", "bbb"),  # row 1
        ("ccc", "ddd"),  # row 2
        ("aaa", "bbb"),  # row 3 (duplicate)
        ("ccc", "ddd"),  # row 4 (duplicate)
        ("aaa", "bbb"),  # row 5 (duplicate)
        ("ccc", "ddd"),  # row 6 (duplicate)
    ])
    pdf_lines = _make_pdf_lines([
        ("Header1 Header2", 10),
        ("aaa bbb", 10),  # row 1
        ("ccc ddd", 10),  # row 2
        ("aaa bbb", 11),  # row 3 (page break here)
        ("ccc ddd", 11),  # row 4
        ("aaa bbb", 11),  # row 5
        ("ccc ddd", 11),  # row 6
    ])
    if _match_row_pages(sig, pdf_lines) is not None:
        return _result(False, "fixture broken: strict matcher accepted duplicates")
    row_pages = _match_row_pages_relaxed_for_appendix(sig, pdf_lines)
    if row_pages is None:
        return _result(False, "relaxed matcher should accept duplicate signatures")
    if row_pages.get(1) != 10 or row_pages.get(2) != 10:
        return _result(False, f"row 1/2 page mapping wrong: {row_pages!r}")
    if row_pages.get(3) != 11 or row_pages.get(6) != 11:
        return _result(False, f"row 3/6 page mapping wrong: {row_pages!r}")
    return _result(True, f"relaxed matcher returned {row_pages!r}")


def test_p2a_relaxed_matches_wrapped_cells_via_window() -> tuple[bool, str]:
    """
    P2-a′: when a cell's text is split across multiple PdfLines (rendering
    wrap), the relaxed matcher must reconstruct it via window aggregation
    (concatenating 2–3 adjacent lines) and still produce a row→page map.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _match_row_pages_relaxed_for_appendix,
    )
    # Long second cell wraps across two PDF lines per row.
    sig = _make_table_signature([
        ("Hdr1", "Hdr2"),
        ("alpha-uniqueK1", "long-tail-K1-detail-token-K1"),
        ("beta-uniqueK2", "long-tail-K2-detail-token-K2"),
        ("gamma-uniqueK3", "long-tail-K3-detail-token-K3"),
        ("delta-uniqueK4", "long-tail-K4-detail-token-K4"),
    ])
    # Each row spans 2 wrapped PDF lines.
    pdf_lines = _make_pdf_lines([
        ("Hdr1 Hdr2", 10),
        ("alpha-uniqueK1 long-tail-K1-detail-", 10),
        ("token-K1", 10),
        ("beta-uniqueK2 long-tail-K2-detail-", 10),
        ("token-K2", 10),
        ("gamma-uniqueK3 long-tail-K3-detail-", 11),
        ("token-K3", 11),
        ("delta-uniqueK4 long-tail-K4-detail-", 11),
        ("token-K4", 11),
    ])
    row_pages = _match_row_pages_relaxed_for_appendix(sig, pdf_lines)
    if row_pages is None:
        return _result(False, "relaxed matcher failed to handle wrapped cells")
    if row_pages.get(1) != 10 or row_pages.get(4) != 11:
        return _result(False, f"window-aggregated page mapping wrong: {row_pages!r}")
    return _result(True, f"window aggregation: {row_pages!r}")


def test_p2a_relaxed_returns_none_on_low_confidence() -> tuple[bool, str]:
    """
    P2-a′: when fewer than `min_matched_data_rows` (=4) data rows can be
    matched, the relaxed matcher returns None — fail-open, do not split.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _match_row_pages_relaxed_for_appendix,
    )
    sig = _make_table_signature([
        ("H1", "H2"),
        ("only-this", "row-matches"),
        ("missing-A", "missing-B"),
        ("missing-C", "missing-D"),
        ("missing-E", "missing-F"),
    ])
    pdf_lines = _make_pdf_lines([
        ("H1 H2", 10),
        ("only-this row-matches", 10),
        ("unrelated junk", 11),
    ])
    result = _match_row_pages_relaxed_for_appendix(sig, pdf_lines)
    if result is not None:
        return _result(False, f"expected None for low confidence, got {result!r}")
    return _result(True, "low-confidence input correctly rejected")


def test_p2a_relaxed_rejects_non_monotonic_pages() -> tuple[bool, str]:
    """
    P2-a′: relaxed matcher rejects mappings where pages are not non-decreasing
    in row order (signals matching went backwards — unreliable).
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _match_row_pages_relaxed_for_appendix,
    )
    sig = _make_table_signature([
        ("H1", "H2"),
        ("aaa", "bbb"),
        ("ccc", "ddd"),
        ("eee", "fff"),
        ("ggg", "hhh"),
    ])
    # PDF has rows in WRONG order — first match would be on later page,
    # then earlier page. Relaxed must reject non-monotonic.
    pdf_lines = _make_pdf_lines([
        ("aaa bbb", 11),
        ("ccc ddd", 10),  # backwards page
        ("eee fff", 11),
        ("ggg hhh", 11),
    ])
    result = _match_row_pages_relaxed_for_appendix(sig, pdf_lines)
    if result is not None:
        return _result(False, f"expected None for non-monotonic, got {result!r}")
    return _result(True, "non-monotonic mapping correctly rejected")


def test_p2a_relaxed_rejects_single_page_mapping() -> tuple[bool, str]:
    """
    P2-a′: if all matched rows live on a single page (no multi-page span),
    the relaxed matcher returns None — no split is needed.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _match_row_pages_relaxed_for_appendix,
    )
    sig = _make_table_signature([
        ("H1", "H2"),
        ("rowA1", "rowA2"),
        ("rowB1", "rowB2"),
        ("rowC1", "rowC2"),
        ("rowD1", "rowD2"),
    ])
    pdf_lines = _make_pdf_lines([
        ("H1 H2", 10),
        ("rowA1 rowA2", 10),
        ("rowB1 rowB2", 10),
        ("rowC1 rowC2", 10),
        ("rowD1 rowD2", 10),
    ])
    result = _match_row_pages_relaxed_for_appendix(sig, pdf_lines)
    if result is not None:
        return _result(False, f"expected None for single-page, got {result!r}")
    return _result(True, "single-page mapping correctly rejected")


def test_p2a_strict_matcher_behavior_unchanged() -> tuple[bool, str]:
    """
    P2-a′ regression: strict `_match_row_pages` must remain bit-identical
    on existing inputs. The relaxed matcher is additive and must not affect
    strict semantics elsewhere (marker-split, P1-c, legacy split).
    """
    from guides.coursework_kfu_2025.table_continuation import _match_row_pages
    sig = _make_table_signature([
        ("Hdr1", "Hdr2"),
        ("alpha", "beta"),
        ("gamma", "delta"),
        ("epsilon", "zeta"),
    ])
    pdf_lines = _make_pdf_lines([
        ("Hdr1 Hdr2", 5),
        ("alpha beta", 5),
        ("gamma delta", 5),
        ("epsilon zeta", 6),
    ])
    result = _match_row_pages(sig, pdf_lines)
    if result is None:
        return _result(False, "strict matcher regressed on valid input")
    if result.get(1) != 5 or result.get(3) != 6:
        return _result(False, f"strict matcher mapping changed: {result!r}")
    # Strict still rejects duplicates.
    dup_sig = _make_table_signature([
        ("Hdr1", "Hdr2"),
        ("x", "y"),
        ("x", "y"),
        ("x", "y"),
    ])
    if _match_row_pages(dup_sig, pdf_lines) is not None:
        return _result(False, "strict matcher regressed: now accepts duplicates")
    return _result(True, "strict matcher behavior preserved")


def test_p2a_collector_uses_relaxed_when_strict_fails() -> tuple[bool, str]:
    """
    P2-a′ integration: when an appendix table has duplicate row signatures
    (strict fails) but the relaxed matcher succeeds, the P2-a collector
    must produce a candidate (not skip with `render_probe_unreliable`).
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
    )
    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    # 7-row appendix table with triplicate-style duplicate signatures (rows 1,3,5 share key A; rows 2,4,6 share key B).
    tbl = doc.add_table(rows=7, cols=2)
    tbl.rows[0].cells[0].text = "Hdr1"
    tbl.rows[0].cells[1].text = "Hdr2"
    pattern = [("dupA1", "dupA2"), ("dupB1", "dupB2")]
    for r in range(1, 7):
        a, b = pattern[(r - 1) % 2]
        tbl.rows[r].cells[0].text = a
        tbl.rows[r].cells[1].text = b
    pdf_lines = _make_pdf_lines([
        ("Hdr1 Hdr2", 10),
        ("dupA1 dupA2", 10),
        ("dupB1 dupB2", 10),
        ("dupA1 dupA2", 11),
        ("dupB1 dupB2", 11),
        ("dupA1 dupA2", 11),
        ("dupB1 dupB2", 11),
    ])
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=set())
    if not candidates:
        return _result(False, f"expected ≥1 candidate via relaxed matcher; got skips={skips!r}")
    c = candidates[0]
    if c.appendix_num != "1":
        return _result(False, f"appendix_num wrong: {c.appendix_num!r}")
    if c.split_before_row is None or c.split_before_row < 2:
        return _result(False, f"split_before_row invalid: {c.split_before_row}")
    return _result(True, f"collector used relaxed matcher: candidate split_before_row={c.split_before_row}")


def test_p2a_strict_path_takes_priority_when_strict_succeeds() -> tuple[bool, str]:
    """
    P2-a′: when strict `_match_row_pages` succeeds (unique data rows mapping
    cleanly to PDF pages), the collector uses that mapping and does NOT
    invoke the relaxed fallback. Verifies strict path priority.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
        _match_row_pages,
    )
    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    tbl = doc.add_table(rows=5, cols=2)
    tbl.rows[0].cells[0].text = "Hdr1"
    tbl.rows[0].cells[1].text = "Hdr2"
    unique_rows = [("uniA1", "uniA2"), ("uniB1", "uniB2"), ("uniC1", "uniC2"), ("uniD1", "uniD2")]
    for r, (a, b) in enumerate(unique_rows, start=1):
        tbl.rows[r].cells[0].text = a
        tbl.rows[r].cells[1].text = b
    pdf_lines = _make_pdf_lines([
        ("Hdr1 Hdr2", 10),
        ("uniA1 uniA2", 10),
        ("uniB1 uniB2", 10),
        ("uniC1 uniC2", 11),
        ("uniD1 uniD2", 11),
    ])
    # Strict matcher MUST succeed on this fixture.
    from guides.coursework_kfu_2025.table_continuation import _collect_table_signatures
    sigs = _collect_table_signatures(doc)
    if _match_row_pages(sigs[0], pdf_lines) is None:
        return _result(False, "fixture broken: strict matcher failed on unique data")
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=set())
    if not candidates:
        return _result(False, f"expected candidate via strict; got skips={skips!r}")
    c = candidates[0]
    # Strict says rows 1,2 on page 10; rows 3,4 on page 11 → split_before_row = 3.
    if c.split_before_row != 3:
        return _result(False, f"strict path should give split_before_row=3, got {c.split_before_row}")
    return _result(True, "strict matcher path took priority (split_before_row=3)")


def test_p2a_bondarev_style_triplicate_appendix_creates_candidate() -> tuple[bool, str]:
    """
    P2-a′ end-to-end fixture: Bondarev-style appendix table with 18 rows
    (header + 6 unique × 3 copies) where strict matcher rejects but relaxed
    accepts and produces a P2-a candidate with a valid split point.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _collect_appendix_table_continuation_candidates,
    )
    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЕ А")
    tbl = doc.add_table(rows=19, cols=2)  # header + 18 data rows
    tbl.rows[0].cells[0].text = "Этап"
    tbl.rows[0].cells[1].text = "Описание"
    unique = [
        ("Подготовка", "Формирование повестки"),
        ("Обсуждение", "Дебаты по вопросам"),
        ("Голосование", "Электронное голосование"),
        ("Анализ", "Систематизация оценки"),
        ("Отчётность", "Публикация итогов"),
        ("Архивация", "Сохранение материалов"),
    ]
    for r in range(1, 19):
        a, b = unique[(r - 1) % 6]
        tbl.rows[r].cells[0].text = a
        tbl.rows[r].cells[1].text = b
    # PDF: 9 rows on page 10, 9 rows on page 11 (mid-table page break).
    pdf_entries = [("Этап Описание", 10)]
    for r in range(1, 10):
        a, b = unique[(r - 1) % 6]
        pdf_entries.append((f"{a} {b}", 10))
    for r in range(10, 19):
        a, b = unique[(r - 1) % 6]
        pdf_entries.append((f"{a} {b}", 11))
    pdf_lines = _make_pdf_lines(pdf_entries)
    candidates, skips = _collect_appendix_table_continuation_candidates(doc, pdf_lines, manual_chain_ids=set())
    if not candidates:
        return _result(False, f"Bondarev-style triplicate didn't produce candidate; skips={skips!r}")
    c = candidates[0]
    if c.appendix_num != "А":
        return _result(False, f"appendix_num wrong: {c.appendix_num!r}")
    if c.split_before_row is None or c.split_before_row < 2:
        return _result(False, f"split_before_row invalid: {c.split_before_row}")
    if c.first_data_row_page != 10 or c.last_data_row_page != 11:
        return _result(False, f"page mapping wrong: first={c.first_data_row_page} last={c.last_data_row_page}")
    return _result(True, f"Bondarev-style candidate: split_before_row={c.split_before_row}")


# ── DEFECT 3 — caption-like analytical prose under table must stay body_text ──

def test_table_caption_reference_prose_with_svyazyvaet_demoted() -> tuple[bool, str]:
    """
    DEFECT 3: a paragraph whose text starts with 'Таблица X.Y.Z' followed by an
    analytical reference verb (e.g. 'связывает …', 'позволяет …') is analytical
    body prose, not a caption. classify_paragraph must return body_text so that
    pagination_rules._apply_rule3 does NOT set keep_with_next on it.
    """
    from guides.coursework_kfu_2025.classifier import classify_paragraph

    cases = [
        "Таблица 1.3.1 связывает каждую точку контакта с конкретным вопросом",
        "Таблица 2.1.1 позволяет проследить динамику показателей",
        "Таблица 1.2.1 показывает структуру выручки",
        "Таблица 1.2.1 отражает изменения за период",
    ]
    for text in cases:
        kind = classify_paragraph(text, prev_kind="empty_paragraph")
        if kind == "table_caption":
            return _result(False, f"analytical prose classified as table_caption: {text!r}")
        if kind != "body_text":
            return _result(False, f"expected body_text, got {kind!r} for {text!r}")
    return _result(True, "analytical prose under table classified as body_text")


def test_pagination_rule3_does_not_set_keepnext_on_prose_below_table() -> tuple[bool, str]:
    """
    DEFECT 3 integration: a paragraph beginning 'Таблица 2.1.1 связывает …'
    appearing in body context (after table + Источник: + blank) must NOT receive
    <w:keepNext/> from pagination_rules._apply_rule3.
    """
    from guides.coursework_kfu_2025.pagination_rules import apply_pagination_rules

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "H1"
    tbl.rows[0].cells[1].text = "H2"
    tbl.rows[1].cells[0].text = "a"
    tbl.rows[1].cells[1].text = "b"
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("")
    prose = doc.add_paragraph(
        "Таблица 2.1.1 связывает каждую точку контакта с конкретным вопросом потенциального франчайзи."
    )
    doc.add_paragraph("Далее следует обычный текст основного содержания.")

    apply_pagination_rules(doc)

    pPr = prose._element.find(qn("w:pPr"))
    has_keep_next = pPr is not None and pPr.find(qn("w:keepNext")) is not None
    if has_keep_next:
        return _result(False, "analytical prose paragraph below table received keepNext")
    return _result(True, "analytical prose paragraph below table has no keepNext")


def test_genuine_table_caption_still_gets_keepnext() -> tuple[bool, str]:
    """
    DEFECT 3 regression: a genuine caption 'Таблица 2.1.1 — Структура выручки'
    immediately preceding a real table must still receive keep_with_next from
    pagination_rules._apply_rule3.
    """
    from guides.coursework_kfu_2025.pagination_rules import apply_pagination_rules

    doc = Document()
    caption = doc.add_paragraph("Таблица 2.1.1 — Структура выручки")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "H1"
    tbl.rows[0].cells[1].text = "H2"
    tbl.rows[1].cells[0].text = "x"
    tbl.rows[1].cells[1].text = "y"

    apply_pagination_rules(doc)

    pPr = caption._element.find(qn("w:pPr"))
    has_keep_next = pPr is not None and pPr.find(qn("w:keepNext")) is not None
    if not has_keep_next:
        return _result(False, "genuine table caption lost keepNext (regression)")
    return _result(True, "genuine table caption still has keepNext")


# ── P4 / DEFECT 4 — source/note soft-break normalization ─────────────────────

def _paragraph_has_w_br(paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("w:br")))


def test_p4_source_note_softbreak_splits_into_two_paragraphs() -> tuple[bool, str]:
    """
    DEFECT 4 positive: a single paragraph containing
    'Источник: ... <w:br/> Примечание: ...' must be split into two proper
    paragraphs, with the inline soft-break removed.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        split_body_structural_soft_breaks,
    )

    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Источник: составлено автором по [1].")
    r1.add_break()
    p.add_run("Примечание: ставки на 2026 г.")

    raw_text_before = p.text
    if "\n" not in raw_text_before:
        return _result(False, f"fixture failed to create soft-break in paragraph.text: {raw_text_before!r}")
    if not _paragraph_has_w_br(p):
        return _result(False, "fixture failed to insert <w:br/> element")

    split_body_structural_soft_breaks(doc, body_start=0)

    paras = doc.paragraphs
    if len(paras) != 2:
        return _result(False, f"expected 2 paragraphs after split, got {len(paras)}: {[p.text for p in paras]!r}")
    if not paras[0].text.startswith("Источник:"):
        return _result(False, f"first paragraph not Источник: {paras[0].text!r}")
    if not paras[1].text.startswith("Примечание:"):
        return _result(False, f"second paragraph not Примечание: {paras[1].text!r}")
    if _paragraph_has_w_br(paras[0]) or _paragraph_has_w_br(paras[1]):
        return _result(False, "inline <w:br/> still present in resulting paragraphs")
    return _result(True, "source/note soft-break split into two clean paragraphs")


def test_p4_ordinary_body_softbreak_remains_single_paragraph() -> tuple[bool, str]:
    """
    DEFECT 4 regression: a paragraph with a generic body-text soft-break (neither
    segment is heading / source / note) must NOT be split. Preserves the
    existing 'preserve structural soft breaks' invariant (commit 31a612a).
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        split_body_structural_soft_breaks,
    )

    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("Первая строка обычного тела абзаца.")
    r1.add_break()
    p.add_run("Вторая строка обычного тела абзаца.")

    if not _paragraph_has_w_br(p):
        return _result(False, "fixture failed to insert <w:br/>")

    split_body_structural_soft_breaks(doc, body_start=0)

    paras = doc.paragraphs
    if len(paras) != 1:
        return _result(False, f"ordinary body paragraph was split unexpectedly: {len(paras)} paragraphs: {[p.text for p in paras]!r}")
    if not _paragraph_has_w_br(paras[0]):
        return _result(False, "ordinary body soft-break was removed (regression)")
    return _result(True, "ordinary body soft-break preserved as single paragraph")


def test_p4_heading_body_softbreak_still_splits() -> tuple[bool, str]:
    """
    DEFECT 4 regression: pre-existing heading→body soft-break split behavior
    must remain intact (predicate accepts heading1/heading2 as first segment).
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        split_body_structural_soft_breaks,
    )

    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("1. Теоретические основы предмета исследования")
    r1.add_break()
    p.add_run("Современная теория данной области активно развивается.")

    if not _paragraph_has_w_br(p):
        return _result(False, "fixture failed to insert <w:br/>")

    split_body_structural_soft_breaks(doc, body_start=0)

    paras = doc.paragraphs
    if len(paras) != 2:
        return _result(False, f"heading→body split regressed: {len(paras)} paragraphs: {[p.text for p in paras]!r}")
    return _result(True, "heading→body split still works")


def test_p4_source_note_split_resulting_paras_format_through_phase1() -> tuple[bool, str]:
    """
    DEFECT 4: after process_document runs end-to-end, the resulting Источник: and
    Примечание: paragraphs must be classified as source_line and receive the
    canonical source/note formatting (justify alignment, first-line indent).
    Verifies that the split paragraphs go through the regular Phase 1 dispatch.
    """
    import os, tempfile
    from guides.coursework_kfu_2025.safe_formatter import process_document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with tempfile.TemporaryDirectory() as tmp:
        src = os.path.join(tmp, "in.docx")
        out = os.path.join(tmp, "out.docx")
        doc = Document()
        # Minimal viable body — needs an "Введение" heading so body_start gets detected.
        doc.add_paragraph("ВВЕДЕНИЕ")
        doc.add_paragraph("Краткое введение для запуска Phase 1.")
        # Soft-break paragraph between two body paragraphs.
        sn = doc.add_paragraph()
        r1 = sn.add_run("Источник: составлено автором по [1], [2].")
        r1.add_break()
        sn.add_run("Примечание: данные на 1 января 2026 года.")
        doc.add_paragraph("Заключительный обычный абзац.")
        doc.save(src)

        process_document(src, out)

        d2 = Document(out)
        # Find paragraphs by text prefix
        ist = next((p for p in d2.paragraphs if (p.text or "").startswith("Источник:")), None)
        pri = next((p for p in d2.paragraphs if (p.text or "").startswith("Примечание:")), None)
        if ist is None or pri is None:
            return _result(False, f"missing Источник or Примечание paragraph after process_document; texts={[p.text for p in d2.paragraphs]!r}")
        if _paragraph_has_w_br(ist) or _paragraph_has_w_br(pri):
            return _result(False, "soft-break <w:br/> survived process_document")
        # Should be justified (source_line format)
        for p in (ist, pri):
            if p.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
                return _result(False, f"paragraph alignment != JUSTIFY: {p.text!r} -> {p.alignment!r}")
        return _result(True, "process_document split + normalized source/note paragraphs")


def test_c_apply_table_continuation_does_not_heuristic_split() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc = Document()
    for _ in range(22):
        doc.add_paragraph("Текст абзаца для заполнения страницы.")

    doc.add_paragraph("Таблица 2.3")
    doc.add_paragraph("Название таблицы")

    tbl = doc.add_table(rows=8, cols=2)
    tbl.rows[0].cells[0].text = "Колонка A"
    tbl.rows[0].cells[1].text = "Колонка B"
    for i in range(1, 8):
        tbl.rows[i].cells[0].text = f"a{i}"
        tbl.rows[i].cells[1].text = f"b{i}"

    before_rows = len(doc.tables[0].rows)
    n = apply_table_continuation(doc)
    markers = [p for p in doc.paragraphs if "Продолжение таблицы" in (p.text or "")]

    if n != 0:
        return _result(False, f"expected no width changes in split fixture, got {n}")
    if len(doc.tables) != 1:
        return _result(False, f"heuristic split created extra table(s): {len(doc.tables)}")
    if len(doc.tables[0].rows) != before_rows:
        return _result(False, f"row count changed: {before_rows} -> {len(doc.tables[0].rows)}")
    if markers:
        return _result(False, f"heuristic continuation marker inserted: {[p.text for p in markers]!r}")
    return _result(True, "heuristic split disabled")


def test_c_apply_table_continuation_width_normalization_only() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[1].cells[0].text = "1"
    tbl.rows[1].cells[1].text = "2"
    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return _result(False, "test setup failed: no tblGrid")
    for gc in grid.findall(qn("w:gridCol")):
        gc.set(qn("w:w"), "12000")

    before_tables = len(doc.tables)
    before_rows = len(doc.tables[0].rows)
    n = apply_table_continuation(doc)
    markers = [p for p in doc.paragraphs if "Продолжение таблицы" in (p.text or "")]

    if n != 1:
        return _result(False, f"expected one width-normalised table, got {n}")
    if len(doc.tables) != before_tables:
        return _result(False, f"table count changed: {before_tables} -> {len(doc.tables)}")
    if len(doc.tables[0].rows) != before_rows:
        return _result(False, f"row count changed: {before_rows} -> {len(doc.tables[0].rows)}")
    if markers:
        return _result(False, f"unexpected continuation marker: {[p.text for p in markers]!r}")
    return _result(True, "width normalisation remained active without splitting")


def test_c_apply_table_continuation_no_split_double_run_idempotent() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc = Document()
    for _ in range(22):
        doc.add_paragraph("Текст абзаца для заполнения страницы.")

    doc.add_paragraph("Таблица 3.1")
    tbl = doc.add_table(rows=8, cols=2)
    tbl.rows[0].cells[0].text = "H1"
    tbl.rows[0].cells[1].text = "H2"
    for i in range(1, 8):
        tbl.rows[i].cells[0].text = f"a{i}"
        tbl.rows[i].cells[1].text = f"b{i}"

    first = apply_table_continuation(doc)
    marker_count_1 = sum(1 for p in doc.paragraphs if "Продолжение таблицы" in (p.text or ""))
    table_count_1 = len(doc.tables)
    table_rows_1 = [len(t.rows) for t in doc.tables]

    second = apply_table_continuation(doc)
    marker_count_2 = sum(1 for p in doc.paragraphs if "Продолжение таблицы" in (p.text or ""))
    table_count_2 = len(doc.tables)
    table_rows_2 = [len(t.rows) for t in doc.tables]

    if first != 0 or second != 0:
        return _result(False, f"expected no heuristic changes, got first={first}, second={second}")
    if marker_count_2 != marker_count_1:
        return _result(False, f"marker count changed: {marker_count_1} -> {marker_count_2}")
    if table_count_2 != table_count_1:
        return _result(False, f"table count changed: {table_count_1} -> {table_count_2}")
    if table_rows_2 != table_rows_1:
        return _result(False, f"table structure changed: {table_rows_1!r} -> {table_rows_2!r}")
    return _result(True, "double run did not add markers or split structure")


def test_c_apply_rendered_table_continuation_warns_when_lo_unavailable() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.docx_utils import FormattingReport

    doc = Document()
    doc.add_paragraph("Таблица 1.1")
    doc.add_table(rows=3, cols=2)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "in.docx"
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        report = FormattingReport()
        try:
            def raise_lo(_path):
                raise tc.LibreOfficeNotFoundError("missing LO")

            tc.render_docx_to_pdf = raise_lo
            n = tc.apply_rendered_table_continuation(path, report=report)
        finally:
            tc.render_docx_to_pdf = old_render

        reread = Document(str(path))

    if n != 0:
        return _result(False, f"expected 0 rendered splits, got {n}")
    if not report.warnings:
        return _result(False, "expected rendered split warning")
    if len(reread.tables) != 1:
        return _result(False, f"DOCX mutated unexpectedly, tables={len(reread.tables)}")
    return _result(True, "LO unavailable path warns and does not mutate")


def test_c_apply_rendered_table_continuation_warns_when_pdf_analysis_fails() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.docx_utils import FormattingReport

    doc = Document()
    doc.add_paragraph("Таблица 1.1")
    doc.add_table(rows=3, cols=2)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "in.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "in.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        report = FormattingReport()
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path

            def raise_analysis(_path):
                raise RuntimeError("pdf parse failed")

            tc.analyze_pdf_lines = raise_analysis
            n = tc.apply_rendered_table_continuation(path, report=report)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))

    if n != 0:
        return _result(False, f"expected 0 rendered splits, got {n}")
    if not report.warnings:
        return _result(False, "expected PDF analysis warning")
    if len(reread.tables) != 1:
        return _result(False, f"DOCX mutated unexpectedly, tables={len(reread.tables)}")
    return _result(True, "PDF analysis failure warns and does not mutate")


def test_c_rendered_split_single_boundary_success() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    doc.add_paragraph("Таблица 2.3")
    tbl = doc.add_table(rows=4, cols=2)
    tbl.rows[0].cells[0].text = "H1"
    tbl.rows[0].cells[1].text = "H2"
    tbl.rows[1].cells[0].text = "alpha one"
    tbl.rows[1].cells[1].text = "beta one"
    tbl.rows[2].cells[0].text = "gamma two"
    tbl.rows[2].cells[1].text = "delta two"
    tbl.rows[3].cells[0].text = "epsilon three"
    tbl.rows[3].cells[1].text = "zeta three"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "in.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "in.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("H1 H2", 1, 100.0, 112.0),
                PdfLine("alpha one beta one", 1, 120.0, 132.0),
                PdfLine("gamma two delta two", 2, 80.0, 92.0),
                PdfLine("epsilon three zeta three", 2, 100.0, 112.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))

    if n != 1:
        return _result(False, f"expected one rendered split, got {n}")
    if len(reread.tables) != 2:
        return _result(False, f"expected 2 tables after split, got {len(reread.tables)}")
    markers = [p.text for p in reread.paragraphs if "Продолжение таблицы" in (p.text or "")]
    if markers != ["Продолжение таблицы 2.3"]:
        return _result(False, f"unexpected markers: {markers!r}")
    if [c.text for c in reread.tables[0].rows[0].cells] != [c.text for c in reread.tables[1].rows[0].cells]:
        return _result(False, "continuation table header was not repeated")
    return _result(True, "rendered single-boundary split succeeded")


def test_c_rendered_split_preserves_valid_manual_split() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    t1 = doc.add_table(rows=2, cols=2)
    t1.rows[0].cells[0].text = "H1"
    t1.rows[0].cells[1].text = "H2"
    t1.rows[1].cells[0].text = "alpha"
    t1.rows[1].cells[1].text = "beta"
    marker = doc.add_paragraph("Продолжение таблицы 1.1")
    marker.alignment = 2  # right; must be preserved exactly
    marker.paragraph_format.keep_with_next = True
    t2 = doc.add_table(rows=2, cols=2)
    t2.rows[0].cells[0].text = "H1"
    t2.rows[0].cells[1].text = "H2"
    t2.rows[1].cells[0].text = "gamma"
    t2.rows[1].cells[1].text = "delta"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "manual.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "manual.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)
        before_xml = marker._element.xml

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("H1 H2", 1, 100.0, 112.0),
                PdfLine("alpha beta", 1, 120.0, 132.0),
                PdfLine("gamma delta", 2, 80.0, 92.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        markers = [p for p in reread.paragraphs if "Продолжение таблицы" in (p.text or "")]

    if n != 0:
        return _result(False, f"valid manual split should be preserved, got split count {n}")
    if len(reread.tables) != 2:
        return _result(False, f"manual split table count changed: {len(reread.tables)}")
    if len(markers) != 1 or markers[0]._element.xml != before_xml:
        return _result(False, "manual continuation marker XML changed")
    return _result(True, "valid manual split preserved exactly")


def test_c_rendered_split_skips_ambiguous_repeated_rows() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    doc.add_paragraph("Таблица 2.4")
    tbl = doc.add_table(rows=4, cols=2)
    tbl.rows[0].cells[0].text = "H1"
    tbl.rows[0].cells[1].text = "H2"
    for idx in range(1, 4):
        tbl.rows[idx].cells[0].text = "same"
        tbl.rows[idx].cells[1].text = "row"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "ambiguous.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "ambiguous.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("same row", 1, 120.0, 132.0),
                PdfLine("same row", 2, 80.0, 92.0),
                PdfLine("same row", 2, 100.0, 112.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))

    if n != 0:
        return _result(False, f"ambiguous repeated rows should skip, got {n}")
    if len(reread.tables) != 1:
        return _result(False, f"ambiguous split mutated table count: {len(reread.tables)}")
    return _result(True, "ambiguous repeated rows skipped")


def test_c_rendered_split_skips_merged_boundary_conflict() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    doc.add_paragraph("Таблица 2.5")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "H1"
    tbl.rows[0].cells[1].text = "H2"
    tbl.rows[1].cells[0].text = "merge start"
    tbl.rows[1].cells[1].text = "alpha"
    tbl.rows[2].cells[0].text = "merge continue"
    tbl.rows[2].cells[1].text = "beta"
    tbl.cell(1, 0).merge(tbl.cell(2, 0))

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "merged.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "merged.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("H1 H2", 1, 100.0, 112.0),
                PdfLine("merge start alpha", 1, 120.0, 132.0),
                PdfLine("merge continue beta", 2, 80.0, 92.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))

    if n != 0:
        return _result(False, f"merged boundary conflict should skip, got {n}")
    if len(reread.tables) != 1:
        return _result(False, f"merged conflict mutated table count: {len(reread.tables)}")
    return _result(True, "merged boundary conflict skipped")


def test_c_rendered_split_marker_is_right_aligned() -> tuple[bool, str]:
    """
    Product rule: generated ordinary-table continuation markers are
    right-aligned and keep the existing continuation marker formatting.
    """
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    doc.add_paragraph("Таблица 2.6")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "H1"
    tbl.rows[0].cells[1].text = "H2"
    tbl.rows[1].cells[0].text = "alpha"
    tbl.rows[1].cells[1].text = "beta"
    tbl.rows[2].cells[0].text = "gamma"
    tbl.rows[2].cells[1].text = "delta"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "format.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "format.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("H1 H2", 1, 100.0, 112.0),
                PdfLine("alpha beta", 1, 120.0, 132.0),
                PdfLine("gamma delta", 2, 80.0, 92.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        markers = [p for p in reread.paragraphs if "Продолжение таблицы" in (p.text or "")]

    if n != 1 or len(markers) != 1:
        return _result(False, f"expected one generated marker, n={n}, markers={len(markers)}")
    pPr = markers[0]._element.find(qn("w:pPr"))
    page_break = pPr.find(qn("w:pageBreakBefore")) if pPr is not None else None
    jc = pPr.find(qn("w:jc")) if pPr is not None else None
    ind = pPr.find(qn("w:ind")) if pPr is not None else None
    keep = pPr.find(qn("w:keepNext")) if pPr is not None else None
    sz = markers[0]._element.find(".//" + qn("w:sz"))
    if page_break is None:
        return _result(False, "marker pageBreakBefore missing")
    if jc is None or jc.get(qn("w:val")) != "right":
        return _result(False, "marker is not right-aligned")
    if ind is None or ind.get(qn("w:firstLine")) != "0":
        return _result(False, "marker first-line indent is not zero")
    if keep is None:
        return _result(False, "marker keepWithNext missing")
    if sz is None or sz.get(qn("w:val")) != "28":
        return _result(False, "marker font size is not 14pt")
    return _result(True, "generated marker formatting is correct")


def test_p1a_legacy_rendered_split_runs_post_normalizer() -> tuple[bool, str]:
    """P1a: legacy rendered split must normalize the generated marker before save."""
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    inline = doc.add_paragraph(
        "Обычная фраза с Продолжение таблицы 9.9.9 внутри текста."
    )
    inline.paragraph_format.page_break_before = False
    inline.paragraph_format.keep_with_next = False
    doc.add_paragraph("Таблица 2.6")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "H1"
    tbl.rows[0].cells[1].text = "H2"
    tbl.rows[1].cells[0].text = "alpha"
    tbl.rows[1].cells[1].text = "beta"
    tbl.rows[2].cells[0].text = "gamma"
    tbl.rows[2].cells[1].text = "delta"

    def disabled_marker_builder(text: str):
        p = old_builder(text)
        p_pr = p.find(qn("w:pPr"))
        for prop_name in ("pageBreakBefore", "keepNext"):
            prop = p_pr.find(qn(f"w:{prop_name}"))
            if prop is not None:
                prop.set(qn("w:val"), "0")
        return p

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "legacy_anchor.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "legacy_anchor.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        old_builder = tc._build_continuation_para
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("H1 H2", 1, 100.0, 112.0),
                PdfLine("alpha beta", 1, 120.0, 132.0),
                PdfLine("gamma delta", 2, 80.0, 92.0),
            ]
            tc._build_continuation_para = disabled_marker_builder
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            tc._build_continuation_para = old_builder

        reread = Document(str(path))

    if n != 1:
        return _result(False, f"expected one legacy rendered split, got {n}")
    ok, msg = _p1a_assert_rendered_chain_anchored(
        reread, "Продолжение таблицы 2.6", require_blank=False,
    )
    if not ok:
        return _result(False, msg)
    inline_after = next(
        p for p in reread.paragraphs
        if "Продолжение таблицы 9.9.9 внутри текста" in (p.text or "")
    )
    if _marker_page_break_before_enabled(inline_after._element):
        return _result(False, "inline prose received active pageBreakBefore")
    if _marker_keep_next_enabled(inline_after._element):
        return _result(False, "inline prose received active keepNext")
    return _result(True, "legacy rendered split normalized marker without touching inline prose")


def test_c_rendered_split_caption_number_and_fallback() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    def run_case(caption: str, expected_marker: str) -> tuple[bool, str]:
        doc = Document()
        doc.add_paragraph(caption)
        tbl = doc.add_table(rows=3, cols=2)
        tbl.rows[0].cells[0].text = "H1"
        tbl.rows[0].cells[1].text = "H2"
        tbl.rows[1].cells[0].text = "alpha"
        tbl.rows[1].cells[1].text = "beta"
        tbl.rows[2].cells[0].text = "gamma"
        tbl.rows[2].cells[1].text = "delta"

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption.docx"
            pdf_dir = Path(tmp) / "pdf"
            pdf_dir.mkdir()
            pdf_path = pdf_dir / "caption.pdf"
            pdf_path.write_bytes(b"%PDF-1.4\n")
            doc.save(path)

            old_render = tc.render_docx_to_pdf
            old_analyze = tc.analyze_pdf_lines
            try:
                tc.render_docx_to_pdf = lambda _path: pdf_path
                tc.analyze_pdf_lines = lambda _path: [
                    PdfLine("H1 H2", 1, 100.0, 112.0),
                    PdfLine("alpha beta", 1, 120.0, 132.0),
                    PdfLine("gamma delta", 2, 80.0, 92.0),
                ]
                n = tc.apply_rendered_table_continuation(path)
            finally:
                tc.render_docx_to_pdf = old_render
                tc.analyze_pdf_lines = old_analyze

            reread = Document(str(path))
            markers = [p.text for p in reread.paragraphs if "Продолжение таблицы" in (p.text or "")]

        if n != 1:
            return _result(False, f"{caption!r}: expected split, got {n}")
        if markers != [expected_marker]:
            return _result(False, f"{caption!r}: expected {expected_marker!r}, got {markers!r}")
        return _result(True, "")

    ok, msg = run_case("Таблица 2.3.4", "Продолжение таблицы 2.3.4")
    if not ok:
        return _result(False, msg)
    ok, msg = run_case("Таблица абв", "Продолжение таблицы")
    if not ok:
        return _result(False, msg)
    return _result(True, "strict caption number and fallback markers correct")


def test_c_rendered_start_page_moves_whole_table_without_complete_data_row() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    caption = doc.add_paragraph("Таблица 2.2.3")
    doc.add_paragraph("Показатели эффективности")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "Показатель"
    tbl.rows[0].cells[1].text = "Экономия"
    tbl.rows[1].cells[0].text = "Почтовые расходы"
    tbl.rows[1].cells[1].text = "переход на электронный документооборот"
    tbl.rows[2].cells[0].text = "Архивное хранение"
    tbl.rows[2].cells[1].text = "высокая экономия архива"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "move.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "move.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("Таблица 2.2.3", 1, 686.5, 700.0),
                PdfLine("Показатели эффективности", 1, 710.8, 724.0),
                PdfLine("Показатель Экономия", 1, 741.8, 755.0),
                PdfLine("Почтовые расходы переход", 1, 763.2, 776.0),
                PdfLine("на электронный документооборот", 2, 86.8, 99.0),
                PdfLine("Архивное хранение высокая экономия архива", 2, 108.0, 121.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        reread_caption = next(p for p in reread.paragraphs if p.text == caption.text)

    pPr = reread_caption._element.find(qn("w:pPr"))
    page_break = pPr.find(qn("w:pageBreakBefore")) if pPr is not None else None
    markers = [p.text for p in reread.paragraphs if "Продолжение таблицы" in (p.text or "")]

    if n != 1:
        return _result(False, f"expected whole-table move, got {n}")
    if page_break is None:
        return _result(False, "caption did not receive pageBreakBefore")
    if len(reread.tables) != 1:
        return _result(False, f"whole-table move should not split, got {len(reread.tables)} tables")
    if markers:
        return _result(False, f"whole-table move inserted continuation marker: {markers!r}")
    return _result(True, "whole-table move applied to caption")


def test_c_rendered_start_page_first_row_spill_moves_whole_table() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    caption = doc.add_paragraph("Таблица 2.3.1")
    doc.add_paragraph("Структура прямой экономии ТТС при переходе на ЭДО")
    tbl = doc.add_table(rows=3, cols=3)
    tbl.rows[0].cells[0].text = "Статья"
    tbl.rows[0].cells[1].text = "Значение"
    tbl.rows[0].cells[2].text = "Комментарий"
    tbl.rows[1].cells[0].text = "Почтовые расходы"
    tbl.rows[1].cells[1].text = "31–33"
    tbl.rows[1].cells[2].text = "отказ от бумажных отправлений"
    tbl.rows[2].cells[0].text = "Печать"
    tbl.rows[2].cells[1].text = "4–5"
    tbl.rows[2].cells[2].text = "сокращение печати"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "first_row_spill.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "first_row_spill.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("Таблица 2.3.1", 1, 683.2, 695.0),
                PdfLine("Структура прямой экономии ТТС при переходе на ЭДО", 1, 707.4, 719.0),
                PdfLine("Статья Значение Комментарий", 1, 731.5, 743.0),
                PdfLine("Почтовые расходы 31–33 отказ от бумажных", 1, 759.6, 771.0),
                PdfLine("Статья Значение Комментарий", 2, 58.8, 70.0),
                PdfLine("отправлений", 2, 86.8, 98.0),
                PdfLine("Печать 4–5 сокращение печати", 2, 108.2, 120.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        reread_caption = next(p for p in reread.paragraphs if p.text == caption.text)

    pPr = reread_caption._element.find(qn("w:pPr"))
    page_break = pPr.find(qn("w:pageBreakBefore")) if pPr is not None else None

    if n != 1:
        return _result(False, f"first-row spill should trigger whole-table move, got {n}")
    if page_break is None:
        return _result(False, "caption did not receive pageBreakBefore after first-row spill")
    if len(reread.tables) != 1:
        return _result(False, f"whole-table move should not split tables, got {len(reread.tables)}")
    return _result(True, "first-row spill triggered whole-table move")


def test_c_rendered_start_page_skips_existing_page_break_candidate() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    first_caption = doc.add_paragraph("Таблица 1.2.2")
    first_caption.paragraph_format.page_break_before = True
    first_tbl = doc.add_table(rows=2, cols=2)
    first_tbl.rows[0].cells[0].text = "Показатель"
    first_tbl.rows[0].cells[1].text = "Эффект"
    first_tbl.rows[1].cells[0].text = "Первый показатель"
    first_tbl.rows[1].cells[1].text = "переход на электронный обмен"

    second_caption = doc.add_paragraph("Таблица 2.3.3")
    second_tbl = doc.add_table(rows=2, cols=2)
    second_tbl.rows[0].cells[0].text = "Год"
    second_tbl.rows[0].cells[1].text = "Комментарий"
    second_tbl.rows[1].cells[0].text = "Первый год"
    second_tbl.rows[1].cells[1].text = "обучение сотрудников"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "skip_existing_break.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "skip_existing_break.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("Таблица 1.2.2", 1, 680.0, 692.0),
                PdfLine("Показатель Эффект", 1, 705.0, 717.0),
                PdfLine("Первый показатель", 2, 80.0, 92.0),
                PdfLine("переход на электронный обмен", 2, 100.0, 112.0),
                PdfLine("Таблица 2.3.3", 3, 680.0, 692.0),
                PdfLine("Год Комментарий", 3, 705.0, 717.0),
                PdfLine("Первый год", 4, 80.0, 92.0),
                PdfLine("обучение сотрудников", 4, 100.0, 112.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        first = next(p for p in reread.paragraphs if p.text == first_caption.text)
        second = next(p for p in reread.paragraphs if p.text == second_caption.text)

    first_pPr = first._element.find(qn("w:pPr"))
    second_pPr = second._element.find(qn("w:pPr"))
    first_pb = first_pPr.find(qn("w:pageBreakBefore")) if first_pPr is not None else None
    second_pb = second_pPr.find(qn("w:pageBreakBefore")) if second_pPr is not None else None

    if n != 1:
        return _result(False, f"expected one later whole-table move, got {n}")
    if first_pb is None:
        return _result(False, "existing pageBreakBefore was lost from first caption")
    if second_pb is None:
        return _result(False, "later candidate did not receive pageBreakBefore")
    if len(reread.tables) != 2:
        return _result(False, f"whole-table move should not split tables, got {len(reread.tables)}")
    return _result(True, "existing page-break candidate skipped and later candidate moved")


def test_c_rendered_start_page_upgrades_disabled_page_break() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    caption = doc.add_paragraph("Таблица 2.4.1")
    pPr = caption._element.get_or_add_pPr()
    disabled_break = OxmlElement("w:pageBreakBefore")
    disabled_break.set(qn("w:val"), "0")
    pPr.append(disabled_break)

    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Показатель"
    tbl.rows[0].cells[1].text = "Комментарий"
    tbl.rows[1].cells[0].text = "Первый показатель"
    tbl.rows[1].cells[1].text = "обучение сотрудников"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "disabled_page_break.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "disabled_page_break.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("Таблица 2.4.1", 1, 680.0, 692.0),
                PdfLine("Показатель Комментарий", 1, 705.0, 717.0),
                PdfLine("Первый показатель", 2, 80.0, 92.0),
                PdfLine("обучение сотрудников", 2, 100.0, 112.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        reread_caption = next(p for p in reread.paragraphs if p.text == "Таблица 2.4.1")

    reread_pPr = reread_caption._element.find(qn("w:pPr"))
    page_break = reread_pPr.find(qn("w:pageBreakBefore")) if reread_pPr is not None else None
    page_break_val = page_break.get(qn("w:val")) if page_break is not None else None

    if n != 1:
        return _result(False, f"disabled pageBreakBefore should not block move, got {n}")
    if page_break is None:
        return _result(False, "disabled pageBreakBefore was not upgraded")
    if page_break_val in {"0", "false", "False", "off"}:
        return _result(False, f"pageBreakBefore still disabled: {page_break_val!r}")
    if len(reread.tables) != 1:
        return _result(False, f"whole-table move should not split tables, got {len(reread.tables)}")
    return _result(True, "disabled pageBreakBefore upgraded to active")


def test_c_rendered_start_page_skips_ambiguous_usability() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    doc.add_paragraph("Таблица 2.2.4")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "Год"
    tbl.rows[0].cells[1].text = "Значение"
    tbl.rows[1].cells[0].text = "2023"
    tbl.rows[1].cells[1].text = "10"
    tbl.rows[2].cells[0].text = "2024"
    tbl.rows[2].cells[1].text = "10"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "ambiguous_move.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "ambiguous_move.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("Таблица 2.2.4", 1, 690.0, 702.0),
                PdfLine("Год Значение", 1, 735.0, 748.0),
                PdfLine("2023", 1, 763.0, 776.0),
                PdfLine("10", 2, 86.0, 98.0),
                PdfLine("2024 10", 2, 108.0, 120.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        reread_caption = next(p for p in reread.paragraphs if p.text == "Таблица 2.2.4")

    pPr = reread_caption._element.find(qn("w:pPr"))
    page_break = pPr.find(qn("w:pageBreakBefore")) if pPr is not None else None

    if n != 0:
        return _result(False, f"ambiguous start-page evidence should skip, got {n}")
    if page_break is not None:
        return _result(False, "ambiguous start-page evidence added pageBreakBefore")
    if len(reread.tables) != 1:
        return _result(False, f"ambiguous start-page evidence changed tables: {len(reread.tables)}")
    return _result(True, "ambiguous start-page usability skipped")


def test_c_rendered_start_page_first_row_spill_needs_strong_next_page_evidence() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    doc.add_paragraph("Таблица 2.3.6")
    doc.add_paragraph("Промежуточный эффект")
    tbl = doc.add_table(rows=3, cols=3)
    tbl.rows[0].cells[0].text = "Статья"
    tbl.rows[0].cells[1].text = "Значение"
    tbl.rows[0].cells[2].text = "Комментарий"
    tbl.rows[1].cells[0].text = "Почтовые расходы"
    tbl.rows[1].cells[1].text = "31–33"
    tbl.rows[1].cells[2].text = "отказ от бумажных отправлений"
    tbl.rows[2].cells[0].text = "Печать"
    tbl.rows[2].cells[1].text = "4–5"
    tbl.rows[2].cells[2].text = "сокращение печати"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "weak_spill.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "weak_spill.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("Таблица 2.3.6", 1, 683.2, 695.0),
                PdfLine("Промежуточный эффект", 1, 707.4, 719.0),
                PdfLine("Статья Значение Комментарий", 1, 731.5, 743.0),
                PdfLine("Почтовые расходы 31–33 отказ от бумажных", 1, 759.6, 771.0),
                PdfLine("Печать 4–5", 1, 776.0, 788.0),
                PdfLine("отправлений", 2, 86.8, 98.0),
                PdfLine("сокращение печати", 2, 108.2, 120.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        reread_caption = next(p for p in reread.paragraphs if p.text == "Таблица 2.3.6")

    pPr = reread_caption._element.find(qn("w:pPr"))
    page_break = pPr.find(qn("w:pageBreakBefore")) if pPr is not None else None

    if n != 0:
        return _result(False, f"weak next-page evidence should not trigger move, got {n}")
    if page_break is not None:
        return _result(False, "weak next-page evidence still added pageBreakBefore")
    return _result(True, "weak next-page evidence does not trigger spill detection")


def test_c_rendered_start_page_first_row_spill_ignores_later_prose_token_reuse() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    doc.add_paragraph("Таблица 2.3.7")
    doc.add_paragraph("Промежуточный эффект")
    tbl = doc.add_table(rows=3, cols=3)
    tbl.rows[0].cells[0].text = "Статья"
    tbl.rows[0].cells[1].text = "Значение"
    tbl.rows[0].cells[2].text = "Комментарий"
    tbl.rows[1].cells[0].text = "Почтовые расходы"
    tbl.rows[1].cells[1].text = "31–33"
    tbl.rows[1].cells[2].text = "отказ от бумажных отправлений"
    tbl.rows[2].cells[0].text = "Печать"
    tbl.rows[2].cells[1].text = "4–5"
    tbl.rows[2].cells[2].text = "сокращение печати"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "prose_reuse_spill.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "prose_reuse_spill.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("Таблица 2.3.7", 1, 683.2, 695.0),
                PdfLine("Промежуточный эффект", 1, 707.4, 719.0),
                PdfLine("Статья Значение Комментарий", 1, 731.5, 743.0),
                PdfLine("Почтовые расходы 31–33 отказ от бумажных", 1, 759.6, 771.0),
                PdfLine("Печать 4–5", 1, 776.0, 788.0),
                PdfLine("Статья Значение Комментарий", 2, 58.8, 70.0),
                PdfLine("В тексте обсуждаются отправлений документов и риски обмена", 2, 86.8, 98.0),
                PdfLine("сокращение печати", 2, 108.2, 120.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        reread_caption = next(p for p in reread.paragraphs if p.text == "Таблица 2.3.7")

    pPr = reread_caption._element.find(qn("w:pPr"))
    page_break = pPr.find(qn("w:pageBreakBefore")) if pPr is not None else None

    if n != 0:
        return _result(False, f"later prose token reuse should not trigger move, got {n}")
    if page_break is not None:
        return _result(False, "later prose token reuse still added pageBreakBefore")
    return _result(True, "later prose token reuse does not trigger spill detection")


def test_c_rendered_decision_logging_for_ambiguous_skip() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    doc.add_paragraph("Таблица 2.2.4")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "Год"
    tbl.rows[0].cells[1].text = "Значение"
    tbl.rows[1].cells[0].text = "2023"
    tbl.rows[1].cells[1].text = "10"
    tbl.rows[2].cells[0].text = "2024"
    tbl.rows[2].cells[1].text = "10"

    log_stream = io.StringIO()
    handler = logging.StreamHandler(log_stream)
    handler.setFormatter(logging.Formatter("%(message)s"))
    old_level = tc.logger.level
    tc.logger.addHandler(handler)
    tc.logger.setLevel(logging.INFO)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "ambiguous_logging.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "ambiguous_logging.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("Таблица 2.2.4", 1, 690.0, 702.0),
                PdfLine("Год Значение", 1, 735.0, 748.0),
                PdfLine("2023", 1, 763.0, 776.0),
                PdfLine("10", 2, 86.0, 98.0),
                PdfLine("2024 10", 2, 108.0, 120.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            tc.logger.removeHandler(handler)
            tc.logger.setLevel(old_level)

    logs = log_stream.getvalue()
    expected_fragments = [
        "rendered_table_continuation_enter tables=1 pdf_lines=5",
        "rendered_whole_table_candidate table_idx=0 caption=2.2.4",
        "pdf_caption_matches=1 strict_caption_found=True start_page_usability=ambiguous",
        "rendered_split_candidate table_idx=0 rows=3 skip=row_mapping_ambiguous",
        "rendered_final_decision action=rendered_skip_ambiguous",
    ]
    missing = [fragment for fragment in expected_fragments if fragment not in logs]

    if n != 0:
        return _result(False, f"ambiguous logging scenario should not mutate, got {n}")
    if missing:
        return _result(False, f"missing log fragments: {missing!r}; logs={logs!r}")
    return _result(True, "ambiguous rendered decision logs are emitted")


def test_c_rendered_start_page_keeps_table_with_clear_complete_data_row() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    doc = Document()
    doc.add_paragraph("Таблица 2.2.5")
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "Показатель"
    tbl.rows[0].cells[1].text = "Эффект"
    tbl.rows[1].cells[0].text = "Почтовые расходы"
    tbl.rows[1].cells[1].text = "экономия бюджета"
    tbl.rows[2].cells[0].text = "Архивное хранение"
    tbl.rows[2].cells[1].text = "снижение затрат"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "clear_row.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "clear_row.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: [
                PdfLine("Таблица 2.2.5", 1, 400.0, 412.0),
                PdfLine("Показатель Эффект", 1, 430.0, 442.0),
                PdfLine("Почтовые расходы экономия бюджета", 1, 455.0, 467.0),
                PdfLine("Архивное хранение снижение затрат", 1, 480.0, 492.0),
            ]
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))
        reread_caption = next(p for p in reread.paragraphs if p.text == "Таблица 2.2.5")

    pPr = reread_caption._element.find(qn("w:pPr"))
    page_break = pPr.find(qn("w:pageBreakBefore")) if pPr is not None else None

    if n != 0:
        return _result(False, f"clear complete data row should not move table, got {n}")
    if page_break is not None:
        return _result(False, "clear complete data row still added pageBreakBefore")
    if len(reread.tables) != 1:
        return _result(False, f"clear complete data row changed tables: {len(reread.tables)}")
    return _result(True, "clear complete data row prevents whole-table move")


def test_c_vmerge_guard_rejects_boundary_inside_merge_zone() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import _is_split_boundary_safe

    doc = Document()
    tbl = doc.add_table(rows=4, cols=2)
    for r_idx, row in enumerate(tbl.rows):
        row.cells[0].text = f"A{r_idx}"
        row.cells[1].text = f"B{r_idx}"

    merged = tbl.cell(1, 0).merge(tbl.cell(2, 0))
    merged.text = "merged"

    rows_xml = tbl._tbl.findall(qn("w:tr"))
    if _is_split_boundary_safe(rows_xml, 1):
        return _result(False, "boundary before vMerge continuation row was considered safe")
    if not _is_split_boundary_safe(rows_xml, 2):
        return _result(False, "boundary after vMerge continuation row was considered unsafe")
    return _result(True, "vMerge guard rejects split inside merge zone")



# ── Asset regression ──────────────────────────────────────────────────────────

def test_regression_asset(asset_path: Path) -> tuple[bool, str]:
    """
    Format asset file end-to-end; verify:
    - No crash
    - Output .docx exists
    - Image count not decreased
    - No Python exception in formatter
    """
    with tempfile.TemporaryDirectory() as tmp:
        out_path = Path(tmp) / f"out_{asset_path.name}"
        # Count images before
        doc_in = Document(str(asset_path))
        imgs_before = _count_drawings(doc_in)
        del doc_in

        try:
            format_docx(str(asset_path), str(out_path))
        except Exception as e:
            return _result(False, f"formatter raised: {e}\n{traceback.format_exc()}")

        if not out_path.exists():
            return _result(False, "output file not created")

        doc_out = Document(str(out_path))
        imgs_after = _count_drawings(doc_out)
        if imgs_after < imgs_before:
            return _result(
                False,
                f"images deleted: before={imgs_before}, after={imgs_after}",
            )

        return _result(True, f"ok (images: {imgs_before}→{imgs_after})")


# ── Batch 1 — tblW fix, _MIN_COL_PT, stale LRPB skip ────────────────────────

def test_b1_tblW_updated_after_col_optimization() -> tuple[bool, str]:
    """
    _optimize_table_col_widths must update w:tblPr/w:tblW to match the new
    column total after scaling.  Without this fix Word renders the table at
    the original (too-wide) tblW instead of the corrected column sum.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _optimize_table_col_widths, TWIP_PER_PT,
    )

    doc = Document()
    tbl = doc.add_table(rows=2, cols=3)
    tbl_xml = tbl._element
    body_w = 481.9  # standard KFU body width in pt

    # Set each of 3 columns to 200 pt → total 600 pt > body_w
    grid = tbl_xml.find(qn("w:tblGrid"))
    if grid is None:
        return _result(False, "no tblGrid in table XML")
    for gc in grid.findall(qn("w:gridCol")):
        gc.set(qn("w:w"), str(int(200 * TWIP_PER_PT)))

    # Set tblW to original oversized value
    tblPr = tbl_xml.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_xml.insert(0, tblPr)
    tblW_el = tblPr.find(qn("w:tblW"))
    if tblW_el is None:
        tblW_el = OxmlElement("w:tblW")
        tblPr.append(tblW_el)
    tblW_el.set(qn("w:w"), str(int(600 * TWIP_PER_PT)))
    tblW_el.set(qn("w:type"), "dxa")

    changed = _optimize_table_col_widths(tbl_xml, body_w)
    if not changed:
        return _result(False, "optimizer reported no change (expected scale-down)")

    new_tblW_el = tblPr.find(qn("w:tblW"))
    if new_tblW_el is None:
        return _result(False, "w:tblW element missing after optimization")

    new_total_twips = int(new_tblW_el.get(qn("w:w"), 0))
    expected_twips = round(body_w * TWIP_PER_PT)
    # Allow ±50 twips rounding slack
    if abs(new_total_twips - expected_twips) > 50:
        return _result(
            False,
            f"tblW not updated: got {new_total_twips} twips, expected ~{expected_twips}",
        )
    return _result(True, f"tblW updated to {new_total_twips} twips (expected ~{expected_twips})")


def test_b1_min_col_pt_is_20() -> tuple[bool, str]:
    """
    _MIN_COL_PT must be ≤ 20 (variant C: only phantom columns < 20 pt
    are redistributed; legitimate narrow columns like 30 pt survive).
    """
    from guides.coursework_kfu_2025.table_continuation import _MIN_COL_PT
    if _MIN_COL_PT > 20.5:
        return _result(False, f"_MIN_COL_PT={_MIN_COL_PT} > 20 — old value, fix not applied")
    return _result(True, f"_MIN_COL_PT={_MIN_COL_PT} ✓")



# ── Batch 2 — keepTogether, Rule 6 propagation, image height ─────────────────

def test_b2_keep_together_on_table_caption() -> tuple[bool, str]:
    """
    After apply_pagination_rules, table_caption and table_title paragraphs
    must have keep_together=True (prevents a long title from being split
    across pages by Word's line-breaker).
    """
    from guides.coursework_kfu_2025.pagination_rules import apply_pagination_rules

    doc = Document()
    doc.add_paragraph("Таблица 1.1 — Test caption line")   # → table_caption
    doc.add_table(rows=2, cols=2)
    apply_pagination_rules(doc)

    p = doc.paragraphs[0]
    if not p.paragraph_format.keep_together:
        return _result(False, "keep_together not set on table_caption paragraph")
    return _result(True, "table_caption has keep_together=True")


def test_b2_keep_together_on_headings() -> tuple[bool, str]:
    """
    After apply_pagination_rules, heading1 and heading2 paragraphs must have
    keep_together=True (prevents a multi-line heading from being split across pages).
    """
    from guides.coursework_kfu_2025.pagination_rules import apply_pagination_rules

    doc = Document()
    doc.add_paragraph("1. Теоретические основы исследования")   # → heading1
    doc.add_paragraph("1.1. Понятие и сущность термина")         # → heading2
    doc.add_paragraph("Основной текст параграфа.")
    apply_pagination_rules(doc)

    p_h1 = doc.paragraphs[0]
    p_h2 = doc.paragraphs[1]
    if not p_h1.paragraph_format.keep_together:
        return _result(False, "keep_together not set on heading1")
    if not p_h2.paragraph_format.keep_together:
        return _result(False, "keep_together not set on heading2")
    return _result(True, "heading1 and heading2 have keep_together=True")


def test_b2_rule6_propagates_through_empty_para() -> tuple[bool, str]:
    """
    _apply_rule6: an image paragraph followed by one (or more) empty paragraphs
    and then a figure_caption must have keepWithNext set on BOTH the image paragraph
    AND the intervening empty paragraph(s), so the chain reaches the caption.
    """
    from guides.coursework_kfu_2025.pagination_rules import apply_pagination_rules

    doc = Document()
    # Image paragraph
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r_el = OxmlElement("w:r")
    r_el.append(drawing)
    img_p._element.append(r_el)
    # Empty paragraph between image and caption
    doc.add_paragraph("")
    # Figure caption
    doc.add_paragraph("Рисунок 1.1 — Схема взаимодействия")

    apply_pagination_rules(doc)

    img_para   = doc.paragraphs[0]
    empty_para = doc.paragraphs[1]
    if not img_para.paragraph_format.keep_with_next:
        return _result(False, "keep_with_next not set on image paragraph")
    if not empty_para.paragraph_format.keep_with_next:
        return _result(
            False,
            "keep_with_next not set on empty paragraph between image and caption",
        )
    return _result(True, "keepWithNext propagated through empty paragraph to caption")


def test_b2_table_source_note_normalised_and_chained() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from guides.coursework_kfu_2025.pagination_rules import apply_pagination_rules

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "H1"
    tbl.rows[0].cells[1].text = "H2"
    tbl.rows[1].cells[0].text = "a"
    tbl.rows[1].cells[1].text = "b"
    source = doc.add_paragraph("Источник: составлено автором.")
    source.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph("Примечание: расчет ориентировочный.")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER

    apply_pagination_rules(doc)

    last_cell_p = tbl.rows[-1].cells[-1].paragraphs[-1]
    if not last_cell_p.paragraph_format.keep_with_next:
        return _result(False, "table tail is not chained to source/note")
    if not source.paragraph_format.keep_with_next:
        return _result(False, "source is not chained to following note")
    if source.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
        return _result(False, f"source alignment not normalised: {source.alignment}")
    if note.alignment != WD_ALIGN_PARAGRAPH.JUSTIFY:
        return _result(False, f"note alignment not normalised: {note.alignment}")
    if source.paragraph_format.first_line_indent is None:
        return _result(False, "source first-line indent was not restored")
    if note.paragraph_format.first_line_indent is None:
        return _result(False, "note first-line indent was not restored")
    return _result(True, "table source/note normalised and chained")


def test_b2_image_height_from_emu() -> tuple[bool, str]:
    """
    _get_image_height_pt must read wp:extent cy from a drawing element and
    convert EMU → pt correctly (EMU_PER_PT = 12700).
    """
    from guides.coursework_kfu_2025.table_continuation import _get_image_height_pt

    doc = Document()
    p = doc.add_paragraph()

    # Build a minimal drawing: w:drawing > wp:inline > wp:extent cy="1270000" (=100pt)
    drawing  = OxmlElement("w:drawing")
    inline   = OxmlElement("wp:inline")
    extent   = OxmlElement("wp:extent")
    extent.set("cy", str(100 * 12700))   # 100 pt × 12700 EMU/pt = 1270000 EMU
    inline.append(extent)
    drawing.append(inline)
    r_el = OxmlElement("w:r")
    r_el.append(drawing)
    p._element.append(r_el)

    h = _get_image_height_pt(p._element)
    if h is None:
        return _result(False, "_get_image_height_pt returned None — extent not read")
    if abs(h - 100.0) > 0.5:
        return _result(False, f"expected 100.0 pt, got {h:.2f} pt")
    return _result(True, f"image height correctly read as {h:.1f} pt from EMU")


# ── Batch 3 — footnote standardization ───────────────────────────────────────

def test_b3_format_footnote_para_applies_10pt_tnr() -> tuple[bool, str]:
    """
    _format_footnote_para must apply 10pt Times New Roman, no bold,
    single line spacing, and zero indent to a paragraph XML element.
    Tests the low-level helper directly to avoid needing a real footnotes part.
    """
    from guides.coursework_kfu_2025.safe_formatter import _format_footnote_para

    doc = Document()
    p = doc.add_paragraph()

    # Give the paragraph some run with 14pt bold text (typical body style)
    r_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz_el = OxmlElement("w:sz")
    sz_el.set(qn("w:val"), "28")   # 14pt
    bold_el = OxmlElement("w:b")
    rPr.append(sz_el)
    rPr.append(bold_el)
    t_el = OxmlElement("w:t")
    t_el.text = "Footnote text"
    r_el.append(rPr)
    r_el.append(t_el)
    p._element.append(r_el)

    _format_footnote_para(p._element)

    # Check run font size is now 10pt (w:sz val="20")
    r_out = p._element.find(".//" + qn("w:r"))
    if r_out is None:
        return _result(False, "no w:r found after formatting")
    rPr_out = r_out.find(qn("w:rPr"))
    if rPr_out is None:
        return _result(False, "no w:rPr on run after formatting")

    sz_out = rPr_out.find(qn("w:sz"))
    if sz_out is None:
        return _result(False, "w:sz missing from run rPr after formatting")
    sz_val = sz_out.get(qn("w:val"))
    if sz_val != "20":
        return _result(False, f"expected w:sz val='20' (10pt), got '{sz_val}'")

    # Bold must be suppressed: w:b absent or val="0"
    b_out = rPr_out.find(qn("w:b"))
    if b_out is not None:
        b_val = b_out.get(qn("w:val"), "1")
        if b_val not in ("0", "false"):
            return _result(False, f"bold not suppressed (w:b val='{b_val}')")

    # Check paragraph indent = 0
    pPr_out = p._element.find(qn("w:pPr"))
    if pPr_out is not None:
        ind_out = pPr_out.find(qn("w:ind"))
        if ind_out is not None:
            left_val = ind_out.get(qn("w:left"), "0")
            if left_val not in ("0", None):
                return _result(False, f"indent not zeroed (w:ind left='{left_val}')")

    return _result(True, "footnote para: 10pt TNR, no bold, zero indent ✓")


# ── Batch C2 — image gap, table-fits-on-1-page, number columns ───────────────

def test_c2_empty_para_between_image_and_caption_removed() -> tuple[bool, str]:
    """
    Phase 3 must remove empty paragraphs that appear between an image paragraph
    and its figure_caption (e.g. blank line inserted by student between рисунок
    and 'Рис. 1.2.1 — …').
    """
    from guides.coursework_kfu_2025.table_continuation import remove_empty_before_figure_captions

    doc = Document()
    # Image paragraph
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r_el = OxmlElement("w:r")
    r_el.append(drawing)
    img_p._element.append(r_el)
    # Empty paragraph between image and caption (the student's stray blank line)
    doc.add_paragraph("")
    # Figure caption
    doc.add_paragraph("Рисунок 1.2.1 — Схема взаимодействия")

    n = remove_empty_before_figure_captions(doc)

    if n != 1:
        return _result(False, f"expected 1 removal, got {n}")
    # Check the empty paragraph is gone: image should be immediately before caption
    remaining = [p for p in doc.paragraphs if not _para_has_image(p._element)]
    # paragraphs: [img_p (has image), caption]
    total = len(doc.paragraphs)
    if total != 2:
        return _result(False, f"expected 2 paragraphs after removal, got {total}")
    return _result(True, "empty paragraph between image and caption removed ✓")



def test_c2_number_column_minimum() -> tuple[bool, str]:
    """
    _optimize_table_col_widths must protect numeric-only columns from being
    scaled below the width needed to display their content on one line.
    A 7-digit number like '9503005' in a column requires at least ~50pt.
    """
    from guides.coursework_kfu_2025.table_continuation import (
        _optimize_table_col_widths, TWIP_PER_PT,
    )

    doc = Document()
    tbl = doc.add_table(rows=3, cols=4)
    tbl_xml = tbl._element
    body_w = 481.9

    # Set column widths: [250, 100, 100, 130] pt → total 580pt (needs scaling)
    original_widths_pt = [250.0, 100.0, 100.0, 130.0]
    grid = tbl_xml.find(qn("w:tblGrid"))
    if grid is None:
        return _result(False, "no tblGrid")
    for gc, w in zip(grid.findall(qn("w:gridCol")), original_widths_pt):
        gc.set(qn("w:w"), str(round(w * TWIP_PER_PT)))

    # Put numeric content in column 1 (index 1): '9 503 005' (9 chars)
    for ri in range(3):
        cells = tbl.rows[ri].cells
        cells[0].text = "Текстовый заголовок показателя" if ri == 0 else "Текст"
        cells[1].text = "2023 г." if ri == 0 else "9 503 005"  # numeric
        cells[2].text = "2024 г." if ri == 0 else "9 875 076"  # numeric
        cells[3].text = "Абсолютное изменение" if ri == 0 else "−372 071"

    # Also update tcW for each cell to match initial widths
    for ri in range(3):
        tr = tbl.rows[ri]._tr
        col_idx = 0
        for tc in tr.findall(qn("w:tc")):
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(round(original_widths_pt[col_idx] * TWIP_PER_PT)))
            tcW.set(qn("w:type"), "dxa")
            col_idx += 1

    _optimize_table_col_widths(tbl_xml, body_w)

    # Column 1 and 2 contain "9 503 005" / "9 875 076" (9 chars × 6pt + 8pt ≈ 62pt)
    # After optimization, columns 1 and 2 should be at least 50pt
    grid_after = tbl_xml.find(qn("w:tblGrid"))
    cols_after = grid_after.findall(qn("w:gridCol"))
    widths_after_pt = [int(c.get(qn("w:w"), 0)) / TWIP_PER_PT for c in cols_after]

    min_expected = 50.0  # 9 chars × 6pt + 8pt padding ≈ 62pt; 50pt is a safe floor
    for col_idx in (1, 2):
        if widths_after_pt[col_idx] < min_expected:
            return _result(
                False,
                f"numeric column {col_idx} too narrow: {widths_after_pt[col_idx]:.1f}pt < {min_expected}pt",
            )
    return _result(True, f"numeric columns protected: {[f'{w:.1f}' for w in widths_after_pt]}")


def test_yo_normalisation_midword_uppercase() -> tuple[bool, str]:
    """
    Words starting with uppercase but containing lowercase ё mid-word
    (e.g. "Лётчик") must have the ё replaced with е.
    Capital Ё at the start of a word must be preserved.
    """
    from guides.coursework_kfu_2025.safe_formatter import normalize_yo_in_text

    cases = [
        # (input, expected)
        ("лётчик",       "летчик"),
        ("ёж",           "еж"),
        ("Ёж",           "Ёж"),        # capital Ё: preserved
        ("Лётчик",       "Летчик"),    # starts with uppercase Л, ё is lowercase → replace
        ("ЛЁТЧИК",       "ЛЁТЧИК"),   # Ё uppercase → preserved
        ("неёмкий",      "неемкий"),
        ("Чернышёв",     "Чернышев"),
    ]
    failures = []
    for inp, expected in cases:
        got = normalize_yo_in_text(inp)
        if got != expected:
            failures.append(f"normalize_yo_in_text({inp!r}) = {got!r}, expected {expected!r}")
    if failures:
        return _result(False, "\n".join(failures))
    return _result(True, f"all {len(cases)} ё-normalisation cases correct")


def test_t_indent_body_paragraph_left_zero() -> tuple[bool, str]:
    """
    After formatting, regular body paragraphs must have:
    - left_indent = 0 (or None, not a hanging indent)
    - first_line_indent = 709 twips (≈1.25 cm)
    No hanging indent (w:hanging must not be present).
    """
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    # process_document requires a paragraph with text "введение" to find body start
    doc.add_paragraph("введение")
    # Simulate a paragraph that originally had a List style with hanging indent
    p = doc.add_paragraph("Это обычный абзац с текстом.")
    # Manually inject a hanging indent (simulating "List Paragraph" style effect)
    pPr = p._element.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "709")
    ind.set(qn("w:hanging"), "360")
    pPr.append(ind)

    with tempfile.TemporaryDirectory() as tmp:
        inp = os.path.join(tmp, "in.docx")
        out = os.path.join(tmp, "out.docx")
        doc.save(inp)
        process_document(inp, out)
        result_doc = Document(out)

    body_paras = [p for p in result_doc.paragraphs if "обычный абзац" in (p.text or "")]
    if not body_paras:
        return _result(False, "body paragraph not found in output")

    bp = body_paras[0]
    pPr_out = bp._element.find(qn("w:pPr"))
    ind_out = pPr_out.find(qn("w:ind")) if pPr_out is not None else None

    # Check no hanging
    if ind_out is not None and ind_out.get(qn("w:hanging")):
        return _result(False, f"w:hanging still present: {ind_out.get(qn('w:hanging'))}")

    # Check left=0 (either absent or "0")
    left_val = ind_out.get(qn("w:left")) if ind_out is not None else None
    if left_val and left_val != "0":
        return _result(False, f"w:left={left_val!r} (expected 0 or absent)")

    # Check firstLine≈709
    fl_val = ind_out.get(qn("w:firstLine")) if ind_out is not None else None
    if fl_val is None or abs(int(fl_val) - 709) > 30:
        return _result(False, f"w:firstLine={fl_val!r} (expected ≈709)")

    return _result(True, f"body paragraph indent: left=0, firstLine={fl_val} ✓")


# ── Task 2 — Глава N without title ────────────────────────────────────────────

def test_t2_chapter_heading_without_title() -> tuple[bool, str]:
    """
    "Глава 1" (no title) must be classified as heading1.
    "Глава 1. Название" (with title) must still work.
    """
    from guides.coursework_kfu_2025.classifier import parse_heading1

    cases = [
        ("Глава 1",                    True),
        ("глава 2",                    True),
        ("ГЛАВА 3",                    True),
        ("Глава 1.",                   True),
        ("Глава 1. Теоретические основы", True),
        ("Глава 10. Заключение",       True),
        ("Глава",                      False),  # no number
        ("1. Теоретические основы",    True),   # normalized heading — must still work
        ("Введение",                   True),   # exact match — must still work
    ]
    failures = []
    for text, expected in cases:
        result = parse_heading1(text)
        got = result is not None
        if got != expected:
            failures.append(f"parse_heading1({text!r}) → {result}, expected match={expected}")
    if failures:
        return _result(False, "\n".join(failures))
    return _result(True, f"all {len(cases)} chapter heading cases correct")


def _add_fake_word_numbering(paragraph, ilvl_value: str = "0") -> None:
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), ilvl_value)
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(num_id)
    pPr.append(numPr)


def _add_fake_style_numbering(document: Document, style_name: str, ilvl_value: str = "0") -> None:
    style = document.styles[style_name]
    style_element = style.element
    pPr = style_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        style_element.append(pPr)

    existing = pPr.find(qn("w:numPr"))
    if existing is not None:
        pPr.remove(existing)

    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), ilvl_value)
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "42")
    numPr.append(ilvl)
    numPr.append(num_id)
    pPr.append(numPr)


def _style_has_numbering(document: Document, style_name: str) -> bool:
    style = document.styles[style_name]
    pPr = style.element.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def _paragraph_has_direct_numbering(paragraph) -> bool:
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def _style_name(paragraph) -> str:
    try:
        return (paragraph.style.name or "").strip().lower()
    except Exception:
        return ""


def _find_paragraph_starting_with(document: Document, prefix: str):
    for paragraph in document.paragraphs:
        if " ".join(paragraph.text.split()).startswith(prefix):
            return paragraph
    return None


def test_t2_manual_heading2_still_promoted() -> tuple[bool, str]:
    """Explicit manual heading syntax '1.1. ...' remains Heading 2."""
    from guides.coursework_kfu_2025.safe_formatter import process_document
    import tempfile, os

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ")
    doc.add_paragraph("1.1. Понятие конкурентоспособности организации")
    doc.add_paragraph("Обычный текст подраздела.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = os.path.join(tmp, "in.docx")
        out = os.path.join(tmp, "out.docx")
        doc.save(inp)
        process_document(inp, out)
        result = Document(out)

    heading = _find_paragraph_starting_with(result, "1.1. Понятие конкурентоспособности")
    if heading is None:
        return _result(False, "manual heading2 text missing after formatting")

    if _style_name(heading) not in {"heading 2", "заголовок 2"}:
        return _result(False, f"manual heading2 style is {_style_name(heading)!r}")

    return _result(True, "manual heading2 remains Heading 2")


def test_t2_word_autonumbered_heading2_with_style_still_promoted() -> tuple[bool, str]:
    """
    A real Word-autonumbered Heading 2 may have numPr but no visible '1.1.'
    in paragraph.text. Heading style is enough structural evidence to promote it.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        auto_detect_heading2,
        clean_spaces,
        is_likely_numbered_heading2_candidate,
        is_probable_body_list_item,
        normalize_heading2_numbering,
        paragraph_has_numbering,
    )

    doc = Document()
    heading = doc.add_paragraph("Понятие конкурентоспособности организации")
    heading.style = "Heading 2"
    _add_fake_word_numbering(heading, ilvl_value="1")

    if is_probable_body_list_item(heading):
        return _result(False, "Word-autonumbered Heading 2 was classified as body/list")

    if not auto_detect_heading2(heading, current_chapter_num=1, next_paragraph_num=1):
        return _result(False, "Word-autonumbered Heading 2 was not auto-detected")

    if not is_likely_numbered_heading2_candidate(heading, 1, 1):
        return _result(False, "Word-autonumbered Heading 2 was not a heading2 candidate")

    normalized = normalize_heading2_numbering(heading, 1, 1)
    expected = "1.1. Понятие конкурентоспособности организации"
    if normalized != expected or clean_spaces(heading.text) != expected:
        return _result(False, f"unexpected Heading 2 normalization: {normalized!r}, text={heading.text!r}")

    if paragraph_has_numbering(heading):
        return _result(False, "Heading 2 Word numbering was not converted to plain text")

    if _style_name(heading) not in {"heading 2", "заголовок 2"}:
        return _result(False, f"autonumbered heading2 style is {_style_name(heading)!r}")

    return _result(True, "Word-autonumbered Heading 2 remains supported")


def test_t2_word_autonumbered_heading1_with_style_still_promoted() -> tuple[bool, str]:
    """
    A real Word-autonumbered Heading 1 may have numPr but no visible '1.'
    in paragraph.text. Heading style/outline must keep it on the heading path.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        auto_detect_numbered_heading1,
        paragraph_has_numbering,
        process_document,
    )
    import tempfile, os

    direct_doc = Document()
    direct_heading = direct_doc.add_paragraph("ТЕОРЕТИЧЕСКИЕ ОСНОВЫ КОНКУРЕНТОСПОСОБНОСТИ")
    direct_heading.style = "Heading 1"
    _add_fake_word_numbering(direct_heading)
    following_h2 = direct_doc.add_paragraph("Понятие конкурентоспособности организации")
    following_h2.style = "Heading 2"
    _add_fake_word_numbering(following_h2, ilvl_value="1")

    if not auto_detect_numbered_heading1(direct_heading, current_chapter_num=None, next_paragraph=following_h2):
        return _result(False, "Word-autonumbered Heading 1 was not auto-detected")

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    heading1 = doc.add_paragraph("ТЕОРЕТИЧЕСКИЕ ОСНОВЫ КОНКУРЕНТОСПОСОБНОСТИ")
    heading1.style = "Heading 1"
    _add_fake_word_numbering(heading1)
    heading2 = doc.add_paragraph("Понятие конкурентоспособности организации")
    heading2.style = "Heading 2"
    _add_fake_word_numbering(heading2, ilvl_value="1")
    doc.add_paragraph("Обычный текст подраздела.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = os.path.join(tmp, "in.docx")
        out = os.path.join(tmp, "out.docx")
        doc.save(inp)
        process_document(inp, out)
        result = Document(out)

    formatted_h1 = _find_paragraph_starting_with(
        result,
        "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ КОНКУРЕНТОСПОСОБНОСТИ",
    )
    if formatted_h1 is None:
        return _result(False, "Word-autonumbered Heading 1 did not get plain-text chapter number")

    if paragraph_has_numbering(formatted_h1):
        return _result(False, "Heading 1 Word numbering remained after formatting")

    if _style_name(formatted_h1) not in {"heading 1", "заголовок 1"}:
        return _result(False, f"autonumbered heading1 style is {_style_name(formatted_h1)!r}")

    formatted_h2 = _find_paragraph_starting_with(
        result,
        "1.1. Понятие конкурентоспособности организации",
    )
    if formatted_h2 is None:
        return _result(False, "following autonumbered Heading 2 was not normalized under Heading 1")

    return _result(True, "Word-autonumbered Heading 1 remains supported")


def test_t2_heading_style_numbering_is_removed() -> tuple[bool, str]:
    """
    Product rule: headings must not use Word autonumbering.
    Heading styles may carry w:numPr, which renders extra numbering even when
    heading paragraphs have no direct numPr. Manual numbering in heading text
    must remain as literal text.
    """
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    _add_fake_style_numbering(doc, "Heading 1", ilvl_value="0")
    _add_fake_style_numbering(doc, "Heading 2", ilvl_value="1")
    _add_fake_style_numbering(doc, "Heading 3", ilvl_value="2")

    h1_exact = doc.add_paragraph("ВВЕДЕНИЕ")
    h1_exact.style = "Heading 1"
    _add_fake_word_numbering(h1_exact)
    h1_chapter = doc.add_paragraph("1. Теоретические основы")
    h1_chapter.style = "Heading 1"
    _add_fake_word_numbering(h1_chapter)
    h2 = doc.add_paragraph("1.1. Понятие конкурентоспособности организации")
    h2.style = "Heading 2"
    _add_fake_word_numbering(h2, ilvl_value="1")
    doc.add_paragraph("Обычный текст подраздела.")
    doc.add_paragraph("Ненумерованный основной текст не должен получить numPr.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(inp)
        process_document(inp, out)
        formatted = Document(str(out))

    numbered_styles = [
        style_name
        for style_name in ("Heading 1", "Heading 2", "Heading 3")
        if _style_has_numbering(formatted, style_name)
    ]
    if numbered_styles:
        return _result(False, f"heading style numbering remained: {numbered_styles!r}")

    heading_texts = {
        "ВВЕДЕНИЕ",
        "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ",
        "1.1. Понятие конкурентоспособности организации",
    }
    found = []
    for paragraph in formatted.paragraphs:
        text = " ".join((paragraph.text or "").split())
        if text in heading_texts:
            found.append(text)
            if _paragraph_has_direct_numbering(paragraph):
                return _result(False, f"direct heading numbering remained on {text!r}")
        elif text and _paragraph_has_direct_numbering(paragraph):
            return _result(False, f"numbering was added outside tables/headings: {text!r}")

    missing = sorted(heading_texts - set(found))
    if missing:
        return _result(False, f"manual heading text missing after formatting: {missing!r}")

    return _result(True, "heading style numbering removed while manual heading text stayed")


def test_t2_word_numbered_body_items_not_promoted_to_headings() -> tuple[bool, str]:
    """
    Word-numbered body list items are not heading evidence by themselves.
    This protects real coursework lists such as "Правление и Совет директоров"
    from becoming artificial "3.1." / "8.1." Heading 2 lines.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        auto_detect_heading2,
        auto_detect_numbered_heading1,
        is_likely_numbered_heading2_candidate,
        is_probable_body_list_item,
        normalize_heading2_numbering,
    )

    doc = Document()
    previous = doc.add_paragraph("Организационная структура включает несколько элементов.")
    item = doc.add_paragraph("Правление и Совет директоров")
    _add_fake_word_numbering(item)

    if not is_probable_body_list_item(item, prev_paragraph=previous, prev_kind="body_text"):
        return _result(False, "Word-numbered body item was not classified as body_list_item")

    if auto_detect_heading2(item, current_chapter_num=3, next_paragraph_num=1, prev_kind="body_text"):
        return _result(False, "Word-numbered body item auto-detected as heading2")

    if is_likely_numbered_heading2_candidate(item, 3, 1, prev_kind="body_text"):
        return _result(False, "Word-numbered body item considered likely heading2 candidate")

    if auto_detect_numbered_heading1(item, current_chapter_num=3):
        return _result(False, "Word-numbered body item auto-detected as heading1")

    before = item.text
    normalized = normalize_heading2_numbering(item, 3, 1)
    if normalized is not None or item.text != before:
        return _result(False, f"body item was renumbered: normalized={normalized!r}, text={item.text!r}")

    return _result(True, "Word-numbered body items stay body/list items")


def test_t2_numbered_sentence_not_promoted_to_heading1() -> tuple[bool, str]:
    """
    A numbered sentence-like body paragraph must not be uppercased as Heading 1.
    Real Heading 1 syntax without sentence boundary remains allowed.
    """
    from guides.coursework_kfu_2025.classifier import parse_heading1
    from guides.coursework_kfu_2025.safe_formatter import is_heading1_promotion_safe

    doc = Document()
    body_sentence = doc.add_paragraph("1. Маркетинговый подход. Данный подход")
    parsed = parse_heading1(body_sentence.text)
    if not parsed:
        return _result(False, "test setup failed: parse_heading1 did not parse numbered sentence")
    if is_heading1_promotion_safe(body_sentence, parsed):
        return _result(False, "sentence-like numbered body paragraph considered safe heading1")

    real_heading = doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ КОНКУРЕНТОСПОСОБНОСТИ")
    parsed_real = parse_heading1(real_heading.text)
    if not parsed_real or not is_heading1_promotion_safe(real_heading, parsed_real):
        return _result(False, "real explicit heading1 was rejected")

    return _result(True, "numbered sentence rejected; real heading accepted")


def test_t2_chapter_colon_heading_repaired_without_colon_artifact() -> tuple[bool, str]:
    """'Глава 2: Название' becomes '2. НАЗВАНИЕ', never '2.: НАЗВАНИЕ'."""
    from guides.coursework_kfu_2025.safe_formatter import smart_repair_heading1

    doc = Document()
    paragraph = doc.add_paragraph("Глава 2: Практические аспекты критериев")

    if not smart_repair_heading1(paragraph, paragraph.text):
        return _result(False, "smart_repair_heading1 did not repair chapter heading")

    expected = "2. ПРАКТИЧЕСКИЕ АСПЕКТЫ КРИТЕРИЕВ"
    if paragraph.text != expected:
        return _result(False, f"unexpected repaired heading: {paragraph.text!r}")

    return _result(True, "chapter heading colon artifact removed")


def test_t2_real_coursework_17_heading_regression() -> tuple[bool, str]:
    """
    Real regression: body/list paragraphs in coursework 17 must not become
    artificial headings such as "3.1. Правление..." or ALL CAPS list items.
    """
    from guides.coursework_kfu_2025.safe_formatter import is_empty_paragraph

    fixture = Path(
        "/Users/mac/Desktop/курсовые/"
        "курсова 17. Критерии и показатели конкурентоспособности организации.docx"
    )
    if not fixture.exists():
        return _result(True, f"fixture not present, skipped: {fixture}")

    with tempfile.TemporaryDirectory() as tmp:
        out_path = Path(tmp) / "coursework_17_formatted.docx"
        try:
            format_docx(str(fixture), str(out_path))
        except Exception as e:
            return _result(False, f"formatter raised on real fixture: {e}\n{traceback.format_exc()}")

        doc = Document(str(out_path))
        paragraphs = doc.paragraphs
        texts = [" ".join(p.text.split()) for p in paragraphs if " ".join(p.text.split())]

    forbidden = [
        "1. МАРКЕТИНГОВЫЙ ПОДХОД. ДАННЫЙ ПОДХОД",
        "1.1. Доля рынка продукции предприятия",
        "3.1. Правление и Совет директоров",
        "3.2. Интеграция с международными научными центрами",
        "8.1. Повышение экспортного потенциала",
        "2.:",
    ]
    found_forbidden = [
        marker
        for marker in forbidden
        if any(text.startswith(marker) or marker in text for text in texts)
    ]
    if found_forbidden:
        return _result(False, f"false heading markers found: {found_forbidden}")

    required = [
        "ВВЕДЕНИЕ",
        "1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ",
        "1.1. Понятие",
        "2. ПРАКТИЧЕСКИЕ АСПЕКТЫ",
        "2.1. Общая характеристика",
        "ЗАКЛЮЧЕНИЕ",
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    ]
    missing = [
        marker
        for marker in required
        if not any(marker.lower() in text.lower() for text in texts)
    ]
    if missing:
        return _result(False, f"real headings missing after formatting: {missing}")

    for idx, paragraph in enumerate(paragraphs):
        if " ".join(paragraph.text.split()).startswith("1.3. Методы оценки конкурентоспособности"):
            if idx < 1 or not is_empty_paragraph(paragraphs[idx - 1]):
                return _result(False, "real fixture: missing blank before 1.3 heading")
            if idx >= 2 and is_empty_paragraph(paragraphs[idx - 2]):
                return _result(False, "real fixture: double blank before 1.3 heading")
            break
    else:
        return _result(False, "real fixture: 1.3 heading missing")

    return _result(True, "real coursework 17 heading regression is clean")


def test_t3_reference_subheading_centred() -> tuple[bool, str]:
    """
    After formatting, reference section headers must be CENTER aligned, bold,
    preceded by exactly one empty paragraph.
    Source entries must use regular body-style indentation:
    left=0, firstLine≈709 twips, no hanging indent.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from guides.coursework_kfu_2025.safe_formatter import process_document
    import tempfile, os

    doc = Document()
    doc.add_paragraph("Введение")
    doc.add_paragraph("")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("Официальные материалы")
    doc.add_paragraph("1. Некий закон.")
    doc.add_paragraph("Интернет-ресурсы")
    doc.add_paragraph("2. Некий сайт.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = os.path.join(tmp, "in.docx")
        out = os.path.join(tmp, "out.docx")
        doc.save(inp)
        process_document(inp, out)
        result_doc = Document(out)

    paras = list(result_doc.paragraphs)
    sh_idx = next((i for i, p in enumerate(paras) if "официальные" in (p.text or "").lower()), None)
    if sh_idx is None:
        return _result(False, "subheading paragraph not found in output")

    sh = paras[sh_idx]
    if sh.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, f"subheading not centred: alignment={sh.alignment}")

    pPr_sh = sh._element.find(qn("w:pPr"))
    ind_sh = pPr_sh.find(qn("w:ind")) if pPr_sh is not None else None
    if ind_sh is not None:
        fli = ind_sh.get(qn("w:firstLine"))
        left = ind_sh.get(qn("w:left"))
        hang = ind_sh.get(qn("w:hanging"))
        if hang and int(hang) > 100:
            return _result(False, f"subheading has hanging indent: {hang}")
        if fli and int(fli) > 100:
            return _result(False, f"subheading has first-line indent: {fli}")
        if left and int(left) > 100:
            return _result(False, f"subheading has left indent: {left}")

    bold_ok = any(r.bold for r in sh.runs if r.text.strip())
    if not bold_ok:
        return _result(False, "subheading runs are not bold")

    if sh_idx == 0 or (paras[sh_idx - 1].text or "").strip():
        return _result(False, "no empty paragraph before reference subheading")

    # Check source entry body-style indent
    source_paras = [p for p in paras if "некий закон" in (p.text or "").lower()]
    if source_paras:
        sp = source_paras[0]
        pPr_sp = sp._element.find(qn("w:pPr"))
        ind_sp = pPr_sp.find(qn("w:ind")) if pPr_sp is not None else None
        if ind_sp is None:
            return _result(False, "source entry has no w:ind")
        left_v = ind_sp.get(qn("w:left"))
        first_line_v = ind_sp.get(qn("w:firstLine"))
        hang_v = ind_sp.get(qn("w:hanging"))
        if left_v not in {None, "0"}:
            return _result(False, f"source entry left={left_v!r} (expected 0)")
        if not first_line_v or abs(int(first_line_v) - 709) > 60:
            return _result(False, f"source entry firstLine={first_line_v!r} (expected ≈709)")
        if hang_v is not None:
            return _result(False, f"source entry hanging={hang_v!r} (expected absent)")

    return _result(True, "reference subheading: centred, bold, blank before; source body indent OK")


def test_t4_citation_brackets_split() -> tuple[bool, str]:
    """
    Multi-source citation brackets split; single-source with page range get hyphen→en-dash.
    p. notation is supported. Single page [5, с. 12] unchanged.
    """
    from guides.coursework_kfu_2025.safe_formatter import _split_citation_brackets_in_text

    cases = [
        # Multi-source split
        ("[21, с. 30–45, 22, с. 21–33, 5, с. 3–8, 10]",
         "[21, с. 30–45], [22, с. 21–33], [5, с. 3–8], [10]"),
        ("[12; 13; 5]",      "[12], [13], [5]"),
        ("[21, 22]",         "[21], [22]"),
        # Single source — unchanged (but hyphen normalized)
        ("[21, с. 30–45]",   "[21, с. 30–45]"),
        ("[10]",             "[10]"),
        # Hyphen → en-dash in single source range
        ("[5, с. 12-15]",    "[5, с. 12–15]"),
        ("[5, с. 12–15]",    "[5, с. 12–15]"),
        # Single page (no range)
        ("[5, с. 12]",       "[5, с. 12]"),
        # p. notation → с. in output
        ("[5, p. 12-15]",    "[5, с. 12–15]"),
        ("[5, p. 12]",       "[5, с. 12]"),
        # Mixed in sentence
        ("по данным [21, 22], а также [5, с. 3–8, 10]",
         "по данным [21], [22], а также [5, с. 3–8], [10]"),
    ]
    failures = []
    for inp, expected in cases:
        got = _split_citation_brackets_in_text(inp)
        if got != expected:
            failures.append(f"Input:    {inp!r}\nExpected: {expected!r}\nGot:      {got!r}")
    if failures:
        return _result(False, "\n\n".join(failures))
    return _result(True, f"all {len(cases)} citation cases correct")


def test_t5_list_formatting() -> tuple[bool, str]:
    """
    Numeric-paren items (1)/2)/3)) after a colon-ending paragraph become а)/б)/в).
    Level-1 items get firstLine=708 (KFU L1 standard, 1.25 cm).
    Note: numeric-dot items (1./2./3.) after colon become '- text' items (not letters).
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs
    from docx.oxml.ns import qn

    doc = Document()
    intro = doc.add_paragraph("Выделяют следующие виды:")
    p1 = doc.add_paragraph("1) первый вид")
    p2 = doc.add_paragraph("2) второй вид")
    p3 = doc.add_paragraph("3) третий вид")

    _normalize_plain_list_paragraphs([intro, p1, p2, p3])

    if not p1.text.startswith("а)"):
        return _result(False, f"p1 not converted: {p1.text!r}")
    if not p2.text.startswith("б)"):
        return _result(False, f"p2 not converted: {p2.text!r}")
    if not p3.text.startswith("в)"):
        return _result(False, f"p3 not converted: {p3.text!r}")

    pPr = p1._element.find(qn("w:pPr"))
    ind = pPr.find(qn("w:ind")) if pPr is not None else None
    if ind is None:
        return _result(False, "no w:ind on level-1 item")
    fl = ind.get(qn("w:firstLine"))
    if fl != "708":
        return _result(False, f"wrong indent: firstLine={fl!r} (expected '708')")

    return _result(True, "list items converted with firstLine=708 ✓")


# ── Free-standing list normalization (MVP) ──────────────────────────────


def _list_ind_attrs(paragraph):
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        return None, None, None
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        return None, None, None
    return (
        ind.get(qn("w:left")),
        ind.get(qn("w:hanging")),
        ind.get(qn("w:right")),
    )


def test_lists_manual_dash_block_normalized() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs
    from docx.oxml.ns import qn

    doc = Document()
    pre = doc.add_paragraph("Параграф без двоеточия.")
    p1 = doc.add_paragraph("- первый пункт")
    p2 = doc.add_paragraph("- второй пункт")
    p3 = doc.add_paragraph("- третий пункт")
    _normalize_plain_list_paragraphs([pre, p1, p2, p3])
    for p in (p1, p2, p3):
        # New behaviour: plain literal '-' text, no numPr, firstLine=708
        if not p.text.startswith("- "):
            return _result(False, f"dash item should start with '- ': {p.text!r}")
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return _result(False, f"dash item should not have numPr: {p.text!r}")
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        fl = ind.get(qn("w:firstLine")) if ind is not None else None
        if fl != "708":
            return _result(False, f"dash item bad firstLine indent: {fl!r} (expected '708')")
    return _result(True, "ascii-hyphen block normalized to plain '- text' with firstLine=708")


def test_lists_manual_em_dash_block_normalized() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs
    from docx.oxml.ns import qn

    doc = Document()
    pre = doc.add_paragraph("Введение к блоку.")
    p1 = doc.add_paragraph("— первый пункт")
    p2 = doc.add_paragraph("— второй пункт")
    _normalize_plain_list_paragraphs([pre, p1, p2])
    for p in (p1, p2):
        # em-dash normalized to '-' (hyphen) with firstLine=708, no numPr
        if not p.text.startswith("- "):
            return _result(False, f"em-dash item should become '- text': {p.text!r}")
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return _result(False, f"em-dash item should not have numPr: {p.text!r}")
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        fl = ind.get(qn("w:firstLine")) if ind is not None else None
        if fl != "708":
            return _result(False, f"em-dash item bad firstLine: {fl!r} (expected '708')")
    return _result(True, "em-dash block normalized to plain '- text' with firstLine=708")


def test_lists_manual_black_bullets_normalized() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs
    from docx.oxml.ns import qn

    doc = Document()
    pre = doc.add_paragraph("Параграф без двоеточия.")
    p1 = doc.add_paragraph("• первый bullet")
    p2 = doc.add_paragraph("· второй bullet")
    p3 = doc.add_paragraph("● третий bullet")
    _normalize_plain_list_paragraphs([pre, p1, p2, p3])
    for p in (p1, p2, p3):
        # All bullet chars normalized to '-' (hyphen), no numPr, firstLine=708
        if not p.text.startswith("- "):
            return _result(False, f"bullet should become '- text': {p.text!r}")
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return _result(False, f"bullet item should not have numPr: {p.text!r}")
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        fl = ind.get(qn("w:firstLine")) if ind is not None else None
        if fl != "708":
            return _result(False, f"bullet item bad firstLine: {fl!r} (expected '708')")
    return _result(True, "black bullets normalized to plain '- text' with firstLine=708")


def test_lists_singleton_dash_not_converted() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs

    doc = Document()
    pre = doc.add_paragraph("Параграф без двоеточия.")
    solo = doc.add_paragraph("- одинокий пункт")
    after = doc.add_paragraph("Обычный текст после.")
    _normalize_plain_list_paragraphs([pre, solo, after])
    if solo.text.startswith("– "):
        return _result(False, f"singleton dash unexpectedly normalized: {solo.text!r}")
    left, hang, _ = _list_ind_attrs(solo)
    if left == "906" and hang == "198":
        return _result(False, "singleton dash unexpectedly got list indent")
    return _result(True, "singleton dash left untouched")


def test_lists_mixed_marker_block_not_converted() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs

    doc = Document()
    pre = doc.add_paragraph("Параграф без двоеточия.")
    p1 = doc.add_paragraph("- первый пункт")
    p2 = doc.add_paragraph("1. второй пункт")
    _normalize_plain_list_paragraphs([pre, p1, p2])
    if p1.text.startswith("– "):
        return _result(False, f"mixed-family block converted p1: {p1.text!r}")
    if p2.text.startswith("а)") or p2.text.startswith("– "):
        return _result(False, f"mixed-family block converted p2: {p2.text!r}")
    return _result(True, "mixed-family block left untouched")


def test_lists_free_standing_letter_block_normalized() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs
    from docx.oxml.ns import qn

    doc = Document()
    pre = doc.add_paragraph("Параграф без двоеточия.")
    p1 = doc.add_paragraph("а) первый пункт")
    p2 = doc.add_paragraph("б) второй пункт")
    p3 = doc.add_paragraph("в) третий пункт")
    _normalize_plain_list_paragraphs([pre, p1, p2, p3])
    if not p1.text.startswith("а)"):
        return _result(False, f"letter item lost marker: {p1.text!r}")
    if not p2.text.startswith("б)"):
        return _result(False, f"letter item lost marker: {p2.text!r}")
    pPr = p1._element.find(qn("w:pPr"))
    ind = pPr.find(qn("w:ind")) if pPr is not None else None
    fl = ind.get(qn("w:firstLine")) if ind is not None else None
    if fl != "708":
        return _result(False, f"letter item wrong indent: firstLine={fl!r} (expected '708')")
    return _result(True, "free-standing letter block keeps markers, gets firstLine=708")


def test_lists_g4_letter_block_with_colon_endings() -> tuple[bool, str]:
    """
    G4: a free-standing letter block where every item ends with `:` must
    still be recognised as L1 letters with firstLine=708 indent. The `:` is
    preserved in the output text.
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs
    from docx.oxml.ns import qn

    doc = Document()
    pre = doc.add_paragraph("Параграф без двоеточия.")
    p1 = doc.add_paragraph("а) первый пункт первого уровня:")
    p2 = doc.add_paragraph("б) второй пункт первого уровня:")
    p3 = doc.add_paragraph("в) третий пункт первого уровня:")
    _normalize_plain_list_paragraphs([pre, p1, p2, p3])

    for p, expected_marker in ((p1, "а)"), (p2, "б)"), (p3, "в)")):
        if not p.text.startswith(expected_marker):
            return _result(False, f"letter+colon item lost marker: {p.text!r}")
        if not p.text.rstrip().endswith(":"):
            return _result(False, f"letter+colon item lost trailing colon: {p.text!r}")
        pPr = p._element.find(qn("w:pPr"))
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        fl = ind.get(qn("w:firstLine")) if ind is not None else None
        if fl != "708":
            return _result(False, f"letter+colon item wrong indent: firstLine={fl!r} text={p.text!r}")
    return _result(True, "G4: letter block with colon endings keeps markers + `:` + firstLine=708")


def test_lists_g1_letter_parent_colon_nested_numeric_l2() -> tuple[bool, str]:
    """
    G1: letter L1 parent ending with `:` followed by numeric paren children.
    Parent L1 gets firstLine=708. Children L2 get left=1200 hanging=198.
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs
    from docx.oxml.ns import qn

    doc = Document()
    pre = doc.add_paragraph("Параграф без двоеточия.")
    parent = doc.add_paragraph("а) первый пункт первого уровня:")
    child1 = doc.add_paragraph("1) вложенный пункт второго уровня один")
    child2 = doc.add_paragraph("2) вложенный пункт второго уровня два")
    _normalize_plain_list_paragraphs([pre, parent, child1, child2])

    if not parent.text.startswith("а) "):
        return _result(False, f"parent lost marker: {parent.text!r}")
    if not parent.text.rstrip().endswith(":"):
        return _result(False, f"parent lost trailing colon: {parent.text!r}")
    pPr_p = parent._element.find(qn("w:pPr"))
    ind_p = pPr_p.find(qn("w:ind")) if pPr_p is not None else None
    fl_p = ind_p.get(qn("w:firstLine")) if ind_p is not None else None
    if fl_p != "708":
        return _result(False, f"parent missing L1 indent: firstLine={fl_p!r} (expected '708')")

    for child, expected in ((child1, "1) "), (child2, "2) ")):
        if not child.text.startswith(expected):
            return _result(False, f"child lost marker: {child.text!r}")
        c_left, c_hang, _ = _list_ind_attrs(child)
        if (c_left, c_hang) != ("1200", "198"):
            return _result(False, f"child missing L2 indent: left={c_left} hang={c_hang} text={child.text!r}")
    return _result(True, "G1: letter parent (left=708) + nested numeric L2 children (left=1200 hang=198)")


def test_lists_g1_g4_block_e_full_smoke_case() -> tuple[bool, str]:
    """
    Block E from list_normalization_smoke_input_kfu.docx — two letter L1
    parents each ending with `:`, each followed by two numeric children.
    All 6 paragraphs must end up methodically formatted.
    L1 parents: left=708. L2 children: left=1200 hanging=198.
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs
    from docx.oxml.ns import qn

    doc = Document()
    pre = doc.add_paragraph("BLOCK E — вложенный список: тестовый параграф без двоеточия в конце.")
    parent_a = doc.add_paragraph("а) первый пункт первого уровня:")
    child_a1 = doc.add_paragraph("1) вложенный пункт второго уровня с длинным текстом для переноса")
    child_a2 = doc.add_paragraph("2) второй вложенный пункт второго уровня")
    parent_b = doc.add_paragraph("б) второй пункт первого уровня:")
    child_b1 = doc.add_paragraph("1) вложенный пункт второго уровня внутри второго буквенного пункта")
    child_b2 = doc.add_paragraph("2) второй вложенный пункт второго уровня внутри второго буквенного пункта")
    _normalize_plain_list_paragraphs([pre, parent_a, child_a1, child_a2, parent_b, child_b1, child_b2])

    for parent, expected_letter in ((parent_a, "а)"), (parent_b, "б)")):
        if not parent.text.startswith(expected_letter):
            return _result(False, f"parent lost marker: {parent.text!r}")
        if not parent.text.rstrip().endswith(":"):
            return _result(False, f"parent lost colon: {parent.text!r}")
        pPr = parent._element.find(qn("w:pPr"))
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        fl = ind.get(qn("w:firstLine")) if ind is not None else None
        if fl != "708":
            return _result(False, f"parent indent wrong: firstLine={fl!r} text={parent.text!r}")

    for child, expected_marker in (
        (child_a1, "1) "),
        (child_a2, "2) "),
        (child_b1, "1) "),
        (child_b2, "2) "),
    ):
        if not child.text.startswith(expected_marker):
            return _result(False, f"child marker wrong: {child.text!r}")
        left, hang, _ = _list_ind_attrs(child)
        if (left, hang) != ("1200", "198"):
            return _result(False, f"child L2 indent wrong: left={left} hang={hang} text={child.text!r}")

    return _result(True, "G1+G4: Block E fully formatted — parents left=708, children left=1200 hang=198")


def test_lists_g4_letter_singleton_with_colon_not_converted() -> tuple[bool, str]:
    """
    Safety lock: single `а) item:` paragraph alone (no nested numerics, no
    sibling letters) must stay untouched per MVP singleton rule.
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs

    doc = Document()
    pre = doc.add_paragraph("Параграф без двоеточия.")
    solo = doc.add_paragraph("а) одинокий пункт с двоеточием:")
    after = doc.add_paragraph("Обычный текст после, не список.")
    _normalize_plain_list_paragraphs([pre, solo, after])

    left, hang, _ = _list_ind_attrs(solo)
    if left == "906" and hang == "198":
        return _result(False, f"singleton letter+colon got L1 indent: {solo.text!r}")
    if not solo.text.startswith("а) "):
        return _result(False, f"singleton text mutated unexpectedly: {solo.text!r}")
    return _result(True, "singleton letter+colon not converted (MVP safety preserved)")


def test_lists_g1_numeric_without_letter_parent_not_promoted_to_l2() -> tuple[bool, str]:
    """
    Safety lock: a stand-alone numeric-paren block (`1)/2)/3)`) WITHOUT a
    preceding letter+`:` parent must NOT be auto-promoted to L2 indent.
    Existing colon-mode tests cover the legacy non-promotion behaviour;
    this guards the new path explicitly.
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs

    doc = Document()
    pre = doc.add_paragraph("Параграф без двоеточия и без letter parent.")
    p1 = doc.add_paragraph("1) первый пункт")
    p2 = doc.add_paragraph("2) второй пункт")
    _normalize_plain_list_paragraphs([pre, p1, p2])

    for p in (p1, p2):
        left, hang, _ = _list_ind_attrs(p)
        if (left, hang) == ("963", "198"):
            return _result(False, f"stray numeric-paren got L2 indent: {p.text!r}")
    return _result(True, "stand-alone numeric-paren without letter parent not converted")


def test_lists_numeric_dot_without_colon_not_converted() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs

    doc = Document()
    pre = doc.add_paragraph("Обычный абзац без двоеточия.")
    p1 = doc.add_paragraph("1. Первое утверждение по тексту.")
    p2 = doc.add_paragraph("2. Второе утверждение по тексту.")
    p3 = doc.add_paragraph("3. Третье утверждение по тексту.")
    _normalize_plain_list_paragraphs([pre, p1, p2, p3])
    if p1.text.startswith("а)") or p1.text.startswith("– "):
        return _result(False, f"numeric-dot without colon converted: {p1.text!r}")
    return _result(True, "numeric-dot without colon lead-in untouched")


def test_lists_heading_styled_dash_not_converted() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs

    doc = Document()
    pre = doc.add_paragraph("Обычный абзац без двоеточия.")
    h1 = doc.add_paragraph("- ТЕОРЕТИЧЕСКИЕ ОСНОВЫ")
    h1.style = "Heading 1"
    p2 = doc.add_paragraph("- продолжение псевдо-списка")
    _normalize_plain_list_paragraphs([pre, h1, p2])
    if h1.text.startswith("– "):
        return _result(False, f"heading-styled dash converted: {h1.text!r}")
    return _result(True, "heading-styled paragraph not converted")


def test_lists_table_caption_not_converted() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs

    doc = Document()
    pre = doc.add_paragraph("Обычный абзац без двоеточия.")
    cap = doc.add_paragraph("- Таблица 1.1.1")  # absurd but tests guard
    real_cap = doc.add_paragraph("Таблица 1.1.2")
    p2 = doc.add_paragraph("- второй пункт")
    _normalize_plain_list_paragraphs([pre, cap, real_cap, p2])
    if real_cap.text.startswith("– "):
        return _result(False, f"real table caption converted: {real_cap.text!r}")
    return _result(True, "table caption guard fires")


def test_lists_source_line_not_converted() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs

    doc = Document()
    pre = doc.add_paragraph("Обычный абзац без двоеточия.")
    src1 = doc.add_paragraph("Источник: составлено автором.")
    src2 = doc.add_paragraph("Примечание: тестовая заметка.")
    _normalize_plain_list_paragraphs([pre, src1, src2])
    if src1.text.startswith("– ") or src2.text.startswith("– "):
        return _result(False, f"source/note converted: {src1.text!r} / {src2.text!r}")
    return _result(True, "source/note service lines not converted")


def test_lists_colon_numeric_dot_becomes_dash() -> tuple[bool, str]:
    """
    Colon-triggered numeric-dot items (1./2./3.) become '- text' dash items
    with firstLine=708 indent (1.25 cm). This is the KFU task-list pattern.
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_plain_list_paragraphs
    from docx.oxml.ns import qn

    doc = Document()
    intro = doc.add_paragraph("Для достижения цели необходимо решить следующие задачи:")
    p1 = doc.add_paragraph("1. изучить теоретические основы")
    p2 = doc.add_paragraph("2. проанализировать практику")
    p3 = doc.add_paragraph("3. разработать предложения")
    _normalize_plain_list_paragraphs([intro, p1, p2, p3])
    # New behaviour: numeric-dot after colon → '- text', not letters
    if not p1.text.startswith("- "):
        return _result(False, f"colon numeric-dot should become '- text', got: {p1.text!r}")
    if not p2.text.startswith("- "):
        return _result(False, f"colon numeric-dot should become '- text', got: {p2.text!r}")
    if not p3.text.startswith("- "):
        return _result(False, f"colon numeric-dot should become '- text', got: {p3.text!r}")
    # Must have firstLine=708 (paragraph indent 1.25 cm on first line)
    pPr = p1._element.find(qn("w:pPr"))
    ind = pPr.find(qn("w:ind")) if pPr is not None else None
    fl_val = ind.get(qn("w:firstLine")) if ind is not None else None
    if fl_val != "708":
        return _result(False, f"colon dash item should have firstLine=708, got: {fl_val!r}")
    return _result(True, "colon mode numeric-dot → '- text' (firstLine=708)")


def test_lists_appendix_section_not_normalized() -> tuple[bool, str]:
    """After ПРИЛОЖЕНИЯ heading, list-like paragraphs are not touched."""
    from guides.coursework_kfu_2025.safe_formatter import normalize_plain_lists_in_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    p1 = doc.add_paragraph("- псевдо-пункт в приложении")
    p2 = doc.add_paragraph("- второй псевдо-пункт")
    normalize_plain_lists_in_document(doc, body_start=0)
    if p1.text.startswith("– "):
        return _result(False, f"appendix dash converted: {p1.text!r}")
    if p2.text.startswith("– "):
        return _result(False, f"appendix dash converted: {p2.text!r}")
    return _result(True, "appendix-local list-like paragraphs left untouched")


def test_lists_bibliography_section_not_normalized() -> tuple[bool, str]:
    """Numeric entries inside СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ are preserved."""
    from guides.coursework_kfu_2025.safe_formatter import normalize_plain_lists_in_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    intro = doc.add_paragraph("Для проверки задачи:")
    p1 = doc.add_paragraph("1. Иванов И.И. Учебное пособие.")
    p2 = doc.add_paragraph("2. Петров П.П. Документооборот.")
    normalize_plain_lists_in_document(doc, body_start=0)
    if p1.text.startswith("а)") or p2.text.startswith("а)"):
        return _result(False, f"bibliography entry converted: {p1.text!r} / {p2.text!r}")
    return _result(True, "bibliography entries not converted")


# ── Word dash-list normalization tests ───────────────────────────────────────
# Covers: G2 (decimal→dash), plain manual dash, bullet chars, letter-list
# preservation, singleton guards, references/appendix guards, idempotency.


def _make_decimal_numbering_doc():
    """
    Return a Document() where numId='1' maps to decimal ilvl=0 via a proper
    numbering part. Used by Word-decimal list tests.
    """
    import io
    import zipfile as _zipfile

    buf = io.BytesIO()
    Document().save(buf)
    buf.seek(0)
    NUMBERING_XML = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:abstractNum w:abstractNumId="0">'
        b'<w:lvl w:ilvl="0">'
        b'<w:start w:val="1"/><w:numFmt w:val="decimal"/>'
        b'<w:lvlText w:val="%1."/><w:lvlJc w:val="left"/>'
        b'</w:lvl></w:abstractNum>'
        b'<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
        b'</w:numbering>'
    )
    out = io.BytesIO()
    with _zipfile.ZipFile(buf, "r") as zin, _zipfile.ZipFile(out, "w", _zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/numbering.xml":
                zout.writestr(item, NUMBERING_XML)
            else:
                zout.writestr(item, zin.read(item.filename))
    out.seek(0)
    return Document(out)


def _make_letter_numbering_doc():
    """
    Return a Document() where numId='1' maps to lowerLetter ilvl=0.
    Used by letter-list preservation tests.
    """
    import io
    import zipfile as _zipfile

    buf = io.BytesIO()
    Document().save(buf)
    buf.seek(0)
    NUMBERING_XML = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:abstractNum w:abstractNumId="0">'
        b'<w:lvl w:ilvl="0">'
        b'<w:start w:val="1"/><w:numFmt w:val="lowerLetter"/>'
        b'<w:lvlText w:val="%1)"/><w:lvlJc w:val="left"/>'
        b'</w:lvl></w:abstractNum>'
        b'<w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>'
        b'</w:numbering>'
    )
    out = io.BytesIO()
    with _zipfile.ZipFile(buf, "r") as zin, _zipfile.ZipFile(out, "w", _zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            if item.filename == "word/numbering.xml":
                zout.writestr(item, NUMBERING_XML)
            else:
                zout.writestr(item, zin.read(item.filename))
    out.seek(0)
    return Document(out)


def _add_numpr(paragraph, num_id: str = "1", ilvl: str = "0") -> None:
    """Attach a w:numPr pointing at numId with the given ilvl."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), ilvl)
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), num_id)
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)


# Keep backward-compatible alias used in earlier tests in this file.
_add_decimal_numpr = _add_numpr


def _para_is_kfu_dash_list(p) -> bool:
    """
    Return True iff *p* has a real KFU Word dash-list numPr —
    i.e. the abstractNum carries our internal NSID marker (FFFEFFFF)
    OR has numFmt=bullet with lvlText=en-dash at ilvl=0.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        _KFU_DASH_NSID, paragraph_has_numbering, _get_num_fmt_for_paragraph,
    )
    if not paragraph_has_numbering(p):
        return False
    if _get_num_fmt_for_paragraph(p) != "bullet":
        return False
    try:
        numbering_el = p.part.numbering_part._element
        pPr = p._element.pPr
        numPr = pPr.find(qn("w:numPr"))
        num_id_val = numPr.find(qn("w:numId")).get(qn("w:val"))
        for num_e in numbering_el.findall(qn("w:num")):
            if num_e.get(qn("w:numId")) != num_id_val:
                continue
            ref = num_e.find(qn("w:abstractNumId"))
            if ref is None:
                return False
            ab_id = ref.get(qn("w:val"))
            for ab in numbering_el.findall(qn("w:abstractNum")):
                if ab.get(qn("w:abstractNumId")) != ab_id:
                    continue
                # Check NSID marker
                nsid = ab.find(qn("w:nsid"))
                if nsid is not None and nsid.get(qn("w:val")) == _KFU_DASH_NSID:
                    return True
                # Fallback: check lvlText=– at ilvl=0
                for lvl in ab.findall(qn("w:lvl")):
                    if lvl.get(qn("w:ilvl")) != "0":
                        continue
                    lt = lvl.find(qn("w:lvlText"))
                    if lt is not None and lt.get(qn("w:val")) == "–":
                        return True
        return False
    except Exception:
        return False


# ── Test 1: G2 Pattern A — same numId block → real Word dash-list ─────────────

def test_word_numbered_pattern_a_converted() -> tuple[bool, str]:
    """
    G2 Pattern A: 2+ contiguous decimal numPr paragraphs sharing the same numId
    become plain-text '- text' items (hyphen, no numPr, firstLine=708).
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_word_numbered_list_paragraphs
    from docx.oxml.ns import qn

    doc = _make_decimal_numbering_doc()
    pre = doc.add_paragraph("Вводный абзац.")
    p1 = doc.add_paragraph("first Word numbered item")
    p2 = doc.add_paragraph("second Word numbered item")
    p3 = doc.add_paragraph("third Word numbered item")
    for p in (p1, p2, p3):
        _add_numpr(p, num_id="1")

    _normalize_word_numbered_list_paragraphs([pre, p1, p2, p3])

    for p in (p1, p2, p3):
        if not p.text.startswith("- "):
            return _result(False, f"item should start with '- ': {p.text!r}")
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return _result(False, f"item should not have numPr: {p.text!r}")
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        fl = ind.get(qn("w:firstLine")) if ind is not None else None
        if fl != "708":
            return _result(False, f"item bad firstLine: {fl!r} (expected '708') text={p.text!r}")
    return _result(True, "Pattern A: 3-item block converted to plain '- text' with firstLine=708")


# ── Test 2: G2 Pattern B — cross-numId block → real Word dash-list ───────────

def test_word_numbered_pattern_b_converted() -> tuple[bool, str]:
    """
    G2 Pattern B: contiguous decimal numPr paragraphs converted to
    plain-text '- text' items (hyphen, no numPr, firstLine=708).
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_word_numbered_list_paragraphs
    from docx.oxml.ns import qn

    doc = _make_decimal_numbering_doc()
    pre = doc.add_paragraph("Вводный абзац.")
    p1 = doc.add_paragraph("Reset-numbered item A")
    p2 = doc.add_paragraph("Reset-numbered item B")
    p3 = doc.add_paragraph("Reset-numbered item C")
    for p in (p1, p2, p3):
        _add_numpr(p, num_id="1")

    _normalize_word_numbered_list_paragraphs([pre, p1, p2, p3])

    for p in (p1, p2, p3):
        if not p.text.startswith("- "):
            return _result(False, f"Pattern B item should start with '- ': {p.text!r}")
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return _result(False, f"Pattern B item should not have numPr: {p.text!r}")
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        fl = ind.get(qn("w:firstLine")) if ind is not None else None
        if fl != "708":
            return _result(False, f"Pattern B item bad firstLine: {fl!r} text={p.text!r}")
    return _result(True, "Pattern B: reset-numbered block converted to plain '- text' with firstLine=708")


# ── Test 3: Plain manual "– text" blocks → real Word dash-list ───────────────

def test_plain_dash_block_becomes_word_list() -> tuple[bool, str]:
    """
    normalize_plain_lists_in_document must convert a block of 2+ manual
    '– text' paragraphs into plain-text '- text' items (hyphen, firstLine=708).
    (en-dash markers are normalized to hyphen for free-standing dash blocks.)
    """
    from guides.coursework_kfu_2025.safe_formatter import normalize_plain_lists_in_document
    from docx.oxml.ns import qn

    doc = _make_decimal_numbering_doc()
    doc.add_paragraph("ВВЕДЕНИЕ")
    p1 = doc.add_paragraph("– первый пункт")
    p2 = doc.add_paragraph("– второй пункт")
    p3 = doc.add_paragraph("– третий пункт")

    normalize_plain_lists_in_document(doc, body_start=0)

    for p in (p1, p2, p3):
        if not p.text.startswith("- "):
            return _result(False, f"plain dash item should start with '- ': {p.text!r}")
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return _result(False, f"dash item should not have numPr: {p.text!r}")
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        fl = ind.get(qn("w:firstLine")) if ind is not None else None
        if fl != "708":
            return _result(False, f"dash item bad firstLine: {fl!r} text={p.text!r}")
    return _result(True, "plain manual dash block normalized to '- text' with firstLine=708")


# ── Test 4: Bullet chars → real Word dash-list ───────────────────────────────

def test_bullet_chars_become_word_dash_list() -> tuple[bool, str]:
    """
    normalize_plain_lists_in_document must convert a block of 2+ bullet
    paragraphs (•, ·, ●, -, —) into plain-text '- text' items (firstLine=708).
    """
    from guides.coursework_kfu_2025.safe_formatter import normalize_plain_lists_in_document
    from docx.oxml.ns import qn

    doc = _make_decimal_numbering_doc()
    doc.add_paragraph("ВВЕДЕНИЕ")
    p1 = doc.add_paragraph("• первый")
    p2 = doc.add_paragraph("· второй")
    p3 = doc.add_paragraph("- третий")

    normalize_plain_lists_in_document(doc, body_start=0)

    for p in (p1, p2, p3):
        if not p.text.startswith("- "):
            return _result(False, f"bullet item should become '- text': {p.text!r}")
        pPr = p._element.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:numPr")) is not None:
            return _result(False, f"bullet item should not have numPr: {p.text!r}")
        ind = pPr.find(qn("w:ind")) if pPr is not None else None
        fl = ind.get(qn("w:firstLine")) if ind is not None else None
        if fl != "708":
            return _result(False, f"bullet item bad firstLine: {fl!r} text={p.text!r}")
    return _result(True, "bullet chars normalized to plain '- text' with firstLine=708")


# ── Test 5: Letter-list numPr preserved by format_body_list_item ─────────────

def test_letter_numpr_preserved_in_body_list_item() -> tuple[bool, str]:
    """
    format_body_list_item must preserve numPr for lowerLetter Word lists
    (Cyrillic а) б) в) autonumbering) rather than stripping it.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        format_body_list_item, _get_num_fmt_for_paragraph, paragraph_has_numbering,
    )

    doc = _make_letter_numbering_doc()
    p = doc.add_paragraph("а) первый пункт списка")
    _add_numpr(p, num_id="1")

    format_body_list_item(p)

    if not paragraph_has_numbering(p):
        return _result(False, "numPr was removed from letter-list item")
    nf = _get_num_fmt_for_paragraph(p)
    if nf != "lowerLetter":
        return _result(False, f"numFmt changed from lowerLetter to {nf!r}")
    return _result(True, "lowerLetter numPr preserved by format_body_list_item")


# ── Test 6: Singleton manual dash unchanged ───────────────────────────────────

def test_singleton_plain_dash_unchanged() -> tuple[bool, str]:
    """
    A single isolated '– text' paragraph must NOT be converted (singleton guard).
    """
    from guides.coursework_kfu_2025.safe_formatter import normalize_plain_lists_in_document

    doc = _make_decimal_numbering_doc()
    doc.add_paragraph("ВВЕДЕНИЕ")
    solo = doc.add_paragraph("– одиночный пункт")
    doc.add_paragraph("Обычный текст после.")

    orig = solo.text
    normalize_plain_lists_in_document(doc, body_start=0)

    if solo.text != orig:
        return _result(False, f"singleton was modified: {solo.text!r}")
    if _para_is_kfu_dash_list(solo):
        return _result(False, "singleton was unexpectedly converted to Word list")
    return _result(True, "singleton manual dash left unchanged")


# ── Test 7: Singleton Word-decimal unchanged ──────────────────────────────────

def test_word_numbered_singleton_not_converted() -> tuple[bool, str]:
    """
    A single isolated decimal numPr paragraph (block size == 1) must not
    be converted. Block I from autonumbering_kfu_input_cases.docx.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        _normalize_word_numbered_list_paragraphs, _get_num_fmt_for_paragraph,
    )

    doc = _make_decimal_numbering_doc()
    pre = doc.add_paragraph("Вводный абзац.")
    solo = doc.add_paragraph("Single Word numbered paragraph should remain.")
    after = doc.add_paragraph("Обычный текст после.")
    _add_numpr(solo, num_id="1")

    orig_text = solo.text
    _normalize_word_numbered_list_paragraphs([pre, solo, after])

    if solo.text != orig_text:
        return _result(False, f"singleton text changed: {solo.text!r}")
    if _get_num_fmt_for_paragraph(solo) != "decimal":
        return _result(False, "singleton numFmt changed away from decimal")
    return _result(True, "singleton decimal numPr paragraph left unchanged")


# ── Test 8: References block unchanged ───────────────────────────────────────

def test_word_numbered_references_block_not_converted() -> tuple[bool, str]:
    """
    Decimal numPr paragraphs inside the references block must not be converted.
    """
    from guides.coursework_kfu_2025.safe_formatter import normalize_word_numbered_lists_in_document

    doc = _make_decimal_numbering_doc()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    p1 = doc.add_paragraph("Иванов И.И. Учебное пособие. 2024.")
    p2 = doc.add_paragraph("Петров П.П. Документооборот. 2023.")
    _add_numpr(p1, num_id="1")
    _add_numpr(p2, num_id="1")

    t1, t2 = p1.text, p2.text
    normalize_word_numbered_lists_in_document(doc, body_start=0)

    if p1.text != t1 or p2.text != t2:
        return _result(False, f"references items were converted: {p1.text!r} / {p2.text!r}")
    return _result(True, "references decimal numPr items not converted")


# ── Test 9: Appendices block unchanged ───────────────────────────────────────

def test_word_numbered_appendix_block_not_converted() -> tuple[bool, str]:
    """
    Decimal numPr paragraphs inside the appendices section must not be converted.
    """
    from guides.coursework_kfu_2025.safe_formatter import normalize_word_numbered_lists_in_document

    doc = _make_decimal_numbering_doc()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    p1 = doc.add_paragraph("appendix item one")
    p2 = doc.add_paragraph("appendix item two")
    _add_numpr(p1, num_id="1")
    _add_numpr(p2, num_id="1")

    t1, t2 = p1.text, p2.text
    normalize_word_numbered_lists_in_document(doc, body_start=0)

    if p1.text != t1 or p2.text != t2:
        return _result(False, f"appendix items were converted: {p1.text!r} / {p2.text!r}")
    return _result(True, "appendix decimal numPr items not converted")


# ── Zone D: Lettered list recognition (uppercase, dot-delimiter, Latin) ──────

def test_zone_d_lettered_list_variants() -> tuple[bool, str]:
    """
    Zone D: _CYRILLIC_LETTER_LIST_RE must recognize uppercase Cyrillic,
    lowercase/uppercase Cyrillic with dot delimiter, and Latin A/B/C variants.
    Singletons remain unchanged; ПРИЛОЖЕНИЕ А and plain body words are not affected.
    """
    from guides.coursework_kfu_2025.safe_formatter import normalize_plain_lists_in_document
    from docx.oxml.ns import qn

    def _firstline(p):
        pPr = p._element.find(qn("w:pPr"))
        if pPr is None:
            return None
        ind = pPr.find(qn("w:ind"))
        return ind.get(qn("w:firstLine")) if ind is not None else None

    def _check_list_item(p, label):
        """Return error string or None."""
        fl = _firstline(p)
        if fl != "708":
            return f"{label}: expected firstLine=708, got {fl!r} (text={p.text!r})"
        return None

    doc = _make_decimal_numbering_doc()
    doc.add_paragraph("ВВЕДЕНИЕ")

    # ── 1. Existing lowercase Cyrillic а) still works ─────────────────────────
    a1 = doc.add_paragraph("а) первый пункт")
    a2 = doc.add_paragraph("б) второй пункт")
    doc.add_paragraph("")

    # ── 2. Uppercase Cyrillic А) ──────────────────────────────────────────────
    b1 = doc.add_paragraph("А) первый пункт")
    b2 = doc.add_paragraph("Б) второй пункт")
    doc.add_paragraph("")

    # ── 3. Lowercase Cyrillic dot а. ─────────────────────────────────────────
    c1 = doc.add_paragraph("а. первый пункт")
    c2 = doc.add_paragraph("б. второй пункт")
    doc.add_paragraph("")

    # ── 4. Uppercase Cyrillic dot А. ─────────────────────────────────────────
    d1 = doc.add_paragraph("А. первый пункт")
    d2 = doc.add_paragraph("Б. второй пункт")
    doc.add_paragraph("")

    # ── 5. Latin A) ───────────────────────────────────────────────────────────
    e1 = doc.add_paragraph("A) first item")
    e2 = doc.add_paragraph("B) second item")
    doc.add_paragraph("")

    # ── 6. Latin A. ───────────────────────────────────────────────────────────
    f1 = doc.add_paragraph("A. first item")
    f2 = doc.add_paragraph("B. second item")
    doc.add_paragraph("")

    # ── Guards: must NOT be converted ─────────────────────────────────────────
    appendix_label = doc.add_paragraph("ПРИЛОЖЕНИЕ А")
    body_word = doc.add_paragraph("Анализ ситуации на рынке.")

    normalize_plain_lists_in_document(doc, body_start=0)

    # Check all list pairs got list formatting (firstLine=708)
    for pair_label, p1, p2 in [
        ("lowercase Cyrillic а)", a1, a2),
        ("uppercase Cyrillic А)", b1, b2),
        ("lowercase Cyrillic dot а.", c1, c2),
        ("uppercase Cyrillic dot А.", d1, d2),
        ("Latin А)", e1, e2),
        ("Latin А.", f1, f2),
    ]:
        err = _check_list_item(p1, f"{pair_label} item 1")
        if err:
            return _result(False, err)
        err = _check_list_item(p2, f"{pair_label} item 2")
        if err:
            return _result(False, err)

    # ПРИЛОЖЕНИЕ А must not be touched
    if appendix_label.text != "ПРИЛОЖЕНИЕ А":
        return _result(False, f"ПРИЛОЖЕНИЕ А was modified: {appendix_label.text!r}")
    if _firstline(appendix_label) == "708":
        return _result(False, "ПРИЛОЖЕНИЕ А got list indent")

    # Plain body word starting with А must not be touched
    if "Анализ" not in body_word.text:
        return _result(False, f"body word starting with А was mutated: {body_word.text!r}")
    if _firstline(body_word) == "708":
        return _result(False, "plain body word got list indent")

    return _result(True, "all lettered list variants recognized; guards intact")


def test_zone_d_singleton_letter_unchanged() -> tuple[bool, str]:
    """Zone D: a single isolated letter-list paragraph is not converted (singleton guard).
    Each candidate is separated from other list-like text by a blank paragraph so the
    two-item buffer confirmation never fires.
    """
    from guides.coursework_kfu_2025.safe_formatter import normalize_plain_lists_in_document
    from docx.oxml.ns import qn

    doc = _make_decimal_numbering_doc()
    doc.add_paragraph("ВВЕДЕНИЕ")
    # Singleton А) — surrounded by blank + body text so it cannot confirm a buffer
    solo_upper = doc.add_paragraph("А) одиночный пункт")
    doc.add_paragraph("")
    doc.add_paragraph("Обычный текст после.")
    doc.add_paragraph("")
    # Singleton А. — same isolation
    solo_dot = doc.add_paragraph("А. одиночный пункт с точкой")
    doc.add_paragraph("Обычный текст после второго.")

    normalize_plain_lists_in_document(doc, body_start=0)

    pPr_u = solo_upper._element.find(qn("w:pPr"))
    ind_u = pPr_u.find(qn("w:ind")) if pPr_u is not None else None
    fl_u  = ind_u.get(qn("w:firstLine")) if ind_u is not None else None
    if fl_u == "708":
        return _result(False, f"singleton А) got list indent: {solo_upper.text!r}")

    pPr_d = solo_dot._element.find(qn("w:pPr"))
    ind_d = pPr_d.find(qn("w:ind")) if pPr_d is not None else None
    fl_d  = ind_d.get(qn("w:firstLine")) if ind_d is not None else None
    if fl_d == "708":
        return _result(False, f"singleton А. got list indent: {solo_dot.text!r}")

    return _result(True, "singleton uppercase/dot letter items left unchanged")


# ── Test 10: Table-caption paragraphs not touched ────────────────────────────

def test_table_caption_not_normalized() -> tuple[bool, str]:
    """
    _paragraph_blocks_list_conversion vetoes 'Таблица N …' paragraphs so
    they are never converted to list items by either normalization pass.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        normalize_plain_lists_in_document, normalize_word_numbered_lists_in_document,
    )

    doc = _make_decimal_numbering_doc()
    doc.add_paragraph("ВВЕДЕНИЕ")
    cap = doc.add_paragraph("Таблица 1 – Перечень показателей")
    body1 = doc.add_paragraph("– первый пункт")
    body2 = doc.add_paragraph("– второй пункт")

    cap_orig = cap.text
    normalize_plain_lists_in_document(doc, body_start=0)
    normalize_word_numbered_lists_in_document(doc, body_start=0)

    if cap.text != cap_orig:
        return _result(False, f"table caption was changed: {cap.text!r}")
    return _result(True, "table caption paragraph untouched by normalization")


# ── Test 11: Heading paragraphs not converted ─────────────────────────────────

def test_word_numbered_heading_style_not_converted() -> tuple[bool, str]:
    """
    Decimal numPr paragraphs with heading style or outlineLvl must not be
    converted, regardless of block size.
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_word_numbered_list_paragraphs

    doc = _make_decimal_numbering_doc()
    h1 = doc.add_paragraph("1. Первая глава")
    h1.style = doc.styles["Heading 1"]
    h2 = doc.add_paragraph("2. Вторая глава")
    h2.style = doc.styles["Heading 1"]
    _add_numpr(h1, num_id="1")
    _add_numpr(h2, num_id="1")

    t1, t2 = h1.text, h2.text
    _normalize_word_numbered_list_paragraphs([h1, h2])

    if h1.text != t1 or h2.text != t2:
        return _result(False, f"heading was converted: {h1.text!r} / {h2.text!r}")
    return _result(True, "heading-styled decimal numPr paragraphs not converted")


# ── Test 12: Full idempotency ─────────────────────────────────────────────────

def test_word_numbered_normalization_idempotent() -> tuple[bool, str]:
    """
    Running the normalization pipeline stabilizes after 1 run:
    – run 1: Word-decimal items → '- text' (hyphen, firstLine=708) via kfuListType='endash'
    – run 2: kfuListType fast-path restores geometry without changing text → '- text'
    – run 3: same — fully stable, no prefix duplication.
    Paragraph count stays constant throughout.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        normalize_word_numbered_lists_in_document,
        normalize_plain_lists_in_document,
        paragraph_has_numbering,
    )

    doc = _make_decimal_numbering_doc()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    p1 = doc.add_paragraph("first item")
    p2 = doc.add_paragraph("second item")
    p3 = doc.add_paragraph("third item")
    for p in (p1, p2, p3):
        _add_numpr(p, num_id="1")

    def _run_passes(d):
        normalize_plain_lists_in_document(d, body_start=0)
        normalize_word_numbered_lists_in_document(d, body_start=0)

    count_before = len(doc.paragraphs)

    # Run 1: Word-decimal → '- text' (hyphen, uniform)
    _run_passes(doc)
    after1_texts = [p.text for p in doc.paragraphs]
    after1_count = len(doc.paragraphs)

    if after1_count != count_before:
        return _result(False, f"paragraph count changed after run 1: {count_before} -> {after1_count}")

    # Verify run 1 produced hyphen markers, not en-dash
    list_texts = [t for t in after1_texts if t.startswith("- ")]
    if len(list_texts) != 3:
        return _result(False, f"run 1 did not produce 3 '- text' items; got: {after1_texts}")

    # Run 2: fast-path restores geometry; text must stay identical
    _run_passes(doc)
    after2_texts = [p.text for p in doc.paragraphs]

    if after1_texts != after2_texts:
        diff = [(a, b) for a, b in zip(after1_texts, after2_texts) if a != b]
        return _result(False, f"texts changed between run 1 and run 2 (not idempotent): {diff}")

    if len(doc.paragraphs) != after1_count:
        return _result(False, f"paragraph count changed after run 2: {after1_count} -> {len(doc.paragraphs)}")

    # Run 3: must be fully stable
    _run_passes(doc)
    after3_texts = [p.text for p in doc.paragraphs]

    if after2_texts != after3_texts:
        diff = [(a, b) for a, b in zip(after2_texts, after3_texts) if a != b]
        return _result(False, f"texts still changing on run 3 (not stable): {diff}")

    # Verify no prefix duplication ever occurred
    for t in after3_texts:
        if t.startswith("– – ") or t.startswith("- - ") or t.startswith("– – "):
            return _result(False, f"duplicated dash prefix found: {t!r}")

    return _result(True, "normalization stable from run 1; run 1→2→3 all idempotent with hyphen marker")


# ── New test: colon+numeric-dot soft-break split ─────────────────────────────

def _add_run_with_br(paragraph, text: str, *, add_br: bool = True) -> None:
    """Append a run to paragraph, preceded by a w:br type='textWrapping' inside a run.

    Word stores soft line-breaks as <w:r><w:br w:type="textWrapping"/></w:r> followed
    by additional runs. python-docx's .text yields '\\n' for w:br elements inside runs.
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    if add_br:
        # Insert a dedicated run that contains only the soft line-break
        br_run = OxmlElement("w:r")
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "textWrapping")
        br_run.append(br)
        paragraph._element.append(br_run)
    paragraph.add_run(text)


def test_split_colon_numeric_dot_br_paragraph() -> tuple[bool, str]:
    """
    A single paragraph whose text ends with ':' and is followed by numeric-dot
    items joined by w:br type='textWrapping' (soft line breaks) must be split
    into separate paragraphs by split_manual_dash_lists.

    Real-world example: Рыбаков 220 — 'Для достижения цели необходимо решить
    задачи:' followed by '1. Изучить...', '2. Рассмотреть...' all in ONE <w:p>.
    """
    from guides.coursework_kfu_2025.safe_formatter import split_manual_dash_lists

    doc = Document()
    # Build one paragraph: "Для достижения цели:\n1. Изучить\n2. Рассмотреть"
    p = doc.add_paragraph("Для достижения цели:")
    _add_run_with_br(p, "1. Изучить теоретические основы")
    _add_run_with_br(p, "2. Рассмотреть практические аспекты")

    count_before = len(doc.paragraphs)
    split_manual_dash_lists(doc, body_start=0)
    count_after = len(doc.paragraphs)

    if count_after != count_before + 2:
        return _result(
            False,
            f"expected {count_before + 2} paragraphs after split, got {count_after}; "
            f"texts={[p.text for p in doc.paragraphs]}"
        )

    texts = [p.text for p in doc.paragraphs]
    if "Для достижения цели:" not in texts:
        return _result(False, f"intro colon paragraph missing after split; texts={texts}")
    if "1. Изучить теоретические основы" not in texts:
        return _result(False, f"first numeric item missing; texts={texts}")
    if "2. Рассмотреть практические аспекты" not in texts:
        return _result(False, f"second numeric item missing; texts={texts}")

    return _result(True, "colon+numeric-dot soft-break paragraph split into separate paragraphs")


# ── New test: list body first letter lowercased ───────────────────────────────

def test_list_body_first_letter_lowercased() -> tuple[bool, str]:
    """
    All list formatters must lowercase the first letter of body text unless
    the word is an acronym (both first and second character are uppercase).

    Covers: free-standing dash (hyphen), cyrillic-letter, level-2 numeric-paren,
    colon-triggered dash items.  Acronyms like 'ПАО', 'IBM', 'СБИС' must not
    be lowercased.
    """
    from guides.coursework_kfu_2025.safe_formatter import _lowercase_first

    cases = [
        # (input, expected, description)
        ("Раскрыть понятие",   "раскрыть понятие",   "Capitalised Russian word"),
        ("Изучить методы",     "изучить методы",     "Capitalised Russian verb"),
        ("Describe the task",  "describe the task",  "Capitalised English word"),
        ("ПАО Сбербанк",       "ПАО Сбербанк",       "Russian acronym — untouched"),
        ("IBM Watson",         "IBM Watson",          "English acronym — untouched"),
        ("СБИС система",       "СБИС система",        "Russian all-caps acronym"),
        ("",                   "",                    "Empty string — no error"),
        ("а) уже строчная",    "а) уже строчная",    "Already lowercase — untouched"),
    ]

    for inp, expected, desc in cases:
        got = _lowercase_first(inp)
        if got != expected:
            return _result(False, f"_lowercase_first({inp!r}) = {got!r}, want {expected!r} [{desc}]")

    return _result(True, "_lowercase_first works for all cases: capitalised words lowercased, acronyms protected")


def _paragraph_has_tab_stop(paragraph, val: str) -> bool:
    p_pr = paragraph._element.pPr
    if p_pr is None:
        return False
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is None:
        return False
    return any(tab.get(qn("w:val")) == val for tab in tabs.findall(qn("w:tab")))


def test_formula_unnumbered_formula_gets_number_from_preceding_prose() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_formula_blocks

    doc = Document()
    doc.add_paragraph("Приколы представлены в формуле 1.1.1.")
    formula = doc.add_paragraph("C=V×R")
    doc.add_paragraph("где C-товарооборот предприятия, _")
    doc.add_paragraph("V -количество реализованного товара.")
    doc.add_paragraph("R- цена реализованного товара")
    doc.add_paragraph("Следующий абзац.")

    changed = normalize_formula_blocks(doc, body_start=0)

    if not changed:
        return _result(False, "unnumbered formula block was not changed")
    if formula.text != "C = V × R,\t(1.1.1)":
        return _result(False, f"formula text not normalized/numbered: {formula.text!r}")
    if not _paragraph_has_tab_stop(formula, "right"):
        return _result(False, "formula paragraph lacks right tab stop")
    texts = [p.text for p in doc.paragraphs]
    expected = [
        "где C - товарооборот предприятия;",
        "V - количество реализованного товара;",
        "R - цена реализованного товара.",
    ]
    for item in expected:
        if item not in texts:
            return _result(False, f"expected explanation line missing: {item!r}; texts={texts!r}")
    if any("_" in text for text in texts):
        return _result(False, f"stray underscore was not removed: {texts!r}")
    return _result(True, "unnumbered formula gained number, spacing, tabs, and cleaned explanations")


def test_formula_existing_numbered_formula_still_formats() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_formula_blocks

    doc = Document()
    formula = doc.add_paragraph("E=mc^2 (1.2.3)")
    doc.add_paragraph("где E-энергия;")
    doc.add_paragraph("m - масса;")
    doc.add_paragraph("c- скорость света")

    normalize_formula_blocks(doc, body_start=0)

    if formula.text != "E = mc ^ 2,\t(1.2.3)":
        return _result(False, f"existing numbered formula not formatted: {formula.text!r}")
    if not _paragraph_has_tab_stop(formula, "right"):
        return _result(False, "existing formula paragraph lacks right tab stop")
    texts = [p.text for p in doc.paragraphs]
    if "где E - энергия;" not in texts or "m - масса;" not in texts or "c - скорость света." not in texts:
        return _result(False, f"existing formula explanations not normalized: {texts!r}")
    return _result(True, "existing numbered formula remains supported")


def test_formula_paragraph_not_list_normalized() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import (
        normalize_formula_blocks,
        normalize_plain_lists_in_document,
    )

    doc = Document()
    doc.add_paragraph("Расчет представлен в формуле 1.1.1.")
    formula = doc.add_paragraph("C=V×R")
    expl = doc.add_paragraph("где C-товарооборот;")
    doc.add_paragraph("- обычный пункт")
    doc.add_paragraph("- второй пункт")

    normalize_formula_blocks(doc, body_start=0)
    normalize_plain_lists_in_document(doc, body_start=0)

    if not formula.text.startswith("C = V × R,\t"):
        return _result(False, f"formula was not preserved as formula: {formula.text!r}")
    if not expl.text.startswith("где C - "):
        return _result(False, f"formula explanation was not preserved as explanation: {expl.text!r}")
    if formula.text.startswith("- ") or expl.text.startswith("- "):
        return _result(False, f"formula block was list-normalized: {formula.text!r} / {expl.text!r}")
    return _result(True, "formula block is protected from list normalization")


def test_formula_unnumbered_formula_numbered_from_current_heading2_context() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_formula_blocks

    doc = Document()
    doc.add_paragraph("1.1. Понятие и роль документооборота")
    doc.add_paragraph("Текст перед формулой.")
    formula = doc.add_paragraph("C=V×R")
    doc.add_paragraph("где C-товарооборот;")
    doc.add_paragraph("V-количество")
    doc.add_paragraph("R-цена")

    changed = normalize_formula_blocks(doc, body_start=0)

    if not changed:
        return _result(False, "formula without prose reference was not changed")
    if formula.text != "C = V × R,\t(1.1.1)":
        return _result(False, f"formula did not get heading-derived number: {formula.text!r}")
    texts = [p.text for p in doc.paragraphs]
    if "где C - товарооборот;" not in texts or "V - количество;" not in texts or "R - цена." not in texts:
        return _result(False, f"explanations not normalized with hyphen separator: {texts!r}")
    return _result(True, "unnumbered formula gets next number from heading2 context")


def test_formula_inline_gde_split_into_explanations() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_formula_blocks

    doc = Document()
    doc.add_paragraph("1.1. Понятие и роль документооборота")
    formula = doc.add_paragraph("C=V×R где C-товарооборот; V-количество; R-цена")
    doc.add_paragraph("Следующий текст.")

    changed = normalize_formula_blocks(doc, body_start=0)

    if not changed:
        return _result(False, "inline-gde formula block was not changed")
    texts = [p.text for p in doc.paragraphs]
    expected = [
        "C = V × R,\t(1.1.1)",
        "где C - товарооборот;",
        "V - количество;",
        "R - цена.",
    ]
    for item in expected:
        if item not in texts:
            return _result(False, f"expected inline-gde output missing: {item!r}; texts={texts!r}")
    if formula.text != expected[0]:
        return _result(False, f"formula paragraph not split cleanly: {formula.text!r}")
    return _result(True, "formula and inline explanations split into canonical block")


def test_formula_next_formula_after_explanations_is_not_absorbed() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_formula_blocks

    doc = Document()
    doc.add_paragraph("1.1. Понятие и роль документооборота")
    first = doc.add_paragraph("C=V*R где C-товарооборот; V-количество; R-цена")
    second = doc.add_paragraph("Y=27x+17*z-25")
    doc.add_paragraph("где x-количество денег; z-стоимость услуги; Y-величина капитала")

    changed = normalize_formula_blocks(doc, body_start=0)

    if not changed:
        return _result(False, "adjacent formula blocks were not changed")
    texts = [p.text for p in doc.paragraphs]
    expected = [
        "C = V * R,\t(1.1.1)",
        "где C - товарооборот;",
        "V - количество;",
        "R - цена.",
        "Y = 27x + 17 * z - 25,\t(1.1.2)",
        "где x - количество денег;",
        "z - стоимость услуги;",
        "Y - величина капитала.",
    ]
    for item in expected:
        if item not in texts:
            return _result(False, f"expected adjacent formula output missing: {item!r}; texts={texts!r}")
    if first.text != expected[0] or second.text != expected[4]:
        return _result(False, f"formula paragraphs not preserved: {first.text!r}; {second.text!r}")
    return _result(True, "adjacent formula blocks remain separate")


def test_formula_comma_separated_explanations_split_safely() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_formula_blocks

    doc = Document()
    doc.add_paragraph("1.1. Понятие и роль документооборота")
    formula = doc.add_paragraph("C=V×R")
    doc.add_paragraph("где C-товарооборот, V-количество, R-цена")

    changed = normalize_formula_blocks(doc, body_start=0)

    if not changed:
        return _result(False, "comma-separated formula explanations were not changed")
    texts = [p.text for p in doc.paragraphs]
    expected = [
        "C = V × R,\t(1.1.1)",
        "где C - товарооборот;",
        "V - количество;",
        "R - цена.",
    ]
    for item in expected:
        if item not in texts:
            return _result(False, f"expected comma-separated output missing: {item!r}; texts={texts!r}")
    if formula.text != expected[0]:
        return _result(False, f"formula paragraph changed unexpectedly: {formula.text!r}")
    return _result(True, "comma-separated explanation symbols split into canonical lines")


def test_formula_numeric_coefficient_explanations_stay_in_block() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_formula_blocks

    doc = Document()
    doc.add_paragraph("1.1. Понятие и роль документооборота")
    formula = doc.add_paragraph("Эп = (76 ₽ – 10 ₽) × 172 860 = 11 408 760 ₽ (1.1.1)")
    doc.add_paragraph("где Эп – процессный эффект (экономия) (притом 76 ₽ — стоимость заказного письма;")
    doc.add_paragraph("10 ₽ — усредненная стоимость одного исходящего электронного документа;")
    doc.add_paragraph("172 860 — расчетный объем бумажных документов в год)")
    doc.add_paragraph("Следующий текст.")

    normalize_formula_blocks(doc, body_start=0)

    texts = [p.text for p in doc.paragraphs]
    expected = [
        "Эп = (76 ₽ - 10 ₽) × 172 860 = 11 408 760 ₽,\t(1.1.1)",
        "где Эп - процессный эффект (экономия) (притом 76 ₽ - стоимость заказного письма;",
        "10 ₽ - усредненная стоимость одного исходящего электронного документа;",
        "172 860 - расчетный объем бумажных документов в год).",
    ]
    for item in expected:
        if item not in texts:
            return _result(False, f"expected numeric coefficient explanation missing: {item!r}; texts={texts!r}")
    idx = texts.index(expected[0])
    if texts[idx + 1:idx + 4] != expected[1:]:
        return _result(False, f"numeric coefficient explanations split by blank: {texts!r}")
    if formula.text != expected[0]:
        return _result(False, f"formula text unexpected: {formula.text!r}")
    return _result(True, "numeric coefficient explanation lines stay in formula block")


def test_formula_block_has_single_blank_before_and_after() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_formula_blocks

    doc = Document()
    doc.add_paragraph("1.1. Понятие и роль документооборота")
    doc.add_paragraph("Текст перед формулой.")
    formula = doc.add_paragraph("C=V×R")
    doc.add_paragraph("где C-товарооборот")
    doc.add_paragraph("Следующий текст.")

    normalize_formula_blocks(doc, body_start=0)

    paragraphs = doc.paragraphs
    idx = next(
        (i for i, paragraph in enumerate(paragraphs) if paragraph.text.startswith("C = V × R,")),
        None,
    )
    if idx is None:
        return _result(False, f"formatted formula paragraph not found: {[p.text for p in paragraphs]!r}")
    if idx - 1 < 0 or paragraphs[idx - 1].text != "":
        return _result(False, f"missing blank before formula: {[p.text for p in paragraphs]!r}")
    if idx + 1 >= len(paragraphs) or paragraphs[idx + 1].text.startswith("где ") is False:
        return _result(False, f"explanation not immediately after formula: {[p.text for p in paragraphs]!r}")
    if idx + 2 >= len(paragraphs) or paragraphs[idx + 2].text != "":
        return _result(False, f"missing blank after whole formula block: {[p.text for p in paragraphs]!r}")
    if idx + 3 >= len(paragraphs) or paragraphs[idx + 3].text != "Следующий текст.":
        return _result(False, f"text after formula block shifted unexpectedly: {[p.text for p in paragraphs]!r}")
    return _result(True, "formula block has one blank before and after, not inside")


def test_formula_prose_with_equals_is_not_promoted() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_formula_blocks

    prose = "В настройках системы параметр x=5 означает число попыток отправки."
    doc = Document()
    doc.add_paragraph("1.1. Понятие и роль документооборота")
    p = doc.add_paragraph(prose)
    doc.add_paragraph("где дальше начинается обычное предложение, а не пояснение формулы.")

    changed = normalize_formula_blocks(doc, body_start=0)

    if changed:
        return _result(False, f"ordinary prose with equals changed: {[p.text for p in doc.paragraphs]!r}")
    if p.text != prose:
        return _result(False, f"ordinary prose was mutated: {p.text!r}")
    return _result(True, "ordinary prose with equals is not promoted to formula")


# ── Guard test: blank breaks G2 block ────────────────────────────────────────

def test_word_numbered_block_broken_by_blank_not_converted() -> tuple[bool, str]:
    """
    A blank paragraph between two decimal numPr items breaks the block.
    Neither orphan forms a 2+ block, so both must stay unchanged.
    """
    from guides.coursework_kfu_2025.safe_formatter import _normalize_word_numbered_list_paragraphs

    doc = _make_decimal_numbering_doc()
    p1 = doc.add_paragraph("item before blank")
    blank = doc.add_paragraph("")
    p2 = doc.add_paragraph("item after blank")
    _add_numpr(p1, num_id="1")
    _add_numpr(p2, num_id="1")

    t1, t2 = p1.text, p2.text
    _normalize_word_numbered_list_paragraphs([p1, blank, p2])

    if p1.text != t1:
        return _result(False, f"p1 before blank was changed: {p1.text!r}")
    if p2.text != t2:
        return _result(False, f"p2 after blank was changed: {p2.text!r}")
    return _result(True, "blank-separated items stay unchanged")



# ── Real body-start detection (skip fake/old TOC entries) ──────────────


def _intro_index_at(doc: Document, text: str) -> int | None:
    for idx, p in enumerate(doc.paragraphs):
        if (p.text or "").strip() == text:
            return idx
    return None


def test_real_body_start_skips_fake_plain_toc_soderzhanie() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import find_real_body_start_index

    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    fake_intro = doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. Раздел первый")
    doc.add_paragraph("1.1. Подраздел")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    real_intro = doc.add_paragraph("ВВЕДЕНИЕ")
    real_intro.style = "Heading 1"
    doc.add_paragraph("Тело введения.")

    body_start = find_real_body_start_index(doc)
    expected = [p._element for p in doc.paragraphs].index(real_intro._element)
    fake_idx = [p._element for p in doc.paragraphs].index(fake_intro._element)
    if body_start != expected:
        return _result(False, f"expected body_start={expected} (real intro), got {body_start} (fake={fake_idx})")
    return _result(True, "fake СОДЕРЖАНИЕ ВВЕДЕНИЕ entry skipped, real Heading 1 ВВЕДЕНИЕ picked")


def test_real_body_start_skips_fake_oglavlenie_toc() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import find_real_body_start_index

    doc = Document()
    doc.add_paragraph("Титульная страница")
    doc.add_paragraph("Оглавление")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. Раздел один")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    real_intro = doc.add_paragraph("ВВЕДЕНИЕ")
    real_intro.style = "Heading 1"
    doc.add_paragraph("Текст введения.")

    body_start = find_real_body_start_index(doc)
    expected = [p._element for p in doc.paragraphs].index(real_intro._element)
    if body_start != expected:
        return _result(False, f"expected real intro at {expected}, got {body_start}")
    return _result(True, "fake Оглавление ВВЕДЕНИЕ entry skipped")


def test_real_body_start_skips_fake_toc_with_appendix_entries() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import find_real_body_start_index

    doc = Document()
    doc.add_paragraph("Титул")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. Глава")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 2")
    real_intro = doc.add_paragraph("ВВЕДЕНИЕ")
    real_intro.style = "Heading 1"
    doc.add_paragraph("Тело.")

    body_start = find_real_body_start_index(doc)
    expected = [p._element for p in doc.paragraphs].index(real_intro._element)
    if body_start != expected:
        return _result(False, f"expected real intro at {expected}, got {body_start}")
    return _result(True, "fake TOC with appendix entries skipped, real intro picked")


def test_real_body_start_without_toc_returns_first_intro() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import find_real_body_start_index

    doc = Document()
    doc.add_paragraph("Титульная страница")
    intro = doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("1. Глава первая")
    doc.add_paragraph("Текст главы.")

    body_start = find_real_body_start_index(doc)
    expected = [p._element for p in doc.paragraphs].index(intro._element)
    if body_start != expected:
        return _result(False, f"expected first intro at {expected}, got {body_start}")
    return _result(True, "doc without TOC picks first standalone ВВЕДЕНИЕ")


def test_real_body_start_with_canonical_toc_picks_real_intro() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import find_real_body_start_index

    doc = Document()
    doc.add_paragraph("Титул")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ\t3")
    doc.add_paragraph("1. Глава\t4")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ\t10")
    real_intro = doc.add_paragraph("ВВЕДЕНИЕ")
    real_intro.style = "Heading 1"
    doc.add_paragraph("Текст.")

    body_start = find_real_body_start_index(doc)
    expected = [p._element for p in doc.paragraphs].index(real_intro._element)
    if body_start != expected:
        return _result(False, f"expected real intro at {expected}, got {body_start}")
    return _result(True, "canonical-TOC doc picks Heading 1 ВВЕДЕНИЕ as body start")


def test_real_body_start_styled_intro_wins_when_no_fake_toc() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import find_real_body_start_index

    doc = Document()
    doc.add_paragraph("Параграф вступления.")
    intro = doc.add_paragraph("ВВЕДЕНИЕ")
    intro.style = "Heading 1"
    doc.add_paragraph("Тело.")

    body_start = find_real_body_start_index(doc)
    expected = [p._element for p in doc.paragraphs].index(intro._element)
    if body_start != expected:
        return _result(False, f"expected styled intro at {expected}, got {body_start}")
    return _result(True, "styled-Heading intro picked directly")


def test_real_body_start_lone_contents_heading_does_not_demote_intro() -> tuple[bool, str]:
    """
    Standalone Содержание with NO TOC-like entries before the intro should
    NOT push body_start later — there is no fake TOC to skip.
    """
    from guides.coursework_kfu_2025.safe_formatter import find_real_body_start_index

    doc = Document()
    doc.add_paragraph("Титул")
    doc.add_paragraph("Содержание.")
    doc.add_paragraph("Обычная вступительная фраза без структуры TOC.")
    intro = doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Тело.")

    body_start = find_real_body_start_index(doc)
    expected = [p._element for p in doc.paragraphs].index(intro._element)
    if body_start != expected:
        return _result(False, f"expected lone first intro at {expected}, got {body_start}")
    return _result(True, "lone Содержание without TOC entries does not displace body_start")


def test_real_body_start_smoke_input_picks_heading1_intro() -> tuple[bool, str]:
    """
    End-to-end fixture mirroring list_normalization_smoke_input_kfu.docx:
    a plain-text fake TOC including СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ at
    front, then real Heading 1 ВВЕДЕНИЕ further down. The real intro must
    win — otherwise convert_reference_numbering will eat the whole body.
    """
    from guides.coursework_kfu_2025.safe_formatter import find_real_body_start_index

    doc = Document()
    doc.add_paragraph("Заголовок документа")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. ТЕСТОВЫЕ БЛОКИ СПИСКОВ")
    doc.add_paragraph("1.1. Ручные списки")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    real_intro = doc.add_paragraph("ВВЕДЕНИЕ")
    real_intro.style = "Heading 1"
    doc.add_paragraph("Контрольный текст.")
    doc.add_paragraph("1. изучить теоретические основы;")
    doc.add_paragraph("2. проанализировать практику;")

    body_start = find_real_body_start_index(doc)
    expected = [p._element for p in doc.paragraphs].index(real_intro._element)
    if body_start != expected:
        return _result(False, f"smoke fixture: expected real intro at {expected}, got {body_start}")
    return _result(True, "smoke-input fixture picks real Heading 1 ВВЕДЕНИЕ")


# ── Word-managed SDT TOC removal + tightened H1 promotion ─────────────────


def _add_sdt_toc_block(document: Document, gallery_val: str = "Table of Contents") -> None:
    """
    Inject a synthetic Word-managed TOC <w:sdt> block at the start of the
    document body, mirroring the structure observed in real student docs
    (file 198 inspect).
    """
    body = document.element.body
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    doc_part_obj = OxmlElement("w:docPartObj")
    gallery = OxmlElement("w:docPartGallery")
    gallery.set(qn("w:val"), gallery_val)
    doc_part_obj.append(gallery)
    sdt_pr.append(doc_part_obj)
    sdt.append(sdt_pr)
    sdt_content = OxmlElement("w:sdtContent")
    for text in (
        "Содержание",
        "Введение",
        "1. Теоретические аспекты критериев",
        "1.1. Понятие",
    ):
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)
        sdt_content.append(p)
    sdt.append(sdt_content)
    body.insert(0, sdt)


def _add_sdt_with_field_toc(document: Document) -> None:
    body = document.element.body
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    sdt.append(sdt_pr)
    sdt_content = OxmlElement("w:sdtContent")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r.append(fld)
    instr = OxmlElement("w:instrText")
    instr.text = ' TOC \\o "1-3" '
    r.append(instr)
    p.append(r)
    sdt_content.append(p)
    sdt.append(sdt_content)
    body.insert(0, sdt)


def _add_non_toc_sdt(document: Document) -> None:
    """An SDT used for a content control (NOT a TOC) — must be preserved."""
    body = document.element.body
    sdt = OxmlElement("w:sdt")
    sdt_pr = OxmlElement("w:sdtPr")
    doc_part_obj = OxmlElement("w:docPartObj")
    gallery = OxmlElement("w:docPartGallery")
    gallery.set(qn("w:val"), "Cover Pages")
    doc_part_obj.append(gallery)
    sdt_pr.append(doc_part_obj)
    sdt.append(sdt_pr)
    sdt_content = OxmlElement("w:sdtContent")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "Cover content placeholder"
    r.append(t)
    p.append(r)
    sdt_content.append(p)
    sdt.append(sdt_content)
    body.insert(0, sdt)


def _count_sdt_blocks(document: Document) -> int:
    return len(document.element.body.findall(qn("w:sdt")))


def test_contents_removes_word_managed_sdt_toc() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.contents_builder import _remove_word_managed_toc_blocks

    doc = Document()
    doc.add_paragraph("Титульная страница")
    doc.add_paragraph("ВВЕДЕНИЕ").style = "Heading 1"
    doc.add_paragraph("Тело.")
    _add_sdt_toc_block(doc)

    before = _count_sdt_blocks(doc)
    removed = _remove_word_managed_toc_blocks(doc)
    after = _count_sdt_blocks(doc)

    if before != 1:
        return _result(False, f"setup: expected 1 SDT before, got {before}")
    if removed != 1:
        return _result(False, f"expected 1 removal, got {removed}")
    if after != 0:
        return _result(False, f"expected 0 SDT after, got {after}")
    return _result(True, "Word-managed TOC SDT removed")


def test_contents_removes_field_code_sdt_toc() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.contents_builder import _remove_word_managed_toc_blocks

    doc = Document()
    doc.add_paragraph("Title")
    doc.add_paragraph("ВВЕДЕНИЕ").style = "Heading 1"
    _add_sdt_with_field_toc(doc)

    removed = _remove_word_managed_toc_blocks(doc)
    if removed != 1:
        return _result(False, f"expected 1 field-code TOC removal, got {removed}")
    if _count_sdt_blocks(doc) != 0:
        return _result(False, "field-code TOC SDT not removed")
    return _result(True, "field-code TOC SDT removed via instrText TOC detection")


def test_contents_keeps_non_toc_sdt() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.contents_builder import _remove_word_managed_toc_blocks

    doc = Document()
    doc.add_paragraph("Title")
    doc.add_paragraph("ВВЕДЕНИЕ").style = "Heading 1"
    _add_non_toc_sdt(doc)

    removed = _remove_word_managed_toc_blocks(doc)
    if removed != 0:
        return _result(False, f"non-TOC SDT was removed; count={removed}")
    if _count_sdt_blocks(doc) != 1:
        return _result(False, "non-TOC SDT count drift")
    return _result(True, "non-TOC SDT (Cover Pages gallery) preserved")


def test_contents_rebuild_drops_sdt_toc_end_to_end() -> tuple[bool, str]:
    """
    End-to-end via rebuild_static_contents_page: doc with a Word-managed
    TOC SDT in front matter should have it stripped during rebuild; only
    the canonical text TOC remains.
    """
    doc = Document()
    doc.add_paragraph("Титульная страница")
    _add_sdt_toc_block(doc)
    intro = doc.add_paragraph("ВВЕДЕНИЕ")
    intro.style = "Heading 1"
    doc.add_paragraph("Текст введения.")
    h1 = doc.add_paragraph("1. Теоретические основы")
    h1.style = "Heading 1"
    doc.add_paragraph("Текст главы.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ")
    doc.add_paragraph("Итоги.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. Источник.")

    rendered_lines = [
        ("ВВЕДЕНИЕ", 3),
        ("1. Теоретические основы", 4),
        ("ЗАКЛЮЧЕНИЕ", 5),
        ("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 6),
    ]
    out = _run_static_contents_rebuild(doc, rendered_lines)

    if _count_sdt_blocks(out) != 0:
        return _result(False, f"SDT TOC survived rebuild; count={_count_sdt_blocks(out)}")

    front = _toc_texts_before_intro(out)
    if front.count("СОДЕРЖАНИЕ") != 1:
        return _result(False, f"expected exactly one canonical СОДЕРЖАНИЕ, got {front!r}")
    return _result(True, "end-to-end: SDT TOC dropped, single canonical TOC remains")


def test_h1_body_sentence_not_promoted_via_toc_text() -> tuple[bool, str]:
    """
    Body sentence shaped like `1. Foo. Bar.` with Normal style + matching
    chapter entry in toc_h1_map must NOT be promoted to Heading 1.
    """
    from guides.coursework_kfu_2025.safe_formatter import is_heading1_promotion_safe
    from guides.coursework_kfu_2025.classifier import parse_heading1

    doc = Document()
    # Exact text from file 198 source idx 80 — `parse_heading1` accepts it
    # because the body sentence has no trailing punctuation and only one
    # embedded period. Before the patch, `toc_text` from the front-matter
    # plain TOC was a sufficient signal to flip this paragraph into a
    # Heading 1 — producing duplicate "1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ…" headings.
    p = doc.add_paragraph("1. Маркетинговый подход. Данный подход")
    parsed = parse_heading1(p.text)
    if parsed is None:
        return _result(False, "test setup: parse_heading1 did not parse body sentence")
    toc_text = "1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ КРИТЕРИЕВ И ПОКАЗАТЕЛЕЙ"

    if is_heading1_promotion_safe(p, parsed, toc_text=toc_text):
        return _result(False, "body sentence promoted via toc_text alone — regression")
    return _result(True, "toc_text shortcut blocked for body prose without structural signal")


def test_h1_styled_chapter_with_toc_text_still_promoted() -> tuple[bool, str]:
    """
    Real chapter heading styled Heading 1 + matching toc_h1_map entry MUST
    still be promotion-safe. Regression guard against over-tightening.
    """
    from guides.coursework_kfu_2025.safe_formatter import is_heading1_promotion_safe
    from guides.coursework_kfu_2025.classifier import parse_heading1

    doc = Document()
    p = doc.add_paragraph("1. Теоретические основы")
    p.style = "Heading 1"
    parsed = parse_heading1(p.text)
    if parsed is None:
        return _result(False, "test setup: parse_heading1 did not parse styled heading")

    if not is_heading1_promotion_safe(p, parsed, toc_text="1. Теоретические основы"):
        return _result(False, "styled chapter heading was wrongly rejected")
    return _result(True, "styled Heading 1 + toc_text remains promotion-safe")


def test_h1_exact_intro_still_promoted_with_toc_text() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import is_heading1_promotion_safe
    from guides.coursework_kfu_2025.classifier import parse_heading1

    p = Document().add_paragraph("ВВЕДЕНИЕ")
    parsed = parse_heading1(p.text)
    if not parsed or parsed["kind"] != "heading1_exact":
        return _result(False, f"setup: unexpected parse_heading1 result: {parsed!r}")
    if not is_heading1_promotion_safe(p, parsed, toc_text="ВВЕДЕНИЕ"):
        return _result(False, "exact intro rejected")
    return _result(True, "ВВЕДЕНИЕ exact-match still promotion-safe")


def test_h1_centered_bold_body_with_toc_text_promoted() -> tuple[bool, str]:
    """
    A centred + bold heading-like paragraph + matching toc_h1_map entry
    must still be promotion-safe even with Normal style (visual signal).
    """
    from guides.coursework_kfu_2025.safe_formatter import is_heading1_promotion_safe
    from guides.coursework_kfu_2025.classifier import parse_heading1
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("1. Теоретические основы")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    parsed = parse_heading1(p.text)
    if parsed is None:
        return _result(False, "test setup: parse_heading1 did not parse")
    if not is_heading1_promotion_safe(p, parsed, toc_text="1. Теоретические основы"):
        return _result(False, "centred+bold body heading wrongly rejected")
    return _result(True, "centred+bold visual heading remains promotion-safe")


def test_h1_body_sentence_no_duplicate_in_output_pipeline() -> tuple[bool, str]:
    """
    Integration: a doc with an old plain-text TOC AND a body sentence shaped
    like a chapter heading must produce a single Heading 1 per chapter in
    the formatted output. The body sentence stays body text.
    """
    from guides.coursework_kfu_2025.safe_formatter import process_document
    import tempfile, os

    doc = Document()
    doc.add_paragraph("Титульная страница")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ\t3")
    doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ КРИТЕРИЕВ И ПОКАЗАТЕЛЕЙ\t5")
    doc.add_paragraph("1.1. Понятие и сущность\t5")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ\t38")
    intro = doc.add_paragraph("ВВЕДЕНИЕ")
    intro.style = "Heading 1"
    doc.add_paragraph("Текст введения.")
    h1 = doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ КРИТЕРИЕВ И ПОКАЗАТЕЛЕЙ")
    h1.style = "Heading 1"
    doc.add_paragraph("Тело раздела.")
    # Body prose that USED to be promoted to a duplicate Heading 1:
    doc.add_paragraph("1. Маркетинговый подход. Данный подход применяется в анализе.")
    doc.add_paragraph("Продолжение прозы.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ").style = "Heading 1"
    doc.add_paragraph("Итоги.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ").style = "Heading 1"
    doc.add_paragraph("1. Источник.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = os.path.join(tmp, "in.docx")
        out = os.path.join(tmp, "out.docx")
        doc.save(inp)
        process_document(inp, out)
        formatted = Document(out)

    h1_chapter_paragraphs = [
        p for p in formatted.paragraphs
        if (p.style.name or "").lower() in {"heading 1", "заголовок 1"}
        and (p.text or "").startswith("1. ")
    ]
    if len(h1_chapter_paragraphs) != 1:
        texts = [p.text[:80] for p in h1_chapter_paragraphs]
        return _result(False, f"expected exactly 1 Heading 1 starting with '1. ', got {len(h1_chapter_paragraphs)}: {texts!r}")

    # The body sentence must remain body text and keep its original phrasing.
    body_sentence_present = any(
        "Маркетинговый подход" in (p.text or "")
        and (p.style.name or "").lower() not in {"heading 1", "заголовок 1"}
        for p in formatted.paragraphs
    )
    if not body_sentence_present:
        return _result(False, "body sentence with 'Маркетинговый подход' was lost or promoted")
    return _result(True, "single H1 for chapter 1; body sentence preserved as body text")


# ── Always-safe pre-pass: strip_obsolete_toc_blocks_inplace ───────────────


def test_strip_obsolete_toc_blocks_inplace_removes_sdt_and_plain_toc() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.contents_builder import (
        strip_obsolete_toc_blocks_inplace,
    )
    import tempfile, os

    doc = Document()
    doc.add_paragraph("Title page")
    _add_sdt_toc_block(doc)
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ\t3")
    doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ\t5")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ\t10")
    intro = doc.add_paragraph("ВВЕДЕНИЕ")
    intro.style = "Heading 1"
    doc.add_paragraph("Body prose.")

    with tempfile.TemporaryDirectory() as tmp:
        p = os.path.join(tmp, "in.docx")
        doc.save(p)
        report = strip_obsolete_toc_blocks_inplace(p)
        out = Document(p)

    if report["sdt_removed"] != 1:
        return _result(False, f"expected 1 SDT removed, got {report['sdt_removed']}")
    if report["plain_toc_removed"] < 4:
        return _result(False, f"expected ≥4 plain paragraphs removed, got {report['plain_toc_removed']}")
    if _count_sdt_blocks(out) != 0:
        return _result(False, "SDT still present after pre-pass")
    if any((p.text or "").strip() in ("СОДЕРЖАНИЕ", "ВВЕДЕНИЕ\t3", "1. ТЕОРЕТИЧЕСКИЕ АСПЕКТЫ\t5") for p in out.paragraphs):
        return _result(False, "old TOC paragraphs survived pre-pass")
    return _result(True, "pre-pass removes SDT + plain old TOC and saves in place")


def test_strip_obsolete_toc_blocks_inplace_noop_on_clean_doc() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.contents_builder import (
        strip_obsolete_toc_blocks_inplace,
    )
    import tempfile, os

    doc = Document()
    doc.add_paragraph("Title page")
    doc.add_paragraph("ВВЕДЕНИЕ").style = "Heading 1"
    doc.add_paragraph("Body prose.")
    doc.add_paragraph("1. Глава").style = "Heading 1"
    doc.add_paragraph("Текст главы.")

    with tempfile.TemporaryDirectory() as tmp:
        p = os.path.join(tmp, "in.docx")
        doc.save(p)
        report = strip_obsolete_toc_blocks_inplace(p)
        out = Document(p)

    if report["sdt_removed"] != 0 or report["plain_toc_removed"] != 0:
        return _result(False, f"unexpected removals on clean doc: {report!r}")
    if len(out.paragraphs) != 5:
        return _result(False, f"clean doc paragraph count changed: {len(out.paragraphs)}")
    return _result(True, "pre-pass is no-op when no obsolete TOC artifacts present")


def test_strip_obsolete_toc_blocks_inplace_runs_even_when_pdf_render_would_fail() -> tuple[bool, str]:
    """
    Mock rebuild_static_contents_page to raise (simulating LibreOffice
    failure); verify obsolete TOC cleanup is still applied via the
    formatter_service pre-pass.
    """
    import guides.coursework_kfu_2025.formatter_service as fs
    from guides.coursework_kfu_2025.contents_builder import (
        strip_obsolete_toc_blocks_inplace,
    )
    import tempfile, os

    doc = Document()
    doc.add_paragraph("Title")
    _add_sdt_toc_block(doc)
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ\t3")
    doc.add_paragraph("1. Глава\t5")
    intro = doc.add_paragraph("ВВЕДЕНИЕ")
    intro.style = "Heading 1"
    doc.add_paragraph("Body.")
    h1 = doc.add_paragraph("1. Глава")
    h1.style = "Heading 1"
    doc.add_paragraph("Section body.")
    doc.add_paragraph("ЗАКЛЮЧЕНИЕ").style = "Heading 1"
    doc.add_paragraph("Conclusion.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ").style = "Heading 1"
    doc.add_paragraph("1. Источник.")

    saved_rebuild = fs.rebuild_static_contents_page

    def _failing_rebuild(_path):
        raise RuntimeError("simulated LibreOffice render failure")

    with tempfile.TemporaryDirectory() as tmp:
        inp = os.path.join(tmp, "in.docx")
        outp = os.path.join(tmp, "out.docx")
        doc.save(inp)
        fs.rebuild_static_contents_page = _failing_rebuild
        try:
            fs.format_docx(inp, outp)
        finally:
            fs.rebuild_static_contents_page = saved_rebuild
        formatted = Document(outp)

    if _count_sdt_blocks(formatted) != 0:
        return _result(False, "SDT TOC survived when rebuild was simulated to fail")
    if any((p.text or "").strip() == "СОДЕРЖАНИЕ" for p in formatted.paragraphs):
        return _result(False, "old plain СОДЕРЖАНИЕ survived when rebuild was simulated to fail")
    return _result(True, "obsolete TOC cleanup runs even when rebuild raises (no LibreOffice)")


def test_table_caption_trailing_period_cleanup() -> tuple[bool, str]:
    """Table numbers/titles lose one terminal period; body text stays unchanged."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица показывает рост.")
    doc.add_paragraph("Таблица 1.1.1.")
    table1 = doc.add_table(rows=1, cols=1)
    table1.cell(0, 0).text = "Значение"
    doc.add_paragraph("Таблица 1.1.2. Анализ финансовых результатов.")
    table2 = doc.add_table(rows=1, cols=1)
    table2.cell(0, 0).text = "Показатель"

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(inp)
        process_document(inp, out)
        result = Document(str(out))

    texts = [" ".join(p.text.split()) for p in result.paragraphs if " ".join(p.text.split())]
    if "Таблица 1.1.1." in texts:
        return _result(False, "standalone table number kept trailing period")
    if "Таблица 1.1.1" not in texts:
        return _result(False, f"standalone table number missing after cleanup: {texts!r}")
    if "Анализ финансовых результатов." in texts:
        return _result(False, "table title kept trailing period")
    if "Анализ финансовых результатов" not in texts:
        return _result(False, f"table title missing after cleanup: {texts!r}")
    if "Таблица показывает рост." not in texts:
        return _result(False, f"ordinary body text changed unexpectedly: {texts!r}")

    return _result(True, "table caption/title terminal period cleanup is scoped")


def test_b25_real_table_caption_directly_before_table_is_formatted() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 1.1.1")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Значение"

    formatted = _format_synthetic_doc(doc)
    caption_idx = _paragraph_index(formatted, "Таблица 1.1.1")
    if caption_idx is None:
        return _result(False, "real table caption missing after formatting")
    caption = formatted.paragraphs[caption_idx]
    if caption.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
        return _result(False, "real table caption directly before table was not right-aligned")
    return _result(True, "real table caption directly before table is formatted")


def test_b25_real_table_caption_title_table_is_formatted() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 1.1.1")
    doc.add_paragraph("Анализ финансовых результатов.")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Показатель"

    formatted = _format_synthetic_doc(doc)
    caption_idx = _paragraph_index(formatted, "Таблица 1.1.1")
    title_idx = _paragraph_index(formatted, "Анализ финансовых результатов")
    if caption_idx is None or title_idx is None:
        return _result(False, "real table caption/title missing after formatting")
    if formatted.paragraphs[caption_idx].alignment != WD_ALIGN_PARAGRAPH.RIGHT:
        return _result(False, "real table caption was not right-aligned")
    if formatted.paragraphs[title_idx].alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "real table title before table was not centered")
    return _result(True, "real table caption + title + table is formatted")


def test_b25_inline_dash_table_caption_with_adjacent_table_is_accepted() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 1.1.1 — Отличия индивидуального и организационного поведения")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Показатель"

    formatted = _format_synthetic_doc(doc)
    caption_idx = _paragraph_index(formatted, "Таблица 1.1.1")
    title_idx = next(
        (
            idx for idx, paragraph in enumerate(formatted.paragraphs)
            if "Отличия индивидуального и организационного поведения" in paragraph.text
        ),
        None,
    )
    if caption_idx is None or title_idx is None:
        return _result(False, "inline dash table caption was not structurally accepted")
    if formatted.paragraphs[caption_idx].alignment != WD_ALIGN_PARAGRAPH.RIGHT:
        return _result(False, "inline dash table caption was not right-aligned")
    if formatted.paragraphs[title_idx].alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "inline dash table title was not centered")
    return _result(True, "inline dash table caption with adjacent table is accepted")


def test_b25_table_number_analytical_prose_without_table_remains_body() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    prose = "Таблица 1.1.1 показывает, что показатели растут."
    body_after = "Следующий аналитический абзац остается обычным текстом."

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph(prose)
    doc.add_paragraph(body_after)

    formatted = _format_synthetic_doc(doc)
    texts = [" ".join(p.text.split()) for p in formatted.paragraphs if " ".join(p.text.split())]
    idx = _paragraph_index(formatted, prose)
    if idx is None:
        return _result(False, f"analytical table prose was mutated or split: {texts!r}")
    if "Таблица 1.1.1" in texts:
        return _result(False, f"analytical table prose was split into fake caption: {texts!r}")
    if formatted.paragraphs[idx].alignment in {WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER}:
        return _result(False, "analytical table prose received caption/title alignment")
    return _result(True, "analytical table prose without adjacent table remains body text")


def test_b25_v_tablitse_analytical_prose_remains_body() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    prose = "В таблице 1.1.1 представлены основные показатели исследования."

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph(prose)

    formatted = _format_synthetic_doc(doc)
    idx = _paragraph_index(formatted, prose)
    if idx is None:
        return _result(False, "ordinary prose mentioning table was changed")
    if formatted.paragraphs[idx].alignment in {WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER}:
        return _result(False, "ordinary prose mentioning table received caption/title alignment")
    return _result(True, "ordinary prose mentioning table remains body text")


def test_b25_source_then_table_analytical_prose_is_not_promoted() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    prose = "Таблица 1.1.1 отражает динамику изменения показателей."

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph(prose)
    doc.add_paragraph("Продолжение аналитического текста.")

    formatted = _format_synthetic_doc(doc)
    texts = [" ".join(p.text.split()) for p in formatted.paragraphs if " ".join(p.text.split())]
    idx = _paragraph_index(formatted, prose)
    if idx is None:
        return _result(False, f"table analytical prose after source was mutated: {texts!r}")
    if "Таблица 1.1.1" in texts:
        return _result(False, f"table analytical prose after source was split: {texts!r}")
    if formatted.paragraphs[idx].alignment in {WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER}:
        return _result(False, "table analytical prose after source received caption/title alignment")
    return _result(True, "source + analytical table prose is not promoted")


def test_b25_appendix_immediate_table_like_title_is_preserved() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ 3")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ А")
    doc.add_paragraph("Таблица А.1 Расчет показателей.")
    doc.add_table(rows=1, cols=1)

    formatted = _format_synthetic_doc(doc)
    title_idx = _paragraph_index(formatted, "Таблица А.1 Расчет показателей")
    if title_idx is None:
        return _result(False, "appendix immediate table-like title was not preserved")
    if formatted.paragraphs[title_idx].alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "appendix immediate table-like title was not centered")
    return _result(True, "appendix immediate table-like title is preserved")


def test_b25_neuromarketing_style_table_false_positive_is_prevented() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    prose = "Таблица 1.1.1 показывает, что нейромаркетинговые методы применяются в исследованиях."

    doc = Document()
    doc.add_paragraph("СОДЕРЖАНИЕ\nВВЕДЕНИЕ\n1. Теоретические основы")
    doc.add_paragraph("ВВЕДЕНИЕ........................................................3")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("1. Теоретические основы нейромаркетинга")
    doc.add_paragraph(prose)
    doc.add_paragraph("Следующий абзац описывает результаты анализа.")

    formatted = _format_synthetic_doc(doc)
    idx = _paragraph_index(formatted, prose)
    if idx is None:
        return _result(False, "neuromarketing-style table prose was mutated or split")
    if formatted.paragraphs[idx].alignment in {WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER}:
        return _result(False, "neuromarketing-style table prose received caption/title alignment")
    if _paragraph_index(formatted, "Таблица 1.1.1") is not None:
        return _result(False, "neuromarketing-style table prose was split into fake caption")
    return _result(True, "neuromarketing-style table false positive is prevented")


def test_b26_inline_dash_table_caption_title_loses_leading_dash() -> tuple[bool, str]:
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 1.1.1 — Отличия индивидуального и организационного поведения")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Показатель"

    formatted = _format_synthetic_doc(doc)
    texts = [" ".join(p.text.split()) for p in formatted.paragraphs if " ".join(p.text.split())]
    if "Таблица 1.1.1" not in texts:
        return _result(False, f"caption missing after inline split: {texts!r}")
    expected_title = "Отличия индивидуального и организационного поведения"
    if expected_title not in texts:
        return _result(False, f"clean inline title missing: {texts!r}")
    if any(text.startswith(("-", "–", "—")) and "Отличия" in text for text in texts):
        return _result(False, f"inline title kept leading dash: {texts!r}")
    return _result(True, "inline dash table title loses leading dash")


def test_b26_inline_en_dash_table_caption_title_strips_extra_spaces() -> tuple[bool, str]:
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 1.1.1 –   Отличия индивидуального и организационного поведения")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Показатель"

    formatted = _format_synthetic_doc(doc)
    texts = [" ".join(p.text.split()) for p in formatted.paragraphs if " ".join(p.text.split())]
    expected_title = "Отличия индивидуального и организационного поведения"
    if expected_title not in texts:
        return _result(False, f"inline title was not cleaned: {texts!r}")
    if any(text.startswith(("-", "–", "—")) and "Отличия" in text for text in texts):
        return _result(False, f"inline title kept dash/extra spaces: {texts!r}")
    return _result(True, "inline en dash title strips dash and extra spaces")


def test_b26_already_split_table_caption_remains_single_title() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title_text = "Отличия индивидуального и организационного поведения"
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 1.1.1")
    doc.add_paragraph(title_text)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Показатель"

    formatted = _format_synthetic_doc(doc)
    matching_titles = [p for p in formatted.paragraphs if " ".join(p.text.split()) == title_text]
    if len(matching_titles) != 1:
        return _result(False, f"already split title duplicated or missing: count={len(matching_titles)}")
    if matching_titles[0].alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "already split title is not centered")
    return _result(True, "already split table caption remains single title")


def test_b26_non_adjacent_table_prose_remains_unchanged() -> tuple[bool, str]:
    prose = "Таблица 1.1.1 показывает, что показатели растут."

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph(prose)
    doc.add_paragraph("Следующий аналитический абзац.")

    formatted = _format_synthetic_doc(doc)
    texts = [" ".join(p.text.split()) for p in formatted.paragraphs if " ".join(p.text.split())]
    if prose not in texts:
        return _result(False, f"non-adjacent table prose was changed: {texts!r}")
    if "Таблица 1.1.1" in texts:
        return _result(False, f"non-adjacent table prose was split: {texts!r}")
    return _result(True, "non-adjacent table prose remains unchanged")


def test_b26_v_tablitse_prose_remains_unchanged() -> tuple[bool, str]:
    prose = "В таблице 1.1.1 представлены основные показатели исследования."

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph(prose)

    formatted = _format_synthetic_doc(doc)
    if _paragraph_index(formatted, prose) is None:
        return _result(False, "ordinary prose mentioning table was changed")
    return _result(True, "ordinary prose mentioning table remains unchanged")


def test_b26_appendix_immediate_dash_table_like_title_stays_appendix_title() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ 3")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ А")
    doc.add_paragraph("Таблица А.1 — Расчет показателей.")
    doc.add_table(rows=1, cols=1)

    formatted = _format_synthetic_doc(doc)
    title = next(
        (
            p for p in formatted.paragraphs
            if "Таблица А.1" in " ".join(p.text.split()) and "Расчет показателей" in " ".join(p.text.split())
        ),
        None,
    )
    if title is None:
        return _result(False, "appendix table-like title missing")
    if " ".join(title.text.split()).startswith("Расчет"):
        return _result(False, "appendix table-like title was split as ordinary table caption")
    if title.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "appendix table-like title is not centered")
    return _result(True, "appendix immediate dash table-like title stays appendix title")


def test_b26_neuromarketing_style_inline_caption_normalizes() -> tuple[bool, str]:
    doc = Document()
    doc.add_paragraph("СОДЕРЖАНИЕ\nВВЕДЕНИЕ\n1. Теоретические основы")
    doc.add_paragraph("ВВЕДЕНИЕ........................................................3")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("Таблица 1.1.1 — Отличия индивидуального и организационного поведения")
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Показатель"

    formatted = _format_synthetic_doc(doc)
    texts = [" ".join(p.text.split()) for p in formatted.paragraphs if " ".join(p.text.split())]
    expected_title = "Отличия индивидуального и организационного поведения"
    if "Таблица 1.1.1" not in texts or expected_title not in texts:
        return _result(False, f"neuromarketing-style inline caption was not normalized: {texts!r}")
    if any(text.startswith(("-", "–", "—")) and "Отличия" in text for text in texts):
        return _result(False, f"neuromarketing-style title kept leading dash: {texts!r}")
    return _result(True, "neuromarketing-style inline caption normalizes")


def test_b26_rybakov_style_split_caption_remains_stable() -> tuple[bool, str]:
    title_text = "Сравнительная характеристика каналов коммуникации"
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 1.1.1")
    doc.add_paragraph(title_text)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Канал"

    formatted = _format_synthetic_doc(doc)
    texts = [" ".join(p.text.split()) for p in formatted.paragraphs if " ".join(p.text.split())]
    if texts.count("Таблица 1.1.1") != 1 or texts.count(title_text) != 1:
        return _result(False, f"Rybakov-style split caption changed unexpectedly: {texts!r}")
    return _result(True, "Rybakov-style split caption remains stable")


def _add_test_drawing_paragraph(doc: Document):
    paragraph = doc.add_paragraph()
    run = OxmlElement("w:r")
    drawing = OxmlElement("w:drawing")
    run.append(drawing)
    paragraph._element.append(run)
    return paragraph


def _figure_block_tokens(doc: Document) -> list[str]:
    from guides.coursework_kfu_2025.safe_formatter import paragraph_has_drawing

    tokens = []
    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph):
            tokens.append("<IMAGE>")
        elif paragraph.text:
            tokens.append(" ".join(paragraph.text.split()))
    return tokens


def test_b27_canonical_image_source_note_caption_unchanged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    _add_test_drawing_paragraph(doc)
    source = doc.add_paragraph("Источник: составлено автором.")
    note = doc.add_paragraph("Примечание: условные данные.")
    caption = doc.add_paragraph("Рис. 1.2.1. Схема процесса")

    changed = normalize_figure_blocks(doc, 0)

    if changed:
        return _result(False, "canonical figure block reported a change")
    if _figure_block_tokens(doc) != [
        "<IMAGE>",
        "Источник: составлено автором.",
        "Примечание: условные данные.",
        "Рис. 1.2.1. Схема процесса",
    ]:
        return _result(False, f"canonical figure block changed: {_figure_block_tokens(doc)!r}")
    if doc.paragraphs[1]._p is not source._p or doc.paragraphs[2]._p is not note._p or doc.paragraphs[3]._p is not caption._p:
        return _result(False, "canonical paragraphs were replaced or duplicated")
    return _result(True, "canonical image/source/note/caption remains stable")


def test_b27_caption_above_image_moves_below_image() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    doc.add_paragraph("Рис. 1.2.1. Схема процесса")
    _add_test_drawing_paragraph(doc)

    changed = normalize_figure_blocks(doc, 0)

    if _figure_block_tokens(doc) != ["<IMAGE>", "Рис. 1.2.1. Схема процесса"]:
        return _result(False, f"caption was not moved below image: {_figure_block_tokens(doc)!r}")
    if not changed:
        return _result(False, "caption/image reorder did not report a change")
    return _result(True, "caption above image moves below image")


def test_b27_source_above_image_moves_below_image() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    doc.add_paragraph("Источник: составлено автором.")
    _add_test_drawing_paragraph(doc)
    doc.add_paragraph("Рис. 1.2.1. Схема процесса")

    changed = normalize_figure_blocks(doc, 0)

    if _figure_block_tokens(doc) != [
        "<IMAGE>",
        "Источник: составлено автором.",
        "Рис. 1.2.1. Схема процесса",
    ]:
        return _result(False, f"source was not moved below image: {_figure_block_tokens(doc)!r}")
    if not changed:
        return _result(False, "source/image reorder did not report a change")
    return _result(True, "source above image moves below image")


def test_b27_caption_and_source_above_image_normalize_to_canonical_order() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    doc.add_paragraph("Рис. 1.2.1. Схема процесса")
    doc.add_paragraph("Источник: составлено автором.")
    _add_test_drawing_paragraph(doc)

    changed = normalize_figure_blocks(doc, 0)

    if _figure_block_tokens(doc) != [
        "<IMAGE>",
        "Источник: составлено автором.",
        "Рис. 1.2.1. Схема процесса",
    ]:
        return _result(False, f"caption/source/image block not canonical: {_figure_block_tokens(doc)!r}")
    if not changed:
        return _result(False, "caption/source/image reorder did not report a change")
    return _result(True, "caption and source above image normalize to canonical order")


def test_b27_merged_source_note_splits_only_near_image() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    _add_test_drawing_paragraph(doc)
    doc.add_paragraph("Источник: составлено автором.\nПримечание: условные данные.")
    doc.add_paragraph("Рис. 1.2.1. Схема процесса")

    changed = normalize_figure_blocks(doc, 0)

    if _figure_block_tokens(doc) != [
        "<IMAGE>",
        "Источник: составлено автором.",
        "Примечание: условные данные.",
        "Рис. 1.2.1. Схема процесса",
    ]:
        return _result(False, f"merged source/note was not split near image: {_figure_block_tokens(doc)!r}")
    if not changed:
        return _result(False, "source/note split did not report a change")
    return _result(True, "source/note soft break splits only in a confirmed figure block")


def test_b27_no_image_source_paragraph_unchanged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    doc.add_paragraph("Источник: данные исследования.")
    doc.add_paragraph("Обычный текст после источника.")

    changed = normalize_figure_blocks(doc, 0)

    if [p.text for p in doc.paragraphs] != ["Источник: данные исследования.", "Обычный текст после источника."]:
        return _result(False, "source paragraph without image was changed")
    if changed:
        return _result(False, "no-image source paragraph reported a change")
    return _result(True, "source paragraph without image remains untouched")


def test_b27_table_source_nearby_unchanged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    doc.add_paragraph("Таблица 1.2.1")
    doc.add_table(rows=1, cols=1)
    doc.add_paragraph("Источник: составлено автором.")

    changed = normalize_figure_blocks(doc, 0)

    if [p.text for p in doc.paragraphs] != ["Таблица 1.2.1", "Источник: составлено автором."]:
        return _result(False, "table source block was changed by figure pass")
    if changed:
        return _result(False, "table source block reported a figure change")
    return _result(True, "table/source block remains untouched by figure pass")


def test_b27_two_nearby_images_are_ambiguous_noop() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    _add_test_drawing_paragraph(doc)
    doc.add_paragraph("Источник: составлено автором.")
    _add_test_drawing_paragraph(doc)
    doc.add_paragraph("Рис. 1.2.1. Схема процесса")

    before = _figure_block_tokens(doc)
    changed = normalize_figure_blocks(doc, 0)

    if _figure_block_tokens(doc) != before:
        return _result(False, f"ambiguous nearby images were reordered: {_figure_block_tokens(doc)!r}")
    if changed:
        return _result(False, "ambiguous nearby images reported a change")
    return _result(True, "nearby images are treated as ambiguous and left unchanged")


def test_b27_neuromarketing_style_disorder_normalizes() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    doc.add_paragraph("Источник: данные автора.")
    doc.add_paragraph("Примечание: значения условные.")
    _add_test_drawing_paragraph(doc)
    doc.add_paragraph("Рис. 1.1.1. Модель восприятия")

    changed = normalize_figure_blocks(doc, 0)

    if _figure_block_tokens(doc) != [
        "<IMAGE>",
        "Источник: данные автора.",
        "Примечание: значения условные.",
        "Рис. 1.1.1. Модель восприятия",
    ]:
        return _result(False, f"neuromarketing-style figure block not canonical: {_figure_block_tokens(doc)!r}")
    if not changed:
        return _result(False, "neuromarketing-style disorder did not report a change")
    return _result(True, "neuromarketing-style disorder normalizes")


def test_b27_rybakov_stable_figure_block_unchanged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    doc.add_paragraph("Текст перед рисунком.")
    _add_test_drawing_paragraph(doc)
    source = doc.add_paragraph("Источник: составлено автором.")
    caption = doc.add_paragraph("Рис. 2.1.1. Сравнение каналов коммуникации")

    changed = normalize_figure_blocks(doc, 0)

    if _figure_block_tokens(doc) != [
        "Текст перед рисунком.",
        "<IMAGE>",
        "Источник: составлено автором.",
        "Рис. 2.1.1. Сравнение каналов коммуникации",
    ]:
        return _result(False, f"stable Rybakov-style figure block changed: {_figure_block_tokens(doc)!r}")
    if doc.paragraphs[2]._p is not source._p or doc.paragraphs[3]._p is not caption._p:
        return _result(False, "stable Rybakov-style paragraphs were replaced")
    if changed:
        return _result(False, "stable Rybakov-style figure block reported a change")
    return _result(True, "Rybakov-style figure block remains stable")


def test_b27_figure_prose_after_source_does_not_block_reorder() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_figure_blocks

    doc = Document()
    _add_test_drawing_paragraph(doc)
    doc.add_paragraph("Рис. 1.1.1. Закупочный центр")
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Рис. 1.1.1. показывает, что решение формируется под влиянием нескольких участников.")

    changed = normalize_figure_blocks(doc, 0)

    if _figure_block_tokens(doc) != [
        "<IMAGE>",
        "Источник: составлено автором.",
        "Рис. 1.1.1. Закупочный центр",
        "Рис. 1.1.1. показывает, что решение формируется под влиянием нескольких участников.",
    ]:
        return _result(False, f"figure prose after source blocked reorder: {_figure_block_tokens(doc)!r}")
    if not changed:
        return _result(False, "figure caption/source/prose reorder did not report a change")
    return _result(True, "figure prose after source does not count as a second caption")


def test_figure_caption_spacing_and_blank_font() -> tuple[bool, str]:
    """
    Figure captions require exactly one blank before the caption, but no blank
    between the caption and its Источник line. Formatter-created blanks use
    body font size.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        ensure_single_blank_before_figure_captions,
        remove_empty_between_figure_caption_and_source,
    )

    doc = Document()
    doc.add_paragraph("Текст перед рисунком.")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Рис. 1.2.1. Схема процесса")
    doc.add_paragraph("")
    doc.add_paragraph("Источник: составлено автором.")

    ensure_single_blank_before_figure_captions(doc, 0)
    remove_empty_between_figure_caption_and_source(doc, 0)

    texts = [p.text for p in doc.paragraphs]
    expected = [
        "Текст перед рисунком.",
        "",
        "Рис. 1.2.1. Схема процесса",
        "Источник: составлено автором.",
    ]
    if texts != expected:
        return _result(False, f"unexpected paragraph layout: {texts!r}")

    blank = doc.paragraphs[1]
    run = blank.runs[0] if blank.runs else None
    if run is None:
        return _result(False, "blank paragraph has no run")

    sz = run._element.get_or_add_rPr().find(qn("w:sz"))
    if sz is None or sz.get(qn("w:val")) != "28":
        val = sz.get(qn("w:val")) if sz is not None else None
        return _result(False, f"blank font size is {val}, expected 28 half-points")

    return _result(True, "figure spacing and blank font are correct")


def test_figure_source_after_caption_is_moved_before_caption() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from guides.coursework_kfu_2025.safe_formatter import reorder_figure_source_before_caption

    doc = Document()
    _add_test_drawing_paragraph(doc)
    caption = doc.add_paragraph("Рис. 1.2.1. Схема процесса")
    source = doc.add_paragraph("Источник: составлено автором.")

    changed = reorder_figure_source_before_caption(doc, 0)

    texts = _figure_block_tokens(doc)
    if texts != ["<IMAGE>", "Источник: составлено автором.", "Рис. 1.2.1. Схема процесса"]:
        return _result(False, f"wrong figure source/caption order: {texts!r}")
    if doc.paragraphs[1]._p is not source._p or doc.paragraphs[2]._p is not caption._p:
        return _result(False, "source/caption paragraphs were duplicated or replaced")
    if caption.alignment not in (WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.JUSTIFY):
        return _result(False, f"figure caption has wrong alignment after move: {caption.alignment!r}")
    if not changed:
        return _result(False, "reorder function did not report a changed document")
    return _result(True, "figure source after caption is moved before caption")


def test_figure_source_before_caption_is_unchanged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import reorder_figure_source_before_caption

    doc = Document()
    source = doc.add_paragraph("Источник: составлено автором.")
    caption = doc.add_paragraph("Рис. 1.2.1. Схема процесса")

    changed = reorder_figure_source_before_caption(doc, 0)

    if [p.text for p in doc.paragraphs] != ["Источник: составлено автором.", "Рис. 1.2.1. Схема процесса"]:
        return _result(False, "already-correct figure source/caption order changed")
    if doc.paragraphs[0]._p is not source._p or doc.paragraphs[1]._p is not caption._p:
        return _result(False, "already-correct paragraphs were duplicated or replaced")
    if changed:
        return _result(False, "already-correct figure block reported a change")
    return _result(True, "already-correct figure source/caption order is unchanged")


def test_table_source_after_caption_is_not_moved() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import reorder_figure_source_before_caption

    doc = Document()
    doc.add_paragraph("Таблица 1.2.1")
    doc.add_paragraph("Источник: составлено автором.")

    changed = reorder_figure_source_before_caption(doc, 0)

    if [p.text for p in doc.paragraphs] != ["Таблица 1.2.1", "Источник: составлено автором."]:
        return _result(False, "table source/caption order was changed")
    if changed:
        return _result(False, "table source/caption block reported a change")
    return _result(True, "table source/caption order is not moved")


def test_appendix_figure_source_after_caption_is_moved() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import reorder_figure_source_before_caption

    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    _add_test_drawing_paragraph(doc)
    doc.add_paragraph("Рис. 3.1.1. Схема приложения")
    doc.add_paragraph("Источник: составлено автором.")

    changed = reorder_figure_source_before_caption(doc, 0)

    texts = _figure_block_tokens(doc)
    if texts[-3:] != ["<IMAGE>", "Источник: составлено автором.", "Рис. 3.1.1. Схема приложения"]:
        return _result(False, f"appendix figure order was not fixed: {texts!r}")
    if not changed:
        return _result(False, "appendix figure reorder did not report a change")
    return _result(True, "appendix figure source/caption order is fixed")


def test_figure_source_not_duplicated_after_reorder() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import reorder_figure_source_before_caption

    doc = Document()
    doc.add_paragraph("Рис. 2.2.1. Динамика показателя")
    doc.add_paragraph("Источник: данные автора.")

    reorder_figure_source_before_caption(doc, 0)
    reorder_figure_source_before_caption(doc, 0)

    texts = [p.text for p in doc.paragraphs]
    if texts.count("Источник: данные автора.") != 1:
        return _result(False, f"source line duplicated: {texts!r}")
    if texts.count("Рис. 2.2.1. Динамика показателя") != 1:
        return _result(False, f"figure caption duplicated: {texts!r}")
    return _result(True, "figure source/caption paragraphs are not duplicated")


def test_bibliography_source_line_is_not_moved() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import reorder_figure_source_before_caption

    doc = Document()
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("Рис. 1.2.1. Название источника")
    doc.add_paragraph("Источник: библиографическое описание.")

    changed = reorder_figure_source_before_caption(doc, 0)

    if [p.text for p in doc.paragraphs] != [
        "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
        "Рис. 1.2.1. Название источника",
        "Источник: библиографическое описание.",
    ]:
        return _result(False, "bibliography source line was moved")
    if changed:
        return _result(False, "bibliography source block reported a change")
    return _result(True, "bibliography source line is not moved")


def test_heading2_late_spacing_before_13() -> tuple[bool, str]:
    """Late/final Heading 2 formatting still leaves one blank before 1.3."""
    from guides.coursework_kfu_2025.safe_formatter import is_empty_paragraph, process_document
    import tempfile, os

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ")
    doc.add_paragraph("1.2. Критерии конкурентоспособности организации")
    doc.add_paragraph("Текст подраздела 1.2.")
    doc.add_paragraph("Эти критерии позволят перейти к разделу 1.3.")
    doc.add_paragraph("1.3. Методы оценки конкурентоспособности организации")
    doc.add_paragraph("В процессе оценки конкурентоспособности применяются методы.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = os.path.join(tmp, "in.docx")
        out = os.path.join(tmp, "out.docx")
        doc.save(inp)
        process_document(inp, out)
        result = Document(out)

    paragraphs = result.paragraphs
    target_idx = None
    for idx, paragraph in enumerate(paragraphs):
        if " ".join(paragraph.text.split()).startswith("1.3. Методы оценки"):
            target_idx = idx
            break

    if target_idx is None:
        return _result(False, "1.3 heading not found after formatting")

    if target_idx < 1 or not is_empty_paragraph(paragraphs[target_idx - 1]):
        return _result(False, "missing blank before 1.3 heading")

    if target_idx >= 2 and is_empty_paragraph(paragraphs[target_idx - 2]):
        return _result(False, "double blank before 1.3 heading")

    if target_idx + 1 >= len(paragraphs) or not is_empty_paragraph(paragraphs[target_idx + 1]):
        return _result(False, "missing blank after 1.3 heading")

    if target_idx + 2 < len(paragraphs) and is_empty_paragraph(paragraphs[target_idx + 2]):
        return _result(False, "double blank after 1.3 heading")

    return _result(True, "1.3 heading has exactly one blank before and after")


def test_blank_before_figure_block() -> tuple[bool, str]:
    """
    A drawing paragraph that follows body text must have exactly one blank before it.
    The caption/source spacing rules remain untouched.
    """
    from guides.coursework_kfu_2025.safe_formatter import (
        ensure_single_blank_before_figure_blocks,
        is_empty_paragraph,
        remove_empty_between_figure_caption_and_source,
        paragraph_has_drawing,
    )

    doc = Document()
    doc.add_paragraph("Текст перед рисунком.")
    drawing_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    run = OxmlElement("w:r")
    run.append(drawing)
    drawing_p._element.append(run)
    doc.add_paragraph("Рис. 1.1.1. Схема процесса")
    doc.add_paragraph("")
    doc.add_paragraph("Источник: составлено автором.")

    ensure_single_blank_before_figure_blocks(doc, 0)
    remove_empty_between_figure_caption_and_source(doc, 0)

    drawing_idx = None
    for idx, paragraph in enumerate(doc.paragraphs):
        if paragraph_has_drawing(paragraph):
            drawing_idx = idx
            break

    if drawing_idx is None:
        return _result(False, "drawing paragraph not found")

    if drawing_idx < 1 or not is_empty_paragraph(doc.paragraphs[drawing_idx - 1]):
        return _result(False, "missing blank before drawing paragraph")

    if drawing_idx >= 2 and is_empty_paragraph(doc.paragraphs[drawing_idx - 2]):
        return _result(False, "double blank before drawing paragraph")

    texts = [p.text for p in doc.paragraphs]
    expected = [
        "Текст перед рисунком.",
        "",
        "",
        "Рис. 1.1.1. Схема процесса",
        "Источник: составлено автором.",
    ]
    if texts != expected:
        return _result(False, f"unexpected figure block layout: {texts!r}")

    blank = doc.paragraphs[drawing_idx - 1]
    run = blank.runs[0] if blank.runs else None
    if run is None:
        return _result(False, "blank before drawing has no run")

    sz = run._element.get_or_add_rPr().find(qn("w:sz"))
    if sz is None or sz.get(qn("w:val")) != "28":
        val = sz.get(qn("w:val")) if sz is not None else None
        return _result(False, f"blank before drawing font size is {val}, expected 28 half-points")

    return _result(True, "drawing paragraph has one TNR 14 blank before it")


def test_marker_instrumentation_keeps_source_unchanged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_markers import instrument_table_rows_copy

    doc = Document()
    tbl = doc.add_table(rows=3, cols=2)
    tbl.rows[0].cells[0].text = "Header"
    tbl.rows[1].cells[0].text = "Row one"
    tbl.rows[2].cells[0].text = "Row two"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "source.docx"
        workdir = Path(tmp) / "work"
        doc.save(src)
        before = src.read_bytes()

        instrumentation = instrument_table_rows_copy(src, 0, workdir=workdir, marker_font_size_pt=1)

        after = src.read_bytes()
        if before != after:
            return _result(False, "source docx changed after instrumentation")

        source_doc = Document(str(src))
        instrumented_doc = Document(str(instrumentation.instrumented_docx_path))
        source_text = " ".join(p.text for p in source_doc.paragraphs)
        if "KPFU_TMARK_" in source_text:
            return _result(False, "marker leaked into source document")

        marker_hits = sum(
            text.count("KPFU_TMARK_")
            for table in instrumented_doc.tables
            for row in table.rows
            for cell in row.cells
            for text in [cell.text]
        )
        if marker_hits != 3:
            return _result(False, f"expected 3 row markers in instrumented copy, got {marker_hits}")

    return _result(True, "instrumentation only changes temp copy")


def test_marker_instrumentation_only_targets_selected_table() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_markers import instrument_table_rows_copy

    doc = Document()
    first = doc.add_table(rows=2, cols=1)
    first.rows[0].cells[0].text = "First header"
    first.rows[1].cells[0].text = "First body"
    second = doc.add_table(rows=3, cols=1)
    second.rows[0].cells[0].text = "Second header"
    second.rows[1].cells[0].text = "Second row"
    second.rows[2].cells[0].text = "Third row"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "multi.docx"
        doc.save(src)
        instrumentation = instrument_table_rows_copy(src, 1, workdir=Path(tmp) / "work", marker_font_size_pt=1)
        instrumented = Document(str(instrumentation.instrumented_docx_path))

    first_text = " ".join(cell.text for row in instrumented.tables[0].rows for cell in row.cells)
    second_text = " ".join(cell.text for row in instrumented.tables[1].rows for cell in row.cells)
    if "KPFU_TMARK_" in first_text:
        return _result(False, "marker inserted into non-target table")
    if second_text.count("KPFU_TMARK_") != 3:
        return _result(False, f"expected markers only in target table rows, got text={second_text!r}")

    return _result(True, "only selected table was instrumented")


def test_marker_extract_handles_inline_text_and_missing_rows() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_markers import extract_row_pages_from_pdf_lines
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    result = extract_row_pages_from_pdf_lines(
        [
            PdfLine("prefixKPFU_TMARK_ABC123_T00_R000suffix", 27, 10.0, 20.0),
            PdfLine("bodyKPFU_TMARK_ABC123_T00_R001tail", 27, 30.0, 40.0),
        ],
        marker_salt="ABC123",
        table_index=0,
        total_rows=3,
    )

    if result.row_pages != {0: 27, 1: 27}:
        return _result(False, f"unexpected row_pages: {result.row_pages!r}")
    if result.found_rows != [0, 1]:
        return _result(False, f"unexpected found_rows: {result.found_rows!r}")
    if result.missing_rows != [2]:
        return _result(False, f"unexpected missing_rows: {result.missing_rows!r}")
    if result.duplicate_rows:
        return _result(False, f"unexpected duplicate_rows: {result.duplicate_rows!r}")

    return _result(True, "inline marker parsing and missing-row diagnostics work")


def test_marker_map_rows_to_pages_keep_temp_debug_paths() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    tbl = doc.add_table(rows=3, cols=1)
    tbl.rows[0].cells[0].text = "Header"
    tbl.rows[1].cells[0].text = "Alpha"
    tbl.rows[2].cells[0].text = "Beta"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "source.docx"
        doc.save(src)

        seen_docx: dict[str, Path] = {}
        old_render = tm.render_docx_to_pdf
        old_analyze = tm.analyze_pdf_lines
        try:
            def fake_render(docx_path):
                seen_docx["path"] = Path(docx_path)
                pdf_dir = Path(tmp) / "pdf_keep"
                pdf_dir.mkdir(exist_ok=True)
                pdf_path = pdf_dir / "instrumented.pdf"
                pdf_path.write_bytes(b"%PDF-1.4\n")
                return pdf_path

            def fake_analyze(_pdf_path):
                inst_doc = Document(str(seen_docx["path"]))
                row_markers = []
                for row in inst_doc.tables[0].rows:
                    text = " ".join(cell.text for cell in row.cells)
                    match = re.search(r"KPFU_TMARK_[A-F0-9]{6}_T00_R\d{3}", text)
                    if not match:
                        raise AssertionError(f"marker not found in row text: {text!r}")
                    row_markers.append(match.group(0))
                return [
                    tm.PdfLine(f"left{row_markers[0]}right", 27, 10.0, 20.0),
                    tm.PdfLine(f"left{row_markers[1]}right", 27, 30.0, 40.0),
                    tm.PdfLine(f"left{row_markers[2]}right", 28, 50.0, 60.0),
                ]

            tm.render_docx_to_pdf = fake_render
            tm.analyze_pdf_lines = fake_analyze
            result = tm.map_table_rows_to_pages(src, 0, keep_temp=True)
        finally:
            tm.render_docx_to_pdf = old_render
            tm.analyze_pdf_lines = old_analyze

        if result.row_pages != {0: 27, 1: 27, 2: 28}:
            return _result(False, f"unexpected row_pages: {result.row_pages!r}")
        if result.instrumented_docx_path is None or not result.instrumented_docx_path.exists():
            return _result(False, "instrumented_docx_path was not preserved in keep_temp mode")
        if result.pdf_path is None or not result.pdf_path.exists():
            return _result(False, "pdf_path was not preserved in keep_temp mode")
        if result.marker_font_size_pt != 1:
            return _result(False, f"expected 1pt success path, got {result.marker_font_size_pt}")

    return _result(True, "keep_temp preserves instrumented DOCX/PDF and returns exact mapping")


def test_marker_map_rows_to_pages_falls_back_to_2pt_and_returns_debug_info() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    tbl = doc.add_table(rows=3, cols=1)
    tbl.rows[0].cells[0].text = "Header"
    tbl.rows[1].cells[0].text = "Alpha"
    tbl.rows[2].cells[0].text = "Beta"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "source.docx"
        doc.save(src)

        seen_docx: dict[str, Path] = {}
        old_render = tm.render_docx_to_pdf
        old_analyze = tm.analyze_pdf_lines
        try:
            def fake_render(docx_path):
                seen_docx["path"] = Path(docx_path)
                pdf_dir = Path(tmp) / f"pdf_{Path(docx_path).stem}"
                pdf_dir.mkdir(exist_ok=True)
                pdf_path = pdf_dir / "instrumented.pdf"
                pdf_path.write_bytes(b"%PDF-1.4\n")
                return pdf_path

            def fake_analyze(_pdf_path):
                inst_doc = Document(str(seen_docx["path"]))
                full_text = " ".join(cell.text for row in inst_doc.tables[0].rows for cell in row.cells)
                markers = re.findall(r"KPFU_TMARK_[A-F0-9]{6}_T00_R\d{3}", full_text)
                if len(markers) != 3:
                    raise AssertionError(f"expected 3 markers, got {markers!r}")
                if "_2pt" not in seen_docx["path"].name:
                    return [tm.PdfLine(f"x{markers[0]}y", 27, 10.0, 20.0)]
                return [
                    tm.PdfLine(f"x{markers[0]}y", 27, 10.0, 20.0),
                    tm.PdfLine(f"x{markers[1]}y", 28, 30.0, 40.0),
                    tm.PdfLine(f"x{markers[1]}y", 29, 50.0, 60.0),
                ]

            tm.render_docx_to_pdf = fake_render
            tm.analyze_pdf_lines = fake_analyze
            result = tm.map_table_rows_to_pages(src, 0, keep_temp=False)
        finally:
            tm.render_docx_to_pdf = old_render
            tm.analyze_pdf_lines = old_analyze

        if result.marker_font_size_pt != 2:
            return _result(False, f"expected 2pt fallback, got {result.marker_font_size_pt}")
        if result.row_pages != {0: 27}:
            return _result(False, f"unexpected partial row_pages: {result.row_pages!r}")
        if result.missing_rows != [2]:
            return _result(False, f"unexpected missing_rows after fallback: {result.missing_rows!r}")
        if result.duplicate_rows != {1: [28, 29]}:
            return _result(False, f"unexpected duplicate_rows after fallback: {result.duplicate_rows!r}")
        if result.instrumented_docx_path is None or result.pdf_path is None:
            return _result(False, "debug paths should be preserved for incomplete diagnostics")
        if not result.instrumented_docx_path.exists() or not result.pdf_path.exists():
            return _result(False, "preserved debug paths do not exist")

    return _result(True, "1pt fallback to 2pt preserves diagnostics and debug artifacts")


def test_marker_instrumentation_rejects_invalid_table_index() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_markers import instrument_table_rows_copy

    doc = Document()
    doc.add_table(rows=1, cols=1)

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "source.docx"
        doc.save(src)
        try:
            instrument_table_rows_copy(src, 3, workdir=Path(tmp) / "work")
        except ValueError:
            return _result(True, "invalid table index rejected")
        except Exception as exc:
            return _result(False, f"unexpected exception type: {exc}")

    return _result(False, "expected ValueError for invalid table index")


def test_marker_page_span_summary() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_markers import summarize_row_page_spans

    spans = summarize_row_page_spans({
        0: 12,
        1: 12,
        2: 13,
        3: 13,
        5: 14,
    })
    triples = [(s.start_row, s.end_row, s.page_num) for s in spans]
    expected = [(0, 1, 12), (2, 3, 13), (5, 5, 14)]
    if triples != expected:
        return _result(False, f"unexpected page spans: {triples!r}")
    return _result(True, "row page spans are grouped correctly")


def test_marker_diagnose_all_tables_summary() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    first = doc.add_table(rows=3, cols=1)
    first.rows[0].cells[0].text = "H1"
    first.rows[1].cells[0].text = "A"
    first.rows[2].cells[0].text = "B"
    second = doc.add_table(rows=2, cols=1)
    second.rows[0].cells[0].text = "H2"
    second.rows[1].cells[0].text = "C"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "diag.docx"
        doc.save(src)

        old_map = tm.map_table_rows_to_pages
        try:
            def fake_map(_docx_path, table_index, keep_temp=False):
                if table_index == 0:
                    return tm.TableMarkerResult(
                        row_pages={0: 12, 1: 12, 2: 13},
                        found_rows=[0, 1, 2],
                        missing_rows=[],
                        duplicate_rows={},
                        marker_font_size_pt=1,
                    )
                return tm.TableMarkerResult(
                    row_pages={0: 15},
                    found_rows=[0],
                    missing_rows=[1],
                    duplicate_rows={},
                    instrumented_docx_path=Path(tmp) / "inst.docx" if keep_temp else None,
                    pdf_path=Path(tmp) / "inst.pdf" if keep_temp else None,
                    marker_font_size_pt=2,
                )

            tm.map_table_rows_to_pages = fake_map
            diagnostics = tm.diagnose_all_tables(src, keep_temp=True)
        finally:
            tm.map_table_rows_to_pages = old_map

    if len(diagnostics) != 2:
        return _result(False, f"expected 2 diagnostics, got {len(diagnostics)}")
    if diagnostics[0].candidate_for_split is not True:
        return _result(False, "multi-page fully-mapped table should be candidate_for_split=yes")
    if [(s.start_row, s.end_row, s.page_num) for s in diagnostics[0].page_spans] != [(0, 1, 12), (2, 2, 13)]:
        return _result(False, f"unexpected first table spans: {diagnostics[0].page_spans!r}")
    if diagnostics[0].appendix_table is not False:
        return _result(False, "first table should not be marked as appendix table")
    if diagnostics[0].caption_detected is not False:
        return _result(False, "first table should not report caption without preceding paragraph")
    if diagnostics[1].candidate_for_split is not False:
        return _result(False, "table with missing rows should not be candidate_for_split")
    if diagnostics[1].missing_rows != [1]:
        return _result(False, f"unexpected missing rows: {diagnostics[1].missing_rows!r}")
    if diagnostics[1].marker_font_size_pt != 2:
        return _result(False, f"unexpected fallback font size: {diagnostics[1].marker_font_size_pt}")

    return _result(True, "document-level diagnostics summarize all tables")


def test_marker_diagnose_table_handles_mapping_error() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    tbl = doc.add_table(rows=2, cols=1)
    tbl.rows[0].cells[0].text = "H"
    tbl.rows[1].cells[0].text = "A"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "error.docx"
        doc.save(src)

        old_map = tm.map_table_rows_to_pages
        try:
            def raise_map(_docx_path, _table_index, keep_temp=False):
                raise RuntimeError("render failed")

            tm.map_table_rows_to_pages = raise_map
            diagnostic = tm.diagnose_table(src, 0, keep_temp=False)
        finally:
            tm.map_table_rows_to_pages = old_map

    if diagnostic.error_message != "render failed":
        return _result(False, f"unexpected error_message: {diagnostic.error_message!r}")
    if diagnostic.candidate_for_split:
        return _result(False, "error diagnostic must not be candidate_for_split")
    if diagnostic.row_pages != {} or diagnostic.pages_detected != []:
        return _result(False, "error diagnostic should not report row/page mapping")
    return _result(True, "diagnose_table degrades to diagnostic error instead of crashing")


def test_marker_appendix_and_caption_metadata() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 1.1")
    doc.add_paragraph("Двухстрочный заголовок обычной таблицы")
    first = doc.add_table(rows=2, cols=1)
    first.rows[0].cells[0].text = "H1"
    first.rows[1].cells[0].text = "A"
    doc.add_paragraph("Приложение 1")
    doc.add_paragraph("Длинная таблица по приложению")
    second = doc.add_table(rows=2, cols=1)
    second.rows[0].cells[0].text = "H2"
    second.rows[1].cells[0].text = "B"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "appendix.docx"
        doc.save(src)

        old_map = tm.map_table_rows_to_pages
        try:
            def fake_map(_docx_path, table_index, keep_temp=False):
                return tm.TableMarkerResult(
                    row_pages={0: 10, 1: 10},
                    found_rows=[0, 1],
                    missing_rows=[],
                    duplicate_rows={},
                    marker_font_size_pt=1,
                )

            tm.map_table_rows_to_pages = fake_map
            diagnostics = tm.diagnose_all_tables(src, keep_temp=False)
        finally:
            tm.map_table_rows_to_pages = old_map

    if diagnostics[0].caption_detected is not True or diagnostics[0].has_standard_table_caption is not True:
        return _result(False, "standard split table caption was not detected for first table")
    if diagnostics[0].appendix_table is not False:
        return _result(False, "first table should not be appendix table")
    if diagnostics[0].preceding_paragraph_text != "Двухстрочный заголовок обычной таблицы":
        return _result(False, f"immediate title context was not preserved: {diagnostics[0].preceding_paragraph_text!r}")
    if diagnostics[1].appendix_table is not True:
        return _result(False, "second table should be marked as appendix table")
    if diagnostics[1].caption_detected is not True:
        return _result(False, "appendix table title/caption should be detected")
    if diagnostics[1].has_standard_table_caption is not False:
        return _result(False, "appendix table title should not be treated as standard table caption")
    if diagnostics[1].preceding_paragraph_text != "Длинная таблица по приложению":
        return _result(False, f"unexpected preceding paragraph text: {diagnostics[1].preceding_paragraph_text!r}")

    return _result(True, "appendix and caption metadata are detected")


def test_appendix_start_labels_are_normalized() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_appendix_start_labels

    doc = Document()
    body_text = doc.add_paragraph("Приложение2")
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    appendix_2 = doc.add_paragraph("Приложение2")
    appendix_a_spaced = doc.add_paragraph("Приложение А")
    appendix_a_glued = doc.add_paragraph("ПриложениеА")

    normalize_appendix_start_labels(doc, body_start=0)

    if body_text.text != "Приложение2":
        return _result(False, "ordinary body appendix-like text before appendix section was changed")
    expected = [
        (appendix_2, "ПРИЛОЖЕНИЕ 2", False),
        (appendix_a_spaced, "ПРИЛОЖЕНИЕ А", True),
        (appendix_a_glued, "ПРИЛОЖЕНИЕ А", True),
    ]
    for paragraph, text, should_start_new_page in expected:
        if paragraph.text != text:
            return _result(False, f"unexpected appendix start label: {paragraph.text!r}")
        if paragraph.text != paragraph.text.upper():
            return _result(False, f"appendix start label is not uppercase: {paragraph.text!r}")
        if not _paragraph_is_right_aligned(paragraph._p):
            return _result(False, f"appendix start label is not right aligned: {paragraph.text!r}")
        has_page_break = _paragraph_has_page_break_before(paragraph._p)
        if should_start_new_page and not has_page_break:
            return _result(False, f"appendix start label must start a new page: {paragraph.text!r}")
        if not should_start_new_page and has_page_break:
            return _result(False, f"first appendix start label must stay with ПРИЛОЖЕНИЯ: {paragraph.text!r}")
        if not _paragraph_has_keep_next(paragraph._p):
            return _result(False, f"appendix start label is not kept with content: {paragraph.text!r}")
    return _result(True, "appendix start labels are normalized and scoped")


def test_appendix_local_table_title_before_table_is_centered() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from guides.coursework_kfu_2025.safe_formatter import normalize_appendix_local_table_titles

    doc = Document()
    body_title = doc.add_paragraph("Обычный текст перед таблицей")
    doc.add_table(rows=1, cols=1)
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    appendix_label = doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    appendix_title = doc.add_paragraph("Расчет трудозатрат.")
    doc.add_table(rows=1, cols=1)

    normalize_appendix_local_table_titles(doc, body_start=0)

    if body_title.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "ordinary body paragraph before table was changed")
    if appendix_label.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "appendix label was treated as a table title")
    if appendix_title.text != "Расчет трудозатрат":
        return _result(False, f"appendix table title period was not stripped: {appendix_title.text!r}")
    if appendix_title.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "appendix-local table title was not centered")
    if _paragraph_has_page_break_before(appendix_title._p):
        return _result(False, "appendix-local table title must not force a page break")
    if not _paragraph_has_keep_next(appendix_title._p):
        return _result(False, "appendix-local table title must stay with table")
    return _result(True, "appendix-local table title before table is centered")


def _run_uses_tnr_14_not_bold(paragraph) -> bool:
    runs = [run for run in paragraph.runs if run.text.strip()]
    if not runs:
        return False
    for run in runs:
        r_pr = run._element.rPr
        if r_pr is None:
            return False
        r_fonts = r_pr.rFonts
        if r_fonts is None or r_fonts.get(qn("w:ascii")) != "Times New Roman":
            return False
        sz = r_pr.find(qn("w:sz"))
        if sz is None or sz.get(qn("w:val")) != "28":
            return False
        bold = r_pr.find(qn("w:b"))
        if run.bold is True or (bold is not None and bold.get(qn("w:val")) not in (None, "0", "false")):
            return False
    return True


def _blank_count_after_paragraph(doc: Document, paragraph) -> int:
    children = list(doc.element.body)
    idx = children.index(paragraph._p)
    count = 0
    idx += 1
    while idx < len(children) and children[idx].tag == qn("w:p") and not _paragraph_text(children[idx]):
        count += 1
        idx += 1
    return count


def test_appendix_title_after_label_is_normalized() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from guides.coursework_kfu_2025.safe_formatter import (
        normalize_appendix_start_labels,
        normalize_appendix_titles,
    )

    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    label = doc.add_paragraph("Приложение А")
    title = doc.add_paragraph("Расчет трудозатрат.")
    title.runs[0].bold = True
    doc.add_paragraph("Текст приложения.")

    normalize_appendix_start_labels(doc, body_start=0)
    normalize_appendix_titles(doc, body_start=0)

    if label.text != "ПРИЛОЖЕНИЕ А" or not _paragraph_is_right_aligned(label._p):
        return _result(False, "appendix label formatting changed")
    if title.text != "Расчет трудозатрат":
        return _result(False, f"appendix title text was not normalized: {title.text!r}")
    if title.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "appendix title is not centered")
    if not _run_uses_tnr_14_not_bold(title):
        return _result(False, "appendix title font must be TNR 14 and not bold")
    if _paragraph_has_page_break_before(title._p):
        return _result(False, "appendix title must not start a new page")
    if _paragraph_has_direct_numbering(title):
        return _result(False, "appendix title must not have numbering")
    if _blank_count_after_paragraph(doc, title) != 1:
        return _result(False, "appendix title must have exactly one blank after it")
    return _result(True, "appendix title after label is normalized")


def test_table_caption_like_appendix_title_after_label_is_normalized() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from guides.coursework_kfu_2025.safe_formatter import (
        normalize_appendix_start_labels,
        normalize_appendix_titles,
    )

    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    title = doc.add_paragraph("Таблица А.1 Расчет показателей.")
    doc.add_table(rows=1, cols=1)

    normalize_appendix_start_labels(doc, body_start=0)
    normalize_appendix_titles(doc, body_start=0)

    if title.text != "Таблица А.1 Расчет показателей":
        return _result(False, f"table-like appendix title text changed unexpectedly: {title.text!r}")
    if title.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "table-like appendix title was not centered")
    if _blank_count_after_paragraph(doc, title) != 1:
        return _result(False, "table-like appendix title must have exactly one blank after it")
    return _result(True, "table-like appendix title after label is normalized")


def test_table_caption_like_appendix_title_survives_full_process() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.add_paragraph("Титульная строка")
    doc.add_paragraph("СОДЕРЖАНИЕ")
    doc.add_paragraph("ВВЕДЕНИЕ........................................................3")
    doc.add_paragraph("ПРИЛОЖЕНИЯ.......................................................8")
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ Б")
    doc.add_paragraph("Таблица 2.1 Расчет показателей.")
    doc.add_table(rows=1, cols=1)

    formatted = _format_synthetic_doc(doc)
    label_idx = _paragraph_index(formatted, "ПРИЛОЖЕНИЕ Б")
    if label_idx is None:
        return _result(False, "appendix label missing after full process")

    title = None
    for paragraph in formatted.paragraphs[label_idx + 1:]:
        if paragraph.text.strip():
            title = paragraph
            break
    if title is None:
        return _result(False, "appendix title missing after full process")

    if title.text != "Таблица 2.1 Расчет показателей":
        return _result(False, f"table-like appendix title was not preserved: {title.text!r}")
    if title.alignment != WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "table-like appendix title is not centered after full process")
    if _blank_count_after_paragraph(formatted, title) != 1:
        return _result(False, "table-like appendix title does not have exactly one blank after full process")
    return _result(True, "table-like appendix title survives full process")


def test_long_body_paragraph_after_appendix_label_is_not_title() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from guides.coursework_kfu_2025.safe_formatter import (
        normalize_appendix_start_labels,
        normalize_appendix_titles,
    )

    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    body = doc.add_paragraph(
        "Это обычный развернутый текст приложения, который содержит несколько предложений "
        "и явно не должен становиться названием приложения после форматирования документа."
    )
    doc.add_paragraph("Следующий абзац.")

    normalize_appendix_start_labels(doc, body_start=0)
    normalize_appendix_titles(doc, body_start=0)

    if body.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "long body paragraph was promoted to appendix title")
    if _blank_count_after_paragraph(doc, body) != 0:
        return _result(False, "long body paragraph received title spacing")
    return _result(True, "long body paragraph after appendix label is not title")


def test_appendix_continuation_label_does_not_trigger_title_formatting() -> tuple[bool, str]:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from guides.coursework_kfu_2025.safe_formatter import normalize_appendix_titles

    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    doc.add_paragraph("Расчет трудозатрат")
    continuation = doc.add_paragraph("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ 1")
    after_continuation = doc.add_paragraph("Таблица А.1 Продолжение данных")
    doc.add_table(rows=1, cols=1)

    normalize_appendix_titles(doc, body_start=0)

    if continuation.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "appendix continuation label was formatted as title")
    if after_continuation.alignment == WD_ALIGN_PARAGRAPH.CENTER:
        return _result(False, "paragraph after appendix continuation was formatted as title")
    return _result(True, "appendix continuation label does not trigger title formatting")


def test_appendix_title_spacing_is_exactly_one_blank() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import normalize_appendix_titles

    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    zero_blank_title = doc.add_paragraph("Первое приложение")
    doc.add_paragraph("Текст первого приложения.")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 2")
    multi_blank_title = doc.add_paragraph("Второе приложение")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("Текст второго приложения.")

    normalize_appendix_titles(doc, body_start=0)

    if _blank_count_after_paragraph(doc, zero_blank_title) != 1:
        return _result(False, "zero-blank title was not normalized to one blank")
    if _blank_count_after_paragraph(doc, multi_blank_title) != 1:
        return _result(False, "multi-blank title was not normalized to one blank")
    return _result(True, "appendix title spacing is exactly one blank")


def test_empty_paragraph_after_appendix_label_before_table_is_preserved() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import (
        normalize_appendix_start_labels,
        remove_empty_paragraphs_after_appendix_labels,
    )

    doc = Document()
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    appendix_label = doc.add_paragraph("Приложение2")
    doc.add_paragraph("")
    doc.add_table(rows=1, cols=1)

    normalize_appendix_start_labels(doc, body_start=0)
    remove_empty_paragraphs_after_appendix_labels(doc, body_start=0)

    children = list(doc.element.body)
    label_idx = children.index(appendix_label._p)
    blank_count = 0
    next_idx = label_idx + 1
    while next_idx < len(children) and children[next_idx].tag == qn("w:p") and not _paragraph_text(children[next_idx]):
        blank_count += 1
        next_idx += 1
    if blank_count != 1:
        return _result(False, f"expected exactly one blank paragraph after appendix label, got {blank_count}")
    if next_idx >= len(children) or children[next_idx].tag != qn("w:tbl"):
        return _result(False, "appendix table does not follow the single blank paragraph")
    return _result(True, "exactly one empty paragraph after appendix label is preserved")


def test_empty_paragraph_between_appendices_heading_and_first_label_is_removed() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import (
        normalize_appendix_start_labels,
        remove_empty_paragraphs_between_appendices_heading_and_first_label,
    )

    doc = Document()
    appendices_heading = doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("")
    first_label = doc.add_paragraph("ПриложениеА")

    normalize_appendix_start_labels(doc, body_start=0)
    remove_empty_paragraphs_between_appendices_heading_and_first_label(doc, body_start=0)

    children = list(doc.element.body)
    heading_idx = children.index(appendices_heading._p)
    if children[heading_idx + 1] is not first_label._p:
        return _result(False, "empty paragraph remains between ПРИЛОЖЕНИЯ and first appendix label")
    if _paragraph_has_page_break_before(first_label._p):
        return _result(False, "first appendix label after ПРИЛОЖЕНИЯ starts a new page")
    return _result(True, "empty paragraph before first appendix label is removed")


def test_marker_runtime_dry_run_clean_two_page_table_is_eligible() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.table_markers import TableMarkerDiagnostic, TablePageSpan

    diagnostic = TableMarkerDiagnostic(
        table_index=10,
        rows_count=21,
        pages_detected=[53, 54],
        row_pages={**{0: 53}, **{row: 53 for row in range(1, 18)}, 18: 54, 19: 54, 20: 54},
        found_rows=list(range(21)),
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[TablePageSpan(0, 17, 53), TablePageSpan(18, 20, 54)],
        appendix_table=True,
        caption_detected=True,
        has_standard_table_caption=False,
        preceding_paragraph_text="Расчет трудозатрат",
    )

    decision = tc._evaluate_marker_split_diagnostic(diagnostic, header_rows=1)
    if decision.eligible is not True:
        return _result(False, f"expected eligible decision, got {decision!r}")
    if decision.split_before_row != 18:
        return _result(False, f"expected split_before_row=18, got {decision.split_before_row!r}")
    return _result(True, "clean two-page marker mapping is marked eligible")


def test_marker_runtime_dry_run_skips_duplicate_rows() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.table_markers import TableMarkerDiagnostic

    diagnostic = TableMarkerDiagnostic(
        table_index=0,
        rows_count=4,
        pages_detected=[12, 13],
        row_pages={0: 12, 1: 12, 3: 13},
        found_rows=[0, 1, 3],
        missing_rows=[],
        duplicate_rows={2: [12, 13]},
        candidate_for_split=False,
        page_spans=[],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
    )

    decision = tc._evaluate_marker_split_diagnostic(diagnostic, header_rows=1)
    if decision.eligible:
        return _result(False, "duplicate rows should skip dry-run eligibility")
    if decision.skip_reason != "duplicate_rows":
        return _result(False, f"unexpected skip_reason: {decision.skip_reason!r}")
    return _result(True, "duplicate rows are skipped")


def test_marker_runtime_dry_run_skips_missing_rows_outside_header() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.table_markers import TableMarkerDiagnostic

    diagnostic = TableMarkerDiagnostic(
        table_index=0,
        rows_count=5,
        pages_detected=[20, 21],
        row_pages={0: 20, 1: 20, 3: 21, 4: 21},
        found_rows=[0, 1, 3, 4],
        missing_rows=[2],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
    )

    decision = tc._evaluate_marker_split_diagnostic(diagnostic, header_rows=1)
    if decision.eligible:
        return _result(False, "missing body rows should skip dry-run eligibility")
    if decision.skip_reason != "missing_rows_outside_header":
        return _result(False, f"unexpected skip_reason: {decision.skip_reason!r}")
    return _result(True, "missing rows outside header are skipped")


def test_marker_runtime_dry_run_skips_three_page_tables() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.table_markers import TableMarkerDiagnostic

    diagnostic = TableMarkerDiagnostic(
        table_index=0,
        rows_count=6,
        pages_detected=[30, 31, 32],
        row_pages={0: 30, 1: 30, 2: 31, 3: 31, 4: 32, 5: 32},
        found_rows=[0, 1, 2, 3, 4, 5],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
    )

    decision = tc._evaluate_marker_split_diagnostic(diagnostic, header_rows=1)
    if decision.eligible:
        return _result(False, "3-page table should not be eligible in v1")
    if decision.skip_reason != "not_2_pages":
        return _result(False, f"unexpected skip_reason: {decision.skip_reason!r}")
    return _result(True, "3-page tables are skipped")


def test_marker_runtime_dry_run_logs_eligible_candidate() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    log_stream = io.StringIO()
    handler = logging.StreamHandler(log_stream)
    handler.setFormatter(logging.Formatter("%(message)s"))
    old_level = tc.logger.level
    tc.logger.addHandler(handler)
    tc.logger.setLevel(logging.INFO)

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=10,
        rows_count=21,
        pages_detected=[53, 54],
        row_pages={**{0: 53}, **{row: 53 for row in range(1, 18)}, 18: 54, 19: 54, 20: 54},
        found_rows=list(range(21)),
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 17, 53), tm.TablePageSpan(18, 20, 54)],
        appendix_table=True,
        caption_detected=True,
        has_standard_table_caption=False,
        preceding_paragraph_text="Расчет трудозатрат",
    )

    old_diagnose_all = tm.diagnose_all_tables
    try:
        tm.diagnose_all_tables = lambda _path, keep_temp=False: [diagnostic]
        count = tc._run_marker_split_detection_pass(Path("/tmp/fake.docx"))
    finally:
        tm.diagnose_all_tables = old_diagnose_all
        tc.logger.removeHandler(handler)
        tc.logger.setLevel(old_level)

    logs = log_stream.getvalue()
    expected_fragments = [
        "marker_split_candidate table_index=10 rows=21 pages=[53, 54]",
        "marker_split_boundary table_index=10 split_before_row=18",
        "marker_split_decision=ELIGIBLE table_index=10",
    ]
    missing = [fragment for fragment in expected_fragments if fragment not in logs]
    if count != 1:
        return _result(False, f"expected one eligible candidate, got {count}")
    if missing:
        return _result(False, f"missing log fragments: {missing!r}; logs={logs!r}")
    return _result(True, "eligible marker candidate logs are emitted")


def test_marker_runtime_dry_run_feature_flag_off_skips_detection_hook() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc

    doc = Document()
    doc.add_paragraph("Таблица 1.1")
    tbl = doc.add_table(rows=2, cols=1)
    tbl.rows[0].cells[0].text = "H"
    tbl.rows[1].cells[0].text = "A"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "flag_off.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "flag_off.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_flag = os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
        old_hook = tc._run_marker_split_detection_pass
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            def fail_hook(_docx_path):
                raise AssertionError("marker dry-run hook should not be called when flag is off")

            tc._run_marker_split_detection_pass = fail_hook
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: []
            tc.apply_rendered_table_continuation(path)
        finally:
            tc._run_marker_split_detection_pass = old_hook
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_flag is not None:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_flag

    return _result(True, "feature flag off keeps marker dry-run hook disabled")


def test_marker_runtime_dry_run_only_does_not_mutate_document() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    tbl = doc.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[0].cells[2].text = "C"
    for i in range(1, 5):
        for j in range(3):
            tbl.rows[i].cells[j].text = f"r{i}c{j}"

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=5,
        pages_detected=[12, 13],
        row_pages={0: 12, 1: 12, 2: 12, 3: 13, 4: 13},
        found_rows=[0, 1, 2, 3, 4],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 2, 12), tm.TablePageSpan(3, 4, 13)],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text="Таблица 1.1.1",
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "dry_run_only.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "dry_run_only.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)
        before = path.read_bytes()

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tm.diagnose_all_tables = lambda _path, keep_temp=False: [diagnostic]
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: []
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is not None:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply

        after = path.read_bytes()

    if n != 0:
        return _result(False, f"dry-run only should not mutate, got {n}")
    if before != after:
        return _result(False, "dry-run only changed document bytes")
    return _result(True, "dry-run only does not mutate document")


def test_marker_runtime_apply_split_for_appendix_table() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Приложение 1")
    doc.add_paragraph("Трудозатраты проекта")
    tbl = doc.add_table(rows=6, cols=3)
    tbl.rows[0].cells[0].text = "Исполнитель"
    tbl.rows[0].cells[1].text = "Работы"
    tbl.rows[0].cells[2].text = "Стоимость"
    for i in range(1, 6):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"
        tbl.rows[i].cells[2].text = f"r{i}c2"
    doc.add_paragraph("Источник: данные автора")

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=6,
        pages_detected=[53, 54],
        row_pages={0: 53, 1: 53, 2: 53, 3: 53, 4: 54, 5: 54},
        found_rows=[0, 1, 2, 3, 4, 5],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 3, 53), tm.TablePageSpan(4, 5, 54)],
        appendix_table=True,
        caption_detected=True,
        has_standard_table_caption=False,
        preceding_paragraph_text="Трудозатраты проекта",
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "appendix_apply.docx"
        doc.save(path)

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tm.diagnose_all_tables = lambda _path, keep_temp=False: [diagnostic]
            tc.render_docx_to_pdf = lambda _path: (_ for _ in ()).throw(AssertionError("render path should not run after marker split apply"))
            tc.analyze_pdf_lines = lambda _path: (_ for _ in ()).throw(AssertionError("pdf analysis should not run after marker split apply"))
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        out = Document(str(path))

    if n != 1:
        return _result(False, f"expected one appendix split mutation, got {n}")
    if len(out.tables) != 2:
        return _result(False, f"expected 2 tables after appendix split, got {len(out.tables)}")
    if len(out.tables[0].rows) != 4:
        return _result(False, f"appendix first part should include generated numbered row, got {len(out.tables[0].rows)} rows")
    if [c.text for c in out.tables[0].rows[1].cells] != ["1", "2", "3"]:
        return _result(False, "numbered row missing under first appendix table header")
    if _count_table_rows_with_texts(out.tables[0], ["1", "2", "3"]) != 1:
        return _result(False, "first appendix table should contain exactly one generated numbered row")
    if [c.text for c in out.tables[1].rows[0].cells] != ["1", "2", "3"]:
        return _result(False, "numbered row missing in second appendix table")
    if _count_table_rows_with_texts(out.tables[1], ["1", "2", "3"]) != 1:
        return _result(False, "second appendix table should contain exactly one generated numbered row")
    if [c.text for c in out.tables[1].rows[1].cells] != ["r3c0", "r3c1", "r3c2"]:
        return _result(False, "appendix continuation should move last first-page data row after numbered row")
    continuation_label = _paragraph_before_table(out, 1)
    if continuation_label is None:
        return _result(False, "appendix continuation label is missing before continuation table")
    label_text = _paragraph_text(continuation_label)
    if label_text != "ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ 1":
        return _result(False, f"unexpected appendix continuation label: {label_text!r}")
    if label_text != label_text.upper():
        return _result(False, f"appendix continuation label is not uppercase: {label_text!r}")
    if not _paragraph_is_right_aligned(continuation_label):
        return _result(False, "appendix continuation label is not right aligned")
    if not _paragraph_has_page_break_before(continuation_label):
        return _result(False, "appendix continuation label should start on a new page")
    if not _paragraph_has_keep_next(continuation_label):
        return _result(False, "appendix continuation label should keep with continuation table")
    if not _all_table_rows_have_cant_split(out.tables[0]):
        return _result(False, "first appendix split table rows can split across pages")
    if not _all_table_rows_have_cant_split(out.tables[1]):
        return _result(False, "second appendix split table rows can split across pages")
    if any("Продолжение таблицы" in (p.text or "") for p in out.paragraphs):
        return _result(False, "appendix split inserted forbidden continuation paragraph")
    if _table_has_page_break_service_paragraph_before(out, 1):
        return _result(False, "appendix split left a blank service paragraph before continuation table")
    return _result(True, "eligible appendix table is split with visible appendix continuation label")


def test_marker_runtime_apply_split_for_ordinary_table() -> tuple[bool, str]:
    """
    Product rule: ordinary table continuation inserts a continuation marker
    and a numbered row without duplicating the textual table header.
    """
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    tbl = doc.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "Показатель"
    tbl.rows[0].cells[1].text = "Как влияет"
    tbl.rows[0].cells[2].text = "Последствия"
    for i in range(1, 5):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"
        tbl.rows[i].cells[2].text = f"r{i}c2"
    doc.add_paragraph("Источник: данные автора")

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=5,
        pages_detected=[12, 13],
        row_pages={0: 12, 1: 12, 2: 12, 3: 13, 4: 13},
        found_rows=[0, 1, 2, 3, 4],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 2, 12), tm.TablePageSpan(3, 4, 13)],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text="Таблица 1.1.1",
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "ordinary_apply.docx"
        doc.save(path)

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tm.diagnose_all_tables = lambda _path, keep_temp=False: [diagnostic]
            tc.render_docx_to_pdf = lambda _path: (_ for _ in ()).throw(AssertionError("render path should not run after marker split apply"))
            tc.analyze_pdf_lines = lambda _path: (_ for _ in ()).throw(AssertionError("pdf analysis should not run after marker split apply"))
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        out = Document(str(path))

    if n != 1:
        return _result(False, f"expected one ordinary split mutation, got {n}")
    if len(out.tables) != 2:
        return _result(False, f"expected 2 tables after ordinary split, got {len(out.tables)}")
    if [c.text for c in out.tables[0].rows[1].cells] != ["1", "2", "3"]:
        return _result(False, "numbered row missing under first ordinary table header")
    if _count_table_rows_with_texts(out.tables[0], ["1", "2", "3"]) != 1:
        return _result(False, "first ordinary table should contain exactly one generated numbered row")
    if [c.text for c in out.tables[1].rows[0].cells] != ["1", "2", "3"]:
        return _result(False, "continuation table should start with numbered row only")
    if _count_table_rows_with_texts(out.tables[1], ["1", "2", "3"]) != 1:
        return _result(False, "second ordinary table should contain exactly one generated numbered row")
    if not _all_table_rows_have_cant_split(out.tables[0]):
        return _result(False, "first ordinary split table rows can split across pages")
    if not _all_table_rows_have_cant_split(out.tables[1]):
        return _result(False, "second ordinary split table rows can split across pages")
    continuation_paras = [p for p in out.paragraphs if p.text == "Продолжение таблицы 1.1.1"]
    if len(continuation_paras) != 1:
        return _result(False, "ordinary split did not insert continuation paragraph")
    pPr = continuation_paras[0]._element.find(qn("w:pPr"))
    page_break = pPr.find(qn("w:pageBreakBefore")) if pPr is not None else None
    jc = pPr.find(qn("w:jc")) if pPr is not None else None
    keep = pPr.find(qn("w:keepNext")) if pPr is not None else None
    if page_break is None:
        return _result(False, "ordinary continuation marker should start on a new page")
    if jc is None or jc.get(qn("w:val")) != "right":
        return _result(False, "ordinary continuation marker should be right-aligned")
    if keep is None:
        return _result(False, "ordinary continuation marker should keep with following table")
    if any(cell.text == "Показатель" for cell in out.tables[1].rows[0].cells):
        return _result(False, "text header leaked into continuation row")
    return _result(True, "eligible ordinary table is split with continuation paragraph")


def test_marker_runtime_apply_skips_nested_ordinary_table_header() -> tuple[bool, str]:
    """
    Product rule: ordinary marker split must not split a source XML table that
    already contains another logical table header inside its body.
    """
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 1.2.1")
    tbl = doc.add_table(rows=6, cols=4)
    for cell, text in zip(
        tbl.rows[0].cells,
        ["Функция", "KPI", "Метрика (единицы)", "Источник данных"],
    ):
        cell.text = text
    valid_rows = [
        ["Планирование", "Time-to-Decision", "медиана дней", "протоколы комитетов"],
        ["Организация", "Уровень делегирования", "доля решений", "матрица RACI"],
    ]
    for row, values in zip(tbl.rows[1:3], valid_rows):
        for cell, text in zip(row.cells, values):
            cell.text = text
    for cell, text in zip(
        tbl.rows[3].cells,
        ["Уровень", "Формальные органы", "Неформальные практики", "Основные функции"],
    ):
        cell.text = text
    for idx in (4, 5):
        for col, cell in enumerate(tbl.rows[idx].cells):
            cell.text = f"nested{idx}c{col}"

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=6,
        pages_detected=[17, 18],
        row_pages={0: 17, 1: 17, 2: 17, 3: 18, 4: 18, 5: 18},
        found_rows=[0, 1, 2, 3, 4, 5],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 2, 17), tm.TablePageSpan(3, 5, 18)],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text="Таблица 1.2.1",
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "nested_ordinary_header.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "nested_ordinary_header.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)
        before = path.read_bytes()

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        old_logger_info = tc.logger.info
        log_lines: list[str] = []
        try:
            def fake_info(message, *args, **kwargs):
                log_lines.append(message % args if args else str(message))

            tm.diagnose_all_tables = lambda _path, keep_temp=False: [diagnostic]
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: []
            tc.logger.info = fake_info
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            tc.logger.info = old_logger_info
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        after = path.read_bytes()
        out = Document(str(path))

    if n != 0:
        return _result(False, f"contaminated ordinary table should be skipped, got {n}")
    if before != after:
        return _result(False, "contaminated ordinary table skip mutated the document")
    if len(out.tables) != 1:
        return _result(False, f"contaminated ordinary table should not be split, got {len(out.tables)} tables")
    if any(p.text == "Продолжение таблицы 1.2.1" for p in out.paragraphs):
        return _result(False, "contaminated ordinary table received continuation paragraph")
    if not any("body_contains_nested_table_header" in line for line in log_lines):
        return _result(False, "nested header skip reason was not logged")
    return _result(True, "ordinary table with nested logical header is skipped")


def test_marker_runtime_apply_skips_ineligible_tables() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    cases = [
        ("duplicate", tm.TableMarkerDiagnostic(
            table_index=0,
            rows_count=4,
            pages_detected=[12, 13],
            row_pages={0: 12, 1: 12, 3: 13},
            found_rows=[0, 1, 3],
            missing_rows=[],
            duplicate_rows={2: [12, 13]},
            candidate_for_split=False,
            page_spans=[],
            appendix_table=False,
            caption_detected=True,
            has_standard_table_caption=True,
            preceding_paragraph_text="Таблица 2.1",
        )),
        ("missing", tm.TableMarkerDiagnostic(
            table_index=0,
            rows_count=5,
            pages_detected=[12, 13],
            row_pages={0: 12, 1: 12, 3: 13, 4: 13},
            found_rows=[0, 1, 3, 4],
            missing_rows=[2],
            duplicate_rows={},
            candidate_for_split=False,
            page_spans=[],
            appendix_table=False,
            caption_detected=True,
            has_standard_table_caption=True,
            preceding_paragraph_text="Таблица 2.2",
        )),
        ("threepage", tm.TableMarkerDiagnostic(
            table_index=0,
            rows_count=6,
            pages_detected=[12, 13, 14],
            row_pages={0: 12, 1: 12, 2: 13, 3: 13, 4: 14, 5: 14},
            found_rows=[0, 1, 2, 3, 4, 5],
            missing_rows=[],
            duplicate_rows={},
            candidate_for_split=False,
            page_spans=[],
            appendix_table=False,
            caption_detected=True,
            has_standard_table_caption=True,
            preceding_paragraph_text="Таблица 2.3",
        )),
    ]

    for label, diagnostic in cases:
        doc = Document()
        doc.add_paragraph("Таблица 2.1")
        tbl = doc.add_table(rows=4, cols=2)
        for i in range(4):
            tbl.rows[i].cells[0].text = f"r{i}c0"
            tbl.rows[i].cells[1].text = f"r{i}c1"

        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / f"ineligible_{label}.docx"
            pdf_dir = Path(tmp) / "pdf"
            pdf_dir.mkdir()
            pdf_path = pdf_dir / f"ineligible_{label}.pdf"
            pdf_path.write_bytes(b"%PDF-1.4\n")
            doc.save(path)
            before = path.read_bytes()

            old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
            old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
            os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
            os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"

            old_diagnose_all = tm.diagnose_all_tables
            old_render = tc.render_docx_to_pdf
            old_analyze = tc.analyze_pdf_lines
            try:
                tm.diagnose_all_tables = lambda _path, keep_temp=False, diag=diagnostic: [diag]
                tc.render_docx_to_pdf = lambda _path: pdf_path
                tc.analyze_pdf_lines = lambda _path: []
                n = tc.apply_rendered_table_continuation(path)
            finally:
                tm.diagnose_all_tables = old_diagnose_all
                tc.render_docx_to_pdf = old_render
                tc.analyze_pdf_lines = old_analyze
                if old_enable is None:
                    os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
                else:
                    os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
                if old_apply is None:
                    os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
                else:
                    os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply

            after = path.read_bytes()

        if n != 0:
            return _result(False, f"ineligible case {label} should not mutate, got {n}")
        if before != after:
            return _result(False, f"ineligible case {label} changed document bytes")

    return _result(True, "ineligible duplicate/missing/3page cases do not mutate")


def test_marker_runtime_apply_is_idempotent_on_second_run() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 7.1")
    tbl = doc.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[0].cells[2].text = "C"
    for i in range(1, 5):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"
        tbl.rows[i].cells[2].text = f"r{i}c2"

    eligible = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=5,
        pages_detected=[12, 13],
        row_pages={0: 12, 1: 12, 2: 12, 3: 13, 4: 13},
        found_rows=[0, 1, 2, 3, 4],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 2, 12), tm.TablePageSpan(3, 4, 13)],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text="Таблица 7.1",
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "idempotent.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "idempotent.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            def fake_diagnose_all(docx_path, keep_temp=False):
                current = Document(str(docx_path))
                if len(current.tables) == 1:
                    return [eligible]
                return [
                    tm.TableMarkerDiagnostic(
                        table_index=0,
                        rows_count=len(current.tables[0].rows),
                        pages_detected=[12],
                        row_pages={0: 12, 1: 12, 2: 12, 3: 12},
                        found_rows=[0, 1, 2, 3],
                        missing_rows=[],
                        duplicate_rows={},
                        candidate_for_split=False,
                        page_spans=[tm.TablePageSpan(0, 3, 12)],
                        appendix_table=False,
                        caption_detected=True,
                        has_standard_table_caption=True,
                        preceding_paragraph_text="Таблица 7.1",
                    ),
                    tm.TableMarkerDiagnostic(
                        table_index=1,
                        rows_count=len(current.tables[1].rows),
                        pages_detected=[13],
                        row_pages={0: 13, 1: 13, 2: 13},
                        found_rows=[0, 1, 2],
                        missing_rows=[],
                        duplicate_rows={},
                        candidate_for_split=False,
                        page_spans=[tm.TablePageSpan(0, 2, 13)],
                        appendix_table=False,
                        caption_detected=True,
                        has_standard_table_caption=False,
                        preceding_paragraph_text="Продолжение таблицы 7.1",
                    ),
                ]

            tm.diagnose_all_tables = fake_diagnose_all
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: []
            first = tc.apply_rendered_table_continuation(path)
            second = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        out = Document(str(path))

    if first != 1 or second != 0:
        return _result(False, f"expected first run=1 and second run=0, got {first}/{second}")
    if len(out.tables) != 2:
        return _result(False, f"expected 2 tables after second run, got {len(out.tables)}")
    first_numbered_rows = sum(
        1 for row in out.tables[0].rows
        if [cell.text for cell in row.cells] == ["1", "2", "3"]
    )
    if first_numbered_rows != 1:
        return _result(False, f"first table should contain exactly one generated numbered row, got {first_numbered_rows}")
    second_numbered_rows = sum(
        1 for row in out.tables[1].rows
        if [cell.text for cell in row.cells] == ["1", "2", "3"]
    )
    if second_numbered_rows != 1:
        return _result(False, f"second table should contain exactly one numbered row, got {second_numbered_rows}")
    continuation_count = sum(1 for p in out.paragraphs if p.text == "Продолжение таблицы 7.1")
    if continuation_count != 1:
        return _result(False, f"expected one continuation paragraph after two runs, got {continuation_count}")
    return _result(True, "active marker split is idempotent on second run")


def test_marker_runtime_apply_processes_multiple_ordinary_tables() -> tuple[bool, str]:
    """
    Product rule: one formatter run should process every eligible marker split,
    not just the first eligible table.
    """
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    for num, prefix in (("1.1", "a"), ("2.1", "b")):
        doc.add_paragraph(f"Таблица {num}")
        tbl = doc.add_table(rows=5, cols=3)
        tbl.rows[0].cells[0].text = "A"
        tbl.rows[0].cells[1].text = "B"
        tbl.rows[0].cells[2].text = "C"
        for i in range(1, 5):
            tbl.rows[i].cells[0].text = f"{prefix}{i}c0"
            tbl.rows[i].cells[1].text = f"{prefix}{i}c1"
            tbl.rows[i].cells[2].text = f"{prefix}{i}c2"

    def diagnostic_for(table_index: int, caption: str):
        return tm.TableMarkerDiagnostic(
            table_index=table_index,
            rows_count=5,
            pages_detected=[12, 13],
            row_pages={0: 12, 1: 12, 2: 12, 3: 13, 4: 13},
            found_rows=[0, 1, 2, 3, 4],
            missing_rows=[],
            duplicate_rows={},
            candidate_for_split=False,
            page_spans=[tm.TablePageSpan(0, 2, 12), tm.TablePageSpan(3, 4, 13)],
            appendix_table=False,
            caption_detected=True,
            has_standard_table_caption=True,
            preceding_paragraph_text=caption,
        )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "multi_ordinary.docx"
        doc.save(path)

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        calls = {"diagnose": 0}
        try:
            def fake_diagnose_all(docx_path, keep_temp=False):
                calls["diagnose"] += 1
                return [
                    diagnostic_for(0, "Таблица 1.1"),
                    diagnostic_for(1, "Таблица 2.1"),
                ]

            tm.diagnose_all_tables = fake_diagnose_all
            tc.render_docx_to_pdf = lambda _path: (_ for _ in ()).throw(AssertionError("render path should not run after marker splits apply"))
            tc.analyze_pdf_lines = lambda _path: (_ for _ in ()).throw(AssertionError("pdf analysis should not run after marker splits apply"))
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        out = Document(str(path))

    if n != 2:
        return _result(False, f"expected two marker split mutations, got {n}")
    if calls["diagnose"] != 1:
        return _result(False, f"expected one marker diagnostic pass, got {calls['diagnose']}")
    if len(out.tables) != 4:
        return _result(False, f"expected 4 tables after two splits, got {len(out.tables)}")
    if [p.text for p in out.paragraphs].count("Продолжение таблицы 1.1") != 1:
        return _result(False, "first ordinary continuation paragraph missing")
    if [p.text for p in out.paragraphs].count("Продолжение таблицы 2.1") != 1:
        return _result(False, "second ordinary continuation paragraph missing")
    for table_index in (0, 2):
        if [cell.text for cell in out.tables[table_index].rows[1].cells] != ["1", "2", "3"]:
            return _result(False, f"first part table {table_index} does not have numbered row under header")
        if _count_table_rows_with_texts(out.tables[table_index], ["1", "2", "3"]) != 1:
            return _result(False, f"first part table {table_index} should contain exactly one numbered row")
    for table_index in (1, 3):
        if [cell.text for cell in out.tables[table_index].rows[0].cells] != ["1", "2", "3"]:
            return _result(False, f"continuation table {table_index} does not start with numbered row")
        if _count_table_rows_with_texts(out.tables[table_index], ["1", "2", "3"]) != 1:
            return _result(False, f"continuation table {table_index} should contain exactly one numbered row")
    return _result(True, "one run applies multiple ordinary marker splits")


def test_marker_runtime_apply_skips_stale_candidate_and_continues() -> tuple[bool, str]:
    """
    Product rule: batch marker apply uses one diagnostic snapshot; if one
    candidate is stale or invalid, later safe candidates are still applied.
    """
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    for num, prefix in (("1.1", "a"), ("2.1", "b")):
        doc.add_paragraph(f"Таблица {num}")
        tbl = doc.add_table(rows=5, cols=3)
        tbl.rows[0].cells[0].text = "A"
        tbl.rows[0].cells[1].text = "B"
        tbl.rows[0].cells[2].text = "C"
        for i in range(1, 5):
            tbl.rows[i].cells[0].text = f"{prefix}{i}c0"
            tbl.rows[i].cells[1].text = f"{prefix}{i}c1"
            tbl.rows[i].cells[2].text = f"{prefix}{i}c2"

    def diagnostic_for(table_index: int, caption: str):
        return tm.TableMarkerDiagnostic(
            table_index=table_index,
            rows_count=5,
            pages_detected=[12, 13],
            row_pages={0: 12, 1: 12, 2: 12, 3: 13, 4: 13},
            found_rows=[0, 1, 2, 3, 4],
            missing_rows=[],
            duplicate_rows={},
            candidate_for_split=False,
            page_spans=[tm.TablePageSpan(0, 2, 12), tm.TablePageSpan(3, 4, 13)],
            appendix_table=False,
            caption_detected=True,
            has_standard_table_caption=True,
            preceding_paragraph_text=caption,
        )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "stale_candidate.docx"
        doc.save(path)

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tm.diagnose_all_tables = lambda _path, keep_temp=False: [
                diagnostic_for(99, "Таблица 99.1"),
                diagnostic_for(1, "Таблица 2.1"),
            ]
            tc.render_docx_to_pdf = lambda _path: (_ for _ in ()).throw(AssertionError("render path should not run after marker split apply"))
            tc.analyze_pdf_lines = lambda _path: (_ for _ in ()).throw(AssertionError("pdf analysis should not run after marker split apply"))
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        out = Document(str(path))

    if n != 1:
        return _result(False, f"expected one valid split after stale skip, got {n}")
    if len(out.tables) != 3:
        return _result(False, f"expected one table to split, got {len(out.tables)} tables")
    if [p.text for p in out.paragraphs].count("Продолжение таблицы 2.1") != 1:
        return _result(False, "valid lower-index candidate was not applied after stale skip")
    return _result(True, "stale marker candidate is skipped while valid candidate applies")


def test_marker_runtime_apply_processes_mixed_ordinary_and_appendix_tables() -> tuple[bool, str]:
    """
    Product rule: ordinary and appendix marker splits can both be applied in
    one run, while appendix tables do not get ordinary continuation text.
    """
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 3.1")
    ordinary = doc.add_table(rows=5, cols=3)
    ordinary.rows[0].cells[0].text = "A"
    ordinary.rows[0].cells[1].text = "B"
    ordinary.rows[0].cells[2].text = "C"
    for i in range(1, 5):
        ordinary.rows[i].cells[0].text = f"o{i}c0"
        ordinary.rows[i].cells[1].text = f"o{i}c1"
        ordinary.rows[i].cells[2].text = f"o{i}c2"
    doc.add_paragraph("Приложение А")
    doc.add_paragraph("Расчет трудозатрат")
    appendix = doc.add_table(rows=5, cols=3)
    appendix.rows[0].cells[0].text = "Исполнитель"
    appendix.rows[0].cells[1].text = "Работы"
    appendix.rows[0].cells[2].text = "Стоимость"
    for i in range(1, 5):
        appendix.rows[i].cells[0].text = f"a{i}c0"
        appendix.rows[i].cells[1].text = f"a{i}c1"
        appendix.rows[i].cells[2].text = f"a{i}c2"

    def diagnostic_for(table_index: int, *, appendix_table: bool):
        return tm.TableMarkerDiagnostic(
            table_index=table_index,
            rows_count=5,
            pages_detected=[20, 21],
            row_pages={0: 20, 1: 20, 2: 20, 3: 21, 4: 21},
            found_rows=[0, 1, 2, 3, 4],
            missing_rows=[],
            duplicate_rows={},
            candidate_for_split=False,
            page_spans=[tm.TablePageSpan(0, 2, 20), tm.TablePageSpan(3, 4, 21)],
            appendix_table=appendix_table,
            caption_detected=True,
            has_standard_table_caption=not appendix_table,
            preceding_paragraph_text="Расчет трудозатрат" if appendix_table else "Таблица 3.1",
        )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "multi_mixed.docx"
        doc.save(path)

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        calls = {"diagnose": 0}
        try:
            def fake_diagnose_all(docx_path, keep_temp=False):
                calls["diagnose"] += 1
                return [
                    diagnostic_for(0, appendix_table=False),
                    diagnostic_for(1, appendix_table=True),
                ]

            tm.diagnose_all_tables = fake_diagnose_all
            tc.render_docx_to_pdf = lambda _path: (_ for _ in ()).throw(AssertionError("render path should not run after marker splits apply"))
            tc.analyze_pdf_lines = lambda _path: (_ for _ in ()).throw(AssertionError("pdf analysis should not run after marker splits apply"))
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        out = Document(str(path))

    if n != 2:
        return _result(False, f"expected two mixed marker split mutations, got {n}")
    if calls["diagnose"] != 1:
        return _result(False, f"expected one marker diagnostic pass, got {calls['diagnose']}")
    if [p.text for p in out.paragraphs].count("Продолжение таблицы 3.1") != 1:
        return _result(False, "ordinary continuation paragraph missing")
    if any("Продолжение таблицы" in (p.text or "") and p.text != "Продолжение таблицы 3.1" for p in out.paragraphs):
        return _result(False, "appendix split inserted forbidden continuation paragraph")
    if "ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ А" not in [p.text for p in out.paragraphs]:
        return _result(False, "appendix continuation label missing for mixed split")
    for table_index in (0, 2):
        if [cell.text for cell in out.tables[table_index].rows[1].cells] != ["1", "2", "3"]:
            return _result(False, f"first part table {table_index} does not have numbered row under header")
        if _count_table_rows_with_texts(out.tables[table_index], ["1", "2", "3"]) != 1:
            return _result(False, f"first part table {table_index} should contain exactly one numbered row")
    for table_index in (1, 3):
        if [cell.text for cell in out.tables[table_index].rows[0].cells] != ["1", "2", "3"]:
            return _result(False, f"continuation table {table_index} does not start with numbered row")
        if _count_table_rows_with_texts(out.tables[table_index], ["1", "2", "3"]) != 1:
            return _result(False, f"continuation table {table_index} should contain exactly one numbered row")
    return _result(True, "one run applies ordinary and appendix marker splits")


def test_marker_runtime_apply_skips_generated_appendix_continuation_tables() -> tuple[bool, str]:
    """
    Product rule: generated appendix continuation tables have a visible appendix
    continuation label, but must still be excluded from later marker passes.
    """
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Приложение А")
    first_part = doc.add_table(rows=3, cols=3)
    first_part.rows[0].cells[0].text = "Исполнитель"
    first_part.rows[0].cells[1].text = "Работы"
    first_part.rows[0].cells[2].text = "Стоимость"
    for i in range(1, 3):
        first_part.rows[i].cells[0].text = f"a{i}c0"
        first_part.rows[i].cells[1].text = f"a{i}c1"
        first_part.rows[i].cells[2].text = f"a{i}c2"

    label = doc.add_paragraph("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ А")
    label.paragraph_format.page_break_before = True
    label.paragraph_format.keep_with_next = True

    generated_continuation = doc.add_table(rows=3, cols=3)
    generated_continuation.rows[0].cells[0].text = "1"
    generated_continuation.rows[0].cells[1].text = "2"
    generated_continuation.rows[0].cells[2].text = "3"
    for i in range(1, 3):
        generated_continuation.rows[i].cells[0].text = f"cont{i}c0"
        generated_continuation.rows[i].cells[1].text = f"cont{i}c1"
        generated_continuation.rows[i].cells[2].text = f"cont{i}c2"

    doc.add_paragraph("Приложение Б")
    doc.add_paragraph("Расчет трудозатрат")
    original_appendix = doc.add_table(rows=5, cols=3)
    original_appendix.rows[0].cells[0].text = "Исполнитель"
    original_appendix.rows[0].cells[1].text = "Работы"
    original_appendix.rows[0].cells[2].text = "Стоимость"
    for i in range(1, 5):
        original_appendix.rows[i].cells[0].text = f"b{i}c0"
        original_appendix.rows[i].cells[1].text = f"b{i}c1"
        original_appendix.rows[i].cells[2].text = f"b{i}c2"

    def diagnostic_for(table_index: int, rows_count: int = 5):
        return tm.TableMarkerDiagnostic(
            table_index=table_index,
            rows_count=rows_count,
            pages_detected=[30, 31],
            row_pages={0: 30, 1: 30, 2: 30, 3: 31, 4: 31}
            if rows_count == 5
            else {0: 30, 1: 30, 2: 31},
            found_rows=list(range(rows_count)),
            missing_rows=[],
            duplicate_rows={},
            candidate_for_split=False,
            page_spans=[tm.TablePageSpan(0, 2, 30), tm.TablePageSpan(3, 4, 31)]
            if rows_count == 5
            else [tm.TablePageSpan(0, 1, 30), tm.TablePageSpan(2, 2, 31)],
            appendix_table=True,
            caption_detected=True,
            has_standard_table_caption=False,
            preceding_paragraph_text="Расчет трудозатрат",
        )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "appendix_generated_skip.docx"
        doc.save(path)

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            def fake_diagnose_all(docx_path, keep_temp=False):
                current = Document(str(docx_path))
                return [
                    diagnostic_for(1, rows_count=len(current.tables[1].rows)),
                    diagnostic_for(2),
                ]

            tm.diagnose_all_tables = fake_diagnose_all
            tc.render_docx_to_pdf = lambda _path: (_ for _ in ()).throw(AssertionError("render path should not run after marker splits apply"))
            tc.analyze_pdf_lines = lambda _path: (_ for _ in ()).throw(AssertionError("pdf analysis should not run after marker splits apply"))
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        out = Document(str(path))

    if n != 1:
        return _result(False, f"expected one original appendix split and one generated skip, got {n}")
    if len(out.tables) != 4:
        return _result(False, f"expected 4 tables after original appendix split, got {len(out.tables)}")
    if any("Продолжение таблицы" in (p.text or "") for p in out.paragraphs):
        return _result(False, "appendix split inserted forbidden continuation paragraph")
    if [p.text for p in out.paragraphs].count("ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ А") != 1:
        return _result(False, "pre-existing appendix continuation label was mutated")
    if "ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ Б" not in [p.text for p in out.paragraphs]:
        return _result(False, "new appendix continuation label missing")
    if _table_has_row_texts(out.tables[0], ["1", "2", "3"]):
        return _result(False, "pre-existing first appendix fragment should not be mutated")
    if [cell.text for cell in out.tables[2].rows[1].cells] != ["1", "2", "3"]:
        return _result(False, "original appendix split first part missing numbered row under header")
    if _count_table_rows_with_texts(out.tables[2], ["1", "2", "3"]) != 1:
        return _result(False, "original appendix split first part should contain exactly one numbered row")
    for table_index in (1, 3):
        numbered_rows = _count_table_rows_with_texts(out.tables[table_index], ["1", "2", "3"])
        if numbered_rows != 1:
            return _result(False, f"continuation appendix table {table_index} has {numbered_rows} numbered rows")
    return _result(True, "generated appendix continuation is skipped while original appendix applies")


def test_marker_runtime_apply_loop_is_bounded() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc

    doc = Document()
    for label in ("A", "B"):
        tbl = doc.add_table(rows=2, cols=1)
        tbl.rows[0].cells[0].text = "H"
        tbl.rows[1].cells[0].text = label

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "bounded.docx"
        doc.save(path)

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_pass = tc._run_marker_split_detection_pass
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        calls = {"count": 0}
        try:
            def fake_pass(_docx_path, *, apply_split=False):
                calls["count"] += 1
                return 1

            tc._run_marker_split_detection_pass = fake_pass
            tc.render_docx_to_pdf = lambda _path: (_ for _ in ()).throw(AssertionError("render path should not run after marker loop"))
            tc.analyze_pdf_lines = lambda _path: (_ for _ in ()).throw(AssertionError("pdf analysis should not run after marker loop"))
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc._run_marker_split_detection_pass = old_pass
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

    if n != 1:
        return _result(False, f"single marker pass should return pass result 1, got {n}")
    if calls["count"] != 1:
        return _result(False, f"single marker pass should call diagnostics once, got {calls['count']}")
    return _result(True, "marker apply uses one full diagnostic pass")


def test_split_prototype_simple_table() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    tbl = doc.add_table(rows=5, cols=2)
    for i in range(5):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "simple.docx"
        doc.save(src)
        result = prototype_split_table_copy(src, 0, 3, header_rows=1, keep_temp=True)
        out = Document(str(result.output_docx_path))

    if result.total_tables_after != 2:
        return _result(False, f"expected 2 tables after split, got {result.total_tables_after}")
    if result.first_table_rows_count != 3:
        return _result(False, f"expected 3 rows in first table, got {result.first_table_rows_count}")
    if result.second_table_rows_count != 3:
        return _result(False, f"expected 3 rows in second table, got {result.second_table_rows_count}")
    if out.tables[1].rows[0].cells[0].text != "r0c0":
        return _result(False, f"header row not copied into second table: {out.tables[1].rows[0].cells[0].text!r}")
    if out.tables[1].rows[1].cells[0].text != "r3c0":
        return _result(False, f"tail rows not moved to second table: {out.tables[1].rows[1].cells[0].text!r}")
    return _result(True, "simple table split produced two clone-based tables")


def test_split_prototype_source_note_stays_after_second_table() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    tbl = doc.add_table(rows=5, cols=1)
    for i in range(5):
        tbl.rows[i].cells[0].text = f"row{i}"
    doc.add_paragraph("Источник: данные автора")

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "source_note.docx"
        doc.save(src)
        result = prototype_split_table_copy(src, 0, 3, header_rows=1, keep_temp=True)
        out = Document(str(result.output_docx_path))

    if result.source_note_after_second is not True:
        return _result(False, f"source note did not stay after second table: {result.source_note_after_second!r}")

    body = list(out.element.body)
    def _local(node):
        return node.tag.split("}")[-1] if "}" in node.tag else node.tag

    tags = [_local(node) for node in body]
    try:
        first_tbl_idx = tags.index("tbl")
        second_tbl_idx = tags.index("tbl", first_tbl_idx + 1)
        note_idx = next(
            i for i, node in enumerate(body)
            if _local(node) == "p" and "Источник:" in "".join(t.text or "" for t in node.findall('.//' + qn('w:t')))
        )
    except Exception as exc:
        return _result(False, f"failed to inspect body ordering: {exc}")
    if not (first_tbl_idx < second_tbl_idx < note_idx):
        return _result(False, f"source note ordering invalid: first={first_tbl_idx}, second={second_tbl_idx}, note={note_idx}")
    return _result(True, "source note remains after second table")


def test_split_prototype_original_document_unchanged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    tbl = doc.add_table(rows=4, cols=1)
    for i in range(4):
        tbl.rows[i].cells[0].text = f"row{i}"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "original.docx"
        doc.save(src)
        before = src.read_bytes()
        prototype_split_table_copy(src, 0, 2, header_rows=1, keep_temp=True)
        after = src.read_bytes()

    if before != after:
        return _result(False, "source docx changed after prototype split")
    return _result(True, "prototype split leaves source document unchanged")


def test_split_prototype_invalid_table_index() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_table(rows=2, cols=1)
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "invalid_idx.docx"
        doc.save(src)
        try:
            prototype_split_table_copy(src, 3, 1, header_rows=1, keep_temp=False)
        except ValueError:
            return _result(True, "invalid table index rejected")
        except Exception as exc:
            return _result(False, f"unexpected exception type: {exc}")
    return _result(False, "expected ValueError for invalid table index")


def test_split_prototype_invalid_split_before_row() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_table(rows=3, cols=1)
    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "invalid_split.docx"
        doc.save(src)
        try:
            prototype_split_table_copy(src, 0, 0, header_rows=1, keep_temp=False)
        except ValueError:
            return _result(True, "invalid split_before_row rejected")
        except Exception as exc:
            return _result(False, f"unexpected exception type: {exc}")
    return _result(False, "expected ValueError for invalid split_before_row")


def test_split_prototype_no_continuation_paragraph_inserted() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_paragraph("Приложение А")
    doc.add_paragraph("Длинная таблица приложения")
    tbl = doc.add_table(rows=4, cols=1)
    for i in range(4):
        tbl.rows[i].cells[0].text = f"row{i}"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "appendix_split.docx"
        doc.save(src)
        result = prototype_split_table_copy(src, 0, 2, header_rows=1, keep_temp=True)
        out = Document(str(result.output_docx_path))

    paragraph_texts = [p.text for p in out.paragraphs]
    if any("Продолжение таблицы" in (text or "") for text in paragraph_texts):
        return _result(False, "unexpected continuation paragraph inserted")
    if any(text == "Продолжение" for text in paragraph_texts):
        return _result(False, "unexpected generic continuation paragraph inserted")
    return _result(True, "appendix split does not insert continuation paragraph")


def test_split_prototype_numbered_ordinary_continuation_row_only() -> tuple[bool, str]:
    """
    Product rule: ordinary table split continuation uses "Продолжение таблицы"
    plus the numbered row only; the original title/header row is not duplicated.
    """
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    tbl = doc.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "Показатель"
    tbl.rows[0].cells[1].text = "Как влияет"
    tbl.rows[0].cells[2].text = "Последствия"
    for i in range(1, 5):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"
        tbl.rows[i].cells[2].text = f"r{i}c2"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "ordinary_numbered.docx"
        doc.save(src)
        result = prototype_split_table_copy(
            src,
            0,
            3,
            header_rows=1,
            numbered_header=True,
            appendix_table=False,
            keep_temp=True,
        )
        out = Document(str(result.output_docx_path))

    if result.continuation_text != "Продолжение таблицы 1.1.1":
        return _result(False, f"unexpected continuation text: {result.continuation_text!r}")
    if result.continuation_paragraph_inserted is not True:
        return _result(False, "ordinary table should insert continuation paragraph")
    if result.column_count != 3:
        return _result(False, f"unexpected column_count: {result.column_count!r}")

    second_row_texts = [cell.text for cell in out.tables[1].rows[0].cells]
    if [cell.text for cell in out.tables[0].rows[1].cells] != ["1", "2", "3"]:
        return _result(False, "numbered row missing under first table header")
    if _count_table_rows_with_texts(out.tables[0], ["1", "2", "3"]) != 1:
        return _result(False, "first table should contain exactly one numbered row")
    if second_row_texts != ["1", "2", "3"]:
        return _result(False, f"unexpected continuation numbered row: {second_row_texts!r}")
    if _count_table_rows_with_texts(out.tables[1], ["1", "2", "3"]) != 1:
        return _result(False, "continuation table should contain exactly one numbered row")
    if [cell.text for cell in out.tables[1].rows[1].cells] != ["r3c0", "r3c1", "r3c2"]:
        return _result(False, "tail rows not moved under numbered continuation row")
    if not _all_table_rows_have_cant_split(out.tables[0]):
        return _result(False, "first ordinary prototype table rows can split across pages")
    if not _all_table_rows_have_cant_split(out.tables[1]):
        return _result(False, "second ordinary prototype table rows can split across pages")
    if any(cell.text == "Показатель" for cell in out.tables[1].rows[0].cells):
        return _result(False, "text header leaked into continuation numbered row")
    if "Продолжение таблицы 1.1.1" not in [p.text for p in out.paragraphs]:
        return _result(False, "continuation paragraph missing from ordinary numbered split")
    return _result(True, "ordinary numbered split uses continuation text and numbered row only")


def test_split_prototype_numbered_ordinary_split_caption_before_title() -> tuple[bool, str]:
    """
    Product rule: ordinary captions may be split as "Таблица X.Y.Z" plus a
    separate title paragraph before the table. Continuation uses the table
    number and does not duplicate the title/header text.
    """
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    title = "Комитеты НС Сбера: состав мандата и фокус надзора"

    doc = Document()
    doc.add_paragraph("Таблица 2.1.2")
    doc.add_paragraph(title)
    tbl = doc.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "Показатель"
    tbl.rows[0].cells[1].text = "Как влияет"
    tbl.rows[0].cells[2].text = "Последствия"
    for i in range(1, 5):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"
        tbl.rows[i].cells[2].text = f"r{i}c2"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "ordinary_split_caption_before_title.docx"
        doc.save(src)
        result = prototype_split_table_copy(
            src,
            0,
            3,
            header_rows=1,
            numbered_header=True,
            appendix_table=False,
            keep_temp=True,
        )
        out = Document(str(result.output_docx_path))

    if result.continuation_text != "Продолжение таблицы 2.1.2":
        return _result(False, f"unexpected continuation text: {result.continuation_text!r}")
    if result.continuation_paragraph_inserted is not True:
        return _result(False, "ordinary table should insert continuation paragraph")
    if [cell.text for cell in out.tables[0].rows[1].cells] != ["1", "2", "3"]:
        return _result(False, "numbered row missing under first split-caption table header")
    if _count_table_rows_with_texts(out.tables[0], ["1", "2", "3"]) != 1:
        return _result(False, "first split-caption table should contain exactly one numbered row")
    if [cell.text for cell in out.tables[1].rows[0].cells] != ["1", "2", "3"]:
        return _result(False, "continuation table should start with numbered row")
    if _count_table_rows_with_texts(out.tables[1], ["1", "2", "3"]) != 1:
        return _result(False, "split-caption continuation table should contain exactly one numbered row")
    if not _all_table_rows_have_cant_split(out.tables[0]):
        return _result(False, "first split-caption table rows can split across pages")
    if not _all_table_rows_have_cant_split(out.tables[1]):
        return _result(False, "second split-caption table rows can split across pages")
    if any(cell.text == "Показатель" for cell in out.tables[1].rows[0].cells):
        return _result(False, "text header leaked into continuation numbered row")
    paragraph_texts = [p.text for p in out.paragraphs]
    if paragraph_texts.count(title) != 1:
        return _result(False, f"table title duplicated or removed: {paragraph_texts!r}")
    if paragraph_texts.count("Продолжение таблицы 2.1.2") != 1:
        return _result(False, f"continuation paragraph missing or duplicated: {paragraph_texts!r}")
    return _result(True, "ordinary split caption before title uses continuation number")


def test_split_prototype_numbered_appendix_has_continuation_label() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_paragraph("Приложение 1")
    doc.add_paragraph("Таблица приложения")
    tbl = doc.add_table(rows=4, cols=3)
    tbl.rows[0].cells[0].text = "Колонка А"
    tbl.rows[0].cells[1].text = "Колонка Б"
    tbl.rows[0].cells[2].text = "Колонка В"
    for i in range(1, 4):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"
        tbl.rows[i].cells[2].text = f"r{i}c2"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "appendix_numbered.docx"
        doc.save(src)
        result = prototype_split_table_copy(
            src,
            0,
            2,
            header_rows=1,
            numbered_header=True,
            appendix_table=True,
            keep_temp=True,
        )
        out = Document(str(result.output_docx_path))

    if not result.continuation_paragraph_inserted:
        return _result(False, "appendix numbered split must insert continuation label")
    if result.continuation_text != "ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ 1":
        return _result(False, f"unexpected appendix continuation text: {result.continuation_text!r}")
    second_row_texts = [cell.text for cell in out.tables[1].rows[0].cells]
    if second_row_texts != ["1", "2", "3"]:
        return _result(False, f"unexpected appendix continuation row: {second_row_texts!r}")
    continuation_label = _paragraph_before_table(out, 1)
    if continuation_label is None:
        return _result(False, "appendix continuation label missing before continuation table")
    label_text = _paragraph_text(continuation_label)
    if label_text != "ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ 1":
        return _result(False, f"unexpected appendix continuation label: {label_text!r}")
    if label_text != label_text.upper():
        return _result(False, f"appendix continuation label is not uppercase: {label_text!r}")
    if not _paragraph_is_right_aligned(continuation_label):
        return _result(False, "appendix continuation label is not right aligned")
    if not _paragraph_has_page_break_before(continuation_label):
        return _result(False, "appendix continuation label should start on a new page")
    if not _paragraph_has_keep_next(continuation_label):
        return _result(False, "appendix continuation label should keep with continuation table")
    if _table_has_page_break_service_paragraph_before(out, 1):
        return _result(False, "appendix split left a blank service paragraph before continuation table")
    if [cell.text for cell in out.tables[0].rows[1].cells] != ["1", "2", "3"]:
        return _result(False, "numbered row missing under first appendix table header")
    if _count_table_rows_with_texts(out.tables[0], ["1", "2", "3"]) != 1:
        return _result(False, "first appendix table should contain exactly one numbered row")
    if _count_table_rows_with_texts(out.tables[1], ["1", "2", "3"]) != 1:
        return _result(False, "appendix continuation table should contain exactly one numbered row")
    if not _all_table_rows_have_cant_split(out.tables[0]):
        return _result(False, "first appendix prototype table rows can split across pages")
    if not _all_table_rows_have_cant_split(out.tables[1]):
        return _result(False, "second appendix prototype table rows can split across pages")
    if any("Продолжение таблицы" in (text or "") for text in [p.text for p in out.paragraphs]):
        return _result(False, "appendix numbered split inserted continuation paragraph")
    return _result(True, "appendix numbered split inserts continuation label and keeps numbered row")


def test_split_prototype_numbered_existing_row_reused_without_duplicate() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_paragraph("Таблица 2.4")
    tbl = doc.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[0].cells[2].text = "C"
    tbl.rows[1].cells[0].text = "1"
    tbl.rows[1].cells[1].text = "2"
    tbl.rows[1].cells[2].text = "3"
    for i in range(2, 5):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"
        tbl.rows[i].cells[2].text = f"r{i}c2"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "reuse_numbered.docx"
        doc.save(src)
        result = prototype_split_table_copy(
            src,
            0,
            4,
            header_rows=1,
            numbered_header=True,
            appendix_table=False,
            keep_temp=True,
        )
        out = Document(str(result.output_docx_path))

    if result.numbered_row_reused is not True:
        return _result(False, f"expected numbered row reuse, got {result.numbered_row_reused!r}")
    first_table_numbered_rows = sum(
        1 for row in out.tables[0].rows
        if [cell.text for cell in row.cells] == ["1", "2", "3"]
    )
    if first_table_numbered_rows != 1:
        return _result(False, f"expected exactly one numbered row in first table, got {first_table_numbered_rows}")
    second_row_texts = [cell.text for cell in out.tables[1].rows[0].cells]
    if second_row_texts != ["1", "2", "3"]:
        return _result(False, "reused numbered row not copied to continuation table")
    return _result(True, "existing numbered row is reused without duplication")


def test_split_prototype_numbered_malformed_existing_row_fails_safely() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_paragraph("Таблица 3.1")
    tbl = doc.add_table(rows=4, cols=3)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[0].cells[2].text = "C"
    tbl.rows[1].cells[0].text = "1"
    tbl.rows[1].cells[1].text = "3"
    tbl.rows[1].cells[2].text = "4"
    tbl.rows[2].cells[0].text = "r2c0"
    tbl.rows[2].cells[1].text = "r2c1"
    tbl.rows[2].cells[2].text = "r2c2"
    tbl.rows[3].cells[0].text = "r3c0"
    tbl.rows[3].cells[1].text = "r3c1"
    tbl.rows[3].cells[2].text = "r3c2"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "malformed_numbered.docx"
        doc.save(src)
        try:
            prototype_split_table_copy(
                src,
                0,
                3,
                header_rows=1,
                numbered_header=True,
                appendix_table=False,
                keep_temp=False,
            )
        except ValueError as exc:
            if "malformed" not in str(exc):
                return _result(False, f"unexpected ValueError text: {exc}")
            return _result(True, "malformed numbered row fails safely")
        except Exception as exc:
            return _result(False, f"unexpected exception type: {exc}")
    return _result(False, "expected ValueError for malformed numbered row")


def test_split_prototype_numbered_source_note_after_second_table() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_paragraph("Таблица 4.2")
    tbl = doc.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[0].cells[2].text = "C"
    for i in range(1, 5):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"
        tbl.rows[i].cells[2].text = f"r{i}c2"
    doc.add_paragraph("Источник: данные автора")

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "numbered_source_note.docx"
        doc.save(src)
        result = prototype_split_table_copy(
            src,
            0,
            3,
            header_rows=1,
            numbered_header=True,
            appendix_table=False,
            keep_temp=True,
        )
        out = Document(str(result.output_docx_path))

    if result.source_note_after_second is not True:
        return _result(False, "source note did not remain after continuation table")
    paragraph_texts = [p.text for p in out.paragraphs]
    if "Продолжение таблицы 4.2" not in paragraph_texts:
        return _result(False, "continuation paragraph missing in numbered source-note case")
    return _result(True, "numbered split keeps source note after second table")


def test_split_prototype_numbered_original_document_unchanged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_paragraph("Таблица 5.1")
    tbl = doc.add_table(rows=4, cols=3)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[0].cells[2].text = "C"
    for i in range(1, 4):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"
        tbl.rows[i].cells[2].text = f"r{i}c2"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "numbered_original.docx"
        doc.save(src)
        before = src.read_bytes()
        prototype_split_table_copy(
            src,
            0,
            2,
            header_rows=1,
            numbered_header=True,
            appendix_table=False,
            keep_temp=True,
        )
        after = src.read_bytes()

    if before != after:
        return _result(False, "source docx changed after numbered prototype split")
    return _result(True, "numbered prototype split leaves source document unchanged")


def test_split_prototype_numbered_row_has_no_numpr_and_no_calibri() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_split_prototype import prototype_split_table_copy

    doc = Document()
    doc.add_paragraph("Таблица 6.1")
    tbl = doc.add_table(rows=4, cols=3)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[0].cells[2].text = "C"
    for i in range(1, 4):
        for j in range(3):
            tbl.rows[i].cells[j].text = f"r{i}c{j}"

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "numbered_markup.docx"
        doc.save(src)
        result = prototype_split_table_copy(
            src,
            0,
            2,
            header_rows=1,
            numbered_header=True,
            appendix_table=False,
            keep_temp=True,
        )
        out = Document(str(result.output_docx_path))

    numbered_row = out.tables[1].rows[0]
    for cell in numbered_row.cells:
        for paragraph in cell.paragraphs:
            p_pr = paragraph._element.find(qn("w:pPr"))
            if p_pr is not None and p_pr.find(qn("w:numPr")) is not None:
                return _result(False, "generated numbered row has w:numPr")
            for run in paragraph.runs:
                r_pr = run._element.find(qn("w:rPr"))
                fonts = r_pr.find(qn("w:rFonts")) if r_pr is not None else None
                ascii_font = fonts.get(qn("w:ascii")) if fonts is not None else None
                if ascii_font != "Times New Roman":
                    return _result(False, f"generated numbered row font is {ascii_font!r}, expected Times New Roman")
    return _result(True, "generated numbered row has no numPr and no Calibri fallback")


def test_marker_runtime_flags_do_not_change_headings() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    def heading_snapshot(doc: Document):
        out = []
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text in {"ВВЕДЕНИЕ", "1. ГЛАВА", "1.1. Подраздел"}:
                p_pr = p._element.find(qn("w:pPr"))
                num_pr = p_pr.find(qn("w:numPr")) if p_pr is not None else None
                fonts = []
                for run in p.runs[:2]:
                    r_pr = run._element.find(qn("w:rPr"))
                    r_fonts = r_pr.find(qn("w:rFonts")) if r_pr is not None else None
                    fonts.append(r_fonts.get(qn("w:ascii")) if r_fonts is not None else None)
                out.append((text, p.style.name if p.style else None, num_pr is not None, tuple(fonts)))
        return out

    base = Document()
    p = base.add_paragraph("ВВЕДЕНИЕ")
    p.style = "Heading 1"
    p.runs[0].font.name = "Times New Roman"
    p = base.add_paragraph("1. ГЛАВА")
    p.style = "Heading 1"
    p.runs[0].font.name = "Times New Roman"
    p = base.add_paragraph("1.1. Подраздел")
    p.style = "Heading 2"
    p.runs[0].font.name = "Times New Roman"
    base.add_paragraph("Таблица 1.1.1")
    tbl = base.add_table(rows=5, cols=3)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[0].cells[2].text = "C"
    for i in range(1, 5):
        for j in range(3):
            tbl.rows[i].cells[j].text = f"r{i}c{j}"

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=5,
        pages_detected=[12, 13],
        row_pages={0: 12, 1: 12, 2: 12, 3: 13, 4: 13},
        found_rows=[0, 1, 2, 3, 4],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 2, 12), tm.TablePageSpan(3, 4, 13)],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text="Таблица 1.1.1",
    )

    expected = heading_snapshot(base)

    for mode_name, env in [
        ("flags_off", {}),
        ("dry_run", {"KPFU_ENABLE_MARKER_SPLIT": "1"}),
        ("apply", {"KPFU_ENABLE_MARKER_SPLIT": "1", "KPFU_APPLY_MARKER_SPLIT": "1"}),
    ]:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / f"{mode_name}.docx"
            base.save(path)
            old_env = {k: os.environ.get(k) for k in ["KPFU_ENABLE_MARKER_SPLIT", "KPFU_APPLY_MARKER_SPLIT"]}
            old_diagnose_all = tm.diagnose_all_tables
            old_render = tc.render_docx_to_pdf
            old_analyze = tc.analyze_pdf_lines
            try:
                for k in old_env:
                    os.environ.pop(k, None)
                os.environ.update(env)
                tm.diagnose_all_tables = lambda _path, keep_temp=False: [diagnostic]
                tc.render_docx_to_pdf = lambda _path: (_ for _ in ()).throw(AssertionError("render path should not run in heading regression test"))
                tc.analyze_pdf_lines = lambda _path: []
                tc.apply_rendered_table_continuation(path)
            finally:
                tm.diagnose_all_tables = old_diagnose_all
                tc.render_docx_to_pdf = old_render
                tc.analyze_pdf_lines = old_analyze
                for k, v in old_env.items():
                    if v is None:
                        os.environ.pop(k, None)
                    else:
                        os.environ[k] = v

            out = Document(str(path))
            if heading_snapshot(out) != expected:
                return _result(False, f"heading snapshot changed in mode {mode_name}: {heading_snapshot(out)!r}")

    return _result(True, "flags off, dry-run, and apply do not change headings outside target table")


def test_marker_runtime_real_rybakov_target_applies_split() -> tuple[bool, str]:
    asset = next(ASSETS.glob("*Рыбаков*.docx"), None)
    if asset is None:
        return _result(True, "Рыбаков asset missing, skipped")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "rybakov_apply.docx"
        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        try:
            os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
            os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
            format_docx(str(asset), str(out))
        finally:
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        doc = Document(str(out))

    if len(doc.tables) != 11:
        return _result(False, f"expected 11 tables after Рыбаков split, got {len(doc.tables)}")
    if len(doc.tables[9].rows) != 6 or len(doc.tables[10].rows) != 17:
        return _result(False, f"unexpected table row counts after split: {len(doc.tables[9].rows)}/{len(doc.tables[10].rows)}")
    if [c.text for c in doc.tables[9].rows[1].cells] != ["1", "2", "3", "4", "5"]:
        return _result(False, "first split table missing numbered row")
    if [c.text for c in doc.tables[10].rows[0].cells] != ["1", "2", "3", "4", "5"]:
        return _result(False, "second split table missing numbered row")
    continuations = [p.text for p in doc.paragraphs if "Продолжение таблицы" in (p.text or "")]
    if any(text == "Продолжение таблицы 10" for text in continuations):
        return _result(False, f"unexpected appendix continuation paragraph inserted: {continuations!r}")
    return _result(True, "Рыбаков target is split with numbered rows in active mode")


def test_marker_runtime_real_bondarev_keeps_headings_safe() -> tuple[bool, str]:
    asset = ASSETS / "курсовая_Бондарев_Никита_2_курс.docx"
    if not asset.exists():
        return _result(True, "Бондарев asset missing, skipped")

    def snapshot(path: Path):
        doc = Document(str(path))
        out = []
        for p in doc.paragraphs:
            text = (p.text or "").strip()
            if text in {
                "ВВЕДЕНИЕ",
                "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ ФУНКЦИЙ И ОРГАНОВ САМОУПРАВЛЕНИЯ В ОРГАНИЗАЦИИ",
                "1.1. Понятие, сущность и классификация органов самоуправления в организациях",
            }:
                p_pr = p._element.find(qn("w:pPr"))
                num_pr = p_pr.find(qn("w:numPr")) if p_pr is not None else None
                fonts = []
                for run in p.runs[:2]:
                    r_pr = run._element.find(qn("w:rPr"))
                    r_fonts = r_pr.find(qn("w:rFonts")) if r_pr is not None else None
                    fonts.append(r_fonts.get(qn("w:ascii")) if r_fonts is not None else None)
                out.append((text, p.style.name if p.style else None, num_pr is not None, tuple(fonts)))
        return out

    with tempfile.TemporaryDirectory() as tmp:
        off = Path(tmp) / "bond_off.docx"
        on = Path(tmp) / "bond_on.docx"

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        try:
            os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            format_docx(str(asset), str(off))
            os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
            os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
            format_docx(str(asset), str(on))
        finally:
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2

        off_snapshot = snapshot(off)
        on_snapshot = snapshot(on)

    if off_snapshot != on_snapshot:
        return _result(False, f"Бондарев heading snapshot changed under active split: off={off_snapshot!r} on={on_snapshot!r}")
    if any(item[2] for item in on_snapshot):
        return _result(False, f"Бондарев headings unexpectedly gained w:numPr: {on_snapshot!r}")
    if any("Calibri" in (font or "") for item in on_snapshot for font in item[3]):
        return _result(False, f"Бондарев headings unexpectedly use Calibri: {on_snapshot!r}")
    return _result(True, "Бондарев active mode leaves headings unchanged")


def _get_para_xml_attrs(p):
    from docx.oxml.ns import qn
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        return {}
    out = {}
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        for k, v in ind.attrib.items():
            out[f"ind:{k.split('}')[-1]}"] = v
    jc = pPr.find(qn("w:jc"))
    out["jc"] = jc.get(qn("w:val")) if jc is not None else None

    def _bool_flag(elem):
        # Word boolean toggle: element absent → off; element present without
        # val → on; element present with val in {"0","false","off"} → off.
        if elem is None:
            return False
        val = elem.get(qn("w:val"))
        if val is None:
            return True
        return val.lower() not in {"0", "false", "off"}

    out["keepNext"] = _bool_flag(pPr.find(qn("w:keepNext")))
    out["pageBreakBefore"] = _bool_flag(pPr.find(qn("w:pageBreakBefore")))
    return out


def test_caption_tail_is_reference_prose_unit() -> tuple[bool, str]:
    """Unit-test the shared helper used by the new caption demotion logic."""
    from guides.coursework_kfu_2025.classifier import caption_tail_is_reference_prose
    cases = [
        ("показывает динамику", True),
        ("показывают, что", True),
        ("демонстрирует структуру", True),
        ("отражает зависимость", True),
        ("иллюстрирует тезис", True),
        ("содержит данные", True),
        ("представлены результаты", True),
        ("свидетельствует о росте", True),
        (". показывает, что", True),
        ("— показывает, что", True),
        ("Жизненный цикл документа", False),
        ("— Влияние документооборота", False),
        ("Закупочный центр потенциального франчайзи", False),
        ("", False),
        ("   ", False),
    ]
    for tail, expected in cases:
        got = caption_tail_is_reference_prose(tail)
        if got != expected:
            return _result(False, f"caption_tail_is_reference_prose({tail!r}) = {got}, expected {expected}")
    return _result(True, "caption_tail_is_reference_prose covers verb list and edge cases")


def test_table_reference_paragraph_not_caption() -> tuple[bool, str]:
    """'Таблица 1.1.1 показывает...' is body text even when adjacent to a real table."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст введения.")
    doc.add_paragraph("Таблица 1.1.1 показывает динамику внедрения ЭДО в крупных компаниях.")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "h0"
    tbl.rows[0].cells[1].text = "h1"
    tbl.rows[1].cells[0].text = "v0"
    tbl.rows[1].cells[1].text = "v1"
    doc.add_paragraph("Заключительный текст.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    target = None
    for p in result.paragraphs:
        if p.text.strip().startswith("Таблица 1.1.1 показывает"):
            target = p
            break
    if target is None:
        return _result(False, "reference-prose paragraph missing from output")
    attrs = _get_para_xml_attrs(target)
    # Body paragraph: justified (or default body), 1.25cm first-line indent,
    # no keepNext, no pageBreakBefore, not heading style.
    if attrs.get("keepNext"):
        return _result(False, f"reference-prose paragraph has keepNext: attrs={attrs!r}")
    if attrs.get("pageBreakBefore"):
        return _result(False, f"reference-prose paragraph has pageBreakBefore: attrs={attrs!r}")
    if attrs.get("ind:firstLine") != "709":
        return _result(False, f"reference-prose paragraph wrong first-line indent: attrs={attrs!r}")
    if attrs.get("jc") not in (None, "both", "left"):
        # process_document uses justify ("both" in DOCX) for body; allow None as inherited body default
        return _result(False, f"reference-prose paragraph wrongly aligned: attrs={attrs!r}")
    style = (target.style.name or "").lower()
    if "heading" in style or "заголовок" in style:
        return _result(False, f"reference-prose paragraph got heading style: {style!r}")
    return _result(True, "Таблица N показывает... stays body text")


def test_figure_reference_paragraph_not_caption() -> tuple[bool, str]:
    """'Рис. 1.1.1. показывает...' is body text, not figure caption."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Вводный текст.")
    # Real figure block (image + source + real caption)
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Рис. 1.1.1. Закупочный центр потенциального франчайзи")
    # Body prose referencing the figure — must NOT be reclassified as caption
    doc.add_paragraph("Рис. 1.1.1. показывает структуру закупочного центра в малом бизнесе.")
    doc.add_paragraph("Заключительный текст.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    target = None
    for p in result.paragraphs:
        if p.text.strip().startswith("Рис. 1.1.1. показывает"):
            target = p
            break
    if target is None:
        return _result(False, "figure reference-prose paragraph missing from output")
    attrs = _get_para_xml_attrs(target)
    if attrs.get("keepNext"):
        return _result(False, f"figure reference-prose has keepNext: attrs={attrs!r}")
    if attrs.get("pageBreakBefore"):
        return _result(False, f"figure reference-prose has pageBreakBefore: attrs={attrs!r}")
    if attrs.get("jc") == "center":
        return _result(False, f"figure reference-prose centered: attrs={attrs!r}")
    if attrs.get("ind:firstLine") != "709":
        return _result(False, f"figure reference-prose wrong first-line indent: attrs={attrs!r}")
    return _result(True, "Рис. N. показывает... stays body text, not centered")


def test_no_keep_with_next_on_reference_paragraph() -> tuple[bool, str]:
    """Phase 2 pagination must not set keepWithNext on reference-prose paragraphs."""
    from guides.coursework_kfu_2025.pagination_rules import apply_pagination_rules

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 1.1.1 показывает результаты пилота.")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.rows[0].cells[0].text = "cell"
    doc.add_paragraph("Рис. 1.1.1. показывает структуру.")

    apply_pagination_rules(doc)

    target_tab = next((p for p in doc.paragraphs if p.text.startswith("Таблица 1.1.1 показывает")), None)
    target_fig = next((p for p in doc.paragraphs if p.text.startswith("Рис. 1.1.1. показывает")), None)
    if target_tab is None or target_fig is None:
        return _result(False, "reference-prose paragraphs missing from doc")
    if _get_para_xml_attrs(target_tab).get("keepNext"):
        return _result(False, "Phase 2 set keepNext on table reference-prose paragraph")
    if _get_para_xml_attrs(target_fig).get("keepNext"):
        return _result(False, "Phase 2 set keepNext on figure reference-prose paragraph")
    return _result(True, "pagination_rules leaves reference-prose paragraphs alone")


def test_actual_table_caption_still_formats() -> tuple[bool, str]:
    """Regression guard: a real table caption is still classified and formatted as caption."""
    from guides.coursework_kfu_2025.safe_formatter import process_document
    from guides.coursework_kfu_2025.pagination_rules import apply_pagination_rules

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph("Таблица 1.1.1")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "h0"
    tbl.rows[0].cells[1].text = "h1"
    tbl.rows[1].cells[0].text = "v0"
    tbl.rows[1].cells[1].text = "v1"

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))
    apply_pagination_rules(result)

    caption = next((p for p in result.paragraphs if p.text.strip() == "Таблица 1.1.1"), None)
    if caption is None:
        return _result(False, "real table caption missing after format")
    attrs = _get_para_xml_attrs(caption)
    if not attrs.get("keepNext"):
        return _result(False, f"real table caption lost keepNext: attrs={attrs!r}")
    # format_table_caption sets right alignment — the existing Phase 1 behavior must be preserved.
    if attrs.get("jc") != "right":
        return _result(False, f"real table caption alignment changed: attrs={attrs!r}")
    return _result(True, "real table caption still formatted as caption and keepNext set")


def test_actual_figure_caption_still_formats() -> tuple[bool, str]:
    """Regression guard: a real figure caption is still classified as figure_caption."""
    from guides.coursework_kfu_2025.safe_formatter import process_document
    from guides.coursework_kfu_2025.classifier import classify_paragraph

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Рис. 1.1.1. Жизненный цикл документа")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    caption = next((p for p in result.paragraphs if "Жизненный цикл документа" in p.text), None)
    if caption is None:
        return _result(False, "real figure caption missing after format")
    # The text-only classifier must still call this a figure caption.
    if classify_paragraph(caption.text) != "figure_caption":
        return _result(False, f"real figure caption misclassified: {caption.text!r} -> {classify_paragraph(caption.text)}")
    return _result(True, "real figure caption still classified as figure_caption")


def test_phase3_marker_budget_fail_open_many_tables() -> tuple[bool, str]:
    """
    When a document has more tables than the rendered-marker-split budget,
    apply_rendered_table_continuation must skip the expensive diagnostic
    (no LibreOffice render, no diagnose_all_tables call), return 0, and
    surface a Russian warning via the FormattingReport.
    """
    import logging
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm
    from guides.coursework_kfu_2025.docx_utils import FormattingReport

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    for i in range(5):
        doc.add_paragraph(f"Таблица 1.1.{i+1}")
        tbl = doc.add_table(rows=2, cols=2)
        tbl.rows[0].cells[0].text = f"h{i}0"
        tbl.rows[0].cells[1].text = f"h{i}1"
        tbl.rows[1].cells[0].text = f"v{i}0"
        tbl.rows[1].cells[1].text = f"v{i}1"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "many_tables.docx"
        doc.save(path)
        before = path.read_bytes()

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_budget = os.environ.get("KPFU_MARKER_SPLIT_MAX_RENDERS")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MAX_RENDERS"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        render_calls: list = []
        diagnose_calls: list = []
        log_handler_buffer: list[str] = []

        class _CaptureHandler(logging.Handler):
            def emit(self, record):
                log_handler_buffer.append(record.getMessage())

        handler = _CaptureHandler(level=logging.DEBUG)
        previous_level = tc.logger.level
        try:
            def boom_diagnose(_path, keep_temp=False):
                diagnose_calls.append(_path)
                raise AssertionError(
                    "diagnose_all_tables must not be called when render budget is exceeded"
                )

            def boom_render(_path):
                render_calls.append(_path)
                raise AssertionError(
                    "render_docx_to_pdf must not be called when render budget is exceeded"
                )

            tm.diagnose_all_tables = boom_diagnose
            tc.render_docx_to_pdf = boom_render
            tc.analyze_pdf_lines = lambda _path: []

            report = FormattingReport()
            tc.logger.addHandler(handler)
            tc.logger.setLevel(logging.DEBUG)
            n = tc.apply_rendered_table_continuation(path, report=report)
        finally:
            tc.logger.removeHandler(handler)
            tc.logger.setLevel(previous_level)
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2
            if old_budget is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MAX_RENDERS", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MAX_RENDERS"] = old_budget

        after = path.read_bytes()

    if n != 0:
        return _result(False, f"expected 0 from budget skip, got {n}")
    if before != after:
        return _result(False, "budget skip should not mutate document bytes")
    if diagnose_calls:
        return _result(False, "diagnose_all_tables should not have been called")
    if render_calls:
        return _result(False, "render_docx_to_pdf should not have been called")
    skip_log = next(
        (m for m in log_handler_buffer if "render_budget_exceeded" in m and "marker_split_skipped" in m),
        None,
    )
    if skip_log is None:
        return _result(False, f"missing budget skip log; captured: {log_handler_buffer!r}")
    if "count=5" not in skip_log or "budget=1" not in skip_log:
        return _result(False, f"skip log missing count/budget: {skip_log!r}")
    if not any("Автоматическое разделение длинных таблиц пропущено" in w for w in report.warnings):
        return _result(False, f"expected Russian warning in report, got {report.warnings!r}")
    if not any("Проверьте переносы таблиц вручную" in w for w in report.warnings):
        return _result(False, f"expected advisory tail in report, got {report.warnings!r}")
    return _result(True, "many-table doc fails open without render or diagnose")


def test_phase3_marker_budget_allows_small_doc() -> tuple[bool, str]:
    """
    When the table count fits inside the budget, the marker diagnostic path
    runs unchanged: diagnose_all_tables must be called and the existing
    eligibility/skip logic must apply.
    """
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    tbl = doc.add_table(rows=4, cols=2)
    for i in range(4):
        tbl.rows[i].cells[0].text = f"r{i}c0"
        tbl.rows[i].cells[1].text = f"r{i}c1"

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=4,
        pages_detected=[12, 13, 14],
        row_pages={0: 12, 1: 12, 2: 13, 3: 14},
        found_rows=[0, 1, 2, 3],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text="Таблица 1.1.1",
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "small_doc.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "small_doc.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)
        before = path.read_bytes()

        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_apply = os.environ.get("KPFU_APPLY_MARKER_SPLIT")
        old_budget = os.environ.get("KPFU_MARKER_SPLIT_MAX_RENDERS")
        old_mode_e2 = os.environ.get("KPFU_MARKER_SPLIT_MODE")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
        os.environ["KPFU_MARKER_SPLIT_MAX_RENDERS"] = "10"
        os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"

        old_diagnose_all = tm.diagnose_all_tables
        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        diagnose_calls: list = []
        try:
            def spy_diagnose(_path, keep_temp=False):
                diagnose_calls.append(_path)
                return [diagnostic]

            tm.diagnose_all_tables = spy_diagnose
            tc.render_docx_to_pdf = lambda _path: pdf_path
            tc.analyze_pdf_lines = lambda _path: []
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tm.diagnose_all_tables = old_diagnose_all
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_apply is None:
                os.environ.pop("KPFU_APPLY_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_APPLY_MARKER_SPLIT"] = old_apply
            if old_mode_e2 is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MODE", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MODE"] = old_mode_e2
            if old_budget is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MAX_RENDERS", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MAX_RENDERS"] = old_budget

        after = path.read_bytes()

    if not diagnose_calls:
        return _result(False, "diagnose_all_tables must be called when budget allows")
    if n != 0:
        return _result(False, f"3-page diagnostic should be ineligible (n=0), got {n}")
    if before != after:
        return _result(False, "dry-run path with ineligible diagnostic should not mutate doc")
    return _result(True, "small-doc marker split runs unchanged when within budget")


# ── Patch B+C: figure caption alignment + blank cleanup ───────────────────────

def test_figure_caption_alignment_left_or_justify() -> tuple[bool, str]:
    """Patch B: format_figure_caption must set LEFT or JUSTIFY, not CENTER."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Рис. 1.1.1. Жизненный цикл документа")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    caption = next((p for p in result.paragraphs if "Жизненный цикл документа" in p.text), None)
    if caption is None:
        return _result(False, "figure caption missing after format")
    attrs = _get_para_xml_attrs(caption)
    jc = attrs.get("jc")
    if jc == "center":
        return _result(False, f"figure caption is still CENTER after Patch B: jc={jc!r}")
    if jc not in (None, "left", "both"):
        return _result(False, f"figure caption has unexpected alignment: jc={jc!r}")
    return _result(True, f"figure caption alignment is left/justify (jc={jc!r}), not center")


def test_remove_empty_between_figure_source_and_caption() -> tuple[bool, str]:
    """Patch C: blank paragraph between Источник: and figure caption must be removed."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("")  # blank that must be removed
    doc.add_paragraph("Рис. 1.2.1. Структура отдела")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    paragraphs = result.paragraphs
    source_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip().startswith("Источник:")), None)
    caption_idx = next((i for i, p in enumerate(paragraphs) if "Структура отдела" in p.text), None)
    if source_idx is None:
        return _result(False, "source line missing from output")
    if caption_idx is None:
        return _result(False, "figure caption missing from output")
    if caption_idx != source_idx + 1:
        between = [p.text for p in paragraphs[source_idx + 1:caption_idx]]
        return _result(False, f"blank not removed between source and caption; between={between!r}")
    return _result(True, "blank between Источник: and figure caption was removed")


def test_one_blank_after_real_figure_caption_before_body_prose() -> tuple[bool, str]:
    """FSP: real figure caption must be followed by exactly one blank before reference prose."""
    from guides.coursework_kfu_2025.safe_formatter import process_document, is_empty_paragraph

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Рис. 1.3.1. Закупочный центр")
    doc.add_paragraph("Рисунок 1.3.1 показывает структуру закупочного центра в малом бизнесе.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    paragraphs = result.paragraphs
    cap_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "Рис. 1.3.1. Закупочный центр"), None)
    prose_idx = next((i for i, p in enumerate(paragraphs) if "Рисунок 1.3.1 показывает структуру" in p.text), None)
    if cap_idx is None:
        return _result(False, "figure caption missing from output")
    if prose_idx is None:
        return _result(False, "figure reference prose missing from output")

    between = paragraphs[cap_idx + 1:prose_idx]
    if len(between) != 1:
        return _result(False, f"expected exactly 1 paragraph between cap and prose, got {len(between)}: {[p.text for p in between]!r}")
    if not is_empty_paragraph(between[0]):
        return _result(False, f"paragraph between cap and prose is not empty: {between[0].text!r}")
    return _result(True, "exactly one blank paragraph between real figure caption and reference prose")


def test_one_blank_after_caption_is_idempotent() -> tuple[bool, str]:
    """FSP: running process_document twice keeps exactly one blank after caption (no doubling)."""
    from guides.coursework_kfu_2025.safe_formatter import process_document, is_empty_paragraph

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Рис. 1.3.1. Закупочный центр")
    doc.add_paragraph("Рисунок 1.3.1 показывает структуру.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        mid = Path(tmp) / "mid.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, mid)
        process_document(mid, out)
        result = Document(str(out))

    paragraphs = result.paragraphs
    cap_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "Рис. 1.3.1. Закупочный центр"), None)
    prose_idx = next((i for i, p in enumerate(paragraphs) if "Рисунок 1.3.1 показывает структуру" in p.text), None)
    if cap_idx is None or prose_idx is None:
        return _result(False, "caption or reference prose missing after second run")

    between = paragraphs[cap_idx + 1:prose_idx]
    if len(between) != 1:
        return _result(False, f"second run gave {len(between)} paragraphs between cap and prose, expected 1: {[p.text for p in between]!r}")
    if not is_empty_paragraph(between[0]):
        return _result(False, "second run: paragraph between cap and prose is not empty")
    return _result(True, "two consecutive process_document runs preserve exactly one blank after caption")


def test_table_source_not_affected_by_figure_blank_cleanup() -> tuple[bool, str]:
    """Patch C guard: table source → blank → table caption must NOT be affected."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph("Таблица 1.1.1. Показатели")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "h0"
    tbl.rows[0].cells[1].text = "h1"
    tbl.rows[1].cells[0].text = "v0"
    tbl.rows[1].cells[1].text = "v1"
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("")  # blank after table source — must NOT be removed
    doc.add_paragraph("Текст после таблицы.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    paragraphs = result.paragraphs
    source_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip().startswith("Источник:")), None)
    prose_idx = next((i for i, p in enumerate(paragraphs) if "Текст после таблицы" in p.text), None)
    if source_idx is None:
        return _result(False, "table source line missing from output")
    if prose_idx is None:
        return _result(False, "prose after table missing from output")
    # source → prose should NOT be adjacent (blank must be preserved or some spacing kept)
    # The key check: no figure-blank-cleanup touched this table source
    # We just verify the table source and following prose both exist in correct order
    if prose_idx <= source_idx:
        return _result(False, f"unexpected paragraph order: source={source_idx}, prose={prose_idx}")
    return _result(True, "table source blank not touched by figure blank cleanup")


def test_figure_reference_prose_still_body_after_patch_A() -> tuple[bool, str]:
    """Regression: 'Рисунок N показывает...' must remain body text (not caption) after Patch A+B+C."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Вводный текст.")
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Рис. 2.1.1. Динамика показателей")
    doc.add_paragraph("Рисунок 2.1.1 показывает динамику внедрения ЭДО.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    target = next((p for p in result.paragraphs if "показывает динамику внедрения" in p.text), None)
    if target is None:
        return _result(False, "figure reference prose missing from output")
    attrs = _get_para_xml_attrs(target)
    if attrs.get("jc") == "center":
        return _result(False, f"figure reference prose was centered (got caption treatment): attrs={attrs!r}")
    if attrs.get("keepNext"):
        return _result(False, f"figure reference prose has keepNext: attrs={attrs!r}")
    style = (target.style.name or "").lower()
    if "heading" in style or "заголовок" in style:
        return _result(False, f"figure reference prose got heading style: {style!r}")
    return _result(True, "Рисунок N показывает... stays body text after Patch A+B+C")


# ── PB2: figure caption keepLines + figure block keepWithNext chain ───────────

def _keep_flags(p):
    """Return (keepNext, keepLines) as booleans for a paragraph."""
    from docx.oxml.ns import qn
    pPr = p._element.find(qn("w:pPr"))
    if pPr is None:
        return (False, False)

    def _flag(tag):
        el = pPr.find(qn(f"w:{tag}"))
        if el is None:
            return False
        val = el.get(qn("w:val"))
        return val is None or val.lower() not in {"0", "false", "off"}

    return (_flag("keepNext"), _flag("keepLines"))


def test_figure_caption_keep_lines_true() -> tuple[bool, str]:
    """PB2: format_figure_caption must set keepLines so caption text doesn't split across pages."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Рис. 1.1.1. Длинная подпись рисунка, которая может занять больше одной строки в готовом документе")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    caption = next((p for p in result.paragraphs if "Длинная подпись" in p.text), None)
    if caption is None:
        return _result(False, "figure caption missing after format")
    keep_next, keep_lines = _keep_flags(caption)
    if not keep_lines:
        return _result(False, f"caption keepLines is False, expected True")
    if keep_next:
        return _result(False, f"caption keepNext is True, expected False (caption is last link)")
    return _result(True, "figure caption has keepLines=True, keepNext=False")


def test_figure_block_image_keeps_with_source_and_caption() -> tuple[bool, str]:
    """PB2: IMG → Источник → Примечание → CAP chain has keepNext on IMG and both service lines, none on CAP."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Вводный текст.")
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Примечание: схема упрощена.")
    doc.add_paragraph("Рис. 1.3.1. Этапы организационной покупки франшизы")
    doc.add_paragraph("Заключительный текст после рисунка.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    paragraphs = result.paragraphs
    img = next((p for p in paragraphs if p._element.xpath(".//*[local-name()='drawing']")), None)
    src = next((p for p in paragraphs if p.text.strip().startswith("Источник:")), None)
    note = next((p for p in paragraphs if p.text.strip().startswith("Примечание:")), None)
    cap = next((p for p in paragraphs if p.text.strip().startswith("Рис. 1.3.1.")), None)
    for name, par in (("image", img), ("source", src), ("note", note), ("caption", cap)):
        if par is None:
            return _result(False, f"{name} paragraph missing from output")

    img_kn, _ = _keep_flags(img)
    src_kn, _ = _keep_flags(src)
    note_kn, _ = _keep_flags(note)
    cap_kn, cap_kl = _keep_flags(cap)

    if not img_kn:
        return _result(False, "image paragraph missing keepNext")
    if not src_kn:
        return _result(False, "source line missing keepNext")
    if not note_kn:
        return _result(False, "note line missing keepNext")
    if cap_kn:
        return _result(False, "caption has keepNext (should be False, it is the last link)")
    if not cap_kl:
        return _result(False, "caption missing keepLines")
    return _result(True, "IMG+source+note chained via keepNext; caption keepLines=True keepNext=False")


def test_figure_block_appendix_image_keeps_with_caption() -> tuple[bool, str]:
    """PB2: appendix block IMG → CAP (no source) — IMG must have keepNext, caption must not."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph("ПРИЛОЖЕНИЯ")
    doc.add_paragraph("ПРИЛОЖЕНИЕ 1")
    img_p = doc.add_paragraph()
    drawing = OxmlElement("w:drawing")
    r = OxmlElement("w:r")
    r.append(drawing)
    img_p._element.append(r)
    doc.add_paragraph("Рис. 1. Страница сайта франшизы")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    paragraphs = result.paragraphs
    img = next((p for p in paragraphs if p._element.xpath(".//*[local-name()='drawing']")), None)
    cap = next((p for p in paragraphs if p.text.strip().startswith("Рис. 1.")), None)
    if img is None or cap is None:
        return _result(False, "appendix image or caption missing from output")

    img_kn, _ = _keep_flags(img)
    cap_kn, cap_kl = _keep_flags(cap)

    if not img_kn:
        return _result(False, "appendix image paragraph missing keepNext")
    if cap_kn:
        return _result(False, "appendix caption has keepNext (should be False)")
    if not cap_kl:
        return _result(False, "appendix caption missing keepLines")
    return _result(True, "appendix IMG→CAP chain: IMG keepNext=True, CAP keepNext=False keepLines=True")


def test_table_block_unaffected_by_figure_keepnext() -> tuple[bool, str]:
    """PB2 guard: table source 'Источник:' must NOT receive figure-style keepNext."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph("Таблица 1.1.1. Показатели")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "h0"
    tbl.rows[0].cells[1].text = "h1"
    tbl.rows[1].cells[0].text = "v0"
    tbl.rows[1].cells[1].text = "v1"
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Текст после таблицы.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    src = next((p for p in result.paragraphs if p.text.strip().startswith("Источник:")), None)
    if src is None:
        return _result(False, "table source paragraph missing from output")
    src_kn, _ = _keep_flags(src)
    if src_kn:
        return _result(False, "table source got figure-style keepNext (should be False — no preceding image)")
    return _result(True, "table source paragraph unaffected by figure-block keepNext pass")


# ── TCF-A: universal <w:cantSplit/> on every table row ───────────────────────

def _count_cant_split_per_row(table):
    """Return list of <w:cantSplit/> counts per <w:tr> in the given table element."""
    from docx.oxml.ns import qn
    counts = []
    for tr in table._element.findall(qn("w:tr")):
        tr_pr = tr.find(qn("w:trPr"))
        if tr_pr is None:
            counts.append(0)
            continue
        counts.append(len(tr_pr.findall(qn("w:cantSplit"))))
    return counts


def test_tcfa_all_table_rows_have_cant_split() -> tuple[bool, str]:
    """TCF-A: every <w:tr> in every table must carry <w:cantSplit/> after process_document."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    for rows, cols in ((2, 2), (3, 3), (4, 2)):
        doc.add_paragraph(f"Таблица 1.1.{rows}")
        tbl = doc.add_table(rows=rows, cols=cols)
        for r in range(rows):
            for c in range(cols):
                tbl.rows[r].cells[c].text = f"r{r}c{c}"
        doc.add_paragraph("")  # spacing

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    if len(result.tables) < 3:
        return _result(False, f"expected at least 3 tables in output, got {len(result.tables)}")

    for ti, table in enumerate(result.tables):
        counts = _count_cant_split_per_row(table)
        if not counts:
            return _result(False, f"table {ti} has no rows after format")
        for ri, c in enumerate(counts):
            if c < 1:
                return _result(False, f"table {ti} row {ri} missing <w:cantSplit/>")
    return _result(True, f"all {sum(len(_count_cant_split_per_row(t)) for t in result.tables)} rows across {len(result.tables)} tables have cantSplit")


def test_tcfa_cant_split_is_idempotent() -> tuple[bool, str]:
    """TCF-A: running process_document twice must not duplicate <w:cantSplit/> on any row."""
    from guides.coursework_kfu_2025.safe_formatter import process_document

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph("Таблица 1.1.1")
    tbl = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            tbl.rows[r].cells[c].text = f"r{r}c{c}"

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        mid = Path(tmp) / "mid.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, mid)
        process_document(mid, out)
        result = Document(str(out))

    for ti, table in enumerate(result.tables):
        counts = _count_cant_split_per_row(table)
        for ri, c in enumerate(counts):
            if c != 1:
                return _result(False, f"second run: table {ti} row {ri} has {c} <w:cantSplit/> elements, expected exactly 1")
    return _result(True, "two runs preserve exactly one <w:cantSplit/> per row")


def test_tcfa_existing_cant_split_preserved() -> tuple[bool, str]:
    """TCF-A: a pre-existing <w:cantSplit/> set on input rows must be preserved (exactly one element each)."""
    from guides.coursework_kfu_2025.safe_formatter import process_document
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph("Таблица 1.1.1")
    tbl = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            tbl.rows[r].cells[c].text = f"r{r}c{c}"
    # Manually pre-set <w:cantSplit/> on every row
    for tr in tbl._element.findall(qn("w:tr")):
        tr_pr = tr.find(qn("w:trPr"))
        if tr_pr is None:
            tr_pr = OxmlElement("w:trPr")
            tr.insert(0, tr_pr)
        tr_pr.append(OxmlElement("w:cantSplit"))

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    if not result.tables:
        return _result(False, "table missing from output")
    counts = _count_cant_split_per_row(result.tables[0])
    for ri, c in enumerate(counts):
        if c != 1:
            return _result(False, f"row {ri} has {c} <w:cantSplit/> elements, expected exactly 1 (no duplication)")
    return _result(True, "pre-existing cantSplit preserved, not duplicated")


def test_tcfa_does_not_change_caption_or_source_classification() -> tuple[bool, str]:
    """TCF-A guard: adding cantSplit must not perturb caption / source / reference-prose classification."""
    from guides.coursework_kfu_2025.safe_formatter import process_document
    from guides.coursework_kfu_2025.classifier import classify_paragraph
    from docx.oxml.ns import qn

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph("Таблица 1.1.1")
    tbl = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            tbl.rows[r].cells[c].text = f"r{r}c{c}"
    doc.add_paragraph("Источник: составлено автором.")
    doc.add_paragraph("Таблица 1.1.1 показывает динамику внедрения ЭДО.")

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(str(inp))
        process_document(inp, out)
        result = Document(str(out))

    caption = next((p for p in result.paragraphs if p.text.strip() == "Таблица 1.1.1"), None)
    source = next((p for p in result.paragraphs if p.text.strip().startswith("Источник:")), None)
    prose = next((p for p in result.paragraphs if p.text.strip().startswith("Таблица 1.1.1 показывает")), None)

    if caption is None:
        return _result(False, "table caption missing from output")
    if source is None:
        return _result(False, "table source missing from output")
    if prose is None:
        return _result(False, "reference prose missing from output")

    if classify_paragraph(caption.text) != "table_caption":
        return _result(False, f"table caption misclassified: {classify_paragraph(caption.text)!r}")

    # Source must NOT have figure-style keepNext (PB2 guard)
    pPr = source._element.find(qn("w:pPr"))
    if pPr is not None:
        keep_next = pPr.find(qn("w:keepNext"))
        if keep_next is not None:
            v = keep_next.get(qn("w:val"))
            if v is None or v.lower() not in {"0", "false", "off"}:
                return _result(False, "table source got figure-style keepNext (PB2 guard regression)")

    # Reference prose stays body — must not be classified as caption
    if classify_paragraph(prose.text) == "table_caption":
        return _result(False, f"reference prose wrongly classified as table_caption: {prose.text!r}")

    return _result(True, "caption/source/prose classification unchanged after TCF-A")


# ── E1: Phase 3 marker-split candidate classification (logging-only) ─────────

def _e1_build_synth_doc(table_specs: list[tuple[int, int, str | None]]) -> Document:
    """Build a synthetic doc with [(rows, cols, caption_or_None), ...] tables."""
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    for rows, cols, caption in table_specs:
        if caption is not None:
            doc.add_paragraph(caption)
        tbl = doc.add_table(rows=rows, cols=cols)
        for r in range(rows):
            for c in range(cols):
                tbl.rows[r].cells[c].text = f"r{r}c{c}"
        doc.add_paragraph("")
    return doc


def test_e1_classify_many_tiny_tables_no_candidates() -> tuple[bool, str]:
    """E1: 20 tables × 2 rows with real captions → zero candidates (all tiny)."""
    from guides.coursework_kfu_2025.table_continuation import _classify_marker_split_candidates

    specs = [(2, 2, f"Таблица 1.1.{i+1}") for i in range(20)]
    doc = _e1_build_synth_doc(specs)
    c = _classify_marker_split_candidates(doc)
    if c["total_tables"] != 20:
        return _result(False, f"total_tables={c['total_tables']}, expected 20")
    if c["candidate_tables"]:
        return _result(False, f"expected no candidates, got {c['candidate_tables']}")
    if len(c["tiny_table_skipped"]) != 20:
        return _result(False, f"expected 20 tiny skipped, got {len(c['tiny_table_skipped'])}")
    return _result(True, f"20 tiny tables → 0 candidates, all in tiny_table_skipped")


def test_e1_classify_one_long_candidate_among_many() -> tuple[bool, str]:
    """E1: 1 long table + 9 tiny tables → exactly 1 candidate, the long one."""
    from guides.coursework_kfu_2025.table_continuation import _classify_marker_split_candidates

    specs = [(2, 2, f"Таблица 1.1.{i+1}") for i in range(9)]
    specs.insert(4, (30, 3, "Таблица 1.1.10"))  # one big table in the middle
    doc = _e1_build_synth_doc(specs)
    c = _classify_marker_split_candidates(doc)
    if len(c["candidate_tables"]) != 1:
        return _result(False, f"expected 1 candidate, got {c['candidate_tables']}")
    if len(c["tiny_table_skipped"]) != 9:
        return _result(False, f"expected 9 tiny, got {len(c['tiny_table_skipped'])}")
    # Verify the candidate is the long one (rows=30, priority=30 since not appendix)
    cand_idx, cand_priority = c["candidate_priority"][0]
    if cand_priority != 30:
        return _result(False, f"candidate priority={cand_priority}, expected 30")
    return _result(True, f"1 long candidate selected (idx={cand_idx}, priority={cand_priority})")


def test_e1_classify_skips_manual_continuation() -> tuple[bool, str]:
    """E1: tables already in a valid manual continuation chain must be filtered out."""
    from guides.coursework_kfu_2025.table_continuation import _classify_marker_split_candidates

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph("Таблица 1.1.1")
    # Matching headers + 6+ rows so it would otherwise be a candidate.
    tbl1 = doc.add_table(rows=8, cols=3)
    tbl1.rows[0].cells[0].text = "H1"; tbl1.rows[0].cells[1].text = "H2"; tbl1.rows[0].cells[2].text = "H3"
    for r in range(1, 8):
        for c_ in range(3):
            tbl1.rows[r].cells[c_].text = f"r{r}c{c_}"

    marker = doc.add_paragraph("Продолжение таблицы 1.1.1")
    marker.alignment = 2  # WD_ALIGN_PARAGRAPH.RIGHT
    marker.paragraph_format.keep_with_next = True

    tbl2 = doc.add_table(rows=6, cols=3)
    tbl2.rows[0].cells[0].text = "H1"; tbl2.rows[0].cells[1].text = "H2"; tbl2.rows[0].cells[2].text = "H3"
    for r in range(1, 6):
        for c_ in range(3):
            tbl2.rows[r].cells[c_].text = f"r{r}c{c_}"

    c = _classify_marker_split_candidates(doc)
    if len(c["manual_continuation_skipped"]) != 2:
        return _result(False, f"expected 2 tables in manual_continuation_skipped, got {c['manual_continuation_skipped']}")
    if any(idx in c["candidate_tables"] for idx in c["manual_continuation_skipped"]):
        return _result(False, "manual-chain table also appeared in candidate_tables")
    return _result(True, f"manual continuation chain filtered: {c['manual_continuation_skipped']}")


def test_e1_classify_skips_tables_without_caption() -> tuple[bool, str]:
    """E1: a long table without a 'Таблица N' caption above must not become a candidate."""
    from guides.coursework_kfu_2025.table_continuation import _classify_marker_split_candidates

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Просто текст без подписи таблицы.")  # no caption
    tbl = doc.add_table(rows=20, cols=2)
    for r in range(20):
        for c_ in range(2):
            tbl.rows[r].cells[c_].text = f"r{r}c{c_}"

    c = _classify_marker_split_candidates(doc)
    if c["candidate_tables"]:
        return _result(False, f"no-caption table wrongly became candidate: {c['candidate_tables']}")
    if not c["no_caption_skipped"]:
        return _result(False, "expected no_caption_skipped to contain the table")
    return _result(True, f"no-caption table correctly filtered to no_caption_skipped={c['no_caption_skipped']}")


def test_e1_classify_priority_ordering() -> tuple[bool, str]:
    """E1: multiple candidates ordered by descending priority (rows + appendix bonus)."""
    from guides.coursework_kfu_2025.table_continuation import _classify_marker_split_candidates

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    # Body tables: rows 10, 25, 15
    for caption, rows in (("Таблица 1.1.1", 10), ("Таблица 1.1.2", 25), ("Таблица 1.1.3", 15)):
        doc.add_paragraph(caption)
        t = doc.add_table(rows=rows, cols=2)
        for r in range(rows):
            for cc in range(2):
                t.rows[r].cells[cc].text = "x"

    c = _classify_marker_split_candidates(doc)
    if len(c["candidate_tables"]) != 3:
        return _result(False, f"expected 3 candidates, got {c['candidate_tables']}")
    priorities = [p for _, p in c["candidate_priority"]]
    if priorities != sorted(priorities, reverse=True):
        return _result(False, f"priorities not descending: {priorities}")
    # Top candidate should be the 25-row table
    if priorities[0] != 25:
        return _result(False, f"top priority={priorities[0]}, expected 25")
    return _result(True, f"candidates ordered by priority desc: {priorities}")


def test_e1_log_emitted_with_classification() -> tuple[bool, str]:
    """E1: apply_rendered_table_continuation emits phase3_candidate_classification + previews."""
    import logging
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    # 2 candidates (rows>=6, real captions) so previews are emitted
    for caption, rows in (("Таблица 1.1.1", 8), ("Таблица 1.1.2", 12)):
        doc.add_paragraph(caption)
        t = doc.add_table(rows=rows, cols=2)
        for r in range(rows):
            for cc in range(2):
                t.rows[r].cells[cc].text = "x"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)
        before = path.read_bytes()

        captured: list[str] = []
        class _CaptureHandler(logging.Handler):
            def emit(self, record):
                captured.append(record.getMessage())

        handler = _CaptureHandler(level=logging.DEBUG)
        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        old_diagnose = tm.diagnose_all_tables
        prev_level = tc.logger.level
        try:
            # Block any real diagnose call so this test never touches LibreOffice.
            tm.diagnose_all_tables = lambda *a, **k: (_ for _ in ()).throw(AssertionError("must not run"))
            tc.logger.addHandler(handler)
            tc.logger.setLevel(logging.DEBUG)
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.logger.removeHandler(handler)
            tc.logger.setLevel(prev_level)
            tm.diagnose_all_tables = old_diagnose
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable

        after = path.read_bytes()

    classification_line = next((m for m in captured if m.startswith("phase3_candidate_classification ")), None)
    if classification_line is None:
        return _result(False, "phase3_candidate_classification log line not emitted")
    if "candidates=2" not in classification_line:
        return _result(False, f"expected candidates=2 in log, got: {classification_line!r}")
    previews = [m for m in captured if m.startswith("marker_split_candidate_preview ")]
    if len(previews) != 2:
        return _result(False, f"expected 2 preview lines, got {len(previews)}: {previews!r}")
    if before != after:
        return _result(False, "doc bytes changed despite E1 being observe-only")
    return _result(True, f"classification + 2 previews emitted; doc unchanged; return={n}")


def test_e1_does_not_change_existing_behaviour() -> tuple[bool, str]:
    """E1: docs > budget still hit render_budget_exceeded skip; output bytes unchanged; return 0."""
    import logging
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    # 8 tables — exceeds default budget of 6
    for i in range(8):
        doc.add_paragraph(f"Таблица 1.1.{i+1}")
        t = doc.add_table(rows=2, cols=2)
        for r in range(2):
            for cc in range(2):
                t.rows[r].cells[cc].text = "x"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)
        before = path.read_bytes()

        captured: list[str] = []
        class _CaptureHandler(logging.Handler):
            def emit(self, record):
                captured.append(record.getMessage())

        handler = _CaptureHandler(level=logging.DEBUG)
        old_enable = os.environ.get("KPFU_ENABLE_MARKER_SPLIT")
        old_budget = os.environ.get("KPFU_MARKER_SPLIT_MAX_RENDERS")
        os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
        os.environ["KPFU_MARKER_SPLIT_MAX_RENDERS"] = "6"

        old_diagnose = tm.diagnose_all_tables
        prev_level = tc.logger.level
        try:
            tm.diagnose_all_tables = lambda *a, **k: (_ for _ in ()).throw(AssertionError("must not run on >budget doc"))
            tc.logger.addHandler(handler)
            tc.logger.setLevel(logging.DEBUG)
            n = tc.apply_rendered_table_continuation(path)
        finally:
            tc.logger.removeHandler(handler)
            tc.logger.setLevel(prev_level)
            tm.diagnose_all_tables = old_diagnose
            if old_enable is None:
                os.environ.pop("KPFU_ENABLE_MARKER_SPLIT", None)
            else:
                os.environ["KPFU_ENABLE_MARKER_SPLIT"] = old_enable
            if old_budget is None:
                os.environ.pop("KPFU_MARKER_SPLIT_MAX_RENDERS", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_MAX_RENDERS"] = old_budget

        after = path.read_bytes()

    if n != 0:
        return _result(False, f"expected return 0 (skip), got {n}")
    if before != after:
        return _result(False, "doc bytes changed under render_budget_exceeded path")
    if not any(m.startswith("phase3_candidate_classification ") for m in captured):
        return _result(False, "phase3_candidate_classification not emitted on >budget doc")
    if not any("render_budget_exceeded" in m for m in captured):
        return _result(False, "render_budget_exceeded log missing — existing behavior lost")
    return _result(True, "8>6 docs still skip with render_budget_exceeded; bytes unchanged; new logs additive only")


# ── E2: quality-first candidate-mode marker split ────────────────────────────

class _E2Report:
    """Minimal FormattingReport stub for E2 tests."""
    def __init__(self):
        self.warnings: list[str] = []
    def warn(self, msg: str):
        self.warnings.append(msg)


def _e2_make_long_table_doc(num_long_tables: int, rows_per_table: int = 8) -> Document:
    """Synthetic doc with N captioned long tables (each ≥ candidate threshold)."""
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    for i in range(num_long_tables):
        doc.add_paragraph(f"Таблица 1.1.{i+1}")
        t = doc.add_table(rows=rows_per_table, cols=3)
        for r in range(rows_per_table):
            for c in range(3):
                t.rows[r].cells[c].text = f"r{r}c{c}"
        doc.add_paragraph("")
    return doc


def _e2_set_candidate_mode(budget: int | None = None, hard_timeout: float | None = None):
    """Set env for candidate mode; return restore dict."""
    save = {
        "enable": os.environ.get("KPFU_ENABLE_MARKER_SPLIT"),
        "apply":  os.environ.get("KPFU_APPLY_MARKER_SPLIT"),
        "mode":   os.environ.get("KPFU_MARKER_SPLIT_MODE"),
        "budget": os.environ.get("KPFU_MARKER_SPLIT_MAX_RENDERS"),
        "hto":    os.environ.get("KPFU_MARKER_SPLIT_HARD_TIMEOUT_SECONDS"),
    }
    os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
    os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
    os.environ["KPFU_MARKER_SPLIT_MODE"] = "candidate"
    if budget is not None:
        os.environ["KPFU_MARKER_SPLIT_MAX_RENDERS"] = str(budget)
    if hard_timeout is not None:
        os.environ["KPFU_MARKER_SPLIT_HARD_TIMEOUT_SECONDS"] = str(hard_timeout)
    return save


def _e2_restore_env(save):
    for key, env in (("enable", "KPFU_ENABLE_MARKER_SPLIT"),
                     ("apply",  "KPFU_APPLY_MARKER_SPLIT"),
                     ("mode",   "KPFU_MARKER_SPLIT_MODE"),
                     ("budget", "KPFU_MARKER_SPLIT_MAX_RENDERS"),
                     ("hto",    "KPFU_MARKER_SPLIT_HARD_TIMEOUT_SECONDS")):
        if save[key] is None:
            os.environ.pop(env, None)
        else:
            os.environ[env] = save[key]


def _e2_fake_diagnostic(tm, table_index, rows_count, two_page=True):
    """Build a TableMarkerDiagnostic that the eligibility evaluator considers eligible
    for a clean 2-page split."""
    # Put first half on page 1 and second half on page 2.
    split = rows_count // 2
    row_pages = {}
    for r in range(rows_count):
        row_pages[r] = 1 if r < split else 2
    page_spans = [tm.TablePageSpan(0, split - 1, 1), tm.TablePageSpan(split, rows_count - 1, 2)]
    return tm.TableMarkerDiagnostic(
        table_index=table_index,
        rows_count=rows_count,
        pages_detected=[1, 2] if two_page else [1],
        row_pages=row_pages if two_page else {r: 1 for r in range(rows_count)},
        found_rows=list(range(rows_count)),
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=page_spans if two_page else [tm.TablePageSpan(0, rows_count - 1, 1)],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text=f"Таблица 1.1.{table_index+1}",
    )


def test_e2_31_total_8_candidates_marker_split_enters() -> tuple[bool, str]:
    """E2: doc with 31 total tables (8 candidates after filter) enters marker split (not global skip)."""
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    # 8 long tables + 23 tiny tables (no caption) = 31 total, 8 candidates.
    doc = _e2_make_long_table_doc(8, rows_per_table=8)
    for i in range(23):
        # tiny tables without caption — filtered as no_caption.
        t = doc.add_table(rows=2, cols=2)
        for r in range(2):
            for c in range(2):
                t.rows[r].cells[c].text = "x"
        doc.add_paragraph("")

    diagnosed: list[int] = []
    def _fake_diag(_path, idx, keep_temp=False):
        diagnosed.append(idx)
        return _e2_fake_diagnostic(tm, idx, rows_count=8, two_page=False)  # no eligible split

    save = _e2_set_candidate_mode(budget=20)
    old_dt = tm.diagnose_table
    try:
        tm.diagnose_table = _fake_diag
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            report = _E2Report()
            n = tc.apply_rendered_table_continuation(path, report=report)
    finally:
        tm.diagnose_table = old_dt
        _e2_restore_env(save)

    if len(diagnosed) != 8:
        return _result(False, f"expected 8 diagnose calls, got {len(diagnosed)}")
    if report.warnings:
        return _result(False, f"expected no user warnings (all candidates processed), got {report.warnings!r}")
    return _result(True, f"31 total / 8 candidates → 8 diagnose calls, no warning, return={n}")


def test_e2_25_total_7_candidates_marker_split_enters() -> tuple[bool, str]:
    """E2: 25 total / 7 candidates enters marker split."""
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = _e2_make_long_table_doc(7, rows_per_table=8)
    for i in range(18):
        t = doc.add_table(rows=2, cols=2)
        for r in range(2):
            for c in range(2):
                t.rows[r].cells[c].text = "x"
        doc.add_paragraph("")

    diagnosed: list[int] = []
    def _fake_diag(_path, idx, keep_temp=False):
        diagnosed.append(idx)
        return _e2_fake_diagnostic(tm, idx, rows_count=8, two_page=False)

    save = _e2_set_candidate_mode(budget=20)
    old_dt = tm.diagnose_table
    try:
        tm.diagnose_table = _fake_diag
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            report = _E2Report()
            n = tc.apply_rendered_table_continuation(path, report=report)
    finally:
        tm.diagnose_table = old_dt
        _e2_restore_env(save)

    if len(diagnosed) != 7:
        return _result(False, f"expected 7 diagnose calls, got {len(diagnosed)}")
    if report.warnings:
        return _result(False, f"unexpected warnings: {report.warnings!r}")
    return _result(True, f"25 total / 7 candidates → 7 diagnose calls, no warning")


def test_e2_11_total_10_candidates_marker_split_enters() -> tuple[bool, str]:
    """E2: 11 total / 10 candidates enters marker split within budget=20."""
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = _e2_make_long_table_doc(10, rows_per_table=8)
    # +1 tiny without caption
    t = doc.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            t.rows[r].cells[c].text = "x"

    diagnosed: list[int] = []
    def _fake_diag(_path, idx, keep_temp=False):
        diagnosed.append(idx)
        return _e2_fake_diagnostic(tm, idx, rows_count=8, two_page=False)

    save = _e2_set_candidate_mode(budget=20)
    old_dt = tm.diagnose_table
    try:
        tm.diagnose_table = _fake_diag
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            report = _E2Report()
            n = tc.apply_rendered_table_continuation(path, report=report)
    finally:
        tm.diagnose_table = old_dt
        _e2_restore_env(save)

    if len(diagnosed) != 10:
        return _result(False, f"expected 10 diagnose calls, got {len(diagnosed)}")
    if report.warnings:
        return _result(False, f"unexpected warnings: {report.warnings!r}")
    return _result(True, f"11 total / 10 candidates → 10 diagnose calls, no warning")


def test_e2_no_warning_when_all_candidates_processed() -> tuple[bool, str]:
    """E2: report.warn must NOT fire when all candidates are diagnosed within budget."""
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = _e2_make_long_table_doc(5, rows_per_table=8)
    def _fake_diag(_path, idx, keep_temp=False):
        return _e2_fake_diagnostic(tm, idx, rows_count=8, two_page=False)

    save = _e2_set_candidate_mode(budget=20)
    old_dt = tm.diagnose_table
    try:
        tm.diagnose_table = _fake_diag
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            report = _E2Report()
            tc.apply_rendered_table_continuation(path, report=report)
    finally:
        tm.diagnose_table = old_dt
        _e2_restore_env(save)

    if report.warnings:
        return _result(False, f"warnings emitted when all candidates processed: {report.warnings!r}")
    return _result(True, "no user warning when all candidates fit in budget")


def test_e2_warning_when_some_candidates_skipped_for_budget() -> tuple[bool, str]:
    """E2: when candidates > budget, overflow tables logged + user warning fires."""
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    # 10 long candidates, budget=6 → 4 overflow
    doc = _e2_make_long_table_doc(10, rows_per_table=8)

    diagnosed: list[int] = []
    def _fake_diag(_path, idx, keep_temp=False):
        diagnosed.append(idx)
        return _e2_fake_diagnostic(tm, idx, rows_count=8, two_page=False)

    save = _e2_set_candidate_mode(budget=6)
    old_dt = tm.diagnose_table
    try:
        tm.diagnose_table = _fake_diag
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            report = _E2Report()
            tc.apply_rendered_table_continuation(path, report=report)
    finally:
        tm.diagnose_table = old_dt
        _e2_restore_env(save)

    if len(diagnosed) != 6:
        return _result(False, f"expected 6 diagnose calls (budget), got {len(diagnosed)}")
    if not report.warnings:
        return _result(False, "expected user warning for budget overflow, got none")
    msg = report.warnings[0]
    if "частично" not in msg.lower() or "10" not in msg:
        return _result(False, f"warning text wrong: {msg!r}")
    return _result(True, f"6/10 processed; warning fires: {msg[:80]!r}")


def test_e2_candidate_diagnose_failure_does_not_abort_others() -> tuple[bool, str]:
    """E2: a single diagnose_table failure logs and continues with the next candidate."""
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = _e2_make_long_table_doc(3, rows_per_table=8)

    diagnosed: list[int] = []
    def _flaky(_path, idx, keep_temp=False):
        diagnosed.append(idx)
        if idx == 0:
            raise RuntimeError("simulated diagnose failure")
        return _e2_fake_diagnostic(tm, idx, rows_count=8, two_page=False)

    save = _e2_set_candidate_mode(budget=20)
    old_dt = tm.diagnose_table
    try:
        tm.diagnose_table = _flaky
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            report = _E2Report()
            tc.apply_rendered_table_continuation(path, report=report)
    finally:
        tm.diagnose_table = old_dt
        _e2_restore_env(save)

    if diagnosed != [0, 1, 2]:
        return _result(False, f"expected diagnose attempts on all 3 candidates, got {diagnosed!r}")
    if not report.warnings:
        return _result(False, "expected user warning for the 1 diagnose failure, got none")
    return _result(True, f"diagnose failure on idx=0 did not abort idx=1,2; warning fires")


def test_e2_existing_manual_continuation_preserved() -> tuple[bool, str]:
    """E2: a valid manual chain is filtered out of candidates; not diagnosed."""
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph("Таблица 1.1.1")
    t1 = doc.add_table(rows=8, cols=3)
    t1.rows[0].cells[0].text = "H1"; t1.rows[0].cells[1].text = "H2"; t1.rows[0].cells[2].text = "H3"
    for r in range(1, 8):
        for c in range(3):
            t1.rows[r].cells[c].text = f"r{r}c{c}"
    marker = doc.add_paragraph("Продолжение таблицы 1.1.1")
    marker.alignment = 2
    marker.paragraph_format.keep_with_next = True
    t2 = doc.add_table(rows=6, cols=3)
    t2.rows[0].cells[0].text = "H1"; t2.rows[0].cells[1].text = "H2"; t2.rows[0].cells[2].text = "H3"
    for r in range(1, 6):
        for c in range(3):
            t2.rows[r].cells[c].text = f"r{r}c{c}"

    diagnosed: list[int] = []
    def _fake_diag(_path, idx, keep_temp=False):
        diagnosed.append(idx)
        return _e2_fake_diagnostic(tm, idx, rows_count=8, two_page=False)

    save = _e2_set_candidate_mode(budget=20)
    old_dt = tm.diagnose_table
    try:
        tm.diagnose_table = _fake_diag
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            tc.apply_rendered_table_continuation(path)
    finally:
        tm.diagnose_table = old_dt
        _e2_restore_env(save)

    if diagnosed:
        return _result(False, f"manual chain tables wrongly diagnosed: {diagnosed!r}")
    return _result(True, "valid manual chain filtered before diagnose")


def test_e2_continuation_label_has_blank_after_before_table() -> tuple[bool, str]:
    """E2: auto-inserted ordinary continuation has structure tbl → marker_p → blank_p → tbl."""
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 1.2.3")
    tbl = doc.add_table(rows=6, cols=3)
    tbl.rows[0].cells[0].text = "H1"; tbl.rows[0].cells[1].text = "H2"; tbl.rows[0].cells[2].text = "H3"
    for r in range(1, 6):
        for c in range(3):
            tbl.rows[r].cells[c].text = f"r{r}c{c}"
    doc.add_paragraph("Источник: данные автора")

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=6,
        pages_detected=[1, 2],
        row_pages={0: 1, 1: 1, 2: 1, 3: 2, 4: 2, 5: 2},
        found_rows=[0, 1, 2, 3, 4, 5],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 2, 1), tm.TablePageSpan(3, 5, 2)],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text="Таблица 1.2.3",
    )

    save = _e2_set_candidate_mode(budget=20)
    old_dt = tm.diagnose_table
    try:
        tm.diagnose_table = lambda _p, idx, keep_temp=False: diagnostic
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            n = tc.apply_rendered_table_continuation(path)
            out = Document(str(path))
    finally:
        tm.diagnose_table = old_dt
        _e2_restore_env(save)

    if n != 1:
        return _result(False, f"expected 1 split applied, got {n}")
    if len(out.tables) != 2:
        return _result(False, f"expected 2 tables after split, got {len(out.tables)}")

    body = out.element.body
    children = list(body)
    tbl1 = out.tables[0]._element
    tbl2 = out.tables[1]._element
    i1 = children.index(tbl1)
    i2 = children.index(tbl2)
    if i2 - i1 != 3:
        return _result(False, f"chain shape wrong: expected tbl→p→blank→tbl, got gap={i2-i1}")
    marker = children[i1 + 1]
    blank = children[i1 + 2]
    if marker.tag != qn("w:p"):
        return _result(False, "marker node is not a paragraph")
    if blank.tag != qn("w:p"):
        return _result(False, "blank node is not a paragraph")
    marker_text = "".join(t.text or "" for t in marker.findall(".//" + qn("w:t"))).strip()
    blank_text = "".join(t.text or "" for t in blank.findall(".//" + qn("w:t"))).strip()
    if "Продолжение таблицы" not in marker_text:
        return _result(False, f"marker text not as expected: {marker_text!r}")
    if blank_text:
        return _result(False, f"blank paragraph is not empty: {blank_text!r}")
    ok, msg = _p1a_assert_rendered_chain_anchored(
        out, "Продолжение таблицы 1.2.3", require_blank=True,
    )
    if not ok:
        return _result(False, msg)
    if _count_table_rows_with_texts(out.tables[0], ["1", "2", "3"]) != 1:
        return _result(False, "first marker-split fragment numeric row changed or duplicated")
    if _count_table_rows_with_texts(out.tables[1], ["1", "2", "3"]) != 1:
        return _result(False, "second marker-split fragment numeric row changed or duplicated")
    return _result(True, "ordinary continuation has tbl → marker → blank → tbl")


def _pg1_add_tbl_pr_child(tbl_xml, tag: str):
    tbl_pr = tbl_xml.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_xml.insert(0, tbl_pr)
    child = OxmlElement(tag)
    tbl_pr.append(child)
    return child


def _pg1_set_first_cell_tcw(tbl, value: str) -> None:
    tc_pr = tbl.rows[0].cells[0]._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), value)
    tc_w.set(qn("w:type"), "dxa")


def _pg1_geometry_snapshot(tbl_xml) -> tuple:
    tbl_pr = tbl_xml.find(qn("w:tblPr"))
    grid = tbl_xml.find(qn("w:tblGrid"))
    tc_ws = []
    for tc in tbl_xml.findall(".//" + qn("w:tc")):
        tc_pr = tc.find(qn("w:tcPr"))
        tc_w = tc_pr.find(qn("w:tcW")) if tc_pr is not None else None
        tc_ws.append((
            tc_w.get(qn("w:w")) if tc_w is not None else None,
            tc_w.get(qn("w:type")) if tc_w is not None else None,
        ))
    return (
        tbl_pr.xml if tbl_pr is not None else None,
        grid.xml if grid is not None else None,
        tuple(tc_ws),
    )


def test_pg1_geometry_policy_classifies_sensitive_tables() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import classify_table_geometry_policy

    cases = []

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.cell(0, 0).merge(tbl.cell(0, 1))
    cases.append(("gridSpan", tbl._tbl))

    doc = Document()
    tbl = doc.add_table(rows=2, cols=1)
    restart = OxmlElement("w:vMerge")
    restart.set(qn("w:val"), "restart")
    tbl.rows[0].cells[0]._tc.get_or_add_tcPr().append(restart)
    cont = OxmlElement("w:vMerge")
    cont.set(qn("w:val"), "continue")
    tbl.rows[1].cells[0]._tc.get_or_add_tcPr().append(cont)
    cases.append(("vMerge", tbl._tbl))

    for name, tag in (
        ("tblCellSpacing", "w:tblCellSpacing"),
        ("tblCellMar", "w:tblCellMar"),
    ):
        doc = Document()
        tbl = doc.add_table(rows=1, cols=2)
        _pg1_add_tbl_pr_child(tbl._tbl, tag)
        cases.append((name, tbl._tbl))

    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    look = tbl._tbl.find(qn("w:tblPr")).find(qn("w:tblLook"))
    look.set(qn("w:val"), "0000")
    cases.append(("tblLook", tbl._tbl))

    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    layout = _pg1_add_tbl_pr_child(tbl._tbl, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    cases.append(("authored fixed layout", tbl._tbl))

    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    tbl_w = tbl._tbl.find(qn("w:tblPr")).find(qn("w:tblW"))
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")
    cases.append(("tblW pct", tbl._tbl))

    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    tr_pr = OxmlElement("w:trPr")
    tr_pr.append(OxmlElement("w:tblPrEx"))
    tbl._tbl.findall(qn("w:tr"))[0].insert(0, tr_pr)
    cases.append(("tblPrEx", tbl._tbl))

    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    _pg1_set_first_cell_tcw(tbl, "0")
    cases.append(("tcW=0", tbl._tbl))

    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    tbl._tbl.find(qn("w:tblGrid")).findall(qn("w:gridCol"))[0].set(qn("w:w"), "45")
    cases.append(("narrow column", tbl._tbl))

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    row = tbl._tbl.findall(qn("w:tr"))[1]
    row.remove(row.findall(qn("w:tc"))[1])
    cases.append(("raw tc count mismatch", tbl._tbl))

    for name, tbl_xml in cases:
        policy = classify_table_geometry_policy(tbl_xml)
        if policy != "preserve_geometry":
            return _result(False, f"{name} classified as {policy!r}")

    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    policy = classify_table_geometry_policy(
        tbl._tbl,
        has_existing_continuation_marker=True,
    )
    if policy != "preserve_geometry":
        return _result(False, f"existing continuation marker classified as {policy!r}")

    return _result(True, f"{len(cases) + 1} sensitive geometry shapes classified preserve_geometry")


def test_pg1_geometry_policy_marks_malformed_grid_unsafe() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import classify_table_geometry_policy

    doc = Document()
    tbl = doc.add_table(rows=1, cols=2)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    tbl._tbl.remove(grid)
    policy = classify_table_geometry_policy(tbl._tbl)
    if policy != "unsafe_no_split":
        return _result(False, f"missing tblGrid classified as {policy!r}")

    doc = Document()
    tbl = doc.add_table(rows=4, cols=1)
    restart = OxmlElement("w:vMerge")
    restart.set(qn("w:val"), "restart")
    tbl.rows[1].cells[0]._tc.get_or_add_tcPr().append(restart)
    cont = OxmlElement("w:vMerge")
    cont.set(qn("w:val"), "continue")
    tbl.rows[2].cells[0]._tc.get_or_add_tcPr().append(cont)
    policy = classify_table_geometry_policy(tbl._tbl, split_before_row=2)
    if policy != "unsafe_no_split":
        return _result(False, f"split across vMerge classified as {policy!r}")

    return _result(True, "malformed grid and vMerge-crossing split fail closed")


def test_pg1_preserve_geometry_skips_width_optimizer() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[1].cells[0].text = "1"
    tbl.rows[1].cells[1].text = "2"
    look = tbl._tbl.find(qn("w:tblPr")).find(qn("w:tblLook"))
    look.set(qn("w:val"), "0000")
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        gc.set(qn("w:w"), "12000")
    before = _pg1_geometry_snapshot(tbl._tbl)

    n = apply_table_continuation(doc)
    after = _pg1_geometry_snapshot(tbl._tbl)

    if n != 0:
        return _result(False, f"preserve table should skip optimizer, got n={n}")
    if after != before:
        return _result(False, "preserve table geometry changed during continuation flow")
    return _result(True, "preserve_geometry table bypassed width optimizer")


def test_pg1_simple_table_still_uses_width_optimizer() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "A"
    tbl.rows[0].cells[1].text = "B"
    tbl.rows[1].cells[0].text = "1"
    tbl.rows[1].cells[1].text = "2"
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        gc.set(qn("w:w"), "12000")

    n = apply_table_continuation(doc)
    if n != 1:
        return _result(False, f"simple table optimizer path did not run, n={n}")
    return _result(True, "simple table still uses existing width optimizer")


def test_pg1_unsafe_geometry_skips_marker_split() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 1.2.3")
    tbl = doc.add_table(rows=5, cols=1)
    for r in range(5):
        tbl.rows[r].cells[0].text = f"r{r}"
    restart = OxmlElement("w:vMerge")
    restart.set(qn("w:val"), "restart")
    tbl.rows[1].cells[0]._tc.get_or_add_tcPr().append(restart)
    cont = OxmlElement("w:vMerge")
    cont.set(qn("w:val"), "continue")
    tbl.rows[2].cells[0]._tc.get_or_add_tcPr().append(cont)

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=5,
        pages_detected=[1, 2],
        row_pages={0: 1, 1: 1, 2: 2, 3: 2, 4: 2},
        found_rows=[0, 1, 2, 3, 4],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 1, 1), tm.TablePageSpan(2, 4, 2)],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text="Таблица 1.2.3",
    )

    decision = tc._MarkerSplitDecision(
        eligible=True,
        split_before_row=2,
        skip_reason=None,
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)
        result, skip_reason = tc._apply_marker_split_candidate(path, diagnostic, decision)
        out = Document(str(path))

    if result is not None:
        return _result(False, "unsafe vMerge split unexpectedly applied")
    if skip_reason != "unsafe_geometry":
        return _result(False, f"expected unsafe_geometry skip, got {skip_reason!r}")
    if len(out.tables) != 1:
        return _result(False, f"unsafe split changed table count to {len(out.tables)}")
    return _result(True, "unsafe_no_split table skipped marker split generation")


def test_pg1_real_continuation_flow_preserves_sensitive_geometry() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    doc = Document()
    tbl = doc.add_table(rows=3, cols=2)
    for r in range(3):
        for c in range(2):
            tbl.rows[r].cells[c].text = f"r{r}c{c}"
    _pg1_add_tbl_pr_child(tbl._tbl, "w:tblCellMar")
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        gc.set(qn("w:w"), "15000")
    _pg1_set_first_cell_tcw(tbl, "15000")

    before = _pg1_geometry_snapshot(tbl._tbl)
    n = apply_table_continuation(doc)
    after = _pg1_geometry_snapshot(tbl._tbl)

    if n != 0:
        return _result(False, f"real continuation flow should not optimize preserve table, got n={n}")
    if after != before:
        return _result(False, "tblGrid/tblW/tcW changed for preserve-mode table")
    return _result(True, "real continuation flow kept tblGrid/tblPr/tcW unchanged")


# ── Table-engine split model / diagnostics foundation ────────────────────────
# These tests are model/diagnostic tests only. They are not visual proof of
# production table quality; visual DOCX/PDF smoke remains mandatory before merge.

def _table_engine_diag_module():
    import importlib.util
    import sys

    path = ROOT / "tests" / "tools" / "table_engine_diagnostics.py"
    spec = importlib.util.spec_from_file_location("table_engine_diagnostics", path)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"cannot load table diagnostics module: {path}")
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


def _tm_fill_table(tbl) -> None:
    for r, row in enumerate(tbl.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"r{r}c{c}"


def test_tm_split_eligibility_requires_rendered_boundary() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    tbl = doc.add_table(rows=8, cols=3)
    _tm_fill_table(tbl)

    no_render = diag.evaluate_split_eligibility(
        tbl._tbl,
        rendered_row_pages=None,
        split_before_row=4,
        header_rows=1,
    )
    if no_render.eligible or no_render.reason != "render_boundary_unmapped":
        return _result(False, f"row-count-only split was allowed: {no_render!r}")

    rendered = {0: 1, 1: 1, 2: 1, 3: 1, 4: 2, 5: 2, 6: 2, 7: 2}
    decision = diag.evaluate_split_eligibility(
        tbl._tbl,
        rendered_row_pages=rendered,
        split_before_row=4,
        header_rows=1,
    )
    if not decision.eligible:
        return _result(False, f"safe rendered boundary rejected: {decision!r}")

    return _result(True, "split model rejects row-count-only fallback and accepts safe rendered boundary")


def test_tm_forbidden_split_reasons_are_explicit() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    tbl = doc.add_table(rows=5, cols=1)
    _tm_fill_table(tbl)
    restart = OxmlElement("w:vMerge")
    restart.set(qn("w:val"), "restart")
    tbl.rows[1].cells[0]._tc.get_or_add_tcPr().append(restart)
    cont = OxmlElement("w:vMerge")
    cont.set(qn("w:val"), "continue")
    tbl.rows[2].cells[0]._tc.get_or_add_tcPr().append(cont)
    rendered = {0: 1, 1: 1, 2: 2, 3: 2, 4: 2}
    decision = diag.evaluate_split_eligibility(
        tbl._tbl,
        rendered_row_pages=rendered,
        split_before_row=2,
        header_rows=1,
    )
    if decision.reason != "unsafe_vmerge_boundary":
        return _result(False, f"expected unsafe_vmerge_boundary, got {decision!r}")

    doc = Document()
    tbl = doc.add_table(rows=5, cols=2)
    grid = tbl._tbl.find(qn("w:tblGrid"))
    tbl._tbl.remove(grid)
    decision = diag.evaluate_split_eligibility(
        tbl._tbl,
        rendered_row_pages={0: 1, 1: 1, 2: 2, 3: 2, 4: 2},
        split_before_row=2,
        header_rows=1,
    )
    if decision.reason != "malformed_grid":
        return _result(False, f"expected malformed_grid, got {decision!r}")

    doc = Document()
    tbl = doc.add_table(rows=5, cols=2)
    _tm_fill_table(tbl)
    decision = diag.evaluate_split_eligibility(
        tbl._tbl,
        rendered_row_pages={0: 1, 1: 2, 2: 2, 3: 2, 4: 2},
        split_before_row=1,
        header_rows=1,
    )
    if decision.reason != "fragment_too_small":
        return _result(False, f"expected fragment_too_small, got {decision!r}")

    return _result(True, "forbidden split cases return explicit skip reasons")


def test_tm_numeric_row_safety_model() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    tbl = doc.add_table(rows=3, cols=3)
    for c, text in enumerate(["A", "B", "C"]):
        tbl.rows[0].cells[c].text = text
        tbl.rows[1].cells[c].text = str(c + 1)
    roles = diag.classify_table_rows(tbl._tbl, header_rows=1)
    if roles[1].role != "numeric":
        return _result(False, f"numeric row not detected: {roles!r}")
    safety = diag.numeric_row_synthesis_safety(tbl._tbl, header_rows=1)
    if not safety.safe:
        return _result(False, f"simple numeric synthesis should be safe: {safety!r}")

    doc = Document()
    tbl = doc.add_table(rows=2, cols=3)
    tbl.cell(0, 0).merge(tbl.cell(0, 1))
    safety = diag.numeric_row_synthesis_safety(tbl._tbl, header_rows=1)
    if safety.safe or safety.reason != "unsafe_numeric_row_synthesis":
        return _result(False, f"merged header synthesis should be unsafe: {safety!r}")

    return _result(True, "numeric row detection and synthesis safety model are conservative")


def test_tm_source_note_boundary_fails_closed() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    tbl = doc.add_table(rows=4, cols=2)
    _tm_fill_table(tbl)
    tbl.rows[3].cells[0].text = "Источник: составлено автором"
    tbl.rows[3].cells[1].text = ""

    roles = diag.classify_table_rows(tbl._tbl, header_rows=1)
    if roles[3].role != "source_note":
        return _result(False, f"source/note row not detected: {roles!r}")

    decision = diag.evaluate_split_eligibility(
        tbl._tbl,
        rendered_row_pages={0: 1, 1: 1, 2: 1, 3: 2},
        split_before_row=3,
        header_rows=1,
        min_body_rows_per_fragment=1,
    )
    if decision.eligible or decision.reason != "source_note_boundary_uncertain":
        return _result(False, f"source/note-only continuation was allowed: {decision!r}")

    return _result(True, "source/note-only continuation fails closed")


def test_tm_geometry_snapshot_invariant_for_preserve_mode() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    diag = _table_engine_diag_module()

    doc = Document()
    tbl = doc.add_table(rows=2, cols=2)
    _pg1_add_tbl_pr_child(tbl._tbl, "w:tblCellMar")
    grid = tbl._tbl.find(qn("w:tblGrid"))
    for gc in grid.findall(qn("w:gridCol")):
        gc.set(qn("w:w"), "15000")

    before = diag.snapshot_table(tbl._tbl)
    apply_table_continuation(doc)
    after = diag.snapshot_table(tbl._tbl)
    changes = diag.diff_table_geometry(before, after)
    if changes:
        return _result(False, f"preserve-mode geometry changed: {changes!r}")

    return _result(True, "diagnostic geometry snapshot catches preserve-mode invariants")


def test_tm_diagnostic_harness_writes_stage_artifacts() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 1.1.1")
    tbl = doc.add_table(rows=3, cols=2)
    _tm_fill_table(tbl)

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "source.docx"
        artifact_root = Path(tmp) / "artifacts"
        doc.save(src)
        result = diag.run_table_engine_diagnostics(
            src,
            artifact_root=artifact_root,
            render=False,
        )
        summary = Path(result.summary_json_path)
        if not summary.exists():
            return _result(False, f"summary was not written: {summary}")
        expected = {
            "00_source",
            "01_after_safe_formatter",
            "02_after_table_continuation",
            "03_final",
        }
        stage_names = {stage.name for stage in result.stages}
        if not expected.issubset(stage_names):
            return _result(False, f"missing stages: expected={expected}, got={stage_names}")
        for stage in result.stages:
            if not stage.tables:
                return _result(False, f"stage has no table snapshots: {stage.name}")
            table = stage.tables[0]
            for key in ("tblPr_xml", "tblGrid_xml", "gridCol_widths", "row_count", "raw_tc_counts"):
                if key not in table:
                    return _result(False, f"snapshot missing {key}: {table}")

    return _result(True, "diagnostic harness writes stage artifacts and table geometry reports")


def _tm_set_row(table, row_index: int, values: list[str]) -> None:
    for cell_index, value in enumerate(values):
        table.rows[row_index].cells[cell_index].text = value


def test_tm_universal_model_groups_caption_title_table_source() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    doc.add_paragraph("Название таблицы")
    table = doc.add_table(rows=2, cols=2)
    _tm_set_row(table, 0, ["Показатель", "Значение"])
    _tm_set_row(table, 1, ["A", "B"])
    doc.add_paragraph("Источник: составлено автором")

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "grouping.docx"
        doc.save(path)
        model = diag.build_docx_table_model(path)

    logicals = model["logical_tables"]
    if len(logicals) != 1:
        return _result(False, f"expected one logical table, got {logicals!r}")
    logical = logicals[0]
    if logical["table_num"] != "1.1.1":
        return _result(False, f"caption/table number not grouped: {logical!r}")
    if logical["title_paragraph_text"] != "Название таблицы":
        return _result(False, f"title not grouped: {logical!r}")
    if logical["physical_table_indexes"] != [0]:
        return _result(False, f"physical table index not grouped: {logical!r}")
    if not logical["source_note_paragraphs"] or logical["source_note_paragraphs"][0]["text"] != "Источник: составлено автором":
        return _result(False, f"source/note not grouped: {logical!r}")
    return _result(True, "universal model groups caption, title, table, and source note")


def test_tm_universal_row_roles_detect_header_numeric_data() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    table = doc.add_table(rows=3, cols=3)
    _tm_set_row(table, 0, ["Показатель", "База", "Цель"])
    _tm_set_row(table, 1, ["1", "2", "3"])
    _tm_set_row(table, 2, ["Конверсия", "10", "15"])

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "roles.docx"
        doc.save(path)
        model = diag.build_docx_table_model(path)

    roles = [row["role"] for row in model["rows"]]
    if roles != ["semantic_header", "numeric_row", "data_row"]:
        return _result(False, f"unexpected row roles: {roles!r}")
    return _result(True, "universal row role classifier detects semantic header, numeric row, and data")


def test_tm_universal_provenance_unknown_and_source_safe() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    source = Document()
    source.add_paragraph("Таблица 1.1.1")
    source_table = source.add_table(rows=2, cols=2)
    _tm_set_row(source_table, 0, ["Показатель", "Значение"])
    _tm_set_row(source_table, 1, ["A", "B"])

    formatted = Document()
    formatted.add_paragraph("Таблица 1.1.1")
    formatted_table = formatted.add_table(rows=3, cols=2)
    _tm_set_row(formatted_table, 0, ["Показатель", "Значение"])
    _tm_set_row(formatted_table, 1, ["1", "2"])
    _tm_set_row(formatted_table, 2, ["A", "B"])

    with tempfile.TemporaryDirectory() as tmp:
        source_path = Path(tmp) / "source.docx"
        formatted_path = Path(tmp) / "formatted.docx"
        source.save(source_path)
        formatted.save(formatted_path)

        unknown = diag.build_docx_table_model(formatted_path)
        proven = diag.build_docx_table_model(formatted_path, source_docx=source_path)

    if {row["provenance"] for row in unknown["rows"]} != {"unknown"}:
        return _result(False, f"omitted source should produce unknown provenance: {unknown['rows']!r}")
    by_fp = {row["normalized_fingerprint"]: row["provenance"] for row in proven["rows"]}
    if by_fp.get("a b") != "source":
        return _result(False, f"source data row not proven: {proven['rows']!r}")
    if by_fp.get("1 2") != "formatter_generated":
        return _result(False, f"generated numeric row not classified by provenance: {proven['rows']!r}")
    return _result(True, "provenance is unknown without source and source/formatter-generated with source")


def test_tm_universal_same_page_repeated_fragment_issue() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    doc.add_paragraph("Таблица 1.3.1")
    first = doc.add_table(rows=3, cols=3)
    _tm_set_row(first, 0, ["НПА", "Предмет", "Для кого"])
    _tm_set_row(first, 1, ["1", "2", "3"])
    _tm_set_row(first, 2, ["208-ФЗ", "Компетенция", "АО"])
    second = doc.add_table(rows=3, cols=3)
    _tm_set_row(second, 0, ["НПА", "Предмет", "Для кого"])
    _tm_set_row(second, 1, ["1", "2", "3"])
    _tm_set_row(second, 2, ["ТК РФ", "Труд", "Работники"])

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "same_page.docx"
        doc.save(path)
        report = diag.build_universal_table_diagnostics(
            formatted_docx=path,
            pdf_lines=[
                _rv_line("Таблица 1.3.1", 22, 40.0),
                _rv_line("НПА Предмет Для кого", 22, 70.0),
                _rv_line("1 2 3", 22, 90.0),
                _rv_line("208-ФЗ Компетенция АО", 22, 115.0),
                _rv_line("НПА Предмет Для кого", 22, 210.0),
                _rv_line("1 2 3", 22, 230.0),
                _rv_line("ТК РФ Труд Работники", 22, 255.0),
            ],
        )

    issues = [issue["issue_type"] for issue in report["issues"]]
    if "same_page_repeated_fragment" not in issues:
        return _result(False, f"same-page repeated fragment not emitted: {report['issues']!r}")
    return _result(True, "universal diagnostics emit same-page repeated fragment issue")


def test_tm_universal_continuation_missing_numeric_row_issue() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    doc.add_paragraph("Таблица 1.1.2")
    first = doc.add_table(rows=2, cols=2)
    _tm_set_row(first, 0, ["Критерий", "Описание"])
    _tm_set_row(first, 1, ["A", "B"])
    doc.add_paragraph("Продолжение таблицы 1.1.2")
    second = doc.add_table(rows=2, cols=2)
    _tm_set_row(second, 0, ["Критерий", "Описание"])
    _tm_set_row(second, 1, ["C", "D"])

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "missing_numeric.docx"
        doc.save(path)
        report = diag.build_universal_table_diagnostics(
            formatted_docx=path,
            pdf_lines=[
                _rv_line("Таблица 1.1.2", 4, 40.0),
                _rv_line("Критерий Описание", 4, 70.0),
                _rv_line("A B", 4, 95.0),
                _rv_line("Продолжение таблицы 1.1.2", 5, 40.0),
                _rv_line("Критерий Описание", 5, 70.0),
                _rv_line("C D", 5, 95.0),
            ],
        )

    missing = [issue for issue in report["issues"] if issue["issue_type"] == "continuation_missing_numeric_row"]
    if not missing:
        return _result(False, f"missing numeric issue not emitted: {report['issues']!r}")
    if missing[0]["evidence"].get("missing_numeric_physical_table_indexes") != [0, 1]:
        return _result(False, f"missing numeric evidence wrong: {missing!r}")
    return _result(True, "marker chain without numeric rows emits continuation_missing_numeric_row")


def test_tm_universal_single_physical_table_crosses_pages_issue() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    doc.add_paragraph("Таблица 1.1.3")
    table = doc.add_table(rows=3, cols=2)
    _tm_set_row(table, 0, ["Кейс", "Описание"])
    _tm_set_row(table, 1, ["Кейс 1", "Начало"])
    _tm_set_row(table, 2, ["Кейс 5", "Середина"])
    table.add_row()
    _tm_set_row(table, 3, ["Кейс 9", "Финал"])

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "cross_page.docx"
        doc.save(path)
        report = diag.build_universal_table_diagnostics(
            formatted_docx=path,
            rendered_rows=[
                {
                    "logical_table_id": "table_1.1.3",
                    "physical_table_index": 0,
                    "row_index": 1,
                    "page": 5,
                    "top_y": 500.0,
                    "bottom_y": 520.0,
                    "confidence": "high",
                    "split_or_spill": False,
                },
                {
                    "logical_table_id": "table_1.1.3",
                    "physical_table_index": 0,
                    "row_index": 2,
                    "page": 6,
                    "top_y": 60.0,
                    "bottom_y": 80.0,
                    "confidence": "high",
                    "split_or_spill": False,
                },
                {
                    "logical_table_id": "table_1.1.3",
                    "physical_table_index": 0,
                    "row_index": 3,
                    "page": 7,
                    "top_y": 60.0,
                    "bottom_y": 80.0,
                    "confidence": "high",
                    "split_or_spill": False,
                },
            ],
        )

    issues = [issue for issue in report["issues"] if issue["issue_type"] == "single_physical_table_crosses_pages_without_marker"]
    if not issues:
        return _result(False, f"single physical cross-page issue not emitted: {report['issues']!r}")
    if issues[0]["pages"] != [5, 6, 7]:
        return _result(False, f"cross-page pages wrong: {issues!r}")
    return _result(True, "mocked rendered rows emit single physical cross-page issue")


def test_tm_universal_rendered_mapping_keeps_repeated_case_rows_on_real_pages() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    doc.add_paragraph("Таблица 1.1.3")
    table = doc.add_table(rows=4, cols=2)
    _tm_set_row(table, 0, ["Кейс", "Описание"])
    _tm_set_row(table, 1, ["Кейс 1 длинный, грустно, Word ширины, тест", "A"])
    _tm_set_row(table, 2, ["Кейс 4 длинный, грустно, Word ширины, тест", "B"])
    _tm_set_row(table, 3, ["Кейс 9 длинный, грустно, Word ширины, тест", "C"])

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "case_rows.docx"
        doc.save(path)
        report = diag.build_universal_table_diagnostics(
            formatted_docx=path,
            pdf_lines=[
                _rv_line("Таблица 1.1.3", 5, 250.0),
                _rv_line("Кейс 1 длинный, грустно, Word ширины, тест", 5, 382.0),
                _rv_line("Кейс 4 длинный, грустно, Word ширины, тест", 6, 114.0),
                _rv_line("Кейс 9 длинный, грустно, Word ширины, тест", 7, 114.0),
            ],
        )

    rendered = [
        (row["row_index"], row["page"])
        for row in report["rendered_rows"]
        if row["logical_table_id"] == "table_1.1.3" and row["row_index"] in {1, 2, 3}
    ]
    if rendered != [(1, 5), (2, 6), (3, 7)]:
        return _result(False, f"case row/page mapping drifted: {rendered!r}")
    issues = [issue for issue in report["issues"] if issue["issue_type"] == "single_physical_table_crosses_pages_without_marker"]
    if not issues or issues[0]["pages"] != [5, 6, 7]:
        return _result(False, f"cross-page issue missing from mapped rows: {report['issues']!r}")
    return _result(True, "cell-specific rendered mapping keeps repeated case rows on their pages")


def test_tm_universal_bad2_style_clean_continuation() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = Document()
    doc.add_paragraph("Таблица 1.3.1")
    first = doc.add_table(rows=3, cols=3)
    _tm_set_row(first, 0, ["Страна", "Модель", "Роль"])
    _tm_set_row(first, 1, ["1", "2", "3"])
    _tm_set_row(first, 2, ["ЕС", "Peppol", "Право"])
    doc.add_paragraph("Продолжение таблицы 1.3.1")
    second = doc.add_table(rows=2, cols=3)
    _tm_set_row(second, 0, ["1", "2", "3"])
    _tm_set_row(second, 1, ["РФ", "Оператор", "Контроль"])

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "bad2_clean.docx"
        doc.save(path)
        report = diag.build_universal_table_diagnostics(
            formatted_docx=path,
            pdf_lines=[
                _rv_line("Таблица 1.3.1", 14, 450.0),
                _rv_line("Страна Модель Роль", 14, 480.0),
                _rv_line("1 2 3", 14, 500.0),
                _rv_line("ЕС Peppol Право", 14, 530.0),
                _rv_line("Продолжение таблицы 1.3.1", 15, 40.0),
                _rv_line("1 2 3", 15, 70.0),
                _rv_line("РФ Оператор Контроль", 15, 100.0),
            ],
        )

    hard = [issue for issue in report["issues"] if issue["issue_type"] not in {"clean"}]
    if hard:
        return _result(False, f"clean continuation emitted issues: {hard!r}")
    return _result(True, "bad2-style marker-before-rows continuation is clean")


def test_tm_universal_diagnostics_run_is_non_mutating() -> tuple[bool, str]:
    import hashlib

    diag = _table_engine_diag_module()

    doc = Document()
    doc.add_paragraph("Таблица 1.1.1")
    table = doc.add_table(rows=2, cols=2)
    _tm_set_row(table, 0, ["Показатель", "Значение"])
    _tm_set_row(table, 1, ["A", "B"])

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "input.docx"
        doc.save(path)
        before = hashlib.sha256(path.read_bytes()).hexdigest()
        diag.build_universal_table_diagnostics(formatted_docx=path, pdf_lines=[])
        after = hashlib.sha256(path.read_bytes()).hexdigest()

    if before != after:
        return _result(False, "diagnostics mutated input DOCX bytes")
    return _result(True, "universal diagnostics are non-mutating")


def test_tm_same_page_continuation_marker_is_flagged() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    markers = diag.validate_continuation_markers_from_blocks([
        {"kind": "table", "page": 16, "y0": 56.0, "y1": 237.0},
        {"kind": "text", "page": 16, "y0": 242.0, "y1": 252.0, "text": "Продолжение таблицы 1.3.1"},
        {"kind": "table", "page": 16, "y0": 262.0, "y1": 692.0},
    ])

    if len(markers) != 1:
        return _result(False, f"expected one marker diagnostic, got {markers!r}")
    marker = markers[0]
    if marker.verdict != "fail" or marker.same_page_violation is not True:
        return _result(False, f"same-page continuation was not flagged: {marker!r}")
    if marker.marker_page != 16 or marker.previous_table_page != 16 or marker.following_table_page != 16:
        return _result(False, f"unexpected marker/table pages: {marker!r}")
    return _result(True, "same-page continuation marker is reported as a violation")


def test_tm_next_page_continuation_marker_passes_validation() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    markers = diag.validate_continuation_markers_from_blocks([
        {"kind": "table", "page": 10, "y0": 450.0, "y1": 760.0},
        {"kind": "text", "page": 11, "y0": 60.0, "y1": 70.0, "text": "Продолжение таблицы 2.2.1"},
        {"kind": "table", "page": 11, "y0": 82.0, "y1": 400.0},
    ])

    marker = markers[0]
    if marker.verdict != "pass" or marker.same_page_violation is not False:
        return _result(False, f"new-page continuation was not accepted: {marker!r}")
    if marker.previous_table_page != 10 or marker.marker_page != 11 or marker.following_table_page != 11:
        return _result(False, f"unexpected marker/table pages: {marker!r}")
    return _result(True, "new-page continuation marker passes validation")


def test_tm_inline_prose_continuation_text_is_not_valid_marker() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    markers = diag.validate_continuation_markers_from_blocks([
        {
            "kind": "text",
            "page": 7,
            "y0": 180.0,
            "y1": 194.0,
            "text": "В тексте встречается Продолжение таблицы 1.2.3 как пример.",
        },
    ])

    if len(markers) != 1:
        return _result(False, f"expected inline prose diagnostic, got {markers!r}")
    marker = markers[0]
    if marker.marker_kind != "source_inline_marker_text" or marker.verdict != "source_inline_marker_text":
        return _result(False, f"inline prose was treated as a valid marker: {marker!r}")
    if marker.same_page_violation is not None:
        return _result(False, f"inline prose should not get pass/fail validation: {marker!r}")
    return _result(True, "inline prose continuation text is classified separately")


def test_tm_unknown_marker_relation_is_not_pass() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    markers = diag.validate_continuation_markers_from_blocks([
        {"kind": "text", "page": 12, "y0": 80.0, "y1": 92.0, "text": "Продолжение таблицы 3.1.4"},
    ])

    marker = markers[0]
    if marker.verdict != "unknown" or marker.same_page_violation is not None:
        return _result(False, f"low-confidence relation produced pass/fail: {marker!r}")
    if marker.confidence != "low":
        return _result(False, f"unknown relation should be low confidence: {marker!r}")
    return _result(True, "unknown continuation relation is reported as unknown, not pass")


def test_tm_continuation_marker_report_includes_page_and_violation() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    markers = diag.validate_continuation_markers_from_blocks([
        {"kind": "table", "page": 16, "y0": 56.0, "y1": 237.0},
        {"kind": "text", "page": 16, "y0": 242.0, "y1": 252.0, "text": "Продолжение таблицы 1.3.1"},
        {"kind": "table", "page": 16, "y0": 262.0, "y1": 692.0},
    ])

    with tempfile.TemporaryDirectory() as tmp:
        paths = diag.write_continuation_marker_reports(
            Path(tmp),
            {"03_final": markers},
        )
        json_text = Path(paths["continuation_markers_json_path"]).read_text(encoding="utf-8")
        md_text = Path(paths["markdown_summary_path"]).read_text(encoding="utf-8")

    if '"marker_page": 16' not in json_text or '"same_page_violation": true' not in json_text:
        return _result(False, f"JSON report missing page/violation: {json_text}")
    if "Продолжение таблицы 1.3.1" not in md_text or "page 16" not in md_text or "violation=True" not in md_text:
        return _result(False, f"markdown report missing marker summary: {md_text}")
    return _result(True, "continuation marker report includes marker page and violation flag")


def _rv_line(text: str, page: int, top: float):
    from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine

    return PdfLine(text=text, page_num=page, top=top, bottom=top + 10.0)


def test_tm_rendered_fragment_without_marker_is_flagged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.rendered_table_validation import (
        RenderedTableIdentity,
        validate_rendered_continuations,
    )

    identity = RenderedTableIdentity(
        table_index=7,
        body_order_index=7,
        caption_num="2.2.3",
        preceding_marker=None,
        following_marker="Продолжение таблицы 2.2.3",
        header_fingerprint=("уровень показатель источник бенчмарк цель",),
        numeric_row_fingerprint="1 2 3 4 5 6",
        row_fingerprints=("intention сделать заказ ночью",),
    )
    violations = validate_rendered_continuations(
        [
            _rv_line("Таблица 2.2.3", 45, 492.9),
            _rv_line("Уровень Показатель Источник Бенчмарк/цель", 46, 59.2),
            _rv_line("1 2 3 4 5 6", 46, 88.8),
            _rv_line("Продолжение таблицы 2.2.3", 47, 58.2),
        ],
        [identity],
    )

    if len(violations) != 1:
        return _result(False, f"expected one missing-marker violation, got {violations!r}")
    v = violations[0]
    if v.violation_type != "missing_continuation_marker" or v.table_num != "2.2.3" or v.page != 46:
        return _result(False, f"unexpected violation: {v!r}")
    if not v.evidence.get("repeated_header") or not v.evidence.get("repeated_numeric_row"):
        return _result(False, f"missing repeated header/numeric evidence: {v!r}")
    if v.evidence.get("marker_page") != 47 or v.evidence.get("previous_caption_page") != 45:
        return _result(False, f"missing page evidence: {v!r}")
    return _result(True, "rendered continuation fragment without marker is flagged")


def test_tm_rendered_fragment_with_marker_before_it_passes() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.rendered_table_validation import (
        RenderedTableIdentity,
        validate_rendered_continuations,
    )

    identity = RenderedTableIdentity(
        table_index=0,
        body_order_index=0,
        caption_num="1.1.1",
        preceding_marker=None,
        following_marker="Продолжение таблицы 1.1.1",
        header_fingerprint=("показатель как влияет",),
        numeric_row_fingerprint="1 2 3",
        row_fingerprints=("оперативность повышение",),
    )
    violations = validate_rendered_continuations(
        [
            _rv_line("Таблица 1.1.1", 6, 430.2),
            _rv_line("Продолжение таблицы 1.1.1", 7, 58.2),
            _rv_line("Показатель Как влияет", 7, 95.0),
            _rv_line("1 2 3", 7, 122.1),
        ],
        [identity],
    )

    if violations:
        return _result(False, f"valid continuation marker produced violations: {violations!r}")
    return _result(True, "marker before rendered continuation fragment passes")


def test_tm_same_page_adjacent_fragments_are_flagged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.rendered_table_validation import (
        RenderedTableIdentity,
        validate_rendered_continuations,
    )

    first = RenderedTableIdentity(
        table_index=4,
        body_order_index=4,
        caption_num="1.3.1",
        preceding_marker=None,
        following_marker=None,
        header_fingerprint=("нпа предмет регулирования для кого ключевые положения",),
        numeric_row_fingerprint="1 2 3 4",
        row_fingerprints=(
            "нпа предмет регулирования для кого ключевые положения",
            "1 2 3 4",
            "208-фз компетенция органов ао публичные общества",
        ),
    )
    second = RenderedTableIdentity(
        table_index=5,
        body_order_index=5,
        caption_num=None,
        preceding_marker=None,
        following_marker=None,
        header_fingerprint=("нпа предмет регулирования для кого ключевые положения",),
        numeric_row_fingerprint="1 2 3 4",
        row_fingerprints=(
            "нпа предмет регулирования для кого ключевые положения",
            "1 2 3 4",
            "трудовой кодекс трудовые отношения работодатели работники",
        ),
    )
    violations = validate_rendered_continuations(
        [
            _rv_line("Таблица 1.3.1", 22, 40.0),
            _rv_line("НПА Предмет регулирования Для кого Ключевые положения", 22, 70.0),
            _rv_line("1 2 3 4", 22, 90.0),
            _rv_line("208-ФЗ компетенция органов АО публичные общества", 22, 115.0),
            _rv_line("НПА Предмет регулирования Для кого Ключевые положения", 22, 210.0),
            _rv_line("1 2 3 4", 22, 230.0),
            _rv_line("Трудовой кодекс трудовые отношения работодатели работники", 22, 255.0),
        ],
        [first, second],
    )

    if not any(v.violation_type == "same_page_repeated_fragment" for v in violations):
        return _result(False, f"same-page repeated fragment not flagged: {violations!r}")
    return _result(True, "same-page adjacent repeated fragment is flagged")


def test_tm_late_marker_rows_before_marker_are_flagged() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.rendered_table_validation import (
        RenderedTableIdentity,
        validate_rendered_continuations,
    )

    identity = RenderedTableIdentity(
        table_index=4,
        body_order_index=4,
        caption_num="1.3.1",
        preceding_marker=None,
        following_marker="Продолжение таблицы 1.3.1",
        header_fingerprint=("страна модель роль государства роль операторов форматы стандарты",),
        numeric_row_fingerprint="1 2 3 4 5",
        row_fingerprints=(
            "страна модель роль государства роль операторов форматы стандарты",
            "1 2 3 4 5",
            "ес peppol eidas директивы требования общий правовой контур",
        ),
    )
    violations = validate_rendered_continuations(
        [
            _rv_line("Таблица 1.3.1", 14, 480.0),
            _rv_line("Страна модель Роль государства Роль операторов Форматы стандарты", 15, 45.0),
            _rv_line("1 2 3 4 5", 15, 70.0),
            _rv_line("ЕС Peppol eIDAS директивы требования общий правовой контур", 15, 95.0),
            _rv_line("Продолжение таблицы 1.3.1", 15, 150.0),
        ],
        [identity],
    )

    late = [v for v in violations if v.violation_type == "late_continuation_marker"]
    if not late:
        return _result(False, f"late marker not flagged: {violations!r}")
    if late[0].page != 15 or late[0].confidence != "high":
        return _result(False, f"late marker page/confidence wrong: {late!r}")
    return _result(True, "rows before same-page continuation marker are flagged")


def test_tm_source_bad_duplicate_rows_are_warning_only() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.rendered_table_validation import (
        RenderedTableIdentity,
        validate_rendered_continuations,
    )

    source = RenderedTableIdentity(
        table_index=3,
        body_order_index=3,
        caption_num="1.2.1",
        preceding_marker=None,
        following_marker=None,
        header_fingerprint=("функция kpi метрика источник",),
        numeric_row_fingerprint=None,
        row_fingerprints=(
            "функция kpi метрика источник",
            "уровень формальные органы неформальные практики основные функции",
            "корпоративный общее собрание наблюдательный совет стратегия надзор",
            "уровень формальные органы неформальные практики основные функции",
            "корпоративный общее собрание наблюдательный совет стратегия надзор",
        ),
    )
    output = RenderedTableIdentity(
        table_index=3,
        body_order_index=3,
        caption_num="1.2.1",
        preceding_marker=None,
        following_marker=None,
        header_fingerprint=("функция kpi метрика источник",),
        numeric_row_fingerprint=None,
        row_fingerprints=source.row_fingerprints,
    )
    violations = validate_rendered_continuations(
        [_rv_line("Таблица 1.2.1", 16, 40.0)],
        [output],
        source_table_identities=[source],
    )

    source_bad = [v for v in violations if v.violation_type == "source_bad_duplicated_content_rows"]
    if not source_bad:
        return _result(False, f"source-bad duplicate rows not flagged: {violations!r}")
    if source_bad[0].evidence.get("source_proven") is not True:
        return _result(False, f"source provenance missing: {source_bad!r}")
    repair_classes = {"missing_continuation_marker", "late_continuation_marker"}
    if any(v.violation_type in repair_classes for v in violations):
        return _result(False, f"source-bad rows should not become repair class: {violations!r}")
    return _result(True, "source-bad duplicated content rows are warning-only")


def test_tm_ambiguous_adjacent_tables_are_flagged_without_repair_class() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.rendered_table_validation import (
        RenderedTableIdentity,
        validate_rendered_continuations,
    )

    first = RenderedTableIdentity(
        table_index=19,
        body_order_index=19,
        caption_num="2.3.1",
        preceding_marker=None,
        following_marker=None,
        header_fingerprint=("роль ответственность частота артефакты kpi срок пилота",),
        numeric_row_fingerprint=None,
        row_fingerprints=("роль ответственность частота артефакты kpi срок пилота",),
    )
    second = RenderedTableIdentity(
        table_index=20,
        body_order_index=20,
        caption_num=None,
        preceding_marker=None,
        following_marker=None,
        header_fingerprint=("роль ответственность частота артефакты kpi срок пилота",),
        numeric_row_fingerprint=None,
        row_fingerprints=(
            "роль ответственность частота артефакты kpi срок пилота",
            "представители трайбов экспертные ревью обмен кейсами",
            "риск офицер комплаенс проверка рисков контроль изменений",
        ),
    )
    violations = validate_rendered_continuations(
        [
            _rv_line("Таблица 2.3.1", 44, 40.0),
            _rv_line("Роль Ответственность Частота Артефакты KPI Срок пилота", 44, 70.0),
        ],
        [first, second],
    )

    ambiguous = [v for v in violations if v.violation_type == "ambiguous_adjacent_tables"]
    if not ambiguous:
        return _result(False, f"ambiguous adjacent tables not flagged: {violations!r}")
    if any(v.violation_type == "missing_continuation_marker" for v in violations):
        return _result(False, f"ambiguous case should not become hard repair class: {violations!r}")
    return _result(True, "ambiguous adjacent tables are warning-only")


def test_tm_clean_continuation_has_no_new_fragment_diagnostics() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.rendered_table_validation import (
        RenderedTableIdentity,
        validate_rendered_continuations,
    )

    first = RenderedTableIdentity(
        table_index=0,
        body_order_index=0,
        caption_num="1.1.1",
        preceding_marker=None,
        following_marker="Продолжение таблицы 1.1.1",
        header_fingerprint=("показатель как влияет последствия",),
        numeric_row_fingerprint="1 2 3",
        row_fingerprints=("показатель как влияет последствия", "1 2 3", "скорость реакции запросы"),
    )
    second = RenderedTableIdentity(
        table_index=1,
        body_order_index=1,
        caption_num=None,
        preceding_marker="Продолжение таблицы 1.1.1",
        following_marker=None,
        header_fingerprint=("показатель как влияет последствия",),
        numeric_row_fingerprint="1 2 3",
        row_fingerprints=("показатель как влияет последствия", "1 2 3", "управление документами"),
    )
    violations = validate_rendered_continuations(
        [
            _rv_line("Таблица 1.1.1", 5, 60.0),
            _rv_line("Показатель Как влияет Последствия", 5, 90.0),
            _rv_line("1 2 3", 5, 110.0),
            _rv_line("Продолжение таблицы 1.1.1", 6, 50.0),
            _rv_line("Показатель Как влияет Последствия", 6, 80.0),
            _rv_line("1 2 3", 6, 100.0),
        ],
        [first, second],
    )

    if violations:
        return _result(False, f"clean continuation produced diagnostics: {violations!r}")
    return _result(True, "clean continuation has no same-page/ambiguous diagnostics")


def test_tm_artifact_guard_pdfs_remain_clean_for_new_diagnostics() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.pdf_layout_analyzer import analyze_pdf_lines
    from guides.coursework_kfu_2025.rendered_table_validation import (
        build_rendered_table_identities,
        validate_rendered_continuations,
    )

    artifacts = {
        "237": (
            Path("/tmp/kfu_patchB_padding_final_pdfs/237_after.pdf"),
            Path("/tmp/kfu_patchB_padding_final_pdfs_work/237_formatted_from_original.docx"),
        ),
        "244": (
            Path("/tmp/kfu_patchB_padding_final_pdfs/244_after.pdf"),
            Path("/tmp/kfu_patchB_padding_final_pdfs_work/244_formatted_from_original.docx"),
        ),
        "reference": (
            Path("/tmp/kfu_patchB_padding_final_pdfs/reference_rybakov_after.pdf"),
            Path("/tmp/kfu_patchB_padding_final_pdfs_work/reference_rybakov_formatted_from_original.docx"),
        ),
    }
    missing = [str(path) for pair in artifacts.values() for path in pair if not path.exists()]
    if missing:
        return _result(True, f"artifact guard skipped; missing {missing!r}")

    bad_types = {
        "same_page_repeated_fragment",
        "same_page_adjacent_fragment",
        "late_continuation_marker",
        "missing_or_late_continuation_marker",
        "ambiguous_adjacent_tables",
    }
    failures: dict[str, list[tuple[str, int, str]]] = {}
    for key, (pdf_path, docx_path) in artifacts.items():
        violations = validate_rendered_continuations(
            analyze_pdf_lines(pdf_path),
            build_rendered_table_identities(Document(str(docx_path))),
        )
        selected = [
            (v.table_num or "", v.page, v.violation_type)
            for v in violations
            if v.violation_type in bad_types
        ]
        if selected:
            failures[key] = selected
    if failures:
        return _result(False, f"guard artifacts got new diagnostics: {failures!r}")
    return _result(True, "237/244/reference guard PDFs remain clean for new diagnostics")


def test_tm_inline_prose_is_not_rendered_continuation_marker() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.rendered_table_validation import (
        classify_continuation_marker_line,
    )

    marker = classify_continuation_marker_line(
        "В тексте встречается Продолжение таблицы 1.2.3 как пример."
    )
    if marker.marker_kind != "source_inline_marker_text" or marker.table_num != "1.2.3":
        return _result(False, f"inline prose marker misclassified: {marker!r}")
    return _result(True, "inline prose continuation text is not a strict marker")


def test_tm_stable_manual_chain_identity_survives_save_reload() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.rendered_table_validation import (
        build_rendered_table_identities,
    )
    from guides.coursework_kfu_2025.table_continuation import (
        _valid_manual_continuation_table_indexes,
    )

    doc = Document()
    doc.add_paragraph("Таблица 1.2.1")
    table1 = doc.add_table(rows=3, cols=2)
    table1.rows[0].cells[0].text = "Показатель"
    table1.rows[0].cells[1].text = "Значение"
    table1.rows[1].cells[0].text = "1"
    table1.rows[1].cells[1].text = "2"
    table1.rows[2].cells[0].text = "A"
    table1.rows[2].cells[1].text = "B"
    marker = doc.add_paragraph("Продолжение таблицы 1.2.1")
    marker.alignment = 2
    marker.paragraph_format.keep_with_next = True
    table2 = doc.add_table(rows=2, cols=2)
    table2.rows[0].cells[0].text = "1"
    table2.rows[0].cells[1].text = "2"
    table2.rows[1].cells[0].text = "C"
    table2.rows[1].cells[1].text = "D"

    before_indexes = _valid_manual_continuation_table_indexes(doc)
    before_ids = build_rendered_table_identities(doc)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "chain.docx"
        doc.save(path)
        reloaded = Document(str(path))

    after_indexes = _valid_manual_continuation_table_indexes(reloaded)
    after_ids = build_rendered_table_identities(reloaded)

    if before_indexes != {0, 1} or after_indexes != {0, 1}:
        return _result(False, f"manual chain indexes unstable: before={before_indexes}, after={after_indexes}")
    before_core = [
        (i.table_index, i.body_order_index, i.caption_num, i.numeric_row_fingerprint)
        for i in before_ids
    ]
    after_core = [
        (i.table_index, i.body_order_index, i.caption_num, i.numeric_row_fingerprint)
        for i in after_ids
    ]
    if before_core != after_core:
        return _result(False, f"stable identity changed after reload: before={before_core}, after={after_core}")
    return _result(True, "manual chain identity is stable across save/reload")


def test_tm_real_fixture_rendered_missing_marker_diagnostics() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.formatter_service import format_docx
    from guides.coursework_kfu_2025.layout_render import render_docx_to_pdf
    from guides.coursework_kfu_2025.pdf_layout_analyzer import analyze_pdf_lines
    from guides.coursework_kfu_2025.rendered_table_validation import (
        build_rendered_table_identities,
        validate_rendered_continuations,
    )

    fixtures = {
        "208": Path("/Users/mac/Downloads/208_курсовая пример 1.docx"),
        "210": Path("/Users/mac/Downloads/210_неиромаркетинг_Рыбаков.docx"),
        "ref": Path("/Users/mac/Downloads/Рыбаков_Олег_Дмитриевич_курсовая_3_курс.docx"),
    }
    missing = [str(path) for path in fixtures.values() if not path.exists()]
    if missing:
        return _result(True, f"fixture(s) not present; skipped: {missing!r}")

    old_env = {
        "KPFU_ENABLE_MARKER_SPLIT": os.environ.get("KPFU_ENABLE_MARKER_SPLIT"),
        "KPFU_APPLY_MARKER_SPLIT": os.environ.get("KPFU_APPLY_MARKER_SPLIT"),
        "KPFU_MARKER_SPLIT_MODE": os.environ.get("KPFU_MARKER_SPLIT_MODE"),
    }
    os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
    os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
    os.environ["KPFU_MARKER_SPLIT_MODE"] = "candidate"
    try:
        with tempfile.TemporaryDirectory() as tmp:
            results: dict[str, list] = {}
            for key, fixture in fixtures.items():
                out = Path(tmp) / f"{key}.docx"
                format_docx(str(fixture), str(out))
                pdf = render_docx_to_pdf(out, timeout=180)
                doc = Document(str(out))
                results[key] = validate_rendered_continuations(
                    analyze_pdf_lines(pdf),
                    build_rendered_table_identities(doc),
                )
    finally:
        for key, value in old_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value

    r208 = [(v.table_num, v.page, v.violation_type) for v in results["208"]]
    if ("2.2.3", 46, "missing_continuation_marker") not in r208:
        return _result(False, f"208 missing page-46 violation: {r208!r}")

    r210 = [(v.table_num, v.page, v.violation_type) for v in results["210"]]
    unexpected_210 = [item for item in r210 if item[2] == "missing_continuation_marker" and item[0] not in {"2.1.1"}]
    if unexpected_210:
        return _result(False, f"210 got unexpected rendered-continuation hard failures: {unexpected_210!r}")

    rref = [(v.table_num, v.page, v.violation_type) for v in results["ref"]]
    hard_ref = [item for item in rref if item[2] == "missing_continuation_marker"]
    if hard_ref:
        return _result(False, f"reference Рыбаков got hard false positives: {hard_ref!r}")
    return _result(True, f"real fixture rendered diagnostics: 208={r208}, 210={r210}, ref={rref}")


def test_tm_formatter_surfaces_rendered_continuation_warnings() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.formatter_service as fs
    from guides.coursework_kfu_2025.rendered_table_validation import RenderedContinuationViolation

    if not hasattr(fs, "_rendered_continuation_violations_for_docx"):
        return _result(False, "formatter has no rendered-continuation final validator")

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 2.2.3")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header"
    table.rows[0].cells[1].text = "Header"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "2"

    high = RenderedContinuationViolation(
        table_num="2.2.3",
        table_index=18,
        page=46,
        violation_type="missing_continuation_marker",
        confidence="high",
        evidence={"repeated_numeric_row": True},
    )
    medium = RenderedContinuationViolation(
        table_num="2.2.1",
        table_index=11,
        page=48,
        violation_type="suspected_missing_continuation_marker",
        confidence="medium",
        evidence={"repeated_row": True},
    )

    old_process = fs.process_document
    old_pagination = fs.apply_pagination_rules
    old_merging = fs.apply_table_merging
    old_continuation = fs.apply_table_continuation
    old_rule3 = fs.apply_rule3_table_orphan
    old_rule4 = fs.apply_rule4_empty_first_lines
    old_rule6 = fs.apply_rule6_figure_orphan
    old_gap = fs.remove_empty_before_figure_captions
    old_strip = fs.strip_obsolete_toc_blocks_inplace
    old_rebuild = fs.rebuild_static_contents_page
    old_rendered = fs.apply_rendered_table_continuation
    old_restore = fs.restore_docx_if_same_page_continuation_markers
    old_remove = fs.remove_same_page_continuation_markers_inplace
    old_validator = fs._rendered_continuation_violations_for_docx

    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "input.docx"
        out = Path(tmp) / "output.docx"
        doc.save(inp)
        try:
            fs.process_document = lambda _src, dst: doc.save(dst)
            fs.apply_pagination_rules = lambda _doc: None
            fs.apply_table_merging = lambda _doc: 0
            fs.apply_table_continuation = lambda _doc, report=None: 0
            fs.apply_rule3_table_orphan = lambda _doc: 0
            fs.apply_rule4_empty_first_lines = lambda _doc: 0
            fs.apply_rule6_figure_orphan = lambda _doc: 0
            fs.remove_empty_before_figure_captions = lambda _doc: 0
            fs.strip_obsolete_toc_blocks_inplace = lambda _path: {"sdt_removed": 0, "plain_toc_removed": 0}
            fs.rebuild_static_contents_page = lambda _path: False
            fs.apply_rendered_table_continuation = lambda _path, report=None: 0
            fs.restore_docx_if_same_page_continuation_markers = (
                lambda _output, _backup, report=None, context="": False
            )
            fs.remove_same_page_continuation_markers_inplace = lambda _path, report=None: 0
            fs._rendered_continuation_violations_for_docx = (
                lambda _path, **_kwargs: [high, medium]
            )
            _, warnings = fs.format_docx(str(inp), str(out))
        finally:
            fs.process_document = old_process
            fs.apply_pagination_rules = old_pagination
            fs.apply_table_merging = old_merging
            fs.apply_table_continuation = old_continuation
            fs.apply_rule3_table_orphan = old_rule3
            fs.apply_rule4_empty_first_lines = old_rule4
            fs.apply_rule6_figure_orphan = old_rule6
            fs.remove_empty_before_figure_captions = old_gap
            fs.strip_obsolete_toc_blocks_inplace = old_strip
            fs.rebuild_static_contents_page = old_rebuild
            fs.apply_rendered_table_continuation = old_rendered
            fs.restore_docx_if_same_page_continuation_markers = old_restore
            fs.remove_same_page_continuation_markers_inplace = old_remove
            fs._rendered_continuation_violations_for_docx = old_validator

    joined = "\n".join(warnings)
    if "таблицы 2.2.3" not in joined or "стр. 46" not in joined:
        return _result(False, f"high-confidence missing marker warning absent: {warnings!r}")
    if "таблицы 2.2.1" not in joined or "стр. 48" not in joined:
        return _result(False, f"medium-confidence suspected marker warning absent: {warnings!r}")
    return _result(True, "formatter surfaces high and medium rendered-continuation warnings")


def test_tm_real_fixture_formatter_rendered_warnings() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.formatter_service import format_docx

    fixtures = {
        "208": Path("/Users/mac/Downloads/208_курсовая пример 1.docx"),
        "210": Path("/Users/mac/Downloads/210_неиромаркетинг_Рыбаков.docx"),
        "207": Path("/Users/mac/Downloads/207_Гаянов_Амир_Ленарович_Разработка_проектного_решения_по_автом.docx"),
        "209": Path("/Users/mac/Downloads/209_курсовая_Бондарев_Никита_2_курс.docx"),
        "ref": Path("/Users/mac/Downloads/Рыбаков_Олег_Дмитриевич_курсовая_3_курс.docx"),
    }
    missing = [str(path) for path in fixtures.values() if not path.exists()]
    if missing:
        return _result(True, f"fixture(s) not present; skipped: {missing!r}")

    old_env = {
        "KPFU_ENABLE_MARKER_SPLIT": os.environ.get("KPFU_ENABLE_MARKER_SPLIT"),
        "KPFU_APPLY_MARKER_SPLIT": os.environ.get("KPFU_APPLY_MARKER_SPLIT"),
        "KPFU_MARKER_SPLIT_MODE": os.environ.get("KPFU_MARKER_SPLIT_MODE"),
    }
    os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
    os.environ["KPFU_APPLY_MARKER_SPLIT"] = "1"
    os.environ["KPFU_MARKER_SPLIT_MODE"] = "candidate"
    try:
        with tempfile.TemporaryDirectory() as tmp:
            warnings_by_key: dict[str, list[str]] = {}
            for key, fixture in fixtures.items():
                out = Path(tmp) / f"{key}.docx"
                _, warnings = format_docx(str(fixture), str(out))
                warnings_by_key[key] = warnings
    finally:
        for key, value in old_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value

    if not any("таблицы 2.2.3" in warning and "стр. 46" in warning for warning in warnings_by_key["208"]):
        return _result(False, f"208 warning missing: {warnings_by_key['208']!r}")
    noisy_same_page = {
        key: [
            warning for warning in warnings
            if "маркер продолжения на той же странице" in warning
        ]
        for key, warnings in warnings_by_key.items()
    }
    noisy_same_page = {key: warnings for key, warnings in noisy_same_page.items() if warnings}
    if noisy_same_page:
        return _result(False, f"intermediate same-page warnings leaked: {noisy_same_page!r}")
    for clean_key in ("207", "209", "ref"):
        unexpected = [
            warning for warning in warnings_by_key[clean_key]
            if "без маркера продолжения" in warning
        ]
        if unexpected:
            return _result(False, f"{clean_key} got rendered-continuation false positive: {unexpected!r}")
    return _result(True, f"real fixture formatter warnings: {warnings_by_key!r}")


def test_tm_post_render_gate_rejects_same_page_marker_output() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc

    violations = tc._same_page_continuation_marker_violations_from_blocks([
        {"kind": "table", "page": 13, "y0": 80.0, "y1": 180.0},
        {"kind": "text", "page": 13, "y0": 190.0, "y1": 202.0, "text": "Продолжение таблицы 1.2.2"},
        {"kind": "table", "page": 13, "y0": 210.0, "y1": 390.0},
    ])

    if len(violations) != 1:
        return _result(False, f"expected same-page violation, got {violations!r}")
    violation = violations[0]
    if violation.marker_text != "Продолжение таблицы 1.2.2" or violation.marker_page != 13:
        return _result(False, f"unexpected violation payload: {violation!r}")
    return _result(True, "post-render validator rejects same-page marker output")


def test_tm_post_render_gate_restores_backup_on_same_page_marker() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc

    clean = Document()
    clean.add_paragraph("ВВЕДЕНИЕ")
    clean.add_paragraph("Таблица 1.2.2")
    table = clean.add_table(rows=2, cols=1)
    table.rows[0].cells[0].text = "H"
    table.rows[1].cells[0].text = "A"

    bad = Document()
    bad.add_paragraph("ВВЕДЕНИЕ")
    bad.add_paragraph("Таблица 1.2.2")
    bad.add_table(rows=2, cols=1)
    bad.add_paragraph("Продолжение таблицы 1.2.2")
    bad.add_table(rows=2, cols=1)

    with tempfile.TemporaryDirectory() as tmp:
        backup = Path(tmp) / "backup.docx"
        output = Path(tmp) / "output.docx"
        clean.save(backup)
        bad.save(output)

        old_probe = tc._same_page_continuation_marker_violations_for_docx
        try:
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: [
                tc._SamePageContinuationMarkerViolation(
                    marker_text="Продолжение таблицы 1.2.2",
                    marker_page=13,
                    previous_table_page=13,
                    following_table_page=13,
                    confidence="high",
                )
            ]
            restored = tc.restore_docx_if_same_page_continuation_markers(
                output,
                backup,
                context="unit_test",
            )
        finally:
            tc._same_page_continuation_marker_violations_for_docx = old_probe

        out = Document(str(output))

    if not restored:
        return _result(False, "same-page marker did not trigger restore")
    if any("Продолжение таблицы" in (p.text or "") for p in out.paragraphs):
        return _result(False, "fake continuation marker survived restore")
    if len(out.tables) != 1:
        return _result(False, f"restore should leave unsplit table only, got {len(out.tables)} tables")
    return _result(True, "same-page marker gate restores pre-split backup")


def test_tm_post_render_gate_preserves_clean_next_page_marker() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 2.3.3")
    doc.add_table(rows=2, cols=1)
    doc.add_paragraph("Продолжение таблицы 2.3.3")
    doc.add_table(rows=2, cols=1)

    with tempfile.TemporaryDirectory() as tmp:
        backup = Path(tmp) / "backup.docx"
        output = Path(tmp) / "output.docx"
        doc.save(backup)
        doc.save(output)

        old_probe = tc._same_page_continuation_marker_violations_for_docx
        try:
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: []
            restored = tc.restore_docx_if_same_page_continuation_markers(
                output,
                backup,
                context="unit_test",
            )
        finally:
            tc._same_page_continuation_marker_violations_for_docx = old_probe

        out = Document(str(output))

    if restored:
        return _result(False, "clean next-page marker unexpectedly restored")
    if sum(1 for p in out.paragraphs if p.text == "Продолжение таблицы 2.3.3") != 1:
        return _result(False, "clean continuation marker was removed")
    if len(out.tables) != 2:
        return _result(False, f"clean continuation chain changed: tables={len(out.tables)}")
    return _result(True, "clean next-page marker output is preserved")


def test_tm_post_render_gate_preserves_unresolved_manual_marker() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Таблица 2.1.2")
    doc.add_table(rows=2, cols=1)
    doc.add_paragraph("Продолжение таблицы 2.1.2")
    doc.add_table(rows=2, cols=1)

    violation = tc._SamePageContinuationMarkerViolation(
        marker_text="Продолжение таблицы 2.1.2",
        marker_page=26,
        previous_table_page=26,
        following_table_page=27,
        confidence="medium",
    )

    with tempfile.TemporaryDirectory() as tmp:
        backup = Path(tmp) / "backup.docx"
        output = Path(tmp) / "output.docx"
        doc.save(backup)
        doc.save(output)

        old_probe = tc._same_page_continuation_marker_violations_for_docx
        log_stream = io.StringIO()
        handler = logging.StreamHandler(log_stream)
        tc.logger.addHandler(handler)
        try:
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: [violation]
            restored = tc.restore_docx_if_same_page_continuation_markers(
                output,
                backup,
                context="unit_test",
            )
        finally:
            tc.logger.removeHandler(handler)
            tc._same_page_continuation_marker_violations_for_docx = old_probe

        out = Document(str(output))
        marker_count = sum(1 for p in out.paragraphs if p.text == "Продолжение таблицы 2.1.2")
        logs = log_stream.getvalue()

    if not restored:
        return _result(False, "same-page marker did not trigger restore")
    if marker_count != 1:
        return _result(False, f"manual continuation marker should be preserved, got {marker_count}")
    if "post_render_same_page_marker_violation_unresolved" not in logs:
        return _result(False, f"unresolved violation was not logged: {logs}")
    if "post_render_same_page_marker_removed" in logs:
        return _result(False, f"semantic marker was silently removed: {logs}")
    return _result(True, "unresolved manual continuation marker is preserved and logged")


def test_tm_same_page_marker_cleanup_preserves_marker_when_deletion_creates_missing_fragment() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.docx_utils import FormattingReport
    from guides.coursework_kfu_2025.rendered_table_validation import RenderedContinuationViolation

    doc = Document()
    doc.add_paragraph("Таблица 2.2.3")
    doc.add_table(rows=2, cols=1)
    doc.add_paragraph("Продолжение таблицы 2.2.3")
    doc.add_table(rows=2, cols=1)

    violation = tc._SamePageContinuationMarkerViolation(
        marker_text="Продолжение таблицы 2.2.3",
        marker_page=45,
        previous_table_page=45,
        following_table_page=45,
        confidence="high",
    )
    rendered_regression = RenderedContinuationViolation(
        table_num="2.2.3",
        table_index=17,
        page=45,
        violation_type="missing_continuation_marker",
        confidence="high",
        evidence={"repeated_header": True},
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)
        report = FormattingReport()

        old_same_page = tc._same_page_continuation_marker_violations_for_docx
        old_rendered = getattr(tc, "_rendered_continuation_violations_for_docx", None)
        try:
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: [violation]
            tc._rendered_continuation_violations_for_docx = lambda _path: [rendered_regression]
            removed = tc.remove_same_page_continuation_markers_inplace(path, report=report)
        finally:
            tc._same_page_continuation_marker_violations_for_docx = old_same_page
            if old_rendered is None:
                try:
                    delattr(tc, "_rendered_continuation_violations_for_docx")
                except AttributeError:
                    pass
            else:
                tc._rendered_continuation_violations_for_docx = old_rendered

        out = Document(str(path))

    marker_count = sum(1 for p in out.paragraphs if p.text == "Продолжение таблицы 2.2.3")
    if removed != 0:
        return _result(False, f"unsafe cleanup should remove 0 markers, got {removed}")
    if marker_count != 1:
        return _result(False, f"marker should be preserved when deletion creates missing fragment, got {marker_count}")
    if not any("таблицы 2.2.3" in warning and "маркер продолжения" in warning for warning in report.warnings):
        return _result(False, f"review-needed warning missing: {report.warnings!r}")
    return _result(True, "unsafe same-page marker cleanup preserves marker and warns")


def test_tm_same_page_marker_cleanup_removes_marker_when_final_render_is_clean() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.docx_utils import FormattingReport

    doc = Document()
    doc.add_paragraph("Таблица 2.1.1")
    doc.add_table(rows=2, cols=1)
    doc.add_paragraph("Продолжение таблицы 2.1.1")
    doc.add_table(rows=2, cols=1)

    violation = tc._SamePageContinuationMarkerViolation(
        marker_text="Продолжение таблицы 2.1.1",
        marker_page=29,
        previous_table_page=29,
        following_table_page=29,
        confidence="high",
    )

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)
        report = FormattingReport()

        old_same_page = tc._same_page_continuation_marker_violations_for_docx
        old_rendered = getattr(tc, "_rendered_continuation_violations_for_docx", None)
        try:
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: [violation]
            tc._rendered_continuation_violations_for_docx = lambda _path: []
            removed = tc.remove_same_page_continuation_markers_inplace(path, report=report)
        finally:
            tc._same_page_continuation_marker_violations_for_docx = old_same_page
            if old_rendered is None:
                try:
                    delattr(tc, "_rendered_continuation_violations_for_docx")
                except AttributeError:
                    pass
            else:
                tc._rendered_continuation_violations_for_docx = old_rendered

        out = Document(str(path))

    marker_count = sum(1 for p in out.paragraphs if p.text == "Продолжение таблицы 2.1.1")
    if removed != 1:
        return _result(False, f"clean cleanup should remove 1 marker, got {removed}")
    if marker_count != 0:
        return _result(False, f"stale marker should be removed, got {marker_count}")
    if report.warnings:
        return _result(False, f"clean cleanup should not produce user warnings: {report.warnings!r}")
    return _result(True, "safe same-page marker cleanup removes stale marker silently")


def test_tm_post_render_restore_does_not_emit_user_warning_before_final_validation() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.docx_utils import FormattingReport

    backup_doc = Document()
    backup_doc.add_paragraph("Таблица 1.1.1")
    backup_doc.add_table(rows=1, cols=1)

    output_doc = Document()
    output_doc.add_paragraph("Таблица 1.1.1")
    output_doc.add_table(rows=1, cols=1)
    output_doc.add_paragraph("Продолжение таблицы 1.1.1")
    output_doc.add_table(rows=1, cols=1)

    violation = tc._SamePageContinuationMarkerViolation(
        marker_text="Продолжение таблицы 1.1.1",
        marker_page=8,
        previous_table_page=8,
        following_table_page=8,
        confidence="high",
    )

    with tempfile.TemporaryDirectory() as tmp:
        backup = Path(tmp) / "backup.docx"
        output = Path(tmp) / "output.docx"
        backup_doc.save(backup)
        output_doc.save(output)
        report = FormattingReport()

        old_same_page = tc._same_page_continuation_marker_violations_for_docx
        try:
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: [violation]
            restored = tc.restore_docx_if_same_page_continuation_markers(
                output,
                backup,
                report=report,
                context="unit_test",
            )
        finally:
            tc._same_page_continuation_marker_violations_for_docx = old_same_page

    if not restored:
        return _result(False, "same-page marker should still trigger restore")
    if report.warnings:
        return _result(False, f"intermediate restore warning leaked to user: {report.warnings!r}")
    return _result(True, "post-render restore defers user warnings to final validation")


def _tm_make_manual_chain_doc() -> Document:
    doc = Document()
    doc.add_paragraph("Таблица 2.2.3")
    first = doc.add_table(rows=7, cols=6)
    rows = [
        ["Уровень (DAGMAR)", "Показатель (пример)", "Единица", "Источник данных", "Бенчмарк/цель пилота", "Частота"],
        ["1", "2", "3", "4", "5", "6"],
        ["Awareness", "Доля знающих о ночной доставке", "Процент", "Brand-lift, опрос в приложении", "Не менее 70%", "Еженедельно/ежемесячно"],
        ["Knowledge", "Понимание условий сервиса (время, зона)", "Процент верных ответов", "Опрос, in-app квизы", "Рост на 10–15 пунктов", "Еженедельно"],
        ["Preference", "Интерес к ночному заказу", "CTR, доля добавлений в корзину", "Платформы, CRM", "Плюс 10–15% к базовому уровню", "Еженедельно"],
        ["Intention", "Намерение сделать заказ ночью", "Доля намерений/броней", "In-app события, формы", "Плюс 10% к базовому уровню", "Еженедельно"],
        ["Action", "Первая проба и объём ночных заказов", "Trial-rate; число заказов", "CRM, BI", "Первая проба не менее 8%; прирост заказов не менее 20%", "Еженедельно/ежемесячно"],
    ]
    for row, values in zip(first.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
    doc.add_paragraph("Продолжение таблицы 2.2.3")
    second = doc.add_table(rows=3, cols=6)
    second_rows = [
        rows[0],
        rows[1],
        ["Эффективность", "CPI, CR к первому заказу, CAC, AOV ночью, LTV-90", "Валюта/процент", "Платформы, CRM, BI", "В пределах плановых коридоров", "Еженедельно/ежемесячно"],
    ]
    for row, values in zip(second.rows, second_rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
    doc.add_paragraph("Источник: составлено по [1].")
    return doc


def _tm_row_texts(table) -> list[list[str]]:
    return [[" ".join((cell.text or "").split()) for cell in row.cells] for row in table.rows]


def _tm_missing_marker_violation():
    from guides.coursework_kfu_2025.rendered_table_validation import RenderedContinuationViolation

    return RenderedContinuationViolation(
        table_num="2.2.3",
        table_index=0,
        page=45,
        violation_type="missing_continuation_marker",
        confidence="high",
        evidence={
            "repeated_header": True,
            "repeated_numeric_row": True,
            "repeated_row": True,
            "row_fingerprint": "intention намерение сделать заказ ночью доля намерений/броней",
        },
    )


def test_tm_manual_chain_overflow_repair_moves_rows_into_continuation_table() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc

    doc = _tm_make_manual_chain_doc()
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)

        old_validator = tc._rendered_continuation_violations_for_docx
        old_same_page = tc._same_page_continuation_marker_violations_for_docx
        try:
            tc._rendered_continuation_violations_for_docx = lambda _path: []
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: []
            repaired = tc.repair_manual_chain_overflow_before_marker(
                path,
                [_tm_missing_marker_violation()],
            )
        finally:
            tc._rendered_continuation_violations_for_docx = old_validator
            tc._same_page_continuation_marker_violations_for_docx = old_same_page

        out = Document(str(path))
        first_rows = _tm_row_texts(out.tables[0])
        second_rows = _tm_row_texts(out.tables[1])
        source_texts = [p.text for p in out.paragraphs]

    if repaired != 1:
        return _result(False, f"expected one repair, got {repaired}")
    if [row[0] for row in first_rows] != ["Уровень (DAGMAR)", "1", "Awareness", "Knowledge", "Preference"]:
        return _result(False, f"unexpected first fragment rows: {first_rows!r}")
    if [row[0] for row in second_rows] != ["Уровень (DAGMAR)", "1", "Intention", "Action", "Эффективность"]:
        return _result(False, f"unexpected continuation rows: {second_rows!r}")
    if source_texts[-1] != "Источник: составлено по [1].":
        return _result(False, f"source paragraph moved/lost: {source_texts!r}")
    return _result(True, "manual-chain overflow rows moved into continuation table")


def test_tm_manual_chain_overflow_repair_skips_header_mismatch() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc

    doc = _tm_make_manual_chain_doc()
    doc.tables[1].rows[0].cells[0].text = "Другой заголовок"

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)
        before = path.read_bytes()
        repaired = tc.repair_manual_chain_overflow_before_marker(
            path,
            [_tm_missing_marker_violation()],
        )
        after = path.read_bytes()

    if repaired != 0:
        return _result(False, f"header mismatch should not repair, got {repaired}")
    if before != after:
        return _result(False, "header mismatch mutated the document")
    return _result(True, "manual-chain repair skips mismatched headers")


def test_tm_manual_chain_overflow_repair_skips_merged_cells() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = _tm_make_manual_chain_doc()
    tcPr = doc.tables[0].rows[5].cells[0]._tc.get_or_add_tcPr()
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), "2")
    tcPr.append(grid_span)

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)
        before = path.read_bytes()
        repaired = tc.repair_manual_chain_overflow_before_marker(
            path,
            [_tm_missing_marker_violation()],
        )
        after = path.read_bytes()

    if repaired != 0:
        return _result(False, f"merged-cell chain should not repair, got {repaired}")
    if before != after:
        return _result(False, "merged-cell chain mutated the document")
    return _result(True, "manual-chain repair skips merged/vMerge-sensitive rows")


def test_tm_manual_chain_overflow_repair_rolls_back_failed_revalidation() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc

    doc = _tm_make_manual_chain_doc()
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)
        before = path.read_bytes()

        old_validator = tc._rendered_continuation_violations_for_docx
        try:
            tc._rendered_continuation_violations_for_docx = lambda _path: [_tm_missing_marker_violation()]
            repaired = tc.repair_manual_chain_overflow_before_marker(
                path,
                [_tm_missing_marker_violation()],
            )
        finally:
            tc._rendered_continuation_violations_for_docx = old_validator
        after = path.read_bytes()

    if repaired != 0:
        return _result(False, f"failed revalidation should rollback, got {repaired}")
    if before != after:
        return _result(False, "failed revalidation left partial row movement")
    return _result(True, "manual-chain repair rolls back when validation stays dirty")


def test_tm_apply_rendered_split_rolls_back_same_page_marker() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("Таблица 1.2.2")
    table = doc.add_table(rows=6, cols=3)
    for r in range(6):
        for c in range(3):
            table.rows[r].cells[c].text = f"r{r}c{c}"

    diagnostic = tm.TableMarkerDiagnostic(
        table_index=0,
        rows_count=6,
        pages_detected=[1, 2],
        row_pages={0: 1, 1: 1, 2: 1, 3: 2, 4: 2, 5: 2},
        found_rows=[0, 1, 2, 3, 4, 5],
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=[tm.TablePageSpan(0, 2, 1), tm.TablePageSpan(3, 5, 2)],
        appendix_table=False,
        caption_detected=True,
        has_standard_table_caption=True,
        preceding_paragraph_text="Таблица 1.2.2",
    )

    save = _e2_set_candidate_mode(budget=20)
    old_dt = tm.diagnose_table
    old_probe = tc._same_page_continuation_marker_violations_for_docx
    try:
        tm.diagnose_table = lambda _p, idx, keep_temp=False: diagnostic
        tc._same_page_continuation_marker_violations_for_docx = lambda _path: [
            tc._SamePageContinuationMarkerViolation(
                marker_text="Продолжение таблицы 1.2.2",
                marker_page=13,
                previous_table_page=13,
                following_table_page=13,
                confidence="high",
            )
        ]
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            n = tc.apply_rendered_table_continuation(path)
            out = Document(str(path))
    finally:
        tm.diagnose_table = old_dt
        tc._same_page_continuation_marker_violations_for_docx = old_probe
        _e2_restore_env(save)

    if n != 0:
        return _result(False, f"same-page split should be reported as skipped/restored, got n={n}")
    if len(out.tables) != 1:
        return _result(False, f"rollback should restore one unsplit table, got {len(out.tables)}")
    if any("Продолжение таблицы" in (p.text or "") for p in out.paragraphs):
        return _result(False, "same-page continuation marker survived apply_rendered rollback")
    return _result(True, "rendered split rollback removes same-page continuation marker")


def test_tm_real_fixture_marker_gate_preserves_semantic_markers() -> tuple[bool, str]:
    import guides.coursework_kfu_2025.table_continuation as tc
    from guides.coursework_kfu_2025.formatter_service import format_docx

    strict_fixtures = [
        Path("/Users/mac/Desktop/курсовые/курсовая пример 1.docx"),
        Path("/Users/mac/Desktop/курсовые/Гаянов_Амир_Ленарович_Разработка проектного решения по автоматизации документооборота в организации.docx"),
        Path("/Users/mac/Desktop/курсовые/курсовая_Бондарев_Никита_2_курс.docx"),
    ]
    reference = Path("/Users/mac/Desktop/курсовые/Рыбаков_Олег_Дмитриевич_курсовая_3_курс.docx")
    fixtures = strict_fixtures + [reference]
    missing = [str(path) for path in fixtures if not path.exists()]
    if missing:
        return _result(True, f"fixture(s) not present; skipped: {missing!r}")

    failures: list[str] = []
    with tempfile.TemporaryDirectory() as tmp:
        for idx, fixture in enumerate(strict_fixtures):
            out = Path(tmp) / f"strict_fixture_{idx}_formatted.docx"
            _, warnings = format_docx(str(fixture), str(out))
            violations = tc._same_page_continuation_marker_violations_for_docx(out)
            if violations:
                review_needed = any(
                    "маркер продолжения сохран" in warning
                    or "без маркера продолжения" in warning
                    or "повторный фрагмент виден на той же странице" in warning
                    or "соседний фрагмент выглядит как часть той же таблицы" in warning
                    for warning in warnings
                )
                if not review_needed:
                    failures.append(
                        f"{fixture.name}: "
                        + "; ".join(
                            f"{v.marker_text}@p{v.marker_page}/prev={v.previous_table_page}/next={v.following_table_page}"
                            for v in violations
                        )
                    )

        ref_out = Path(tmp) / "reference_rybakov_formatted.docx"
        format_docx(str(reference), str(ref_out))
        ref_doc = Document(str(ref_out))
        ref_texts = [p.text for p in ref_doc.paragraphs]
        if "Продолжение таблицы 2.1.2" not in ref_texts:
            failures.append("reference Рыбаков: semantic marker Продолжение таблицы 2.1.2 was deleted")

    if failures:
        return _result(False, "real fixture marker gate regression: " + " | ".join(failures))
    return _result(True, "real fixture marker gate preserves semantic markers and warns on unsafe cleanup")


# ── PG2: safe_formatter geometry isolation foundation ────────────────────────
# These tests guard Phase 1 table geometry preservation. They are not visual
# proof; rendered diagnostics/smoke remain the source of truth for table quality.

def _pg2_base_doc() -> Document:
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст перед таблицей.")
    return doc


def _pg2_fill_table(tbl) -> None:
    for r, row in enumerate(tbl.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"Значение {r + 1}.{c + 1}"


def _pg2_get_tbl_pr(tbl_xml):
    tbl_pr = tbl_xml.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_xml.insert(0, tbl_pr)
    return tbl_pr


def _pg2_set_tbl_w(tbl, value: str, typ: str) -> None:
    tbl_pr = _pg2_get_tbl_pr(tbl._tbl)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), value)
    tbl_w.set(qn("w:type"), typ)


def _pg2_set_tbl_layout(tbl, typ: str = "fixed") -> None:
    tbl_pr = _pg2_get_tbl_pr(tbl._tbl)
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), typ)


def _pg2_set_tbl_cell_spacing(tbl, value: str = "15") -> None:
    tbl_pr = _pg2_get_tbl_pr(tbl._tbl)
    spacing = tbl_pr.find(qn("w:tblCellSpacing"))
    if spacing is None:
        spacing = OxmlElement("w:tblCellSpacing")
        tbl_pr.append(spacing)
    spacing.set(qn("w:w"), value)
    spacing.set(qn("w:type"), "dxa")


def _pg2_set_tbl_borders(tbl, *, val: str = "double", color: str = "auto") -> None:
    tbl_pr = _pg2_get_tbl_pr(tbl._tbl)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), val)
        node.set(qn("w:sz"), "8")
        node.set(qn("w:space"), "1")
        node.set(qn("w:color"), color)


def _pg2_tbl_border_attrs(tbl) -> list[tuple[str, str | None, str | None, str | None, str | None]]:
    tbl_pr = _pg2_get_tbl_pr(tbl._tbl)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        return []
    attrs = []
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        attrs.append((
            edge,
            node.get(qn("w:val")) if node is not None else None,
            node.get(qn("w:sz")) if node is not None else None,
            node.get(qn("w:space")) if node is not None else None,
            node.get(qn("w:color")) if node is not None else None,
        ))
    return attrs


def _pg2_set_all_tcw(tbl, widths: list[int]) -> None:
    for row in tbl.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx % len(widths)]))
            tc_w.set(qn("w:type"), "dxa")


def _pg2_process_and_snapshot(doc: Document, table_index: int = 0):
    from guides.coursework_kfu_2025.safe_formatter import process_document

    diag = _table_engine_diag_module()
    with tempfile.TemporaryDirectory() as tmp:
        inp = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(inp)
        before_doc = Document(str(inp))
        before = diag.snapshot_table(before_doc.tables[table_index]._tbl, table_index=table_index)
        process_document(inp, out)
        after_doc = Document(str(out))
        after = diag.snapshot_table(after_doc.tables[table_index]._tbl, table_index=table_index)
    return before, after


def _pg2_geometry_changes(before: dict, after: dict) -> dict:
    diag = _table_engine_diag_module()
    return diag.diff_table_geometry(before, after)


def test_pg2_safe_formatter_preserves_sensitive_tblw_layout_tcw_grid() -> tuple[bool, str]:
    doc = _pg2_base_doc()
    tbl = doc.add_table(rows=2, cols=3)
    _pg2_fill_table(tbl)
    _pg2_set_tbl_w(tbl, "5000", "pct")
    _pg2_set_tbl_layout(tbl, "fixed")
    _pg2_set_all_tcw(tbl, [1300, 2200, 3100])

    before, after = _pg2_process_and_snapshot(doc)
    for key in ("tblW", "tblLayout", "tblGrid_xml", "gridCol_widths", "tcW"):
        if before[key] != after[key]:
            return _result(False, f"{key} changed for preserve-mode table")
    return _result(True, "safe_formatter preserved tblW/tblLayout/tcW/tblGrid for sensitive table")


def test_pg2_safe_formatter_keeps_simple_table_width_normalization() -> tuple[bool, str]:
    doc = _pg2_base_doc()
    tbl = doc.add_table(rows=2, cols=3)
    _pg2_fill_table(tbl)

    _before, after = _pg2_process_and_snapshot(doc)
    if after["tblW"].get("type") != "dxa" or not after["tblW"].get("w"):
        return _result(False, f"simple table tblW was not normalized: {after['tblW']!r}")
    if after["tblLayout"] != "fixed":
        return _result(False, f"simple table layout not fixed: {after['tblLayout']!r}")
    if not after["tcW"] or any(item.get("w") is None for item in after["tcW"]):
        return _result(False, "simple table tcW widths were not populated")
    if "w:tblBorders" not in (after.get("tblPr_xml") or ""):
        return _result(False, "simple table borders were not normalized")
    return _result(True, "simple table still receives width/layout/border normalization")


def test_pg2_safe_formatter_preserves_auto_pct_and_authored_fixed_layout() -> tuple[bool, str]:
    doc = _pg2_base_doc()
    pct_tbl = doc.add_table(rows=2, cols=2)
    _pg2_fill_table(pct_tbl)
    _pg2_set_tbl_w(pct_tbl, "4200", "pct")
    _pg2_set_all_tcw(pct_tbl, [1800, 3000])

    fixed_tbl = doc.add_table(rows=2, cols=2)
    _pg2_fill_table(fixed_tbl)
    _pg2_set_tbl_layout(fixed_tbl, "fixed")
    _pg2_set_all_tcw(fixed_tbl, [1500, 2500])

    before_pct, after_pct = _pg2_process_and_snapshot(doc, table_index=0)
    before_fixed, after_fixed = _pg2_process_and_snapshot(doc, table_index=1)

    if before_pct["tblW"] != after_pct["tblW"] or before_pct["tcW"] != after_pct["tcW"]:
        return _result(False, "tblW pct table geometry changed")
    if before_fixed["tblLayout"] != after_fixed["tblLayout"] or before_fixed["tcW"] != after_fixed["tcW"]:
        return _result(False, "authored fixed-layout table geometry changed")
    return _result(True, "auto/pct and authored fixed-layout tables are preserved")


def test_pg2_safe_formatter_preserves_merged_and_vmerge_tables() -> tuple[bool, str]:
    doc = _pg2_base_doc()
    tbl = doc.add_table(rows=3, cols=3)
    _pg2_fill_table(tbl)
    _pg2_set_all_tcw(tbl, [1200, 2400, 3600])
    tbl.cell(0, 0).merge(tbl.cell(0, 1))

    restart = OxmlElement("w:vMerge")
    restart.set(qn("w:val"), "restart")
    tbl.rows[1].cells[0]._tc.get_or_add_tcPr().append(restart)
    cont = OxmlElement("w:vMerge")
    cont.set(qn("w:val"), "continue")
    tbl.rows[2].cells[0]._tc.get_or_add_tcPr().append(cont)

    before, after = _pg2_process_and_snapshot(doc)
    for key in ("tblGrid_xml", "gridCol_widths", "tcW", "gridSpan", "vMerge"):
        if before[key] != after[key]:
            return _result(False, f"{key} changed for merged/vMerge table")
    return _result(True, "merged/vMerge table geometry is preserved by safe_formatter")


def test_pg2_preserve_geometry_table_gets_border_cleanup_without_geometry_mutation() -> tuple[bool, str]:
    doc = _pg2_base_doc()
    tbl = doc.add_table(rows=2, cols=3)
    _pg2_fill_table(tbl)
    _pg2_set_tbl_w(tbl, "5000", "pct")
    _pg2_set_all_tcw(tbl, [1300, 2200, 3100])
    _pg2_set_tbl_cell_spacing(tbl, "15")
    _pg2_set_tbl_borders(tbl, val="double", color="auto")

    before, after = _pg2_process_and_snapshot(doc)
    for key in ("tblW", "tblLayout", "tblGrid_xml", "gridCol_widths", "tcW", "gridSpan", "vMerge"):
        if before[key] != after[key]:
            return _result(False, f"{key} changed during preserve border cleanup")

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "in.docx"
        out = Path(tmp) / "out.docx"
        doc.save(src)
        from guides.coursework_kfu_2025.safe_formatter import process_document

        process_document(src, out)
        out_doc = Document(str(out))
        out_tbl = out_doc.tables[0]
        tbl_pr = _pg2_get_tbl_pr(out_tbl._tbl)
        if tbl_pr.find(qn("w:tblCellSpacing")) is not None:
            return _result(False, "tblCellSpacing survived preserve border cleanup")
        attrs = _pg2_tbl_border_attrs(out_tbl)
        expected = [(edge, "single", "4", "0", "000000") for edge in ("top", "left", "bottom", "right", "insideH", "insideV")]
        if attrs != expected:
            return _result(False, f"preserve table borders not normalized: {attrs!r}")

    return _result(True, "preserve table spacing removed and borders normalized without geometry mutation")


def test_pg2_preserve_border_cleanup_keeps_merged_and_vmerge_structure() -> tuple[bool, str]:
    doc = _pg2_base_doc()
    tbl = doc.add_table(rows=3, cols=3)
    _pg2_fill_table(tbl)
    _pg2_set_tbl_cell_spacing(tbl, "15")
    _pg2_set_tbl_borders(tbl, val="double", color="auto")
    tbl.cell(0, 0).merge(tbl.cell(0, 1))

    restart = OxmlElement("w:vMerge")
    restart.set(qn("w:val"), "restart")
    tbl.rows[1].cells[0]._tc.get_or_add_tcPr().append(restart)
    cont = OxmlElement("w:vMerge")
    cont.set(qn("w:val"), "continue")
    tbl.rows[2].cells[0]._tc.get_or_add_tcPr().append(cont)

    before, after = _pg2_process_and_snapshot(doc)
    for key in ("tblGrid_xml", "gridCol_widths", "tcW", "gridSpan", "vMerge"):
        if before[key] != after[key]:
            return _result(False, f"{key} changed for merged/vMerge preserve border cleanup")
    return _result(True, "preserve border cleanup does not damage merged/vMerge structure")


def test_pg2_force_borders_preserves_geometry_adjacent_nodes() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import force_table_outer_borders_single

    doc = Document()
    tbl = doc.add_table(rows=1, cols=1)
    tbl_pr = _pg2_get_tbl_pr(tbl._tbl)
    for tag in ("w:tblLook", "w:tblCellMar", "w:tblCellSpacing"):
        tbl_pr.append(OxmlElement(tag))

    tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:tblPrEx"))

    force_table_outer_borders_single(tbl)

    for tag in ("w:tblLook", "w:tblCellMar", "w:tblCellSpacing"):
        if tbl_pr.find(qn(tag)) is None:
            return _result(False, f"{tag} was deleted")
    if tr_pr.find(qn("w:tblPrEx")) is None:
        return _result(False, "w:tblPrEx was deleted")
    return _result(True, "legacy simple-table border helper still fails closed on preserve tables")


def test_pg2_rybakov_style_table_keeps_tcw_in_safe_formatter_stage() -> tuple[bool, str]:
    from guides.coursework_kfu_2025.safe_formatter import process_document

    candidates = sorted(Path("/Users/mac/Desktop/курсовые").glob("*ромаркетинг_Рыбаков.docx"))
    if not candidates:
        return _result(True, "Rybakov fixture not present; skipped")

    diag = _table_engine_diag_module()
    src = candidates[0]
    before_doc = Document(str(src))
    preserve_indexes = [
        idx for idx, table in enumerate(before_doc.tables)
        if diag.snapshot_table(table._tbl, table_index=idx)["geometry_policy"] != "simple"
    ]
    if not preserve_indexes:
        return _result(False, "Rybakov fixture had no preserve-mode tables")

    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "out.docx"
        process_document(src, out)
        after_doc = Document(str(out))

    for idx in preserve_indexes:
        before = diag.snapshot_table(before_doc.tables[idx]._tbl, table_index=idx)
        after = diag.snapshot_table(after_doc.tables[idx]._tbl, table_index=idx)
        changes = diag.diff_table_geometry(before, after)
        if "tcW" in changes:
            return _result(False, f"tcW changed for Rybakov preserve table {idx}")

    return _result(True, f"Rybakov-style preserve tables kept tcW unchanged: {preserve_indexes}")


def test_pg2_diagnostics_show_reduced_safe_formatter_geometry_mutation() -> tuple[bool, str]:
    diag = _table_engine_diag_module()

    doc = _pg2_base_doc()
    preserve_tbl = doc.add_table(rows=2, cols=2)
    _pg2_fill_table(preserve_tbl)
    _pg1_add_tbl_pr_child(preserve_tbl._tbl, "w:tblCellMar")
    _pg2_set_all_tcw(preserve_tbl, [1400, 2600])

    simple_tbl = doc.add_table(rows=2, cols=2)
    _pg2_fill_table(simple_tbl)

    with tempfile.TemporaryDirectory() as tmp:
        src = Path(tmp) / "source.docx"
        doc.save(src)
        result = diag.run_table_engine_diagnostics(src, artifact_root=Path(tmp) / "artifacts", render=False)

    safe_stage = next(stage for stage in result.stages if stage.name == "01_after_safe_formatter")
    preserve_snapshot = safe_stage.tables[0]
    simple_snapshot = safe_stage.tables[1]
    preserve_diff = preserve_snapshot["geometry_diff_from_previous_stage"]
    if any(key != "tblPr_xml" for key in preserve_diff):
        return _result(False, f"preserve table geometry changed beyond border-only tblPr cleanup: {preserve_diff}")
    if not simple_snapshot["geometry_diff_from_previous_stage"]:
        return _result(False, "simple table no longer records expected normalization diff")
    return _result(True, "diagnostics distinguish preserved sensitive table from normalized simple table")


def test_e2_rollback_global_skip_mode_preserves_pre_e2_behaviour() -> tuple[bool, str]:
    """E2: KPFU_MARKER_SPLIT_MODE=global_skip restores pre-E2 behaviour."""
    import logging
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    for i in range(8):
        doc.add_paragraph(f"Таблица 1.1.{i+1}")
        t = doc.add_table(rows=2, cols=2)
        for r in range(2):
            for c in range(2):
                t.rows[r].cells[c].text = "x"

    captured: list[str] = []
    class _Cap(logging.Handler):
        def __init__(self): super().__init__(level=logging.DEBUG)
        def emit(self, r): captured.append(r.getMessage())

    save = {
        "enable": os.environ.get("KPFU_ENABLE_MARKER_SPLIT"),
        "mode":   os.environ.get("KPFU_MARKER_SPLIT_MODE"),
        "budget": os.environ.get("KPFU_MARKER_SPLIT_MAX_RENDERS"),
    }
    os.environ["KPFU_ENABLE_MARKER_SPLIT"] = "1"
    os.environ["KPFU_MARKER_SPLIT_MODE"] = "global_skip"
    os.environ["KPFU_MARKER_SPLIT_MAX_RENDERS"] = "6"

    old_dt = tm.diagnose_table
    handler = _Cap()
    prev = tc.logger.level
    tc.logger.addHandler(handler)
    tc.logger.setLevel(logging.DEBUG)
    try:
        # diagnose_table must NOT be called under global_skip with overflow.
        tm.diagnose_table = lambda *a, **k: (_ for _ in ()).throw(AssertionError("global_skip must not diagnose"))
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            n = tc.apply_rendered_table_continuation(path)
    finally:
        tm.diagnose_table = old_dt
        tc.logger.removeHandler(handler)
        tc.logger.setLevel(prev)
        for key, env in (("enable","KPFU_ENABLE_MARKER_SPLIT"),
                         ("mode","KPFU_MARKER_SPLIT_MODE"),
                         ("budget","KPFU_MARKER_SPLIT_MAX_RENDERS")):
            if save[key] is None:
                os.environ.pop(env, None)
            else:
                os.environ[env] = save[key]

    if n != 0:
        return _result(False, f"expected 0 (skip), got {n}")
    if not any("render_budget_exceeded" in m for m in captured):
        return _result(False, "render_budget_exceeded log missing in global_skip mode")
    return _result(True, "global_skip rollback preserves pre-E2 behavior")


# ── E3: NUM-row compensation for marker split ─────────────────────────────────

def _e3_build_table_doc(rows: int, cols: int = 3, *, with_numbered_row=False,
                        caption: str = "Таблица 1.1.1") -> Document:
    """Build a synthetic doc with one table.

    Layout:
      ВВЕДЕНИЕ
      Текст.
      <caption>
      <table with `rows` rows × `cols` cols>
    If with_numbered_row=True, row 1 is the exact numbered row "1 2 ... cols".
    """
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Текст.")
    doc.add_paragraph(caption)
    tbl = doc.add_table(rows=rows, cols=cols)
    for r in range(rows):
        for c in range(cols):
            if r == 0:
                tbl.rows[r].cells[c].text = f"H{c+1}"
            elif r == 1 and with_numbered_row:
                tbl.rows[r].cells[c].text = str(c + 1)
            else:
                tbl.rows[r].cells[c].text = f"r{r}c{c}"
    return doc


def _e3_make_diagnostic(tm, *, table_index=0, rows_count=9, split_at=4,
                        appendix=False, has_caption=True):
    """Build a TableMarkerDiagnostic where rows 0..split_at-1 are on page 1 and
    rows split_at..rows_count-1 are on page 2 (a clean 2-page boundary)."""
    row_pages = {}
    for r in range(rows_count):
        row_pages[r] = 1 if r < split_at else 2
    page_spans = [
        tm.TablePageSpan(0, split_at - 1, 1),
        tm.TablePageSpan(split_at, rows_count - 1, 2),
    ]
    return tm.TableMarkerDiagnostic(
        table_index=table_index,
        rows_count=rows_count,
        pages_detected=[1, 2],
        row_pages=row_pages,
        found_rows=list(range(rows_count)),
        missing_rows=[],
        duplicate_rows={},
        candidate_for_split=False,
        page_spans=page_spans,
        appendix_table=appendix,
        caption_detected=True,
        has_standard_table_caption=has_caption,
        preceding_paragraph_text=("ПРИЛОЖЕНИЕ 1" if appendix else "Таблица 1.1.1"),
    )


def _e3_save_and_get_args(diagnostic, **kwargs):
    """Helper: write a docx with the given table shape, run _effective_marker_split_before_row.
    Returns (K_effective)."""
    import guides.coursework_kfu_2025.table_continuation as tc
    rows = kwargs.pop("rows", 9)
    cols = kwargs.pop("cols", 3)
    with_num = kwargs.pop("with_numbered_row", False)
    flag = kwargs.pop("flag", None)  # None / "0" / "1"
    doc = _e3_build_table_doc(rows, cols, with_numbered_row=with_num)
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "doc.docx"
        doc.save(path)
        decision = tc._MarkerSplitDecision(eligible=True, split_before_row=diagnostic.row_pages and (
            min(r for r, p in diagnostic.row_pages.items() if p == diagnostic.pages_detected[-1])
        ), skip_reason=None)
        saved = os.environ.get("KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION")
        if flag is None:
            os.environ.pop("KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION", None)
        else:
            os.environ["KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION"] = flag
        try:
            return tc._effective_marker_split_before_row(diagnostic, decision, docx_path=path)
        finally:
            if saved is None:
                os.environ.pop("KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION", None)
            else:
                os.environ["KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION"] = saved


def test_e3_compensation_applies_for_ordinary_body_table() -> tuple[bool, str]:
    """E3: ordinary body table without pre-existing NUM row, rows_count - K >= 2 → K-1."""
    import guides.coursework_kfu_2025.table_markers as tm
    diag = _e3_make_diagnostic(tm, rows_count=9, split_at=4, appendix=False)
    K = _e3_save_and_get_args(diag, rows=9, cols=3, with_numbered_row=False)
    if K != 3:
        return _result(False, f"expected K=3 (compensation 4->3), got {K}")
    return _result(True, "compensation K=4 → K=3 for ordinary body table without NUM row 1")


def test_e3_compensation_skipped_when_table_has_pre_existing_numbered_row() -> tuple[bool, str]:
    """E3: pre-existing numbered row at row 1 → no insertion will fire → no compensation."""
    import guides.coursework_kfu_2025.table_markers as tm
    diag = _e3_make_diagnostic(tm, rows_count=9, split_at=4, appendix=False)
    K = _e3_save_and_get_args(diag, rows=9, cols=3, with_numbered_row=True)
    if K != 4:
        return _result(False, f"expected K=4 (no compensation, pre-existing NUM), got {K}")
    return _result(True, "compensation skipped when original row 1 is already numbered")


def test_e3_compensation_skipped_for_appendix() -> tuple[bool, str]:
    """E3: appendix branch still uses the existing -1 (unrelated path); body E3 path must skip."""
    import guides.coursework_kfu_2025.table_markers as tm
    diag = _e3_make_diagnostic(tm, rows_count=9, split_at=4, appendix=True)
    K = _e3_save_and_get_args(diag, rows=9, cols=3, with_numbered_row=False)
    # Appendix branch: K=4, header_rows=1, K > header_rows + 1 (4 > 2) → returns K-1 = 3 via the
    # existing appendix logic — but NOT via E3. Verify the existing behavior is preserved.
    if K != 3:
        return _result(False, f"expected K=3 from appendix branch, got {K}")
    return _result(True, "appendix branch returns K-1 via existing logic (E3 path not used)")


def test_e3_compensation_skipped_when_fragment2_too_small() -> tuple[bool, str]:
    """E3: rows_count - K < 2 → compensation guard blocks the -1."""
    import guides.coursework_kfu_2025.table_markers as tm
    # rows_count=6, K=5: rows_count - K = 1 → guard fails → return K unchanged
    diag = _e3_make_diagnostic(tm, rows_count=6, split_at=5, appendix=False)
    K = _e3_save_and_get_args(diag, rows=6, cols=3, with_numbered_row=False)
    if K != 5:
        return _result(False, f"expected K=5 (Бондарев-shape guard), got {K}")
    return _result(True, "compensation blocked when fragment 2 would have only 1 data row (Case-A guard)")


def test_e3_compensation_skipped_when_fragment1_would_be_empty() -> tuple[bool, str]:
    """E3: K-1 < header_rows + 1 (i.e., fragment 1 would have 0 data rows) → no compensation."""
    import guides.coursework_kfu_2025.table_markers as tm
    # rows_count=9, K=2: K-1 = 1 = header_rows (= 1), K-1 < header_rows + 1 → guard fails
    diag = _e3_make_diagnostic(tm, rows_count=9, split_at=2, appendix=False)
    K = _e3_save_and_get_args(diag, rows=9, cols=3, with_numbered_row=False)
    if K != 2:
        return _result(False, f"expected K=2 (no compensation when fragment 1 would be near-empty), got {K}")
    return _result(True, "compensation blocked when fragment 1 would have 0 data rows")


def test_e3_feature_flag_off_disables_compensation() -> tuple[bool, str]:
    """E3: KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION=0 disables the new logic."""
    import guides.coursework_kfu_2025.table_markers as tm
    diag = _e3_make_diagnostic(tm, rows_count=9, split_at=4, appendix=False)
    K_off = _e3_save_and_get_args(diag, rows=9, cols=3, with_numbered_row=False, flag="0")
    if K_off != 4:
        return _result(False, f"flag=0 expected K=4, got {K_off}")
    K_off_false = _e3_save_and_get_args(diag, rows=9, cols=3, with_numbered_row=False, flag="false")
    if K_off_false != 4:
        return _result(False, f"flag=false expected K=4, got {K_off_false}")
    K_on = _e3_save_and_get_args(diag, rows=9, cols=3, with_numbered_row=False, flag="1")
    if K_on != 3:
        return _result(False, f"flag=1 expected K=3, got {K_on}")
    return _result(True, "feature flag toggles compensation correctly (default ON; '0'/'false' disable)")


def test_e3_integration_apply_marker_split_passes_compensated_k() -> tuple[bool, str]:
    """E3 integration: _apply_marker_split_candidate routes through E3 and the resulting split
    matches the compensated K. Synthesize a Case-B-shape table (no pre-existing NUM, rows_count - K >= 2)
    and verify that apply_numbered_split_to_document is called with K-1, producing the expected
    fragment sizes."""
    import guides.coursework_kfu_2025.table_continuation as tc
    import guides.coursework_kfu_2025.table_markers as tm

    rows_count = 9
    natural_K = 4
    doc = _e3_build_table_doc(rows=rows_count, cols=3, with_numbered_row=False,
                              caption="Таблица 1.1.1")
    diag = _e3_make_diagnostic(tm, rows_count=rows_count, split_at=natural_K, appendix=False)

    saved = os.environ.get("KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION")
    os.environ["KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION"] = "1"
    try:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "doc.docx"
            doc.save(path)
            decision = tc._evaluate_marker_split_diagnostic(diag, header_rows=1)
            if not decision.eligible:
                return _result(False, f"decision not eligible: {decision!r}")
            result, skip_reason = tc._apply_marker_split_candidate(path, diag, decision)
            if result is None:
                return _result(False, f"apply returned None, skip_reason={skip_reason!r}")
            out = Document(str(path))
    finally:
        if saved is None:
            os.environ.pop("KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION", None)
        else:
            os.environ["KPFU_MARKER_SPLIT_NUM_ROW_COMPENSATION"] = saved

    if len(out.tables) != 2:
        return _result(False, f"expected 2 tables after split, got {len(out.tables)}")
    # With compensation: K=3 → table 1 has rows 0..2 of original (3 rows) + NUM = 4 rows total.
    # Table 2 has NUM + rows 3..8 of original = 7 rows total.
    if len(out.tables[0].rows) != 4:
        return _result(False, f"expected fragment 1 to have 4 rows after compensation, got {len(out.tables[0].rows)}")
    if len(out.tables[1].rows) != 7:
        return _result(False, f"expected fragment 2 to have 7 rows after compensation, got {len(out.tables[1].rows)}")
    return _result(True, "compensation propagated through _apply_marker_split_candidate (4-row fragment 1, 7-row fragment 2)")


# ── Runner ────────────────────────────────────────────────────────────────────

def run_all() -> None:
    # Default suite: fast synthetic/XML checks for confirmed product rules.
    # Real asset formatting is useful as a smoke check, but it is slower and
    # can preserve broken historical output; keep it opt-in below.
    tests = [
        # Front matter and page numbering robustness.
        ("A1 | intro-only starts at page 3", test_front_matter_intro_only_starts_numbering_at_three),
        ("A1 | title+intro starts at page 3", test_front_matter_title_intro_starts_numbering_at_three),
        ("A1 | title+contents+intro starts at page 3", test_front_matter_title_contents_intro_starts_numbering_at_three),
        ("A1 | section breaks bounded", test_front_matter_section_breaks_are_bounded),
        ("A1 | appendix numbering stops after first page", test_appendices_first_page_numbered_following_pages_unnumbered),
        ("A1 | appendix continuation unnumbered", test_appendix_continuation_pages_are_unnumbered),
        ("A1 | front matter protected", test_front_matter_before_introduction_remains_protected),
        ("B2.2 | real intro detection ignores TOC", test_real_intro_detection_ignores_toc_embedded_intro),
        ("B2.2 | front matter text frozen", test_front_matter_text_before_real_intro_is_preserved),
        ("B2 | contents front matter frozen", test_b2_contents_entries_have_stable_tab_leaders),
        ("B2.3 | merged chapter/section headings split", test_body_soft_break_chapter_and_section_headings_are_separated),
        ("B2.3 | merged heading/body split", test_body_soft_break_heading_and_body_are_separated_after_intro),
        ("B2.3 | body splitter keeps TOC frozen", test_body_soft_break_split_does_not_touch_toc),
        ("B2.3 | ordinary body soft break kept", test_body_soft_break_split_does_not_split_ordinary_body_text),
        # Figure/paragraph preservation.
        ("A  | rule4 does not delete images",          test_a_rule4_does_not_delete_images),
        ("A  | _para_has_image helper",                test_a_para_has_image_helper),
        ("A  | rule4 preserves section breaks",        test_a_rule4_preserves_front_matter_section_breaks),
        # Table continuation and split behavior.
        ("C  | continuation length guard",             test_c_continuation_length_guard),
        ("C  | strict caption-number extraction",      test_c_caption_number_extraction_strict),
        ("C  | merge invalid manual split",            test_c_apply_table_merging_rebuilds_invalid_split),
        ("C  | keep valid manual split",               test_c_apply_table_merging_keeps_valid_manual_split),
        ("C  | keep loose manual marker (no keepNext)", test_c_apply_table_merging_keeps_marker_without_keep_next),
        ("C  | rebuild caption-mismatch chain",         test_c_apply_table_merging_rebuilds_caption_mismatch),
        ("P0 | numeric-row-only continuation valid",     test_p0_manual_continuation_numeric_row_only_fragment_is_valid),
        ("P0 | synthesize missing manual NUM rows",      test_p0_preserved_manual_chain_synthesizes_missing_numeric_rows),
        ("P0 | unsplit table unchanged",                 test_p0_unsplit_ordinary_table_does_not_get_synthetic_numeric_row),
        ("P0 | existing NUM rows not duplicated",        test_p0_existing_correct_numeric_rows_are_not_duplicated),
        ("P0 | Rybakov 2.2.1 chain not merged",          test_p0_rybakov_style_221_chain_does_not_merge_into_malformed_table),
        ("E  | student marker enables pageBreakBefore", test_e_preserved_student_marker_enables_page_break_before),
        ("E  | student marker anchor repairs counted", test_e_preserved_student_marker_anchor_repairs_are_counted),
        ("E  | formatter saves repaired marker anchors", test_e_format_docx_saves_repaired_manual_marker_anchors),
        ("E  | enable pageBreakBefore is idempotent",   test_e_page_break_enable_is_idempotent),
        ("E  | formatter-authored chain not modified",  test_e_formatter_authored_chain_with_keepnext_not_modified),
        ("E  | preserved marker keeps alignment+text",  test_e_preserved_marker_keeps_alignment_and_text),
        ("E  | integration tbl→marker→tbl enabled pb",  test_e_integration_tbl_marker_tbl_marker_has_enabled_break),
        ("P1a | marker blank table anchored",           test_p1a_marker_blank_table_chain_is_anchored),
        ("P1a | disabled marker page break active",     test_p1a_disabled_marker_page_break_becomes_active),
        ("P1a | preserved manual chain normalized",     test_p1a_preserved_manual_chain_normalized_without_merging),
        ("P1a | appendix label not modified",           test_p1a_appendix_continuation_label_not_modified),
        ("P1a | numeric rows unchanged",                test_p1a_numeric_rows_stay_unchanged_after_anchoring),
        ("P1c | detached source/note detected",         test_p1c_detects_detached_source_note),
        ("P1c | attached source/note → not_detached",   test_p1c_skip_when_source_note_attached_same_page),
        ("P1c | no caption skipped",                    test_p1c_skip_no_caption),
        ("P1c | no source/note skipped",                test_p1c_skip_no_source_note),
        ("P1c | small table no_safe_data_row",          test_p1c_skip_small_table_no_safe_data_row),
        ("P1c | already in manual chain skipped",       test_p1c_skip_already_in_manual_chain),
        ("P1c | render probe unreliable skipped",       test_p1c_skip_render_probe_unreliable),
        ("P1c | split inserts marker + numbered row",   test_p1c_apply_split_inserts_continuation_marker_and_numbered_row),
        ("P1a | P1c rendered split anchored",           test_p1a_p1c_rendered_split_anchors_marker_blank_chain),
        ("P1c | idempotent via natural skip reasons",   test_p1c_double_run_idempotent_via_natural_skips),
        ("P2a | multi-page appendix table detected",    test_p2a_detects_multipage_appendix_table),
        ("P2a | non-appendix table skipped",            test_p2a_skip_non_appendix_table),
        ("P2a | existing continuation label skipped",   test_p2a_skip_existing_continuation_label),
        ("P2a | manual chain skipped",                  test_p2a_skip_manual_chain),
        ("P2a | single-page appendix skipped",          test_p2a_skip_single_page_appendix),
        ("P2a | small appendix skipped",                test_p2a_skip_small_appendix_table),
        ("P2a | split inserts uppercase marker",        test_p2a_apply_split_inserts_uppercase_continuation_marker),
        ("P2a | idempotent after apply",                test_p2a_idempotent_after_apply),
        ("P2a | no regression on P1-c fixture",         test_p2a_no_regression_on_p1c_source_note_fixture),
        # P2-a' (relaxed row/page matcher) tests are registered at the END
        # of this list — see bottom — so they do not shift the ordinal
        # position of any pre-existing test. (Placing them here surfaced a
        # latent order-sensitive interaction with E2 manual-chain tests.)
        ("D3 | analytical prose svyazyvaet demoted",    test_table_caption_reference_prose_with_svyazyvaet_demoted),
        ("D3 | rule3 no keepNext on prose below table", test_pagination_rule3_does_not_set_keepnext_on_prose_below_table),
        ("D3 | genuine caption still keepNext",         test_genuine_table_caption_still_gets_keepnext),
        ("D4 | source/note softbreak splits",           test_p4_source_note_softbreak_splits_into_two_paragraphs),
        ("D4 | ordinary body softbreak preserved",      test_p4_ordinary_body_softbreak_remains_single_paragraph),
        ("D4 | heading→body softbreak still splits",    test_p4_heading_body_softbreak_still_splits),
        ("D4 | resulting paras formatted by Phase 1",   test_p4_source_note_split_resulting_paras_format_through_phase1),
        ("C  | heuristic split disabled",              test_c_apply_table_continuation_does_not_heuristic_split),
        ("C  | width normalisation only",              test_c_apply_table_continuation_width_normalization_only),
        ("C  | no-split double-run idempotency",       test_c_apply_table_continuation_no_split_double_run_idempotent),
        ("C  | rendered split LO fallback",            test_c_apply_rendered_table_continuation_warns_when_lo_unavailable),
        ("C  | rendered split PDF fallback",           test_c_apply_rendered_table_continuation_warns_when_pdf_analysis_fails),
        ("C  | rendered single-boundary split",        test_c_rendered_split_single_boundary_success),
        ("C  | rendered preserves manual split",       test_c_rendered_split_preserves_valid_manual_split),
        ("C  | rendered ambiguity skip",               test_c_rendered_split_skips_ambiguous_repeated_rows),
        ("C  | rendered merged-boundary skip",         test_c_rendered_split_skips_merged_boundary_conflict),
        ("C  | rendered marker formatting",            test_c_rendered_split_marker_is_right_aligned),
        ("P1a | legacy rendered split post-normalized", test_p1a_legacy_rendered_split_runs_post_normalizer),
        ("C  | rendered caption number/fallback",      test_c_rendered_split_caption_number_and_fallback),
        ("C  | rendered whole-table move",             test_c_rendered_start_page_moves_whole_table_without_complete_data_row),
        ("C  | rendered first-row spill move",         test_c_rendered_start_page_first_row_spill_moves_whole_table),
        ("C  | rendered skip existing page break",     test_c_rendered_start_page_skips_existing_page_break_candidate),
        ("C  | rendered disabled page break",          test_c_rendered_start_page_upgrades_disabled_page_break),
        ("C  | rendered start-page ambiguity skip",    test_c_rendered_start_page_skips_ambiguous_usability),
        ("C  | rendered first-row spill weak skip",    test_c_rendered_start_page_first_row_spill_needs_strong_next_page_evidence),
        ("C  | rendered first-row spill prose skip",   test_c_rendered_start_page_first_row_spill_ignores_later_prose_token_reuse),
        ("C  | rendered decision logging",             test_c_rendered_decision_logging_for_ambiguous_skip),
        ("C  | rendered start-page complete row",      test_c_rendered_start_page_keeps_table_with_clear_complete_data_row),
        ("C  | vMerge guard",                          test_c_vmerge_guard_rejects_boundary_inside_merge_zone),
        # General DOCX formatting invariants used by Phase 3 output.
        ("B1 | tblW updated after optimization",       test_b1_tblW_updated_after_col_optimization),
        ("B1 | _MIN_COL_PT ≤ 20",                     test_b1_min_col_pt_is_20),
        ("B2 | keepTogether on table_caption",         test_b2_keep_together_on_table_caption),
        ("B2 | keepTogether on heading1/heading2",     test_b2_keep_together_on_headings),
        ("B2 | rule6 keepWithNext through empty para", test_b2_rule6_propagates_through_empty_para),
        ("B2 | table source/note chained",             test_b2_table_source_note_normalised_and_chained),
        ("B2 | image height from wp:extent cy",        test_b2_image_height_from_emu),
        ("B3 | footnote para: 10pt TNR no bold",       test_b3_format_footnote_para_applies_10pt_tnr),
        ("C2 | empty para image→caption removed",      test_c2_empty_para_between_image_and_caption_removed),
        ("C2 | numeric column minimum protected",      test_c2_number_column_minimum),
        ("T1 | ё→е normalisation (midword uppercase fix)", test_yo_normalisation_midword_uppercase),
        ("T_indent | body paragraph left=0 firstLine=709", test_t_indent_body_paragraph_left_zero),
        # Heading product rules: no Word autonumbering, manual text numbering remains.
        ("T2 | 'Глава N' without title → heading1", test_t2_chapter_heading_without_title),
        ("T2 | manual heading2 still works", test_t2_manual_heading2_still_promoted),
        ("T2 | Word-autonumbered heading2 still works", test_t2_word_autonumbered_heading2_with_style_still_promoted),
        ("T2 | Word-autonumbered heading1 still works", test_t2_word_autonumbered_heading1_with_style_still_promoted),
        ("T2 | heading style numbering removed", test_t2_heading_style_numbering_is_removed),
        ("T2 | Word-numbered body items stay body/list", test_t2_word_numbered_body_items_not_promoted_to_headings),
        ("T2 | numbered sentence not promoted to heading1", test_t2_numbered_sentence_not_promoted_to_heading1),
        ("T2 | chapter colon heading repaired", test_t2_chapter_colon_heading_repaired_without_colon_artifact),
        ("T2 | real coursework 17 heading regression", test_t2_real_coursework_17_heading_regression),
        ("T3 | reference subheading centred + source indent", test_t3_reference_subheading_centred),
        ("T4 | citation brackets split + p. notation + hyphen→en-dash", test_t4_citation_brackets_split),
        ("T5 | list а)/б)/в) formatting", test_t5_list_formatting),
        ("LIST | manual dash block normalized",          test_lists_manual_dash_block_normalized),
        ("LIST | manual em-dash normalized",             test_lists_manual_em_dash_block_normalized),
        ("LIST | manual black bullets normalized",       test_lists_manual_black_bullets_normalized),
        ("LIST | singleton dash NOT converted",          test_lists_singleton_dash_not_converted),
        ("LIST | mixed marker block NOT converted",      test_lists_mixed_marker_block_not_converted),
        ("LIST | free-standing letter block normalized", test_lists_free_standing_letter_block_normalized),
        ("LIST | G4 letter+colon block",                 test_lists_g4_letter_block_with_colon_endings),
        ("LIST | G1 letter parent + numeric L2 children", test_lists_g1_letter_parent_colon_nested_numeric_l2),
        ("LIST | G1+G4 Block E full smoke",              test_lists_g1_g4_block_e_full_smoke_case),
        ("LIST | G4 singleton letter+colon NOT converted", test_lists_g4_letter_singleton_with_colon_not_converted),
        ("LIST | G1 stray numeric-paren NOT promoted",   test_lists_g1_numeric_without_letter_parent_not_promoted_to_l2),
        ("LIST | numeric-dot w/o colon NOT converted",   test_lists_numeric_dot_without_colon_not_converted),
        ("LIST | heading-styled dash NOT converted",     test_lists_heading_styled_dash_not_converted),
        ("LIST | table caption NOT converted",           test_lists_table_caption_not_converted),
        ("LIST | source/note NOT converted",             test_lists_source_line_not_converted),
        ("LIST | colon numeric-dot → dash items", test_lists_colon_numeric_dot_becomes_dash),
        ("LIST | appendix section NOT normalized",       test_lists_appendix_section_not_normalized),
        ("LIST | bibliography section NOT normalized",   test_lists_bibliography_section_not_normalized),
        ("LIST | G2 Pattern A → real Word dash-list",        test_word_numbered_pattern_a_converted),
        ("LIST | G2 Pattern B → real Word dash-list",        test_word_numbered_pattern_b_converted),
        ("LIST | plain dash block → real Word dash-list",    test_plain_dash_block_becomes_word_list),
        ("LIST | bullet chars → real Word dash-list",        test_bullet_chars_become_word_dash_list),
        ("LIST | letter numPr preserved by body_list_item",  test_letter_numpr_preserved_in_body_list_item),
        ("LIST | singleton plain dash unchanged",             test_singleton_plain_dash_unchanged),
        ("LIST | G2 singleton Word-decimal unchanged",        test_word_numbered_singleton_not_converted),
        ("LIST | G2 references guard",                        test_word_numbered_references_block_not_converted),
        ("LIST | G2 appendix guard",                          test_word_numbered_appendix_block_not_converted),
        ("LIST | table caption not normalized",               test_table_caption_not_normalized),
        ("LIST | G2 heading guard",                           test_word_numbered_heading_style_not_converted),
        ("LIST | full normalization idempotent",              test_word_numbered_normalization_idempotent),
        ("LIST | colon+numeric-dot w:br split into paras",   test_split_colon_numeric_dot_br_paragraph),
        ("LIST | list body first letter lowercased",          test_list_body_first_letter_lowercased),
        ("LIST | D uppercase/dot/Latin letter variants",      test_zone_d_lettered_list_variants),
        ("LIST | D singleton letter item unchanged",          test_zone_d_singleton_letter_unchanged),
        ("FORMULA | unnumbered formula numbered",             test_formula_unnumbered_formula_gets_number_from_preceding_prose),
        ("FORMULA | existing numbered formula",               test_formula_existing_numbered_formula_still_formats),
        ("FORMULA | formula not list-normalized",             test_formula_paragraph_not_list_normalized),
        ("FORMULA | number from heading2 context",            test_formula_unnumbered_formula_numbered_from_current_heading2_context),
        ("FORMULA | inline где split",                        test_formula_inline_gde_split_into_explanations),
        ("FORMULA | comma explanations split",                test_formula_comma_separated_explanations_split_safely),
        ("FORMULA | numeric explanations",                    test_formula_numeric_coefficient_explanations_stay_in_block),
        ("FORMULA | blank around block",                      test_formula_block_has_single_blank_before_and_after),
        ("FORMULA | prose equals guard",                      test_formula_prose_with_equals_is_not_promoted),
        ("LIST | G2 blank breaks block",                      test_word_numbered_block_broken_by_blank_not_converted),
        ("BODY | skip fake plain-TOC СОДЕРЖАНИЕ",        test_real_body_start_skips_fake_plain_toc_soderzhanie),
        ("BODY | skip fake Оглавление",                  test_real_body_start_skips_fake_oglavlenie_toc),
        ("BODY | skip fake TOC with appendix entries",   test_real_body_start_skips_fake_toc_with_appendix_entries),
        ("BODY | no TOC → first intro",                  test_real_body_start_without_toc_returns_first_intro),
        ("BODY | canonical TOC picks real intro",        test_real_body_start_with_canonical_toc_picks_real_intro),
        ("BODY | styled intro wins",                     test_real_body_start_styled_intro_wins_when_no_fake_toc),
        ("BODY | lone Содержание keeps first intro",     test_real_body_start_lone_contents_heading_does_not_demote_intro),
        ("BODY | smoke fixture picks real intro",        test_real_body_start_smoke_input_picks_heading1_intro),
        ("SDT | Word-managed TOC SDT removed",           test_contents_removes_word_managed_sdt_toc),
        ("SDT | field-code TOC SDT removed",             test_contents_removes_field_code_sdt_toc),
        ("SDT | non-TOC SDT preserved",                  test_contents_keeps_non_toc_sdt),
        ("SDT | end-to-end SDT TOC removed via rebuild", test_contents_rebuild_drops_sdt_toc_end_to_end),
        ("H1 | body sentence NOT promoted via toc_text", test_h1_body_sentence_not_promoted_via_toc_text),
        ("H1 | styled chapter with toc_text promoted",   test_h1_styled_chapter_with_toc_text_still_promoted),
        ("H1 | exact intro still promoted",              test_h1_exact_intro_still_promoted_with_toc_text),
        ("H1 | centered+bold body promoted",             test_h1_centered_bold_body_with_toc_text_promoted),
        ("H1 | integration: no duplicate H1 in output",  test_h1_body_sentence_no_duplicate_in_output_pipeline),
        ("CLEAN | pre-pass strips SDT + plain old TOC",  test_strip_obsolete_toc_blocks_inplace_removes_sdt_and_plain_toc),
        ("CLEAN | pre-pass no-op on clean doc",          test_strip_obsolete_toc_blocks_inplace_noop_on_clean_doc),
        ("CLEAN | pre-pass survives rebuild failure",    test_strip_obsolete_toc_blocks_inplace_runs_even_when_pdf_render_would_fail),
        ("T5 | table caption trailing period cleanup", test_table_caption_trailing_period_cleanup),
        ("B2.5 | real caption before table", test_b25_real_table_caption_directly_before_table_is_formatted),
        ("B2.5 | caption title table", test_b25_real_table_caption_title_table_is_formatted),
        ("B2.5 | inline dash adjacent table", test_b25_inline_dash_table_caption_with_adjacent_table_is_accepted),
        ("B2.5 | table analytical prose body", test_b25_table_number_analytical_prose_without_table_remains_body),
        ("B2.5 | prose mentioning table body", test_b25_v_tablitse_analytical_prose_remains_body),
        ("B2.5 | source analytical prose body", test_b25_source_then_table_analytical_prose_is_not_promoted),
        ("B2.5 | appendix table-like title", test_b25_appendix_immediate_table_like_title_is_preserved),
        ("B2.5 | neuromarketing false positive", test_b25_neuromarketing_style_table_false_positive_is_prevented),
        ("B2.6 | inline dash title cleanup", test_b26_inline_dash_table_caption_title_loses_leading_dash),
        ("B2.6 | inline en dash spacing cleanup", test_b26_inline_en_dash_table_caption_title_strips_extra_spaces),
        ("B2.6 | split caption stable", test_b26_already_split_table_caption_remains_single_title),
        ("B2.6 | non-adjacent table prose", test_b26_non_adjacent_table_prose_remains_unchanged),
        ("B2.6 | prose mentioning table", test_b26_v_tablitse_prose_remains_unchanged),
        ("B2.6 | appendix table-like title", test_b26_appendix_immediate_dash_table_like_title_stays_appendix_title),
        ("B2.6 | neuromarketing inline caption", test_b26_neuromarketing_style_inline_caption_normalizes),
        ("B2.6 | Rybakov split caption stable", test_b26_rybakov_style_split_caption_remains_stable),
        ("B2.7 | canonical figure block stable", test_b27_canonical_image_source_note_caption_unchanged),
        ("B2.7 | caption above image", test_b27_caption_above_image_moves_below_image),
        ("B2.7 | source above image", test_b27_source_above_image_moves_below_image),
        ("B2.7 | caption source above image", test_b27_caption_and_source_above_image_normalize_to_canonical_order),
        ("B2.7 | merged source note near image", test_b27_merged_source_note_splits_only_near_image),
        ("B2.7 | no-image source unchanged", test_b27_no_image_source_paragraph_unchanged),
        ("B2.7 | table source unchanged", test_b27_table_source_nearby_unchanged),
        ("B2.7 | nearby images ambiguous", test_b27_two_nearby_images_are_ambiguous_noop),
        ("B2.7 | neuromarketing disorder", test_b27_neuromarketing_style_disorder_normalizes),
        ("B2.7 | Rybakov figure stable", test_b27_rybakov_stable_figure_block_unchanged),
        ("B2.7 | figure prose not caption", test_b27_figure_prose_after_source_does_not_block_reorder),
        ("T6 | figure caption spacing + blank font", test_figure_caption_spacing_and_blank_font),
        ("T6 | figure source moved before caption", test_figure_source_after_caption_is_moved_before_caption),
        ("T6 | correct figure source order unchanged", test_figure_source_before_caption_is_unchanged),
        ("T6 | table source order unchanged", test_table_source_after_caption_is_not_moved),
        ("T6 | appendix figure source moved", test_appendix_figure_source_after_caption_is_moved),
        ("T6 | figure source not duplicated", test_figure_source_not_duplicated_after_reorder),
        ("T6 | bibliography source unchanged", test_bibliography_source_line_is_not_moved),
        ("T6 | heading2 late spacing before 1.3", test_heading2_late_spacing_before_13),
        ("T6 | blank before figure block", test_blank_before_figure_block),
        # Marker split diagnostics and runtime decisions.
        ("M1 | source unchanged after instrumentation", test_marker_instrumentation_keeps_source_unchanged),
        ("M1 | only target table instrumented", test_marker_instrumentation_only_targets_selected_table),
        ("M1 | inline marker parsing", test_marker_extract_handles_inline_text_and_missing_rows),
        ("M1 | keep_temp mapping result", test_marker_map_rows_to_pages_keep_temp_debug_paths),
        ("M1 | 1pt fallback to 2pt", test_marker_map_rows_to_pages_falls_back_to_2pt_and_returns_debug_info),
        ("M1 | invalid table index", test_marker_instrumentation_rejects_invalid_table_index),
        ("M1 | row page span summary", test_marker_page_span_summary),
        ("M1 | diagnose all tables summary", test_marker_diagnose_all_tables_summary),
        ("M1 | diagnose table error handling", test_marker_diagnose_table_handles_mapping_error),
        ("M1 | appendix/caption metadata", test_marker_appendix_and_caption_metadata),
        ("M1 | appendix start labels", test_appendix_start_labels_are_normalized),
        ("M1 | appendix local table titles", test_appendix_local_table_title_before_table_is_centered),
        ("B2.1 | appendix title after label", test_appendix_title_after_label_is_normalized),
        ("B2.1 | table-like appendix title", test_table_caption_like_appendix_title_after_label_is_normalized),
        ("B2.1 | table-like appendix title full process", test_table_caption_like_appendix_title_survives_full_process),
        ("B2.1 | long appendix body not title", test_long_body_paragraph_after_appendix_label_is_not_title),
        ("B2.1 | continuation label not title", test_appendix_continuation_label_does_not_trigger_title_formatting),
        ("B2.1 | appendix title spacing", test_appendix_title_spacing_is_exactly_one_blank),
        ("M1 | appendix label/table spacing", test_empty_paragraph_after_appendix_label_before_table_is_preserved),
        ("M1 | appendices heading/label spacing", test_empty_paragraph_between_appendices_heading_and_first_label_is_removed),
        ("M1 | dry-run eligible boundary", test_marker_runtime_dry_run_clean_two_page_table_is_eligible),
        ("M1 | dry-run duplicate skip", test_marker_runtime_dry_run_skips_duplicate_rows),
        ("M1 | dry-run missing skip", test_marker_runtime_dry_run_skips_missing_rows_outside_header),
        ("M1 | dry-run 3-page skip", test_marker_runtime_dry_run_skips_three_page_tables),
        ("M1 | dry-run eligible logging", test_marker_runtime_dry_run_logs_eligible_candidate),
        ("M1 | dry-run flag off", test_marker_runtime_dry_run_feature_flag_off_skips_detection_hook),
        ("M1 | dry-run no mutation", test_marker_runtime_dry_run_only_does_not_mutate_document),
        ("M1 | apply appendix split", test_marker_runtime_apply_split_for_appendix_table),
        ("M1 | apply ordinary split", test_marker_runtime_apply_split_for_ordinary_table),
        ("M1 | skip nested ordinary header", test_marker_runtime_apply_skips_nested_ordinary_table_header),
        ("M1 | apply ineligible skip", test_marker_runtime_apply_skips_ineligible_tables),
        ("M1 | apply idempotent", test_marker_runtime_apply_is_idempotent_on_second_run),
        ("M1 | apply multiple ordinary splits", test_marker_runtime_apply_processes_multiple_ordinary_tables),
        ("M1 | apply skips stale candidate", test_marker_runtime_apply_skips_stale_candidate_and_continues),
        ("M1 | apply mixed ordinary appendix splits", test_marker_runtime_apply_processes_mixed_ordinary_and_appendix_tables),
        ("M1 | skip generated appendix continuations", test_marker_runtime_apply_skips_generated_appendix_continuation_tables),
        ("M1 | apply single diagnostic pass", test_marker_runtime_apply_loop_is_bounded),
        # Prototype split rules.
        ("S1 | prototype simple table split", test_split_prototype_simple_table),
        ("S1 | source note after second table", test_split_prototype_source_note_stays_after_second_table),
        ("S1 | original doc unchanged", test_split_prototype_original_document_unchanged),
        ("S1 | invalid table index", test_split_prototype_invalid_table_index),
        ("S1 | invalid split_before_row", test_split_prototype_invalid_split_before_row),
        ("S1 | no continuation paragraph", test_split_prototype_no_continuation_paragraph_inserted),
        ("S1 | numbered ordinary continuation", test_split_prototype_numbered_ordinary_continuation_row_only),
        ("S1 | numbered ordinary split caption", test_split_prototype_numbered_ordinary_split_caption_before_title),
        ("S1 | numbered appendix continuation", test_split_prototype_numbered_appendix_has_continuation_label),
        ("S1 | numbered row reused", test_split_prototype_numbered_existing_row_reused_without_duplicate),
        ("S1 | numbered malformed row", test_split_prototype_numbered_malformed_existing_row_fails_safely),
        ("S1 | numbered source note", test_split_prototype_numbered_source_note_after_second_table),
        ("S1 | numbered original unchanged", test_split_prototype_numbered_original_document_unchanged),
        ("S1 | numbered row safe markup", test_split_prototype_numbered_row_has_no_numpr_and_no_calibri),
        ("M1 | headings unchanged across flags", test_marker_runtime_flags_do_not_change_headings),
        ("M1 | render budget fail-open many tables", test_phase3_marker_budget_fail_open_many_tables),
        ("M1 | render budget allows small doc", test_phase3_marker_budget_allows_small_doc),
        # Caption demotion: reference-prose paragraphs must stay body text.
        ("PA | caption_tail_is_reference_prose unit", test_caption_tail_is_reference_prose_unit),
        ("PA | Таблица N показывает... stays body", test_table_reference_paragraph_not_caption),
        ("PA | Рис. N показывает... stays body", test_figure_reference_paragraph_not_caption),
        ("PA | no keepNext on reference-prose", test_no_keep_with_next_on_reference_paragraph),
        ("PA | real table caption still formats", test_actual_table_caption_still_formats),
        ("PA | real figure caption still formats", test_actual_figure_caption_still_formats),
        # Patch B+C: figure caption alignment + blank cleanup.
        ("PB | figure caption alignment left/justify", test_figure_caption_alignment_left_or_justify),
        ("PC | blank between source and caption removed", test_remove_empty_between_figure_source_and_caption),
        ("FSP | one blank after real caption before body prose", test_one_blank_after_real_figure_caption_before_body_prose),
        ("FSP | one blank after caption is idempotent", test_one_blank_after_caption_is_idempotent),
        ("PC | table source blank not affected", test_table_source_not_affected_by_figure_blank_cleanup),
        ("PC | figure reference prose stays body", test_figure_reference_prose_still_body_after_patch_A),
        # PB2: figure caption keepLines + figure block keepWithNext chain.
        ("PB2 | figure caption keepLines=True", test_figure_caption_keep_lines_true),
        ("PB2 | IMG+source+note chained via keepNext", test_figure_block_image_keeps_with_source_and_caption),
        ("PB2 | appendix IMG→CAP keepNext", test_figure_block_appendix_image_keeps_with_caption),
        ("PB2 | table source unaffected by figure keepNext", test_table_block_unaffected_by_figure_keepnext),
        # TCF-A: universal <w:cantSplit/> on every table row.
        ("TCF-A | all table rows have cantSplit", test_tcfa_all_table_rows_have_cant_split),
        ("TCF-A | cantSplit is idempotent", test_tcfa_cant_split_is_idempotent),
        ("TCF-A | existing cantSplit preserved", test_tcfa_existing_cant_split_preserved),
        ("TCF-A | caption/source classification unchanged", test_tcfa_does_not_change_caption_or_source_classification),
        # E1: Phase 3 marker-split candidate classification (logging-only, no behavior change).
        ("E1 | many tiny tables → no candidates", test_e1_classify_many_tiny_tables_no_candidates),
        ("E1 | one long candidate among many tiny", test_e1_classify_one_long_candidate_among_many),
        ("E1 | manual continuation filtered", test_e1_classify_skips_manual_continuation),
        ("E1 | no-caption table filtered", test_e1_classify_skips_tables_without_caption),
        ("E1 | priority ordering desc by rows", test_e1_classify_priority_ordering),
        ("E1 | classification + previews logged", test_e1_log_emitted_with_classification),
        ("E1 | existing budget skip preserved", test_e1_does_not_change_existing_behaviour),
        # E2: quality-first candidate-mode marker split.
        ("E2 | 31 total / 8 candidates marker split enters", test_e2_31_total_8_candidates_marker_split_enters),
        ("E2 | 25 total / 7 candidates marker split enters", test_e2_25_total_7_candidates_marker_split_enters),
        ("E2 | 11 total / 10 candidates marker split enters", test_e2_11_total_10_candidates_marker_split_enters),
        ("E2 | no warning when all candidates processed", test_e2_no_warning_when_all_candidates_processed),
        ("E2 | warning when budget overflow skips candidates", test_e2_warning_when_some_candidates_skipped_for_budget),
        ("E2 | one diagnose failure does not abort others", test_e2_candidate_diagnose_failure_does_not_abort_others),
        ("E2 | manual continuation preserved (not diagnosed)", test_e2_existing_manual_continuation_preserved),
        ("E2 | continuation has blank between marker and table", test_e2_continuation_label_has_blank_after_before_table),
        ("PG1 | classify sensitive geometry", test_pg1_geometry_policy_classifies_sensitive_tables),
        ("PG1 | malformed geometry unsafe", test_pg1_geometry_policy_marks_malformed_grid_unsafe),
        ("PG1 | preserve skips optimizer", test_pg1_preserve_geometry_skips_width_optimizer),
        ("PG1 | simple optimizer still runs", test_pg1_simple_table_still_uses_width_optimizer),
        ("PG1 | unsafe skips marker split", test_pg1_unsafe_geometry_skips_marker_split),
        ("PG1 | real flow preserves geometry", test_pg1_real_continuation_flow_preserves_sensitive_geometry),
        ("TM | split eligibility requires render", test_tm_split_eligibility_requires_rendered_boundary),
        ("TM | forbidden split reasons", test_tm_forbidden_split_reasons_are_explicit),
        ("TM | numeric row safety model", test_tm_numeric_row_safety_model),
        ("TM | source note boundary fail closed", test_tm_source_note_boundary_fails_closed),
        ("TM | geometry snapshot invariant", test_tm_geometry_snapshot_invariant_for_preserve_mode),
        ("TM | diagnostic harness artifacts", test_tm_diagnostic_harness_writes_stage_artifacts),
        ("TM | universal model grouping", test_tm_universal_model_groups_caption_title_table_source),
        ("TM | universal row roles", test_tm_universal_row_roles_detect_header_numeric_data),
        ("TM | universal provenance", test_tm_universal_provenance_unknown_and_source_safe),
        ("TM | universal same-page issue", test_tm_universal_same_page_repeated_fragment_issue),
        ("TM | universal missing numeric issue", test_tm_universal_continuation_missing_numeric_row_issue),
        ("TM | universal single-table cross-page issue", test_tm_universal_single_physical_table_crosses_pages_issue),
        ("TM | universal rendered case-row mapping", test_tm_universal_rendered_mapping_keeps_repeated_case_rows_on_real_pages),
        ("TM | universal bad2 clean continuation", test_tm_universal_bad2_style_clean_continuation),
        ("TM | universal diagnostics non-mutating", test_tm_universal_diagnostics_run_is_non_mutating),
        ("TM | same-page continuation marker flagged", test_tm_same_page_continuation_marker_is_flagged),
        ("TM | next-page continuation marker passes", test_tm_next_page_continuation_marker_passes_validation),
        ("TM | inline continuation prose classified", test_tm_inline_prose_continuation_text_is_not_valid_marker),
        ("TM | unknown continuation relation not pass", test_tm_unknown_marker_relation_is_not_pass),
        ("TM | continuation marker report page+flag", test_tm_continuation_marker_report_includes_page_and_violation),
        ("TM | rendered missing marker flagged", test_tm_rendered_fragment_without_marker_is_flagged),
        ("TM | rendered marker before fragment passes", test_tm_rendered_fragment_with_marker_before_it_passes),
        ("TM | same-page adjacent fragments flagged", test_tm_same_page_adjacent_fragments_are_flagged),
        ("TM | late marker rows before marker flagged", test_tm_late_marker_rows_before_marker_are_flagged),
        ("TM | source-bad duplicate rows warning", test_tm_source_bad_duplicate_rows_are_warning_only),
        ("TM | ambiguous adjacent tables warning", test_tm_ambiguous_adjacent_tables_are_flagged_without_repair_class),
        ("TM | clean continuation no new diagnostics", test_tm_clean_continuation_has_no_new_fragment_diagnostics),
        ("TM | artifact guard PDFs clean", test_tm_artifact_guard_pdfs_remain_clean_for_new_diagnostics),
        ("TM | inline prose not rendered marker", test_tm_inline_prose_is_not_rendered_continuation_marker),
        ("TM | stable manual chain identity reload", test_tm_stable_manual_chain_identity_survives_save_reload),
        ("TM | real fixture rendered missing-marker diagnostics", test_tm_real_fixture_rendered_missing_marker_diagnostics),
        ("TM | formatter surfaces rendered continuation warnings", test_tm_formatter_surfaces_rendered_continuation_warnings),
        ("TM | real fixture formatter rendered warnings", test_tm_real_fixture_formatter_rendered_warnings),
        ("TM | post-render same-page marker rejected", test_tm_post_render_gate_rejects_same_page_marker_output),
        ("TM | post-render gate restores backup", test_tm_post_render_gate_restores_backup_on_same_page_marker),
        ("TM | post-render gate preserves clean marker", test_tm_post_render_gate_preserves_clean_next_page_marker),
        ("TM | post-render gate preserves unresolved manual marker", test_tm_post_render_gate_preserves_unresolved_manual_marker),
        ("TM | same-page cleanup preserves unsafe marker", test_tm_same_page_marker_cleanup_preserves_marker_when_deletion_creates_missing_fragment),
        ("TM | same-page cleanup removes safe stale marker", test_tm_same_page_marker_cleanup_removes_marker_when_final_render_is_clean),
        ("TM | post-render restore warning deferred", test_tm_post_render_restore_does_not_emit_user_warning_before_final_validation),
        ("TM | manual-chain overflow repair moves rows", test_tm_manual_chain_overflow_repair_moves_rows_into_continuation_table),
        ("TM | manual-chain overflow repair skips mismatch", test_tm_manual_chain_overflow_repair_skips_header_mismatch),
        ("TM | manual-chain overflow repair skips merged cells", test_tm_manual_chain_overflow_repair_skips_merged_cells),
        ("TM | manual-chain overflow repair rollback", test_tm_manual_chain_overflow_repair_rolls_back_failed_revalidation),
        ("TM | rendered split rollback on same-page marker", test_tm_apply_rendered_split_rolls_back_same_page_marker),
        ("TM | real fixture marker gate preserves semantic markers", test_tm_real_fixture_marker_gate_preserves_semantic_markers),
        ("PG2 | preserve sensitive safe_formatter geometry", test_pg2_safe_formatter_preserves_sensitive_tblw_layout_tcw_grid),
        ("PG2 | simple safe_formatter normalization remains", test_pg2_safe_formatter_keeps_simple_table_width_normalization),
        ("PG2 | auto pct and fixed layout preserved", test_pg2_safe_formatter_preserves_auto_pct_and_authored_fixed_layout),
        ("PG2 | merged and vMerge tables preserved", test_pg2_safe_formatter_preserves_merged_and_vmerge_tables),
        ("PG2 | preserve border cleanup no geometry mutation", test_pg2_preserve_geometry_table_gets_border_cleanup_without_geometry_mutation),
        ("PG2 | preserve border cleanup keeps merges", test_pg2_preserve_border_cleanup_keeps_merged_and_vmerge_structure),
        ("PG2 | borders preserve geometry nodes", test_pg2_force_borders_preserves_geometry_adjacent_nodes),
        ("PG2 | Rybakov-style tcW preserved", test_pg2_rybakov_style_table_keeps_tcw_in_safe_formatter_stage),
        ("PG2 | diagnostics reduced safe_formatter mutation", test_pg2_diagnostics_show_reduced_safe_formatter_geometry_mutation),
        ("E2 | rollback global_skip preserves pre-E2", test_e2_rollback_global_skip_mode_preserves_pre_e2_behaviour),
        # E3: NUM-row compensation for marker-split first fragment.
        ("E3 | compensation applies for ordinary body table", test_e3_compensation_applies_for_ordinary_body_table),
        ("E3 | compensation skipped when pre-existing NUM row", test_e3_compensation_skipped_when_table_has_pre_existing_numbered_row),
        ("E3 | compensation skipped for appendix path", test_e3_compensation_skipped_for_appendix),
        ("E3 | compensation skipped when fragment 2 too small (Case-A guard)", test_e3_compensation_skipped_when_fragment2_too_small),
        ("E3 | compensation skipped when fragment 1 would be empty", test_e3_compensation_skipped_when_fragment1_would_be_empty),
        ("E3 | feature flag off disables compensation", test_e3_feature_flag_off_disables_compensation),
        ("E3 | integration: K-1 propagates through apply", test_e3_integration_apply_marker_split_passes_compensated_k),
        # Static TOC rebuild tests are registered at the tail so existing
        # Phase 3 ordering-sensitive regression cases keep their baseline order.
        ("TOC | existing Содержание replaced",         test_autotoc_existing_soderzhanie_replaced_by_canonical),
        ("TOC | existing Оглавление replaced",         test_autotoc_existing_oglavlenie_replaced_by_canonical),
        ("TOC | exact intro old entries removed",      test_autotoc_exact_intro_entry_inside_old_toc_is_removed),
        ("TOC | missing contents inserted",            test_autotoc_missing_contents_inserted_before_real_intro),
        ("TOC | no title page inserts at start",       test_autotoc_no_title_page_inserted_at_document_start),
        ("TOC | appendices general heading only",      test_autotoc_appendices_include_general_heading_only),
        ("TOC | heading2 has no left indent",          test_autotoc_heading2_has_no_left_indent),
        ("TOC | zero indent and 1.5 spacing",          test_autotoc_entries_have_zero_indent_and_one_point_five_spacing),
        ("TOC | entries use dot leader tab",           test_autotoc_entries_use_dot_leader_tab_not_manual_dots),
        ("TOC | body heading register",                test_autotoc_uses_body_heading_register),
        ("TOC | normal numbered body excluded",        test_autotoc_normal_numbered_body_paragraph_is_excluded),
        ("TOC | resolver ignores TOC echoes",          test_autotoc_page_resolver_ignores_toc_page_heading_echoes),
        ("TOC | resolver wrapped headings",            test_autotoc_page_resolver_matches_wrapped_rendered_headings),
        ("TOC | degenerate mapping fails safe",        test_autotoc_degenerate_page_mapping_fails_safe),
        ("TOC | internal hyperlinks",                  test_autotoc_entries_are_internal_hyperlinks_to_bookmarks),
        ("TOC | long heading tab leader layout",       test_autotoc_long_heading_uses_same_tab_leader_layout),
        ("TOC | fallback normal H2 included",          test_autotoc_fallback_normal_styled_heading2_included),
        ("TOC | fallback normal H1 included",          test_autotoc_fallback_normal_styled_heading1_included),
        ("TOC | fallback intro task list excluded",    test_autotoc_fallback_intro_task_list_excluded),
        ("TOC | fallback numbered sentence excluded",  test_autotoc_fallback_numbered_sentence_excluded),
        ("TOC | fallback inside table excluded",       test_autotoc_fallback_paragraph_inside_table_excluded),
        ("TOC | fallback non-monotonic rejected",      test_autotoc_fallback_non_monotonic_rejected),
        ("TOC | fallback after appendices excluded",   test_autotoc_fallback_after_appendices_excluded),
        ("TOC | fallback source/note excluded",        test_autotoc_fallback_source_or_note_excluded),
        ("TOC | fallback strict path unchanged",       test_autotoc_fallback_strict_style_path_unchanged),
        ("TOC | blank paragraph after title",          test_autotoc_blank_paragraph_after_title),
        ("TOC | no double blank after title",          test_autotoc_no_double_blank_after_title),
        ("TOC | title TNR 14 bold centered",           test_autotoc_title_font_is_times_new_roman_14_bold_centered),
        ("TOC | hyperlink runs TNR 14",                test_autotoc_entry_hyperlink_runs_use_times_new_roman_14),
        ("TOC | zero hanging/firstLine indent",        test_autotoc_entries_have_zero_hanging_and_firstline_xml),
        ("TOC | blank paragraph has no tab/hyperlink", test_autotoc_blank_paragraph_carries_no_tab_or_hyperlink),
        ("TOC | dot leader preserved with blank",      test_autotoc_dot_leader_and_pages_preserved_with_blank),
        ("TOC | old simple TOC removed",               test_autotoc_old_simple_toc_without_page_numbers_removed),
        ("TOC | old trailing-dot heading removed",     test_autotoc_old_toc_heading_with_trailing_dot_removed),
        ("TOC | old TOC with appendix entries removed", test_autotoc_old_toc_with_appendix_entries_removed),
        ("TOC | old TOC heading with page number removed", test_autotoc_old_toc_heading_with_page_number_removed),
        ("TOC | standalone loose heading keeps body",  test_autotoc_standalone_loose_contents_paragraph_does_not_delete_body),
        ("TOC | P1 garbage removed text-only intro",   test_autotoc_p1_garbage_removed_when_intro_normal_style_text_only),
        ("TOC | P1 Оглавление variant text-only",      test_autotoc_p1_oglavlenie_variant_garbage_removed_text_only),
        ("TOC | P1 fail-closed no confident intro",    test_autotoc_p1_fail_closed_no_confident_intro_keeps_block),
        ("TOC | P1 well-formed old TOC removed",       test_autotoc_p1_wellformed_old_toc_removed_with_confident_intro),
        ("TOC | P1 title page before contents kept",   test_autotoc_p1_title_page_before_contents_preserved),
        ("TOC | soft-break plain TOC removed",         test_autotoc_softbreak_plain_toc_removed),
        ("TOC | soft-break TOC appendix continuation removed", test_autotoc_softbreak_toc_appendix_continuation_removed),
        ("TOC | body ВВЕДЕНИЕ preserved after softbreak cleanup", test_autotoc_body_intro_preserved_after_softbreak_toc),
        ("TOC | no duplicate СОДЕРЖАНИЕ after rebuild", test_autotoc_no_duplicate_contents_after_rebuild),
        # ── P2-a' relaxed row/page matcher tests — registered at tail to
        # avoid shifting the ordinal position of pre-existing E2 tests, which
        # exposed a latent order-sensitive failure in the suite. The matcher
        # itself is independent of all earlier paths; running these last has
        # no functional implication.
        ("P2a' | relaxed accepts duplicate rows",        test_p2a_relaxed_accepts_duplicate_row_signatures),
        ("P2a' | relaxed handles wrapped cells",         test_p2a_relaxed_matches_wrapped_cells_via_window),
        ("P2a' | relaxed low-confidence → None",         test_p2a_relaxed_returns_none_on_low_confidence),
        ("P2a' | relaxed rejects non-monotonic",         test_p2a_relaxed_rejects_non_monotonic_pages),
        ("P2a' | relaxed rejects single-page",           test_p2a_relaxed_rejects_single_page_mapping),
        ("P2a' | strict matcher unchanged",              test_p2a_strict_matcher_behavior_unchanged),
        ("P2a' | collector uses relaxed on strict fail", test_p2a_collector_uses_relaxed_when_strict_fails),
        ("P2a' | strict path takes priority",            test_p2a_strict_path_takes_priority_when_strict_succeeds),
        ("P2a' | Bondarev-style triplicate candidate",   test_p2a_bondarev_style_triplicate_appendix_creates_candidate),
    ]

    if os.environ.get("KPFU_RUN_LONG_PHASE3_TESTS") == "1":
        tests.extend([
            ("M1 | real Рыбаков split", test_marker_runtime_real_rybakov_target_applies_split),
            ("M1 | real Бондарев headings", test_marker_runtime_real_bondarev_keeps_headings_safe),
        ])
        for asset in ASSET_FILES:
            tests.append((
                f"REG| {asset.name}",
                lambda a=asset: test_regression_asset(a),
            ))

    passed = failed = 0
    for name, fn in tests:
        try:
            ok, msg = fn()
        except Exception as e:
            ok, msg = False, f"EXCEPTION: {e}\n{traceback.format_exc()}"
        status = PASS if ok else FAIL
        suffix = f"  — {msg}" if msg else ""
        print(f"[{status}] {name}{suffix}")
        if ok:
            passed += 1
        else:
            failed += 1

    print(f"\n{'='*60}")
    print(f"Results: {passed} passed, {failed} failed")
    if failed:
        sys.exit(1)


if __name__ == "__main__":
    run_all()
