from __future__ import annotations

import re
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from guides.coursework_kfu_2025.safe_formatter import (
    canonical_numbered_reference_subheading_text,
    canonical_reference_block_heading_text,
    canonical_reference_subheading_text,
    cleanup_reference_subheadings_layout,
    ensure_blank_before_reference_subheadings,
    ensure_single_blank_after_references_heading,
    process_document,
    strip_leading_reference_number,
)


def _paragraph_texts(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs]


def _add_fake_word_numbering(paragraph) -> None:
    pPr = paragraph._element.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(num_id)
    pPr.append(numPr)


def test_reference_subheading_spacing() -> tuple[bool, str]:
    """Product rule: reference block subheadings keep a predictable one-blank layout."""
    doc = Document()
    doc.add_paragraph("Список использованных источников")
    doc.add_paragraph("нормативные правовые акты")
    doc.add_paragraph("1. Федеральный закон ...")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("3. статьи")
    doc.add_paragraph("2. Иванов И.И. Статья ...")
    doc.add_paragraph("3. Петров П.П. Источник ...")
    doc.add_paragraph("ДИССЕРТАЦИИ")
    doc.add_paragraph("4. Сидоров С.С. Диссертация ...")

    body_start = 0
    ensure_blank_before_reference_subheadings(doc, body_start)
    ensure_single_blank_after_references_heading(doc, body_start)
    cleanup_reference_subheadings_layout(doc, body_start)

    expected = [
        "Список использованных источников",
        "",
        "Нормативные правовые акты",
        "1. Федеральный закон ...",
        "",
        "Статьи",
        "2. Иванов И.И. Статья ...",
        "3. Петров П.П. Источник ...",
        "",
        "Диссертации",
        "4. Сидоров С.С. Диссертация ...",
    ]
    actual = _paragraph_texts(doc)
    if actual != expected:
        return False, f"unexpected paragraph layout:\nexpected={expected!r}\nactual={actual!r}"

    return True, "reference subheadings have exactly one blank before them"


def test_reference_subheading_detection_is_strict() -> tuple[bool, str]:
    """Product rule: reference entries are not reclassified as subheadings by loose text matches."""
    cases = [
        "1. Статьи",
        "• Статьи",
        "- Статьи",
        "Статьи и монографии",
        "Материалы интернет-сайтов: сайты",
        "статьи в периодических изданиях",
    ]
    for text in cases:
        if canonical_reference_subheading_text(text) is not None:
            return False, f"false reference subheading detected: {text!r}"

    if canonical_reference_subheading_text("статьи") != "Статьи":
        return False, "exact case-insensitive subheading was not detected"
    if (
        canonical_reference_subheading_text("статьи в периодических изданиях и сборниках")
        != "Статьи в периодических изданиях и сборниках"
    ):
        return False, "new exact reference subheading was not detected"
    if (
        canonical_numbered_reference_subheading_text("1. диссертации, авторефераты диссертаций")
        != "Диссертации, авторефераты диссертаций"
    ):
        return False, "new numbered reference subheading was not recovered"

    return True, "reference subheading detection is exact-match only"


def test_reference_block_heading_detection_is_flexible() -> tuple[bool, str]:
    """Product rule: known reference block headings survive minor wording and typo drift."""
    cases = {
        "Официальные материалы": "Официальные материалы",
        "Нормативные правовые акты": "Нормативные правовые акты",
        "Книги, монографии и диссертации": "Книги, монографии и диссертации",
        "Книги, монографии, диссертации": "Книги, монографии, диссертации",
        "Научные статьи": "Научные статьи",
        "Электронные ресурсы": "Электронные ресурсы",
        "Учебники и учебные пособия": "Учебники и учебные пособия",
        "Интернет-ресурсы": "Интернет-ресурсы",
        "3. Книги, монографии, диссертации": "Книги, монографии, диссертации",
        "Научные стати": "Научные статьи",
        "Электроные ресуры": "Электронные ресурсы",
    }
    for text, expected in cases.items():
        actual = canonical_reference_block_heading_text(text)
        if actual != expected:
            return False, f"reference block heading not detected: {text!r} -> {actual!r}, expected {expected!r}"

    false_cases = [
        (
            "Иванов И. И. Научные статьи как объект библиографического анализа "
            "// Вестник экономики. — 2024. — № 2. — С. 11–18."
        ),
        (
            "Федеральный закон от 27.07.2006 № 152-ФЗ «О персональных данных» "
            "// Собрание законодательства Российской Федерации. — 2006. — № 31."
        ),
        "ГОСТ Р 7.0.5-2008. Библиографическая ссылка. Общие требования и правила составления.",
        "Петров П. П. Электронные ресурсы в деятельности предприятия. URL: https://example.com",
    ]
    for text in false_cases:
        actual = canonical_reference_block_heading_text(text)
        if actual is not None:
            return False, f"real source entry was promoted to reference block heading: {text!r} -> {actual!r}"

    return True, "reference block heading detection is flexible and conservative"


def test_reference_short_centered_bold_heading_stays_unnumbered() -> tuple[bool, str]:
    """Product rule: short centered/bold reference-like headings are not numbered as sources."""
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Краткий текст введения.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    heading = doc.add_paragraph("Электронные источники")
    heading.alignment = 1
    run = heading.runs[0]
    run.bold = True
    doc.add_paragraph("Иванов И. И. Электронная коммерция: учебное пособие.")
    doc.add_paragraph("Научные стати")
    doc.add_paragraph("2. Петров П. П. Статья о цифровой экономике.")

    with tempfile.TemporaryDirectory() as tmp:
        input_path = Path(tmp) / "in.docx"
        output_path = Path(tmp) / "out.docx"
        doc.save(str(input_path))

        process_document(input_path, output_path)
        out_doc = Document(str(output_path))

    texts = _paragraph_texts(out_doc)
    try:
        refs_idx = texts.index("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    except ValueError:
        return False, "references heading missing after formatting"

    nonempty = [text for text in texts[refs_idx + 1:] if text]
    expected = [
        "Электронные источники",
        "1. Иванов И. И. Электронная коммерция: учебное пособие.",
        "Научные статьи",
        "2. Петров П. П. Статья о цифровой экономике.",
    ]
    actual = nonempty[:len(expected)]
    if actual != expected:
        return False, f"unexpected reference flow:\nexpected={expected!r}\nactual={actual!r}"

    for text in ("Электронные источники", "Научные статьи"):
        paragraph = next((p for p in out_doc.paragraphs if p.text == text), None)
        if paragraph is None:
            return False, f"missing reference block heading: {text!r}"
        if paragraph.alignment != 1:
            return False, f"reference block heading is not centered: {text!r}"

    return True, "short centered/bold and fuzzy headings remain unnumbered"


def test_numbered_reference_entries_are_not_headings() -> tuple[bool, str]:
    """Product rule: numbered reference entries remain body/list text, not headings."""
    false_heading = (
        "1. ТЕОРЕТИЧЕСКИЕ ОСНОВЫ КОММУНИКАЦИОННОЙ ПОЛИТИКИ "
        "В СИСТЕМЕ МАРКЕТИНГА ПРЕДПРИЯТИЯ"
    )

    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Краткий текст введения.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("1. Монографии и учебники")
    doc.add_paragraph("2. Бельзецкий А. И. Маркетология: монография.")
    doc.add_paragraph(false_heading)
    doc.add_paragraph(
        "4. Закон РФ от 07.02.1992 № 2300-1 «О защите прав потребителей» "
        "[Электронный ресурс]. — URL: https://example.com/(дата обращения: 06.03.2026)."
    )
    doc.add_paragraph("")
    doc.add_paragraph("5. статьи")
    doc.add_paragraph("6. Иванов И. И. Название статьи.")

    with tempfile.TemporaryDirectory() as tmp:
        input_path = Path(tmp) / "in.docx"
        output_path = Path(tmp) / "out.docx"
        doc.save(str(input_path))

        process_document(input_path, output_path)
        out_doc = Document(str(output_path))

    texts = _paragraph_texts(out_doc)
    try:
        refs_idx = texts.index("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    except ValueError:
        return False, "references heading missing after formatting"

    if texts[refs_idx + 1] != "":
        return False, "missing single blank after references heading"
    if texts[refs_idx + 2] != "Монографии и учебники":
        return False, f"numbered block heading was not recovered: {texts[refs_idx + 2]!r}"
    if not texts[refs_idx + 3].startswith("1. "):
        return False, f"first real source lost numbering: {texts[refs_idx + 3]!r}"
    if not texts[refs_idx + 4].startswith("2. "):
        return False, f"numbered reference entry lost numbering: {texts[refs_idx + 4]!r}"
    if "коммуникационной политики" not in texts[refs_idx + 4].lower():
        return False, f"numbered reference entry text changed unexpectedly: {texts[refs_idx + 4]!r}"
    if texts[refs_idx + 5].startswith("3. ") is False:
        return False, "next numbered reference entry changed unexpectedly"
    if "https://example.com/ (дата обращения" not in texts[refs_idx + 5]:
        return False, f"URL spacing was not normalized: {texts[refs_idx + 5]!r}"
    if texts[refs_idx + 6] != "":
        return False, "missing single blank before real reference subheading"
    if texts[refs_idx + 7] != "Статьи":
        return False, f"numbered reference subheading was not canonicalized: {texts[refs_idx + 7]!r}"
    if texts[refs_idx + 8] == "":
        return False, "unexpected blank after reference subheading"

    block_heading_para = out_doc.paragraphs[refs_idx + 2]
    if block_heading_para.alignment != 1:
        return False, "recovered block heading is not centered"

    false_heading_para = out_doc.paragraphs[refs_idx + 4]
    style_name = (false_heading_para.style.name or "").lower()
    if "heading" in style_name or "заголовок" in style_name:
        return False, f"numbered reference entry got heading style: {false_heading_para.style.name!r}"

    for offset in (3, 4, 5):
        pPr = out_doc.paragraphs[refs_idx + offset]._element.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        attrs = ind.attrib if ind is not None else {}
        if attrs.get(qn("w:left")) != "0":
            return False, f"reference entry has non-zero left indent: {attrs}"
        if attrs.get(qn("w:firstLine")) != "709":
            return False, f"reference entry does not have first-line indent 1.25 cm: {attrs}"
        if attrs.get(qn("w:hanging")) is not None:
            return False, f"reference entry still has hanging indent: {attrs}"

    hyperlinks = out_doc.paragraphs[refs_idx + 5]._element.findall(".//" + qn("w:hyperlink"))
    if not hyperlinks:
        return False, "plain URL was not converted to a DOCX hyperlink"

    return True, "numbered reference entries stay body text inside references"


def test_reference_old_numbering_cleanup() -> tuple[bool, str]:
    """Product rule: old manual/Word reference numbering is removed before clean numbering."""
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Краткий текст введения.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1) Иванов И. И. Учебное пособие.")
    doc.add_paragraph("[2] Петров П. П. Научная статья.")
    doc.add_paragraph("10.Семенов С. С. Источник без пробела после старого номера.")
    word_numbered = doc.add_paragraph("Сидоров С. С. Монография.")
    _add_fake_word_numbering(word_numbered)
    doc.add_paragraph("Ненумерованный источник.")
    doc.add_paragraph(
        "Гражданский кодекс Российской Федерации. Часть 1. Статья 10. "
        "Пределы осуществления гражданских прав."
    )

    with tempfile.TemporaryDirectory() as tmp:
        input_path = Path(tmp) / "in.docx"
        output_path = Path(tmp) / "out.docx"
        doc.save(str(input_path))

        process_document(input_path, output_path)
        out_doc = Document(str(output_path))

    texts = _paragraph_texts(out_doc)
    try:
        refs_idx = texts.index("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    except ValueError:
        return False, "references heading missing after formatting"

    entries = [text for text in texts[refs_idx + 1:] if text]
    expected = [
        "1. Иванов И. И. Учебное пособие.",
        "2. Петров П. П. Научная статья.",
        "3. Семенов С. С. Источник без пробела после старого номера.",
        "4. Сидоров С. С. Монография.",
        "5. Ненумерованный источник.",
        (
            "6. Гражданский кодекс Российской Федерации. Часть 1. Статья 10. "
            "Пределы осуществления гражданских прав."
        ),
    ]
    actual = entries[:len(expected)]
    if actual != expected:
        return False, f"unexpected reference numbering:\nexpected={expected!r}\nactual={actual!r}"

    duplicate_markers = ("1. 1", "1. [", "2. [", "1. 1)", "2. [2]")
    for entry in actual:
        if any(marker in entry for marker in duplicate_markers):
            return False, f"duplicated numbering remained: {entry!r}"

    legal_entry = entries[5]
    if "Часть 1. Статья 10." not in legal_entry:
        return False, f"legal/article numbering was corrupted: {legal_entry!r}"

    word_para = next((p for p in out_doc.paragraphs if "Сидоров С. С." in p.text), None)
    if word_para is None:
        return False, "Word-numbered reference entry missing"
    pPr = word_para._element.find(qn("w:pPr"))
    if pPr is not None and pPr.find(qn("w:numPr")) is not None:
        return False, "Word automatic numbering remained on reference entry"

    return True, "old reference numbering is stripped before clean sequential numbering"


def test_zone_j_new_reference_subheadings() -> tuple[bool, str]:
    """Zone J: newly added reference subheadings are unnumbered; entries after them are sequential."""
    # Unit checks: canonical lookups
    new_exact_cases = {
        "нормативно-правовые акты": "Нормативно-правовые акты",
        "Нормативно-правовые акты": "Нормативно-правовые акты",
        "законы и нормативные акты": "Законы и нормативные акты",
        "учебники и статьи": "Учебники и статьи",
        "учебная литература": "Учебная литература",
        "монографии": "Монографии",
        "интернет-источники": "Интернет-источники",
        "Интернет-источники": "Интернет-источники",
        "интернет источники": "Интернет-источники",
        "официальные сайты": "Официальные сайты",
        "материалы судебной практики": "Материалы судебной практики",
        "судебная практика": "Судебная практика",
        "иностранные источники": "Иностранные источники",
        "зарубежные источники": "Зарубежные источники",
        "Зарубежные источники": "Зарубежные источники",
    }
    for text, expected in new_exact_cases.items():
        actual = canonical_reference_block_heading_text(text)
        if actual != expected:
            return False, f"Zone J subheading not detected: {text!r} -> {actual!r}, expected {expected!r}"

    # Integration: process a document with several Zone J subheadings interspersed with entries
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Краткий текст.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("Нормативно-правовые акты")
    doc.add_paragraph("Федеральный закон от 01.01.2020 № 1-ФЗ.")
    doc.add_paragraph("Судебная практика")
    doc.add_paragraph("Решение Арбитражного суда от 01.03.2021.")
    doc.add_paragraph("Зарубежные источники")
    doc.add_paragraph("Smith J. Global Economy. London, 2022.")
    doc.add_paragraph("Интернет-источники")
    doc.add_paragraph("Официальный сайт ЦБ РФ. URL: https://cbr.ru")

    with tempfile.TemporaryDirectory() as tmp:
        input_path = Path(tmp) / "in.docx"
        output_path = Path(tmp) / "out.docx"
        doc.save(str(input_path))
        process_document(input_path, output_path)
        out_doc = Document(str(output_path))

    texts = _paragraph_texts(out_doc)
    try:
        refs_idx = texts.index("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    except ValueError:
        return False, "references heading missing after formatting"

    nonempty = [t for t in texts[refs_idx + 1:] if t]

    subheadings = {"Нормативно-правовые акты", "Судебная практика", "Зарубежные источники", "Интернет-источники"}
    for sh in subheadings:
        if sh not in nonempty:
            return False, f"Zone J subheading missing from output: {sh!r}"

    numbered_entries = [t for t in nonempty if t[0].isdigit() and ". " in t[:4]]
    for entry in numbered_entries:
        if entry[:2] in {"Но", "Су", "За", "Ин"}:
            return False, f"subheading was numbered as a source entry: {entry!r}"

    for i, entry in enumerate(numbered_entries):
        expected_prefix = f"{i + 1}. "
        if not entry.startswith(expected_prefix):
            return False, f"sequential numbering broken: {entry!r} expected prefix {expected_prefix!r}"

    return True, "Zone J subheadings are unnumbered; entries are sequentially numbered"


def test_zone_k_duplicate_reference_numbering() -> tuple[bool, str]:
    """Zone K: duplicate leading reference numbers are stripped before sequential renumbering."""
    # Unit: strip_leading_reference_number removes repeated leading prefixes
    unit_cases = [
        ("1. 1. Федеральный закон от 01.01.2020", "Федеральный закон от 01.01.2020"),
        ("3. 3. Книга по экономике", "Книга по экономике"),
        ("12. 12. Статья в журнале", "Статья в журнале"),
        ("1. Иванов И. И. Источник.", "Иванов И. И. Источник."),  # single prefix also stripped
        ("Иванов И. И. Источник.", "Иванов И. И. Источник."),    # no prefix unchanged
        ("1) 1) Источник номер один", "Источник номер один"),     # paren-style double
        ("[1] [1] Источник номер один", "Источник номер один"),   # bracket-style double
    ]
    for text, expected in unit_cases:
        actual = strip_leading_reference_number(text)
        if actual != expected:
            return False, f"strip_leading_reference_number({text!r}) = {actual!r}, expected {expected!r}"

    # Unit: legal content numbers inside the title are NOT stripped
    safe_cases = [
        "Гражданский кодекс Российской Федерации. Часть 1. Статья 10.",
        "ГОСТ Р 7.0.5-2008. Библиографическая ссылка.",
        "https://example.com/path/1/resource",
    ]
    for text in safe_cases:
        actual = strip_leading_reference_number(text)
        if actual != text:
            return False, f"safe content was mutated: {text!r} -> {actual!r}"

    # Integration: full process_document with double-prefixed entries
    doc = Document()
    doc.add_paragraph("ВВЕДЕНИЕ")
    doc.add_paragraph("Краткий текст.")
    doc.add_paragraph("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    doc.add_paragraph("1. 1. Федеральный закон от 01.01.2020 № 1-ФЗ.")
    doc.add_paragraph("3. 3. Книга по экономике.")
    doc.add_paragraph("Интернет-источники")
    doc.add_paragraph("12. 12. Официальный сайт ЦБ РФ. URL: https://cbr.ru")
    doc.add_paragraph("Нормальный источник без дублирования.")

    with tempfile.TemporaryDirectory() as tmp:
        input_path = Path(tmp) / "in.docx"
        output_path = Path(tmp) / "out.docx"
        doc.save(str(input_path))
        process_document(input_path, output_path)
        out_doc = Document(str(output_path))

    texts = _paragraph_texts(out_doc)
    try:
        refs_idx = texts.index("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    except ValueError:
        return False, "references heading missing after formatting"

    nonempty = [t for t in texts[refs_idx + 1:] if t]

    # Zone J subheading must remain unnumbered
    if "Интернет-источники" not in nonempty:
        return False, "Zone J subheading Интернет-источники missing"
    sh_idx = nonempty.index("Интернет-источники")
    if nonempty[sh_idx][0].isdigit():
        return False, f"subheading was numbered: {nonempty[sh_idx]!r}"

    # All numbered entries must be sequentially 1./2./3./4. with no double prefix
    numbered = [t for t in nonempty if t and t[0].isdigit()]
    for i, entry in enumerate(numbered):
        expected_prefix = f"{i + 1}. "
        if not entry.startswith(expected_prefix):
            return False, f"sequential numbering broken at position {i+1}: {entry!r}"
        # Must not contain a second leading number immediately after the prefix
        rest = entry[len(expected_prefix):]
        if rest and rest[0].isdigit() and re.match(r"^\d+[.)]\s", rest):
            return False, f"duplicate number survived in entry: {entry!r}"

    return True, "duplicate leading reference numbers stripped; sequential numbering correct"


def main() -> int:
    tests = [
        ("reference subheading spacing", test_reference_subheading_spacing),
        ("strict reference subheading detection", test_reference_subheading_detection_is_strict),
        ("flexible reference block heading detection", test_reference_block_heading_detection_is_flexible),
        ("short centered bold reference heading", test_reference_short_centered_bold_heading_stays_unnumbered),
        ("numbered reference entries", test_numbered_reference_entries_are_not_headings),
        ("old reference numbering cleanup", test_reference_old_numbering_cleanup),
        ("Zone J new reference subheadings", test_zone_j_new_reference_subheadings),
        ("Zone K duplicate reference numbering", test_zone_k_duplicate_reference_numbering),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        status = "PASS" if ok else "FAIL"
        print(f"[{status}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
