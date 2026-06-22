"""
Stage A — rendered table layout acceptance gate tests.

These tests pin the behaviour added by the "gate rendered table layout defects"
batch: a structured, severity-bearing acceptance check over the rendered PDF so
that visibly broken table layout can no longer be reported as clean/GO.

They also pin the Stage 0 conservative-table-mode switch in formatter_service.

Run: python3 tests/test_rendered_table_acceptance_gate.py
"""

from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine  # noqa: E402
from guides.coursework_kfu_2025.rendered_table_validation import (  # noqa: E402
    RenderedTableIdentity,
    TableLayoutBlocker,
    evaluate_table_layout_acceptance,
)
import guides.coursework_kfu_2025.formatter_service as fs  # noqa: E402


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


# --------------------------------------------------------------------------- #
# Fixture builders
# --------------------------------------------------------------------------- #

def _line(text: str, page: int, top: float, bottom: float | None = None) -> PdfLine:
    return PdfLine(text=text, page_num=page, top=top, bottom=bottom if bottom is not None else top + 12.0)


def _identity(
    *,
    table_index: int,
    body_order_index: int,
    caption_num: str | None,
    header: tuple[str, ...] = (),
    numeric: str | None = None,
    rows: tuple[str, ...] = (),
    preceding_marker: str | None = None,
    following_marker: str | None = None,
) -> RenderedTableIdentity:
    return RenderedTableIdentity(
        table_index=table_index,
        body_order_index=body_order_index,
        caption_num=caption_num,
        preceding_marker=preceding_marker,
        following_marker=following_marker,
        header_fingerprint=header,
        numeric_row_fingerprint=numeric,
        row_fingerprints=rows,
    )


def _fails(blockers: list[TableLayoutBlocker]) -> list[TableLayoutBlocker]:
    return [b for b in blockers if b.severity == "fail"]


def _types(blockers: list[TableLayoutBlocker]) -> set[str]:
    return {b.blocker_type for b in blockers}


_H_214 = "канал отчётности периодичность объект аудита эскалации follow-up"
_R_214A = "наблюдательный совет по годовому плану процессы рекомендации"
_R_214B = "правление по запросу исполнение отчёты"
_R_214C = "внешний аудитор по требованиям процедуры предоставление"


# --------------------------------------------------------------------------- #
# 1. same-page continuation marker => fail
# --------------------------------------------------------------------------- #

def test_same_page_continuation_marker_is_fail() -> tuple[bool, str]:
    pdf_lines = [
        _line("Таблица 2.1.4", 32, 90),
        _line("СВА: подотчетность и взаимодействие", 32, 110),
        _line(_H_214, 32, 130),
        _line("1 2 3 4", 32, 150),
        _line(_R_214A, 32, 170),
        _line(_R_214B, 32, 200),
        _line("Продолжение таблицы 2.1.4", 32, 300),
        _line(_H_214, 32, 320),
        _line("1 2 3 4", 32, 340),
        _line(_R_214C, 32, 360),
    ]
    identity = _identity(
        table_index=4, body_order_index=4, caption_num="2.1.4",
        header=(_H_214,), numeric="1 2 3 4",
        rows=(_H_214, "1 2 3 4", _R_214A, _R_214B, _R_214C),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    same_page = [b for b in _fails(blockers) if b.blocker_type == "same_page_continuation"]
    if not same_page:
        return _result(False, f"expected same_page_continuation fail, got {_types(blockers)!r}")
    if same_page[0].page != 32 or same_page[0].table_num != "2.1.4":
        return _result(False, f"wrong page/table on blocker: {same_page[0]!r}")
    return _result(True, "same-page continuation marker is a fail blocker")


# --------------------------------------------------------------------------- #
# 2. next-page continuation marker => pass (no fail)
# --------------------------------------------------------------------------- #

def test_next_page_continuation_marker_is_pass() -> tuple[bool, str]:
    pdf_lines = [
        _line("Таблица 2.1.4", 32, 90),
        _line(_H_214, 32, 130),
        _line("1 2 3 4", 32, 150),
        _line(_R_214A, 32, 170),
        _line(_R_214B, 32, 600),
        # continuation correctly on the next page, at the top
        _line("Продолжение таблицы 2.1.4", 33, 70),
        _line(_H_214, 33, 95),
        _line("1 2 3 4", 33, 115),
        _line(_R_214C, 33, 140),
    ]
    identity = _identity(
        table_index=4, body_order_index=4, caption_num="2.1.4",
        header=(_H_214,), numeric="1 2 3 4",
        rows=(_H_214, "1 2 3 4", _R_214A, _R_214B, _R_214C),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    if _fails(blockers):
        return _result(False, f"valid next-page continuation produced fails: {_fails(blockers)!r}")
    return _result(True, "next-page continuation marker produces no fail blocker")


# --------------------------------------------------------------------------- #
# 3. orphaned header row (header on its own page, data on next) => fail
# --------------------------------------------------------------------------- #

_H_231 = "роль ответственность частота артефакты kpi срок пилота"
_R_231A = "лидер совета повестка приоритизация ежемесячно план работ"
_R_231B = "куратор соответствие решениям ежеквартально резюме комитета"


def test_orphaned_header_row_is_fail() -> tuple[bool, str]:
    pdf_lines = [
        _line("Таблица 2.3.1", 43, 690),
        _line("Совет практик: регламент и метрики пилота", 43, 710),
        _line(_H_231, 43, 740),  # header alone at bottom of p43
        _line(_R_231A, 44, 90),  # all data on p44
        _line(_R_231B, 44, 140),
    ]
    identity = _identity(
        table_index=15, body_order_index=15, caption_num="2.3.1",
        header=(_H_231,),
        rows=(_H_231, _R_231A, _R_231B),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    orphan = [b for b in _fails(blockers) if b.blocker_type == "orphaned_header_row"]
    if not orphan:
        return _result(False, f"expected orphaned_header_row fail, got {_types(blockers)!r}")
    if orphan[0].page != 43:
        return _result(False, f"orphaned_header_row should flag p43, got {orphan[0].page}")
    return _result(True, "orphaned header row is a fail blocker")


def test_normal_table_start_with_data_is_not_orphan() -> tuple[bool, str]:
    # header AND a real data row on the same first page => not an orphan
    pdf_lines = [
        _line("Таблица 2.3.1", 43, 300),
        _line(_H_231, 43, 330),
        _line(_R_231A, 43, 360),
        _line(_R_231B, 44, 90),
    ]
    identity = _identity(
        table_index=15, body_order_index=15, caption_num="2.3.1",
        header=(_H_231,),
        rows=(_H_231, _R_231A, _R_231B),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    if [b for b in _fails(blockers) if b.blocker_type == "orphaned_header_row"]:
        return _result(False, "false orphaned_header_row on a table with data on its first page")
    return _result(True, "table with data on first page is not an orphan")


# --------------------------------------------------------------------------- #
# 4. fragment grid mismatch => fail
# --------------------------------------------------------------------------- #

def _add_grid_table(doc: Document, widths: list[int], rows: int = 2) -> None:
    n = len(widths)
    table = doc.add_table(rows=rows, cols=n)
    tbl = table._tbl
    grid = tbl.find(qn("w:tblGrid"))
    for col_el, w in zip(grid.findall(qn("w:gridCol")), widths):
        col_el.set(qn("w:w"), str(w))


def test_fragment_grid_mismatch_is_fail() -> tuple[bool, str]:
    doc = Document()
    _add_grid_table(doc, [2000, 2000, 2000, 2000, 2000])  # 5 cols
    doc.add_paragraph("")  # only a blank between the two fragments
    _add_grid_table(doc, [2500, 2500, 2500, 2500])        # 4 cols -> mismatch
    blockers = evaluate_table_layout_acceptance([], [], doc=doc)
    mismatch = [b for b in _fails(blockers) if b.blocker_type == "fragment_grid_mismatch"]
    if not mismatch:
        return _result(False, f"expected fragment_grid_mismatch fail, got {_types(blockers)!r}")
    return _result(True, "fragment grid mismatch is a fail blocker")


def test_matching_adjacent_grids_are_not_mismatch() -> tuple[bool, str]:
    doc = Document()
    _add_grid_table(doc, [2000, 2000, 2000])
    doc.add_paragraph("")
    _add_grid_table(doc, [2000, 2000, 2000])
    blockers = evaluate_table_layout_acceptance([], [], doc=doc)
    if [b for b in blockers if b.blocker_type == "fragment_grid_mismatch"]:
        return _result(False, "false fragment_grid_mismatch on identical adjacent grids")
    return _result(True, "identical adjacent grids are not a mismatch")


def test_adjacent_fragment_width_drift_is_review() -> tuple[bool, str]:
    # same column count but the fragments' column widths drift -> Rule 8 says a
    # split must preserve widths; flag for review (Bondarev 2.3.1/2.3.3 class).
    doc = Document()
    _add_grid_table(doc, [3000, 1500, 1500])
    doc.add_paragraph("")
    _add_grid_table(doc, [1500, 1500, 3000])
    blockers = evaluate_table_layout_acceptance([], [], doc=doc)
    drift = [b for b in blockers if b.blocker_type == "fragment_grid_mismatch"]
    if not drift:
        return _result(False, "expected fragment_grid_mismatch (width drift) blocker")
    if drift[0].severity != "needs_human_review":
        return _result(False, f"width drift should be needs_human_review, got {drift[0].severity}")
    return _result(True, "adjacent fragment width drift flagged for review")


# --------------------------------------------------------------------------- #
# 5. appendix label not on a new page => fail; on a new page => pass
# --------------------------------------------------------------------------- #

def test_appendix_label_mid_page_is_fail() -> tuple[bool, str]:
    pdf_lines = [
        _line("Этап Описание Документы Порог кворум Метрика", 62, 90),
        _line("Подготовка формирование повестки", 62, 130),
        _line("Отчётность публикация итогов", 62, 300),
        _line("ПРИЛОЖЕНИЕ Б", 62, 400),  # mid-page, content above it
        _line("Сьел дедал", 62, 440),
    ]
    blockers = evaluate_table_layout_acceptance(pdf_lines, [])
    appendix = [b for b in _fails(blockers) if b.blocker_type == "appendix_label_not_on_new_page"]
    if not appendix:
        return _result(False, f"expected appendix_label_not_on_new_page fail, got {_types(blockers)!r}")
    if appendix[0].page != 62:
        return _result(False, f"appendix blocker should flag p62, got {appendix[0].page}")
    return _result(True, "mid-page appendix label is a fail blocker")


def test_appendix_label_on_new_page_is_pass() -> tuple[bool, str]:
    pdf_lines = [
        _line("ПРИЛОЖЕНИЯ", 61, 50),       # parent section heading is allowed above
        _line("ПРИЛОЖЕНИЕ А", 61, 75),     # label at the top of a fresh page
        _line("Таблица А ...", 61, 110),
        _line("Этап Описание Документы", 61, 140),
    ]
    blockers = evaluate_table_layout_acceptance(pdf_lines, [])
    if [b for b in _fails(blockers) if b.blocker_type == "appendix_label_not_on_new_page"]:
        return _result(False, "false appendix_label_not_on_new_page on a label at page top")
    return _result(True, "appendix label on a new page produces no blocker")


# --------------------------------------------------------------------------- #
# 5b. table crosses pages without a continuation marker => fail (Demo 1.1.3)
# --------------------------------------------------------------------------- #

_H_113 = "группа признаков пример дефекта почему плохо статус"
_R_113A = "кейс один длинный текст таблицы внезапно очень широкий"
_R_113B = "кейс два другой длинный текст таблицы снова очень широкий"


def test_table_crosses_pages_without_marker_is_fail() -> tuple[bool, str]:
    pdf_lines = [
        _line("Таблица 1.1.3", 5, 500),
        _line(_H_113, 5, 520),
        _line(_R_113A, 5, 560),
        _line(_R_113B, 6, 90),  # data continues on next page, no marker anywhere
    ]
    identity = _identity(
        table_index=2, body_order_index=2, caption_num="1.1.3",
        header=(_H_113,),
        rows=(_H_113, _R_113A, _R_113B),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    cross = [b for b in _fails(blockers) if b.blocker_type == "single_table_crosses_pages_without_marker"]
    if not cross:
        return _result(False, f"expected cross-page-without-marker fail, got {_types(blockers)!r}")
    return _result(True, "table crossing pages without a marker is a fail blocker")


def test_table_crosses_pages_with_marker_is_ok() -> tuple[bool, str]:
    pdf_lines = [
        _line("Таблица 1.1.3", 5, 500),
        _line(_H_113, 5, 520),
        _line(_R_113A, 5, 560),
        _line("Продолжение таблицы 1.1.3", 6, 70),
        _line(_R_113B, 6, 110),
    ]
    identity = _identity(
        table_index=2, body_order_index=2, caption_num="1.1.3",
        header=(_H_113,),
        rows=(_H_113, _R_113A, _R_113B),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    if [b for b in _fails(blockers) if b.blocker_type == "single_table_crosses_pages_without_marker"]:
        return _result(False, "valid marked continuation flagged as cross-page-without-marker")
    return _result(True, "marked cross-page continuation is accepted")


# --------------------------------------------------------------------------- #
# 6. clean document => no fail blockers (Rybakov guard)
# --------------------------------------------------------------------------- #

def test_clean_single_page_table_has_no_fail() -> tuple[bool, str]:
    pdf_lines = [
        _line("Таблица 1.1.1", 8, 90),
        _line("Заголовок один два три", 8, 120),
        _line("строка данных альфа бета гамма", 8, 150),
        _line("строка данных дельта эпсилон дзета", 8, 180),
    ]
    identity = _identity(
        table_index=0, body_order_index=0, caption_num="1.1.1",
        header=("заголовок один два три",),
        rows=("заголовок один два три",
              "строка данных альфа бета гамма",
              "строка данных дельта эпсилон дзета"),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    if _fails(blockers):
        return _result(False, f"clean single-page table produced fails: {_fails(blockers)!r}")
    return _result(True, "clean single-page table has no fail blockers")


# --------------------------------------------------------------------------- #
# 7. Stage 0 — conservative table mode default
# --------------------------------------------------------------------------- #

def test_conservative_mode_is_default_off() -> tuple[bool, str]:
    saved = os.environ.pop("KPFU_RENDERED_TABLE_CONTINUATION", None)
    try:
        if fs._rendered_table_continuation_enabled():
            return _result(False, "rendered table continuation must be OFF by default")
        os.environ["KPFU_RENDERED_TABLE_CONTINUATION"] = "1"
        if not fs._rendered_table_continuation_enabled():
            return _result(False, "rendered table continuation must turn ON when flag set")
        os.environ["KPFU_RENDERED_TABLE_CONTINUATION"] = "0"
        if fs._rendered_table_continuation_enabled():
            return _result(False, "explicit 0 must keep rendered table continuation OFF")
    finally:
        os.environ.pop("KPFU_RENDERED_TABLE_CONTINUATION", None)
        if saved is not None:
            os.environ["KPFU_RENDERED_TABLE_CONTINUATION"] = saved
    return _result(True, "conservative table mode is default-off and flag-controlled")


# --------------------------------------------------------------------------- #
# 8. squeeze heuristic surfaces a non-clean signal (needs_human_review)
# --------------------------------------------------------------------------- #

def test_severe_squeeze_is_flagged_for_review() -> tuple[bool, str]:
    # "6 месяцев" wrapping as 6 / месяце / в on every row produces many
    # ultra-short fragment lines inside the table region.
    pdf_lines = [_line("Таблица 2.3.1", 44, 60)]
    top = 80.0
    for _ in range(6):
        pdf_lines.append(_line("6", 44, top)); top += 12
        pdf_lines.append(_line("месяце", 44, top)); top += 12
        pdf_lines.append(_line("в", 44, top)); top += 12
    identity = _identity(
        table_index=15, body_order_index=15, caption_num="2.3.1",
        header=("роль ответственность частота",),
        rows=("роль ответственность частота", "лидер совета повестка ежемесячно"),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    squeeze = [b for b in blockers if b.blocker_type == "cell_text_overflow_or_illegible_squeeze"]
    if not squeeze:
        return _result(False, f"expected squeeze review signal, got {_types(blockers)!r}")
    if squeeze[0].severity not in {"fail", "needs_human_review"}:
        return _result(False, f"squeeze severity unexpected: {squeeze[0].severity}")
    return _result(True, "severe squeeze surfaces a non-clean review signal")


# --------------------------------------------------------------------------- #
# B. same-page repeated header (no marker) => fail
# --------------------------------------------------------------------------- #

_H_REP = "роль ответственность частота артефакты"
_R_REPA = "лидер совета повестка приоритизация ежемесячно план работ"
_R_REPB = "куратор соответствие решениям ежеквартально резюме комитета"


def test_same_page_repeated_header_is_fail() -> tuple[bool, str]:
    pdf_lines = [
        _line("Таблица 2.3.1", 44, 60),
        _line(_H_REP, 44, 90),     # header occurrence 1
        _line(_R_REPA, 44, 140),
        _line(_H_REP, 44, 300),    # header occurrence 2 on the SAME page
        _line(_R_REPB, 44, 340),
    ]
    identity = _identity(
        table_index=18, body_order_index=18, caption_num="2.3.1",
        header=(_H_REP,), rows=(_H_REP, _R_REPA, _R_REPB),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    rep = [b for b in _fails(blockers) if b.blocker_type == "same_page_repeated_header"]
    if not rep:
        return _result(False, f"expected same_page_repeated_header fail, got {_types(blockers)!r}")
    if rep[0].page != 44:
        return _result(False, f"repeated header should flag p44, got {rep[0].page}")
    return _result(True, "same-page repeated header is a fail blocker")


def test_header_repeated_on_next_page_is_ok() -> tuple[bool, str]:
    pdf_lines = [
        _line("Таблица 2.3.1", 43, 60),
        _line(_H_REP, 43, 90),
        _line(_R_REPA, 43, 140),
        _line("Продолжение таблицы 2.3.1", 44, 60),
        _line(_H_REP, 44, 90),     # header repeats on a NEW page — allowed
        _line(_R_REPB, 44, 140),
    ]
    identity = _identity(
        table_index=18, body_order_index=18, caption_num="2.3.1",
        header=(_H_REP,), rows=(_H_REP, _R_REPA, _R_REPB),
    )
    blockers = evaluate_table_layout_acceptance(pdf_lines, [identity])
    if [b for b in blockers if b.blocker_type == "same_page_repeated_header"]:
        return _result(False, "header repeated on a continuation page wrongly flagged")
    return _result(True, "header repeated on a new page produces no blocker")


# --------------------------------------------------------------------------- #
# D. source-proven duplicated rows downgrade a hard fail to needs_human_review
# --------------------------------------------------------------------------- #

_DUP = "повторяющаяся содержательная строка таблицы один два"
_OTH = "иная содержательная строка данных три четыре"


def test_source_bad_duplication_downgrades_fail() -> tuple[bool, str]:
    pdf_lines = [
        _line("Таблица 1.2.1", 16, 80),
        _line("заголовок один два три", 16, 110),
        _line(_DUP, 16, 140),
        _line(_OTH, 16, 600),
        _line(_DUP, 17, 90),   # duplicated meaningful row continues to p17 (no marker)
    ]
    fmt = _identity(
        table_index=3, body_order_index=3, caption_num="1.2.1",
        header=("заголовок один два три",),
        rows=("заголовок один два три", _DUP, _DUP, _OTH),
    )
    src = _identity(
        table_index=3, body_order_index=3, caption_num="1.2.1",
        header=("заголовок один два три",),
        rows=("заголовок один два три", _DUP, _DUP, _OTH),
    )
    # without source: hard fail
    no_src = evaluate_table_layout_acceptance(pdf_lines, [fmt])
    if not [b for b in _fails(no_src) if b.blocker_type == "single_table_crosses_pages_without_marker"]:
        return _result(False, "expected a cross-page fail without source context")
    # with source proving duplication: downgraded to needs_human_review
    with_src = evaluate_table_layout_acceptance(pdf_lines, [fmt], source_identities=[src])
    cross = [b for b in with_src if b.blocker_type == "single_table_crosses_pages_without_marker"]
    if not cross:
        return _result(False, "cross-page blocker disappeared entirely (must stay visible)")
    if cross[0].severity != "needs_human_review" or not cross[0].evidence.get("source_bad"):
        return _result(False, f"source-bad fail not downgraded: {cross[0].severity} {cross[0].evidence}")
    if _fails(with_src):
        return _result(False, f"source-bad table must leave no hard fail: {_fails(with_src)!r}")
    return _result(True, "source-proven duplication downgrades fail to needs_human_review")


# --------------------------------------------------------------------------- #
# E. fragment grid mismatch is attributed to a table number / page
# --------------------------------------------------------------------------- #

def test_fragment_grid_mismatch_is_attributed() -> tuple[bool, str]:
    doc = Document()
    doc.add_paragraph("Таблица 5.5")
    _add_grid_table(doc, [2000, 2000, 2000, 2000, 2000])  # 5 cols
    doc.add_paragraph("")
    _add_grid_table(doc, [2500, 2500, 2500, 2500])        # 4 cols -> mismatch
    identities = [
        _identity(table_index=0, body_order_index=0, caption_num="5.5",
                  header=("a b c",), rows=("a b c",)),
        _identity(table_index=1, body_order_index=1, caption_num=None,
                  header=("a b c",), rows=("a b c",)),
    ]
    pdf_lines = [_line("Таблица 5.5", 3, 80)]
    blockers = evaluate_table_layout_acceptance(pdf_lines, identities, doc=doc)
    mm = [b for b in blockers if b.blocker_type == "fragment_grid_mismatch"]
    if not mm:
        return _result(False, "expected fragment_grid_mismatch")
    if mm[0].table_num != "5.5" or mm[0].page != 3:
        return _result(False, f"grid mismatch not attributed: table={mm[0].table_num} page={mm[0].page}")
    return _result(True, "fragment grid mismatch attributed to table/page")


def main() -> int:
    tests = [
        ("same-page continuation is fail", test_same_page_continuation_marker_is_fail),
        ("next-page continuation is pass", test_next_page_continuation_marker_is_pass),
        ("orphaned header row is fail", test_orphaned_header_row_is_fail),
        ("table with first-page data is not orphan", test_normal_table_start_with_data_is_not_orphan),
        ("fragment grid mismatch is fail", test_fragment_grid_mismatch_is_fail),
        ("matching grids are not mismatch", test_matching_adjacent_grids_are_not_mismatch),
        ("fragment width drift is review", test_adjacent_fragment_width_drift_is_review),
        ("cross-page without marker is fail", test_table_crosses_pages_without_marker_is_fail),
        ("cross-page with marker is ok", test_table_crosses_pages_with_marker_is_ok),
        ("same-page repeated header is fail", test_same_page_repeated_header_is_fail),
        ("header repeated on next page ok", test_header_repeated_on_next_page_is_ok),
        ("source-bad duplication downgrades", test_source_bad_duplication_downgrades_fail),
        ("grid mismatch attributed", test_fragment_grid_mismatch_is_attributed),
        ("appendix label mid-page is fail", test_appendix_label_mid_page_is_fail),
        ("appendix label on new page is pass", test_appendix_label_on_new_page_is_pass),
        ("clean single-page table has no fail", test_clean_single_page_table_has_no_fail),
        ("conservative mode default off", test_conservative_mode_is_default_off),
        ("severe squeeze flagged for review", test_severe_squeeze_is_flagged_for_review),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        status = "PASS" if ok else "FAIL"
        print(f"[{status}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
