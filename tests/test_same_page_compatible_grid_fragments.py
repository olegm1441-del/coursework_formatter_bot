from __future__ import annotations

import sys
import tempfile
from pathlib import Path
from types import SimpleNamespace

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import guides.coursework_kfu_2025.table_continuation as tc


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def _set_grid(table, widths: list[str]) -> None:
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), width)
        grid.append(col)
    tbl.insert(0, grid)


def _set_border(table, value: str = "single") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), value)
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    tbl_pr.append(borders)


def _set_margins(table, value: str = "15") -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is not None:
        tbl_pr.remove(margins)
    margins = OxmlElement("w:tblCellMar")
    for name in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{name}")
        margin.set(qn("w:w"), value)
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tbl_pr.append(margins)


def _set_cell_widths(table, widths: list[str]) -> None:
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), widths[idx])
            tc_w.set(qn("w:type"), "dxa")


def _fill_row(row, values: list[str]) -> None:
    for idx, value in enumerate(values):
        row.cells[idx].text = value


def _add_table(doc: Document, rows: list[list[str]], grid: list[str]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row, values in zip(table.rows, rows):
        _fill_row(row, values)
    _set_grid(table, grid)
    _set_cell_widths(table, grid)
    _set_border(table)
    _set_margins(table)
    return table


def _make_compatible_doc(
    *,
    source_numeric: bool = False,
    formatted_numeric: bool = True,
    source_bad_duplicate: bool = False,
    source_note: bool = True,
    merged_second: bool = False,
) -> tuple[Document, Document]:
    header = ["Параметр", "2021", "2022", "2023"]
    numeric = ["1", "2", "3", "4"]
    first_data = [
        ["Выручка", "10", "12", "15"],
        ["Себестоимость", "7", "8", "9"],
        ["Валовая прибыль", "3", "4", "6"],
    ]
    second_data = [
        ["Коммерческие расходы", "1", "1", "2"],
        ["Чистая прибыль", "2", "3", "4"],
    ]
    if source_bad_duplicate:
        second_data = [first_data[-1], second_data[-1]]

    source = Document()
    source.add_paragraph("Таблица 1.3.1")
    source.add_paragraph("Динамика финансовых показателей")
    source_first = [header]
    if source_numeric:
        source_first.append(numeric)
    source_first.extend(first_data)
    _add_table(source, source_first, ["2000", "2639", "2112", "2877"])
    source.add_paragraph("Продолжение табл. 1.3.1")
    source_second = [header]
    if source_numeric:
        source_second.append(numeric)
    source_second.extend(second_data)
    source_second_table = _add_table(source, source_second, ["1814", "2670", "2390", "2754"])
    if merged_second:
        tc_pr = source_second_table.rows[0].cells[0]._tc.get_or_add_tcPr()
        grid_span = OxmlElement("w:gridSpan")
        grid_span.set(qn("w:val"), "2")
        tc_pr.append(grid_span)
    if source_note:
        source.add_paragraph("Источник: составлено автором.")

    formatted = Document()
    formatted.add_paragraph("Таблица 1.3.1")
    formatted.add_paragraph("Динамика финансовых показателей")
    formatted_first = [header]
    formatted_second_rows = [header]
    if formatted_numeric:
        formatted_first.append(numeric)
        formatted_second_rows.append(numeric)
    formatted_first.extend(first_data)
    formatted_second_rows.extend(second_data)
    _add_table(formatted, formatted_first, ["2000", "2639", "2112", "2877"])
    formatted_second = _add_table(formatted, formatted_second_rows, ["1814", "2670", "2390", "2754"])
    if merged_second:
        tc_pr = formatted_second.rows[0].cells[0]._tc.get_or_add_tcPr()
        grid_span = OxmlElement("w:gridSpan")
        grid_span.set(qn("w:val"), "2")
        tc_pr.append(grid_span)
    if source_note:
        formatted.add_paragraph("Источник: составлено автором.")
    return source, formatted


def _save_pair(source: Document, formatted: Document, tmp: str) -> tuple[Path, Path]:
    source_path = Path(tmp) / "source.docx"
    formatted_path = Path(tmp) / "formatted.docx"
    source.save(source_path)
    formatted.save(formatted_path)
    return source_path, formatted_path


def _rows(doc: Document, table_index: int = 0) -> list[list[str]]:
    return [
        [" ".join(cell.text.split()) for cell in row.cells]
        for row in doc.tables[table_index].rows
    ]


def _row_widths(doc: Document, table_index: int = 0) -> list[list[str | None]]:
    widths: list[list[str | None]] = []
    for row in doc.tables[table_index].rows:
        row_widths: list[str | None] = []
        for cell in row.cells:
            tc_pr = cell._tc.tcPr
            tc_w = tc_pr.find(qn("w:tcW")) if tc_pr is not None else None
            row_widths.append(tc_w.get(qn("w:w")) if tc_w is not None else None)
        widths.append(row_widths)
    return widths


def _row_has_tbl_header(row) -> bool:
    tr_pr = row._tr.trPr
    return tr_pr is not None and tr_pr.find(qn("w:tblHeader")) is not None


def _row_has_table_cell_pagination_flags(row) -> bool:
    for paragraph in row._tr.findall(".//" + qn("w:p")):
        p_pr = paragraph.find(qn("w:pPr"))
        if p_pr is None:
            continue
        if p_pr.find(qn("w:pageBreakBefore")) is not None:
            return True
        if p_pr.find(qn("w:keepNext")) is not None:
            return True
    return False


def _source_note_after_only_table(doc: Document) -> bool:
    children = list(doc.element.body)
    table_positions = [idx for idx, child in enumerate(children) if child.tag == qn("w:tbl")]
    if len(table_positions) != 1:
        return False
    table_pos = table_positions[0]
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t"))).strip()
        if text.startswith("Источник:"):
            return idx > table_pos
    return False


def _same_page_violation():
    return SimpleNamespace(
        table_num="1.3.1",
        table_index=0,
        page=22,
        violation_type="same_page_repeated_fragment",
        confidence="high",
        evidence={"following_table_index": 1},
    )


def _run_repair(
    source: Document,
    formatted: Document,
    *,
    after_target_remains: bool = False,
    regressions: list | None = None,
) -> tuple[int, Document]:
    with tempfile.TemporaryDirectory() as tmp:
        source_path, formatted_path = _save_pair(source, formatted, tmp)
        old_rendered = tc._rendered_continuation_violations_for_docx
        old_marker = tc._same_page_continuation_marker_violations_for_docx
        old_regressions = tc._rendered_continuation_deletion_regressions
        calls = {"n": 0}
        try:
            def fake_rendered(_path):
                calls["n"] += 1
                if calls["n"] == 1 or after_target_remains:
                    return [_same_page_violation()]
                return []

            tc._rendered_continuation_violations_for_docx = fake_rendered
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: []
            tc._rendered_continuation_deletion_regressions = lambda _path: regressions or []
            changed = tc.normalize_compatible_grid_same_page_repeated_fragments_inplace(
                formatted_path,
                source_docx_path=source_path,
            )
        finally:
            tc._rendered_continuation_violations_for_docx = old_rendered
            tc._same_page_continuation_marker_violations_for_docx = old_marker
            tc._rendered_continuation_deletion_regressions = old_regressions
        return changed, Document(str(formatted_path))


def test_compatible_grid_same_page_fragments_merge() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc()
    changed, reread = _run_repair(source, formatted)
    rows = _rows(reread)
    actual_data = [row[0] for row in rows[1:] if row[0] != "1"]
    expected_data = [
        "Выручка",
        "Себестоимость",
        "Валовая прибыль",
        "Коммерческие расходы",
        "Чистая прибыль",
    ]
    if changed != 1:
        return _result(False, f"expected one compatible-grid repair, got {changed}")
    if len(reread.tables) != 1:
        return _result(False, f"expected one survivor table, got {len(reread.tables)}")
    if actual_data != expected_data:
        return _result(False, f"data rows not preserved in order: {actual_data!r}")
    return _result(True, "compatible-grid same-page fragments merge")


def test_appended_rows_are_adapted_to_survivor_grid() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc()
    changed, reread = _run_repair(source, formatted)
    widths = _row_widths(reread)
    expected = ["2000", "2639", "2112", "2877"]

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if widths[-1] != expected or widths[-2] != expected:
        return _result(False, f"appended rows were not adapted to survivor widths: {widths!r}")
    return _result(True, "appended rows use survivor table widths")


def test_duplicate_header_and_generated_numeric_from_second_fragment_are_removed() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc()
    changed, reread = _run_repair(source, formatted)
    row_texts = [" | ".join(row) for row in _rows(reread)]
    numeric = "1 | 2 | 3 | 4"

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if row_texts.count("Параметр | 2021 | 2022 | 2023") != 1:
        return _result(False, f"duplicate second-fragment header remained: {row_texts!r}")
    if row_texts.count(numeric) != 1:
        return _result(False, f"duplicate generated numeric row remained: {row_texts!r}")
    return _result(True, "duplicate header and generated numeric row removed")


def test_same_page_merge_clears_repeat_header_metadata() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc()
    changed, reread = _run_repair(source, formatted)
    rows = _rows(reread)

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if rows[0] != ["Параметр", "2021", "2022", "2023"]:
        return _result(False, f"top semantic header was not preserved: {rows[0]!r}")
    if rows[1] != ["1", "2", "3", "4"]:
        return _result(False, f"top numeric row was not preserved: {rows[1]!r}")
    table = reread.tables[0]
    if any(_row_has_tbl_header(row) for row in table.rows):
        return _result(False, "repeat-header metadata remained on merged same-page table")
    if any(_row_has_table_cell_pagination_flags(row) for row in table.rows):
        return _result(False, "table-cell pagination metadata remained on merged same-page table")
    return _result(True, "same-page merge clears repeat-header/pagination metadata")


def test_source_numeric_rows_skip_compatible_repair() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc(source_numeric=True)
    changed, reread = _run_repair(source, formatted)

    if changed != 0:
        return _result(False, f"source-proven numeric rows should skip repair, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, f"skipped source-numeric repair should preserve both tables, got {len(reread.tables)}")
    return _result(True, "source-proven numeric rows are not treated as generated")


def test_source_note_remains_after_merged_table() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc()
    changed, reread = _run_repair(source, formatted)

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if not _source_note_after_only_table(reread):
        return _result(False, "source note is not after the merged table")
    return _result(True, "source note remains after merged compatible-grid table")


def test_source_bad_duplicate_meaningful_rows_skip() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc(source_bad_duplicate=True)
    changed, reread = _run_repair(source, formatted)

    if changed != 0:
        return _result(False, f"source-bad meaningful duplicate should skip, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, f"source-bad skip should preserve both tables, got {len(reread.tables)}")
    return _result(True, "source-bad duplicate meaningful rows are skipped")


def test_no_numeric_row_variant_skips_for_this_patch() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc(formatted_numeric=False)
    changed, reread = _run_repair(source, formatted)

    if changed != 0:
        return _result(False, f"no-numeric same-page variant is out of scope, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, "no-numeric skip should preserve both fragments")
    return _result(True, "no-numeric same-page variant is skipped by this bounded patch")


def test_single_physical_cross_page_case_is_skipped() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc()
    while len(formatted.tables) > 1:
        formatted.tables[-1]._tbl.getparent().remove(formatted.tables[-1]._tbl)
    changed, reread = _run_repair(source, formatted)

    if changed != 0:
        return _result(False, f"single physical table must not be repaired here, got {changed}")
    if len(reread.tables) != 1:
        return _result(False, "single physical table shape changed")
    return _result(True, "single-physical cross-page class is skipped")


def test_merged_cell_complexity_skips() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc(merged_second=True)
    changed, reread = _run_repair(source, formatted)

    if changed != 0:
        return _result(False, f"merged-cell complexity should skip, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, "merged-cell skip did not preserve original fragments")
    return _result(True, "merged-cell complexity is skipped")


def test_rollback_preserves_original_when_validation_worsens() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc()
    changed, reread = _run_repair(source, formatted, after_target_remains=True)

    if changed != 0:
        return _result(False, f"failed validation should roll back, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, "rollback did not preserve original fragments")
    return _result(True, "failed post-render validation rolls back compatible-grid repair")


def test_same_page_compatible_merge_does_not_insert_marker_or_page_break() -> tuple[bool, str]:
    source, formatted = _make_compatible_doc()
    changed, reread = _run_repair(source, formatted)
    caption = next(p for p in reread.paragraphs if p.text == "Таблица 1.3.1")
    p_pr = caption._element.find(qn("w:pPr"))
    page_break = p_pr.find(qn("w:pageBreakBefore")) if p_pr is not None else None

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if page_break is not None:
        return _result(False, "same-page compatible merge must not insert pageBreakBefore")
    if any("Продолжение таблицы 1.3.1" in paragraph.text for paragraph in reread.paragraphs):
        return _result(False, "same-page compatible merge must not insert continuation marker")
    return _result(True, "compatible same-page merge does not insert marker or page break")


def main() -> int:
    tests = [
        ("compatible-grid merge", test_compatible_grid_same_page_fragments_merge),
        ("adapt survivor grid", test_appended_rows_are_adapted_to_survivor_grid),
        ("remove duplicate header/numeric", test_duplicate_header_and_generated_numeric_from_second_fragment_are_removed),
        ("clear repeat metadata", test_same_page_merge_clears_repeat_header_metadata),
        ("preserve source numeric by skip", test_source_numeric_rows_skip_compatible_repair),
        ("source note placement", test_source_note_remains_after_merged_table),
        ("source-bad duplicate skip", test_source_bad_duplicate_meaningful_rows_skip),
        ("no numeric variant skip", test_no_numeric_row_variant_skips_for_this_patch),
        ("single physical cross-page skip", test_single_physical_cross_page_case_is_skipped),
        ("merged-cell skip", test_merged_cell_complexity_skips),
        ("rollback on validation failure", test_rollback_preserves_original_when_validation_worsens),
        ("no marker/page break", test_same_page_compatible_merge_does_not_insert_marker_or_page_break),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        status = "PASS" if ok else "FAIL"
        print(f"[{status}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
