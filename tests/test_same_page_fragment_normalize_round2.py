from __future__ import annotations

import sys
import tempfile
from pathlib import Path
from types import SimpleNamespace

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import guides.coursework_kfu_2025.table_continuation as tc


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def _set_grid(table, widths: list[str]) -> None:
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), width)
        grid.append(col)
    tbl.insert(0, grid)


def _set_border(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    tbl_pr.append(borders)


def _set_margins(table, value: str = "113") -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is not None:
        tbl_pr.remove(margins)
    margins = OxmlElement("w:tblCellMar")
    for name in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{name}")
        margin.set(qn("w:w"), "15" if name in {"top", "bottom"} else value)
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tbl_pr.append(margins)


def _set_cell_widths(table, widths: list[str]) -> None:
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), widths[idx])
            tc_w.set(qn("w:type"), "dxa")


def _add_table(doc: Document, rows: list[list[str]], grid: list[str]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
    _set_grid(table, grid)
    _set_cell_widths(table, grid)
    _set_border(table)
    _set_margins(table)
    return table


def _rows(doc: Document, table_index: int = 0) -> list[list[str]]:
    return [
        [" ".join(cell.text.split()) for cell in row.cells]
        for row in doc.tables[table_index].rows
    ]


def _numeric_row_count(rows: list[list[str]]) -> int:
    return sum(
        1
        for row in rows
        if len(row) >= 2 and row == [str(idx) for idx in range(1, len(row) + 1)]
    )


def _table_texts(doc: Document) -> list[str]:
    return [
        " / ".join(" | ".join(row) for row in _rows(doc, table_idx))
        for table_idx in range(len(doc.tables))
    ]


def _all_table_text(doc: Document) -> str:
    return " / ".join(_table_texts(doc))


def _same_page_violation(table_num: str, table_index: int, following_table_index: int, *, numeric_count: int = 2):
    return SimpleNamespace(
        table_num=table_num,
        table_index=table_index,
        page=34,
        violation_type="same_page_repeated_fragment",
        confidence="high",
        evidence={
            "following_table_index": following_table_index,
            "repeated_header": True,
            "repeated_numeric_row_count": numeric_count,
            "adjacent_fragment_proof": "following_row_on_caption_page",
        },
    )


def _make_two_fragment_table_docs(table_num: str, rows_a: list[list[str]], rows_b: list[list[str]], *, numeric: bool = True):
    header = rows_a[0]
    number_row = [str(idx) for idx in range(1, len(header) + 1)]
    base_a = ["1700", "2600", "2500", "2800", "1450", "1050"]
    base_b = ["1650", "2700", "2400", "2850", "1500", "1000"]
    grid_a = base_a[: len(header)]
    grid_b = base_b[: len(header)]

    source = Document()
    source.add_paragraph(f"Таблица {table_num}")
    source.add_paragraph(f"Название таблицы {table_num}")
    _add_table(source, rows_a, grid_a)
    source.add_paragraph(f"Продолжение табл. {table_num}")
    _add_table(source, rows_b, grid_b)
    source.add_paragraph("Источник: составлено автором.")

    formatted = Document()
    formatted.add_paragraph(f"Таблица {table_num}")
    formatted.add_paragraph(f"Название таблицы {table_num}")
    first_rows = [header]
    second_rows = [header]
    if numeric:
        first_rows.append(number_row)
        second_rows.append(number_row)
    first_rows.extend(rows_a[1:])
    second_rows.extend(rows_b[1:])
    _add_table(formatted, first_rows, grid_a)
    _add_table(formatted, second_rows, grid_b)
    formatted.add_paragraph("Источник: составлено автором.")
    return source, formatted


def _append_doc(dst: Document, src: Document) -> None:
    body = dst.element.body
    for child in list(src.element.body):
        if child.tag == qn("w:sectPr"):
            continue
        body.append(child)


def _make_two_target_docs() -> tuple[Document, Document]:
    header = ["Элемент", "Формат", "Функции", "Эффект"]
    source_a, formatted_a = _make_two_fragment_table_docs(
        "2.1.5",
        [header, ["Советы", "Сессии", "Приоритизация", "Решения"]],
        [header, ["Форумы", "Площадки", "Обсуждения", "Согласование"]],
    )
    source_b, formatted_b = _make_two_fragment_table_docs(
        "2.2.3",
        [header, ["Первая линия", "Контроль", "Менеджмент", "Метрики"]],
        [header, ["Третья линия", "Аудит", "Совет", "Проверки"]],
    )
    _append_doc(source_a, source_b)
    _append_doc(formatted_a, formatted_b)
    return source_a, formatted_a


def _save_pair(source: Document, formatted: Document, tmp: str) -> tuple[Path, Path]:
    source_path = Path(tmp) / "source.docx"
    formatted_path = Path(tmp) / "formatted.docx"
    source.save(source_path)
    formatted.save(formatted_path)
    return source_path, formatted_path


def _source_note_follows_final_table(doc: Document) -> bool:
    body = [child for child in doc.element.body if child.tag != qn("w:sectPr")]
    table_positions = [idx for idx, child in enumerate(body) if child.tag == qn("w:tbl")]
    if not table_positions:
        return False
    final_table_pos = table_positions[-1]
    for child in body[final_table_pos + 1:]:
        if child.tag != qn("w:p"):
            continue
        text = " ".join("".join(node.text or "" for node in child.iter(qn("w:t"))).split())
        if not text:
            continue
        return text.startswith("Источник:")
    return False


def test_reprobes_after_successful_merge_so_later_table_indexes_shift() -> tuple[bool, str]:
    source, formatted = _make_two_target_docs()
    with tempfile.TemporaryDirectory() as tmp:
        source_path, formatted_path = _save_pair(source, formatted, tmp)
        old_rendered = tc._rendered_continuation_violations_for_docx
        old_marker = tc._same_page_continuation_marker_violations_for_docx
        old_regressions = tc._rendered_continuation_deletion_regressions
        try:
            def fake_rendered(path):
                table_count = len(Document(str(path)).tables)
                if table_count >= 4:
                    return [
                        _same_page_violation("2.1.5", 0, 1),
                        _same_page_violation("2.2.3", 2, 3),
                    ]
                if table_count == 3:
                    return [_same_page_violation("2.2.3", 1, 2)]
                return []

            tc._rendered_continuation_violations_for_docx = fake_rendered
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: []
            tc._rendered_continuation_deletion_regressions = lambda _path: []
            changed = tc.normalize_compatible_grid_same_page_repeated_fragments_inplace(
                formatted_path,
                source_docx_path=source_path,
            )
        finally:
            tc._rendered_continuation_violations_for_docx = old_rendered
            tc._same_page_continuation_marker_violations_for_docx = old_marker
            tc._rendered_continuation_deletion_regressions = old_regressions
        reread = Document(str(formatted_path))

    texts = _table_texts(reread)
    if changed != 2:
        return _result(False, f"expected both shifted-index candidates to merge, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, f"expected two merged survivor tables, got {len(reread.tables)}: {texts!r}")
    if any(text.count("Элемент | Формат | Функции | Эффект") != 1 for text in texts):
        return _result(False, f"duplicate headers remained after round2 merge: {texts!r}")
    if any(_numeric_row_count(_rows(reread, idx)) != 1 for idx in range(len(reread.tables))):
        return _result(False, f"duplicate numeric rows remained after round2 merge: {texts!r}")
    return _result(True, "compatible-grid normalizer re-probes after successful merge")


def test_no_numeric_same_page_variant_removes_only_second_header() -> tuple[bool, str]:
    header = ["Роль", "Ответственность", "Частота", "Артефакт", "KPI", "Срок"]
    source, formatted = _make_two_fragment_table_docs(
        "2.3.1",
        [header, ["Собственник", "Приоритеты", "Еженедельно", "План", "SLA", "Q1"]],
        [header, ["Представители", "Ревью", "Ежемесячно", "Отчет", "NPS", "Q2"]],
        numeric=False,
    )
    with tempfile.TemporaryDirectory() as tmp:
        source_path, formatted_path = _save_pair(source, formatted, tmp)
        old_rendered = tc._rendered_continuation_violations_for_docx
        old_marker = tc._same_page_continuation_marker_violations_for_docx
        old_regressions = tc._rendered_continuation_deletion_regressions
        old_orphan = tc._same_table_start_orphan_remains
        try:
            def fake_rendered(path):
                doc = Document(str(path))
                if len(doc.tables) == 2 and _rows(doc, 1)[0] == header:
                    return [_same_page_violation("2.3.1", 0, 1, numeric_count=0)]
                return []

            tc._rendered_continuation_violations_for_docx = fake_rendered
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: []
            tc._rendered_continuation_deletion_regressions = lambda _path: []
            tc._same_table_start_orphan_remains = lambda _path, _idx: False
            changed = tc.normalize_compatible_grid_same_page_repeated_fragments_inplace(
                formatted_path,
                source_docx_path=source_path,
            )
        finally:
            tc._rendered_continuation_violations_for_docx = old_rendered
            tc._same_page_continuation_marker_violations_for_docx = old_marker
            tc._rendered_continuation_deletion_regressions = old_regressions
            tc._same_table_start_orphan_remains = old_orphan
        reread = Document(str(formatted_path))

    texts = _table_texts(reread)
    joined = _all_table_text(reread)
    if changed != 1:
        return _result(False, f"expected no-numeric header cleanup, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, f"header cleanup should preserve both physical fragments: {texts!r}")
    if _rows(reread, 0)[0] != header:
        return _result(False, f"top header was not preserved: {texts!r}")
    if _rows(reread, 1)[0] == header:
        return _result(False, f"second-fragment duplicate header remained: {texts!r}")
    if joined.count("Роль | Ответственность | Частота | Артефакт | KPI | Срок") != 1:
        return _result(False, f"expected exactly one semantic header after cleanup: {texts!r}")
    if joined.count("Собственник | Приоритеты | Еженедельно | План | SLA | Q1") != 1:
        return _result(False, f"first real data row lost/duplicated: {texts!r}")
    if joined.count("Представители | Ревью | Ежемесячно | Отчет | NPS | Q2") != 1:
        return _result(False, f"second real data row lost/duplicated: {texts!r}")
    if any(_numeric_row_count(_rows(reread, idx)) for idx in range(len(reread.tables))):
        return _result(False, f"no-numeric cleanup synthesized a numeric row: {texts!r}")
    if not _source_note_follows_final_table(reread):
        return _result(False, "source note no longer follows the final table")
    return _result(True, "no-numeric same-page cleanup removes only the second duplicate header")


def test_no_numeric_header_cleanup_rolls_back_on_table_start_orphan() -> tuple[bool, str]:
    header = ["Роль", "Ответственность", "Частота", "Артефакт", "KPI", "Срок"]
    source, formatted = _make_two_fragment_table_docs(
        "2.3.1",
        [header, ["Собственник", "Приоритеты", "Еженедельно", "План", "SLA", "Q1"]],
        [header, ["Представители", "Ревью", "Ежемесячно", "Отчет", "NPS", "Q2"]],
        numeric=False,
    )
    with tempfile.TemporaryDirectory() as tmp:
        source_path, formatted_path = _save_pair(source, formatted, tmp)
        old_rendered = tc._rendered_continuation_violations_for_docx
        old_marker = tc._same_page_continuation_marker_violations_for_docx
        old_regressions = tc._rendered_continuation_deletion_regressions
        old_orphan = tc._same_table_start_orphan_remains
        try:
            def fake_rendered(path):
                doc = Document(str(path))
                if len(doc.tables) == 2 and _rows(doc, 1)[0] == header:
                    return [_same_page_violation("2.3.1", 0, 1, numeric_count=0)]
                return []

            tc._rendered_continuation_violations_for_docx = fake_rendered
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: []
            tc._rendered_continuation_deletion_regressions = lambda _path: []
            tc._same_table_start_orphan_remains = lambda _path, _idx: True
            changed = tc.normalize_compatible_grid_same_page_repeated_fragments_inplace(
                formatted_path,
                source_docx_path=source_path,
            )
        finally:
            tc._rendered_continuation_violations_for_docx = old_rendered
            tc._same_page_continuation_marker_violations_for_docx = old_marker
            tc._rendered_continuation_deletion_regressions = old_regressions
            tc._same_table_start_orphan_remains = old_orphan
        reread = Document(str(formatted_path))

    joined = _all_table_text(reread)
    if changed != 0:
        return _result(False, f"expected rollback to report no repair, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, f"rollback should preserve both physical fragments: {_table_texts(reread)!r}")
    if joined.count("Роль | Ответственность | Частота | Артефакт | KPI | Срок") != 2:
        return _result(False, f"rollback did not restore the second header: {_table_texts(reread)!r}")
    return _result(True, "no-numeric header cleanup rolls back when rendered orphan guard fails")


def main() -> int:
    tests = [
        ("reprobe shifted indexes", test_reprobes_after_successful_merge_so_later_table_indexes_shift),
        ("no-numeric header cleanup", test_no_numeric_same_page_variant_removes_only_second_header),
        ("no-numeric cleanup orphan rollback", test_no_numeric_header_cleanup_rolls_back_on_table_start_orphan),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        status = "PASS" if ok else "FAIL"
        print(f"[{status}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
