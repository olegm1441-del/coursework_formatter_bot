"""Stage C — same-page manual-chain cleanup for grid-INCOMPATIBLE fragments.

Pins the safety core of `cleanup_same_page_incompatible_chains_inplace`:
- the second fragment's duplicate header/numeric leading rows are removed;
- meaningful data rows are never removed;
- the candidate only fires on a proven duplicate-header second fragment with
  distinct data and a matching caption, and refuses source-proven duplication.

End-to-end rendered behaviour (markers actually disappear for Bondarev
1.3.1/2.1.4/2.1.5) is covered by the format-from-source smoke.

Run: python3 tests/test_same_page_incompatible_cleanup.py
"""
from __future__ import annotations

import sys
from pathlib import Path
from types import SimpleNamespace

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
import guides.coursework_kfu_2025.table_continuation as tc  # noqa: E402


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def _mk_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    grid = table._tbl.find(qn("w:tblGrid"))
    for col, w in zip(grid.findall(qn("w:gridCol")), widths):
        col.set(qn("w:w"), str(w))
    for r, vals in zip(table.rows, rows):
        for c, val in zip(r.cells, vals):
            c.text = val
    return table


def _all_fps(table):
    # every row fingerprint, header-position-independent
    return [tc._docx_row_fingerprint(r) for r in table.rows]


def test_remove_duplicate_leading_rows_preserves_data() -> tuple[bool, str]:
    doc = Document()
    first = _mk_table(doc, [["Орган", "Компетенции", "Регламент"],
                            ["1", "2", "3"],
                            ["ОСА", "устав", "ежегодно"]], [3000, 3000, 3000])
    second = _mk_table(doc, [["Орган", "Компетенции", "Регламент"],
                             ["1", "2", "3"],
                             ["Правление", "операционное", "по плану"],
                             ["Секретарь", "раскрытие", "постоянно"]], [2000, 4000, 3000])
    before_first = _all_fps(first)
    p_fp = tc._docx_row_fingerprint(second.rows[2])
    s_fp = tc._docx_row_fingerprint(second.rows[3])
    removed = tc._remove_second_fragment_duplicate_leading_rows(first, second)
    if removed != 2:
        return _result(False, f"expected 2 leading dup rows removed, got {removed}")
    if _all_fps(first) != before_first:
        return _result(False, "first fragment rows changed")
    remaining = _all_fps(second)
    if remaining != [p_fp, s_fp]:
        return _result(False, f"data rows not preserved exactly/in order: {remaining}")
    return _result(True, "duplicate header+numeric removed, both data rows preserved in order")


def test_remove_never_empties_or_drops_data() -> tuple[bool, str]:
    doc = Document()
    first = _mk_table(doc, [["H1", "H2"], ["ряд", "данных"]], [3000, 3000])
    # second has dup header then a single data row — must keep the data row
    second = _mk_table(doc, [["H1", "H2"], ["другой", "ряд"]], [3000, 3000])
    data_fp = tc._docx_row_fingerprint(second.rows[1])
    removed = tc._remove_second_fragment_duplicate_leading_rows(first, second)
    if removed != 1:
        return _result(False, f"expected 1 dup header removed, got {removed}")
    if _all_fps(second) != [data_fp]:
        return _result(False, f"second fragment lost/changed its only data row: {_all_fps(second)}")
    return _result(True, "removal stops at data row and never empties the table")


def _violation(first_idx, second_idx, table_num="1.3.1"):
    return SimpleNamespace(
        table_num=table_num,
        table_index=first_idx,
        violation_type="same_page_repeated_fragment",
        confidence="high",
        evidence={"following_table_index": second_idx},
    )


def _candidate_doc():
    doc = Document()
    doc.add_paragraph("Таблица 1.3.1")
    _mk_table(doc, [["Орган", "Компетенции", "Регламент"],
                    ["ОСА", "устав", "ежегодно"]], [3000, 3000, 3000])
    doc.add_paragraph("Продолжение таблицы 1.3.1")
    _mk_table(doc, [["Орган", "Компетенции", "Регламент"],
                    ["Правление", "операционное", "по плану"]], [2000, 4000, 3000])
    return doc


def test_candidate_fires_on_incompatible_marker_chain() -> tuple[bool, str]:
    saved = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: False
    try:
        doc = _candidate_doc()
        cand = tc._incompatible_grid_same_page_cleanup_candidate_from_rendered(
            doc, _violation(0, 1), source_docx_path=None)
    finally:
        tc._source_has_meaningful_duplicate_for_table = saved
    if cand is None:
        return _result(False, "candidate did not fire on a valid incompatible marker chain")
    table_num, first_idx, second_idx, marker_para = cand
    if (table_num, first_idx, second_idx) != ("1.3.1", 0, 1):
        return _result(False, f"wrong candidate tuple: {(table_num, first_idx, second_idx)}")
    if marker_para is None:
        return _result(False, "marker paragraph not located for cleanup")
    return _result(True, "candidate fires on incompatible-grid marker chain")


def test_candidate_skips_source_bad() -> tuple[bool, str]:
    saved = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: True  # source-proven dup
    try:
        doc = _candidate_doc()
        cand = tc._incompatible_grid_same_page_cleanup_candidate_from_rendered(
            doc, _violation(0, 1), source_docx_path=None)
    finally:
        tc._source_has_meaningful_duplicate_for_table = saved
    if cand is not None:
        return _result(False, "candidate must skip source-proven duplicated tables")
    return _result(True, "candidate skips source-bad duplicated tables")


def test_candidate_skips_when_second_starts_with_data() -> tuple[bool, str]:
    saved = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: False
    try:
        doc = Document()
        doc.add_paragraph("Таблица 1.3.1")
        _mk_table(doc, [["Орган", "Компетенции", "Регламент"],
                        ["ОСА", "устав", "ежегодно"]], [3000, 3000, 3000])
        doc.add_paragraph("Продолжение таблицы 1.3.1")
        # second starts directly with data (no repeated header) — must NOT touch
        _mk_table(doc, [["Правление", "операционное", "по плану"],
                        ["Секретарь", "раскрытие", "постоянно"]], [3000, 3000, 3000])
        cand = tc._incompatible_grid_same_page_cleanup_candidate_from_rendered(
            doc, _violation(0, 1), source_docx_path=None)
    finally:
        tc._source_has_meaningful_duplicate_for_table = saved
    if cand is not None:
        return _result(False, "candidate must skip when second fragment has no repeated header")
    return _result(True, "candidate skips when second fragment starts with data")


def _chain_doc(num="1.3.1", second_first_row=None):
    """caption -> table(header+data) -> 'Продолжение таблицы N' -> table(dup header+data)."""
    doc = Document()
    doc.add_paragraph(f"Таблица {num}")
    _mk_table(doc, [["Орган", "Функция", "Срок"],
                    ["совет", "надзор за правлением", "ежегодно"]], [3000, 3000, 3000])
    doc.add_paragraph(f"Продолжение таблицы {num}")
    second = second_first_row or ["Орган", "Функция", "Срок"]
    _mk_table(doc, [second,
                    ["правление", "оперативное управление", "ежеквартально"]], [2200, 3600, 3200])
    return doc


def test_continuation_candidate_found_from_table_num() -> tuple[bool, str]:
    # candidate is located from the same_page_continuation table number via its
    # marker — independent of any `same_page_repeated_fragment` signal
    saved = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: False
    try:
        doc = _chain_doc("1.3.1")
        cand = tc._same_page_continuation_cleanup_candidate(doc, "1.3.1", source_docx_path=None)
    finally:
        tc._source_has_meaningful_duplicate_for_table = saved
    if cand is None:
        return _result(False, "candidate not found from table number + marker")
    first_idx, second_idx, marker_para = cand
    if (first_idx, second_idx) != (0, 1) or marker_para is None:
        return _result(False, f"wrong candidate: {(first_idx, second_idx, marker_para is not None)}")
    return _result(True, "same-page continuation candidate found from acceptance-blocker table number")


def test_continuation_candidate_requires_marker() -> tuple[bool, str]:
    saved = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: False
    try:
        doc = Document()
        doc.add_paragraph("Таблица 1.3.1")
        _mk_table(doc, [["Орган", "Функция"], ["совет", "надзор"]], [3000, 3000])
        # no marker, just a following table
        _mk_table(doc, [["Орган", "Функция"], ["правление", "управление"]], [3000, 3000])
        cand = tc._same_page_continuation_cleanup_candidate(doc, "1.3.1", source_docx_path=None)
    finally:
        tc._source_has_meaningful_duplicate_for_table = saved
    if cand is not None:
        return _result(False, "candidate must require a continuation marker")
    return _result(True, "no marker -> no continuation cleanup candidate")


def test_continuation_candidate_skips_source_bad() -> tuple[bool, str]:
    saved = tc._source_has_meaningful_duplicate_for_table
    tc._source_has_meaningful_duplicate_for_table = lambda *a, **k: True
    try:
        cand = tc._same_page_continuation_cleanup_candidate(_chain_doc("1.3.1"), "1.3.1", source_docx_path=None)
    finally:
        tc._source_has_meaningful_duplicate_for_table = saved
    if cand is not None:
        return _result(False, "source-bad table must not be a cleanup candidate")
    return _result(True, "source-bad table skipped by continuation cleanup candidate")


def main() -> int:
    tests = [
        ("remove dup leading rows preserves data", test_remove_duplicate_leading_rows_preserves_data),
        ("continuation candidate from table_num", test_continuation_candidate_found_from_table_num),
        ("continuation candidate requires marker", test_continuation_candidate_requires_marker),
        ("continuation candidate skips source-bad", test_continuation_candidate_skips_source_bad),
        ("removal never empties/drops data", test_remove_never_empties_or_drops_data),
        ("candidate fires on incompatible chain", test_candidate_fires_on_incompatible_marker_chain),
        ("candidate skips source-bad", test_candidate_skips_source_bad),
        ("candidate skips second-starts-with-data", test_candidate_skips_when_second_starts_with_data),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        print(f"[{'PASS' if ok else 'FAIL'}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
