from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import guides.coursework_kfu_2025.table_continuation as tc


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def _set_grid(table, widths: list[str]) -> None:
    tbl = table._tbl
    grid = tbl.tblGrid
    if grid is not None:
        tbl.remove(grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), width)
        grid.append(col)
    tbl.insert(0, grid)


def _set_border(table, value: str = "single") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is not None:
        tbl_pr.remove(borders)
    borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{name}")
        border.set(qn("w:val"), value)
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), "000000")
        borders.append(border)
    tbl_pr.append(borders)


def _set_margins(table, value: str = "15") -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is not None:
        tbl_pr.remove(margins)
    margins = OxmlElement("w:tblCellMar")
    for name in ("top", "left", "bottom", "right"):
        margin = OxmlElement(f"w:{name}")
        margin.set(qn("w:w"), value)
        margin.set(qn("w:type"), "dxa")
        margins.append(margin)
    tbl_pr.append(margins)


def _fill_row(row, values: list[str]) -> None:
    for idx, value in enumerate(values):
        row.cells[idx].text = value


def _add_table(doc: Document, rows: list[list[str]], grid: list[str] | None = None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row, values in zip(table.rows, rows):
        _fill_row(row, values)
    _set_grid(table, grid or ["2000", "2500", "3000"])
    _set_border(table)
    _set_margins(table)
    return table


def _make_same_page_doc(
    *,
    exact_grid: bool = True,
    source_numeric: bool = False,
    source_bad_duplicate: bool = False,
    with_source_note: bool = True,
) -> tuple[Document, Document]:
    header = ["Орган", "Компетенции", "Регламент"]
    numeric = ["1", "2", "3"]
    first_data = [
        ["ОСА", "Устав и дивиденды", "Закон 208-ФЗ"],
        ["НС", "Стратегия и контроль", "Положение о НС"],
    ]
    second_data = [
        ["Правление", "Операционное управление", "Положение о Правлении"],
        ["Корпоративный секретарь", "Права акционеров", "Положение о секретаре"],
    ]
    if source_bad_duplicate:
        second_data = [first_data[-1], second_data[-1]]

    source = Document()
    source.add_paragraph("Таблица 2.1.1")
    source.add_paragraph("Карта органов управления")
    source_first = [header]
    if source_numeric:
        source_first.append(numeric)
    source_first.extend(first_data)
    _add_table(source, source_first)
    source.add_paragraph("Продолжение табл. 2.1.1")
    source_second = [header]
    if source_numeric:
        source_second.append(numeric)
    source_second.extend(second_data)
    _add_table(source, source_second)
    if with_source_note:
        source.add_paragraph("Источник: составлено автором.")

    formatted = Document()
    formatted.add_paragraph("Таблица 2.1.1")
    formatted.add_paragraph("Карта органов управления")
    _add_table(formatted, [header, numeric, *first_data])
    formatted.add_paragraph("Продолжение таблицы 2.1.1")
    second_grid = ["2000", "2500", "3000"] if exact_grid else ["1900", "2600", "3000"]
    _add_table(formatted, [header, numeric, *second_data], grid=second_grid)
    if with_source_note:
        formatted.add_paragraph("Источник: составлено автором.")

    return source, formatted


def _save_pair(source: Document, formatted: Document, tmp: str) -> tuple[Path, Path]:
    source_path = Path(tmp) / "source.docx"
    formatted_path = Path(tmp) / "formatted.docx"
    source.save(source_path)
    formatted.save(formatted_path)
    return source_path, formatted_path


def _table_rows(doc: Document, table_index: int = 0) -> list[list[str]]:
    return [
        [" ".join(cell.text.split()) for cell in row.cells]
        for row in doc.tables[table_index].rows
    ]


def _numeric_row_count(rows: list[list[str]]) -> int:
    count = 0
    for row in rows:
        if row == [str(idx) for idx in range(1, len(row) + 1)]:
            count += 1
    return count


def _source_note_after_only_table(doc: Document) -> bool:
    children = list(doc.element.body)
    table_positions = [idx for idx, child in enumerate(children) if child.tag == qn("w:tbl")]
    if len(table_positions) != 1:
        return False
    table_pos = table_positions[0]
    for idx, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.findall(".//" + qn("w:t"))).strip()
        if text.startswith("Источник:"):
            return idx > table_pos
    return False


def _marker_violation() -> tc._SamePageContinuationMarkerViolation:
    return tc._SamePageContinuationMarkerViolation(
        marker_text="Продолжение таблицы 2.1.1",
        marker_page=29,
        previous_table_page=29,
        following_table_page=29,
        confidence="high",
    )


def _run_repair(
    source: Document,
    formatted: Document,
    *,
    after_marker_remains: bool = False,
) -> tuple[int, Document]:
    with tempfile.TemporaryDirectory() as tmp:
        source_path, formatted_path = _save_pair(source, formatted, tmp)
        old_marker = tc._same_page_continuation_marker_violations_for_docx
        old_rendered = tc._rendered_continuation_violations_for_docx
        old_regressions = tc._rendered_continuation_deletion_regressions
        try:
            calls = {"n": 0}

            def fake_marker(_path):
                calls["n"] += 1
                if calls["n"] == 1 or after_marker_remains:
                    return [_marker_violation()]
                return []

            tc._same_page_continuation_marker_violations_for_docx = fake_marker
            tc._rendered_continuation_violations_for_docx = lambda _path: []
            tc._rendered_continuation_deletion_regressions = lambda _path: []
            changed = tc.normalize_exact_grid_same_page_repeated_fragments_inplace(
                formatted_path,
                source_docx_path=source_path,
            )
        finally:
            tc._same_page_continuation_marker_violations_for_docx = old_marker
            tc._rendered_continuation_violations_for_docx = old_rendered
            tc._rendered_continuation_deletion_regressions = old_regressions
        return changed, Document(str(formatted_path))


def test_exact_grid_same_page_fragment_merges_safely() -> tuple[bool, str]:
    source, formatted = _make_same_page_doc()
    changed, reread = _run_repair(source, formatted)
    rows = _table_rows(reread)
    row_texts = [" | ".join(row) for row in rows]

    if changed != 1:
        return _result(False, f"expected one exact-grid repair, got {changed}")
    if len(reread.tables) != 1:
        return _result(False, f"expected one merged table, got {len(reread.tables)}")
    expected_data = ["ОСА", "НС", "Правление", "Корпоративный секретарь"]
    actual_data = [row[0] for row in rows[1:] if row[0] != "1"]
    if actual_data != expected_data:
        return _result(False, f"data rows not preserved in order: {actual_data!r}")
    if any("Продолжение таблицы" in paragraph.text for paragraph in reread.paragraphs):
        return _result(False, "same-page continuation marker remained")
    if row_texts.count("Орган | Компетенции | Регламент") != 1:
        return _result(False, f"duplicate header remained: {row_texts!r}")
    return _result(True, "exact-grid same-page fragments merged")


def test_generated_numeric_row_from_second_fragment_is_removed() -> tuple[bool, str]:
    source, formatted = _make_same_page_doc()
    changed, reread = _run_repair(source, formatted)
    rows = _table_rows(reread)

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if _numeric_row_count(rows) != 1:
        return _result(False, f"expected only the first fragment numeric row to remain, got {rows!r}")
    return _result(True, "duplicate formatter-generated numeric row removed after same-page merge")


def test_source_numeric_rows_are_preserved_by_skipping_repair() -> tuple[bool, str]:
    source, formatted = _make_same_page_doc(source_numeric=True)
    changed, reread = _run_repair(source, formatted)
    all_rows = [_table_rows(reread, idx) for idx in range(len(reread.tables))]

    if changed != 0:
        return _result(False, f"source numeric row should make repair skip, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, f"skipped repair should preserve both tables, got {len(reread.tables)}")
    if sum(_numeric_row_count(rows) for rows in all_rows) != 2:
        return _result(False, f"source numeric rows were not preserved: {all_rows!r}")
    return _result(True, "source-proven numeric rows are preserved")


def test_source_note_remains_after_merged_table() -> tuple[bool, str]:
    source, formatted = _make_same_page_doc()
    changed, reread = _run_repair(source, formatted)

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if not _source_note_after_only_table(reread):
        return _result(False, "source note is not after the merged table")
    return _result(True, "source note remains after merged table")


def test_grid_mismatch_same_page_fragments_are_skipped() -> tuple[bool, str]:
    source, formatted = _make_same_page_doc(exact_grid=False)
    changed, reread = _run_repair(source, formatted)

    if changed != 0:
        return _result(False, f"grid mismatch must skip, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, f"grid mismatch skip should preserve both tables, got {len(reread.tables)}")
    return _result(True, "grid mismatch same-page fragments are skipped")


def test_source_bad_duplicate_meaningful_rows_are_skipped() -> tuple[bool, str]:
    source, formatted = _make_same_page_doc(source_bad_duplicate=True)
    changed, reread = _run_repair(source, formatted)

    if changed != 0:
        return _result(False, f"source-bad duplicate data should skip, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, f"source-bad skip should preserve both tables, got {len(reread.tables)}")
    return _result(True, "source-bad meaningful duplicates are skipped")


def test_rollback_preserves_original_when_validation_worsens() -> tuple[bool, str]:
    source, formatted = _make_same_page_doc()
    changed, reread = _run_repair(source, formatted, after_marker_remains=True)

    if changed != 0:
        return _result(False, f"failed post-render validation should roll back, got {changed}")
    if len(reread.tables) != 2:
        return _result(False, "rollback did not preserve original table fragments")
    if not any("Продолжение таблицы" in paragraph.text for paragraph in reread.paragraphs):
        return _result(False, "rollback did not restore original marker")
    return _result(True, "failed validation rolls back exact-grid repair")


def test_same_page_merge_does_not_use_page_breaks_or_duplicate_numeric_rows() -> tuple[bool, str]:
    source, formatted = _make_same_page_doc()
    changed, reread = _run_repair(source, formatted)
    caption = next(p for p in reread.paragraphs if p.text == "Таблица 2.1.1")
    p_pr = caption._element.find(qn("w:pPr"))
    page_break = p_pr.find(qn("w:pageBreakBefore")) if p_pr is not None else None

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if page_break is not None:
        return _result(False, "same-page merge must not insert pageBreakBefore")
    if _numeric_row_count(_table_rows(reread)) != 1:
        return _result(False, "same-page merge must not duplicate numeric rows")
    return _result(True, "same-page merge does not use page breaks or duplicate numeric rows")


def main() -> int:
    tests = [
        ("exact-grid merge", test_exact_grid_same_page_fragment_merges_safely),
        ("remove second generated numeric row", test_generated_numeric_row_from_second_fragment_is_removed),
        ("preserve source numeric rows", test_source_numeric_rows_are_preserved_by_skipping_repair),
        ("source note placement", test_source_note_remains_after_merged_table),
        ("grid mismatch skip", test_grid_mismatch_same_page_fragments_are_skipped),
        ("source-bad duplicate skip", test_source_bad_duplicate_meaningful_rows_are_skipped),
        ("rollback on validation failure", test_rollback_preserves_original_when_validation_worsens),
        ("no page break or duplicate numeric", test_same_page_merge_does_not_use_page_breaks_or_duplicate_numeric_rows),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        status = "PASS" if ok else "FAIL"
        print(f"[{status}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
