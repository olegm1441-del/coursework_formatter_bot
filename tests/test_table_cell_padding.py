from __future__ import annotations

import sys
import tempfile
from pathlib import Path
from types import SimpleNamespace

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from guides.coursework_kfu_2025.safe_formatter import format_tables
import guides.coursework_kfu_2025.table_continuation as tc


TARGET_PADDING_DXA = "113"


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def _add_table(doc: Document):
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Показатель"
    table.rows[0].cells[1].text = "Значение"
    table.rows[1].cells[0].text = "Выручка"
    table.rows[1].cells[1].text = "12"
    return table


def _tbl_cell_mar(table):
    tbl_pr = table._tbl.tblPr
    return tbl_pr.find(qn("w:tblCellMar")) if tbl_pr is not None else None


def _set_tbl_margin(table, side: str, value: str, margin_type: str = "dxa") -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    node = margins.find(qn(f"w:{side}"))
    if node is None:
        node = OxmlElement(f"w:{side}")
        margins.append(node)
    node.set(qn("w:w"), value)
    node.set(qn("w:type"), margin_type)


def _tbl_margin_value(table, side: str) -> tuple[str | None, str | None]:
    margins = _tbl_cell_mar(table)
    if margins is None:
        return None, None
    node = margins.find(qn(f"w:{side}"))
    if node is None:
        return None, None
    return node.get(qn("w:w")), node.get(qn("w:type"))


def _set_cell_margin(cell, side: str, value: str, margin_type: str = "dxa") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    node = margins.find(qn(f"w:{side}"))
    if node is None:
        node = OxmlElement(f"w:{side}")
        margins.append(node)
    node.set(qn("w:w"), value)
    node.set(qn("w:type"), margin_type)


def _cell_margin_value(cell, side: str) -> tuple[str | None, str | None]:
    tc_pr = cell._tc.tcPr
    margins = tc_pr.find(qn("w:tcMar")) if tc_pr is not None else None
    if margins is None:
        return None, None
    node = margins.find(qn(f"w:{side}"))
    if node is None:
        return None, None
    return node.get(qn("w:w")), node.get(qn("w:type"))


def _make_preserve_geometry_table(doc: Document):
    table = _add_table(doc)
    tc_pr = table.rows[0].cells[0]._tc.get_or_add_tcPr()
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), "2")
    tc_pr.append(grid_span)
    return table


def _add_rows_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row, values in zip(table.rows, rows):
        for cell, value in zip(row.cells, values):
            cell.text = value
    return table


def _same_page_violation():
    return SimpleNamespace(
        table_num="1.3.1",
        table_index=0,
        page=22,
        violation_type="same_page_repeated_fragment",
        confidence="high",
        evidence={"following_table_index": 1},
    )


def _compatible_source_and_formatted_docs() -> tuple[Document, Document]:
    header = ["Показатель", "2021", "2022"]
    numeric = ["1", "2", "3"]

    source = Document()
    source.add_paragraph("Таблица 1.3.1")
    source.add_paragraph("Динамика показателей")
    _add_rows_table(source, [header, ["Выручка", "10", "12"]])
    _add_rows_table(source, [header, ["Прибыль", "3", "4"]])
    source.add_paragraph("Источник: составлено автором.")

    formatted = Document()
    formatted.add_paragraph("Таблица 1.3.1")
    formatted.add_paragraph("Динамика показателей")
    _add_rows_table(formatted, [header, numeric, ["Выручка", "10", "12"]])
    _add_rows_table(formatted, [header, numeric, ["Прибыль", "3", "4"]])
    formatted.add_paragraph("Источник: составлено автором.")
    return source, formatted


def test_table_without_tbl_cell_mar_gets_horizontal_padding() -> tuple[bool, str]:
    doc = Document()
    table = _add_table(doc)
    format_tables(doc)
    left = _tbl_margin_value(table, "left")
    right = _tbl_margin_value(table, "right")
    return _result(
        left == (TARGET_PADDING_DXA, "dxa") and right == (TARGET_PADDING_DXA, "dxa"),
        f"left={left}, right={right}",
    )


def test_tiny_table_margins_become_target_horizontal_padding() -> tuple[bool, str]:
    doc = Document()
    table = _add_table(doc)
    _set_tbl_margin(table, "left", "15")
    _set_tbl_margin(table, "right", "15")
    format_tables(doc)
    left = _tbl_margin_value(table, "left")
    right = _tbl_margin_value(table, "right")
    return _result(
        left == (TARGET_PADDING_DXA, "dxa") and right == (TARGET_PADDING_DXA, "dxa"),
        f"left={left}, right={right}",
    )


def test_existing_top_bottom_margins_are_preserved() -> tuple[bool, str]:
    doc = Document()
    table = _add_table(doc)
    _set_tbl_margin(table, "top", "44")
    _set_tbl_margin(table, "bottom", "55")
    format_tables(doc)
    return _result(
        _tbl_margin_value(table, "top") == ("44", "dxa")
        and _tbl_margin_value(table, "bottom") == ("55", "dxa")
        and _tbl_margin_value(table, "left") == (TARGET_PADDING_DXA, "dxa")
        and _tbl_margin_value(table, "right") == (TARGET_PADDING_DXA, "dxa"),
        f"top={_tbl_margin_value(table, 'top')}, bottom={_tbl_margin_value(table, 'bottom')}, "
        f"left={_tbl_margin_value(table, 'left')}, right={_tbl_margin_value(table, 'right')}",
    )


def test_cell_margin_overrides_do_not_defeat_horizontal_padding() -> tuple[bool, str]:
    doc = Document()
    table = _add_table(doc)
    cell = table.rows[1].cells[0]
    _set_cell_margin(cell, "left", "0")
    _set_cell_margin(cell, "right", "999")
    _set_cell_margin(cell, "top", "22")
    format_tables(doc)
    return _result(
        _tbl_margin_value(table, "left") == (TARGET_PADDING_DXA, "dxa")
        and _tbl_margin_value(table, "right") == (TARGET_PADDING_DXA, "dxa")
        and _cell_margin_value(cell, "left") == (None, None)
        and _cell_margin_value(cell, "right") == (None, None)
        and _cell_margin_value(cell, "top") == ("22", "dxa"),
        f"table_left={_tbl_margin_value(table, 'left')}, table_right={_tbl_margin_value(table, 'right')}, "
        f"cell_left={_cell_margin_value(cell, 'left')}, cell_right={_cell_margin_value(cell, 'right')}, "
        f"cell_top={_cell_margin_value(cell, 'top')}",
    )


def test_preserve_geometry_table_still_gets_padding() -> tuple[bool, str]:
    doc = Document()
    table = _make_preserve_geometry_table(doc)
    format_tables(doc)
    return _result(
        _tbl_margin_value(table, "left") == (TARGET_PADDING_DXA, "dxa")
        and _tbl_margin_value(table, "right") == (TARGET_PADDING_DXA, "dxa"),
        f"left={_tbl_margin_value(table, 'left')}, right={_tbl_margin_value(table, 'right')}",
    )


def test_preserve_geometry_existing_top_bottom_margins_are_preserved() -> tuple[bool, str]:
    doc = Document()
    table = _make_preserve_geometry_table(doc)
    _set_tbl_margin(table, "top", "66")
    _set_tbl_margin(table, "bottom", "77")
    format_tables(doc)
    return _result(
        _tbl_margin_value(table, "top") == ("66", "dxa")
        and _tbl_margin_value(table, "bottom") == ("77", "dxa")
        and _tbl_margin_value(table, "left") == (TARGET_PADDING_DXA, "dxa")
        and _tbl_margin_value(table, "right") == (TARGET_PADDING_DXA, "dxa"),
        f"top={_tbl_margin_value(table, 'top')}, bottom={_tbl_margin_value(table, 'bottom')}, "
        f"left={_tbl_margin_value(table, 'left')}, right={_tbl_margin_value(table, 'right')}",
    )


def test_same_page_merge_survivor_keeps_padding() -> tuple[bool, str]:
    source, formatted = _compatible_source_and_formatted_docs()
    format_tables(formatted)
    with tempfile.TemporaryDirectory() as tmp:
        source_path = Path(tmp) / "source.docx"
        formatted_path = Path(tmp) / "formatted.docx"
        source.save(source_path)
        formatted.save(formatted_path)

        old_rendered = tc._rendered_continuation_violations_for_docx
        old_marker = tc._same_page_continuation_marker_violations_for_docx
        old_regressions = tc._rendered_continuation_deletion_regressions
        calls = {"n": 0}
        try:
            def fake_rendered(_path):
                calls["n"] += 1
                return [_same_page_violation()] if calls["n"] == 1 else []

            tc._rendered_continuation_violations_for_docx = fake_rendered
            tc._same_page_continuation_marker_violations_for_docx = lambda _path: []
            tc._rendered_continuation_deletion_regressions = lambda _path: []
            changed = tc.normalize_compatible_grid_same_page_repeated_fragments_inplace(
                formatted_path,
                source_docx_path=source_path,
            )
        finally:
            tc._rendered_continuation_violations_for_docx = old_rendered
            tc._same_page_continuation_marker_violations_for_docx = old_marker
            tc._rendered_continuation_deletion_regressions = old_regressions

        result = Document(str(formatted_path))

    table = result.tables[0] if result.tables else None
    left = _tbl_margin_value(table, "left") if table is not None else None
    right = _tbl_margin_value(table, "right") if table is not None else None
    return _result(
        changed == 1
        and len(result.tables) == 1
        and left == (TARGET_PADDING_DXA, "dxa")
        and right == (TARGET_PADDING_DXA, "dxa"),
        f"changed={changed}, tables={len(result.tables)}, left={left}, right={right}",
    )


def main() -> int:
    tests = [
        ("no tblCellMar", test_table_without_tbl_cell_mar_gets_horizontal_padding),
        ("tiny margins", test_tiny_table_margins_become_target_horizontal_padding),
        ("preserve top/bottom", test_existing_top_bottom_margins_are_preserved),
        ("cell overrides", test_cell_margin_overrides_do_not_defeat_horizontal_padding),
        ("preserve geometry", test_preserve_geometry_table_still_gets_padding),
        ("preserve geometry top/bottom", test_preserve_geometry_existing_top_bottom_margins_are_preserved),
        ("same-page merge survivor padding", test_same_page_merge_survivor_keeps_padding),
    ]
    failed = 0
    for name, test in tests:
        ok, message = test()
        if ok:
            print(f"[PASS] {name} — {message}")
        else:
            failed += 1
            print(f"[FAIL] {name} — {message}")
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
