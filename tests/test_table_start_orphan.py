from __future__ import annotations

import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml.ns import qn

from guides.coursework_kfu_2025.pdf_layout_analyzer import PdfLine
import guides.coursework_kfu_2025.table_continuation as tc


def _result(ok: bool, message: str) -> tuple[bool, str]:
    return ok, message


def _is_blank_paragraph_node(node) -> bool:
    if node.tag != qn("w:p"):
        return False
    return not "".join(text.text or "" for text in node.findall(".//" + qn("w:t"))).strip()


def _caption_index(doc: Document, text: str) -> int:
    for idx, child in enumerate(doc.element.body):
        if child.tag != qn("w:p"):
            continue
        if "".join(t.text or "" for t in child.findall(".//" + qn("w:t"))).strip() == text:
            return idx
    raise AssertionError(f"caption not found: {text!r}")


def _blank_count_before_caption(doc: Document, caption_text: str) -> int:
    idx = _caption_index(doc, caption_text)
    children = list(doc.element.body)
    count = 0
    j = idx - 1
    while j >= 0 and _is_blank_paragraph_node(children[j]):
        count += 1
        j -= 1
    return count


def _caption_has_active_page_break(doc: Document, caption_text: str) -> bool:
    para = next(p for p in doc.paragraphs if p.text == caption_text)
    p_pr = para._element.find(qn("w:pPr"))
    page_break = p_pr.find(qn("w:pageBreakBefore")) if p_pr is not None else None
    if page_break is None:
        return False
    return page_break.get(qn("w:val")) not in {"0", "false", "False", "off"}


def _numeric_row_count(doc: Document, table_index: int = 0) -> int:
    count = 0
    for row in doc.tables[table_index].rows:
        values = [" ".join(cell.text.split()) for cell in row.cells]
        if len(values) >= 2 and values == [str(i) for i in range(1, len(values) + 1)]:
            count += 1
    return count


def _continuation_markers(doc: Document) -> list[str]:
    return [p.text for p in doc.paragraphs if "Продолжение таблицы" in p.text]


def _repaired_lines(*, with_numeric: bool = False) -> list[PdfLine]:
    lines = [
        PdfLine("Таблица 1.1.3", 2, 58.0, 70.0),
        PdfLine("Длинная таблица с признаками хаоса в курсовой", 2, 82.0, 94.0),
        PdfLine("Группа признаков Пример дефекта Статус", 2, 106.0, 118.0),
    ]
    if with_numeric:
        lines.append(PdfLine("1 2 3", 2, 124.0, 136.0))
    lines.extend(
        [
            PdfLine("Кейс 1 строка 1: текст таблицы повторяется тест", 2, 142.0, 156.0),
            PdfLine("Кейс 2 строка 2: текст таблицы повторяется тест", 2, 172.0, 186.0),
        ]
    )
    return lines


def _run_with_rendered_lines(
    doc: Document,
    lines: list[PdfLine],
    name: str = "case.docx",
    after_lines: list[PdfLine] | None = None,
) -> tuple[int, Document]:
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / name
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "case.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            calls = {"n": 0}

            def fake_analyze(_path):
                calls["n"] += 1
                if calls["n"] <= 3 or after_lines is None:
                    return lines
                return after_lines

            tc.analyze_pdf_lines = fake_analyze
            changed = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        return changed, Document(str(path))


def _add_table_start_orphan_doc(*, with_numeric: bool = False, preexisting_blank: bool = False) -> Document:
    doc = Document()
    doc.add_paragraph("Перед таблицей текст.")
    if preexisting_blank:
        doc.add_paragraph("")
    doc.add_paragraph("Таблица 1.1.3")
    doc.add_paragraph("Длинная таблица с признаками хаоса в курсовой")
    table = doc.add_table(rows=4 if with_numeric else 3, cols=3)
    table.rows[0].cells[0].text = "Группа признаков"
    table.rows[0].cells[1].text = "Пример дефекта"
    table.rows[0].cells[2].text = "Статус"
    data_start = 2 if with_numeric else 1
    if with_numeric:
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "2"
        table.rows[1].cells[2].text = "3"
    table.rows[data_start].cells[0].text = "Кейс 1"
    table.rows[data_start].cells[1].text = "строка 1: текст таблицы повторяется"
    table.rows[data_start].cells[2].text = "тест"
    table.rows[data_start + 1].cells[0].text = "Кейс 2"
    table.rows[data_start + 1].cells[1].text = "строка 2: текст таблицы повторяется"
    table.rows[data_start + 1].cells[2].text = "тест"
    doc.add_paragraph("Источник: составлено автором.")
    return doc


def _orphan_lines(*, with_numeric: bool = False) -> list[PdfLine]:
    lines = [
        PdfLine("Таблица 1.1.3", 1, 644.6, 658.6),
        PdfLine("Длинная таблица с признаками хаоса в курсовой", 1, 668.7, 682.0),
        PdfLine("Группа признаков Пример дефекта Статус", 1, 693.1, 706.9),
    ]
    if with_numeric:
        lines.append(PdfLine("1 2 3", 1, 716.0, 728.0))
    lines.extend(
        [
            PdfLine("Кейс 1 строка 1: текст таблицы повторяется тест", 2, 58.5, 72.3),
            PdfLine("Кейс 2 строка 2: текст таблицы повторяется тест", 2, 86.1, 99.9),
        ]
    )
    return lines


def test_table_start_orphan_inserts_two_blank_paragraphs() -> tuple[bool, str]:
    changed, reread = _run_with_rendered_lines(
        _add_table_start_orphan_doc(),
        _orphan_lines(),
        after_lines=_repaired_lines(),
    )

    if changed != 1:
        return _result(False, f"expected one table-start orphan repair, got {changed}")
    blanks = _blank_count_before_caption(reread, "Таблица 1.1.3")
    if blanks != 2:
        return _result(False, f"expected exactly two blank paragraphs before caption, got {blanks}")
    if _caption_has_active_page_break(reread, "Таблица 1.1.3"):
        return _result(False, "table-start orphan repair should not use pageBreakBefore")
    return _result(True, "table-start orphan moved by two blank paragraphs")


def test_table_start_orphan_keeps_exactly_two_blanks_with_existing_separator() -> tuple[bool, str]:
    changed, reread = _run_with_rendered_lines(
        _add_table_start_orphan_doc(preexisting_blank=True),
        _orphan_lines(),
        after_lines=_repaired_lines(),
    )

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    blanks = _blank_count_before_caption(reread, "Таблица 1.1.3")
    if blanks != 2:
        return _result(False, f"expected exactly two blank paragraphs before caption, got {blanks}")
    return _result(True, "existing separator normalized to exactly two blank paragraphs")


def test_table_start_orphan_does_not_split_or_add_marker() -> tuple[bool, str]:
    changed, reread = _run_with_rendered_lines(
        _add_table_start_orphan_doc(),
        _orphan_lines(),
        after_lines=_repaired_lines(),
    )

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if len(reread.tables) != 1:
        return _result(False, f"repair must not split the table, got {len(reread.tables)} tables")
    markers = _continuation_markers(reread)
    if markers:
        return _result(False, f"repair inserted continuation marker(s): {markers!r}")
    return _result(True, "table-start orphan repair does not split or add marker")


def test_table_start_orphan_does_not_synthesize_numeric_row() -> tuple[bool, str]:
    changed, reread = _run_with_rendered_lines(
        _add_table_start_orphan_doc(),
        _orphan_lines(),
        after_lines=_repaired_lines(),
    )

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    if _numeric_row_count(reread) != 0:
        return _result(False, "ordinary non-split table received synthetic numeric row")
    return _result(True, "ordinary non-split table keeps no numeric row")


def test_table_start_orphan_preserves_existing_numeric_row() -> tuple[bool, str]:
    changed, reread = _run_with_rendered_lines(
        _add_table_start_orphan_doc(with_numeric=True),
        _orphan_lines(with_numeric=True),
        after_lines=_repaired_lines(with_numeric=True),
    )

    if changed != 1:
        return _result(False, f"expected repair, got {changed}")
    numeric_count = _numeric_row_count(reread)
    if numeric_count != 1:
        return _result(False, f"expected existing numeric row preserved once, got {numeric_count}")
    return _result(True, "source numeric row is preserved and not duplicated")


def test_table_start_orphan_skips_without_caption_identity() -> tuple[bool, str]:
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Показатель"
    table.rows[0].cells[1].text = "Статус"
    table.rows[1].cells[0].text = "Кейс 1"
    table.rows[1].cells[1].text = "тест"

    changed, reread = _run_with_rendered_lines(
        doc,
        [
            PdfLine("Показатель Статус", 1, 700.0, 714.0),
            PdfLine("Кейс 1 тест", 2, 60.0, 74.0),
        ],
        name="no_caption.docx",
    )

    if changed != 0:
        return _result(False, f"uncaptioned table-start shape should skip safely, got {changed}")
    if len(reread.tables) != 1:
        return _result(False, f"uncaptioned skip should not split table, got {len(reread.tables)}")
    return _result(True, "uncaptioned table-start shape skipped")


def test_table_start_orphan_rolls_back_when_validation_fails() -> tuple[bool, str]:
    doc = _add_table_start_orphan_doc()
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "rollback.docx"
        pdf_dir = Path(tmp) / "pdf"
        pdf_dir.mkdir()
        pdf_path = pdf_dir / "case.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n")
        doc.save(path)

        old_render = tc.render_docx_to_pdf
        old_analyze = tc.analyze_pdf_lines
        try:
            tc.render_docx_to_pdf = lambda _path: pdf_path
            calls = {"n": 0}

            def fake_analyze(_path):
                calls["n"] += 1
                return _orphan_lines()

            tc.analyze_pdf_lines = fake_analyze
            changed = tc.apply_rendered_table_continuation(path)
        finally:
            tc.render_docx_to_pdf = old_render
            tc.analyze_pdf_lines = old_analyze

        reread = Document(str(path))

    if changed != 0:
        return _result(False, f"failed validation should roll back repair, got {changed}")
    blanks = _blank_count_before_caption(reread, "Таблица 1.1.3")
    if blanks != 0:
        return _result(False, f"rollback should keep original DOCX, found {blanks} inserted blanks")
    return _result(True, "failed post-render validation rolls back table-start repair")


def test_table_start_orphan_diagnostics_preempts_generic_cross_page() -> tuple[bool, str]:
    from tests.tools.table_engine_diagnostics import build_universal_table_diagnostics

    doc = _add_table_start_orphan_doc()
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "diagnostics.docx"
        doc.save(path)
        report = build_universal_table_diagnostics(
            formatted_docx=path,
            source_docx=path,
            pdf_lines=_orphan_lines(),
        )

    issue_types = [issue["issue_type"] for issue in report["issues"]]
    if "table_start_orphan" not in issue_types:
        return _result(False, f"missing table_start_orphan diagnostic: {issue_types!r}")
    if "single_physical_table_crosses_pages_without_marker" in issue_types:
        return _result(False, f"table-start orphan should preempt generic cross-page: {issue_types!r}")
    return _result(True, "diagnostics classify table-start orphan before generic cross-page")


def main() -> int:
    tests = [
        ("detect/repair inserts two blanks", test_table_start_orphan_inserts_two_blank_paragraphs),
        ("existing separator stays exactly two blanks", test_table_start_orphan_keeps_exactly_two_blanks_with_existing_separator),
        ("does not split or add marker", test_table_start_orphan_does_not_split_or_add_marker),
        ("does not synthesize numeric row", test_table_start_orphan_does_not_synthesize_numeric_row),
        ("preserves existing numeric row", test_table_start_orphan_preserves_existing_numeric_row),
        ("skips without caption identity", test_table_start_orphan_skips_without_caption_identity),
        ("rolls back on failed validation", test_table_start_orphan_rolls_back_when_validation_fails),
        ("diagnostics classify table-start orphan", test_table_start_orphan_diagnostics_preempts_generic_cross_page),
    ]
    failed = 0
    for name, fn in tests:
        ok, msg = fn()
        status = "PASS" if ok else "FAIL"
        print(f"[{status}] {name} — {msg}")
        if not ok:
            failed += 1
    return 1 if failed else 0


if __name__ == "__main__":
    raise SystemExit(main())
