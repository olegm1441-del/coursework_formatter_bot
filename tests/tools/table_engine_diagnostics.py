from __future__ import annotations

import argparse
import hashlib
import json
import re
import shutil
import sys
from dataclasses import asdict, dataclass
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


_CAPTION_RE = re.compile(
    r"^\s*Таблица\s+([0-9]+(?:\.[0-9]+)*|[А-ЯA-Z]\.\d+(?:\.\d+)*)\b",
    re.IGNORECASE,
)
_STRICT_MARKER_RE = re.compile(
    r"^\s*Продолжение\s+таблицы\s+([0-9]+(?:\.[0-9]+)*|[А-ЯA-Z]\.\d+(?:\.\d+)*)\s*$",
    re.IGNORECASE,
)
_SOURCE_NOTE_RE = re.compile(r"^\s*(Источник|Примечание)\s*:", re.IGNORECASE)


@dataclass(frozen=True)
class SplitDecision:
    eligible: bool
    reason: str


@dataclass(frozen=True)
class RowRole:
    row_index: int
    role: str


@dataclass(frozen=True)
class NumericRowSafety:
    safe: bool
    reason: str | None = None


@dataclass(frozen=True)
class ContinuationMarkerDiagnostic:
    text: str
    marker_page: int | None
    marker_kind: str
    verdict: str
    same_page_violation: bool | None
    previous_table_page: int | None
    following_table_page: int | None
    confidence: str


@dataclass(frozen=True)
class DiagnosticStage:
    name: str
    docx_path: str
    tables: list[dict]


@dataclass(frozen=True)
class DiagnosticRunResult:
    artifact_root: str
    summary_json_path: str
    stages: list[DiagnosticStage]


@dataclass(frozen=True)
class LogicalTable:
    logical_table_id: str
    table_num: str | None
    caption_paragraph_text: str | None
    caption_paragraph_index: int | None
    title_paragraph_text: str | None
    title_paragraph_index: int | None
    physical_table_indexes: list[int]
    continuation_marker_paragraphs: list[dict]
    source_note_paragraphs: list[dict]
    appendix_context: str | None
    source_provenance_available: bool


@dataclass(frozen=True)
class PhysicalTable:
    docx_table_index: int
    body_order_index: int
    logical_table_id: str
    column_count: int
    grid_signature: str
    width_signature: str
    border_signature: str
    cell_margin_signature: str
    has_merged_cells: bool
    row_count: int


@dataclass(frozen=True)
class TableRow:
    logical_table_id: str
    physical_table_index: int
    row_index: int
    cell_texts: list[str]
    normalized_fingerprint: str
    role: str
    provenance: str


@dataclass(frozen=True)
class RenderedRow:
    logical_table_id: str
    physical_table_index: int
    row_index: int
    page: int
    top_y: float | None
    bottom_y: float | None
    confidence: str
    split_or_spill: bool


@dataclass(frozen=True)
class TableIssue:
    issue_type: str
    table_num: str | None
    pages: list[int]
    evidence: dict
    severity: str
    safe_repair_class: str | None
    skip_reason: str | None = None


@dataclass(frozen=True)
class StyleNote:
    note_type: str
    table_num: str | None
    physical_table_index: int
    evidence: dict


def _cell_texts(row_xml) -> list[str]:
    out = []
    for cell in row_xml.findall(qn("w:tc")):
        text = " ".join((node.text or "") for node in cell.findall(".//" + qn("w:t")))
        out.append(" ".join(text.split()))
    return out


def _is_numeric_row(values: list[str]) -> bool:
    return len(values) >= 2 and values == [str(i) for i in range(1, len(values) + 1)]


def _normalize_text(text: str | None) -> str:
    return " ".join((text or "").replace("\xa0", " ").split()).lower()


def _row_fingerprint_from_cells(values: list[str]) -> str:
    return " ".join(_normalize_text(value) for value in values if _normalize_text(value))


def _text_tokens(text: str) -> set[str]:
    return {
        token
        for token in re.findall(r"[0-9a-zа-яё]+", _normalize_text(text))
        if len(token) > 1 or token.isdigit()
    }


def _token_overlap(left: str, right: str) -> float:
    left_tokens = _text_tokens(left)
    right_tokens = _text_tokens(right)
    if not left_tokens or not right_tokens:
        return 0.0
    return len(left_tokens & right_tokens) / max(len(left_tokens), len(right_tokens))


def _cell_line_score(cell_text: str, line_text: str) -> float:
    cell_tokens = _text_tokens(cell_text)
    line_tokens = _text_tokens(line_text)
    if not cell_tokens or not line_tokens:
        return 0.0
    if len(cell_tokens) <= 4:
        return len(cell_tokens & line_tokens) / len(cell_tokens)
    return len(cell_tokens & line_tokens) / max(len(cell_tokens), len(line_tokens))


def _caption_num(text: str | None) -> str | None:
    match = _CAPTION_RE.match(" ".join((text or "").split()))
    return match.group(1) if match else None


def _strict_marker(text: str | None) -> tuple[str, str] | None:
    cleaned = " ".join((text or "").split())
    match = _STRICT_MARKER_RE.match(cleaned)
    if not match:
        return None
    return cleaned, match.group(1)


def _is_source_note(text: str | None) -> bool:
    return bool(_SOURCE_NOTE_RE.match(" ".join((text or "").split())))


def _short_hash(text: str | None) -> str:
    if not text:
        return "none"
    return hashlib.sha1(text.encode("utf-8")).hexdigest()[:12]


def _element_xml_text(el) -> str | None:
    if el is None:
        return None
    xml = getattr(el, "xml", None)
    if xml is not None:
        return str(xml)
    return etree.tostring(el, encoding="unicode")


def _row_has_gridspan(row_xml) -> bool:
    return any(cell.find(qn("w:tcPr")) is not None and cell.find(qn("w:tcPr")).find(qn("w:gridSpan")) is not None for cell in row_xml.findall(qn("w:tc")))


def _row_has_vmerge_continue(row_xml) -> bool:
    for cell in row_xml.findall(qn("w:tc")):
        tc_pr = cell.find(qn("w:tcPr"))
        if tc_pr is None:
            continue
        vm = tc_pr.find(qn("w:vMerge"))
        if vm is not None and vm.get(qn("w:val")) in {None, "continue"}:
            return True
    return False


def classify_table_rows(tbl_xml, *, header_rows: int = 1) -> list[RowRole]:
    roles: list[RowRole] = []
    for idx, row in enumerate(tbl_xml.findall(qn("w:tr"))):
        values = _cell_texts(row)
        joined = " ".join(values).strip().lower()
        if _is_numeric_row(values):
            role = "numeric"
        elif joined.startswith(("источник:", "примечание:")):
            role = "source_note"
        elif idx < header_rows:
            role = "header"
        else:
            role = "body"
        roles.append(RowRole(idx, role))
    return roles


def numeric_row_synthesis_safety(tbl_xml, *, header_rows: int = 1) -> NumericRowSafety:
    rows = tbl_xml.findall(qn("w:tr"))
    if not rows or header_rows <= 0:
        return NumericRowSafety(False, "unsafe_numeric_row_synthesis")
    for row in rows[:header_rows]:
        if _row_has_gridspan(row):
            return NumericRowSafety(False, "unsafe_numeric_row_synthesis")
    return NumericRowSafety(True)


def evaluate_split_eligibility(
    tbl_xml,
    *,
    rendered_row_pages: dict[int, int] | None,
    split_before_row: int,
    header_rows: int = 1,
    min_body_rows_per_fragment: int = 1,
) -> SplitDecision:
    if rendered_row_pages is None:
        return SplitDecision(False, "render_boundary_unmapped")
    if tbl_xml.find(qn("w:tblGrid")) is None:
        return SplitDecision(False, "malformed_grid")
    rows = tbl_xml.findall(qn("w:tr"))
    if split_before_row < header_rows + min_body_rows_per_fragment:
        return SplitDecision(False, "fragment_too_small")
    if split_before_row >= len(rows):
        return SplitDecision(False, "fragment_too_small")
    if _row_has_vmerge_continue(rows[split_before_row]):
        return SplitDecision(False, "unsafe_vmerge_boundary")
    roles = classify_table_rows(tbl_xml, header_rows=header_rows)
    if roles[split_before_row].role == "source_note":
        return SplitDecision(False, "source_note_boundary_uncertain")
    first_pages = {rendered_row_pages.get(i) for i in range(0, split_before_row)}
    second_pages = {rendered_row_pages.get(i) for i in range(split_before_row, len(rows))}
    if None in first_pages or None in second_pages or not first_pages or not second_pages:
        return SplitDecision(False, "render_boundary_unmapped")
    if max(first_pages) >= min(second_pages):
        return SplitDecision(False, "render_boundary_unmapped")
    return SplitDecision(True, "eligible")


def _xml(el) -> str | None:
    return _element_xml_text(el)


def _tbl_pr(tbl_xml):
    return tbl_xml.find(qn("w:tblPr"))


def _tbl_w(tbl_xml) -> dict:
    tbl_pr = _tbl_pr(tbl_xml)
    node = tbl_pr.find(qn("w:tblW")) if tbl_pr is not None else None
    return {
        "w": node.get(qn("w:w")) if node is not None else None,
        "type": node.get(qn("w:type")) if node is not None else None,
    }


def _tbl_layout(tbl_xml) -> str | None:
    tbl_pr = _tbl_pr(tbl_xml)
    node = tbl_pr.find(qn("w:tblLayout")) if tbl_pr is not None else None
    return node.get(qn("w:type")) if node is not None else None


def _grid_widths(tbl_xml) -> list[str | None]:
    grid = tbl_xml.find(qn("w:tblGrid"))
    if grid is None:
        return []
    return [col.get(qn("w:w")) for col in grid.findall(qn("w:gridCol"))]


def _tcw(tbl_xml) -> list[dict]:
    out = []
    for cell in tbl_xml.findall(".//" + qn("w:tc")):
        tc_pr = cell.find(qn("w:tcPr"))
        node = tc_pr.find(qn("w:tcW")) if tc_pr is not None else None
        out.append({
            "w": node.get(qn("w:w")) if node is not None else None,
            "type": node.get(qn("w:type")) if node is not None else None,
        })
    return out


def _geometry_policy(tbl_xml) -> str:
    try:
        from guides.coursework_kfu_2025.table_continuation import classify_table_geometry_policy

        return classify_table_geometry_policy(tbl_xml)
    except Exception:
        return "unknown"


def snapshot_table(tbl_xml, *, table_index: int = 0) -> dict:
    rows = tbl_xml.findall(qn("w:tr"))
    return {
        "table_index": table_index,
        "tblPr_xml": _xml(_tbl_pr(tbl_xml)),
        "tblGrid_xml": _xml(tbl_xml.find(qn("w:tblGrid"))),
        "gridCol_widths": _grid_widths(tbl_xml),
        "row_count": len(rows),
        "raw_tc_counts": [len(row.findall(qn("w:tc"))) for row in rows],
        "tblW": _tbl_w(tbl_xml),
        "tblLayout": _tbl_layout(tbl_xml),
        "tcW": _tcw(tbl_xml),
        "gridSpan": [cell.find(qn("w:tcPr")).find(qn("w:gridSpan")).get(qn("w:val"))
                     for cell in tbl_xml.findall(".//" + qn("w:tc"))
                     if cell.find(qn("w:tcPr")) is not None and cell.find(qn("w:tcPr")).find(qn("w:gridSpan")) is not None],
        "vMerge": [cell.find(qn("w:tcPr")).find(qn("w:vMerge")).get(qn("w:val"))
                   for cell in tbl_xml.findall(".//" + qn("w:tc"))
                   if cell.find(qn("w:tcPr")) is not None and cell.find(qn("w:tcPr")).find(qn("w:vMerge")) is not None],
        "geometry_policy": _geometry_policy(tbl_xml),
    }


def diff_table_geometry(before: dict, after: dict) -> dict:
    keys = (
        "tblPr_xml", "tblGrid_xml", "gridCol_widths", "row_count",
        "raw_tc_counts", "tblW", "tblLayout", "tcW", "gridSpan", "vMerge",
        "geometry_policy",
    )
    return {key: {"before": before.get(key), "after": after.get(key)} for key in keys if before.get(key) != after.get(key)}


def _paragraph_maps(doc: Document):
    para_by_xml = {p._element: p for p in doc.paragraphs}
    para_index_by_xml = {p._element: idx for idx, p in enumerate(doc.paragraphs)}
    return para_by_xml, para_index_by_xml


def _body_table_order(doc: Document) -> dict[object, int]:
    out: dict[object, int] = {}
    counter = 0
    for child in doc.element.body:
        if child.tag == qn("w:tbl"):
            out[child] = counter
            counter += 1
    return out


def _body_table_indexes(doc: Document) -> dict[object, int]:
    return {table._tbl: idx for idx, table in enumerate(doc.tables)}


def _appendix_context(table_num: str | None) -> str | None:
    if table_num and re.match(r"^[А-ЯA-Z]\.", table_num, flags=re.IGNORECASE):
        return "appendix"
    return None


def _grid_signature(tbl_xml) -> str:
    widths = _grid_widths(tbl_xml)
    return "|".join(width or "" for width in widths)


def _width_signature(tbl_xml) -> str:
    tbl_w = _tbl_w(tbl_xml)
    layout = _tbl_layout(tbl_xml) or ""
    first_row = tbl_xml.find(qn("w:tr"))
    tcw_values: list[str] = []
    if first_row is not None:
        for cell in first_row.findall(qn("w:tc")):
            tc_pr = cell.find(qn("w:tcPr"))
            node = tc_pr.find(qn("w:tcW")) if tc_pr is not None else None
            tcw_values.append(
                f"{node.get(qn('w:w')) if node is not None else ''}:"
                f"{node.get(qn('w:type')) if node is not None else ''}"
            )
    return json.dumps(
        {"tblW": tbl_w, "layout": layout, "first_row_tcW": tcw_values},
        ensure_ascii=False,
        sort_keys=True,
    )


def _border_signature(tbl_xml) -> str:
    tbl_pr = _tbl_pr(tbl_xml)
    borders = tbl_pr.find(qn("w:tblBorders")) if tbl_pr is not None else None
    return _short_hash(_element_xml_text(borders))


def _cell_margin_signature(tbl_xml) -> str:
    tbl_pr = _tbl_pr(tbl_xml)
    margins = tbl_pr.find(qn("w:tblCellMar")) if tbl_pr is not None else None
    return _short_hash(_element_xml_text(margins))


def _has_merged_cells(tbl_xml) -> bool:
    for cell in tbl_xml.findall(".//" + qn("w:tc")):
        tc_pr = cell.find(qn("w:tcPr"))
        if tc_pr is None:
            continue
        if tc_pr.find(qn("w:gridSpan")) is not None or tc_pr.find(qn("w:vMerge")) is not None:
            return True
    return False


def _source_row_counts_by_table(source_docx: Path | None) -> dict[str, dict[str, int]]:
    if source_docx is None:
        return {}
    doc = Document(str(source_docx))
    logicals, _, rows = _build_docx_model_parts(doc, source_counts_by_num=None, source_available=False)
    table_num_by_id = {logical.logical_table_id: logical.table_num for logical in logicals}
    out: dict[str, dict[str, int]] = {}
    for row in rows:
        table_num = table_num_by_id.get(row.logical_table_id)
        if not table_num or not row.normalized_fingerprint:
            continue
        out.setdefault(table_num, {})
        out[table_num][row.normalized_fingerprint] = out[table_num].get(row.normalized_fingerprint, 0) + 1
    return out


def _row_provenance(
    *,
    table_num: str | None,
    fingerprint: str,
    role: str,
    source_counts_by_num: dict[str, dict[str, int]] | None,
    source_available: bool,
) -> str:
    if not source_available:
        return "unknown"
    if table_num and source_counts_by_num and source_counts_by_num.get(table_num, {}).get(fingerprint, 0) > 0:
        return "source"
    if role in {"duplicate_header_artifact", "duplicate_numeric_artifact"}:
        return "continuation_artifact"
    if role == "numeric_row":
        return "formatter_generated"
    return "unknown"


def _classify_logical_rows(
    doc: Document,
    logicals: list[LogicalTable],
    source_counts_by_num: dict[str, dict[str, int]] | None,
    *,
    source_available: bool,
) -> list[TableRow]:
    table_num_by_id = {logical.logical_table_id: logical.table_num for logical in logicals}
    table_position_by_index: dict[int, int] = {}
    first_header_by_id: dict[str, str] = {}
    first_numeric_by_id: dict[str, str] = {}
    for logical in logicals:
        for pos, table_index in enumerate(logical.physical_table_indexes):
            table_position_by_index[table_index] = pos
            if table_index >= len(doc.tables):
                continue
            for row_xml in doc.tables[table_index]._tbl.findall(qn("w:tr")):
                cells = _cell_texts(row_xml)
                fp = _row_fingerprint_from_cells(cells)
                if not fp:
                    continue
                if _is_numeric_row(cells):
                    first_numeric_by_id.setdefault(logical.logical_table_id, fp)
                else:
                    first_header_by_id.setdefault(logical.logical_table_id, fp)
                    break

    rows: list[TableRow] = []
    logical_id_by_table_index = {
        table_index: logical.logical_table_id
        for logical in logicals
        for table_index in logical.physical_table_indexes
    }
    for table_index, table in enumerate(doc.tables):
        logical_id = logical_id_by_table_index.get(table_index, f"anonymous_{table_index}")
        table_num = table_num_by_id.get(logical_id)
        physical_pos = table_position_by_index.get(table_index, 0)
        first_header = first_header_by_id.get(logical_id)
        first_numeric = first_numeric_by_id.get(logical_id)
        for row_index, row_xml in enumerate(table._tbl.findall(qn("w:tr"))):
            cells = _cell_texts(row_xml)
            fp = _row_fingerprint_from_cells(cells)
            if not fp:
                role = "ambiguous_content"
            elif _is_numeric_row(cells):
                role = (
                    "duplicate_numeric_artifact"
                    if physical_pos > 0 and first_numeric and fp == first_numeric
                    else "numeric_row"
                )
            elif physical_pos > 0 and first_header and _token_overlap(fp, first_header) >= 0.85:
                role = "duplicate_header_artifact"
            elif row_index == 0:
                role = "semantic_header"
            elif _is_source_note(" ".join(cells)):
                role = "ambiguous_content"
            else:
                role = "data_row"
            rows.append(
                TableRow(
                    logical_table_id=logical_id,
                    physical_table_index=table_index,
                    row_index=row_index,
                    cell_texts=cells,
                    normalized_fingerprint=fp,
                    role=role,
                    provenance=_row_provenance(
                        table_num=table_num,
                        fingerprint=fp,
                        role=role,
                        source_counts_by_num=source_counts_by_num,
                        source_available=source_available,
                    ),
                )
            )
    return rows


def _build_docx_model_parts(
    doc: Document,
    *,
    source_counts_by_num: dict[str, dict[str, int]] | None,
    source_available: bool,
) -> tuple[list[LogicalTable], list[PhysicalTable], list[TableRow]]:
    para_by_xml, para_index_by_xml = _paragraph_maps(doc)
    table_index_by_xml = _body_table_indexes(doc)
    body_table_order = _body_table_order(doc)

    logical_builders: list[dict] = []
    active: dict | None = None
    active_has_table = False
    active_has_title = False

    def new_logical(table_num: str | None, text: str | None, para_index: int | None) -> dict:
        logical_id = f"table_{table_num}" if table_num else f"anonymous_{len(logical_builders)}"
        if any(item["logical_table_id"] == logical_id for item in logical_builders):
            logical_id = f"{logical_id}_{len(logical_builders)}"
        item = {
            "logical_table_id": logical_id,
            "table_num": table_num,
            "caption_paragraph_text": text,
            "caption_paragraph_index": para_index,
            "title_paragraph_text": None,
            "title_paragraph_index": None,
            "physical_table_indexes": [],
            "continuation_marker_paragraphs": [],
            "source_note_paragraphs": [],
            "appendix_context": _appendix_context(table_num),
            "source_provenance_available": source_available,
        }
        logical_builders.append(item)
        return item

    for child in doc.element.body:
        if child.tag == qn("w:p"):
            para = para_by_xml.get(child)
            text = " ".join(((para.text if para is not None else "") or "").split())
            para_index = para_index_by_xml.get(child)
            if not text:
                continue
            table_num = _caption_num(text)
            if table_num:
                active = new_logical(table_num, text, para_index)
                active_has_table = False
                active_has_title = False
                continue
            marker = _strict_marker(text)
            if marker:
                marker_text, marker_table_num = marker
                if active is None or (
                    active.get("table_num") not in {None, marker_table_num}
                    and active.get("physical_table_indexes")
                ):
                    active = next(
                        (
                            item for item in reversed(logical_builders)
                            if item.get("table_num") == marker_table_num
                        ),
                        None,
                    )
                    if active is None:
                        active = new_logical(marker_table_num, None, None)
                active["continuation_marker_paragraphs"].append(
                    {"paragraph_index": para_index, "text": marker_text}
                )
                active_has_table = bool(active["physical_table_indexes"])
                active_has_title = bool(active.get("title_paragraph_text"))
                continue
            if active is not None and _is_source_note(text) and active_has_table:
                active["source_note_paragraphs"].append(
                    {"paragraph_index": para_index, "text": text}
                )
                continue
            if active is not None and not active_has_table and not active_has_title:
                active["title_paragraph_text"] = text
                active["title_paragraph_index"] = para_index
                active_has_title = True
                continue
            if active_has_table:
                active = None
                active_has_table = False
                active_has_title = False
            continue

        if child.tag != qn("w:tbl"):
            continue
        table_index = table_index_by_xml.get(child)
        if table_index is None:
            continue
        if active is None:
            active = new_logical(None, None, None)
            active_has_title = False
        active["physical_table_indexes"].append(table_index)
        active_has_table = True

    logicals = [LogicalTable(**item) for item in logical_builders]
    logical_id_by_table_index = {
        table_index: logical.logical_table_id
        for logical in logicals
        for table_index in logical.physical_table_indexes
    }

    physicals: list[PhysicalTable] = []
    for idx, table in enumerate(doc.tables):
        tbl_xml = table._tbl
        rows = tbl_xml.findall(qn("w:tr"))
        col_count = len(rows[0].findall(qn("w:tc"))) if rows else 0
        physicals.append(
            PhysicalTable(
                docx_table_index=idx,
                body_order_index=body_table_order.get(tbl_xml, idx),
                logical_table_id=logical_id_by_table_index.get(idx, f"anonymous_{idx}"),
                column_count=col_count,
                grid_signature=_grid_signature(tbl_xml),
                width_signature=_width_signature(tbl_xml),
                border_signature=_border_signature(tbl_xml),
                cell_margin_signature=_cell_margin_signature(tbl_xml),
                has_merged_cells=_has_merged_cells(tbl_xml),
                row_count=len(rows),
            )
        )

    rows = _classify_logical_rows(
        doc,
        logicals,
        source_counts_by_num,
        source_available=source_available,
    )
    return logicals, physicals, rows


def build_docx_table_model(
    formatted_docx: Path,
    *,
    source_docx: Path | None = None,
) -> dict:
    source_counts = _source_row_counts_by_table(source_docx)
    doc = Document(str(formatted_docx))
    logicals, physicals, rows = _build_docx_model_parts(
        doc,
        source_counts_by_num=source_counts,
        source_available=source_docx is not None,
    )
    return {
        "logical_tables": [asdict(item) for item in logicals],
        "physical_tables": [asdict(item) for item in physicals],
        "rows": [asdict(item) for item in rows],
    }


def _rendered_rows_from_pdf_lines(model: dict, pdf_lines: list) -> list[RenderedRow]:
    by_page: dict[int, str] = {}
    lines_by_page: dict[int, list] = {}
    for line in pdf_lines:
        page = int(getattr(line, "page_num", 0) or 0)
        by_page.setdefault(page, "")
        by_page[page] = f"{by_page[page]} {_normalize_text(getattr(line, 'text', ''))}".strip()
        lines_by_page.setdefault(page, []).append(line)

    rendered: list[RenderedRow] = []
    for row in model["rows"]:
        fp = row["normalized_fingerprint"]
        if not fp:
            continue
        numeric_exact = bool(re.match(r"^(?:\d+\s+)+\d+$", fp))
        page_scores: list[tuple[int, float]] = []
        for page, page_text in by_page.items():
            if numeric_exact:
                score = 1.0 if any(_normalize_text(getattr(line, "text", "")) == fp for line in lines_by_page[page]) else 0.0
            else:
                cell_scores = [
                    _cell_line_score(cell, getattr(line, "text", ""))
                    for cell in row.get("cell_texts", [])
                    if len(_text_tokens(cell)) >= 2
                    for line in lines_by_page[page]
                ]
                if cell_scores:
                    score = max(cell_scores)
                else:
                    score = _token_overlap(fp, page_text)
            if score >= 0.55:
                page_scores.append((page, score))
        if not page_scores:
            continue
        max_score = max(score for _, score in page_scores)
        matched_pages = [
            (page, score)
            for page, score in page_scores
            if score >= max_score - 0.02 and score >= 0.55
        ]
        for page, score in matched_pages:
            row_tokens = _text_tokens(fp)
            candidate_lines = [
                line for line in lines_by_page.get(page, [])
                if row_tokens & _text_tokens(getattr(line, "text", ""))
            ]
            top_y = min((float(getattr(line, "top", 0.0)) for line in candidate_lines), default=None)
            bottom_y = max((float(getattr(line, "bottom", 0.0)) for line in candidate_lines), default=None)
            rendered.append(
                RenderedRow(
                    logical_table_id=row["logical_table_id"],
                    physical_table_index=row["physical_table_index"],
                    row_index=row["row_index"],
                    page=page,
                    top_y=top_y,
                    bottom_y=bottom_y,
                    confidence="high" if score >= 0.8 else "medium",
                    split_or_spill=len(matched_pages) > 1,
                )
            )
    return rendered


def _logical_by_id(model: dict) -> dict[str, dict]:
    return {item["logical_table_id"]: item for item in model["logical_tables"]}


def _rows_for_logical(model: dict, logical_id: str) -> list[dict]:
    return [row for row in model["rows"] if row["logical_table_id"] == logical_id]


def _rendered_pages_for_logical(rendered_rows: list[dict], logical_id: str) -> list[int]:
    return sorted({int(row["page"]) for row in rendered_rows if row["logical_table_id"] == logical_id and row.get("page")})


def _caption_page_bounds(pdf_lines: list) -> dict[str, tuple[int, int | None]]:
    captions: list[tuple[int, float, str]] = []
    for line in pdf_lines:
        table_num = _caption_num(getattr(line, "text", ""))
        if not table_num:
            continue
        captions.append((
            int(getattr(line, "page_num", 0) or 0),
            float(getattr(line, "top", 0.0) or 0.0),
            table_num,
        ))
    captions.sort()
    bounds: dict[str, tuple[int, int | None]] = {}
    for idx, (page, _top, table_num) in enumerate(captions):
        if table_num in bounds:
            continue
        next_page = None
        for later_page, _later_top, later_num in captions[idx + 1:]:
            if later_num != table_num:
                next_page = later_page
                break
        bounds[table_num] = (page, next_page)
    return bounds


def _strict_marker_pages(pdf_lines: list) -> dict[str, list[int]]:
    pages: dict[str, set[int]] = {}
    for line in pdf_lines:
        marker = _strict_marker(getattr(line, "text", ""))
        if not marker:
            continue
        _text, table_num = marker
        pages.setdefault(table_num, set()).add(int(getattr(line, "page_num", 0) or 0))
    return {table_num: sorted(values) for table_num, values in pages.items()}


def _bounded_rendered_rows(
    rendered_rows: list[dict],
    logical_id: str,
    table_num: str | None,
    page_bounds: dict[str, tuple[int, int | None]] | None,
) -> list[dict]:
    rows = [row for row in rendered_rows if row["logical_table_id"] == logical_id and row.get("page")]
    if not table_num or not page_bounds or table_num not in page_bounds:
        return rows
    start, end = page_bounds[table_num]
    return [
        row for row in rows
        if int(row["page"]) >= start and (end is None or int(row["page"]) < end)
    ]


def _bounded_pages(
    rendered_rows: list[dict],
    logical_id: str,
    table_num: str | None,
    page_bounds: dict[str, tuple[int, int | None]] | None,
) -> list[int]:
    return sorted({int(row["page"]) for row in _bounded_rendered_rows(rendered_rows, logical_id, table_num, page_bounds)})


def _issue_from_rendered_violation(violation) -> TableIssue:
    issue_type = violation.violation_type
    skip_reason = None
    safe_repair_class = None
    severity = "warning"
    if issue_type == "same_page_adjacent_fragment":
        issue_type = "ambiguous_table_chain"
        skip_reason = "same_page_adjacency_without_repeated_artifact"
    elif issue_type == "same_page_repeated_fragment":
        safe_repair_class = "same_page_fragment_normalize"
        severity = "fail"
    elif issue_type == "source_bad_duplicated_content_rows":
        safe_repair_class = None
        skip_reason = "source_content_proven"
        severity = "warning"
    elif issue_type == "late_continuation_marker":
        safe_repair_class = "continuation_marker_reposition"
        severity = "fail"
    elif issue_type in {"ambiguous_adjacent_tables", "suspected_missing_continuation_marker"}:
        issue_type = "ambiguous_table_chain"
        skip_reason = "rendered_or_docx_evidence_incomplete"
        severity = "warning"
    elif issue_type in {"missing_continuation_marker", "missing_or_late_continuation_marker"}:
        safe_repair_class = "rendered_continuation_repair"
        severity = "fail"
    return TableIssue(
        issue_type=issue_type,
        table_num=violation.table_num,
        pages=[violation.page] if violation.page else [],
        evidence={
            **dict(violation.evidence),
            "original_violation_type": violation.violation_type,
            "table_index": violation.table_index,
        },
        severity=severity,
        safe_repair_class=safe_repair_class,
        skip_reason=skip_reason,
    )


def _detect_missing_numeric_rows(
    model: dict,
    rendered_rows: list[dict],
    page_bounds: dict[str, tuple[int, int | None]] | None = None,
    marker_pages: dict[str, list[int]] | None = None,
) -> list[TableIssue]:
    issues: list[TableIssue] = []
    for logical in model["logical_tables"]:
        if not logical["continuation_marker_paragraphs"]:
            continue
        missing: list[int] = []
        for table_index in logical["physical_table_indexes"]:
            table_rows = [row for row in model["rows"] if row["physical_table_index"] == table_index]
            has_numeric = any(row["role"] in {"numeric_row", "duplicate_numeric_artifact"} for row in table_rows)
            if not has_numeric:
                missing.append(table_index)
        if not missing:
            continue
        pages = _bounded_pages(rendered_rows, logical["logical_table_id"], logical["table_num"], page_bounds)
        if logical["table_num"] and marker_pages:
            pages = sorted(set(pages) | set(marker_pages.get(logical["table_num"], [])))
        issues.append(
            TableIssue(
                issue_type="continuation_missing_numeric_row",
                table_num=logical["table_num"],
                pages=pages,
                evidence={
                    "missing_numeric_physical_table_indexes": missing,
                    "marker_paragraphs": logical["continuation_marker_paragraphs"],
                },
                severity="fail",
                safe_repair_class="numeric_row_synthesis",
                skip_reason=None,
            )
        )
    return issues


def _line_matches_row_cells(row: dict, line_texts: list[str]) -> bool:
    joined = " ".join(line_texts)
    cell_texts = [text for text in row.get("cell_texts", []) if _normalize_text(text)]
    if not cell_texts:
        return False
    if _row_fingerprint_from_cells(cell_texts) and _row_fingerprint_from_cells(cell_texts) in joined:
        return True
    hits = 0
    for cell_text in cell_texts:
        if _cell_line_score(cell_text, joined) >= 0.75:
            hits += 1
    required = min(2, len(cell_texts))
    return hits >= required


def _detect_table_start_orphans(
    model: dict,
    pdf_lines: list,
    page_bounds: dict[str, tuple[int, int | None]] | None = None,
) -> list[TableIssue]:
    issues: list[TableIssue] = []
    if not pdf_lines:
        return issues

    lines_by_page: dict[int, list] = {}
    for line in pdf_lines:
        lines_by_page.setdefault(int(getattr(line, "page_num", 0) or 0), []).append(line)

    rows_by_logical: dict[str, list[dict]] = {}
    for row in model["rows"]:
        rows_by_logical.setdefault(row["logical_table_id"], []).append(row)

    for logical in model["logical_tables"]:
        table_num = logical["table_num"]
        if not table_num:
            continue
        if len(logical["physical_table_indexes"]) != 1 or logical["continuation_marker_paragraphs"]:
            continue

        caption_lines = [
            line for line in pdf_lines
            if _caption_num(getattr(line, "text", "")) == table_num
        ]
        if len(caption_lines) != 1:
            continue

        caption_line = caption_lines[0]
        start_page = int(getattr(caption_line, "page_num", 0) or 0)
        if page_bounds and table_num in page_bounds:
            bounded_start, _bounded_end = page_bounds[table_num]
            if start_page != bounded_start:
                continue

        same_page_lines = [
            line for line in lines_by_page.get(start_page, [])
            if float(getattr(line, "top", 0.0) or 0.0) > float(getattr(caption_line, "top", 0.0) or 0.0)
        ]
        same_page_texts = [_normalize_text(getattr(line, "text", "")) for line in same_page_lines]
        if not same_page_texts:
            continue

        logical_rows = rows_by_logical.get(logical["logical_table_id"], [])
        header_rows = [
            row for row in logical_rows
            if row["role"] in {"semantic_header", "numeric_row", "duplicate_numeric_artifact"}
        ]
        data_rows = [row for row in logical_rows if row["role"] == "data_row"]
        if not header_rows or not data_rows:
            continue
        if not any(_line_matches_row_cells(row, same_page_texts[:8]) for row in header_rows):
            continue
        if any(_line_matches_row_cells(row, same_page_texts) for row in data_rows):
            continue

        next_page = start_page + 1
        next_page_lines = [
            _normalize_text(getattr(line, "text", ""))
            for line in lines_by_page.get(next_page, [])
            if float(getattr(line, "top", 0.0) or 0.0) <= 220.0
        ]
        first_data_row = next((row for row in data_rows if _line_matches_row_cells(row, next_page_lines)), None)
        if first_data_row is None:
            continue

        issues.append(
            TableIssue(
                issue_type="table_start_orphan",
                table_num=table_num,
                pages=[start_page, next_page],
                evidence={
                    "logical_table_id": logical["logical_table_id"],
                    "physical_table_index": logical["physical_table_indexes"][0],
                    "caption_page": start_page,
                    "first_data_page": next_page,
                    "first_data_row_index": first_data_row["row_index"],
                    "first_data_row_fingerprint": first_data_row["normalized_fingerprint"][:90],
                    "repair": "insert_two_blank_paragraphs_before_caption",
                },
                severity="fail",
                safe_repair_class="table_start_orphan_move",
                skip_reason=None,
            )
        )
    return issues


def _detect_single_physical_cross_page(
    model: dict,
    rendered_rows: list[dict],
    page_bounds: dict[str, tuple[int, int | None]] | None = None,
    skip_logical_ids: set[str] | None = None,
) -> list[TableIssue]:
    issues: list[TableIssue] = []
    skip_logical_ids = skip_logical_ids or set()
    for logical in model["logical_tables"]:
        if logical["logical_table_id"] in skip_logical_ids:
            continue
        table_num = logical["table_num"]
        if not table_num:
            continue
        if len(logical["physical_table_indexes"]) != 1 or logical["continuation_marker_paragraphs"]:
            continue
        rows = [
            row for row in _bounded_rendered_rows(rendered_rows, logical["logical_table_id"], table_num, page_bounds)
            if not row.get("split_or_spill")
        ]
        pages = sorted({int(row["page"]) for row in rows})
        if len(pages) < 3:
            continue
        if len(pages) > 4 or pages != list(range(pages[0], pages[-1] + 1)):
            continue
        min_row_by_page = {
            page: min(int(row["row_index"]) for row in rows if int(row["page"]) == page)
            for page in pages
        }
        if [min_row_by_page[page] for page in pages] != sorted(min_row_by_page.values()):
            continue
        issues.append(
            TableIssue(
                issue_type="single_physical_table_crosses_pages_without_marker",
                table_num=table_num,
                pages=pages,
                evidence={
                    "physical_table_index": logical["physical_table_indexes"][0],
                    "rendered_pages": pages,
                },
                severity="fail",
                safe_repair_class="rendered_cross_page_split",
                skip_reason=None,
            )
        )
    return issues


def _detect_ordinary_generated_numeric_row(
    model: dict,
    rendered_rows: list[dict],
    page_bounds: dict[str, tuple[int, int | None]] | None = None,
) -> list[TableIssue]:
    issues: list[TableIssue] = []
    for logical in model["logical_tables"]:
        if len(logical["physical_table_indexes"]) != 1 or logical["continuation_marker_paragraphs"]:
            continue
        pages = _bounded_pages(rendered_rows, logical["logical_table_id"], logical["table_num"], page_bounds)
        if len(pages) > 1:
            continue
        for row in _rows_for_logical(model, logical["logical_table_id"]):
            if row["role"] != "numeric_row" or row["provenance"] != "formatter_generated":
                continue
            issues.append(
                TableIssue(
                    issue_type="ordinary_table_has_generated_numeric_row_after_merge",
                    table_num=logical["table_num"],
                    pages=pages,
                    evidence={
                        "physical_table_index": row["physical_table_index"],
                        "row_index": row["row_index"],
                        "row_fingerprint": row["normalized_fingerprint"],
                        "provenance": row["provenance"],
                    },
                    severity="warning",
                    safe_repair_class="numeric_artifact_cleanup_after_merge",
                    skip_reason=None,
                )
            )
            break
    return issues


def _detect_source_bad_duplicates(model: dict) -> list[TableIssue]:
    logical_by_id = _logical_by_id(model)
    issues: list[TableIssue] = []
    for logical in model["logical_tables"]:
        counts: dict[str, int] = {}
        for row in _rows_for_logical(model, logical["logical_table_id"]):
            if row["role"] != "data_row" or row["provenance"] != "source":
                continue
            fp = row["normalized_fingerprint"]
            counts[fp] = counts.get(fp, 0) + 1
        duplicates = [fp for fp, count in counts.items() if count > 1]
        if not duplicates:
            continue
        issues.append(
            TableIssue(
                issue_type="source_bad_duplicated_content_rows",
                table_num=logical_by_id[logical["logical_table_id"]]["table_num"],
                pages=[],
                evidence={
                    "duplicate_row_count": len(duplicates),
                    "row_fingerprint": duplicates[0][:90],
                    "source_proven": True,
                },
                severity="warning",
                safe_repair_class=None,
                skip_reason="source_content_proven",
            )
        )
    return issues


def _style_notes(model: dict) -> list[StyleNote]:
    logical_by_id = _logical_by_id(model)
    notes: list[StyleNote] = []
    for physical in model["physical_tables"]:
        # A missing explicit cell margin is a visual/readability note only in
        # this diagnostics phase. It is not a hard validator failure.
        if physical.get("cell_margin_signature") != "none":
            continue
        logical = logical_by_id.get(physical["logical_table_id"], {})
        notes.append(
            StyleNote(
                note_type="table_cell_padding_not_explicit",
                table_num=logical.get("table_num"),
                physical_table_index=physical["docx_table_index"],
                evidence={"target_left_right_margin_cm": 0.3},
            )
        )
    return notes


def build_universal_table_diagnostics(
    *,
    formatted_docx: Path,
    pdf: Path | None = None,
    source_docx: Path | None = None,
    pdf_lines: list | None = None,
    rendered_rows: list[dict] | None = None,
) -> dict:
    model = build_docx_table_model(formatted_docx, source_docx=source_docx)

    if pdf_lines is None and pdf is not None:
        from guides.coursework_kfu_2025.pdf_layout_analyzer import analyze_pdf_lines

        pdf_lines = analyze_pdf_lines(pdf)
    pdf_lines = pdf_lines or []

    if rendered_rows is None:
        rendered = _rendered_rows_from_pdf_lines(model, pdf_lines)
        rendered_rows_payload = [asdict(item) for item in rendered]
    else:
        rendered_rows_payload = rendered_rows
    page_bounds = _caption_page_bounds(pdf_lines) if pdf_lines else None
    marker_pages = _strict_marker_pages(pdf_lines) if pdf_lines else None

    issues: list[TableIssue] = []
    if pdf_lines:
        from guides.coursework_kfu_2025.rendered_table_validation import (
            build_rendered_table_identities,
            validate_rendered_continuations,
        )

        formatted_doc = Document(str(formatted_docx))
        source_identities = None
        if source_docx is not None:
            source_identities = build_rendered_table_identities(Document(str(source_docx)))
        for violation in validate_rendered_continuations(
            pdf_lines,
            build_rendered_table_identities(formatted_doc),
            source_table_identities=source_identities,
        ):
            issues.append(_issue_from_rendered_violation(violation))

    issues.extend(_detect_missing_numeric_rows(model, rendered_rows_payload, page_bounds, marker_pages))
    table_start_orphans = _detect_table_start_orphans(model, pdf_lines, page_bounds)
    issues.extend(table_start_orphans)
    table_start_orphan_ids = {
        issue.evidence.get("logical_table_id")
        for issue in table_start_orphans
        if issue.evidence.get("logical_table_id")
    }
    issues.extend(
        _detect_single_physical_cross_page(
            model,
            rendered_rows_payload,
            page_bounds,
            skip_logical_ids=table_start_orphan_ids,
        )
    )
    issues.extend(_detect_ordinary_generated_numeric_row(model, rendered_rows_payload, page_bounds))
    issues.extend(_detect_source_bad_duplicates(model))

    seen: set[tuple] = set()
    deduped: list[TableIssue] = []
    for issue in issues:
        if issue.issue_type == "source_bad_duplicated_content_rows":
            key = (issue.issue_type, issue.table_num)
        else:
            key = (
                issue.issue_type,
                issue.table_num,
                tuple(issue.pages),
                issue.safe_repair_class,
                issue.skip_reason,
                json.dumps(issue.evidence, ensure_ascii=False, sort_keys=True),
            )
        if key in seen:
            continue
        seen.add(key)
        deduped.append(issue)
    if not deduped:
        deduped.append(
            TableIssue(
                issue_type="clean",
                table_num=None,
                pages=[],
                evidence={},
                severity="none",
                safe_repair_class=None,
                skip_reason=None,
            )
        )

    return {
        "schema_version": 1,
        "source_docx": str(source_docx) if source_docx is not None else None,
        "formatted_docx": str(formatted_docx),
        "pdf": str(pdf) if pdf is not None else None,
        **model,
        "rendered_rows": rendered_rows_payload,
        "issues": [asdict(item) for item in deduped],
        "style_notes": [asdict(item) for item in _style_notes(model)],
    }


def write_universal_table_diagnostics(
    *,
    formatted_docx: Path,
    pdf: Path,
    out: Path,
    source_docx: Path | None = None,
) -> dict:
    report = build_universal_table_diagnostics(
        source_docx=source_docx,
        formatted_docx=formatted_docx,
        pdf=pdf,
    )
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_text(
        json.dumps(report, ensure_ascii=False, indent=2, sort_keys=True),
        encoding="utf-8",
    )
    return report


def _stage(name: str, path: Path, previous: DiagnosticStage | None = None) -> DiagnosticStage:
    doc = Document(str(path))
    tables = []
    prev_tables = previous.tables if previous is not None else []
    for idx, table in enumerate(doc.tables):
        snap = snapshot_table(table._tbl, table_index=idx)
        if idx < len(prev_tables):
            snap["geometry_diff_from_previous_stage"] = diff_table_geometry(prev_tables[idx], snap)
        else:
            snap["geometry_diff_from_previous_stage"] = {}
        tables.append(snap)
    return DiagnosticStage(name=name, docx_path=str(path), tables=tables)


def run_table_engine_diagnostics(
    docx_path: Path,
    *,
    artifact_root: Path,
    render: bool = False,
) -> DiagnosticRunResult:
    from guides.coursework_kfu_2025.safe_formatter import process_document
    from guides.coursework_kfu_2025.table_continuation import apply_table_continuation

    artifact_root.mkdir(parents=True, exist_ok=True)
    source = artifact_root / "00_source.docx"
    safe = artifact_root / "01_after_safe_formatter.docx"
    phase3 = artifact_root / "02_after_table_continuation.docx"
    final = artifact_root / "03_final.docx"

    shutil.copy2(docx_path, source)
    process_document(source, safe)
    shutil.copy2(safe, phase3)
    doc = Document(str(phase3))
    apply_table_continuation(doc)
    doc.save(str(phase3))
    shutil.copy2(phase3, final)

    stages: list[DiagnosticStage] = []
    for name, path in [
        ("00_source", source),
        ("01_after_safe_formatter", safe),
        ("02_after_table_continuation", phase3),
        ("03_final", final),
    ]:
        stages.append(_stage(name, path, stages[-1] if stages else None))

    summary = artifact_root / "summary.json"
    summary.write_text(
        json.dumps({"stages": [asdict(stage) for stage in stages]}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return DiagnosticRunResult(str(artifact_root), str(summary), stages)


def _page(block: dict) -> int:
    return int(block.get("page") or 0)


def _y0(block: dict) -> float:
    return float(block.get("y0") or 0.0)


def _y1(block: dict) -> float:
    return float(block.get("y1") or 0.0)


def validate_continuation_markers_from_blocks(blocks: list[dict]) -> list[ContinuationMarkerDiagnostic]:
    from guides.coursework_kfu_2025.rendered_table_validation import classify_continuation_marker_line

    tables = sorted(
        [block for block in blocks if (block.get("kind") or "").lower() == "table"],
        key=lambda block: (_page(block), _y0(block)),
    )
    out: list[ContinuationMarkerDiagnostic] = []
    for block in blocks:
        if (block.get("kind") or "").lower() != "text":
            continue
        marker = classify_continuation_marker_line(str(block.get("text") or ""))
        if marker.marker_kind == "none":
            continue
        marker_page = _page(block) or None
        if marker.marker_kind == "source_inline_marker_text":
            out.append(ContinuationMarkerDiagnostic(marker.text, marker_page, marker.marker_kind, marker.marker_kind, None, None, None, "low"))
            continue
        previous = [table for table in tables if _page(table) < _page(block) or (_page(table) == _page(block) and _y1(table) <= _y0(block) + 2.0)]
        following = [table for table in tables if _page(table) > _page(block) or (_page(table) == _page(block) and _y0(table) >= _y1(block) - 2.0)]
        prev_page = _page(previous[-1]) if previous else None
        next_page = _page(following[0]) if following else None
        if prev_page is None or next_page is None:
            out.append(ContinuationMarkerDiagnostic(marker.text, marker_page, marker.marker_kind, "unknown", None, prev_page, next_page, "low"))
        elif prev_page == marker_page:
            out.append(ContinuationMarkerDiagnostic(marker.text, marker_page, marker.marker_kind, "fail", True, prev_page, next_page, "high" if next_page == marker_page else "medium"))
        else:
            out.append(ContinuationMarkerDiagnostic(marker.text, marker_page, marker.marker_kind, "pass", False, prev_page, next_page, "high"))
    return out


def write_continuation_marker_reports(artifact_root: Path, stages: dict[str, list[ContinuationMarkerDiagnostic]]) -> dict[str, str]:
    artifact_root.mkdir(parents=True, exist_ok=True)
    json_path = artifact_root / "continuation_markers.json"
    md_path = artifact_root / "continuation_markers.md"
    payload = {name: [asdict(item) for item in items] for name, items in stages.items()}
    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = []
    for name, items in stages.items():
        lines.append(f"## {name}")
        for item in items:
            lines.append(f"- {item.text}: page {item.marker_page}, violation={item.same_page_violation}, verdict={item.verdict}")
    md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return {
        "continuation_markers_json_path": str(json_path),
        "markdown_summary_path": str(md_path),
    }


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(
        description="Build non-mutating structured DOCX/PDF table diagnostics JSON."
    )
    parser.add_argument("--source-docx", type=Path, default=None)
    parser.add_argument("--formatted-docx", type=Path, required=True)
    parser.add_argument("--pdf", type=Path, required=True)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args(argv)

    if args.source_docx is not None and not args.source_docx.exists():
        parser.error(f"--source-docx does not exist: {args.source_docx}")
    if not args.formatted_docx.exists():
        parser.error(f"--formatted-docx does not exist: {args.formatted_docx}")
    if not args.pdf.exists():
        parser.error(f"--pdf does not exist: {args.pdf}")

    write_universal_table_diagnostics(
        source_docx=args.source_docx,
        formatted_docx=args.formatted_docx,
        pdf=args.pdf,
        out=args.out,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
